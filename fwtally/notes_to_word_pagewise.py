import argparse
import io
import json
import re
import time
from pathlib import Path

import fitz
from deep_translator import GoogleTranslator
from docx import Document

BASE_DIR = Path(__file__).resolve().parent
PDF_PATH = BASE_DIR / "Tally Notes.pdf"
OUT_DOCX = BASE_DIR / "Tally_Notes_English_Pagewise.docx"
PROGRESS_JSON = BASE_DIR / "Tally_Notes_English_Pagewise.progress.json"


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ").replace("\ufeff", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def looks_garbled(text: str) -> bool:
    if not text:
        return True
    patterns = [r"\[kk", r"\bog\b", r"dgyk", r"lEcf", r"O;fDr", r"vk;"]
    hits = sum(1 for p in patterns if re.search(p, text))
    return hits >= 2


def apply_manual_legacy_fixes(text: str) -> str:
    fixes = [
        (
            r"og \[kkrs tks fdfl O;fDr ;k laLFkk ds uke ls cuk, tkrs gS\] os Personal\s*Account ¼O;fDrxr \[kkrs½ dgykrs gSA tSls&",
            "Those accounts that are created in the name of a person or an institution are called Personal Accounts (individual accounts), such as:",
        ),
        (
            r"og \[kkrs tks fdfl oLrq ;k lEifRr vkfn ls lEcfU\/kr gksrs gSa Real\s*Account ¼okLrfod \[kkrs½ dgykrs gSaA tSls&",
            "Those accounts related to assets, property, or things are called Real Accounts, such as:",
        ),
        (
            r"og \[kkrs tks ykHk&gkfu\] vk;&O;;\] vkSj Ø;&foØ; ls lEcfU\/kr gksrs gSa\s*Nominal Account ¼ukeek= ds \[kkrs½ dgykrs gSaA tSls&",
            "Those accounts related to profit-loss, income-expense, and purchase-sales are called Nominal Accounts, such as:",
        ),
        (r"¼ ikus okyk ½", "(the receiver)"),
        (r"¼ nsus okyk ½", "(the giver)"),
        (r"¼ vkus okyk ½", "(what comes in)"),
        (r"¼ tkus okyk ½", "(what goes out)"),
        (r"¼ gkfu\] \[kPkkZ ½", "(loss and expenses)"),
        (r"¼ ykHk\] vk; ½", "(profit and income)"),
        (
            r"lHkh izdkj ds iw¡th \[kkrs ¼ tc O;kikj vkjEHk djrs gS ½\s*Capital Account ds vUrxZr \[kksys tkrs gSA tSls&",
            "All capital-related accounts opened at the start of business are grouped under Capital Account, such as:",
        ),
        (
            r"O;kikj \}kjk fy, x, _\.k O;kikj dk nkf;Ro gSA bl izdkj\s*ds \[kkrs Loans Liabilities ds vUrxZr cuk, tkrs gSA tSls&",
            "Loan amounts taken by the business are liabilities, and such accounts are grouped under Loans Liabilities, such as:",
        ),
        (
            r"Pkkyw nkf;Ro ds \[kkrs Current Liabilities ds vUrxZr \[kksys\s*tkrs gSA tSls&vfxze fy;k gqvk _\.k\] fofo/k ysunkjA",
            "Current obligation accounts are opened under Current Liabilities, such as short-term borrowings and sundry creditors.",
        ),
        (
            r"jke us jkgqy ls 5000 dk Computer m/kkj \[kjhnk A",
            "Ram purchased a computer worth 5000 on credit from Rahul.",
        ),
        (
            r"jke us vius edku dks fxjoh j\[k dj 10000 dk jkgqy ls yksu fy;kA",
            "Ram mortgaged his house and took a loan of 10000 from Rahul.",
        ),
    ]

    out = text
    for pattern, replacement in fixes:
        out = re.sub(pattern, replacement, out, flags=re.IGNORECASE)
    return out


def ocr_space_image_bytes(image_bytes: bytes) -> str:
    import requests

    url = "https://api.ocr.space/parse/image"
    payload = {
        "apikey": "helloworld",
        "language": "hin",
        "isOverlayRequired": False,
        "OCREngine": 2,
        "detectOrientation": True,
        "scale": True,
    }
    files = {"filename": ("page.png", image_bytes, "image/png")}
    r = requests.post(url, data=payload, files=files, timeout=120)
    r.raise_for_status()
    data = r.json()
    if data.get("IsErroredOnProcessing"):
        msg = data.get("ErrorMessage") or data.get("ErrorDetails") or "OCR processing error"
        if isinstance(msg, list):
            msg = " | ".join(str(x) for x in msg)
        raise RuntimeError(str(msg))
    parsed = data.get("ParsedResults") or []
    out = []
    for item in parsed:
        out.append(item.get("ParsedText", ""))
    return clean_text("\n".join(out))


def extract_page_text(page: fitz.Page, force_ocr: bool = False) -> tuple[str, str]:
    native = clean_text(page.get_text("text"))
    need_ocr = force_ocr or (len(native) < 120) or looks_garbled(native)

    if need_ocr:
        pix = page.get_pixmap(dpi=300, alpha=False)
        png_bytes = pix.tobytes("png")
        ocr_txt = ""
        for attempt in range(3):
            try:
                ocr_txt = ocr_space_image_bytes(png_bytes)
                break
            except Exception:
                time.sleep(1.8 * (attempt + 1))
        if force_ocr and ocr_txt:
            best = ocr_txt
            source = "ocr"
        else:
            best = ocr_txt if len(ocr_txt) > len(native) else native
            source = "ocr" if len(ocr_txt) >= len(native) else "native"
        return (best if best else "[Unclear page text]", source)

    return native, "native"


def chunks_for_translation(text: str, max_chars: int = 3500):
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunk = ""
    for para in paras:
        candidate = para if not chunk else f"{chunk}\n\n{para}"
        if len(candidate) <= max_chars:
            chunk = candidate
        else:
            if chunk:
                yield chunk
            chunk = para
    if chunk:
        yield chunk


def translate_to_english(text: str) -> str:
    text = apply_manual_legacy_fixes(text)
    translator = GoogleTranslator(source="auto", target="en")
    chunks = list(chunks_for_translation(text))
    out = []
    for c in chunks:
        translated = None
        for attempt in range(3):
            try:
                translated = translator.translate(c)
                break
            except Exception:
                time.sleep(1.2 * (attempt + 1))
        if translated is None:
            translated = "[Translation failed for chunk]\n" + c
        out.append(translated)
        time.sleep(0.1)
    return clean_text("\n\n".join(out))


def load_progress() -> dict:
    if PROGRESS_JSON.exists():
        try:
            return json.loads(PROGRESS_JSON.read_text(encoding="utf-8"))
        except Exception:
            return {"processed_pages": []}
    return {"processed_pages": []}


def save_progress(progress: dict) -> None:
    PROGRESS_JSON.write_text(json.dumps(progress, ensure_ascii=False, indent=2), encoding="utf-8")


def append_page_to_docx(page_no: int, source: str, content: str) -> None:
    if OUT_DOCX.exists():
        doc = Document(str(OUT_DOCX))
    else:
        doc = Document()
        doc.add_heading("Tally Notes - English (Pagewise)", level=1)

    doc.add_page_break()
    doc.add_heading(f"Page {page_no} ({source})", level=2)

    for para in [p.strip() for p in re.split(r"\n\n+", content) if p.strip()]:
        doc.add_paragraph(para)

    doc.save(str(OUT_DOCX))


def main() -> int:
    parser = argparse.ArgumentParser(description="Process Tally Notes one page at a time and append to a single Word file.")
    parser.add_argument("--page", type=int, required=True, help="1-based page number to process")
    parser.add_argument("--force-ocr", action="store_true", help="Force OCR for this page")
    parser.add_argument("--reprocess", action="store_true", help="Allow processing even if page was already processed")
    args = parser.parse_args()

    doc = fitz.open(PDF_PATH)
    total = len(doc)
    if args.page < 1 or args.page > total:
        print(f"Invalid page: {args.page}. Valid range is 1 to {total}.")
        return 1

    progress = load_progress()
    processed = set(progress.get("processed_pages", []))
    if args.page in processed and not args.reprocess:
        print(f"Page {args.page} already processed. Use --reprocess to force append again.")
        return 0

    page = doc[args.page - 1]
    print(f"Processing page {args.page}/{total}...", flush=True)
    raw_text, source = extract_page_text(page, force_ocr=args.force_ocr)
    english = translate_to_english(raw_text)
    append_page_to_docx(args.page, source, english)

    if args.page not in processed:
        progress.setdefault("processed_pages", []).append(args.page)
        progress["processed_pages"] = sorted(set(progress["processed_pages"]))
        save_progress(progress)

    doc.close()

    print(f"Appended page {args.page} to: {OUT_DOCX}")
    print(f"Updated progress: {PROGRESS_JSON}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
