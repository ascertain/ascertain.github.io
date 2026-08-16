from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HTML_PATH = Path(__file__).resolve().parent / "Tally_Book_BY_PARNAMI.html"
DOCX_PATH = Path(__file__).resolve().parent / "Tally_Book_BY_PARNAMI.docx"


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_runs_from_inline(paragraph, node: Tag | NavigableString) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    if node.name == "a":
        add_hyperlink(paragraph, node.get_text(" ", strip=True), node.get("href", ""))
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    for child in node.children:
        add_runs_from_inline(paragraph, child)


def write_paragraph(doc: Document, text: str, style: str | None = None, align=None) -> None:
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.add_run(text)


def handle_list(doc: Document, tag: Tag, ordered: bool = False) -> None:
    items = tag.find_all("li", recursive=False)
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph(style="List Bullet")
        if ordered:
            paragraph = doc.add_paragraph(style="List Number")
        add_runs_from_inline(paragraph, item)


def handle_table(doc: Document, tag: Tag) -> None:
    rows = tag.find_all("tr", recursive=False)
    if not rows:
        return

    first_row_cells = rows[0].find_all(["th", "td"], recursive=False)
    table = doc.add_table(rows=0, cols=max(1, len(first_row_cells)))
    table.style = "Table Grid"

    for row_tag in rows:
        cells = row_tag.find_all(["th", "td"], recursive=False)
        row = table.add_row().cells
        for index, cell_tag in enumerate(cells[: len(row)]):
            row[index].text = cell_tag.get_text(" ", strip=True)


def handle_pre(doc: Document, tag: Tag) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(tag.get_text("\n", strip=False))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def handle_svg(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("[Diagram/chart retained in HTML version. Refer to the HTML file for the full visual.]")
    run.italic = True


def handle_tag(doc: Document, tag: Tag, chapter_started: bool = False) -> None:
    if tag.name == "div":
        for child in tag.children:
            if isinstance(child, Tag):
                handle_tag(doc, child, chapter_started)
        return

    if tag.name == "h2":
        if chapter_started:
            doc.add_page_break()
        doc.add_heading(tag.get_text(" ", strip=True), level=1)
        return

    if tag.name == "h3":
        doc.add_page_break()
        doc.add_heading(tag.get_text(" ", strip=True), level=2)
        return

    if tag.name == "p":
        paragraph = doc.add_paragraph()
        for child in tag.children:
            add_runs_from_inline(paragraph, child)
        return

    if tag.name == "ul":
        handle_list(doc, tag, ordered=False)
        return

    if tag.name == "ol":
        handle_list(doc, tag, ordered=True)
        return

    if tag.name == "table":
        handle_table(doc, tag)
        return

    if tag.name == "pre":
        handle_pre(doc, tag)
        return

    if tag.name == "svg":
        handle_svg(doc)
        return


def build_docx(html_path: Path, docx_path: Path) -> None:
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    title = soup.title.get_text(strip=True) if soup.title else docx_path.stem
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)

    cover = soup.select_one("section.cover")
    if cover:
        subtitle = cover.find("p")
        if subtitle:
            write_paragraph(document, subtitle.get_text(" ", strip=True), align=WD_ALIGN_PARAGRAPH.CENTER)

        for meta in cover.select(".meta"):
            write_paragraph(document, meta.get_text(" ", strip=True), align=WD_ALIGN_PARAGRAPH.CENTER)

        chart = cover.select_one(".chart")
        if chart:
            document.add_paragraph()
            document.add_heading(chart.find("h3").get_text(" ", strip=True), level=2)
            for bar in chart.select(".bar"):
                label = bar.find("label")
                if label:
                    document.add_paragraph(label.get_text(" ", strip=True), style="List Bullet")

    document.add_page_break()

    chapters = soup.select("section.chapter")
    for chapter_index, chapter in enumerate(chapters):
        if chapter_index > 0:
            document.add_page_break()

        for child in chapter.children:
            if isinstance(child, Tag):
                handle_tag(document, child, chapter_started=False)

    sources = soup.select_one("section.sources")
    if sources:
        document.add_page_break()
        document.add_heading("Internet Sources Used", level=1)
        source_list = sources.find("ul")
        if source_list:
            for item in source_list.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style="List Bullet")
                for child in item.children:
                    add_runs_from_inline(paragraph, child)
        note = sources.find(class_="callout")
        if note:
            write_paragraph(document, note.get_text(" ", strip=True))

    document.save(docx_path)


def main() -> int:
    build_docx(HTML_PATH, DOCX_PATH)
    print(f"Created: {DOCX_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())