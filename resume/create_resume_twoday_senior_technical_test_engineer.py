from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Senior_Technical_Test_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Senior_Technical_Test_Engineer_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "test automation",
    "test architecture",
    "regression",
    "Page Object Model",
    "POM",
    "test orchestration",
    "smoke",
    "functional",
    "integration",
    "component",
    "user acceptance",
    "UAT",
    "acceptance criteria",
    "Definition of Done",
    "API testing",
    "Postman",
    "Selenium",
    "Playwright",
    "RestAssured",
    "Karate",
    "CI/CD",
    "continuous delivery",
    "TDD",
    "Test Driven Development",
    "test strategy",
    "test environment",
    "test data",
    "Agile",
    "Scrum",
    "SAFe",
    "PI planning",
    "sprint planning",
    "backlog refinement",
    "cross-functional",
    "full-stack",
    "built-in quality",
    "scalable",
    "mentor",
    "coaching",
    "AI-assisted",
    "AI-driven",
    "Claude",
    "Copilot",
    "Docker",
    "Kubernetes",
    "GitHub Actions",
    "Jenkins",
    "large-scale",
    "public sector",
    "regulated",
    "IKEA",
    "LEGO",
    "Truecaller",
    "ownership",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior Technical Test Engineer  |  Test Automation, Architecture & Agile Delivery  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior Technical Test Engineer with 15+ years of hands-on experience in "
        "large-scale, business-critical development programmes. Combines test automation, "
        "test architecture (Page Object Model, scalable frameworks), and a cross-functional "
        "role supporting team processes and delivery flow. Develops and maintains automated "
        "regression tests, designs test cases for user stories (Definition of Done), and "
        "performs functional, integration, component, and user acceptance testing. Strong "
        "API testing experience (Postman, Selenium, Playwright, RestAssured). Integrates "
        "testing into CI/CD pipelines for continuous delivery. Experienced in Scrum, SAFe, "
        "PI planning, backlog refinement, and TDD. Mentors engineers on testing practices. "
        "AI-assisted testing with Claude and Copilot. Nordics-based (Malmö) — experience "
        "at IKEA (8+ years), Truecaller, and LEGO across enterprise and e-commerce.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Test Automation & Architecture: ",
            "Developed and maintained automated regression tests ensuring stability "
            "across releases. Designed scalable test architecture using Page Object "
            "Model and structured test frameworks. Built test orchestration and "
            "automation setups (smoke and regression testing). Playwright, Selenium, "
            "RestAssured, Karate — across web, mobile, and API layers.",
        ),
        (
            "Test Design & Execution: ",
            "Designed, implemented, and executed test cases for user stories — ensuring "
            "Definition of Done is met. Performed functional, integration, component, "
            "and user acceptance testing. Validated solutions against acceptance criteria "
            "and quality standards across large-scale programmes.",
        ),
        (
            "Agile Collaboration & Delivery: ",
            "Active in backlog refinement, sprint planning, PI planning (SAFe). "
            "Collaborated closely with developers to ensure built-in quality. "
            "Supported planning, prioritisation, and delivery flow. Comfortable in "
            "cross-functional full-stack teams. Scrum and SAFe experience.",
        ),
        (
            "Test Environments & CI/CD: ",
            "Established and maintained test environments. Ensured reliable test data. "
            "Integrated testing into CI/CD pipelines (GitHub Actions, Jenkins, Docker, "
            "Kubernetes) for continuous delivery. TDD practices. Contributed to "
            "improving overall test strategy and quality practices.",
        ),
        (
            "Mentoring & AI-Assisted Testing: ",
            "Mentored QA engineers and developers on testing best practices, "
            "test architecture, and TDD. AI-assisted testing daily — Claude and "
            "Copilot for test generation, code reviews, and exploratory testing.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Test Automation: ",
            "Playwright, Selenium, Appium, RestAssured, Karate, TestNG, "
            "Postman, E2E, API testing, regression, smoke, component testing",
        ),
        (
            "Test Architecture: ",
            "Page Object Model (POM), scalable test frameworks, test "
            "orchestration, TDD, structured design patterns",
        ),
        (
            "CI/CD & Environments: ",
            "GitHub Actions, Jenkins, Docker, Kubernetes, Terraform, "
            "test environment management, test data, continuous delivery",
        ),
        (
            "Agile & Process: ",
            "Scrum, SAFe, PI planning, sprint planning, backlog refinement, "
            "Definition of Done, acceptance criteria, delivery flow",
        ),
        (
            "AI-Assisted Testing: ",
            "Claude, Copilot, Gemini — AI test generation, code review, "
            "exploratory testing, workflow automation",
        ),
        (
            "Languages & Tools: ",
            "Java, C#, TypeScript, Python, SQL, Maven, TestRail, "
            "Jira, Confluence, Git, Grafana",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Senior Technical Test Engineer / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Large-Scale SaaS Platform, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Test automation & architecture — developed and maintained automated regression "
        "tests (Playwright, RestAssured, Karate) ensuring stability across releases. "
        "Designed scalable test architecture using Page Object Model and structured "
        "frameworks. Test orchestration for smoke and regression testing.",
        "Test design & execution — designed, implemented, and executed test cases for "
        "user stories, ensuring Definition of Done. Functional, integration, component, "
        "and user acceptance testing. Validated acceptance criteria and quality standards.",
        "Agile collaboration — active in backlog refinement, sprint planning, PI planning. "
        "Collaborated closely with developers for built-in quality. Supported planning, "
        "prioritisation, and delivery flow in cross-functional full-stack team. Scrum.",
        "Test environments & CI/CD — established and maintained test environments. "
        "Reliable test data. Integrated testing into CI/CD pipelines (GitHub Actions, "
        "Docker, Kubernetes, Terraform) for continuous delivery. TDD practices.",
        "Mentored engineers on testing practices and test architecture. AI-assisted "
        "testing daily (Claude, Copilot) — 30% velocity improvement. Contributed "
        "to improving overall test strategy and quality practices.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior Technical Test Engineer / QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Large-Scale SaaS — 300M+ Users, Multi-Platform", bold=True, size=10)
    tc_bullets = [
        "Test automation across web (Selenium), mobile (Appium — iOS/Android), and "
        "API (RestAssured, Postman). Automated regression tests, smoke tests, scalable "
        "test architecture (POM). Functional, integration, and UAT for a 300M+ user "
        "SaaS platform.",
        "Agile delivery (Scrum/Kanban) — backlog refinement, sprint planning. "
        "Cross-functional collaboration with developers. CI/CD integration. "
        "Mentored engineers on testing practices. Built-in quality culture.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Technical Test Engineer / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  Large-Scale Enterprise & E-Commerce — Web, Mobile, API", bold=True, size=10)
    lego_bullets = [
        "LEGO: Scalable test automation frameworks — Selenium, RestAssured, Karate, "
        "Appium. Page Object Model architecture. Automated regression and smoke "
        "testing. Functional, integration, component, UAT. Led 8–10 engineers. "
        "Test orchestration. Agile (Scrum), sprint planning, backlog refinement.",
        "IKEA App, Genesys, Verint (2018–2021): Test automation across large-scale "
        "enterprise programmes. Test architecture, CI/CD integration, test "
        "environments. Cross-functional delivery with developers. Definition of "
        "Done. Mentored developers on testing practices and TDD.",
        "Business-critical systems — ensured quality, stability, and reliability "
        "across releases. Worked in complex system landscapes with multiple "
        "integration points. Agile ceremonies, delivery flow, quality practices.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior Test Engineer / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Government — Large-Scale, Regulated Environments", bold=True, size=10)
    fin_bullets = [
        "Test automation in large-scale regulated environments — Selenium, RestAssured, "
        "TestNG. Scalable test architecture (POM). Automated regression, functional, "
        "integration, UAT. CI/CD. API testing (Postman). Test environments and "
        "test data management. SQL, Java, C#.",
        "Mentored 15+ engineers on testing practices. Led teams of 10+. Agile (Scrum). "
        "Cross-functional collaboration with developers. Business-critical banking "
        "systems. Backlog refinement, Definition of Done, delivery flow.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Test automation & architecture — Page Object Model, scalable frameworks, "
        "test orchestration (smoke, regression) across IKEA (30+ markets), "
        "Truecaller (300M+ users), and LEGO.",
        "Large-scale business-critical programmes — ensured quality, stability, "
        "and reliability across releases in complex system landscapes. Enterprise "
        "and public-sector-comparable regulated environments.",
        "CI/CD integration — GitHub Actions, Jenkins, Docker, Kubernetes. Test "
        "environments, test data, continuous delivery. TDD practices.",
        "AI-assisted testing — Claude, Copilot daily. AI test generation, "
        "exploratory testing. 30% velocity improvement. Team rollout.",
        "Mentored engineers across every organization — testing practices, "
        "test architecture, TDD. Cross-functional, built-in quality culture.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic  |  Danish — Conversational", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Senior Technical Test Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior Technical Test Engineer | Test Automation, Architecture &amp; Agile Delivery | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Senior Technical Test Engineer with 15+ years in <span class="hl">large-scale</span>, business-critical development programmes. Combines <span class="hl">test automation</span>, <span class="hl">test architecture</span> (<span class="hl">Page Object Model</span>, <span class="hl">scalable</span> frameworks), and a <span class="hl">cross-functional</span> role supporting team processes and delivery flow. Automated <span class="hl">regression</span> tests, <span class="hl">functional</span>, <span class="hl">integration</span>, <span class="hl">component</span>, and <span class="hl">UAT</span>. <span class="hl">API testing</span> (<span class="hl">Postman</span>, <span class="hl">Selenium</span>, <span class="hl">Playwright</span>, <span class="hl">RestAssured</span>). <span class="hl">CI/CD</span> for <span class="hl">continuous delivery</span>. <span class="hl">Scrum</span>, <span class="hl">SAFe</span>, <span class="hl">PI planning</span>, <span class="hl">TDD</span>. <span class="hl">Mentor</span>s engineers. <span class="hl">AI-assisted</span> testing. Nordics-based — <span class="hl">IKEA</span> (8+ years), <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Test Automation &amp; Architecture:</b> Automated <span class="hl">regression</span> tests. <span class="hl">Scalable</span> <span class="hl">test architecture</span> (<span class="hl">POM</span>). <span class="hl">Test orchestration</span> (<span class="hl">smoke</span>, <span class="hl">regression</span>). <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span>.<br>
  <b>Test Design &amp; Execution:</b> Test cases for user stories — <span class="hl">Definition of Done</span>. <span class="hl">Functional</span>, <span class="hl">integration</span>, <span class="hl">component</span>, <span class="hl">UAT</span>. <span class="hl">Acceptance criteria</span>.<br>
  <b>Agile Collaboration:</b> <span class="hl">Backlog refinement</span>, <span class="hl">sprint planning</span>, <span class="hl">PI planning</span> (<span class="hl">SAFe</span>). <span class="hl">Built-in quality</span> with developers. <span class="hl">Cross-functional</span> <span class="hl">full-stack</span> teams.<br>
  <b>Test Environments &amp; CI/CD:</b> <span class="hl">Test environment</span>s, <span class="hl">test data</span>. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). <span class="hl">TDD</span>. <span class="hl">Continuous delivery</span>.<br>
  <b>Mentoring &amp; AI:</b> <span class="hl">Mentor</span>ed engineers on testing practices. <span class="hl">AI-assisted</span> testing (<span class="hl">Claude</span>, <span class="hl">Copilot</span>).</p>

  <div class="section">Technical Skills</div>
  <p><b>Test Automation:</b> Playwright, Selenium, Appium, RestAssured, Karate, TestNG, Postman, E2E, API, regression, smoke, component<br>
  <b>Test Architecture:</b> Page Object Model, scalable frameworks, test orchestration, TDD, structured patterns<br>
  <b>CI/CD &amp; Environments:</b> GitHub Actions, Jenkins, Docker, Kubernetes, Terraform, test environments, test data, continuous delivery<br>
  <b>Agile &amp; Process:</b> Scrum, SAFe, PI planning, sprint planning, backlog refinement, Definition of Done, acceptance criteria<br>
  <b>AI-Assisted Testing:</b> Claude, Copilot, Gemini — AI test generation, code review, exploratory testing<br>
  <b>Languages &amp; Tools:</b> Java, C#, TypeScript, Python, SQL, Maven, TestRail, Jira, Confluence, Git</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Senior Technical Test Engineer / Test Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Large-Scale SaaS Platform, 30+ Markets</div>
  <ul>
    <li><span class="hl">Test automation</span> &amp; <span class="hl">test architecture</span> — <span class="hl">Playwright</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span>. <span class="hl">Page Object Model</span>. Automated <span class="hl">regression</span> and <span class="hl">smoke</span> tests. <span class="hl">Test orchestration</span>. <span class="hl">Scalable</span> frameworks.</li>
    <li>Test cases for user stories — <span class="hl">Definition of Done</span>. <span class="hl">Functional</span>, <span class="hl">integration</span>, <span class="hl">component</span>, <span class="hl">UAT</span>. <span class="hl">Acceptance criteria</span>.</li>
    <li><span class="hl">Agile</span> — <span class="hl">backlog refinement</span>, <span class="hl">sprint planning</span>, <span class="hl">PI planning</span>. <span class="hl">Built-in quality</span> with developers. <span class="hl">Cross-functional</span> <span class="hl">full-stack</span> team.</li>
    <li><span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). <span class="hl">Test environment</span>s, <span class="hl">test data</span>. <span class="hl">TDD</span>. <span class="hl">Continuous delivery</span>.</li>
    <li><span class="hl">Mentor</span>ed engineers. <span class="hl">AI-assisted</span> testing (<span class="hl">Claude</span>, <span class="hl">Copilot</span>). 30% velocity. <span class="hl">Test strategy</span> improvement.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior Technical Test Engineer / QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Large-Scale SaaS — 300M+ Users, Multi-Platform</div>
  <ul>
    <li><span class="hl">Test automation</span> — <span class="hl">Selenium</span>, Appium, <span class="hl">RestAssured</span>, <span class="hl">Postman</span>. <span class="hl">Test architecture</span> (<span class="hl">POM</span>). <span class="hl">Regression</span>, <span class="hl">smoke</span>, <span class="hl">functional</span>, <span class="hl">integration</span>, <span class="hl">UAT</span>. 300M+ users.</li>
    <li><span class="hl">Agile</span> (<span class="hl">Scrum</span>) — <span class="hl">backlog refinement</span>, <span class="hl">sprint planning</span>. <span class="hl">Cross-functional</span>. <span class="hl">CI/CD</span>. <span class="hl">Mentor</span>ed engineers. <span class="hl">Built-in quality</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Technical Test Engineer / Test Lead</div>
  <div class="job-sub">2016 – 2021 | Large-Scale Enterprise &amp; E-Commerce — Web, Mobile, API</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Scalable</span> <span class="hl">test automation</span> — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span>, Appium. <span class="hl">POM</span> architecture. <span class="hl">Regression</span>, <span class="hl">smoke</span>, <span class="hl">functional</span>, <span class="hl">integration</span>, <span class="hl">component</span>, <span class="hl">UAT</span>. 8–10 engineers. <span class="hl">Test orchestration</span>.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Large-scale</span> enterprise programmes. <span class="hl">Test architecture</span>, <span class="hl">CI/CD</span>, <span class="hl">test environment</span>s. <span class="hl">Definition of Done</span>. <span class="hl">Mentor</span>ed developers on <span class="hl">TDD</span>.</li>
    <li>Business-critical systems — quality, stability, reliability across releases. <span class="hl">Agile</span> (<span class="hl">Scrum</span>), <span class="hl">backlog refinement</span>, delivery flow.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior Test Engineer / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Government — Large-Scale, Regulated Environments</div>
  <ul>
    <li><span class="hl">Test automation</span> in <span class="hl">large-scale</span> <span class="hl">regulated</span> environments — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>. <span class="hl">Test architecture</span> (<span class="hl">POM</span>). <span class="hl">Regression</span>, <span class="hl">functional</span>, <span class="hl">integration</span>, <span class="hl">UAT</span>. <span class="hl">CI/CD</span>. <span class="hl">API testing</span> (<span class="hl">Postman</span>).</li>
    <li><span class="hl">Mentor</span>ed 15+ engineers. <span class="hl">Agile</span> (<span class="hl">Scrum</span>). <span class="hl">Cross-functional</span>. Business-critical banking. <span class="hl">Backlog refinement</span>, <span class="hl">Definition of Done</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Test automation</span> &amp; <span class="hl">test architecture</span> — <span class="hl">POM</span>, <span class="hl">scalable</span> frameworks, <span class="hl">test orchestration</span> across <span class="hl">IKEA</span> (30+ markets), <span class="hl">Truecaller</span> (300M+), <span class="hl">LEGO</span>.</li>
    <li><span class="hl">Large-scale</span> business-critical programmes — quality, stability, reliability. Enterprise &amp; <span class="hl">regulated</span> environments.</li>
    <li><span class="hl">CI/CD</span> — <span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>. <span class="hl">Test environment</span>s, <span class="hl">test data</span>, <span class="hl">continuous delivery</span>. <span class="hl">TDD</span>.</li>
    <li><span class="hl">AI-assisted</span> testing — <span class="hl">Claude</span>, <span class="hl">Copilot</span>. 30% velocity. Team rollout.</li>
    <li><span class="hl">Mentor</span>ed engineers across every organization — testing practices, <span class="hl">test architecture</span>, <span class="hl">TDD</span>, <span class="hl">built-in quality</span> culture.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic  |  Danish — Conversational</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
