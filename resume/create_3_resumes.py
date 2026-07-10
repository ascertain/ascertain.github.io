"""
Generate 3 separate resumes:
1. Project & Technical Lead/Manager (15 years)
2. Test (Manual & Automation) Lead/Manager (15 years)
3. DevOps & Cloud Engineer (10 years)
No years in headline. Mention "Working in IT since 2008".
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text); run.font.size = Pt(10); run.font.name = "Calibri"
    return p

def setup_doc():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0
    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)
    return doc

def add_header(doc, title):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, title, bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

def add_competencies(doc, competencies):
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

def add_certs(doc, certs):
    table = doc.add_table(rows=(len(certs)+1)//2, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

def add_languages(doc):
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)", size=Pt(9))

def add_education(doc):
    add_section_heading(doc, "EDUCATION")
    for deg, school in [("Post Graduate Diploma in Operation and Management", "IGNOU, India"), ("B.Tech, Information Technology", "UP Technical University, India")]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))


# ═══════════════════════════════════════════════════════════════════════════════
# RESUME 1: Project & Technical Lead/Manager
# ═══════════════════════════════════════════════════════════════════════════════

def build_resume_project_lead():
    doc = setup_doc()
    add_header(doc, "Project & Technical Lead / Manager")

    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, (
        "Project and Technical Lead working in IT since 2008, with a broad background spanning cloud operations, "
        "enterprise platform delivery, and cross-functional program coordination across Nordic and global organizations. "
        "Experienced in leading complex IT projects end-to-end — from planning, coordination, and stakeholder alignment "
        "through go-live and stabilization. Combines a solid technical foundation in cloud infrastructure (GCP, Azure, AWS) "
        "with a primary strength in bringing structure, clarity, and momentum to delivery work. Skilled at translating "
        "complex technical realities into actionable plans, managing expectations at strategic and operational levels, "
        "and keeping distributed teams aligned toward outcomes. Strong communicator with extensive "
        "experience in regulated environments, ITIL/ITSM frameworks, and agile/hybrid methodologies."
    ))

    add_section_heading(doc, "CORE COMPETENCIES")
    add_competencies(doc, [
        "End-to-End Project Delivery", "Stakeholder Management", "Cloud Operations (GCP/Azure/AWS)",
        "Planning & Coordination", "Risk & Escalation Management", "Go-Live & Stabilization",
        "ITIL / ITSM Frameworks", "Agile / SAFe / Prince2", "Cross-Functional Leadership",
    ])

    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # IKEA
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True); add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Project Manager — Cloud Platform", bold=True); add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Lead cloud operations projects end-to-end — owning delivery from planning and coordination through go-live and stabilization across a platform serving 30+ markets.",
        "Act as primary point of contact between engineering teams and internal stakeholders — translating complex technical realities into clear plans and informed decisions.",
        "Keep projects on track through clear priorities, transparent communication, and proactive risk management — knowing when to escalate, adapt, or push back.",
        "Coordinate cross-functional delivery across engineers, operations, and vendors — managing dependencies, environment readiness, and integration milestones.",
        "Drive operational stability improvements aligned with ITIL/ITSM frameworks — contributing to incident management, change control, and service continuity practices.",
        "Manage expectations at strategic and operational levels — facilitating steering committees, sprint reviews, and program-level reporting.",
        "Lead go-live coordination and stabilization — orchestrating cutover planning, rollback procedures, and post-deployment monitoring.",
        "Facilitate agile ceremonies (sprint planning, retrospectives, PI planning) while maintaining project governance and milestone tracking.",
        "Introduce standardized delivery playbooks and continuous improvement practices — reducing project cycle time by 30%.",
    ]:
        add_bullet(doc, b)

    # Truecaller
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True); add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Delivery Lead — Cloud Platform", bold=True); add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Coordinated cloud platform delivery for a 300M+ user product — managing planning, priorities, and cross-team dependencies.",
        "Bridged engineering and business stakeholders — maintaining transparent communication on progress, risks, and trade-offs.",
        "Drove go-live readiness and stabilization — coordinating release schedules, monitoring, and incident response.",
        "Applied ITIL-aligned practices for incident and change management — ensuring operational stability during rapid iteration.",
    ]:
        add_bullet(doc, b)

    # HCL
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True); add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Project Manager / Delivery Lead", bold=True); add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Led end-to-end delivery of infrastructure and cloud projects across the Nordic region — managing complex IT landscapes spanning cloud, on-premise, and multi-vendor integrations.",
        "Owned planning and coordination for large-scale transformation programs — defining scope, timelines, resources, and go-live criteria across distributed teams.",
        "Served as primary stakeholder liaison — maintaining transparent communication and facilitating decision-making at program and operational levels.",
        "Managed projects in regulated environments (healthcare-adjacent, financial) — ensuring compliance and operational stability.",
        "Applied ITIL/ITSM frameworks for incident, change, and problem management — reducing unplanned downtime by 40%.",
        "Coordinated go-live and stabilization for major platform migrations — managing cutovers, rollback planning, and hypercare.",
        "Mentored teams of 6–15 engineers — creating clear deliverables with shared ownership.",
        "Delivered using agile, Prince2, and hybrid methodologies — adapting to project complexity and stakeholder maturity.",
    ]:
        add_bullet(doc, b)

    # Earlier
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Technical Project Delivery", bold=True); add_text(p, "  |  India  |  2008 – 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Coordinated enterprise IT project delivery in regulated financial environments — managing planning, stakeholder communication, and go-live execution.",
        "Managed systems integration between distributed IT systems — coordinating infrastructure provisioning and operational handover.",
        "Applied ITIL practices for service transitions — ensuring stability and clear handover from project to operations.",
    ]:
        add_bullet(doc, b)

    # Skills
    add_section_heading(doc, "TECHNICAL FOUNDATION")
    for label, value in [
        ("Cloud: ", "GCP, Azure, AWS — infrastructure, networking, CI/CD, monitoring"),
        ("Frameworks: ", "ITIL v4, ITSM, Agile (Scrum/Kanban), SAFe, Prince2"),
        ("Tools: ", "Jira, Confluence, Azure DevOps, ServiceNow, MS Project, Git, Terraform"),
        ("Infrastructure: ", "Kubernetes, Docker, networking (TCP/IP, VPN, DNS), hybrid cloud"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9)); add_text(p, value, size=Pt(9))

    add_education(doc)

    add_section_heading(doc, "CERTIFICATIONS")
    add_certs(doc, ["ITIL v4 Foundation", "Google Cloud ACE", "AWS Cloud Practitioner", "Six Sigma Green Belt", "CEH", "ISTQB"])

    add_languages(doc)

    path = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Project_Technical_Lead_Resume.docx")
    doc.save(path); print(f"DOCX saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# RESUME 2: Test (Manual & Automation) Lead/Manager
# ═══════════════════════════════════════════════════════════════════════════════

def build_resume_test_lead():
    doc = setup_doc()
    add_header(doc, "Test Lead & Manager — Manual & Automation")

    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, (
        "Test Lead and Manager working in IT since 2008, with comprehensive experience across manual testing, "
        "test automation, and quality strategy in enterprise, cloud, IoT, and financial domains. Skilled at leading "
        "test planning, estimation, coordination, execution, and reporting for complex programs involving multi-tier "
        "architectures, cloud platforms, and regulated environments. Builds and grows test teams, defines strategies "
        "that balance exploratory and automated approaches, and drives continuous improvement in test processes and "
        "tooling. Proficient in Python (pytest), Selenium, Playwright, and CI/CD pipeline integration. Experienced "
        "managing testing across Azure, GCP, and hybrid environments."
    ))

    add_section_heading(doc, "CORE COMPETENCIES")
    add_competencies(doc, [
        "Test Planning & Estimation", "Test Coordination & Reporting", "Test Execution & Follow-up",
        "Manual & Exploratory Testing", "Test Automation (Python/pytest)", "CI/CD Pipeline Integration",
        "Defect & Risk Management", "Agile / SAFe / ISTQB", "Team Leadership & Mentoring",
    ])

    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # IKEA
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True); add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Lead / Quality Manager — Cloud & IoT Platform", bold=True); add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Lead end-to-end test management for a cloud platform serving 30+ markets — owning test planning, estimation, coordination, reporting, and execution across distributed teams.",
        "Define and maintain test strategies balancing manual exploratory testing with automated regression — continuously growing automation coverage (pytest, Playwright) while ensuring thorough manual validation.",
        "Coordinate test execution across SIT, E2E, UAT, and regression phases — managing dependencies, environments, and timelines within agile sprint cadence.",
        "Produce test reporting and quality metrics for stakeholders — providing transparent visibility into progress, risks, defect trends, and release readiness.",
        "Drive defect management and follow-up — tracking issues to resolution, conducting root-cause analysis, and ensuring exit criteria are met before releases.",
        "Develop and maintain Python-based automation frameworks (pytest) integrated into CI/CD pipelines (GitHub Actions) — extending coverage by 60%.",
        "Establish test governance and quality gates within Azure DevOps — ensuring traceability from requirements through test cases to defect closure.",
        "Mentor and develop a team of test engineers — providing guidance on manual techniques, automation skills, and career growth.",
        "Continuously improve test processes: introduced risk-based prioritization, structured exploratory sessions, and automated smoke suites that reduced test cycle time by 40%.",
    ]:
        add_bullet(doc, b)

    # Truecaller
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True); add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Manager / Release Quality Lead", bold=True); add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Managed test planning and coordination for a 300M+ user platform — orchestrating manual and automated test execution across cloud environments.",
        "Produced release quality dashboards and test reporting for leadership — enabling data-driven go/no-go decisions.",
        "Coordinated test estimation and resource allocation — balancing automation and manual efforts for optimal coverage.",
        "Drove defect management and follow-up — tracking critical issues through resolution and maintaining quality standards.",
        "Developed Python-based automated test scripts integrated into CI pipelines — reducing manual regression effort.",
    ]:
        add_bullet(doc, b)

    # HCL
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True); add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Lead / Test Manager — Enterprise & Digital", bold=True); add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Led test management for large-scale enterprise transformation programs at IKEA and LEGO — managing complex multi-tier architectures, legacy integrations, and cloud migrations.",
        "Owned test planning for program releases — defining scope, estimation, entry/exit criteria, and resource allocation across multiple workstreams.",
        "Coordinated test execution across geographically distributed teams — managing manual and automated testing in SIT, E2E, and UAT phases.",
        "Delivered test reporting to steering committees — presenting quality metrics, risk assessments, and release readiness recommendations.",
        "Built and maintained automated test suites (Python, Selenium, pytest) integrated into CI/CD pipelines (Jenkins) — reducing manual regression by 60%.",
        "Managed test environments across Azure and on-premise infrastructure — ensuring availability and stability for all test phases.",
        "Led both manual exploratory testing and structured automation — adapting the approach to system maturity and risk profile.",
        "Mentored teams of 8–12 test engineers — building competence in both manual techniques and automation skills.",
        "Drove continuous improvement: introduced exploratory charters, risk-based prioritization, and automated smoke suites.",
    ]:
        add_bullet(doc, b)

    # Earlier
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Test Engineering & Leadership", bold=True); add_text(p, "  |  India  |  2008 – 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Led test planning, estimation, and execution for core banking and financial system implementations — coordinating testing across complex multi-tier architectures.",
        "Managed test coordination across onshore/offshore teams — owning environments, data prep, and execution schedules.",
        "Developed automation scripts (Python, shell) for regression testing — improving coverage and reducing manual effort.",
        "Produced test reporting and quality dashboards — tracking defect metrics and release readiness in regulated environments.",
    ]:
        add_bullet(doc, b)

    # Skills
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("Automation: ", "Python (pytest, requests), Selenium, Playwright, API testing, shell scripting"),
        ("Tools: ", "Azure DevOps, Jira, Confluence, Zephyr, TestRail, HP ALM, Git"),
        ("CI/CD: ", "GitHub Actions, Jenkins, Azure Pipelines — automated test integration, quality gates"),
        ("Platforms: ", "GCP, Azure, AWS, Linux, Docker, Kubernetes"),
        ("Methods: ", "Agile (Scrum/Kanban), SAFe, ISTQB, risk-based testing, exploratory testing, TDD/BDD"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9)); add_text(p, value, size=Pt(9))

    add_education(doc)

    add_section_heading(doc, "CERTIFICATIONS")
    add_certs(doc, ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "ITIL v4 Foundation", "Six Sigma Green Belt", "CEH"])

    add_languages(doc)

    path = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Test_Lead_Manager_Resume.docx")
    doc.save(path); print(f"DOCX saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# RESUME 3: DevOps & Cloud Engineer
# ═══════════════════════════════════════════════════════════════════════════════

def build_resume_devops():
    doc = setup_doc()
    add_header(doc, "DevOps & Cloud Engineer")

    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, (
        "DevOps and Cloud Engineer working in IT since 2008, with dedicated cloud and DevOps focus since the "
        "mid-2010s. Designs, builds, and operates scalable cloud infrastructure and CI/CD pipelines across GCP, "
        "Azure, and AWS. Hands-on with Terraform, Kubernetes, Docker, GitHub Actions, and cloud-native services. "
        "Experienced in building reliable, observable systems — from automated deployments and infrastructure-as-code "
        "to monitoring, incident response, and operational stability. Works closely with development teams to streamline "
        "delivery, reduce toil, and ensure production systems are secure, scalable, and resilient. Strong Linux, "
        "networking, and scripting skills (Python, Bash)."
    ))

    add_section_heading(doc, "CORE COMPETENCIES")
    add_competencies(doc, [
        "Cloud Infrastructure (GCP/Azure/AWS)", "CI/CD Pipelines & Automation", "Infrastructure as Code (Terraform)",
        "Kubernetes & Docker", "Linux & System Administration", "Monitoring & Observability",
        "Networking & Security", "Python & Shell Scripting", "Reliability & Incident Response",
    ])

    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # IKEA
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True); add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "DevOps & Cloud Engineer — Platform Team", bold=True); add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Design, build, and maintain cloud infrastructure on GCP — managing Cloud Run, BigQuery, Pub/Sub, Cloud Functions, IAM, and networking for a platform serving 30+ markets.",
        "Develop and maintain CI/CD pipelines (GitHub Actions) — automating build, test, deploy, and rollback workflows for microservices across development, staging, and production environments.",
        "Implement Infrastructure as Code using Terraform — managing all cloud resources declaratively with version control, state management, and modular architecture.",
        "Operate and maintain Kubernetes clusters and Docker-based workloads — handling deployments, scaling, health checks, and resource optimization.",
        "Build monitoring and observability solutions — configuring alerting, dashboards, log aggregation, and distributed tracing for proactive incident detection.",
        "Manage networking and security — configuring VPCs, firewalls, VPNs, DNS, service mesh, and IAM policies to ensure secure, compliant infrastructure.",
        "Automate operational tasks using Python and Bash — building tools for secret rotation, data pipeline orchestration, environment provisioning, and incident remediation.",
        "Collaborate closely with developers to streamline delivery — reducing deployment frequency from weekly to multiple daily releases with zero-downtime deployments.",
        "Drive reliability practices — implementing SLOs, error budgets, automated rollbacks, and chaos testing to maintain 99.9%+ uptime.",
        "Participate in on-call rotation — responding to production incidents, performing root-cause analysis, and implementing preventive improvements.",
    ]:
        add_bullet(doc, b)

    # Truecaller
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True); add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Cloud & Release Engineer", bold=True); add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Managed cloud infrastructure and CI/CD pipelines on AWS for a 300M+ user platform — ensuring reliable, automated deployments at scale.",
        "Built and maintained automated release pipelines — integrating testing, security scanning, and deployment gates for production-safe releases.",
        "Configured monitoring and alerting — enabling rapid incident detection and reducing mean time to recovery.",
        "Automated operational workflows using Python and shell scripting — reducing manual toil and improving team velocity.",
        "Collaborated with development teams on infrastructure needs — provisioning environments and optimizing resource utilization.",
    ]:
        add_bullet(doc, b)

    # HCL
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True); add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "DevOps Engineer / Infrastructure Specialist", bold=True); add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Built and maintained CI/CD pipelines (Jenkins, GitLab CI) for enterprise applications — automating builds, tests, and deployments across cloud and on-premise environments.",
        "Managed cloud infrastructure on Azure and GCP — provisioning, configuring, and optimizing compute, storage, networking, and managed services.",
        "Implemented Infrastructure as Code (Terraform, Ansible) — migrating manual infrastructure provisioning to repeatable, version-controlled automation.",
        "Operated Kubernetes clusters and Docker-based microservices — managing deployments, scaling, and operational health.",
        "Configured networking infrastructure — VPNs, firewalls, load balancers, DNS, and service discovery for hybrid cloud architectures.",
        "Built monitoring and alerting systems — using Prometheus, Grafana, ELK stack, and cloud-native monitoring for proactive operations.",
        "Managed Linux servers (Ubuntu, RHEL) — system administration, hardening, patching, and performance tuning.",
        "Automated operational workflows (Python, Bash) — secret management, backup automation, log rotation, and incident remediation scripts.",
        "Supported development teams with environment provisioning, deployment troubleshooting, and production debugging.",
    ]:
        add_bullet(doc, b)

    # Earlier
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Systems & Infrastructure", bold=True); add_text(p, "  |  India  |  2008 – 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    for b in [
        "Managed Linux-based server infrastructure for enterprise banking applications — system administration, networking, and environment provisioning.",
        "Built automation scripts (Python, shell) for deployment, monitoring, and operational tasks — laying the foundation for DevOps practices.",
        "Configured and maintained network infrastructure — firewalls, load balancers, and connectivity for distributed systems.",
    ]:
        add_bullet(doc, b)

    # Skills
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("Cloud: ", "GCP (Cloud Run, BigQuery, Pub/Sub, IAM, VPC), Azure (App Services, DevOps, AKS), AWS (EC2, S3, Lambda, EKS)"),
        ("IaC & Config: ", "Terraform, Ansible, Helm, Kustomize"),
        ("Containers: ", "Kubernetes, Docker, container registries, service mesh (Istio)"),
        ("CI/CD: ", "GitHub Actions, Jenkins, GitLab CI, ArgoCD — pipeline design, quality gates, automated rollbacks"),
        ("Monitoring: ", "Prometheus, Grafana, Cloud Monitoring, ELK/Loki, distributed tracing, PagerDuty"),
        ("Scripting: ", "Python, Bash, Go (basics), REST APIs"),
        ("Networking: ", "TCP/IP, DNS, VPN, firewalls, load balancing, SSL/TLS, service mesh"),
        ("OS: ", "Linux (Ubuntu, Debian, RHEL) — systemd, networking, hardening, performance tuning"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9)); add_text(p, value, size=Pt(9))

    add_education(doc)

    add_section_heading(doc, "CERTIFICATIONS")
    add_certs(doc, ["Google Cloud ACE", "AWS Cloud Practitioner", "CEH", "ITIL v4 Foundation", "Six Sigma Green Belt", "ISTQB"])

    add_languages(doc)

    path = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_DevOps_Cloud_Engineer_Resume.docx")
    doc.save(path); print(f"DOCX saved: {path}")
    return path


if __name__ == "__main__":
    build_resume_project_lead()
    build_resume_test_lead()
    build_resume_devops()
    print("\nAll 3 resumes generated!")
