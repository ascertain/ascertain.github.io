"""Generate Senior Software Engineer resume – IKEA Internal Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib, html

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Senior_Software_Engineer_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Senior Software Engineer with 8+ years of experience designing, developing, and maintaining "
        "scalable digital solutions in cloud environments. Deep IKEA domain expertise across Customer Support "
        "systems (IKEA Customer Connect, VCS, CSSP, Genesys, Verint, IKEA App, Startcus). "
        "Proven track record building vendor-agnostic middleware architectures impacting 32+ markets, "
        "driving test-driven development, and leading DevOps transformation from legacy monoliths to "
        "event-driven microservices. Strong believer in simplicity, togetherness, and cost-consciousness — "
        "committed to giving co-workers the best user experience possible while continuously improving "
        "digital products through data-driven decisions. Experienced in mentoring team members, "
        "collaborating cross-functionally, and translating business needs into robust technical solutions."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Technologies")
    skills_data = [
        ("Frontend", "React · TypeScript · Next.js · Responsive Web Apps · Accessibility (a11y) · Performance Optimization"),
        ("Backend & APIs", "Node.js · TypeScript · RESTful API Design · Middleware Architecture · Event-Driven · Golang (familiar)"),
        ("Cloud & DevOps", "GCP (Cloud Run, Cloud Functions, BigQuery, GKE, Pub/Sub) · AWS (Lambda, S3, EKS, CloudWatch) · Terraform · Docker · Kubernetes · GitHub Actions · CI/CD"),
        ("Testing & Quality", "Playwright (UI + Accessibility) · Vitest · Jest · TDD · Code Reviews · Test Automation Frameworks"),
        ("Data & Observability", "ETL Pipelines · Python · SQL · BigQuery · Grafana · Cloud Monitoring · Logging & Alerting"),
        ("Tools & Practices", "Ingka DevOps Tooling · Jira · Confluence · Agile/Scrum · Release Management · Vendor Management"),
        ("Certifications", "AWS Cloud Practitioner · Google Cloud Associate Cloud Engineer · ISTQB CTFL"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.2)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA Customer Connect - VCS Team Lead Acting
    add_role(doc, "Team Lead Acting — Visual Customer Support (VCS)", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2023 – Present")
    bullets_1 = [
        "Designed, developed, and maintained scalable and secure digital solutions for Visual Customer Support (VCS) in a cloud-native environment using TypeScript (React, Node.js) and GCP.",
        "Developed vendor-agnostic middleware API layer enabling plug-and-play integration of video support solutions — architected for robustness so any vendor can be swapped without disruption across 32 markets. (Ref: https://allen.ingka.com/catalog/default/api/vcs | Architecture: https://confluence.build.ingka.ikea.com/spaces/VCS/pages/860570582/VCS+-+Backend+Architecture)",
        "Ensured good code quality and test-driven development — authored Playwright scripts for comprehensive UI and accessibility testing, achieving 90%+ coverage on critical user flows.",
        "Built and streamlined CI/CD pipelines using Ingka DevOps tooling (GitHub Actions, Terraform), ensuring infrastructure-as-code practices and thorough code reviews on every PR.",
        "Served as Release Expert for the VCS platform — managing release cycles, deployment coordination, and rollback strategies across multiple environments and markets.",
        "Mentored and coached team members on best practices in software engineering, clean code principles, and IKEA development standards.",
        "Collaborated closely with cross-domain teams (VIAM, ICM, BOKA, IKEA App, Chatbot) to ensure seamless integration within the Customer Support ecosystem.",
        "Active interaction with Business Owner — delivered bi-weekly progress updates, facilitated technical discussions, and ensured alignment between engineering delivery and business priorities.",
        "Provided technical inputs to Product Manager (Cecilia) and Specialists (Anna Kack) on product roadmap, feasibility assessments, and technical debt prioritization.",
        "Managed vendor relationships and consultant coordination — created release plans, aligned delivery timelines, and ensured vendor accountability.",
        "Developed API and Data Layer (ETL pipeline) for VCS analytics — shared processed data with the central data team for cross-organizational insights. (Ref: https://datacatalog.ingka.com/data/128/ | Docs: https://confluence.build.ingka.ikea.com/spaces/VCS/pages/1153056481/VCS+-+Data+Layer)",
        "Leveraged event-driven architecture patterns for real-time data processing and implemented application observability using Cloud Monitoring, logging, and alerting.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - IKEA CSSP
    add_role(doc, "Senior Software Engineer — CSSP (Customer Support Staff Planning)", "IKEA, Ingka Digital", "Malmö, Sweden", "2022 – 2023")
    bullets_2 = [
        "Developed and maintained the Customer Support Staff Planning (CSSP) system integrating with Verint and Genesys platforms for workforce optimization across global contact centers.",
        "Built RESTful APIs and React-based interfaces for scheduling, forecasting, and real-time adherence monitoring used by planning teams worldwide.",
        "Optimized backend performance through caching, indexing, and query optimization — reducing response times by 40% for high-traffic planning endpoints.",
        "Wrote Infrastructure as Code (Terraform) for cloud resources, optimizing for cost efficiency and scalability within IKEA's cloud governance framework.",
        "Set up and managed CI/CD pipelines using GitHub Actions, enabling automated testing, security scanning, and continuous deployment.",
        "Collaborated with Verint and Genesys specialists to design integration architectures that aligned with IKEA's event-driven transformation strategy.",
        "Helped identify, measure, and automate data points to improve overall software engineering process and product quality.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - Truecaller
    add_role(doc, "Release Expert / Software Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_3 = [
        "Served as Release Expert for a globally distributed application serving 300M+ users — managing release trains, feature flags, and phased rollouts.",
        "Developed and maintained cloud-native microservices (Node.js, TypeScript) with event-driven communication patterns.",
        "Built CI/CD automation and release coordination tooling, reducing release cycle time by 50%.",
        "Collaborated cross-functionally with product, engineering, and QA teams to ensure release quality and on-time delivery.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - HCLTech for IKEA & LEGO
    add_role(doc, "Software Engineer / Technical Lead", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_4 = [
        "Developed and maintained digital solutions for IKEA smart-home (DIRIGERA/TRÅDFRI) and LEGO connected products — full-stack TypeScript (React, Node.js) and Python.",
        "Designed and deployed RESTful APIs serving multiple frontend applications and third-party integrations across the IKEA ecosystem.",
        "Built serverless and containerized applications on GCP and AWS, leveraging event-driven architectures for real-time IoT data processing.",
        "Implemented CI/CD pipelines (Jenkins, GitHub Actions) with automated testing, security scanning, and Terraform-based infrastructure provisioning.",
        "Coached and mentored junior and mid-level engineers — established code review practices, pair programming sessions, and knowledge sharing forums.",
        "Collaborated with IKEA product managers on technical roadmap inputs, feasibility analysis, and architecture decisions.",
        "Worked on Startcus Spare Parts system — developed interfaces for spare parts ordering and tracking integrated with IKEA's logistics backend.",
        "Integrated with IKEA App ecosystem — building APIs and webhooks consumed by mobile and web frontends across multiple markets.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # Role 5 - Earlier
    add_role(doc, "Software Engineer", "Earlier Career (Wipro, CSC)", "India", "2008 – 2013")
    bullets_5 = [
        "Developed web applications and backend services for enterprise clients using JavaScript, Python, and Java.",
        "Built data integration pipelines and automated reporting solutions for financial and retail domains.",
        "Participated in migration of legacy monolithic applications to service-oriented architectures.",
    ]
    for b in bullets_5:
        bullet(doc, b)

    # ─── IKEA Projects ─────────────────────────────────────────────────────
    add_heading_block(doc, "Key IKEA Projects & Systems")
    projects = [
        "IKEA Customer Connect — VCS (Visual Customer Support): Vendor-agnostic video support platform enabling face-to-face customer assistance across 32 markets via omnichannel (self-service, contact center, in-store).",
        "CSSP (Customer Support Staff Planning): Workforce management integration with Verint and Genesys for scheduling, forecasting, and real-time adherence across global Customer Support Centers.",
        "IKEA App Integration: APIs and event-driven services consumed by the IKEA mobile app for customer support flows, returns, and self-service features.",
        "Chatbot & Self-Service: Technical integration enabling automated customer support through chatbot flows connected to backend systems.",
        "Startcus — Spare Parts System: Digital solution for spare parts ordering, tracking, and fulfillment integrated with IKEA's logistics and warehouse systems.",
        "Genesys & Verint Platform: Integration engineering for contact center telephony (Genesys) and workforce optimization (Verint) within IKEA's Customer Support landscape.",
    ]
    for proj in projects:
        bullet(doc, proj, bold_prefix=proj.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Togetherness: Thrive in cross-functional collaboration — actively work with VIAM, ICM, BOKA, App teams, and business stakeholders to find the best solutions together.",
        "Simplicity: Advocate for simple, maintainable architectures — vendor-agnostic design that reduces complexity and enables teams to move fast.",
        "Cost-Consciousness: Optimize cloud resources for cost efficiency, build reusable solutions, and avoid over-engineering.",
        "Leading by Example: Mentor team members, drive code quality standards, and model the engineering practices expected across the organization.",
        "Constant Improvement: Continuously measure, learn, and iterate — use data to improve both the software engineering process and the digital products we deliver.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:20%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Senior Software Engineer with 8+ years of experience designing, developing, and maintaining scalable digital solutions in cloud environments. Deep IKEA domain expertise across Customer Support systems (IKEA Customer Connect, VCS, CSSP, Genesys, Verint, IKEA App, Startcus). Proven track record building vendor-agnostic middleware architectures impacting 32+ markets, driving test-driven development, and leading DevOps transformation from legacy monoliths to event-driven microservices. Strong believer in simplicity, togetherness, and cost-consciousness — committed to giving co-workers the best user experience possible while continuously improving digital products through data-driven decisions. Experienced in mentoring team members, collaborating cross-functionally, and translating business needs into robust technical solutions.</p>

<h2>KEY SKILLS &amp; TECHNOLOGIES</h2>
<table>
<tr><td class="cat">Frontend</td><td>React · TypeScript · Next.js · Responsive Web Apps · Accessibility (a11y) · Performance Optimization</td></tr>
<tr><td class="cat">Backend &amp; APIs</td><td>Node.js · TypeScript · RESTful API Design · Middleware Architecture · Event-Driven · Golang (familiar)</td></tr>
<tr><td class="cat">Cloud &amp; DevOps</td><td>GCP (Cloud Run, Cloud Functions, BigQuery, GKE, Pub/Sub) · AWS (Lambda, S3, EKS, CloudWatch) · Terraform · Docker · Kubernetes · GitHub Actions · CI/CD</td></tr>
<tr><td class="cat">Testing &amp; Quality</td><td>Playwright (UI + Accessibility) · Vitest · Jest · TDD · Code Reviews · Test Automation Frameworks</td></tr>
<tr><td class="cat">Data &amp; Observability</td><td>ETL Pipelines · Python · SQL · BigQuery · Grafana · Cloud Monitoring · Logging &amp; Alerting</td></tr>
<tr><td class="cat">Tools &amp; Practices</td><td>Ingka DevOps Tooling · Jira · Confluence · Agile/Scrum · Release Management · Vendor Management</td></tr>
<tr><td class="cat">Certifications</td><td>AWS Cloud Practitioner · Google Cloud Associate Cloud Engineer · ISTQB CTFL</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Team Lead Acting — Visual Customer Support (VCS) <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2023 – Present</span></p>
<ul>
<li>Designed, developed, and maintained scalable and secure digital solutions for Visual Customer Support (VCS) in a cloud-native environment using TypeScript (React, Node.js) and GCP.</li>
<li>Developed vendor-agnostic middleware API layer enabling plug-and-play integration of video support solutions — architected for robustness so any vendor can be swapped without disruption across 32 markets. (Ref: <a href="https://allen.ingka.com/catalog/default/api/vcs">API Catalog</a> | <a href="https://confluence.build.ingka.ikea.com/spaces/VCS/pages/860570582/VCS+-+Backend+Architecture">Architecture</a>)</li>
<li>Ensured good code quality and test-driven development — authored Playwright scripts for comprehensive UI and accessibility testing, achieving 90%+ coverage on critical user flows.</li>
<li>Built and streamlined CI/CD pipelines using Ingka DevOps tooling (GitHub Actions, Terraform), ensuring infrastructure-as-code practices and thorough code reviews on every PR.</li>
<li>Served as Release Expert for the VCS platform — managing release cycles, deployment coordination, and rollback strategies across multiple environments and markets.</li>
<li>Mentored and coached team members on best practices in software engineering, clean code principles, and IKEA development standards.</li>
<li>Collaborated closely with cross-domain teams (VIAM, ICM, BOKA, IKEA App, Chatbot) to ensure seamless integration within the Customer Support ecosystem.</li>
<li>Active interaction with Business Owner — delivered bi-weekly progress updates, facilitated technical discussions, and ensured alignment between engineering delivery and business priorities.</li>
<li>Provided technical inputs to Product Manager (Cecilia) and Specialists (Anna Kack) on product roadmap, feasibility assessments, and technical debt prioritization.</li>
<li>Managed vendor relationships and consultant coordination — created release plans, aligned delivery timelines, and ensured vendor accountability.</li>
<li>Developed API and Data Layer (ETL pipeline) for VCS analytics — shared processed data with the central data team for cross-organizational insights. (Ref: <a href="https://datacatalog.ingka.com/data/128/">Data Catalog</a> | <a href="https://confluence.build.ingka.ikea.com/spaces/VCS/pages/1153056481/VCS+-+Data+Layer">Data Layer Docs</a>)</li>
<li>Leveraged event-driven architecture patterns for real-time data processing and implemented application observability using Cloud Monitoring, logging, and alerting.</li>
</ul>

<p class="role">Senior Software Engineer — CSSP (Customer Support Staff Planning) <span class="meta">&nbsp;|&nbsp; IKEA, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – 2023</span></p>
<ul>
<li>Developed and maintained the Customer Support Staff Planning (CSSP) system integrating with Verint and Genesys platforms for workforce optimization across global contact centers.</li>
<li>Built RESTful APIs and React-based interfaces for scheduling, forecasting, and real-time adherence monitoring used by planning teams worldwide.</li>
<li>Optimized backend performance through caching, indexing, and query optimization — reducing response times by 40% for high-traffic planning endpoints.</li>
<li>Wrote Infrastructure as Code (Terraform) for cloud resources, optimizing for cost efficiency and scalability within IKEA's cloud governance framework.</li>
<li>Set up and managed CI/CD pipelines using GitHub Actions, enabling automated testing, security scanning, and continuous deployment.</li>
<li>Collaborated with Verint and Genesys specialists to design integration architectures that aligned with IKEA's event-driven transformation strategy.</li>
<li>Helped identify, measure, and automate data points to improve overall software engineering process and product quality.</li>
</ul>

<p class="role">Release Expert / Software Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Served as Release Expert for a globally distributed application serving 300M+ users — managing release trains, feature flags, and phased rollouts.</li>
<li>Developed and maintained cloud-native microservices (Node.js, TypeScript) with event-driven communication patterns.</li>
<li>Built CI/CD automation and release coordination tooling, reducing release cycle time by 50%.</li>
<li>Collaborated cross-functionally with product, engineering, and QA teams to ensure release quality and on-time delivery.</li>
</ul>

<p class="role">Software Engineer / Technical Lead <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Developed and maintained digital solutions for IKEA smart-home (DIRIGERA/TRÅDFRI) and LEGO connected products — full-stack TypeScript (React, Node.js) and Python.</li>
<li>Designed and deployed RESTful APIs serving multiple frontend applications and third-party integrations across the IKEA ecosystem.</li>
<li>Built serverless and containerized applications on GCP and AWS, leveraging event-driven architectures for real-time IoT data processing.</li>
<li>Implemented CI/CD pipelines (Jenkins, GitHub Actions) with automated testing, security scanning, and Terraform-based infrastructure provisioning.</li>
<li>Coached and mentored junior and mid-level engineers — established code review practices, pair programming sessions, and knowledge sharing forums.</li>
<li>Collaborated with IKEA product managers on technical roadmap inputs, feasibility analysis, and architecture decisions.</li>
<li>Worked on Startcus Spare Parts system — developed interfaces for spare parts ordering and tracking integrated with IKEA's logistics backend.</li>
<li>Integrated with IKEA App ecosystem — building APIs and webhooks consumed by mobile and web frontends across multiple markets.</li>
</ul>

<p class="role">Software Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (Wipro, CSC) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Developed web applications and backend services for enterprise clients using JavaScript, Python, and Java.</li>
<li>Built data integration pipelines and automated reporting solutions for financial and retail domains.</li>
<li>Participated in migration of legacy monolithic applications to service-oriented architectures.</li>
</ul>

<h2>KEY IKEA PROJECTS &amp; SYSTEMS</h2>
<ul>
<li><strong>IKEA Customer Connect — VCS (Visual Customer Support):</strong> Vendor-agnostic video support platform enabling face-to-face customer assistance across 32 markets via omnichannel (self-service, contact center, in-store).</li>
<li><strong>CSSP (Customer Support Staff Planning):</strong> Workforce management integration with Verint and Genesys for scheduling, forecasting, and real-time adherence across global Customer Support Centers.</li>
<li><strong>IKEA App Integration:</strong> APIs and event-driven services consumed by the IKEA mobile app for customer support flows, returns, and self-service features.</li>
<li><strong>Chatbot &amp; Self-Service:</strong> Technical integration enabling automated customer support through chatbot flows connected to backend systems.</li>
<li><strong>Startcus — Spare Parts System:</strong> Digital solution for spare parts ordering, tracking, and fulfillment integrated with IKEA's logistics and warehouse systems.</li>
<li><strong>Genesys &amp; Verint Platform:</strong> Integration engineering for contact center telephony (Genesys) and workforce optimization (Verint) within IKEA's Customer Support landscape.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Togetherness:</strong> Thrive in cross-functional collaboration — actively work with VIAM, ICM, BOKA, App teams, and business stakeholders to find the best solutions together.</li>
<li><strong>Simplicity:</strong> Advocate for simple, maintainable architectures — vendor-agnostic design that reduces complexity and enables teams to move fast.</li>
<li><strong>Cost-Consciousness:</strong> Optimize cloud resources for cost efficiency, build reusable solutions, and avoid over-engineering.</li>
<li><strong>Leading by Example:</strong> Mentor team members, drive code quality standards, and model the engineering practices expected across the organization.</li>
<li><strong>Constant Improvement:</strong> Continuously measure, learn, and iterate — use data to improve both the software engineering process and the digital products we deliver.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
