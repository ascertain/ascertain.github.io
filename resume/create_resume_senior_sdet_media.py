from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Senior_SDET_Media_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Senior_SDET_Media_Resume.doc"

ACCENT = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "15+ years",
    "WebRTC",
    "SIP",
    "media",
    "real-time communication",
    "test automation framework",
    "automation framework",
    "shift-left",
    "C#",
    "TypeScript",
    "Java",
    "Playwright",
    "Selenium",
    "Appium",
    "CI/CD",
    "GCP",
    "AWS",
    "VDI",
    "Windows 10",
    "Windows 11",
    "test strategy",
    "quality strategy",
    "SDET",
    "mentoring",
    "cross-functional",
    "code review",
    "media backend",
    "media frontend",
    "media technologies",
    "DevOps",
    "Git",
    "Docker",
    "Kubernetes",
    "test coverage",
    "test lead",
    "quality",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior SDET | Media & Real-Time Communication | Test Automation & Quality Leadership",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif",
        size=9.5,
        color=TEXT_MUTED,
    )
    add_text(
        contact,
        "\nPortfolio: ascertain.github.io/kashif/",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="7F8FA6", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior SDET with 15+ years of experience leading test strategy, framework architecture, "
        "and quality programmes across media and real-time communication platforms. Currently owning "
        "the quality strategy for a WebRTC-based platform — architecting scalable test automation "
        "frameworks, driving shift-left testing adoption, and coordinating cross-functional delivery "
        "across engineering, product, and QA teams. Proven track record in test lead roles — "
        "mentoring SDET teams, designing automation architectures (Playwright, Selenium, Appium), "
        "establishing CI/CD-integrated quality pipelines, and ensuring media backend/frontend "
        "reliability. Experienced across GCP/AWS cloud environments and VDI setups. Passionate about "
        "building high-quality media solutions through programme leadership, framework design, and "
        "engineering excellence.",
        size=10,
    )

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Media Technologies: ",
            "WebRTC, SIP, real-time communication systems, audio/video quality validation, "
            "media backend & frontend testing, codec/stream verification",
        ),
        (
            "Test Automation: ",
            "Playwright (TypeScript), Selenium, Appium (mobile), Karate, API testing, E2E automation, "
            "test automation framework design & maintenance, scalable test suites",
        ),
        (
            "Programming & Architecture: ",
            "C#, TypeScript, Java, Python — framework design, architecture patterns, code reviews",
        ),
        (
            "Quality & Strategy: ",
            "Shift-left testing, test strategy design, quality strategy ownership, test coverage analysis, "
            "TDD/BDD, UAT, BAT, CI/CD integration, continuous testing",
        ),
        (
            "Cloud & Infrastructure: ",
            "GCP (BigQuery, Cloud Run, Pub/Sub), AWS, VDI environments, Docker, "
            "Terraform, Git, GitHub Actions",
        ),
        (
            "Leadership: ",
            "SDET mentoring, cross-functional collaboration, test process improvement, coordination "
            "across teams/products, Agile/Scrum, stakeholder management",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")
    jobs = [
        (
            "IKEA IT AB, Malmö, Sweden",
            "Senior SDET / Software Engineer — Media & Real-Time Communication | Mar 2022 – Present",
            [
                "Own the quality strategy for VCS (Visual Collaboration Solution) — a WebRTC-based "
                "real-time communication platform integrating CoBrowser, Real-Time Transcription, and "
                "AI-powered Call Quality Assessment across media backend and frontend components.",
                "Design, implement, and maintain a scalable test automation framework using Playwright "
                "(TypeScript) — covering E2E, API, and media-specific test scenarios with full CI/CD "
                "integration via GitHub Actions on GCP infrastructure.",
                "Validate and troubleshoot WebRTC media streams — testing audio, video, screen-sharing, "
                "and SIP-based call flows across browsers, devices, and OS configurations including "
                "VDI environments.",
                "Drive shift-left testing practices across the engineering organization — collaborating "
                "with developers, product managers, and QA engineers to integrate quality early in the "
                "development lifecycle.",
                "Perform code reviews, mentor SDETs, and ensure adherence to testing best practices "
                "across the team. Coordinate with multiple products for seamless integration quality.",
                "Built a data layer on GCP (BigQuery, Pub/Sub, Cloud Functions) for media telemetry, "
                "test metrics, and quality analytics. Recognized as Exceptional Performer.",
            ],
        ),
        (
            "Truecaller, Sweden",
            "Release & Automation Engineer | Sep 2021 – Feb 2022",
            [
                "Led test automation for Android and iOS apps serving 300M+ users — testing media "
                "features (VoIP calls, audio quality) across 50+ device models and OS versions.",
                "Designed test strategy for release readiness across diverse device configurations. "
                "Built automated suites using Appium and Selenium with CI/CD pipeline integration.",
                "Collaborated cross-functionally on quality strategy, defect triage, and release coordination.",
            ],
        ),
        (
            "LEGO and IKEA Group (via HCLTech), Denmark & Sweden",
            "Senior SDET / Test Lead | 2016 – 2021",
            [
                "Led test automation framework design for LEGO digital commerce and IKEA App — "
                "building scalable frameworks using Selenium, Appium, and Karate (C#/Java) with "
                "full CI/CD integration.",
                "Owned quality strategy for multi-platform applications — driving shift-left testing, "
                "test coverage analysis, and cross-team coordination across backend and frontend.",
                "Served as test lead managing 8+ SDETs — mentoring on automation best practices, "
                "performing code reviews, and driving Agile delivery across global teams.",
            ],
        ),
        (
            "HCLTech / Enterprise Programs",
            "Automation Lead / SDET | Dec 2013 – 2014",
            [
                "Led automation framework design and SDET initiatives across enterprise programmes — "
                "scalable C#/Java frameworks with CI/CD pipelines, reducing manual effort by 70%.",
                "Mentored 15+ engineers on test automation architecture, shift-left practices, and "
                "quality-driven development across distributed global teams.",
            ],
        ),
        (
            "Samin TekMindz India Pvt. Ltd. — Banking & Government",
            "SDET | 2011 – 2013",
            [
                "Designed and maintained test automation for SW Global — AMA Accra Metropolitan "
                "Assembly, Ghana payment system — a complex web + POS integration. Validated "
                "real-time transaction flows across hardware/software.",
                "Drove test process improvements, authored test strategies, and ensured quality "
                "delivery for government payment platform across multiple environments.",
            ],
        ),
        (
            "Earlier Career — Banking & Government",
            "Software Test Engineer / SDET | 2008 – 2011",
            [
                "Built test automation frameworks for core banking (Finacle) and government platforms — "
                "validating real-time transaction systems across web, mobile, and backend components.",
                "Led UAT/BAT cycles with full test coverage tracking. Tested Android and iOS banking "
                "apps across multiple device configurations and Windows environments.",
            ],
        ),
    ]
    for company, title, bullets in jobs:
        cp = document.add_paragraph()
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after = Pt(0)
        add_text(cp, company, bold=True, size=10, color=SECTION_COLOR)

        tp = document.add_paragraph()
        tp.paragraph_format.space_after = Pt(0)
        add_text(tp, title, bold=True, size=10)

        for bullet in bullets:
            add_highlighted_bullet(document, bullet)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Designed and own the quality strategy for a WebRTC-based real-time communication platform — "
        "ensuring media backend and frontend reliability across production releases.",
        "Built scalable test automation frameworks (Playwright/TypeScript) with CI/CD — achieving "
        "3x reliability improvement and 50% faster pipeline execution.",
        "Validated real-time media (WebRTC, SIP) across browsers, devices, and VDI environments — "
        "99.5% media quality compliance.",
        "Drove shift-left testing adoption across engineering — reducing defect escape rate by 60% "
        "through early quality integration.",
        "Led and mentored SDET teams of 8-15+ engineers across multiple programmes — driving "
        "automation best practices and cross-functional quality culture.",
    ]
    for achievement in achievements:
        add_highlighted_bullet(document, achievement)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\nGoogle Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\nITIL Foundation\nCertified Ethical Hacker (CEH)\n"
        "Six Sigma Green Belt\nUiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Senior SDET Media Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #1f4e79; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior SDET | Media &amp; Real-Time Communication | Test Automation &amp; Quality Leadership</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif<br>Portfolio: ascertain.github.io/kashif/</div>

  <div class="section">Profile</div>
  <p>Senior <span class="hl">SDET</span> with <span class="hl">15+ years</span> of experience leading <span class="hl">test strategy</span>, framework architecture, and <span class="hl">quality</span> programmes across <span class="hl">media</span> and <span class="hl">real-time communication</span> platforms. Currently owning the <span class="hl">quality strategy</span> for a <span class="hl">WebRTC</span>-based platform — architecting scalable <span class="hl">test automation frameworks</span>, driving <span class="hl">shift-left</span> testing adoption, and coordinating <span class="hl">cross-functional</span> delivery. Proven track record in <span class="hl">test lead</span> roles — <span class="hl">mentoring</span> <span class="hl">SDET</span> teams, designing automation architectures (<span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Appium</span>), establishing <span class="hl">CI/CD</span>-integrated quality pipelines, and ensuring <span class="hl">media backend</span>/frontend reliability. Experienced across <span class="hl">GCP</span>/<span class="hl">AWS</span> and <span class="hl">VDI</span> environments.</p>

  <div class="section">Core Competencies</div>
  <p><b>Media Technologies:</b> WebRTC, SIP, real-time communication, audio/video quality validation, media backend &amp; frontend testing, codec/stream verification<br>
  <b>Test Automation:</b> Playwright (TypeScript), Selenium, Appium, Karate, API testing, E2E automation, framework design &amp; maintenance<br>
  <b>Programming &amp; Architecture:</b> C#, TypeScript, Java, Python — framework design, architecture patterns, code reviews<br>
  <b>Quality &amp; Strategy:</b> Shift-left testing, test strategy, quality strategy, test coverage analysis, TDD/BDD, UAT, BAT, CI/CD, continuous testing<br>
  <b>Cloud &amp; Infrastructure:</b> GCP (BigQuery, Cloud Run, Pub/Sub), AWS, VDI, Docker, Terraform, Git, GitHub Actions<br>
  <b>Leadership:</b> SDET mentoring, cross-functional collaboration, test process improvement, multi-team coordination, Agile/Scrum</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö, Sweden</div>
  <div class="job-sub">Senior SDET / Software Engineer — Media &amp; Real-Time Communication | Mar 2022 – Present</div>
  <ul>
    <li>Own <span class="hl">quality strategy</span> for VCS — a <span class="hl">WebRTC</span>-based <span class="hl">real-time communication</span> platform integrating CoBrowser, Real-Time Transcription, and AI-powered Call Quality Assessment across <span class="hl">media backend</span> and frontend.</li>
    <li>Design, implement, and maintain a scalable <span class="hl">test automation framework</span> using <span class="hl">Playwright</span> (<span class="hl">TypeScript</span>) — E2E, API, and media-specific tests with <span class="hl">CI/CD</span> via GitHub Actions on <span class="hl">GCP</span>.</li>
    <li>Validate and troubleshoot <span class="hl">WebRTC</span> media streams — audio, video, screen-sharing, and <span class="hl">SIP</span>-based call flows across browsers, devices, and <span class="hl">VDI</span> environments.</li>
    <li>Drive <span class="hl">shift-left</span> testing practices — collaborating with developers, product managers, and QA engineers for early quality integration.</li>
    <li>Perform <span class="hl">code review</span>s, mentor <span class="hl">SDET</span>s, coordinate across products for integration <span class="hl">quality</span>. Built data layer on <span class="hl">GCP</span> for media telemetry. Recognized as Exceptional Performer.</li>
  </ul>

  <div class="job-title">Truecaller, Sweden</div>
  <div class="job-sub">Release &amp; Automation Engineer | Sep 2021 – Feb 2022</div>
  <ul>
    <li>Led <span class="hl">test automation</span> for Android/iOS apps (300M+ users) — testing <span class="hl">media</span> features (VoIP, audio quality) across 50+ devices and OS versions.</li>
    <li>Designed <span class="hl">test strategy</span> for release readiness. Built suites using <span class="hl">Appium</span> and <span class="hl">Selenium</span> with <span class="hl">CI/CD</span>. <span class="hl">Cross-functional</span> quality coordination.</li>
  </ul>

  <div class="job-title">LEGO and IKEA Group (via HCLTech), Denmark &amp; Sweden</div>
  <div class="job-sub">Senior SDET / Test Lead | 2016 – 2021</div>
  <ul>
    <li>Led <span class="hl">test automation framework</span> design — <span class="hl">Selenium</span>, <span class="hl">Appium</span>, Karate (<span class="hl">C#</span>/<span class="hl">Java</span>) with <span class="hl">CI/CD</span>. Owned <span class="hl">quality strategy</span> driving <span class="hl">shift-left</span> and <span class="hl">test coverage</span> analysis.</li>
    <li><span class="hl">Test lead</span> managing 8+ <span class="hl">SDET</span>s — <span class="hl">mentoring</span>, <span class="hl">code review</span>s, and Agile delivery across global teams.</li>
  </ul>

  <div class="job-title">HCLTech / Enterprise Programs</div>
  <div class="job-sub">Automation Lead / SDET | Dec 2013 – 2014</div>
  <ul>
    <li>Led <span class="hl">automation framework</span> design (<span class="hl">C#</span>/<span class="hl">Java</span>) with <span class="hl">CI/CD</span> — 70% manual effort reduction. Mentored 15+ engineers on <span class="hl">shift-left</span> and <span class="hl">test automation</span>.</li>
  </ul>

  <div class="job-title">Samin TekMindz India Pvt. Ltd. — Banking &amp; Government</div>
  <div class="job-sub">SDET | 2011 – 2013</div>
  <ul>
    <li>Designed <span class="hl">test automation</span> for SW Global — AMA Accra Metropolitan Assembly, Ghana payment system — complex web + POS integration. Validated real-time transaction flows.</li>
    <li>Drove test process improvements, authored <span class="hl">test strategy</span>, and ensured <span class="hl">quality</span> delivery across multiple environments.</li>
  </ul>

  <div class="job-title">Earlier Career — Banking &amp; Government</div>
  <div class="job-sub">Software Test Engineer / SDET | 2008 – 2011</div>
  <ul>
    <li>Built <span class="hl">test automation frameworks</span> for core banking (Finacle) — real-time transaction systems, mobile apps, <span class="hl">Windows</span> environments. Led UAT/BAT with full <span class="hl">test coverage</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Own <span class="hl">quality strategy</span> for <span class="hl">WebRTC</span> <span class="hl">real-time communication</span> platform — <span class="hl">media backend</span> and frontend reliability.</li>
    <li>Built scalable <span class="hl">test automation frameworks</span> (<span class="hl">Playwright</span>/<span class="hl">TypeScript</span>) — 3x reliability, 50% faster <span class="hl">CI/CD</span> pipelines.</li>
    <li>Validated <span class="hl">WebRTC</span>/<span class="hl">SIP</span> media across browsers, devices, and <span class="hl">VDI</span> environments — 99.5% media quality compliance.</li>
    <li>Drove <span class="hl">shift-left</span> testing adoption — 60% reduction in defect escape rate.</li>
    <li>Led and mentored <span class="hl">SDET</span> teams of 8-15+ engineers — <span class="hl">cross-functional</span> quality culture.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>ITIL Foundation<br>Certified Ethical Hacker<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
