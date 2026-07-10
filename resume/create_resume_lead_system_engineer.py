"""Generate Lead System Engineer resume – IKEA Internal Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Lead_System_Engineer_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Lead System Engineer with 12+ years of experience driving technical leadership, "
        "system design, and engineering excellence across enterprise-scale platforms within the IKEA ecosystem. "
        "Combines deep technical expertise in cloud infrastructure, solution architecture, and distributed systems "
        "with a passion for coaching, mentoring, and enabling product teams to deliver at their best. "
        "Proven track record of harmonizing technical landscapes — championing design patterns, reusability, "
        "reliability, scalability, and security by design across sub-domains and product areas. "
        "Skilled at aligning technical roadmaps with strategic direction, collaborating closely with architects, "
        "engineering managers, and leadership to foster an engineering culture rooted in continuous improvement. "
        "Strong ambassador for engineering best practices both within IKEA and in the broader market. "
        "Thrives on enabling others to succeed through a collaborative, co-creative approach — "
        "with a DevOps mindset and a focus on operational resilience. "
        "Deeply aligned with IKEA values: togetherness, diversity, and leading by example."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Technical Leadership", "Holistic system oversight · Design pattern adoption · Technical roadmap alignment · Engineering culture ambassador · Cross-domain harmonization · Driving complex technical changes"),
        ("Coaching & Mentoring", "Team enablement · Knowledge sharing · Engineering community building · Code reviews · Pair programming · Fostering curiosity and continuous learning"),
        ("Enterprise Infrastructure", "Cloud computing (GCP, AWS) · Distributed systems · Microservices · Event-driven architecture · API design · Reliability & scalability engineering · Security by design"),
        ("Solution Architecture", "System design & refactoring · Modularization · Integration patterns · Data pipelines · Platform & SaaS landscape · Home-grown solutions · Resiliency design"),
        ("DevOps & Platform", "Terraform · Docker · Kubernetes · CI/CD (GitHub Actions) · Monitoring & observability · Infrastructure as Code · GitOps · Operational resilience"),
        ("Stakeholder Engagement", "Cross-functional collaboration · Communicating technical matters to diverse audiences · Influencing senior stakeholders · Driving change in large organizations"),
        ("Agile & Ways of Working", "Product team model · Scrum/Kanban · DevOps working set-up · Self-organizing teams · Iterative delivery · IKEA ways of working"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS
    add_role(doc, "Technical Lead / Team Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Maintain a holistic technical view across the sub-domain — ensuring solutioning and operations align across product teams, driving consistency in architecture decisions and engineering practices.",
        "Champion harmonization of the technical landscape — driving adoption of shared design patterns, reusable components, and engineering standards to enhance reliability, scalability, and security by design.",
        "Lead, coach, and mentor a team of engineers — building engagement around technical direction, fostering curiosity, and enabling teams to navigate challenges with resilience and creativity.",
        "Act as an ambassador for engineering and security best practices across the sub-domain and the broader IKEA organization — influencing engineering culture and contributing to community forums.",
        "Collaborate closely with Enterprise Architecture to design, refactor, and modularize products — aligning to enterprise guidelines for scalability, integration, resiliency, and reusability.",
        "Align technical roadmap with strategic direction — bridging business objectives, architecture goals, and operational realities into actionable engineering initiatives.",
        "Drive complex technical changes across products — from system design and solution architecture to hands-on implementation (Python, TypeScript, GCP, BigQuery, Terraform, Kubernetes).",
        "Provide IKEA business model knowledge and guide teams on how it reflects in the technology landscape — enabling efficient information flows and data-driven decisions for 32 markets.",
        "Work closely with product-background colleagues to augment ideation with deep technical knowledge — contributing to exploration and adoption of new technologies improving system performance.",
        "Communicate technical matters differently depending on stakeholders — from engineers to senior leadership — building strong relations across domains and disciplines.",
        "Champion DevOps mindset and continuous growth — guiding teams toward operational excellence through CI/CD, observability, Infrastructure as Code, and iterative improvement.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Platform Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Designed and operated scalable platform infrastructure supporting a globally distributed application (300M+ users) — ensuring reliability, performance, and security at scale.",
        "Drove harmonization of deployment pipelines and infrastructure patterns — improving consistency and reducing operational overhead across engineering teams.",
        "Collaborated cross-functionally with product and engineering teams in an agile, DevOps-centric environment — enabling rapid iteration and continuous delivery.",
        "Contributed to cloud infrastructure design and CI/CD automation — reducing deployment time by 50% through pipeline optimization and Infrastructure as Code.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Technical Lead — Systems & Infrastructure", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led and mentored a team of 8–12 engineers — providing technical direction, coaching, building engagement around strategic and tactical engineering goals, and fostering continuous learning.",
        "Maintained holistic technical oversight across multiple products — driving design pattern adoption, reusability, and harmonization of the technical landscape across IKEA and LEGO platforms.",
        "Designed and implemented enterprise-scale solutions combining Platforms, SaaS, and home-grown systems — deep experience in solution architecture for complex, integrated environments.",
        "Drove complex technical changes — system redesigns, cloud migrations (on-prem to GCP/AWS), and modularization initiatives improving scalability, reliability, and maintainability.",
        "Collaborated closely with architects, product teams, and business stakeholders — translating strategic direction into technical roadmaps and actionable engineering initiatives.",
        "Championed engineering best practices across the organization — design patterns, security by design, automated testing, observability, documentation, and operational resilience.",
        "Built strong relations with stakeholders across domains and disciplines — communicating technical matters at different levels from engineers to senior leadership.",
        "Worked within the IKEA ecosystem (IKEA App, Customer Support, Spare Parts, IoT) — providing extensive IKEA business model knowledge and guiding teams on efficient information flows.",
        "Drove adoption of DevOps practices, CI/CD pipelines, and Infrastructure as Code — establishing operational foundations that teams could build upon independently.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_4 = [
        "Built enterprise software solutions across financial and retail domains — gaining foundational experience in system design, integration patterns, and distributed systems.",
        "Developed and maintained platform components, APIs, and integration layers using Java, Python, and SQL — supporting complex business workflows.",
        "Participated in system modernization and migration projects — moving from legacy architectures to scalable, cloud-ready platforms.",
        "Worked in agile, cross-functional teams — building collaborative engineering habits and iterative delivery practices.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── IKEA Ecosystem Experience ─────────────────────────────────────────
    add_heading_block(doc, "IKEA Ecosystem Experience")
    ikea_items = [
        "IKEA Customer Connect (VCS): End-to-end technical leadership — system design, cloud infrastructure (GCP), data pipelines, API platforms serving 32 markets within Customer Support.",
        "Engineering Culture: Active contributor to IKEA's engineering community — championing best practices, design patterns, knowledge sharing, and fostering curiosity across product teams.",
        "IKEA Ways of Working: Product team model, Agile/Scrum, DevOps working set-up, Ingka DevOps tooling, cross-functional collaboration with architects, product managers, and leadership.",
        "Platform & SaaS Landscape: Deep experience navigating IKEA's technology landscape combining Platforms, SaaS (Genesys, Verint, Salesforce), and home-grown solutions.",
        "IKEA App & Connected Products: Infrastructure and data engineering for IoT telemetry and smart-home product systems (DIRIGERA/TRÅDFRI ecosystem).",
        "Technical Roadmap Alignment: Partnering with architects and leadership to align engineering initiatives with IKEA's strategic direction — ensuring scalability, security, and business alignment.",
        "Operational Resilience: Driving reliability engineering, monitoring, observability, and incident response practices within the IKEA ecosystem.",
    ]
    for item in ikea_items:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Togetherness: Build strong cross-functional relationships — collaborate with architects, product teams, and leadership to enable high-standard engineering across the organization.",
        "Leading by Example: Coach, mentor, and inspire engineers — championing engineering excellence, curiosity, and a culture of continuous learning.",
        "Diversity & Inclusion: Embrace diversity, equality, and inclusion — creating a workplace where everyone feels valued, inspired, and empowered to contribute.",
        "Simplicity: Drive harmonization and reuse — design pragmatic, maintainable solutions with clear standards that teams can adopt and build upon.",
        "Cost-Consciousness: Optimise infrastructure and operations for cost-efficiency — balancing performance, reliability, and sustainability.",
        "Daring to be Different: Challenge norms, foster curiosity, and champion innovation — exploring new technologies to improve system performance and enable business opportunities.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Lead System Engineer with 12+ years of experience driving technical leadership, system design, and engineering excellence across enterprise-scale platforms within the IKEA ecosystem. Combines deep technical expertise in cloud infrastructure, solution architecture, and distributed systems with a passion for coaching, mentoring, and enabling product teams to deliver at their best. Proven track record of harmonizing technical landscapes — championing design patterns, reusability, reliability, scalability, and security by design across sub-domains and product areas. Skilled at aligning technical roadmaps with strategic direction, collaborating closely with architects, engineering managers, and leadership to foster an engineering culture rooted in continuous improvement. Strong ambassador for engineering best practices both within IKEA and in the broader market. Thrives on enabling others to succeed through a collaborative, co-creative approach — with a DevOps mindset and a focus on operational resilience. Deeply aligned with IKEA values: togetherness, diversity, and leading by example.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Technical Leadership</td><td>Holistic system oversight · Design pattern adoption · Technical roadmap alignment · Engineering culture ambassador · Cross-domain harmonization · Driving complex technical changes</td></tr>
<tr><td class="cat">Coaching &amp; Mentoring</td><td>Team enablement · Knowledge sharing · Engineering community building · Code reviews · Pair programming · Fostering curiosity and continuous learning</td></tr>
<tr><td class="cat">Enterprise Infrastructure</td><td>Cloud computing (GCP, AWS) · Distributed systems · Microservices · Event-driven architecture · API design · Reliability &amp; scalability engineering · Security by design</td></tr>
<tr><td class="cat">Solution Architecture</td><td>System design &amp; refactoring · Modularization · Integration patterns · Data pipelines · Platform &amp; SaaS landscape · Home-grown solutions · Resiliency design</td></tr>
<tr><td class="cat">DevOps &amp; Platform</td><td>Terraform · Docker · Kubernetes · CI/CD (GitHub Actions) · Monitoring &amp; observability · Infrastructure as Code · GitOps · Operational resilience</td></tr>
<tr><td class="cat">Stakeholder Engagement</td><td>Cross-functional collaboration · Communicating technical matters to diverse audiences · Influencing senior stakeholders · Driving change in large organizations</td></tr>
<tr><td class="cat">Agile &amp; Ways of Working</td><td>Product team model · Scrum/Kanban · DevOps working set-up · Self-organizing teams · Iterative delivery · IKEA ways of working</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Technical Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Maintain a holistic technical view across the sub-domain — ensuring solutioning and operations align across product teams, driving consistency in architecture decisions and engineering practices.</li>
<li>Champion harmonization of the technical landscape — driving adoption of shared design patterns, reusable components, and engineering standards to enhance reliability, scalability, and security by design.</li>
<li>Lead, coach, and mentor a team of engineers — building engagement around technical direction, fostering curiosity, and enabling teams to navigate challenges with resilience and creativity.</li>
<li>Act as an ambassador for engineering and security best practices across the sub-domain and the broader IKEA organization — influencing engineering culture and contributing to community forums.</li>
<li>Collaborate closely with Enterprise Architecture to design, refactor, and modularize products — aligning to enterprise guidelines for scalability, integration, resiliency, and reusability.</li>
<li>Align technical roadmap with strategic direction — bridging business objectives, architecture goals, and operational realities into actionable engineering initiatives.</li>
<li>Drive complex technical changes across products — from system design and solution architecture to hands-on implementation (Python, TypeScript, GCP, BigQuery, Terraform, Kubernetes).</li>
<li>Provide IKEA business model knowledge and guide teams on how it reflects in the technology landscape — enabling efficient information flows and data-driven decisions for 32 markets.</li>
<li>Work closely with product-background colleagues to augment ideation with deep technical knowledge — contributing to exploration and adoption of new technologies improving system performance.</li>
<li>Communicate technical matters differently depending on stakeholders — from engineers to senior leadership — building strong relations across domains and disciplines.</li>
<li>Champion DevOps mindset and continuous growth — guiding teams toward operational excellence through CI/CD, observability, Infrastructure as Code, and iterative improvement.</li>
</ul>

<p class="role">Platform Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Designed and operated scalable platform infrastructure supporting a globally distributed application (300M+ users) — ensuring reliability, performance, and security at scale.</li>
<li>Drove harmonization of deployment pipelines and infrastructure patterns — improving consistency and reducing operational overhead across engineering teams.</li>
<li>Collaborated cross-functionally with product and engineering teams in an agile, DevOps-centric environment — enabling rapid iteration and continuous delivery.</li>
<li>Contributed to cloud infrastructure design and CI/CD automation — reducing deployment time by 50% through pipeline optimization and Infrastructure as Code.</li>
</ul>

<p class="role">Technical Lead — Systems &amp; Infrastructure <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led and mentored a team of 8–12 engineers — providing technical direction, coaching, building engagement around strategic and tactical engineering goals, and fostering continuous learning.</li>
<li>Maintained holistic technical oversight across multiple products — driving design pattern adoption, reusability, and harmonization of the technical landscape across IKEA and LEGO platforms.</li>
<li>Designed and implemented enterprise-scale solutions combining Platforms, SaaS, and home-grown systems — deep experience in solution architecture for complex, integrated environments.</li>
<li>Drove complex technical changes — system redesigns, cloud migrations (on-prem to GCP/AWS), and modularization initiatives improving scalability, reliability, and maintainability.</li>
<li>Collaborated closely with architects, product teams, and business stakeholders — translating strategic direction into technical roadmaps and actionable engineering initiatives.</li>
<li>Championed engineering best practices across the organization — design patterns, security by design, automated testing, observability, documentation, and operational resilience.</li>
<li>Built strong relations with stakeholders across domains and disciplines — communicating technical matters at different levels from engineers to senior leadership.</li>
<li>Worked within the IKEA ecosystem (IKEA App, Customer Support, Spare Parts, IoT) — providing extensive IKEA business model knowledge and guiding teams on efficient information flows.</li>
<li>Drove adoption of DevOps practices, CI/CD pipelines, and Infrastructure as Code — establishing operational foundations that teams could build upon independently.</li>
</ul>

<p class="role">Software Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built enterprise software solutions across financial and retail domains — gaining foundational experience in system design, integration patterns, and distributed systems.</li>
<li>Developed and maintained platform components, APIs, and integration layers using Java, Python, and SQL — supporting complex business workflows.</li>
<li>Participated in system modernization and migration projects — moving from legacy architectures to scalable, cloud-ready platforms.</li>
<li>Worked in agile, cross-functional teams — building collaborative engineering habits and iterative delivery practices.</li>
</ul>

<h2>IKEA ECOSYSTEM EXPERIENCE</h2>
<ul>
<li><strong>IKEA Customer Connect (VCS):</strong> End-to-end technical leadership — system design, cloud infrastructure (GCP), data pipelines, API platforms serving 32 markets within Customer Support.</li>
<li><strong>Engineering Culture:</strong> Active contributor to IKEA's engineering community — championing best practices, design patterns, knowledge sharing, and fostering curiosity across product teams.</li>
<li><strong>IKEA Ways of Working:</strong> Product team model, Agile/Scrum, DevOps working set-up, Ingka DevOps tooling, cross-functional collaboration with architects, product managers, and leadership.</li>
<li><strong>Platform &amp; SaaS Landscape:</strong> Deep experience navigating IKEA's technology landscape combining Platforms, SaaS (Genesys, Verint, Salesforce), and home-grown solutions.</li>
<li><strong>IKEA App &amp; Connected Products:</strong> Infrastructure and data engineering for IoT telemetry and smart-home product systems (DIRIGERA/TRÅDFRI ecosystem).</li>
<li><strong>Technical Roadmap Alignment:</strong> Partnering with architects and leadership to align engineering initiatives with IKEA's strategic direction — ensuring scalability, security, and business alignment.</li>
<li><strong>Operational Resilience:</strong> Driving reliability engineering, monitoring, observability, and incident response practices within the IKEA ecosystem.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Togetherness:</strong> Build strong cross-functional relationships — collaborate with architects, product teams, and leadership to enable high-standard engineering across the organization.</li>
<li><strong>Leading by Example:</strong> Coach, mentor, and inspire engineers — championing engineering excellence, curiosity, and a culture of continuous learning.</li>
<li><strong>Diversity &amp; Inclusion:</strong> Embrace diversity, equality, and inclusion — creating a workplace where everyone feels valued, inspired, and empowered to contribute.</li>
<li><strong>Simplicity:</strong> Drive harmonization and reuse — design pragmatic, maintainable solutions with clear standards that teams can adopt and build upon.</li>
<li><strong>Cost-Consciousness:</strong> Optimise infrastructure and operations for cost-efficiency — balancing performance, reliability, and sustainability.</li>
<li><strong>Daring to be Different:</strong> Challenge norms, foster curiosity, and champion innovation — exploring new technologies to improve system performance and enable business opportunities.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
