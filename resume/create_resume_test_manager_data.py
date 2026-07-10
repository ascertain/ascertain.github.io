"""
Generate a tailored resume for Test Manager (Data/BI/GDPR/Year-End Testing).
Focus: E2E test scenarios, test cycle optimization, GDPR/pseudonymization, test data management,
SAFe Agile, BI environments, stakeholder management, data pipelines, regulated industries.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Test_Manager_Data_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Test_Manager_Data_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "test strategy", "test data management", "GDPR", "pseudonymization",
    "data masking", "data privacy", "data pipelines",
    "SAFe", "Agile", "Scrum", "Kanban",
    "stakeholder management", "stakeholder",
    "E2E", "end-to-end", "test cycle",
    "entry and exit criteria", "quality gates",
    "defect management", "triage",
    "BI", "BigQuery", "SQL",
    "regression testing", "integration testing",
    "test environments", "test milestones",
    "Jira", "Zephyr", "Confluence",
    "regulated", "banking", "financial",
    "continuous improvement", "process improvement",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Test Manager — Data Platforms & Enterprise BI", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Test Manager with 10+ years of experience in test lifecycle management across data-intensive "
        "and regulated environments (banking, financial services, enterprise BI). Proven track record "
        "of managing Agile/SAFe teams, defining E2E test strategies, and optimizing test cycle frequency "
        "from annual to continuous delivery. Strong expertise in test data management, GDPR compliance, "
        "data masking/pseudonymization, and data pipeline validation. Experienced in leading stakeholder "
        "management across multiple departments, gathering business requirements, analyzing as-is "
        "processes, and recommending improvements for performance and efficiency. Skilled at defining "
        "entry/exit criteria, quality gates, managing test environments, and overseeing defect management "
        "and triage. Structured, analytical, and detail-oriented with a proactive focus on risk "
        "mitigation and continuous improvement."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Test Strategy & Delivery Planning", "Test Data Management & GDPR/Masking", "SAFe Agile Team Leadership (8+ yrs)",
        "E2E Test Scenarios & Cycle Optimization", "Stakeholder Management (Multi-Dept)", "BI & Data Pipeline Validation",
        "Defect Management & Triage", "Entry/Exit Criteria & Quality Gates", "Regulated Industries (Banking/Finance)",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Manager / Team Lead — Data Platform & BI Testing", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Define and maintain the overall test strategy and delivery plan for a data-intensive platform across 30+ markets — aligned with project roadmap and Agile/SAFe ceremonies.",
        "Lead E2E test scenario collection and execution across data pipelines (BigQuery, Pub/Sub, ETL), APIs, and BI layers — ensuring comprehensive test coverage for all business processes.",
        "Drive test cycle optimization: improved test frequency from quarterly cycles to continuous delivery by implementing automated regression, parallel execution, and streamlined test data provisioning.",
        "Own test data management and enablement: ensure the latest, validated data is always available to business users in the test layer — coordinating data refresh, masking, and provisioning.",
        "Ensure GDPR compliance in test environments: implemented pseudonymization and data masking for PI fields across all test data layers, addressing privacy requirements for 30+ market datasets.",
        "Collaborate with stakeholders across multiple departments to gather and document business requirements — analyzing as-is systems and proposing enhancements for test efficiency.",
        "Define entry and exit criteria, test milestones, and quality gates for each delivery phase — reporting progress and quality metrics to programme leadership.",
        "Manage test environments in alignment with environment owners — coordinating provisioning, refresh cycles, and shared environment governance.",
        "Oversee defect management, triage, and resolution tracking — driving cross-team defect resolution and maintaining traceability to business requirements.",
        "Develop and present data-driven reports and recommendations to business teams — leveraging test metrics to improve cycle frequency and decision-making quality.",
        "Lead Agile team (SAFe) through PI planning, sprint ceremonies, and backlog refinement — balancing regular delivery with strategic platform improvement initiatives.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Release & Test Manager", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed E2E test cycles for a 300M+ user platform — defining test strategy, entry/exit criteria, and quality gates for each release phase.",
        "Coordinated test data provisioning across shared environments — ensuring data freshness, GDPR compliance, and availability for parallel test initiatives.",
        "Led stakeholder management across product, engineering, and operations — facilitating workshops to explore testing needs and prioritize improvements.",
        "Drove process improvements: reduced test cycle duration by 40% through automation, environment optimization, and parallel test execution strategies.",
        "Oversaw defect management and triage — tracking resolution across squads and reporting quality status to leadership with data-driven recommendations.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior Test Manager — Enterprise BI & E-Commerce Platforms", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed Agile test teams (SAFe/Scrum) delivering E2E test coverage for enterprise-scale BI environments with multiple stakeholders and parallel initiatives across IKEA and LEGO.",
        "Defined and maintained test strategies, delivery plans, and quality gates for complex platform programmes — including data warehouse, ETL pipeline, and BI reporting validation.",
        "Led test data management for large-scale environments: data masking, pseudonymization, refresh scheduling, and GDPR-compliant handling of personally identifiable information.",
        "Analyzed business processes, identified areas for improvement, and recommended test framework enhancements — supporting platform migration and modernization initiatives.",
        "Managed shared test environments across multiple teams — coordinating provisioning, environment refresh, and governance with infrastructure and operations teams.",
        "Supported year-end testing cycles for financial/billing systems — ensuring data integrity, regression coverage, and readiness for production cutover.",
        "Prepared functional specifications, workflows, and technical documents for development teams — translating stakeholder requirements into testable acceptance criteria.",
        "Led workshops and facilitated discussions with business users to explore testing needs — working across technical and non-technical teams to align priorities.",
        "Tracked project milestones, deliverables, and outcomes — reporting progress to stakeholders through structured dashboards and quality trend analysis.",
        "Oversaw defect management lifecycle (Jira, Zephyr, HP ALM): triage, root-cause analysis, resolution tracking, and stakeholder escalation management.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Banking & Financial Services", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Manager / Technical Lead — Core Banking & BI", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed E2E test cycles for core banking implementations (Finacle CBS) in a regulated financial services environment — defining test strategy, milestones, and quality gates.",
        "Led test data management for banking systems: data masking of customer PII, test data provisioning, and validation of data migration/ETL pipeline integrity.",
        "Coordinated year-end and quarter-end test cycles for billing/financial modules — ensuring data accuracy and regulatory compliance across 20+ branch deployments.",
        "Delivered data-driven reports on test progress, defect trends, and quality metrics to programme leadership — supporting informed Go/No-go decisions.",
        "Collaborated with stakeholders across business, operations, and IT to gather requirements and propose process improvements for test efficiency.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("Test Management: ", "Jira, Zephyr, Azure DevOps, HP ALM (Quality Center), Confluence, TestRail"),
        ("Data & BI: ", "BigQuery, SQL, data pipelines (ETL/ELT), Pub/Sub, data warehouse testing, BI report validation"),
        ("Data Privacy: ", "GDPR compliance, pseudonymization, data masking, PII handling, test data provisioning & refresh"),
        ("Methodologies: ", "SAFe (PI Planning), Scrum, Kanban, Waterfall, ISTQB, risk-based testing"),
        ("Automation & DevOps: ", "Python, Playwright, Selenium, CI/CD quality gates, GitHub Actions, Docker, Terraform"),
        ("Cloud: ", "GCP (BigQuery, Cloud Run, Pub/Sub, IAM), AWS — shared test environment management"),
        ("Reporting: ", "Test metrics dashboards, defect trend analysis, data-driven recommendations, stakeholder reporting"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PG Diploma, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "Six Sigma Green Belt",
        "ITIL v4 Foundation",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Test Manager &mdash; Data Platforms &amp; Enterprise BI</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Test Manager with 10+ years of experience in test lifecycle management across data-intensive and regulated environments (banking, financial services, enterprise BI). Proven track record of managing Agile/SAFe teams, defining E2E test strategies, and optimizing test cycle frequency from annual to continuous delivery. Strong expertise in test data management, GDPR compliance, data masking/pseudonymization, and data pipeline validation. Experienced in leading stakeholder management across multiple departments, gathering business requirements, analyzing as-is processes, and recommending improvements for performance and efficiency. Skilled at defining entry/exit criteria, quality gates, managing test environments, and overseeing defect management and triage. Structured, analytical, and detail-oriented with a proactive focus on risk mitigation and continuous improvement.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Test Strategy &amp; Delivery Planning</td><td>&bull; Test Data Management &amp; GDPR/Masking</td><td>&bull; SAFe Agile Team Leadership (8+ yrs)</td></tr>
<tr><td>&bull; E2E Test Scenarios &amp; Cycle Optimization</td><td>&bull; Stakeholder Management (Multi-Dept)</td><td>&bull; BI &amp; Data Pipeline Validation</td></tr>
<tr><td>&bull; Defect Management &amp; Triage</td><td>&bull; Entry/Exit Criteria &amp; Quality Gates</td><td>&bull; Regulated Industries (Banking/Finance)</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Test Manager / Team Lead &mdash; Data Platform &amp; BI Testing <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Define and maintain the overall test strategy and delivery plan for a data-intensive platform across 30+ markets &mdash; aligned with project roadmap and Agile/SAFe ceremonies.</li>
<li>Lead E2E test scenario collection and execution across data pipelines (BigQuery, Pub/Sub, ETL), APIs, and BI layers &mdash; ensuring comprehensive test coverage for all business processes.</li>
<li>Drive test cycle optimization: improved test frequency from quarterly cycles to continuous delivery by implementing automated regression, parallel execution, and streamlined test data provisioning.</li>
<li>Own test data management and enablement: ensure the latest, validated data is always available to business users in the test layer &mdash; coordinating data refresh, masking, and provisioning.</li>
<li>Ensure GDPR compliance in test environments: implemented pseudonymization and data masking for PI fields across all test data layers, addressing privacy requirements for 30+ market datasets.</li>
<li>Collaborate with stakeholders across multiple departments to gather and document business requirements &mdash; analyzing as-is systems and proposing enhancements for test efficiency.</li>
<li>Define entry and exit criteria, test milestones, and quality gates for each delivery phase &mdash; reporting progress and quality metrics to programme leadership.</li>
<li>Manage test environments in alignment with environment owners &mdash; coordinating provisioning, refresh cycles, and shared environment governance.</li>
<li>Oversee defect management, triage, and resolution tracking &mdash; driving cross-team defect resolution and maintaining traceability to business requirements.</li>
<li>Develop and present data-driven reports and recommendations to business teams &mdash; leveraging test metrics to improve cycle frequency and decision-making quality.</li>
<li>Lead Agile team (SAFe) through PI planning, sprint ceremonies, and backlog refinement &mdash; balancing regular delivery with strategic platform improvement initiatives.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Release &amp; Test Manager <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Managed E2E test cycles for a 300M+ user platform &mdash; defining test strategy, entry/exit criteria, and quality gates for each release phase.</li>
<li>Coordinated test data provisioning across shared environments &mdash; ensuring data freshness, GDPR compliance, and availability for parallel test initiatives.</li>
<li>Led stakeholder management across product, engineering, and operations &mdash; facilitating workshops to explore testing needs and prioritize improvements.</li>
<li>Drove process improvements: reduced test cycle duration by 40% through automation, environment optimization, and parallel test execution strategies.</li>
<li>Oversaw defect management and triage &mdash; tracking resolution across squads and reporting quality status to leadership with data-driven recommendations.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Senior Test Manager &mdash; Enterprise BI &amp; E-Commerce Platforms <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Managed Agile test teams (SAFe/Scrum) delivering E2E test coverage for enterprise-scale BI environments with multiple stakeholders and parallel initiatives across IKEA and LEGO.</li>
<li>Defined and maintained test strategies, delivery plans, and quality gates for complex platform programmes &mdash; including data warehouse, ETL pipeline, and BI reporting validation.</li>
<li>Led test data management for large-scale environments: data masking, pseudonymization, refresh scheduling, and GDPR-compliant handling of personally identifiable information.</li>
<li>Analyzed business processes, identified areas for improvement, and recommended test framework enhancements &mdash; supporting platform migration and modernization initiatives.</li>
<li>Managed shared test environments across multiple teams &mdash; coordinating provisioning, environment refresh, and governance with infrastructure and operations teams.</li>
<li>Supported year-end testing cycles for financial/billing systems &mdash; ensuring data integrity, regression coverage, and readiness for production cutover.</li>
<li>Prepared functional specifications, workflows, and technical documents for development teams &mdash; translating stakeholder requirements into testable acceptance criteria.</li>
<li>Led workshops and facilitated discussions with business users to explore testing needs &mdash; working across technical and non-technical teams to align priorities.</li>
<li>Tracked project milestones, deliverables, and outcomes &mdash; reporting progress to stakeholders through structured dashboards and quality trend analysis.</li>
<li>Oversaw defect management lifecycle (Jira, Zephyr, HP ALM): triage, root-cause analysis, resolution tracking, and stakeholder escalation management.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Banking &amp; Financial Services</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Managed E2E test cycles for core banking implementations (Finacle CBS) in a regulated financial services environment &mdash; defining test strategy, milestones, and quality gates.</li>
<li>Led test data management for banking systems: data masking of customer PII, test data provisioning, and validation of data migration/ETL pipeline integrity.</li>
<li>Coordinated year-end and quarter-end test cycles for billing/financial modules &mdash; ensuring data accuracy and regulatory compliance across 20+ branch deployments.</li>
<li>Delivered data-driven reports on test progress, defect trends, and quality metrics to programme leadership &mdash; supporting informed Go/No-go decisions.</li>
<li>Collaborated with stakeholders across business, operations, and IT to gather requirements and propose process improvements for test efficiency.</li>
</ul>

<h2>TECHNICAL SKILLS</h2>
<p class="tech-line"><b>Test Management:</b> Jira, Zephyr, Azure DevOps, HP ALM (Quality Center), Confluence, TestRail</p>
<p class="tech-line"><b>Data &amp; BI:</b> BigQuery, SQL, data pipelines (ETL/ELT), Pub/Sub, data warehouse testing, BI report validation</p>
<p class="tech-line"><b>Data Privacy:</b> GDPR compliance, pseudonymization, data masking, PII handling, test data provisioning &amp; refresh</p>
<p class="tech-line"><b>Methodologies:</b> SAFe (PI Planning), Scrum, Kanban, Waterfall, ISTQB, risk-based testing</p>
<p class="tech-line"><b>Automation &amp; DevOps:</b> Python, Playwright, Selenium, CI/CD quality gates, GitHub Actions, Docker, Terraform</p>
<p class="tech-line"><b>Cloud:</b> GCP (BigQuery, Cloud Run, Pub/Sub, IAM), AWS &mdash; shared test environment management</p>
<p class="tech-line"><b>Reporting:</b> Test metrics dashboards, defect trend analysis, data-driven recommendations, stakeholder reporting</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; Six Sigma Green Belt</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
