"""
Generate a tailored resume for Senior Test Manager.
Focus: 10+ years test management, pension/financial domain, Azure,
complex IT landscapes, test planning/estimation/coordination/reporting/execution.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Senior_Test_Manager_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Senior_Test_Manager_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "test planning", "test estimation", "test coordination",
    "test reporting", "test execution", "test follow-up",
    "test management", "test manager", "test strategy",
    "pension", "financial", "insurance",
    "Azure", "cloud", "complex IT landscape",
    "transformation", "migration",
    "stakeholder", "communication",
    "risk management", "defect management",
    "regression", "UAT", "SIT", "E2E",
    "CI/CD", "DevOps",
    "agile", "SAFe", "Scrum",
    "Jira", "Confluence", "Azure DevOps",
    "test automation", "quality assurance",
    "cross-functional", "governance",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Senior Test Manager", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Senior Test Manager with 15+ years of experience leading test management activities across "
        "complex IT transformation programs in financial services, pension, and enterprise domains. "
        "Proven expertise in test planning, test estimation, test coordination, test reporting, and "
        "test execution within large-scale Azure-based environments. Experienced in managing testing "
        "across complex IT landscapes involving legacy system migrations, cloud transformations, and "
        "multi-vendor integrations. Strong communicator who bridges technical teams, business stakeholders, "
        "and program leadership — ensuring transparency, risk visibility, and quality governance. "
        "Track record of establishing test strategies, optimizing test processes, and delivering high-quality "
        "releases on time in regulated, mission-critical environments."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Test Planning & Estimation", "Test Coordination & Reporting", "Test Execution & Follow-up",
        "Azure / Cloud Test Environments", "Complex IT Landscape Management", "Pension & Financial Domain",
        "Stakeholder Communication", "Risk & Defect Management", "Transformation Programs",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior Test Manager / Quality Lead — Cloud Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Lead end-to-end test management for a large-scale cloud transformation program — responsible for test planning, test estimation, test coordination, test reporting, and test execution across distributed teams.",
        "Define and maintain comprehensive test strategies for a complex IT landscape spanning Azure cloud services, microservices, legacy integrations, and third-party vendor systems serving 30+ markets.",
        "Perform test estimation and resource planning for releases — coordinating test activities across SIT, E2E, regression, and UAT phases within tight delivery timelines.",
        "Produce regular test reporting and status updates for program leadership and stakeholders — providing transparent visibility into test progress, risks, blockers, and quality metrics.",
        "Drive test follow-up and defect management — tracking defects to resolution, escalating risks, and ensuring no-go criteria are met before production releases.",
        "Manage and coordinate test execution across cross-functional agile teams (Scrum/SAFe) — aligning test cycles with sprint cadence and PI planning ceremonies.",
        "Establish test governance frameworks and quality gates within Azure DevOps — ensuring traceability from requirements through test cases to defect closure.",
        "Communicate effectively with business stakeholders, developers, architects, and program managers — translating technical risks into business impact and driving informed release decisions.",
        "Lead continuous improvement of test processes: introduced risk-based test prioritization, optimized regression suites, and reduced overall test cycle time by 40%.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Manager / Release Quality Lead", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed test planning and test coordination for a 300M+ user platform — orchestrating test execution across multiple teams and cloud environments (AWS/Azure).",
        "Produced test reporting dashboards and release quality metrics for leadership — enabling data-driven go/no-go decisions for production deployments.",
        "Coordinated test estimation and resource allocation across sprint cycles — balancing automation and manual testing efforts for optimal coverage.",
        "Drove defect management and test follow-up — tracking critical issues, coordinating hotfixes, and maintaining release quality standards.",
        "Communicated test status, risks, and blockers to cross-functional stakeholders — ensuring alignment between QA, development, product, and operations teams.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Manager / Test Lead — Enterprise & Digital Platforms", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led test management for large-scale enterprise transformation programs at IKEA and LEGO — managing complex IT landscapes with legacy systems, ERP integrations, e-commerce, and cloud migrations.",
        "Owned comprehensive test planning for program releases — defining test scope, test estimation, entry/exit criteria, environment needs, and resource allocation across multiple workstreams.",
        "Coordinated test execution across geographically distributed teams — managing dependencies, resolving blockers, and ensuring test environments were provisioned and stable.",
        "Delivered regular test reporting to program steering committees and business stakeholders — presenting quality metrics, risk assessments, and release readiness recommendations.",
        "Managed test follow-up and defect governance — tracking defect trends, conducting root-cause analysis, and driving resolution with development teams to meet release criteria.",
        "Established and managed test environments across Azure and on-premise infrastructure — coordinating with infrastructure teams to ensure environment availability for SIT, E2E, and UAT phases.",
        "Managed pension and financial system testing during IKEA's HR/benefits platform transformation — validating complex pension calculation logic, regulatory compliance, and data migration accuracy.",
        "Drove test automation strategy and CI/CD integration — building automated regression suites that fed into Azure DevOps pipelines, reducing manual regression effort by 60%.",
        "Collaborated closely with architects, business analysts, and Product Owners — translating business requirements into test strategies and ensuring acceptance criteria were testable and complete.",
        "Mentored and developed a team of 8–12 test engineers — providing guidance on test techniques, domain knowledge, and career growth within the testing discipline.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Financial & Enterprise Domain Testing", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Lead / Senior Test Engineer", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed test planning, test estimation, and test execution for core banking (Finacle CBS) and pension/insurance system implementations — coordinating testing across complex multi-tier IT architectures.",
        "Led test coordination across onshore/offshore teams — managing test environments, data preparation, and execution schedules for large financial transformation programs.",
        "Produced test reporting and quality dashboards for project stakeholders — tracking defect metrics, SLA compliance, and release readiness in regulated financial environments.",
        "Validated pension fund calculations, regulatory reporting, and data migration accuracy — ensuring compliance with financial regulations and business rules.",
        "Drove defect management and test follow-up processes — establishing defect triage, escalation paths, and resolution tracking that improved defect closure rates by 35%.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS & TOOLS")
    for label, value in [
        ("Cloud & Platforms: ", "Azure (DevOps, Pipelines, Test Plans, App Services), AWS, GCP, hybrid/on-premise environments"),
        ("Test Management: ", "Azure DevOps, Jira, Confluence, Zephyr, TestRail, HP ALM/Quality Center, qTest"),
        ("Methodologies: ", "Agile (Scrum/Kanban), SAFe, Waterfall, V-Model, ISTQB, risk-based testing, exploratory testing"),
        ("Test Types: ", "SIT, E2E, UAT, regression, performance, data migration, integration, API testing"),
        ("Automation & CI/CD: ", "Python (pytest), Selenium, Playwright, Jenkins, GitHub Actions, Azure Pipelines"),
        ("Domain Expertise: ", "Pension systems, financial services, insurance, retail/e-commerce, HR/benefits platforms"),
        ("Reporting: ", "Test dashboards, KPI/quality metrics, defect trend analysis, release readiness assessments"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PGDOM, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "ITIL v4 Foundation",
        "Certified Ethical Hacker (CEH)",
        "Six Sigma Green Belt",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Senior Test Manager</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Senior Test Manager with 15+ years of experience leading test management activities across complex IT transformation programs in financial services, pension, and enterprise domains. Proven expertise in test planning, test estimation, test coordination, test reporting, and test execution within large-scale Azure-based environments. Experienced in managing testing across complex IT landscapes involving legacy system migrations, cloud transformations, and multi-vendor integrations. Strong communicator who bridges technical teams, business stakeholders, and program leadership &mdash; ensuring transparency, risk visibility, and quality governance. Track record of establishing test strategies, optimizing test processes, and delivering high-quality releases on time in regulated, mission-critical environments.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Test Planning &amp; Estimation</td><td>&bull; Test Coordination &amp; Reporting</td><td>&bull; Test Execution &amp; Follow-up</td></tr>
<tr><td>&bull; Azure / Cloud Test Environments</td><td>&bull; Complex IT Landscape Management</td><td>&bull; Pension &amp; Financial Domain</td></tr>
<tr><td>&bull; Stakeholder Communication</td><td>&bull; Risk &amp; Defect Management</td><td>&bull; Transformation Programs</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Senior Test Manager / Quality Lead &mdash; Cloud Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Lead end-to-end test management for a large-scale cloud transformation program &mdash; responsible for test planning, test estimation, test coordination, test reporting, and test execution across distributed teams.</li>
<li>Define and maintain comprehensive test strategies for a complex IT landscape spanning Azure cloud services, microservices, legacy integrations, and third-party vendor systems serving 30+ markets.</li>
<li>Perform test estimation and resource planning for releases &mdash; coordinating test activities across SIT, E2E, regression, and UAT phases within tight delivery timelines.</li>
<li>Produce regular test reporting and status updates for program leadership and stakeholders &mdash; providing transparent visibility into test progress, risks, blockers, and quality metrics.</li>
<li>Drive test follow-up and defect management &mdash; tracking defects to resolution, escalating risks, and ensuring no-go criteria are met before production releases.</li>
<li>Manage and coordinate test execution across cross-functional agile teams (Scrum/SAFe) &mdash; aligning test cycles with sprint cadence and PI planning ceremonies.</li>
<li>Establish test governance frameworks and quality gates within Azure DevOps &mdash; ensuring traceability from requirements through test cases to defect closure.</li>
<li>Communicate effectively with business stakeholders, developers, architects, and program managers &mdash; translating technical risks into business impact and driving informed release decisions.</li>
<li>Lead continuous improvement of test processes: introduced risk-based test prioritization, optimized regression suites, and reduced overall test cycle time by 40%.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Test Manager / Release Quality Lead <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Managed test planning and test coordination for a 300M+ user platform &mdash; orchestrating test execution across multiple teams and cloud environments (AWS/Azure).</li>
<li>Produced test reporting dashboards and release quality metrics for leadership &mdash; enabling data-driven go/no-go decisions for production deployments.</li>
<li>Coordinated test estimation and resource allocation across sprint cycles &mdash; balancing automation and manual testing efforts for optimal coverage.</li>
<li>Drove defect management and test follow-up &mdash; tracking critical issues, coordinating hotfixes, and maintaining release quality standards.</li>
<li>Communicated test status, risks, and blockers to cross-functional stakeholders &mdash; ensuring alignment between QA, development, product, and operations teams.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Test Manager / Test Lead &mdash; Enterprise &amp; Digital Platforms <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Led test management for large-scale enterprise transformation programs at IKEA and LEGO &mdash; managing complex IT landscapes with legacy systems, ERP integrations, e-commerce, and cloud migrations.</li>
<li>Owned comprehensive test planning for program releases &mdash; defining test scope, test estimation, entry/exit criteria, environment needs, and resource allocation across multiple workstreams.</li>
<li>Coordinated test execution across geographically distributed teams &mdash; managing dependencies, resolving blockers, and ensuring test environments were provisioned and stable.</li>
<li>Delivered regular test reporting to program steering committees and business stakeholders &mdash; presenting quality metrics, risk assessments, and release readiness recommendations.</li>
<li>Managed test follow-up and defect governance &mdash; tracking defect trends, conducting root-cause analysis, and driving resolution with development teams to meet release criteria.</li>
<li>Established and managed test environments across Azure and on-premise infrastructure &mdash; coordinating with infrastructure teams to ensure environment availability for SIT, E2E, and UAT phases.</li>
<li>Managed pension and financial system testing during IKEA's HR/benefits platform transformation &mdash; validating complex pension calculation logic, regulatory compliance, and data migration accuracy.</li>
<li>Drove test automation strategy and CI/CD integration &mdash; building automated regression suites that fed into Azure DevOps pipelines, reducing manual regression effort by 60%.</li>
<li>Collaborated closely with architects, business analysts, and Product Owners &mdash; translating business requirements into test strategies and ensuring acceptance criteria were testable and complete.</li>
<li>Mentored and developed a team of 8&ndash;12 test engineers &mdash; providing guidance on test techniques, domain knowledge, and career growth within the testing discipline.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Financial &amp; Enterprise Domain Testing</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Managed test planning, test estimation, and test execution for core banking (Finacle CBS) and pension/insurance system implementations &mdash; coordinating testing across complex multi-tier IT architectures.</li>
<li>Led test coordination across onshore/offshore teams &mdash; managing test environments, data preparation, and execution schedules for large financial transformation programs.</li>
<li>Produced test reporting and quality dashboards for project stakeholders &mdash; tracking defect metrics, SLA compliance, and release readiness in regulated financial environments.</li>
<li>Validated pension fund calculations, regulatory reporting, and data migration accuracy &mdash; ensuring compliance with financial regulations and business rules.</li>
<li>Drove defect management and test follow-up processes &mdash; establishing defect triage, escalation paths, and resolution tracking that improved defect closure rates by 35%.</li>
</ul>

<h2>TECHNICAL SKILLS &amp; TOOLS</h2>
<p class="tech-line"><b>Cloud &amp; Platforms:</b> Azure (DevOps, Pipelines, Test Plans, App Services), AWS, GCP, hybrid/on-premise environments</p>
<p class="tech-line"><b>Test Management:</b> Azure DevOps, Jira, Confluence, Zephyr, TestRail, HP ALM/Quality Center, qTest</p>
<p class="tech-line"><b>Methodologies:</b> Agile (Scrum/Kanban), SAFe, Waterfall, V-Model, ISTQB, risk-based testing, exploratory testing</p>
<p class="tech-line"><b>Test Types:</b> SIT, E2E, UAT, regression, performance, data migration, integration, API testing</p>
<p class="tech-line"><b>Automation &amp; CI/CD:</b> Python (pytest), Selenium, Playwright, Jenkins, GitHub Actions, Azure Pipelines</p>
<p class="tech-line"><b>Domain Expertise:</b> Pension systems, financial services, insurance, retail/e-commerce, HR/benefits platforms</p>
<p class="tech-line"><b>Reporting:</b> Test dashboards, KPI/quality metrics, defect trend analysis, release readiness assessments</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PGDOM, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; Certified Ethical Hacker (CEH)</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
