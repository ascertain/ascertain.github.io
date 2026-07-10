"""
Resume: Senior Quality Assurance Engineer – Emirates Group (EXTERNAL, Dubai)
Focus: Quality engineering, automation frameworks (UI/API), CI/CD, DevSecOps,
       cross-team coordination, testing strategy, cloud/containers, exploratory testing.
EXTERNAL = show diverse company experience.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Senior_QA_Engineer_Emirates_Resume"
BRAND = RGBColor(0xC8, 0x00, 0x2E)  # Emirates red


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="C8002E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Senior Quality Assurance Engineer with 14+ years of experience in quality engineering across digital, "
        "e-commerce, and enterprise platforms. Extensive hands-on expertise in building UI and API automation "
        "frameworks (Selenium, Cypress, Postman, Python, JavaScript), implementing CI/CD pipelines with DevSecOps "
        "practices, and driving quality gates across multi-team release trains. Proven track record leading QA "
        "for complex customer journey platforms — from booking flows to backend service integrations. Strong in "
        "context-driven and exploratory testing, cross-dependency management, test environment governance, and "
        "guiding engineering teams to adopt automation-first practices. Experienced with cloud platforms (GCP, AWS), "
        "containers (Docker, Kubernetes), and modern DevOps toolchains."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "UI & API Automation Frameworks", "CI/CD & DevSecOps Pipelines", "Quality Gate Governance",
        "Context-Driven & Exploratory Testing", "Cross-Team Release Coordination", "Test Environment & Data Mgmt",
        "Selenium / Cypress / Playwright", "Python, JavaScript, Java (testing)", "Cloud & Container Testing (GCP/AWS)",
        "Metrics & Dashboard Reporting", "E-Commerce / Digital Platforms", "Coaching Engineers on Test Automation",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TOOLS & TECHNOLOGIES ──
    add_heading_block(doc, "Tools & Technologies")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    skills = (
        "Automation: Selenium (WebDriver, Grid), Cypress, Playwright, Cucumber/Gherkin, JUnit, TestNG, Pytest, Vitest  •  "
        "Languages: Python, JavaScript/TypeScript, Java (test automation), SQL  •  "
        "CI/CD: Jenkins, GitHub Actions, Azure DevOps pipelines  •  "
        "SCM: Git, GitHub, Bitbucket  •  "
        "API Testing: Postman, REST Assured, k6, JMeter  •  "
        "Cloud & Containers: GCP (GKE, Cloud Run), AWS (EC2, Lambda), Docker, Kubernetes  •  "
        "Monitoring: Grafana, OpenTelemetry, test metrics dashboards  •  "
        "Collaboration: Jira, Confluence, TestRail, HP ALM"
    )
    r = tp.add_run(skills)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- Ingka Digital / IKEA ---
    role_header(doc, "Senior QA Engineer & Team Lead – Digital Customer Platforms",
                "Ingka Digital (IKEA)", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Lead quality engineering across 5+ integrated digital platforms supporting the end-to-end customer journey — from service booking to delivery and post-purchase interactions.")
    bullet(doc, "Design, implement, and maintain UI and API automation frameworks (Cypress, Vitest, Postman) achieving 80%+ automated regression coverage across microservices and frontend applications.")
    bullet(doc, "Implement and govern CI/CD quality gates: automated test suites execute on every pull request via GitHub Actions; enforce code coverage thresholds and security scanning (DevSecOps).")
    bullet(doc, "Work closely with architects and software engineers to understand the technical landscape; translate complex functional and non-functional requirements into comprehensive test strategies.")
    bullet(doc, "Provide early feedback to product owners on risks in requirements and design; apply context-driven and exploratory testing heuristics to uncover edge cases before development.")
    bullet(doc, "Manage non-production environments and test data provisioning; ensure environments are stable and available to support testing across multiple feature teams.")
    bullet(doc, "Liaise with other release train teams to coordinate testing of cross-dependencies; manage integration test schedules and shared environment bookings.")
    bullet(doc, "Configure test tools and dashboards to collect metrics for programme-wide visibility: pass/fail rates, defect trends, automation coverage, and quality gate compliance.")
    bullet(doc, "Support and guide software engineers to write and maintain automated functional tests; conduct pairing sessions and define coding guidelines for test code quality.")
    bullet(doc, "Deploy and test services on cloud-native infrastructure (GCP Cloud Run, GKE/Kubernetes, Docker containers); validate deployments in containerised environments.")

    # --- Truecaller ---
    role_header(doc, "QA Lead – Platform Quality Engineering",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Built API automation frameworks for microservices-based SaaS platform; implemented contract testing and integration test suites across distributed services.")
    bullet(doc, "Established quality gates in CI/CD pipelines (Jenkins); ensured automated tests ran on every deployment with clear pass/fail reporting.")
    bullet(doc, "Applied exploratory testing techniques on complex user flows; identified critical defects in production-like environments before release.")
    bullet(doc, "Worked with VMs and containers on cloud (GCP); validated service behaviour in containerised deployments.")

    # --- HCLTech ---
    role_header(doc, "Senior QA Engineer & Test Lead – Enterprise Digital Platforms",
                "HCLTech (for IKEA & LEGO)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Led quality engineering for large-scale e-commerce and digital platforms spanning 10+ integrated systems — covering the full customer journey from product discovery to order fulfilment.")
    bullet(doc, "Developed and maintained end-to-end automation frameworks using Selenium (WebDriver, Grid), Cucumber/Gherkin, JUnit, and Python — achieving significant regression cycle time reduction.")
    bullet(doc, "Implemented CI/CD integration for test suites (Jenkins, GitHub Actions); automated smoke, regression, and integration tests as quality gates in deployment pipelines.")
    bullet(doc, "Managed test environments and test data across on-prem and cloud infrastructure; coordinated environment availability for multiple feature teams and release trains.")
    bullet(doc, "Guided software engineers in adopting test automation: defined coding standards for test code, conducted workshops on BDD (Cucumber/Gherkin), and reviewed automation PRs.")
    bullet(doc, "Applied context-driven testing and exploratory techniques using testing heuristics (SFDPOT, HICCUPPS) to uncover risks in complex business process flows.")
    bullet(doc, "Led cross-team testing of dependencies across release trains; managed integration test schedules and defect triage for cross-system issues.")
    bullet(doc, "Configured and managed test tools (HP ALM, Jira, TestRail) to support metrics collection for programme dashboards: defect density, test coverage, automation ROI.")
    bullet(doc, "Provided input on business process flows and scenario mapping; partnered with product owners to define acceptance criteria and validate business solutions.")

    # --- India ---
    role_header(doc, "Software QA Engineer",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Developed automated test frameworks for web applications using Selenium, Java, and Python; progressed from manual testing to full automation engineering.")
    bullet(doc, "Built foundation in testing methodologies (functional, regression, integration, performance) and SCM tools (Git, SVN); worked in Agile and Waterfall environments.")
    bullet(doc, "Gained experience across diverse technology stacks (Java, .NET, Oracle, web services) in digital and enterprise application testing.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "ISTQB Certified Tester – Foundation Level",
        "Google Cloud Associate Cloud Engineer",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
        "Six Sigma Green Belt",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – UP Technical University  |  ").font.size = Pt(9)
    r2 = p.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – IGNOU").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)  |  Arabic (Basic exposure)").font.size = Pt(9)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#C8002E;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#C8002E;font-size:11pt;border-bottom:1px solid #C8002E;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#C8002E;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Senior Quality Assurance Engineer with 14+ years of experience in quality engineering across digital,
e-commerce, and enterprise platforms. Extensive hands-on expertise in building UI and API automation
frameworks (Selenium, Cypress, Postman, Python, JavaScript), implementing CI/CD pipelines with DevSecOps
practices, and driving quality gates across multi-team release trains. Proven track record leading QA
for complex customer journey platforms — from booking flows to backend service integrations. Strong in
context-driven and exploratory testing, cross-dependency management, test environment governance, and
guiding engineering teams to adopt automation-first practices. Experienced with cloud platforms (GCP, AWS),
containers (Docker, Kubernetes), and modern DevOps toolchains.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ UI &amp; API Automation Frameworks</td><td>✓ CI/CD &amp; DevSecOps Pipelines</td><td>✓ Quality Gate Governance</td></tr>
<tr><td>✓ Context-Driven &amp; Exploratory Testing</td><td>✓ Cross-Team Release Coordination</td><td>✓ Test Environment &amp; Data Mgmt</td></tr>
<tr><td>✓ Selenium / Cypress / Playwright</td><td>✓ Python, JavaScript, Java (testing)</td><td>✓ Cloud &amp; Container Testing (GCP/AWS)</td></tr>
<tr><td>✓ Metrics &amp; Dashboard Reporting</td><td>✓ E-Commerce / Digital Platforms</td><td>✓ Coaching Engineers on Test Automation</td></tr>
</table>

<h2>TOOLS &amp; TECHNOLOGIES</h2>
<p class="tools">Automation: Selenium (WebDriver, Grid), Cypress, Playwright, Cucumber/Gherkin, JUnit, TestNG, Pytest, Vitest &bull;
Languages: Python, JavaScript/TypeScript, Java (test automation), SQL &bull;
CI/CD: Jenkins, GitHub Actions, Azure DevOps pipelines &bull;
SCM: Git, GitHub, Bitbucket &bull;
API Testing: Postman, REST Assured, k6, JMeter &bull;
Cloud &amp; Containers: GCP (GKE, Cloud Run), AWS (EC2, Lambda), Docker, Kubernetes &bull;
Monitoring: Grafana, OpenTelemetry, test metrics dashboards &bull;
Collaboration: Jira, Confluence, TestRail, HP ALM</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Senior QA Engineer &amp; Team Lead – Digital Customer Platforms &nbsp;|&nbsp; Ingka Digital (IKEA), Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Lead quality engineering across 5+ integrated digital platforms supporting the end-to-end customer journey — from service booking to delivery and post-purchase interactions.</li>
<li>Design, implement, and maintain UI and API automation frameworks (Cypress, Vitest, Postman) achieving 80%+ automated regression coverage across microservices and frontend applications.</li>
<li>Implement and govern CI/CD quality gates: automated test suites execute on every pull request via GitHub Actions; enforce code coverage thresholds and security scanning (DevSecOps).</li>
<li>Work closely with architects and software engineers to understand the technical landscape; translate complex functional and non-functional requirements into comprehensive test strategies.</li>
<li>Provide early feedback to product owners on risks in requirements and design; apply context-driven and exploratory testing heuristics to uncover edge cases before development.</li>
<li>Manage non-production environments and test data provisioning; ensure environments are stable and available to support testing across multiple feature teams.</li>
<li>Liaise with other release train teams to coordinate testing of cross-dependencies; manage integration test schedules and shared environment bookings.</li>
<li>Configure test tools and dashboards to collect metrics for programme-wide visibility: pass/fail rates, defect trends, automation coverage, and quality gate compliance.</li>
<li>Support and guide software engineers to write and maintain automated functional tests; conduct pairing sessions and define coding guidelines for test code quality.</li>
<li>Deploy and test services on cloud-native infrastructure (GCP Cloud Run, GKE/Kubernetes, Docker containers); validate deployments in containerised environments.</li>
</ul>

<p class="role">QA Lead – Platform Quality Engineering &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Built API automation frameworks for microservices-based SaaS platform; implemented contract testing and integration test suites across distributed services.</li>
<li>Established quality gates in CI/CD pipelines (Jenkins); ensured automated tests ran on every deployment with clear pass/fail reporting.</li>
<li>Applied exploratory testing techniques on complex user flows; identified critical defects in production-like environments before release.</li>
<li>Worked with VMs and containers on cloud (GCP); validated service behaviour in containerised deployments.</li>
</ul>

<p class="role">Senior QA Engineer &amp; Test Lead – Enterprise Digital Platforms &nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Led quality engineering for large-scale e-commerce and digital platforms spanning 10+ integrated systems — covering the full customer journey from product discovery to order fulfilment.</li>
<li>Developed and maintained end-to-end automation frameworks using Selenium (WebDriver, Grid), Cucumber/Gherkin, JUnit, and Python — achieving significant regression cycle time reduction.</li>
<li>Implemented CI/CD integration for test suites (Jenkins, GitHub Actions); automated smoke, regression, and integration tests as quality gates in deployment pipelines.</li>
<li>Managed test environments and test data across on-prem and cloud infrastructure; coordinated environment availability for multiple feature teams and release trains.</li>
<li>Guided software engineers in adopting test automation: defined coding standards for test code, conducted workshops on BDD (Cucumber/Gherkin), and reviewed automation PRs.</li>
<li>Applied context-driven testing and exploratory techniques using testing heuristics (SFDPOT, HICCUPPS) to uncover risks in complex business process flows.</li>
<li>Led cross-team testing of dependencies across release trains; managed integration test schedules and defect triage for cross-system issues.</li>
<li>Configured and managed test tools (HP ALM, Jira, TestRail) to support metrics collection for programme dashboards: defect density, test coverage, automation ROI.</li>
<li>Provided input on business process flows and scenario mapping; partnered with product owners to define acceptance criteria and validate business solutions.</li>
</ul>

<p class="role">Software QA Engineer &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Developed automated test frameworks for web applications using Selenium, Java, and Python; progressed from manual testing to full automation engineering.</li>
<li>Built foundation in testing methodologies (functional, regression, integration, performance) and SCM tools (Git, SVN); worked in Agile and Waterfall environments.</li>
<li>Gained experience across diverse technology stacks (Java, .NET, Oracle, web services) in digital and enterprise application testing.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>ISTQB Certified Tester – Foundation Level</li>
<li>Google Cloud Associate Cloud Engineer</li>
<li>AWS Certified Cloud Practitioner</li>
<li>Certified Ethical Hacker (CEH)</li>
<li>Six Sigma Green Belt</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>B.Tech Information Technology</b> – UP Technical University &nbsp;|&nbsp; <b>PGDOM</b> – IGNOU</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Hindi/Urdu (Native) &nbsp;|&nbsp; Swedish (Basic) &nbsp;|&nbsp; Arabic (Basic exposure)</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
