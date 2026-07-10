"""
Resume: Test Manager – Capgemini UK (EXTERNAL)
Focus: Test Strategy (SIT/UAT/E2E), offshore team management, Agile ceremonies,
       stakeholder management, risk identification, reporting, multi-system integration testing.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import pathlib, html

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Test_Manager_Capgemini_Resume"
BLUE = RGBColor(0x00, 0x3D, 0x6B)  # Capgemini-ish deep blue


# ─── helpers ────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shading)


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    # underline via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="003D6B"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX builder ──────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact = "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif"
    r2 = p2.add_run(contact)
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Results-driven Test Manager with 14+ years of experience defining and executing test strategies "
        "across system integration, user acceptance, and end-to-end testing for large-scale, multi-system programmes. "
        "Proven ability to manage distributed onshore/offshore test teams, lead Agile ceremonies (Scrum of Scrums, "
        "coordination meetings), and deliver concise stakeholder reporting. Skilled at building assurance frameworks "
        "that give business stakeholders full confidence in solution quality through iterative delivery. "
        "Strong track record in risk identification, dependency management, environments coordination, and "
        "cross-functional stakeholder engagement from technical teams to executive sponsors."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Test Strategy (SIT / UAT / E2E)", "Offshore & Distributed Team Mgmt", "Agile Ceremonies & Scrum of Scrums",
        "Stakeholder & Relationship Mgmt", "Risk Identification & Escalation", "Environments & Dependency Mgmt",
        "Test Analyst Management", "Defect & Conflict Resolution", "Project-Level Test Reporting",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(5.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- IKEA / Ingka Digital ---
    role_header(doc, "Test Manager / Team Lead – Visual Customer Support",
                "Ingka Digital (IKEA)", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Defined and refined end-to-end test strategies spanning 5+ integrated systems (VCS, CSSP, Genesys, Verint, Chatbot) ensuring SIT, UAT, and regression cycles deliver measurable quality gates.")
    bullet(doc, "Managed a distributed test team (onshore + offshore resources) with responsibility for test planning, execution tracking, defect triage, and daily coordination across product groups.")
    bullet(doc, "Led Scrum of Scrums and cross-team coordination ceremonies; managed dependencies and environment bookings across multiple delivery trains.")
    bullet(doc, "Built an iterative assurance framework that provided business stakeholders with clear confidence in solution quality at each release increment.")
    bullet(doc, "Owned project-level test reporting: status dashboards, risk heat-maps, and executive-level summaries enabling data-driven Go/No-Go decisions.")
    bullet(doc, "Identified and escalated risks early; maintained RAID logs and ensured mitigations were documented and actioned.")
    bullet(doc, "Managed UAT coordination with business SMEs, facilitating test scenario workshops and sign-off processes.")
    bullet(doc, "Worked directly with Product Group teams to ensure test approaches were fit for purpose and clearly documented.")

    # --- Truecaller ---
    role_header(doc, "QA Lead – Platform Testing",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Developed integration test strategies for microservices-based architecture spanning multiple technology stacks.")
    bullet(doc, "Coordinated E2E test execution across backend, API, and mobile teams; resolved cross-team defects and conflicts.")
    bullet(doc, "Introduced structured test reporting cadence providing concise visibility to product and engineering leadership.")

    # --- HCLTech (IKEA & LEGO) ---
    role_header(doc, "Test Manager – Enterprise Programmes",
                "HCLTech (for IKEA & LEGO)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Led test strategy definition for large-scale system integration programmes across 10+ interconnected systems (e-commerce, supply chain, order management, CRM platforms).")
    bullet(doc, "Managed a team of 8–12 Test Analysts (onshore & offshore); responsible for hiring, coaching, workload allocation, and performance reviews.")
    bullet(doc, "Drove end-to-end testing across integrated systems including Salesforce-adjacent CRM platforms, SAP, and custom web applications.")
    bullet(doc, "Established UAT governance: test scenario design with business SMEs, acceptance criteria workshops, and structured sign-off protocols.")
    bullet(doc, "Owned environments management and dependency tracking across multiple technology teams; ensured test environments were provisioned and stable.")
    bullet(doc, "Facilitated Agile ceremonies (Scrum of Scrums, retrospectives, sprint demos) and ensured testing was embedded in every sprint.")
    bullet(doc, "Delivered weekly/fortnightly project-level test reports to programme boards highlighting progress, blockers, risks, and mitigation actions.")
    bullet(doc, "Resolved cross-team defects and conflicting priorities in conjunction with Product Group test teams and development leads.")
    bullet(doc, "Collaborated closely with offshore delivery centres (India) on test execution, shift-left automation, and knowledge transfer.")

    # --- India roles ---
    role_header(doc, "Senior Test Engineer / Test Lead",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Progressed from Test Engineer to Test Lead; managed small test teams and owned test planning for web and enterprise applications.")
    bullet(doc, "Gained foundational experience in manual testing, test case design, defect lifecycle management, and stakeholder communication.")
    bullet(doc, "Worked across diverse technologies (Java, .NET, Oracle, web services) building a strong multi-platform testing skill set.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "ISTQB Certified Tester – Foundation Level",
        "AWS Certified Cloud Practitioner",
        "Google Cloud Associate Cloud Engineer",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("PGDOM")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – IGNOU  |  ").font.size = Pt(9)
    r2 = p.add_run("B.Tech Information Technology")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – UP Technical University").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Basic)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    # ── WHY CAPGEMINI ──
    add_heading_block(doc, "Why Capgemini")
    why = (
        "I am drawn to Capgemini's commitment to ethical business practices (13 consecutive years as World's Most Ethical Companies®) "
        "and the emphasis on continuous learning and innovation. I thrive in environments where technology is used to reimagine what's possible "
        "and am excited to contribute my test management expertise to programmes that grow clients' businesses while building a more inclusive future."
    )
    pw = doc.add_paragraph(why)
    pw.paragraph_format.space_before = Pt(3)
    for run in pw.runs:
        run.font.size = Pt(9)
        run.italic = True

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) builder ─────────────────────────────────────────────────────
def build_doc():
    h = html.escape
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#003D6B;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#003D6B;font-size:11pt;border-bottom:1px solid #003D6B;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#003D6B;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .certs li,.edu{font-size:9pt}
    .why{font-style:italic;font-size:9pt;margin-top:4px}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Results-driven Test Manager with 14+ years of experience defining and executing test strategies
across system integration, user acceptance, and end-to-end testing for large-scale, multi-system programmes.
Proven ability to manage distributed onshore/offshore test teams, lead Agile ceremonies (Scrum of Scrums,
coordination meetings), and deliver concise stakeholder reporting. Skilled at building assurance frameworks
that give business stakeholders full confidence in solution quality through iterative delivery.
Strong track record in risk identification, dependency management, environments coordination, and
cross-functional stakeholder engagement from technical teams to executive sponsors.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Test Strategy (SIT / UAT / E2E)</td><td>✓ Offshore &amp; Distributed Team Mgmt</td><td>✓ Agile Ceremonies &amp; Scrum of Scrums</td></tr>
<tr><td>✓ Stakeholder &amp; Relationship Mgmt</td><td>✓ Risk Identification &amp; Escalation</td><td>✓ Environments &amp; Dependency Mgmt</td></tr>
<tr><td>✓ Test Analyst Management</td><td>✓ Defect &amp; Conflict Resolution</td><td>✓ Project-Level Test Reporting</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Test Manager / Team Lead – Visual Customer Support &nbsp;|&nbsp; Ingka Digital (IKEA), Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Defined and refined end-to-end test strategies spanning 5+ integrated systems (VCS, CSSP, Genesys, Verint, Chatbot) ensuring SIT, UAT, and regression cycles deliver measurable quality gates.</li>
<li>Managed a distributed test team (onshore + offshore resources) with responsibility for test planning, execution tracking, defect triage, and daily coordination across product groups.</li>
<li>Led Scrum of Scrums and cross-team coordination ceremonies; managed dependencies and environment bookings across multiple delivery trains.</li>
<li>Built an iterative assurance framework that provided business stakeholders with clear confidence in solution quality at each release increment.</li>
<li>Owned project-level test reporting: status dashboards, risk heat-maps, and executive-level summaries enabling data-driven Go/No-Go decisions.</li>
<li>Identified and escalated risks early; maintained RAID logs and ensured mitigations were documented and actioned.</li>
<li>Managed UAT coordination with business SMEs, facilitating test scenario workshops and sign-off processes.</li>
<li>Worked directly with Product Group teams to ensure test approaches were fit for purpose and clearly documented.</li>
</ul>

<p class="role">QA Lead – Platform Testing &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Developed integration test strategies for microservices-based architecture spanning multiple technology stacks.</li>
<li>Coordinated E2E test execution across backend, API, and mobile teams; resolved cross-team defects and conflicts.</li>
<li>Introduced structured test reporting cadence providing concise visibility to product and engineering leadership.</li>
</ul>

<p class="role">Test Manager – Enterprise Programmes &nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Led test strategy definition for large-scale system integration programmes across 10+ interconnected systems (e-commerce, supply chain, order management, CRM platforms).</li>
<li>Managed a team of 8–12 Test Analysts (onshore &amp; offshore); responsible for hiring, coaching, workload allocation, and performance reviews.</li>
<li>Drove end-to-end testing across integrated systems including Salesforce-adjacent CRM platforms, SAP, and custom web applications.</li>
<li>Established UAT governance: test scenario design with business SMEs, acceptance criteria workshops, and structured sign-off protocols.</li>
<li>Owned environments management and dependency tracking across multiple technology teams; ensured test environments were provisioned and stable.</li>
<li>Facilitated Agile ceremonies (Scrum of Scrums, retrospectives, sprint demos) and ensured testing was embedded in every sprint.</li>
<li>Delivered weekly/fortnightly project-level test reports to programme boards highlighting progress, blockers, risks, and mitigation actions.</li>
<li>Resolved cross-team defects and conflicting priorities in conjunction with Product Group test teams and development leads.</li>
<li>Collaborated closely with offshore delivery centres (India) on test execution, shift-left automation, and knowledge transfer.</li>
</ul>

<p class="role">Senior Test Engineer / Test Lead &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Progressed from Test Engineer to Test Lead; managed small test teams and owned test planning for web and enterprise applications.</li>
<li>Gained foundational experience in manual testing, test case design, defect lifecycle management, and stakeholder communication.</li>
<li>Worked across diverse technologies (Java, .NET, Oracle, web services) building a strong multi-platform testing skill set.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>ISTQB Certified Tester – Foundation Level</li>
<li>AWS Certified Cloud Practitioner</li>
<li>Google Cloud Associate Cloud Engineer</li>
<li>Six Sigma Green Belt</li>
<li>Certified Ethical Hacker (CEH)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>PGDOM</b> – IGNOU &nbsp;|&nbsp; <b>B.Tech Information Technology</b> – UP Technical University</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Basic) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

<h2>WHY CAPGEMINI</h2>
<p class="why">I am drawn to Capgemini's commitment to ethical business practices (13 consecutive years as World's Most Ethical Companies®)
and the emphasis on continuous learning and innovation. I thrive in environments where technology is used to reimagine what's possible
and am excited to contribute my test management expertise to programmes that grow clients' businesses while building a more inclusive future.</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


# ─── Main ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
