"""
Resume: Head of Quality Engineering / QA Manager — Studio/Product Environment
Focus: Proactive quality systems, observability, cross-discipline quality ownership, organisational change.
Malmö/Stockholm. DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Head_of_Quality_Engineering_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Quality Engineering leader with ", False),
        ("16+ years", True),
        (" evolving how organisations ", False),
        ("build confidence in what they ship", True),
        (". Track record of moving quality systems from ", False),
        ("reactive to proactive", True),
        (", improving ", False),
        ("production observability", True),
        (", and making ", False),
        ("quality a shared capability", True),
        (" across Engineering, Design, and Operations. ", False),
        ("Systems thinker", True),
        (" who identifies ", False),
        ("systemic bottlenecks", True),
        (" and ", False),
        ("recurring failure patterns", True),
        (" \u2014 then drives meaningful ", False),
        ("organisational change", True),
        (" to address them. Values ", False),
        ("learning over blame", True),
        (", ", False),
        ("observability over process for process\u2019s sake", True),
        (", and builds quality systems that ", False),
        ("work for the humans inside them", True),
        (". People manager who fosters ", False),
        ("psychological safety", True),
        (" and ", False),
        ("continuous learning", True),
        (".", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Proactive Quality Systems", "Production Observability", "Cross-Discipline Quality Ownership",
        "QA Tooling & Modernisation", "Metrics, Telemetry & Reporting", "Systemic Problem Identification",
        "Organisational Change Leadership", "Psychological Safety & Inclusion", "People Management & Hiring",
        "Continuous Improvement", "Release Readiness Assessment", "Embedded + Operational QA Collaboration",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL & QUALITY TOOLKIT ──
    add_section_heading(doc, "Technical & Quality Toolkit")
    skills_data = [
        ("Quality Systems:", " Test strategy design, verification workflows, shift-left practices, risk-based testing"),
        ("Observability:", " Grafana, quality dashboards, defect telemetry, release readiness metrics, SLOs/SLIs"),
        ("Automation & Tooling:", " Python, Playwright, Selenium, Pytest, Robot Framework, custom test infrastructure"),
        ("CI/CD & DevOps:", " GitHub Actions, Jenkins, Docker, Kubernetes, GitOps, automated quality gates"),
        ("Test Management:", " Jira, XRAY, TestRail, Confluence, defect tracking & trend analysis"),
        ("Platforms:", " GCP (Cloud Run, GKE, Pub/Sub, BigQuery), AWS, Terraform, microservices"),
        ("Practices:", " Agile/Scrum, Kanban, blameless retrospectives, learning reviews, OKRs"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Quality Engineering Lead / Team Lead (Acting)",
        "Mar 2022 \u2013 Present",
        "Omni-Channel Platform \u2014 Engineering, Design & LiveOps | 30+ Markets | Cloud-Native SaaS")
    add_bullet(doc, "Drive continuous improvement across QA practices, workflows, and collaboration models \u2014 evolving quality from reactive firefighting to proactive, systemic confidence-building.",
        ["continuous improvement", "QA practices, workflows", "reactive", "proactive", "systemic confidence-building"])
    add_bullet(doc, "Champion quality ownership across disciplines (Engineering, Design, LiveOps) \u2014 reducing dependency on late discovery and shifting verification left.",
        ["quality ownership across disciplines", "reducing dependency on late discovery", "shifting verification left"])
    add_bullet(doc, "Improve operational visibility into quality risks, verification pressure, and release readiness through metrics, telemetry, and Grafana dashboards \u2014 actionable insights for stakeholders.",
        ["operational visibility", "quality risks", "release readiness", "metrics, telemetry", "actionable insights"])
    add_bullet(doc, "Identify systemic bottlenecks and recurring failure patterns through defect trend analysis and root cause investigations \u2014 implement structural solutions.",
        ["systemic bottlenecks", "recurring failure patterns", "defect trend analysis", "structural solutions"])
    add_bullet(doc, "Evaluate and modernise QA tooling, test infrastructure, and verification workflows \u2014 Python-based frameworks, automated quality gates in CI/CD pipelines.",
        ["modernise QA tooling", "test infrastructure", "verification workflows", "automated quality gates"])
    add_bullet(doc, "Build strong collaboration between embedded QA engineers, operational QA functions, and external QA partners \u2014 unified quality approach.",
        ["embedded QA", "operational QA", "QA partners", "unified quality approach"])
    add_bullet(doc, "Support, mentor and grow team members \u2014 setting clear expectations, fostering psychological safety, encouraging continuous learning.",
        ["mentor and grow", "psychological safety", "continuous learning"])
    add_bullet(doc, "Own people management: performance reviews, goal-setting, career development, hiring \u2014 building a diverse, high-performing quality engineering team.",
        ["people management", "performance", "goal-setting", "career development", "hiring"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Quality Automation Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Production Observability & Release Quality")
    add_bullet(doc, "Improved release readiness and production observability \u2014 quality gates, deployment verification, and release stability monitoring.",
        ["release readiness", "production observability", "quality gates", "release stability"])
    add_bullet(doc, "Identified recurring release failure patterns; implemented process improvements reducing reactive incident response.",
        ["recurring release failure patterns", "process improvements", "reactive incident response"])
    add_bullet(doc, "Collaborated cross-functionally with Engineering and Ops to make quality a shared responsibility across the delivery pipeline.",
        ["cross-functionally", "quality a shared responsibility"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 QA Lead / Technical Specialist",
        "2013 \u2013 2021",
        "E-Commerce & Digital Platforms \u2014 Multi-Discipline Delivery | Embedded + Operational QA")
    add_bullet(doc, "Drove organisational change in how teams collaborated around quality \u2014 not just how QA teams operated, but how Engineering and Product shared ownership.",
        ["organisational change", "collaborated around quality", "shared ownership"])
    add_bullet(doc, "Evolved quality systems from manual, reactive validation to proactive, automated confidence-building across 10+ integrated systems.",
        ["reactive validation", "proactive, automated confidence-building"])
    add_bullet(doc, "Built and modernised test infrastructure: Python/Java frameworks, CI pipelines (Jenkins, GitHub Actions), automated verification workflows.",
        ["modernised test infrastructure", "CI pipelines", "automated verification workflows"])
    add_bullet(doc, "Established quality metrics and telemetry \u2014 defect density, automation coverage, cycle time, release stability \u2014 driving data-informed improvement.",
        ["quality metrics and telemetry", "data-informed improvement"])
    add_bullet(doc, "Identified systemic bottlenecks across programmes; structured problem-solving to address recurring failure patterns and reduce verification pressure.",
        ["systemic bottlenecks", "recurring failure patterns", "verification pressure"])
    add_bullet(doc, "Managed distributed teams (8\u201312 engineers across embedded QA and operational QA); mentored individuals, set expectations, fostered learning culture.",
        ["distributed teams", "embedded QA", "operational QA", "mentored", "learning culture"])
    add_bullet(doc, "Balanced structure with flexibility \u2014 adapting quality practices to complex, adaptive environments where answers weren\u2019t always obvious.",
        ["structure with flexibility", "complex, adaptive environments"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 QA Engineer / Consultant",
        "2008 \u2013 2013",
        "Core Banking Systems \u2014 Regulated Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "System-level quality validation for Finacle Core Banking \u2014 complex transaction flows, regulatory compliance, integration verification.",
        ["System-level quality validation", "complex transaction flows", "integration verification"])
    add_bullet(doc, "Built automated test infrastructure (Python, Java, Selenium); early experience identifying systemic quality patterns in regulated domains.",
        ["automated test infrastructure", "systemic quality patterns"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "Six Sigma Green Belt", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    lp2 = doc.add_paragraph()
    lp2.paragraph_format.space_before = Pt(3)
    r = lp2.add_run("Location: ")
    r.bold = True
    r.font.size = Pt(10)
    lp2.add_run("Malm\u00f6 (available for hybrid Malm\u00f6/Stockholm office)").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
