"""Generate QA/Test Manager resume – Nordea Pension External Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_QA_Test_Manager_Nordea_Resume"

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x5E)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "00005E")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x00, 0x5E)

    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # Summary
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "ISTQB Certified Quality Engineer & Test Manager with 14+ years of experience delivering high-quality results "
        "in large, complex agile IT environments. Proven expertise across the full testing lifecycle — analysis, test design, "
        "planning, execution, and reporting — supporting multiple applications and business areas simultaneously. "
        "Structured, quality-driven professional who lives and breathes quality, keeps cool under pressure, and meets "
        "harsh deadlines without compromising standards. Strong communicator who collaborates across different business areas "
        "and levels, challenges constructively, and continuously improves ways of working and processes. "
        "Team player with ownership mindset — experienced in supporting Head of QA/Test Management with diverse tasks, "
        "thinking outside the box, and voicing opinions while maintaining a collaborative, respectful approach. "
        "Extensive experience across retail, IoT, consumer apps, and enterprise systems in complex agile IT setups "
        "spanning multiple development areas."
    )
    r.font.size = Pt(10)

    # Skills
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Test & QA", "Manual testing · Test management · Test analysis · Test design · Test planning · Test execution · Test reporting · Regression testing · Acceptance testing · Exploratory testing"),
        ("Test Management", "Test strategy · Test scope · Risk-based testing · Defect management · Quality gates · Go/no-go decisions · Entry/exit criteria · KPI tracking · Process improvement"),
        ("Tools", "Jira · Confluence · TestRail · Zephyr · Azure DevOps · HP ALM · Postman · Playwright · SQL · Python · Git"),
        ("Agile & Methods", "Agile/Scrum/Kanban · Large-scale agile (SAFe) · DevOps · CI/CD · Continuous improvement · Cross-team collaboration · Sprint planning · Retrospectives"),
        ("Communication", "Stakeholder communication at all levels · Cross-business area collaboration · Status reporting · Risk escalation · Clear and structured documentation"),
        ("Certifications", "ISTQB Certified Tester (CTFL) · Six Sigma Green Belt · AWS Cloud Practitioner · Google Cloud ACE"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.2)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "F0F4FA")

    # Experience
    add_heading_block(doc, "Professional Experience")

    add_role(doc, "Quality Engineer / Test Manager / Team Lead Acting", "Ingka Digital (IKEA Group)", "Malmö, Sweden", "2022 – Present")
    for b in [
        "Deliver high-quality results across multiple applications supporting business operations in 32 markets — responsible for full testing lifecycle: analysis, test design, planning, execution, and reporting.",
        "Lead test management for complex, integrated digital solutions in a large-scale agile IT setup — coordinating across different business areas and development teams.",
        "Define test strategies, test designs, and quality gates — ensuring comprehensive coverage across APIs, frontends, integrations, and end-to-end business workflows.",
        "Collaborate across different business areas (Customer Support, Operations, IT, Finance) — communicating clearly at all levels from developers to senior leadership.",
        "Continuously improve ways of working, processes, and test practices — driving retrospectives, identifying improvement areas, and implementing changes within the QA team.",
        "Support Head of QA responsibilities — contributing to team planning, competence development, process governance, and cross-area coordination.",
        "Execute manual and automated testing (Playwright, Postman, SQL queries) — validating complex business logic, data flows, and system integrations under tight deadlines.",
        "Report test progress, risks, and quality status clearly — enabling informed go/no-go decisions for releases with confidence.",
        "Challenge constructively when quality standards are at risk — voicing concerns professionally while proposing solutions.",
    ]:
        bullet(doc, b)

    add_role(doc, "Quality & Release Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    for b in [
        "Managed test planning, execution, and reporting for a globally distributed application (300M+ users) — working in a fast-paced agile environment with harsh deadlines.",
        "Defined test scope and acceptance criteria — collaborating across engineering, product, and QA teams in a complex IT setup.",
        "Maintained composure under pressure — coordinating release quality for high-visibility deployments with minimal rollback tolerance.",
        "Drove quality improvements and process optimization — reducing defect leakage and improving test cycle efficiency.",
    ]:
        bullet(doc, b)

    add_role(doc, "Test Manager / Quality Lead", "HCLTech", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    for b in [
        "Led test management for large, complex agile IT setups — analysis, test design, planning, execution, and reporting across enterprise platforms (IoT, connected products, enterprise apps) for 8+ years.",
        "Managed QA across multiple development areas simultaneously — coordinating quality efforts for diverse application portfolios in complex integrated environments.",
        "Defined test strategies and test designs for multi-system landscapes — ensuring quality coverage across APIs, middleware, databases, frontends, and third-party integrations.",
        "Collaborated across different business areas and stakeholder levels — communicating test status, risks, and quality recommendations clearly and in a structured manner.",
        "Continuously improved ways of working and QA processes — introducing automation, optimizing test cycles, and establishing best practices adopted across teams.",
        "Supported Test Management leadership with diverse tasks — team coordination, competence planning, tool evaluation, and process governance.",
        "Led a team of 8–12 testers and engineers — fostering team spirit, coaching on test design, and building a culture of quality ownership and collaboration.",
        "Handled pressured situations and harsh deadlines with composure — delivering quality outcomes through structured prioritization and risk-based testing approaches.",
        "Worked in a large-scale agile/DevOps setup — sprint planning, daily stand-ups, retrospectives, and continuous delivery pipelines.",
    ]:
        bullet(doc, b)

    add_role(doc, "Test Engineer / QA Engineer", "HCL, Ultimate Digital, Marlabs, TekMindz", "India", "2008 – 2013")
    for b in [
        "Executed manual testing across enterprise applications (financial, retail) — test analysis, test design, execution, defect reporting in complex multi-tier systems.",
        "Defined test cases and test designs for business-critical workflows — ensuring thorough coverage and structured defect lifecycle management.",
        "Worked in agile and waterfall delivery models — adapting testing approach to project context and timelines.",
        "Built early test automation (Python, Selenium) — transitioning repetitive manual checks into automated regression suites.",
        "Collaborated with developers, business analysts, and project managers — communicating test results and quality risks at different levels.",
    ]:
        bullet(doc, b)

    # Values
    add_heading_block(doc, "Values & Working Style")
    for b in [
        "Collaboration: Team player who works across business areas and development teams — building trust and shared quality ownership.",
        "Ownership: Take responsibility for quality outcomes — driving initiatives, flagging risks, and ensuring follow-through.",
        "Passion: Live and breathe quality — staying updated on QA trends, tools, and best practices; continuously seeking improvement.",
        "Courage: Not afraid to voice opinions, challenge constructively, and push back when quality standards are at risk — always with respect and a solution-oriented mindset.",
    ]:
        bullet(doc, b)

    # Education
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # Languages
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent) · Danish (Basic — actively learning)")
    r.font.size = Pt(10)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")

def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#00005E;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#00005E;font-size:10.5pt;border-bottom:1px solid #00005E;padding-bottom:2px;margin-top:10px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#F0F4FA;font-weight:bold;width:18%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>ISTQB Certified Quality Engineer &amp; Test Manager with 14+ years of experience delivering high-quality results in large, complex agile IT environments. Proven expertise across the full testing lifecycle — analysis, test design, planning, execution, and reporting — supporting multiple applications and business areas simultaneously. Structured, quality-driven professional who lives and breathes quality, keeps cool under pressure, and meets harsh deadlines without compromising standards. Strong communicator who collaborates across different business areas and levels, challenges constructively, and continuously improves ways of working and processes. Team player with ownership mindset — experienced in supporting Head of QA/Test Management with diverse tasks, thinking outside the box, and voicing opinions while maintaining a collaborative, respectful approach. Extensive experience across retail, IoT, consumer apps, and enterprise systems in complex agile IT setups spanning multiple development areas.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Test &amp; QA</td><td>Manual testing · Test management · Test analysis · Test design · Test planning · Test execution · Test reporting · Regression testing · Acceptance testing · Exploratory testing</td></tr>
<tr><td class="cat">Test Management</td><td>Test strategy · Test scope · Risk-based testing · Defect management · Quality gates · Go/no-go decisions · Entry/exit criteria · KPI tracking · Process improvement</td></tr>
<tr><td class="cat">Tools</td><td>Jira · Confluence · TestRail · Zephyr · Azure DevOps · HP ALM · Postman · Playwright · SQL · Python · Git</td></tr>
<tr><td class="cat">Agile &amp; Methods</td><td>Agile/Scrum/Kanban · Large-scale agile (SAFe) · DevOps · CI/CD · Continuous improvement · Cross-team collaboration · Sprint planning · Retrospectives</td></tr>
<tr><td class="cat">Communication</td><td>Stakeholder communication at all levels · Cross-business area collaboration · Status reporting · Risk escalation · Clear and structured documentation</td></tr>
<tr><td class="cat">Certifications</td><td>ISTQB Certified Tester (CTFL) · Six Sigma Green Belt · AWS Cloud Practitioner · Google Cloud ACE</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Quality Engineer / Test Manager / Team Lead Acting <span class="meta">&nbsp;|&nbsp; Ingka Digital (IKEA Group) &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Deliver high-quality results across multiple applications supporting business operations in 32 markets — responsible for full testing lifecycle: analysis, test design, planning, execution, and reporting.</li>
<li>Lead test management for complex, integrated digital solutions in a large-scale agile IT setup — coordinating across different business areas and development teams.</li>
<li>Define test strategies, test designs, and quality gates — ensuring comprehensive coverage across APIs, frontends, integrations, and end-to-end business workflows.</li>
<li>Collaborate across different business areas (Customer Support, Operations, IT, Finance) — communicating clearly at all levels from developers to senior leadership.</li>
<li>Continuously improve ways of working, processes, and test practices — driving retrospectives, identifying improvement areas, and implementing changes within the QA team.</li>
<li>Support Head of QA responsibilities — contributing to team planning, competence development, process governance, and cross-area coordination.</li>
<li>Execute manual and automated testing (Playwright, Postman, SQL queries) — validating complex business logic, data flows, and system integrations under tight deadlines.</li>
<li>Report test progress, risks, and quality status clearly — enabling informed go/no-go decisions for releases with confidence.</li>
<li>Challenge constructively when quality standards are at risk — voicing concerns professionally while proposing solutions.</li>
</ul>

<p class="role">Quality &amp; Release Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Managed test planning, execution, and reporting for a globally distributed application (300M+ users) — working in a fast-paced agile environment with harsh deadlines.</li>
<li>Defined test scope and acceptance criteria — collaborating across engineering, product, and QA teams in a complex IT setup.</li>
<li>Maintained composure under pressure — coordinating release quality for high-visibility deployments with minimal rollback tolerance.</li>
<li>Drove quality improvements and process optimization — reducing defect leakage and improving test cycle efficiency.</li>
</ul>

<p class="role">Test Manager / Quality Lead <span class="meta">&nbsp;|&nbsp; HCLTech &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led test management for large, complex agile IT setups — analysis, test design, planning, execution, and reporting across enterprise platforms (IoT, connected products, enterprise apps) for 8+ years.</li>
<li>Managed QA across multiple development areas simultaneously — coordinating quality efforts for diverse application portfolios in complex integrated environments.</li>
<li>Defined test strategies and test designs for multi-system landscapes — ensuring quality coverage across APIs, middleware, databases, frontends, and third-party integrations.</li>
<li>Collaborated across different business areas and stakeholder levels — communicating test status, risks, and quality recommendations clearly and in a structured manner.</li>
<li>Continuously improved ways of working and QA processes — introducing automation, optimizing test cycles, and establishing best practices adopted across teams.</li>
<li>Supported Test Management leadership with diverse tasks — team coordination, competence planning, tool evaluation, and process governance.</li>
<li>Led a team of 8–12 testers and engineers — fostering team spirit, coaching on test design, and building a culture of quality ownership and collaboration.</li>
<li>Handled pressured situations and harsh deadlines with composure — delivering quality outcomes through structured prioritization and risk-based testing approaches.</li>
<li>Worked in a large-scale agile/DevOps setup — sprint planning, daily stand-ups, retrospectives, and continuous delivery pipelines.</li>
</ul>

<p class="role">Test Engineer / QA Engineer <span class="meta">&nbsp;|&nbsp; HCL, Ultimate Digital, Marlabs, TekMindz &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Executed manual testing across enterprise applications (financial, retail) — test analysis, test design, execution, defect reporting in complex multi-tier systems.</li>
<li>Defined test cases and test designs for business-critical workflows — ensuring thorough coverage and structured defect lifecycle management.</li>
<li>Worked in agile and waterfall delivery models — adapting testing approach to project context and timelines.</li>
<li>Built early test automation (Python, Selenium) — transitioning repetitive manual checks into automated regression suites.</li>
<li>Collaborated with developers, business analysts, and project managers — communicating test results and quality risks at different levels.</li>
</ul>

<h2>VALUES &amp; WORKING STYLE</h2>
<ul>
<li><strong>Collaboration:</strong> Team player who works across business areas and development teams — building trust and shared quality ownership.</li>
<li><strong>Ownership:</strong> Take responsibility for quality outcomes — driving initiatives, flagging risks, and ensuring follow-through.</li>
<li><strong>Passion:</strong> Live and breathe quality — staying updated on QA trends, tools, and best practices; continuously seeking improvement.</li>
<li><strong>Courage:</strong> Not afraid to voice opinions, challenge constructively, and push back when quality standards are at risk — always with respect and a solution-oriented mindset.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent) · Danish (Basic — actively learning)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

if __name__ == "__main__":
    build_docx()
    build_doc()
