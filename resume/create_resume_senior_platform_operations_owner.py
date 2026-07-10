"""Generate Senior Platform Operations Owner resume – IKEA Internal (Marketing Activation) – MAX 2 PAGES."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Senior_Platform_Operations_Owner_Resume"

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(9.5)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(9.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(9.5)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(5)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.7)
        section.bottom_margin = Cm(0.7)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.line_spacing = 1.0

    # Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(0)
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9)

    # Summary
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    r = p.add_run(
        "Senior Platform Operations professional with 12+ years of experience in IT operations, platform management, "
        "and engineering leadership within the IKEA ecosystem. Deep expertise in managing operations and maintenance of "
        "core systems — combining Platforms, SaaS, and home-grown solutions in complex enterprise landscapes. "
        "Proven ability to set up and drive platform operations, operational monitoring, support processes, and "
        "continuous improvement across product teams with a DevOps mindset. "
        "Skilled at influencing senior stakeholders, driving organizational change, and aligning engineering deliveries "
        "across teams. Collaborative leader who enables product teams to deliver effectively while stepping in hands-on when needed. "
        "Strong IKEA business model knowledge and deep understanding of how technology solutions are designed and operated "
        "across IKEA's marketing and digital landscape. Aligned with IKEA values — togetherness, simplicity, and continuous growth."
    )
    r.font.size = Pt(9.5)

    # Skills
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Platform Operations", "IT Operations (support, application, infrastructure) · Operations & maintenance of core systems · Operational monitoring · Platform support ownership · Incident management · SLA governance"),
        ("Leadership & Influence", "Aligning deliveries across product teams · Driving change in large organizations · Senior stakeholder influence · Supplier management · Engineering Manager collaboration · Best practice establishment"),
        ("DevOps & Tooling", "CI/CD · Monitoring & observability (Datadog, Cloud Monitoring) · Infrastructure as Code (Terraform) · Docker · Kubernetes · GitOps · Alerting · Runbooks · On-call processes"),
        ("Cloud & Platforms", "GCP (Cloud Run, Cloud Functions, Pub/Sub, BigQuery, GCS) · Azure · AWS · SaaS integration · Platform architecture · Scalability & performance"),
        ("Agile & Delivery", "Agile/Scrum · DevOps working set-up · Iterative development · Cross-team coordination · Roll-out management · Ways of working development"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.5)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9)
        set_cell_shading(c0, "E8F4FD")

    # Experience
    add_heading_block(doc, "Professional Experience")

    add_role(doc, "Platform Operations Lead / Team Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    for b in [
        "Own platform operations and support for enterprise digital products — accountable for operational monitoring, incident management, platform stability, and continuous improvement across product teams serving 32 IKEA markets.",
        "Align engineering deliveries across product teams — identifying synergies, establishing operational best practices, and driving platform operations initiatives to improve technical effectiveness.",
        "Set up and drive operations for new platform solutions — developing ways of working with Engineering Managers and suppliers to support development, roll-out, and scaling across all markets.",
        "Drive operational monitoring and efficiency — implementing observability (Cloud Monitoring, alerting, dashboards), reducing MTTR by 60%, and establishing proactive incident prevention.",
        "Manage operations and maintenance of core systems combining Platforms, SaaS (Genesys, Verint), and home-grown solutions — ensuring performance, scalability, and reliability at enterprise scale.",
        "Act as ambassador for operational best practices across the sub-domain and organization — influencing senior stakeholders and driving change in how we approach platform operations.",
        "Apply IKEA business model knowledge to guide how technology solutions are designed and operated — ensuring alignment with business needs and customer experience goals.",
        "Work closely with suppliers and vendor partners — managing relationships, SLAs, and co-creating operational improvements for platform stability and growth.",
    ]:
        bullet(doc, b)

    add_role(doc, "Platform Operations Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    for b in [
        "Managed platform operations for a globally distributed system (300M+ users) — monitoring, incident response, and performance optimization ensuring 99.9%+ uptime.",
        "Drove operational improvements through automation — CI/CD pipelines, Infrastructure as Code, and monitoring enhancements reducing manual operations by 50%.",
        "Collaborated with product teams in a DevOps setup — enabling iterative delivery while maintaining operational stability and scalability.",
    ]:
        bullet(doc, b)

    add_role(doc, "Technical Lead — Platform & Operations", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    for b in [
        "Led IT operations (support, application, infrastructure) for enterprise platforms — managing operations and maintenance of core systems across IKEA and LEGO for 8+ years.",
        "Set up and drove platform support processes, operational monitoring, incident management, and continuous improvement — establishing best practices adopted across multiple product teams.",
        "Managed technology landscape combining Platforms, SaaS, and home-grown solutions — ensuring reliability, scalability, and performance across complex integrated environments.",
        "Influenced senior stakeholders and drove organizational change — aligning operational practices, introducing DevOps working methods, and improving cross-team efficiency.",
        "Led a team of 8–12 engineers — coaching on operational excellence, developing ways of working, and enabling teams to take ownership of platform stability.",
        "Worked closely with suppliers and vendors — managing relationships, contracts, and co-driving operational improvements for platform services.",
        "Applied deep IKEA business model knowledge to guide platform design and operations — ensuring solutions supported business processes and customer experience at scale.",
    ]:
        bullet(doc, b)

    add_role(doc, "Software & Operations Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    for b in [
        "Built operational foundations across enterprise applications — application support, system monitoring, incident management, and continuous improvement.",
        "Managed operations of business-critical systems — ensuring uptime, performance, and user satisfaction through proactive monitoring and support processes.",
    ]:
        bullet(doc, b)

    # IKEA Ecosystem
    add_heading_block(doc, "IKEA Ecosystem Experience")
    for item in [
        "IKEA Customer Connect (VCS): Platform operations ownership — monitoring, support, incident management, and continuous improvement for digital products across 32 markets.",
        "Platform & SaaS Landscape: Deep experience operating IKEA's technology landscape combining Platforms, SaaS (Genesys, Verint, Salesforce), and home-grown solutions.",
        "IKEA Ways of Working: Product team model, DevOps setup, Agile/Scrum, cross-functional collaboration with Engineering Managers, suppliers, and stakeholders.",
        "Operational Excellence: Establishing best practices for monitoring, alerting, incident response, and platform support — acting as ambassador across the organization.",
        "Business Model Knowledge: Understanding how IKEA's business model translates into technology operations — guiding teams on efficient, scalable, customer-centric platform design.",
    ]:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # Values
    add_heading_block(doc, "IKEA Culture & Values")
    for v in [
        "Togetherness: Collaborative, co-creative approach — enabling product teams to succeed and stepping in hands-on when needed.",
        "Continuous Growth: DevOps mindset, continuous improvement, and operational excellence — always raising the bar for platform stability and efficiency.",
        "Simplicity: Establish clear, practical operational processes that teams can adopt — reducing complexity while maintaining reliability.",
    ]:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # Education
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    p.space_after = Pt(1)
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = p.add_run(" — IGNOU  |  ")
    r2.font.size = Pt(9.5)
    r3 = p.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(9.5)
    r4 = p.add_run(" — UP Technical University")
    r4.font.size = Pt(9.5)

    # Languages
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    r = p.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(9.5)
    r2 = p.add_run("English (Fluent)")
    r2.font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:9.5pt;margin:0.7cm 1.2cm;line-height:1.25}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:0}
.contact{text-align:center;font-size:9pt;margin-bottom:8px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:10px;margin-bottom:4px}
.role{font-weight:bold;margin-top:6px;margin-bottom:2px} .meta{color:#444;font-size:9pt}
ul{margin:2px 0 4px 16px;padding:0} li{margin-bottom:1px;font-size:9.5pt}
table{width:100%;border-collapse:collapse;font-size:9pt} td{padding:2px 5px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:20%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Senior Platform Operations professional with 12+ years of experience in IT operations, platform management, and engineering leadership within the IKEA ecosystem. Deep expertise in managing operations and maintenance of core systems — combining Platforms, SaaS, and home-grown solutions in complex enterprise landscapes. Proven ability to set up and drive platform operations, operational monitoring, support processes, and continuous improvement across product teams with a DevOps mindset. Skilled at influencing senior stakeholders, driving organizational change, and aligning engineering deliveries across teams. Collaborative leader who enables product teams to deliver effectively while stepping in hands-on when needed. Strong IKEA business model knowledge and deep understanding of how technology solutions are designed and operated across IKEA's marketing and digital landscape. Aligned with IKEA values — togetherness, simplicity, and continuous growth.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Platform Operations</td><td>IT Operations (support, application, infrastructure) · Operations &amp; maintenance of core systems · Operational monitoring · Platform support ownership · Incident management · SLA governance</td></tr>
<tr><td class="cat">Leadership</td><td>Aligning deliveries across product teams · Driving change in large organizations · Senior stakeholder influence · Supplier management · Engineering Manager collaboration · Best practice establishment</td></tr>
<tr><td class="cat">DevOps &amp; Tooling</td><td>CI/CD · Monitoring &amp; observability (Datadog, Cloud Monitoring) · Terraform · Docker · Kubernetes · GitOps · Alerting · Runbooks · On-call processes</td></tr>
<tr><td class="cat">Cloud &amp; Platforms</td><td>GCP (Cloud Run, Functions, Pub/Sub, BigQuery) · Azure · AWS · SaaS integration · Platform architecture · Scalability &amp; performance</td></tr>
<tr><td class="cat">Agile &amp; Delivery</td><td>Agile/Scrum · DevOps working set-up · Iterative development · Cross-team coordination · Roll-out management · Ways of working</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Platform Operations Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Own platform operations and support for enterprise digital products — accountable for operational monitoring, incident management, platform stability, and continuous improvement across product teams serving 32 IKEA markets.</li>
<li>Align engineering deliveries across product teams — identifying synergies, establishing operational best practices, and driving platform operations initiatives to improve technical effectiveness.</li>
<li>Set up and drive operations for new platform solutions — developing ways of working with Engineering Managers and suppliers to support development, roll-out, and scaling across all markets.</li>
<li>Drive operational monitoring and efficiency — implementing observability (Cloud Monitoring, alerting, dashboards), reducing MTTR by 60%, and establishing proactive incident prevention.</li>
<li>Manage operations and maintenance of core systems combining Platforms, SaaS (Genesys, Verint), and home-grown solutions — ensuring performance, scalability, and reliability at enterprise scale.</li>
<li>Act as ambassador for operational best practices across the sub-domain and organization — influencing senior stakeholders and driving change in how we approach platform operations.</li>
<li>Apply IKEA business model knowledge to guide how technology solutions are designed and operated — ensuring alignment with business needs and customer experience goals.</li>
<li>Work closely with suppliers and vendor partners — managing relationships, SLAs, and co-creating operational improvements for platform stability and growth.</li>
</ul>

<p class="role">Platform Operations Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Managed platform operations for a globally distributed system (300M+ users) — monitoring, incident response, and performance optimization ensuring 99.9%+ uptime.</li>
<li>Drove operational improvements through automation — CI/CD pipelines, Infrastructure as Code, and monitoring enhancements reducing manual operations by 50%.</li>
<li>Collaborated with product teams in a DevOps setup — enabling iterative delivery while maintaining operational stability and scalability.</li>
</ul>

<p class="role">Technical Lead — Platform &amp; Operations <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led IT operations (support, application, infrastructure) for enterprise platforms — managing operations and maintenance of core systems across IKEA and LEGO for 8+ years.</li>
<li>Set up and drove platform support processes, operational monitoring, incident management, and continuous improvement — establishing best practices adopted across multiple product teams.</li>
<li>Managed technology landscape combining Platforms, SaaS, and home-grown solutions — ensuring reliability, scalability, and performance across complex integrated environments.</li>
<li>Influenced senior stakeholders and drove organizational change — aligning operational practices, introducing DevOps working methods, and improving cross-team efficiency.</li>
<li>Led a team of 8–12 engineers — coaching on operational excellence, developing ways of working, and enabling teams to take ownership of platform stability.</li>
<li>Worked closely with suppliers and vendors — managing relationships, contracts, and co-driving operational improvements for platform services.</li>
<li>Applied deep IKEA business model knowledge to guide platform design and operations — ensuring solutions supported business processes and customer experience at scale.</li>
</ul>

<p class="role">Software &amp; Operations Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built operational foundations across enterprise applications — application support, system monitoring, incident management, and continuous improvement.</li>
<li>Managed operations of business-critical systems — ensuring uptime, performance, and user satisfaction through proactive monitoring and support processes.</li>
</ul>

<h2>IKEA ECOSYSTEM EXPERIENCE</h2>
<ul>
<li><strong>IKEA Customer Connect (VCS):</strong> Platform operations ownership — monitoring, support, incident management, and continuous improvement for digital products across 32 markets.</li>
<li><strong>Platform &amp; SaaS Landscape:</strong> Deep experience operating IKEA's technology landscape combining Platforms, SaaS (Genesys, Verint, Salesforce), and home-grown solutions.</li>
<li><strong>IKEA Ways of Working:</strong> Product team model, DevOps setup, Agile/Scrum, cross-functional collaboration with Engineering Managers, suppliers, and stakeholders.</li>
<li><strong>Operational Excellence:</strong> Establishing best practices for monitoring, alerting, incident response, and platform support — acting as ambassador across the organization.</li>
<li><strong>Business Model Knowledge:</strong> Understanding how IKEA's business model translates into technology operations — guiding teams on efficient, scalable, customer-centric platform design.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES</h2>
<ul>
<li><strong>Togetherness:</strong> Collaborative, co-creative approach — enabling product teams to succeed and stepping in hands-on when needed.</li>
<li><strong>Continuous Growth:</strong> DevOps mindset, continuous improvement, and operational excellence — always raising the bar for platform stability and efficiency.</li>
<li><strong>Simplicity:</strong> Establish clear, practical operational processes that teams can adopt — reducing complexity while maintaining reliability.</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU &nbsp;|&nbsp; <strong>B.Tech, Information Technology</strong> — UP Technical University &nbsp;|&nbsp; <strong>Languages:</strong> English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

if __name__ == "__main__":
    build_docx()
    build_doc()
