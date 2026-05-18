from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Operative_Test_Lead_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Operative_Test_Lead_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "video analytics",
    "video",
    "analytics",
    "AI",
    "computer vision",
    "automation",
    "test lead",
    "test planning",
    "test systems",
    "end-to-end",
    "E2E",
    "lab",
    "field",
    "network",
    "camera",
    "hardware",
    "system configuration",
    "hands-on",
    "leadership",
    "team",
    "stakeholder",
    "alignment",
    "continuous improvement",
    "logs",
    "analysis",
    "coverage",
    "framework",
    "Playwright",
    "Selenium",
    "CI/CD",
    "Python",
    "quality",
    "coaching",
    "mentoring",
    "ownership",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Operative Test Lead  |  Video Analytics, Automation & Hands-On Technical Leadership  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Hands-on Test Lead with 15+ years combining technical depth and team leadership "
        "across video, real-time communication, and AI-driven platforms. Designs, builds, "
        "and maintains realistic test systems across lab and field environments — automation, "
        "system configuration, and direct product interaction. Leads small, high-performing "
        "teams with ownership, trust, and autonomy. Acts as main point of contact — "
        "coordinates with stakeholders, shapes plans from ambiguous objectives, and ensures "
        "alignment across development, QA, and product teams. Strong analytical mindset — "
        "investigates system behaviour, logs, and test results to identify issues, patterns, "
        "and improvement opportunities. Continuously improves testing methods, tools, and "
        "ways of working. Experience with video analytics, AI/ML-based systems, and "
        "network-connected hardware. Pragmatic, self-driven, and comfortable navigating "
        "uncertainty. Leads through action — grounded leadership style that keeps the team "
        "focused and moving forward.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Hands-On Technical Leadership: ",
            "Designed, built, and maintained realistic test systems at IKEA (VCS — video "
            "collaboration on network-connected devices), Truecaller, and LEGO. Automation, "
            "system configuration, infrastructure setup (lab & production environments). "
            "Direct product interaction. Comfortable writing code, debugging systems, "
            "and solving complex technical problems hands-on while guiding the team.",
        ),
        (
            "Video, Analytics & AI Testing: ",
            "End-to-end ownership of IKEA's Visual Collaboration Solution (VCS) — video "
            "streams, real-time communication, network-connected hardware. Validated "
            "real-world performance across 30+ markets. AI-assisted testing and analytics. "
            "Experience with video pipelines, media quality, and system-level validation. "
            "Familiar with computer vision and AI/ML testing patterns.",
        ),
        (
            "Team Alignment & Stakeholder Coordination: ",
            "Main point of contact for team — coordinating with stakeholders, shaping plans "
            "from complex/unclear objectives, ensuring alignment across teams. Led teams "
            "of 4–15 engineers. Coaching, mentoring, competence development. Natural leadership "
            "style based on trust, ownership, and autonomy. Collaborative and informal.",
        ),
        (
            "Test Planning, Execution & Analysis: ",
            "Led test planning and execution across complex scenarios — coverage strategy, "
            "risk-based testing, quality gates. Analysed system behaviour, logs, and "
            "test results to identify issues and improvement opportunities. End-to-end "
            "perspective — from unit through system/integration to field validation.",
        ),
        (
            "Continuous Improvement: ",
            "Translated ambiguous challenges into actionable work. Continuously improved "
            "testing methods, tools, and ways of working. Led Selenium-to-Playwright "
            "migration (3x faster, 50% CI reduction). Built reusable frameworks and CoE "
            "practices. Six Sigma Green Belt — process improvement mindset.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Test Leadership: ",
            "Hands-on test lead. Test planning, strategy, coverage. "
            "Quality gates. Risk-based testing. Lab & field environments. "
            "System configuration. End-to-end validation. Real-world performance",
        ),
        (
            "Automation & Frameworks: ",
            "Playwright, Selenium, Cypress, Appium. API testing (REST, GraphQL). "
            "Framework design & architecture. CI/CD (GitHub Actions, Jenkins). "
            "TDD/BDD. AI-assisted test generation. Performance testing",
        ),
        (
            "Video, Analytics & AI: ",
            "Video analytics, real-time communication, media quality. "
            "Network-connected devices/hardware. Computer vision & AI/ML testing. "
            "System-level validation. Log analysis. Pattern identification",
        ),
        (
            "Team & Stakeholder: ",
            "Team leadership (4–15 engineers). Coaching, mentoring, competence "
            "development. Stakeholder alignment. Cross-team coordination. "
            "Shaping plans from ambiguous objectives. Ownership & autonomy",
        ),
        (
            "DevOps & Cloud: ",
            "GCP (Cloud Run, BigQuery, Cloud Functions), AWS. Docker, Kubernetes, "
            "Terraform. Infrastructure as Code. Grafana observability. "
            "System monitoring and diagnostics",
        ),
        (
            "Programming: ",
            "Python, TypeScript, C#, Java, Node.js. REST APIs. SQL. "
            "Bash/scripting. Git. Jira, Confluence, TestRail",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Team Lead / Test Lead (VCS Video Platform)", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  Visual Collaboration Solution — Video, Network Devices, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Hands-on technical leadership — end-to-end ownership of IKEA's Visual "
        "Collaboration Solution (VCS): video streams, real-time communication on "
        "network-connected hardware. Designed, built, and maintained test systems across "
        "lab and production environments. Automation, system configuration, direct product "
        "interaction. Scaled from 2K to 50K usage across 30+ global markets.",
        "Team leadership — led team of engineers and consultants. Coaching, mentoring, "
        "competence development. Fostered ownership, trust, and autonomy. Acted as main "
        "point of contact — coordinated with stakeholders, shaped plans from ambiguous "
        "objectives, ensured alignment across development, QA, and product teams.",
        "Test planning & execution — led test activities ensuring strong coverage, quality, "
        "and efficient execution in complex video/network scenarios. End-to-end perspective "
        "from unit through system integration to real-world field validation. Quality gates "
        "and acceptance criteria. Risk-based testing.",
        "Analysis & problem solving — investigated system behaviour, logs, video analytics "
        "data, and test results to identify issues, patterns, and improvement opportunities. "
        "Grafana dashboards for observability. AI-assisted analytics.",
        "Continuous improvement — led Selenium-to-Playwright migration (3x faster execution, "
        "50% CI pipeline reduction). Continuously improved testing methods, tools, and ways "
        "of working. Built reusable automation frameworks. CoE practices — guidelines, "
        "quality gates, knowledge sharing.",
        "Automation & DevOps — Playwright, API testing, CI/CD (GitHub Actions). Terraform, "
        "Docker on GCP. Data pipelines (BigQuery, Cloud Functions). Python, TypeScript. "
        "AI-assisted — 30% velocity improvement. Recognized as Exceptional Performer.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Release & Automation Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users, Real-Time Systems", bold=True, size=10)
    tc_bullets = [
        "Hands-on test leadership at global scale — owned release readiness, test coverage, "
        "and quality for a real-time communication platform serving 300M+ users. Built "
        "automated tools and workflows. CI pipelines, deployment automation. System-level "
        "validation. Log analysis and diagnostics. AWS cloud infrastructure.",
        "Team coordination — backlog prioritization, cross-functional alignment, resource "
        "planning. Feature flag management — controlled automated rollouts, risk analysis, "
        "data-driven go/no-go decisions. Stakeholder interface. Agile (Scrum/Kanban).",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "HCLTech — LEGO & IKEA Group, Denmark & Sweden — Technical Specialist / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2013 – 2021  |  E-Commerce, Mobile & Enterprise — Lab & Field Testing", bold=True, size=10)
    lego_bullets = [
        "LEGO & IKEA (2017–2021): Led test scope, planning, and execution across iOS, "
        "Android, and web platforms. Designed test systems for lab and real-world "
        "environments. Multi-partner coordination. Automation strategy across mobile "
        "and web. Tested on physical devices — beta OS releases, new hardware. Release "
        "readiness and quality gate decisions. Onshore-offshore team leadership (8–10).",
        "SDET Lead (2013–2017): Hands-on framework design — Selenium (C#, NUnit), CI "
        "integration. Full SDLC ownership. System behaviour investigation, log analysis, "
        "test result patterns. Translated ambiguous challenges into actionable test plans. "
        "Coached and mentored team members. Continuous improvement of methods and tools.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "Banking & Enterprise — Test Engineer / Consultant", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2013  |  Core Banking Systems — Hardware Integration, System Validation", bold=True, size=10)
    fin_bullets = [
        "System-level testing for banking platforms — hardware integration (biometric "
        "authentication devices with Finacle CBS). End-to-end validation, system "
        "configuration, automated test scripts. ETL automation (Pentaho). Log analysis, "
        "data validation (SQL, UNIX). Regulated environments, security compliance.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Scaled IKEA VCS (video platform) from 2K to 50K across 30+ markets — full "
        "end-to-end ownership of test systems, automation, and quality. Exceptional Performer.",
        "Led Selenium-to-Playwright migration — 3x faster execution, 50% CI reduction. "
        "Continuously improved testing methods, tools, and ways of working.",
        "Built and led high-performing teams with ownership, trust, and autonomy — coaching, "
        "mentoring, competence development. Shaped plans from ambiguous objectives.",
        "Designed realistic test systems across lab and field environments — video streams, "
        "network-connected devices, real-time communication. System configuration and "
        "hardware interaction.",
        "Acted as main point of contact — stakeholder alignment, cross-team coordination, "
        "and clear communication of insights to drive quality improvements.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Six Sigma Green Belt\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Operative Test Lead Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Operative Test Lead | Video Analytics, Automation &amp; Hands-On Technical Leadership | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Hands-on</span> Test Lead with 15+ years combining technical depth and <span class="hl">team</span> <span class="hl">leadership</span> across <span class="hl">video</span>, real-time communication, and <span class="hl">AI</span>-driven platforms. Designs, builds, and maintains realistic <span class="hl">test systems</span> across <span class="hl">lab</span> and <span class="hl">field</span> environments — <span class="hl">automation</span>, <span class="hl">system configuration</span>, and direct product interaction. Leads small, high-performing <span class="hl">team</span>s with <span class="hl">ownership</span>, trust, and autonomy. Acts as main point of contact — coordinates with <span class="hl">stakeholder</span>s, shapes plans from ambiguous objectives, ensures <span class="hl">alignment</span> across development, QA, and product teams. Strong analytical mindset — investigates system behaviour, <span class="hl">logs</span>, and test results to identify issues, patterns, and improvement opportunities. Continuously improves testing methods, tools, and ways of working. Experience with <span class="hl">video analytics</span>, <span class="hl">AI</span>/ML-based systems, and <span class="hl">network</span>-connected <span class="hl">hardware</span>. Pragmatic, self-driven, comfortable navigating uncertainty.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Hands-On Technical Leadership:</b> Designed, built, maintained realistic <span class="hl">test systems</span> at IKEA (VCS — <span class="hl">video</span> collaboration on <span class="hl">network</span>-connected devices), Truecaller, LEGO. <span class="hl">Automation</span>, <span class="hl">system configuration</span>, infrastructure (<span class="hl">lab</span> &amp; production). Direct product interaction. Writes code, debugs systems, solves complex problems <span class="hl">hands-on</span> while guiding the <span class="hl">team</span>.<br><br>
  <b>Video, Analytics &amp; AI Testing:</b> <span class="hl">End-to-end</span> ownership of IKEA's VCS — <span class="hl">video</span> streams, real-time communication, <span class="hl">network</span>-connected <span class="hl">hardware</span>. Validated real-world performance across 30+ markets. <span class="hl">AI</span>-assisted testing and <span class="hl">analytics</span>. <span class="hl">Video</span> pipelines, media quality, system-level validation. Familiar with <span class="hl">computer vision</span> and <span class="hl">AI</span>/ML testing patterns.<br><br>
  <b>Team Alignment &amp; Stakeholder Coordination:</b> Main point of contact — coordinating with <span class="hl">stakeholder</span>s, shaping plans from complex/unclear objectives, ensuring <span class="hl">alignment</span>. Led teams of 4–15. <span class="hl">Coaching</span>, <span class="hl">mentoring</span>. Natural <span class="hl">leadership</span> based on trust, <span class="hl">ownership</span>, autonomy.<br><br>
  <b>Test Planning, Execution &amp; Analysis:</b> Led <span class="hl">test planning</span> across complex scenarios — <span class="hl">coverage</span>, risk-based testing, <span class="hl">quality</span> gates. Analysed system behaviour, <span class="hl">logs</span>, test results for issues and patterns. <span class="hl">End-to-end</span> — unit through system/integration to <span class="hl">field</span> validation.<br><br>
  <b>Continuous Improvement:</b> Translated ambiguous challenges into actionable work. Improved testing methods, tools, ways of working. <span class="hl">Selenium</span>-to-<span class="hl">Playwright</span> migration (3x faster, 50% CI reduction). Built reusable <span class="hl">framework</span>s. <span class="hl">Six Sigma</span> Green Belt.</p>

  <div class="section">Core Competencies</div>
  <p><b>Test Leadership:</b> Hands-on test lead. Test planning, strategy, coverage. Quality gates. Risk-based testing. Lab &amp; field environments. System configuration. End-to-end validation<br>
  <b>Automation &amp; Frameworks:</b> Playwright, Selenium, Cypress, Appium. API testing (REST, GraphQL). Framework design. CI/CD (GitHub Actions, Jenkins). TDD/BDD. AI-assisted. Performance testing<br>
  <b>Video, Analytics &amp; AI:</b> Video analytics, real-time communication, media quality. Network-connected devices/hardware. Computer vision &amp; AI/ML testing. System-level validation. Log analysis<br>
  <b>Team &amp; Stakeholder:</b> Team leadership (4–15). Coaching, mentoring, competence development. Stakeholder alignment. Cross-team coordination. Ownership &amp; autonomy<br>
  <b>DevOps &amp; Cloud:</b> GCP, AWS. Docker, Kubernetes, Terraform. GitHub Actions. Grafana observability. System monitoring &amp; diagnostics<br>
  <b>Programming:</b> Python, TypeScript, C#, Java, Node.js. REST APIs. SQL. Bash. Git. Jira, Confluence, TestRail</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Team Lead / Test Lead (VCS Video Platform)</div>
  <div class="job-sub">Mar 2022 – Present | Visual Collaboration Solution — Video, Network Devices, 30+ Markets</div>
  <ul>
    <li><span class="hl">Hands-on</span> technical <span class="hl">leadership</span> — <span class="hl">end-to-end</span> ownership of VCS: <span class="hl">video</span> streams, real-time communication on <span class="hl">network</span>-connected <span class="hl">hardware</span>. Designed, built, maintained <span class="hl">test systems</span> across <span class="hl">lab</span> and production. <span class="hl">Automation</span>, <span class="hl">system configuration</span>, direct product interaction. Scaled 2K → 50K across 30+ markets.</li>
    <li><span class="hl">Team</span> <span class="hl">leadership</span> — engineers + consultants. <span class="hl">Coaching</span>, <span class="hl">mentoring</span>, competence development. <span class="hl">Ownership</span>, trust, autonomy. Main point of contact — <span class="hl">stakeholder</span> <span class="hl">alignment</span>, shaped plans from ambiguous objectives.</li>
    <li><span class="hl">Test planning</span> &amp; execution — strong <span class="hl">coverage</span>, <span class="hl">quality</span>, efficient execution in complex <span class="hl">video</span>/<span class="hl">network</span> scenarios. <span class="hl">End-to-end</span>. <span class="hl">Quality</span> gates. Risk-based testing.</li>
    <li><span class="hl">Analysis</span> — investigated system behaviour, <span class="hl">logs</span>, <span class="hl">video analytics</span> data, test results. Identified issues, patterns, improvements. <span class="hl">Grafana</span> observability. <span class="hl">AI</span>-assisted <span class="hl">analytics</span>.</li>
    <li><span class="hl">Continuous improvement</span> — <span class="hl">Selenium</span>-to-<span class="hl">Playwright</span> migration (3x faster, 50% CI reduction). Improved methods, tools, ways of working. Reusable <span class="hl">framework</span>s. CoE practices.</li>
    <li><span class="hl">Automation</span> &amp; DevOps — <span class="hl">Playwright</span>, API testing, <span class="hl">CI/CD</span> (GitHub Actions). <span class="hl">Terraform</span>, <span class="hl">Docker</span> on GCP. <span class="hl">Python</span>, TypeScript. 30% velocity improvement. Exceptional Performer.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Release &amp; Automation Engineer</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users, Real-Time Systems</div>
  <ul>
    <li><span class="hl">Hands-on</span> test <span class="hl">leadership</span> at global scale — release readiness, <span class="hl">coverage</span>, <span class="hl">quality</span> for real-time communication (300M+ users). Built automated tools. <span class="hl">CI/CD</span> pipelines. System validation. <span class="hl">Log</span> <span class="hl">analysis</span>. AWS.</li>
    <li><span class="hl">Team</span> coordination — backlog prioritization, <span class="hl">alignment</span>, resource planning. Feature flags — data-driven rollouts. <span class="hl">Stakeholder</span> interface. Agile.</li>
  </ul>

  <div class="job-title">HCLTech — LEGO &amp; IKEA, Denmark &amp; Sweden — Technical Specialist / Test Lead</div>
  <div class="job-sub">2013 – 2021 | E-Commerce, Mobile &amp; Enterprise — Lab &amp; Field Testing</div>
  <ul>
    <li>LEGO &amp; IKEA (2017–21): <span class="hl">Test planning</span> &amp; execution across iOS, Android, web. Designed <span class="hl">test systems</span> for <span class="hl">lab</span> and real-world environments. Tested on physical devices — beta OS, new <span class="hl">hardware</span>. <span class="hl">Automation</span> strategy. <span class="hl">Quality</span> gates. <span class="hl">Team</span> <span class="hl">leadership</span> (8–10).</li>
    <li>SDET Lead (2013–17): <span class="hl">Hands-on</span> <span class="hl">framework</span> design — <span class="hl">Selenium</span> (C#, NUnit), CI. System behaviour, <span class="hl">log</span> <span class="hl">analysis</span>, test result patterns. Translated ambiguous challenges into actionable plans. <span class="hl">Coaching</span> &amp; <span class="hl">mentoring</span>. <span class="hl">Continuous improvement</span>.</li>
  </ul>

  <div class="job-title">Banking &amp; Enterprise — Test Engineer / Consultant</div>
  <div class="job-sub">2008 – 2013 | Core Banking — Hardware Integration, System Validation</div>
  <ul>
    <li>System-level testing — <span class="hl">hardware</span> integration (biometric devices + Finacle CBS). <span class="hl">End-to-end</span> validation, <span class="hl">system configuration</span>, automated scripts. ETL <span class="hl">automation</span>. <span class="hl">Log</span> <span class="hl">analysis</span>, data validation (SQL, UNIX). Security compliance.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Scaled IKEA VCS (<span class="hl">video</span>) 2K → 50K across 30+ markets — full <span class="hl">end-to-end</span> <span class="hl">ownership</span>. Exceptional Performer.</li>
    <li><span class="hl">Selenium</span>-to-<span class="hl">Playwright</span> — 3x faster, 50% CI reduction. <span class="hl">Continuous improvement</span> of methods and tools.</li>
    <li>Built high-performing <span class="hl">team</span>s with <span class="hl">ownership</span>, trust, autonomy — <span class="hl">coaching</span>, <span class="hl">mentoring</span>. Shaped plans from ambiguity.</li>
    <li>Designed realistic <span class="hl">test systems</span> — <span class="hl">video</span>, <span class="hl">network</span> devices, <span class="hl">lab</span> &amp; <span class="hl">field</span>. <span class="hl">System configuration</span>, <span class="hl">hardware</span> interaction.</li>
    <li>Main point of contact — <span class="hl">stakeholder</span> <span class="hl">alignment</span>, cross-<span class="hl">team</span> coordination, driving <span class="hl">quality</span>.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>Six Sigma Green Belt<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>UiPath RPA Certified</td>
    </tr>
  </table>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
