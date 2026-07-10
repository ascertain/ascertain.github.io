"""
Resume: Platform Manager – Integration Services, Inter IKEA (INTERNAL)
Focus: Integration platform leadership, platform-as-a-product mindset, event-driven architecture,
       API management, people leadership, DevOps, modernisation, stakeholder communication.
NOTE: User is NOT a hardcore coder but has 8 years in IKEA ecosystem as an engineer across areas.
      Frame as a leader who bridges technical depth with service delivery and people management.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Platform_Manager_Integration_Services_Resume"
IKEA_BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = IKEA_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = IKEA_BLUE
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = IKEA_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Technology leader with 8+ years in the IKEA ecosystem, combining hands-on integration engineering experience "
        "with people leadership and a strong platform-as-a-product mindset. Experienced in managing teams that build and "
        "operate integration services connecting systems and data across IKEA — driving modernisation towards event-driven "
        "architectures and API-first patterns. Skilled at translating complex technical topics into simple business language, "
        "building trusted relationships across organisational boundaries, and aligning platform capabilities with evolving "
        "product team needs. Brings practical understanding of enterprise integration patterns, DevOps automation, cloud "
        "platforms, and service delivery — paired with a genuine passion for growing high-performing teams, optimising costs, "
        "and fostering a culture of continuous improvement rooted in IKEA values."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Platform-as-a-Product Leadership", "Integration (Events / APIs / Data)", "People & Team Development",
        "DevOps & Automation", "Stakeholder Communication", "Modernisation & Roadmapping",
        "Budget & Cost Optimisation", "IKEA Ecosystem (8+ years)", "Service Delivery (ITIL aware)",
        "Agile & Scaled Delivery (OKRs)", "Cloud & On-Prem Platforms", "Developer Experience & Self-Service",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TECHNICAL & DOMAIN KNOWLEDGE ──
    add_heading_block(doc, "Technical & Domain Knowledge")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    skills = (
        "Integration Patterns: Event-driven architecture, REST/GraphQL APIs, Pub/Sub messaging, data replication, batch & real-time  •  "
        "Cloud Platforms: GCP (Pub/Sub, Cloud Run, BigQuery, GKE), hybrid cloud/on-prem  •  "
        "DevOps & CI/CD: GitHub Actions, Docker, Kubernetes, Terraform, ArgoCD  •  "
        "API Management: OpenAPI, API gateways, versioning, contract governance  •  "
        "Development: TypeScript, Python, Node.js; familiarity with Java  •  "
        "Service Operations: ITIL awareness, incident management, SLA tracking  •  "
        "Metrics & Value Tracking: OKRs, velocity metrics, delivery KPIs  •  "
        "IKEA Understanding: Franchise model, Inter IKEA / Ingka structure, cross-unit collaboration"
    )
    r = tp.add_run(skills)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- Ingka Digital / IKEA ---
    role_header(doc, "Team Lead & Platform Engineer – Customer Connect Integration",
                "Ingka Digital / IKEA", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Lead a cross-functional engineering team responsible for integration services that connect 5+ IKEA customer contact platforms (VCS, CSSP, Genesys, Verint, Chatbot) — treating the integration layer as a product consumed by multiple teams.")
    bullet(doc, "Drive platform-as-a-product mindset: measure success by the velocity and value delivered to consuming teams; provide self-service capabilities, documentation, and best practices to reduce friction.")
    bullet(doc, "Own the integration modernisation roadmap: migrating from legacy point-to-point integrations towards event-driven architecture (GCP Pub/Sub) and well-governed API contracts (OpenAPI).")
    bullet(doc, "Communicate platform roadmaps and technical plans to business and product stakeholders in simple, accessible language — building trust and alignment across organisational boundaries.")
    bullet(doc, "Foster strong DevOps practices and automation: CI/CD pipelines (GitHub Actions), infrastructure-as-code (Terraform), containerised deployments (Cloud Run, GKE) — shortening delivery cycles.")
    bullet(doc, "Grow team capability through competence development, knowledge sharing, and creating conditions for engineers to innovate and learn continuously.")
    bullet(doc, "Manage budget planning and cost monitoring for platform services; optimise cloud spend while maintaining reliability and performance for mission-critical integrations.")
    bullet(doc, "Collaborate with platform networks and Enterprise Architecture to ensure alignment with IKEA's Technology Architecture and strategic direction.")
    bullet(doc, "Drive developer experience improvements: better tooling, clearer documentation, self-service onboarding — enabling other teams to integrate faster and more independently.")
    bullet(doc, "Navigate ambiguity in a fast-changing landscape; balance modernisation ambitions with operational stability and day-to-day service delivery.")

    # --- Truecaller ---
    role_header(doc, "Platform Engineer",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Worked on platform integration services in a high-scale SaaS environment; gained experience with event-driven architectures and distributed system operations at scale.")
    bullet(doc, "Contributed to developer experience improvements and CI/CD automation; understood how platform teams serve internal consumers effectively.")

    # --- HCLTech for IKEA ---
    role_header(doc, "Senior Engineer & Integration Lead – IKEA Enterprise Platforms",
                "HCLTech (for IKEA)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Spent 8 years embedded in IKEA's integration landscape, working across supply chain, e-commerce, order management, and customer platform domains — gaining in-depth understanding of how systems and data connect across IKEA.")
    bullet(doc, "Led integration development connecting 10+ enterprise systems; designed event-driven and API-based integration patterns that enabled data sharing across product and platform teams.")
    bullet(doc, "Managed a team of 8–12 engineers (onshore & offshore): competence development, succession planning, workload allocation, and performance coaching — growing high-performing teams.")
    bullet(doc, "Drove integration modernisation initiatives: migrating legacy batch integrations to real-time event streaming and API-first approaches, reducing complexity and improving reliability.")
    bullet(doc, "Owned service delivery for integration services: SLA tracking, incident management, operational reporting — with ITIL-aligned practices for IT service operations.")
    bullet(doc, "Managed budgets and optimised costs across integration services; made data-driven decisions on build-vs-buy, cloud migration, and resource allocation.")
    bullet(doc, "Facilitated Agile delivery using OKRs and velocity metrics to track value; drove continuous improvement through retrospectives and process optimisation.")
    bullet(doc, "Built trusted relationships with stakeholders across Inter IKEA and Ingka — translating complex integration challenges into business language that non-technical leaders could act on.")
    bullet(doc, "Championed DevOps practices and automation: pipeline design, automated testing, deployment automation — improving delivery speed and reliability across the integration landscape.")

    # --- India ---
    role_header(doc, "Software Engineer",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Developed enterprise web applications and integration components (Java, .NET); built foundational understanding of service-oriented architecture and system integration.")
    bullet(doc, "Gained experience across full software lifecycle: requirements, design, development, testing, and production support in globally distributed delivery models.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "Google Cloud Associate Cloud Engineer",
        "AWS Certified Cloud Practitioner",
        "ISTQB Certified Tester",
        "Six Sigma Green Belt (Process Improvement)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("PGDOM")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – IGNOU  |  ").font.size = Pt(9)
    r2 = p.add_run("B.Tech Information Technology")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – UP Technical University").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Conversational)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    # ── WHY THIS ROLE ──
    add_heading_block(doc, "Why This Role")
    why = (
        "Integration is the connective tissue of IKEA's digital landscape — and after 8+ years working across "
        "different parts of the ecosystem, I've seen first-hand how powerful it is when systems share data seamlessly "
        "through well-designed events and APIs. I'm passionate about treating integration services as a product: "
        "measuring success by how much friction we remove for consuming teams, not just by uptime metrics. "
        "This role brings together everything I care about — leading people, modernising platforms, simplifying "
        "complexity, and enabling the many teams across IKEA to move faster. I understand how IKEA works as a "
        "franchise system, I know where the integration pain points live, and I'm energised by the opportunity "
        "to drive the shift towards a simpler, event-driven integration landscape."
    )
    pw = doc.add_paragraph(why)
    pw.paragraph_format.space_before = Pt(3)
    for run in pw.runs:
        run.font.size = Pt(9)
        run.italic = True

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#0051BA;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#0051BA;font-size:11pt;border-bottom:1px solid #0051BA;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#0051BA;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    .why{font-style:italic;font-size:9pt;margin-top:4px}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Technology leader with 8+ years in the IKEA ecosystem, combining hands-on integration engineering experience
with people leadership and a strong platform-as-a-product mindset. Experienced in managing teams that build and
operate integration services connecting systems and data across IKEA — driving modernisation towards event-driven
architectures and API-first patterns. Skilled at translating complex technical topics into simple business language,
building trusted relationships across organisational boundaries, and aligning platform capabilities with evolving
product team needs. Brings practical understanding of enterprise integration patterns, DevOps automation, cloud
platforms, and service delivery — paired with a genuine passion for growing high-performing teams, optimising costs,
and fostering a culture of continuous improvement rooted in IKEA values.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Platform-as-a-Product Leadership</td><td>✓ Integration (Events / APIs / Data)</td><td>✓ People &amp; Team Development</td></tr>
<tr><td>✓ DevOps &amp; Automation</td><td>✓ Stakeholder Communication</td><td>✓ Modernisation &amp; Roadmapping</td></tr>
<tr><td>✓ Budget &amp; Cost Optimisation</td><td>✓ IKEA Ecosystem (8+ years)</td><td>✓ Service Delivery (ITIL aware)</td></tr>
<tr><td>✓ Agile &amp; Scaled Delivery (OKRs)</td><td>✓ Cloud &amp; On-Prem Platforms</td><td>✓ Developer Experience &amp; Self-Service</td></tr>
</table>

<h2>TECHNICAL &amp; DOMAIN KNOWLEDGE</h2>
<p class="tools">Integration Patterns: Event-driven architecture, REST/GraphQL APIs, Pub/Sub messaging, data replication, batch &amp; real-time &bull;
Cloud Platforms: GCP (Pub/Sub, Cloud Run, BigQuery, GKE), hybrid cloud/on-prem &bull;
DevOps &amp; CI/CD: GitHub Actions, Docker, Kubernetes, Terraform, ArgoCD &bull;
API Management: OpenAPI, API gateways, versioning, contract governance &bull;
Development: TypeScript, Python, Node.js; familiarity with Java &bull;
Service Operations: ITIL awareness, incident management, SLA tracking &bull;
Metrics &amp; Value Tracking: OKRs, velocity metrics, delivery KPIs &bull;
IKEA Understanding: Franchise model, Inter IKEA / Ingka structure, cross-unit collaboration</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Team Lead &amp; Platform Engineer – Customer Connect Integration &nbsp;|&nbsp; Ingka Digital / IKEA, Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Lead a cross-functional engineering team responsible for integration services that connect 5+ IKEA customer contact platforms (VCS, CSSP, Genesys, Verint, Chatbot) — treating the integration layer as a product consumed by multiple teams.</li>
<li>Drive platform-as-a-product mindset: measure success by the velocity and value delivered to consuming teams; provide self-service capabilities, documentation, and best practices to reduce friction.</li>
<li>Own the integration modernisation roadmap: migrating from legacy point-to-point integrations towards event-driven architecture (GCP Pub/Sub) and well-governed API contracts (OpenAPI).</li>
<li>Communicate platform roadmaps and technical plans to business and product stakeholders in simple, accessible language — building trust and alignment across organisational boundaries.</li>
<li>Foster strong DevOps practices and automation: CI/CD pipelines (GitHub Actions), infrastructure-as-code (Terraform), containerised deployments (Cloud Run, GKE) — shortening delivery cycles.</li>
<li>Grow team capability through competence development, knowledge sharing, and creating conditions for engineers to innovate and learn continuously.</li>
<li>Manage budget planning and cost monitoring for platform services; optimise cloud spend while maintaining reliability and performance for mission-critical integrations.</li>
<li>Collaborate with platform networks and Enterprise Architecture to ensure alignment with IKEA's Technology Architecture and strategic direction.</li>
<li>Drive developer experience improvements: better tooling, clearer documentation, self-service onboarding — enabling other teams to integrate faster and more independently.</li>
<li>Navigate ambiguity in a fast-changing landscape; balance modernisation ambitions with operational stability and day-to-day service delivery.</li>
</ul>

<p class="role">Platform Engineer &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Worked on platform integration services in a high-scale SaaS environment; gained experience with event-driven architectures and distributed system operations at scale.</li>
<li>Contributed to developer experience improvements and CI/CD automation; understood how platform teams serve internal consumers effectively.</li>
</ul>

<p class="role">Senior Engineer &amp; Integration Lead – IKEA Enterprise Platforms &nbsp;|&nbsp; HCLTech (for IKEA), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Spent 8 years embedded in IKEA's integration landscape, working across supply chain, e-commerce, order management, and customer platform domains — gaining in-depth understanding of how systems and data connect across IKEA.</li>
<li>Led integration development connecting 10+ enterprise systems; designed event-driven and API-based integration patterns that enabled data sharing across product and platform teams.</li>
<li>Managed a team of 8–12 engineers (onshore &amp; offshore): competence development, succession planning, workload allocation, and performance coaching — growing high-performing teams.</li>
<li>Drove integration modernisation initiatives: migrating legacy batch integrations to real-time event streaming and API-first approaches, reducing complexity and improving reliability.</li>
<li>Owned service delivery for integration services: SLA tracking, incident management, operational reporting — with ITIL-aligned practices for IT service operations.</li>
<li>Managed budgets and optimised costs across integration services; made data-driven decisions on build-vs-buy, cloud migration, and resource allocation.</li>
<li>Facilitated Agile delivery using OKRs and velocity metrics to track value; drove continuous improvement through retrospectives and process optimisation.</li>
<li>Built trusted relationships with stakeholders across Inter IKEA and Ingka — translating complex integration challenges into business language that non-technical leaders could act on.</li>
<li>Championed DevOps practices and automation: pipeline design, automated testing, deployment automation — improving delivery speed and reliability across the integration landscape.</li>
</ul>

<p class="role">Software Engineer &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Developed enterprise web applications and integration components (Java, .NET); built foundational understanding of service-oriented architecture and system integration.</li>
<li>Gained experience across full software lifecycle: requirements, design, development, testing, and production support in globally distributed delivery models.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>Google Cloud Associate Cloud Engineer</li>
<li>AWS Certified Cloud Practitioner</li>
<li>ISTQB Certified Tester</li>
<li>Six Sigma Green Belt (Process Improvement)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>PGDOM</b> – IGNOU &nbsp;|&nbsp; <b>B.Tech Information Technology</b> – UP Technical University</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Conversational) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

<h2>WHY THIS ROLE</h2>
<p class="why">Integration is the connective tissue of IKEA's digital landscape — and after 8+ years working across
different parts of the ecosystem, I've seen first-hand how powerful it is when systems share data seamlessly
through well-designed events and APIs. I'm passionate about treating integration services as a product:
measuring success by how much friction we remove for consuming teams, not just by uptime metrics.
This role brings together everything I care about — leading people, modernising platforms, simplifying
complexity, and enabling the many teams across IKEA to move faster. I understand how IKEA works as a
franchise system, I know where the integration pain points live, and I'm energised by the opportunity
to drive the shift towards a simpler, event-driven integration landscape.</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
