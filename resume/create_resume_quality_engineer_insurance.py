"""
Resume: Quality Engineer / Test Automation — Insurance/Financial Services domain
Scaled agile (Tribes, Squads, Big Room Planning), Java, Playwright, Selenium, CI/CD.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Quality_Engineer_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx+len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context):
    p = doc.add_paragraph()
    p.space_before = Pt(7)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malmö, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Quality Engineering professional with ", False),
        ("16+ years", True),
        (" designing and implementing ", False),
        ("automated tests across functional and non-functional areas", True),
        (" (API, UI, integration, regression, performance). Proven track record of ", False),
        ("building and evolving test frameworks", True),
        (" and reusable components that ", False),
        ("scale across squads", True),
        (" in ", False),
        ("scaled agile setups", True),
        (" (Tribes, Squads, Big Room Planning). Strong ", False),
        ("Java", True),
        (" and Python developer with hands-on experience in ", False),
        ("Playwright, Selenium", True),
        (", and modern CI/CD pipelines (", False),
        ("Azure DevOps, GitHub Actions, Jenkins", True),
        ("). Passionate about ", False),
        ("coaching teams", True),
        (" and establishing ", False),
        ("quality telemetry", True),
        (" that generates actionable insights. Background in ", False),
        ("financial services", True),
        (" (Core Banking) and complex enterprise domains.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Test Automation Architecture", "Java & Python (Strong)", "CI/CD Pipeline Design",
        "Framework Building & Scaling", "Quality Telemetry & Metrics", "Coaching & Enabling Teams",
        "API / UI / Integration Testing", "Scaled Agile (Squads/Tribes)", "Cross-Functional Collaboration",
        "Performance & Regression Testing", "Defect Trends & RCA", "Community of Practice",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills_data = [
        ("Languages:", " Java (strong), Python, TypeScript/JavaScript, SQL, Bash"),
        ("Test Frameworks:", " Playwright, Selenium, Appium, JUnit, TestNG, Pytest, Robot Framework, Fitnesse"),
        ("CI/CD:", " Azure DevOps, GitHub Actions, Jenkins, Docker, Kubernetes"),
        ("API & Performance:", " REST Assured, Postman, k6, JMeter, contract testing"),
        ("Cloud & Infra:", " GCP (Cloud Run, GKE, Pub/Sub), Azure, Docker, Kubernetes"),
        ("Quality Telemetry:", " Grafana, test coverage dashboards, defect trend analysis, release stability metrics"),
        ("Agile & Tools:", " Jira, Confluence, TestRail, SAFe/Scaled Agile, Big Room Planning"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB, Malmö — Team Lead (Acting) / SDET",
        "Mar 2022 – Present",
        "IKEA App & Customer Connect Platform — Omni-Channel, 30+ Markets | Scaled Agile (Tribes & Squads)")
    add_bullet(doc, "Design and implement automated tests across API, UI, integration, and regression for 5+ integrated platforms in a scaled agile setup (Tribes, Squads, Big Room Planning).",
        ["automated tests across API, UI, integration, and regression", "scaled agile setup", "Tribes, Squads, Big Room Planning"])
    add_bullet(doc, "Build and evolve test frameworks, templates, and reusable components (Java, Python, Playwright, Selenium) that scale across multiple squads.",
        ["Build and evolve test frameworks", "reusable components", "Java", "Playwright", "Selenium", "scale across multiple squads"])
    add_bullet(doc, "Provide hands-on support to squads and tech leads — strengthen quality engineering setup, code reviews, pair programming on test automation.",
        ["hands-on support to squads and tech leads", "quality engineering setup"])
    add_bullet(doc, "Establish quality telemetry: test coverage dashboards, defect trend analysis, release stability metrics — actionable insights for continuous improvement.",
        ["quality telemetry", "test coverage", "defect trend analysis", "release stability", "actionable insights"])
    add_bullet(doc, "Participate in and drive cross-squad QA Community of Practice — share patterns, align on standards, propagate best practices.",
        ["QA Community of Practice", "share patterns", "align on standards"])
    add_bullet(doc, "Coach and enable colleagues to take ownership of quality in their daily work — automation-first mindset across the organisation.",
        ["Coach and enable colleagues", "ownership of quality", "automation-first mindset"])
    add_bullet(doc, "CI/CD pipeline design and maintenance (GitHub Actions, Azure DevOps) — automated quality gates, security scanning, deployment verification.",
        ["CI/CD pipeline", "GitHub Actions", "Azure DevOps", "automated quality gates"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm — Release & Automation Engineer",
        "Sep 2021 – Feb 2022",
        "Communication Platform — 300M+ Users | Microservices Architecture")
    add_bullet(doc, "Managed release pipelines and CI/CD infrastructure (Jenkins) for microservices platform — reliable multi-environment deployments.",
        ["release pipelines", "CI/CD infrastructure", "Jenkins"])
    add_bullet(doc, "Integrated automated test suites into pipelines; quality gates ensuring release stability with fast feedback loops.",
        ["automated test suites", "quality gates", "release stability"])
    add_bullet(doc, "Coordinated releases across distributed teams; analysed defect trends, drove targeted coverage improvements.",
        ["distributed teams", "defect trends"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech — IKEA & LEGO Group, Denmark & Sweden — Technical Specialist / SDET Lead",
        "2013 – 2021",
        "E-Commerce, Mobile & Enterprise — Multi-Partner Delivery | Distributed Teams (Onshore/Offshore)")
    add_bullet(doc, "Architected end-to-end test automation frameworks (Java, Selenium, Cucumber, Python) scaling across 10+ integrated systems — 40% regression cycle reduction.",
        ["test automation frameworks", "Java", "Selenium", "Cucumber", "40% regression cycle reduction"])
    add_bullet(doc, "Built reusable test components and templates adopted across multiple delivery teams — standardised quality practices at scale.",
        ["reusable test components", "standardised quality practices at scale"])
    add_bullet(doc, "Drove test transformation: manual-first to automation-first across programmes — improved coverage, efficiency, and release stability.",
        ["test transformation", "automation-first", "release stability"])
    add_bullet(doc, "Established quality metrics and KPIs (defect density, automation coverage, cycle time); root cause analysis for systemic issues.",
        ["quality metrics and KPIs", "root cause analysis"])
    add_bullet(doc, "Managed distributed international teams (8–12 engineers, cross-functional); coached on automation practices and coding standards.",
        ["distributed international teams", "cross-functional", "coached"])
    add_bullet(doc, "CI/CD integration (Jenkins, GitHub Actions, SVN, Amazon EC2); performance testing for high-traffic e-commerce platforms.",
        ["CI/CD integration", "Jenkins", "performance testing"])
    add_bullet(doc, "Worked in scaled agile delivery model with cross-functional squads, sprint ceremonies, and planning across multiple locations.",
        ["scaled agile", "cross-functional squads", "multiple locations"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise — SDET / Consultant",
        "2008 – 2013",
        "Finacle CBS, Core Banking — Regulated Financial Services (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "Quality engineering for Finacle Core Banking System — accounts, transactions, loans, compliance in regulated financial services environment.",
        ["Finacle Core Banking", "regulated financial services"])
    add_bullet(doc, "Built automated test frameworks in Java and Selenium; data validation tools, regression suites for banking transaction flows.",
        ["Java", "Selenium", "regression suites", "banking transaction flows"])
    add_bullet(doc, "Morpho BAS 2FA: Integration testing for biometric authentication — security flows, device integration, transaction verification.",
        ["Integration testing", "biometric authentication", "security flows"])
    add_bullet(doc, "Experience with complex business domains requiring regulatory compliance, audit trails, and high reliability standards.",
        ["complex business domains", "regulatory compliance"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "CEH", "Six Sigma Green Belt"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
