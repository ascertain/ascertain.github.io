"""
Generate a tailored resume for Software Quality Engineer (Offer & Project Quality).
Focus: SW Quality Fundamentals, 8D/problem solving, Agile quality, OLM, defect analysis,
customer escalations, cross-functional quality leadership, V&V, CI/CD quality gates.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_SW_Quality_Engineer_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_SW_Quality_Engineer_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "8D", "5 Why", "Ishikawa", "root-cause analysis",
    "Software Quality", "SW Quality", "quality process",
    "Agile", "PI planning", "Scrum", "Kanban",
    "Verification and Validation", "V&V",
    "defect", "SQI", "quality metrics",
    "customer escalation", "Voice of Customer",
    "Offer Lifecycle", "OLM", "milestone",
    "cross-functional", "lessons learned",
    "risk identification", "risk mitigation",
    "CI/CD", "release management",
    "IEC62443", "cyber security",
    "Python", "TypeScript", "JavaScript", "C#", ".Net",
    "Early Warning System", "EWS",
    "Go/No go", "quality gates",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Software Quality Engineer — Offer & Project Quality", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Software Quality professional with 16+ years of experience in enterprise software development, "
        "testing, and quality assurance. Proven ability to deploy SW Quality Fundamentals, manage "
        "customer escalations using 8D methodology, and drive cross-functional quality improvements "
        "through problem-solving methods (5 Why, Ishikawa, statistical analysis). Experienced in "
        "embedding quality into Agile project lifecycles — supporting PI planning, defining quality "
        "goals for milestones, and preparing Go/No go decisions. Strong track record of managing "
        "defect analysis (SQI, dSQI), driving lessons learned processes, and implementing Early "
        "Warning Systems. Skilled at leading without direct authority, building cross-team "
        "relationships, and influencing stakeholders to achieve customer satisfaction and robustness "
        "targets. Hands-on programming background (Python, TypeScript, JavaScript, C#) combined "
        "with Software Verification and Validation expertise and Release Management in Agile environments."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "SW Quality Fundamentals & Processes", "8D / 5 Why / Ishikawa Problem Solving", "Customer Escalation Management",
        "Agile Quality (PI Planning, Sprints)", "Offer Lifecycle Management (OLM)", "SW Verification & Validation",
        "Defect Analysis (SQI/dSQI Metrics)", "Cross-Functional Quality Leadership", "Release Management & Quality Gates",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Quality & Test Lead — Enterprise SaaS Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Deploy SW Quality Fundamentals across a customer-facing platform serving 30+ markets — defining quality processes, standards, and best practices for the engineering organization.",
        "Manage customer escalations using structured 8D methodology — driving root-cause analysis, containment actions, and permanent corrective measures with cross-functional teams.",
        "Drive regular defect analysis and SW performance monitoring through quality metrics (SQI, escape rate, defect density) — reporting trends and improvement actions to stakeholders.",
        "Ensure quality is built into each project by embedding quality goals in Agile ceremonies (PI planning, sprint reviews) and defining entry/exit criteria for each milestone.",
        "Capture and analyze Voice of Customer feedback from field quality data — feeding insights into offer development for continuous improvement and robustness.",
        "Implement lessons learned processes: synthesize field performance data, drive feedback loops to impacted stakeholders, and update quality best practices accordingly.",
        "Define Early Warning Systems (EWS) for new releases — collaborating with Tech Support, engineering, and product management to detect quality degradation early.",
        "Prepare quality summaries for milestone reviews and propose Go/No go decisions based on quality targets, risk assessment, and customer impact analysis.",
        "Support risk identification, assessment, and mitigation planning within projects — assisting project managers in maintaining updated risk registers.",
        "Lead quality improvements without direct reports — influencing cross-functional teams (R&D, Tech Support, Product) through quality authority and strong relationships.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Release Quality & Automation Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed release quality gates for a 300M+ user platform — enforcing quality criteria, conducting defect triage, and ensuring release readiness through structured Go/No go decisions.",
        "Drove problem-solving during production incidents: applied 8D and root-cause analysis methods to resolve critical defects and prevent recurrence.",
        "Implemented quality metrics dashboards (defect leakage, regression pass rate, field escape rate) to support data-driven quality decisions.",
        "Coordinated SW Verification and Validation activities across microservices — ensuring integration quality and performance benchmarks were met before each release.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior Quality Analyst / Test Lead — Enterprise Software", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Deployed quality processes across enterprise software projects (e-commerce, PLM, IoT platforms) — ensuring SW Quality Fundamentals were embedded in development lifecycle.",
        "Managed customer escalations and critical defects using 8D, 5 Why, and Ishikawa analysis — driving cross-functional resolution teams and implementing corrective actions.",
        "Defined quality goals during project lifecycle: test coverage targets, defect density thresholds, and customer satisfaction metrics aligned with Offer Lifecycle Management.",
        "Supported Agile transformation by integrating quality activities into PI planning, sprint ceremonies, and continuous integration pipelines — including quality gates and automated checks.",
        "Conducted Verification and Validation for inhouse and brand-labeled software across IKEA (30+ markets) and LEGO (global digital platform).",
        "Led risk identification and mitigation within projects — maintaining risk registers, conducting FMEA-style assessments, and reporting to project leadership.",
        "Drove lessons learned sessions after each major release — synthesizing field data, documenting root causes, and implementing process improvements.",
        "Built and maintained automated test frameworks (Selenium, API automation, Python, JavaScript) to improve regression quality and reduce manual verification effort by 40%.",
        "Managed defect lifecycle using Jira and HP ALM — triage, prioritisation, trend analysis, and stakeholder reporting through quality dashboards.",
        "Operated in regulated environments with strict quality standards — applying ISTQB methodologies and ensuring compliance with engineering best practices.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise IT & Banking", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Quality Analyst / Technical Lead", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led SW quality assurance for core banking implementations — defining test strategies, managing defect resolution, and ensuring Verification & Validation in regulated environments.",
        "Applied problem-solving methods (8D, 5 Why) to production incidents impacting financial transactions — driving permanent corrective actions across cross-functional teams.",
        "Managed quality processes for multi-phase rollouts: milestone reviews, Go/No go decisions, and post-deployment quality monitoring across 20+ branch deployments.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("Languages: ", "Python, TypeScript, JavaScript, C#, .Net, Java, HTML, XML, SQL"),
        ("Quality Tools: ", "Jira, Azure DevOps, HP ALM (Quality Center), Confluence, Zephyr, SonarQube"),
        ("Automation: ", "Selenium, Playwright, Pytest, API automation, CI/CD quality gates (GitHub Actions, Jenkins)"),
        ("Methodologies: ", "8D, 5 Why, Ishikawa, FMEA, ISTQB, V-Model, Agile (Scrum/Kanban/SAFe), Waterfall"),
        ("Quality Metrics: ", "SQI, dSQI, defect density, escape rate, regression pass rate, customer satisfaction KPIs"),
        ("Cloud & DevOps: ", "GCP, AWS, Docker, Terraform, CI/CD pipelines, release management"),
        ("Standards: ", "IEC62443 (awareness), ISTQB, ITIL, ISO 9001 quality management principles"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PG Diploma, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Six Sigma Green Belt",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "ITIL v4 Foundation",
        "Certified Ethical Hacker (CEH)",
        "AWS Certified Cloud Practitioner",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Software Quality Engineer &mdash; Offer &amp; Project Quality</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Software Quality professional with 16+ years of experience in enterprise software development, testing, and quality assurance. Proven ability to deploy SW Quality Fundamentals, manage customer escalations using 8D methodology, and drive cross-functional quality improvements through problem-solving methods (5 Why, Ishikawa, statistical analysis). Experienced in embedding quality into Agile project lifecycles &mdash; supporting PI planning, defining quality goals for milestones, and preparing Go/No go decisions. Strong track record of managing defect analysis (SQI, dSQI), driving lessons learned processes, and implementing Early Warning Systems. Skilled at leading without direct authority, building cross-team relationships, and influencing stakeholders to achieve customer satisfaction and robustness targets. Hands-on programming background (Python, TypeScript, JavaScript, C#) combined with Software Verification and Validation expertise and Release Management in Agile environments.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; SW Quality Fundamentals &amp; Processes</td><td>&bull; 8D / 5 Why / Ishikawa Problem Solving</td><td>&bull; Customer Escalation Management</td></tr>
<tr><td>&bull; Agile Quality (PI Planning, Sprints)</td><td>&bull; Offer Lifecycle Management (OLM)</td><td>&bull; SW Verification &amp; Validation</td></tr>
<tr><td>&bull; Defect Analysis (SQI/dSQI Metrics)</td><td>&bull; Cross-Functional Quality Leadership</td><td>&bull; Release Management &amp; Quality Gates</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Quality &amp; Test Lead &mdash; Enterprise SaaS Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Deploy SW Quality Fundamentals across a customer-facing platform serving 30+ markets &mdash; defining quality processes, standards, and best practices for the engineering organization.</li>
<li>Manage customer escalations using structured 8D methodology &mdash; driving root-cause analysis, containment actions, and permanent corrective measures with cross-functional teams.</li>
<li>Drive regular defect analysis and SW performance monitoring through quality metrics (SQI, escape rate, defect density) &mdash; reporting trends and improvement actions to stakeholders.</li>
<li>Ensure quality is built into each project by embedding quality goals in Agile ceremonies (PI planning, sprint reviews) and defining entry/exit criteria for each milestone.</li>
<li>Capture and analyze Voice of Customer feedback from field quality data &mdash; feeding insights into offer development for continuous improvement and robustness.</li>
<li>Implement lessons learned processes: synthesize field performance data, drive feedback loops to impacted stakeholders, and update quality best practices accordingly.</li>
<li>Define Early Warning Systems (EWS) for new releases &mdash; collaborating with Tech Support, engineering, and product management to detect quality degradation early.</li>
<li>Prepare quality summaries for milestone reviews and propose Go/No go decisions based on quality targets, risk assessment, and customer impact analysis.</li>
<li>Support risk identification, assessment, and mitigation planning within projects &mdash; assisting project managers in maintaining updated risk registers.</li>
<li>Lead quality improvements without direct reports &mdash; influencing cross-functional teams (R&amp;D, Tech Support, Product) through quality authority and strong relationships.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Release Quality &amp; Automation Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Managed release quality gates for a 300M+ user platform &mdash; enforcing quality criteria, conducting defect triage, and ensuring release readiness through structured Go/No go decisions.</li>
<li>Drove problem-solving during production incidents: applied 8D and root-cause analysis methods to resolve critical defects and prevent recurrence.</li>
<li>Implemented quality metrics dashboards (defect leakage, regression pass rate, field escape rate) to support data-driven quality decisions.</li>
<li>Coordinated SW Verification and Validation activities across microservices &mdash; ensuring integration quality and performance benchmarks were met before each release.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Senior Quality Analyst / Test Lead &mdash; Enterprise Software <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Deployed quality processes across enterprise software projects (e-commerce, PLM, IoT platforms) &mdash; ensuring SW Quality Fundamentals were embedded in development lifecycle.</li>
<li>Managed customer escalations and critical defects using 8D, 5 Why, and Ishikawa analysis &mdash; driving cross-functional resolution teams and implementing corrective actions.</li>
<li>Defined quality goals during project lifecycle: test coverage targets, defect density thresholds, and customer satisfaction metrics aligned with Offer Lifecycle Management.</li>
<li>Supported Agile transformation by integrating quality activities into PI planning, sprint ceremonies, and continuous integration pipelines &mdash; including quality gates and automated checks.</li>
<li>Conducted Verification and Validation for inhouse and brand-labeled software across IKEA (30+ markets) and LEGO (global digital platform).</li>
<li>Led risk identification and mitigation within projects &mdash; maintaining risk registers, conducting FMEA-style assessments, and reporting to project leadership.</li>
<li>Drove lessons learned sessions after each major release &mdash; synthesizing field data, documenting root causes, and implementing process improvements.</li>
<li>Built and maintained automated test frameworks (Selenium, API automation, Python, JavaScript) to improve regression quality and reduce manual verification effort by 40%.</li>
<li>Managed defect lifecycle using Jira and HP ALM &mdash; triage, prioritisation, trend analysis, and stakeholder reporting through quality dashboards.</li>
<li>Operated in regulated environments with strict quality standards &mdash; applying ISTQB methodologies and ensuring compliance with engineering best practices.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise IT &amp; Banking</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Led SW quality assurance for core banking implementations &mdash; defining test strategies, managing defect resolution, and ensuring Verification &amp; Validation in regulated environments.</li>
<li>Applied problem-solving methods (8D, 5 Why) to production incidents impacting financial transactions &mdash; driving permanent corrective actions across cross-functional teams.</li>
<li>Managed quality processes for multi-phase rollouts: milestone reviews, Go/No go decisions, and post-deployment quality monitoring across 20+ branch deployments.</li>
</ul>

<h2>TECHNICAL SKILLS</h2>
<p class="tech-line"><b>Languages:</b> Python, TypeScript, JavaScript, C#, .Net, Java, HTML, XML, SQL</p>
<p class="tech-line"><b>Quality Tools:</b> Jira, Azure DevOps, HP ALM (Quality Center), Confluence, Zephyr, SonarQube</p>
<p class="tech-line"><b>Automation:</b> Selenium, Playwright, Pytest, API automation, CI/CD quality gates (GitHub Actions, Jenkins)</p>
<p class="tech-line"><b>Methodologies:</b> 8D, 5 Why, Ishikawa, FMEA, ISTQB, V-Model, Agile (Scrum/Kanban/SAFe), Waterfall</p>
<p class="tech-line"><b>Quality Metrics:</b> SQI, dSQI, defect density, escape rate, regression pass rate, customer satisfaction KPIs</p>
<p class="tech-line"><b>Cloud &amp; DevOps:</b> GCP, AWS, Docker, Terraform, CI/CD pipelines, release management</p>
<p class="tech-line"><b>Standards:</b> IEC62443 (awareness), ISTQB, ITIL, ISO 9001 quality management principles</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Six Sigma Green Belt</td></tr>
<tr><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; Certified Ethical Hacker (CEH)</td><td>&bull; AWS Certified Cloud Practitioner</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
