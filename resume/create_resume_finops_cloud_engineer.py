"""Generate FinOps Cloud Cost Management Engineer resume – LEGO Group style role."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib, html

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_FinOps_Cloud_Engineer_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Cloud & Data Engineer with hands-on experience across AWS, Azure, and GCP, "
        "specializing in building cost-optimized data pipelines, ETL solutions, and automation "
        "using Python and cloud-native services. Strong intersection of engineering and finance—"
        "proven ability to analyze cloud spend, drive rightsizing recommendations, and build "
        "unified cost visibility platforms. Experienced in PowerBI reporting, chatbot development, "
        "and deploying scalable solutions that translate infrastructure data into actionable "
        "business insights. Passionate about FinOps culture and continuous cost optimization."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Technologies")
    skills_data = [
        ("Cloud Platforms", "AWS (EC2, S3, Lambda, EKS, CloudWatch, Cost Explorer, RDS, Glue) · Azure (ADF, Functions, DevOps, Cost Management) · GCP (BigQuery, Cloud Functions, GKE, Cloud Run)"),
        ("Data & ETL", "Python · Pandas · SQL · BigQuery · Airflow · dbt · ETL/ELT pipelines · REST API integrations · Data Lakes"),
        ("FinOps & Cost", "Cloud cost analysis · Rightsizing · Reserved Instances · Savings Plans · Budget forecasting · Tagging strategies · Showback/Chargeback"),
        ("Visualization", "PowerBI (DAX, reports, dashboards) · Grafana · Looker Studio"),
        ("Automation & Apps", "PowerApps · Power Automate · Chatbots (Dialogflow, custom Python bots) · REST APIs · FastAPI · Flask"),
        ("DevOps & IaC", "Terraform · Docker · Kubernetes · GitHub Actions · CI/CD · Helm · ArgoCD"),
        ("Certifications", "AWS Cloud Practitioner · Azure Fundamentals (AZ-900) · Google Cloud ACE · FinOps Certified Practitioner · ISTQB CTFL"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.2)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "F2F7FC")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA
    add_role(doc, "Cloud & Data Engineer", "IKEA (Ingka Digital)", "Malmö, Sweden", "Mar 2022 – Present")
    bullets_1 = [
        "Designed and maintained ETL pipelines ingesting data from 20+ SaaS and on-prem sources into BigQuery, processing 5TB+ daily.",
        "Built unified cloud cost dashboards in PowerBI, providing cost visibility across AWS and GCP for 50+ engineering teams.",
        "Developed Python automation for rightsizing recommendations—identified $400K/year savings through underutilized compute analysis.",
        "Created PowerApps-based self-service tooling for teams to tag resources, request budgets, and view cost anomalies.",
        "Built and deployed chatbot (Dialogflow + Cloud Functions) enabling engineers to query cost data and receive alerts conversationally.",
        "Drove FinOps practices: established tagging governance, showback reports, and monthly cost review cadence with stakeholders.",
        "Deployed data pipelines using Airflow + Python on GKE, extending coverage to new data clouds and scopes.",
        "Collaborated with finance stakeholders to build budget forecasting models and chargeback allocation across business units.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Cloud & DevOps Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Built and maintained cloud infrastructure (AWS/GCP) supporting 300M+ users globally, optimizing for cost and scalability.",
        "Developed cost monitoring and alerting using AWS Cost Explorer APIs and custom Python scripts for budget threshold breaches.",
        "Automated infrastructure provisioning with Terraform, reducing deployment time by 60% and ensuring cost-efficient resource allocation.",
        "Created PowerBI reports for infrastructure cost metrics and SLA compliance visibility for engineering leadership.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Data & Platform Engineer", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Built and operated ETL pipelines (Python) on AWS and GCP for IoT and enterprise data—processing multi-TB datasets from distributed sources.",
        "Designed PowerBI reports surfacing infrastructure cost metrics, resource utilization, and optimization opportunities for leadership.",
        "Implemented infrastructure-as-code (Terraform) for AWS and GCP environments, reducing provisioning time by 70%.",
        "Developed cost monitoring automation using Cloud APIs, alerting on budget threshold breaches and anomalous spend patterns.",
        "Automated CI/CD pipelines (Jenkins, GitHub Actions) for data pipeline deployment across dev/staging/prod environments.",
        "Built Power Automate flows for automated cost anomaly notifications to cost center owners.",
        "Managed cloud infrastructure (EC2, S3, RDS, Lambda, GCE, GCS) for enterprise clients, optimizing for cost and performance.",
        "Developed Python-based automation tools for resource scheduling, auto-scaling policies, and unused resource cleanup—saving 25% monthly spend.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software & Infrastructure Engineer", "Earlier Career (Wipro, CSC)", "India", "2008 – 2013")
    bullets_4 = [
        "Developed data ingestion scripts pulling from REST APIs, databases, and flat files into centralized data stores.",
        "Created monitoring dashboards (Grafana, CloudWatch) providing real-time visibility into infrastructure health and cost trends.",
        "Built Python-based automation for report generation, data extraction, and infrastructure management tasks.",
        "Supported migration of on-prem workloads to AWS, including cost estimation and architecture recommendations.",
        "Managed enterprise application infrastructure, driving cost optimization through rightsizing and resource consolidation.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── Projects ──────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Projects")
    projects = [
        "Unified Cost Platform: End-to-end cost management solution integrating AWS, GCP, and Azure billing data into a single PowerBI workspace with drill-down by team, service, and environment.",
        "FinOps Chatbot: Dialogflow-based chatbot enabling non-technical stakeholders to query cost trends, get optimization suggestions, and submit budget requests via natural language.",
        "Data Pipeline Scaling: Extended ingestion framework from 5 to 25+ data sources (Salesforce, ServiceNow, Jira, custom APIs) using modular Python connectors with Airflow orchestration.",
        "PowerApps Cost Portal: Self-service application for engineering teams to view allocated costs, submit tagging corrections, and track savings initiatives.",
    ]
    for proj in projects:
        bullet(doc, proj, bold_prefix=proj.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    h = html.escape
    content = f"""<html><head><meta charset="utf-8">
<style>
body{{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}}
h1{{text-align:center;color:#1F4E79;font-size:18pt;margin-bottom:2px}}
.contact{{text-align:center;font-size:9.5pt;margin-bottom:10px}}
h2{{color:#1F4E79;font-size:10.5pt;border-bottom:1px solid #1F4E79;padding-bottom:2px;margin-top:12px}}
.role{{font-weight:bold;margin-top:8px}} .meta{{color:#444;font-size:9.5pt}}
ul{{margin:2px 0 4px 18px;padding:0}} li{{margin-bottom:2px}}
table{{width:100%;border-collapse:collapse;font-size:9.5pt}} td{{padding:2px 6px;vertical-align:top}}
.cat{{background:#F2F7FC;font-weight:bold;width:20%}}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Cloud &amp; Data Engineer with hands-on experience across AWS, Azure, and GCP, specializing in building cost-optimized data pipelines, ETL solutions, and automation using Python and cloud-native services. Strong intersection of engineering and finance—proven ability to analyze cloud spend, drive rightsizing recommendations, and build unified cost visibility platforms. Experienced in PowerBI reporting, chatbot development, and deploying scalable solutions that translate infrastructure data into actionable business insights. Passionate about FinOps culture and continuous cost optimization.</p>

<h2>KEY SKILLS &amp; TECHNOLOGIES</h2>
<table>
<tr><td class="cat">Cloud Platforms</td><td>AWS (EC2, S3, Lambda, EKS, CloudWatch, Cost Explorer, RDS, Glue) · Azure (ADF, Functions, DevOps, Cost Management) · GCP (BigQuery, Cloud Functions, GKE, Cloud Run)</td></tr>
<tr><td class="cat">Data &amp; ETL</td><td>Python · Pandas · SQL · BigQuery · Airflow · dbt · ETL/ELT pipelines · REST API integrations · Data Lakes</td></tr>
<tr><td class="cat">FinOps &amp; Cost</td><td>Cloud cost analysis · Rightsizing · Reserved Instances · Savings Plans · Budget forecasting · Tagging strategies · Showback/Chargeback</td></tr>
<tr><td class="cat">Visualization</td><td>PowerBI (DAX, reports, dashboards) · Grafana · Looker Studio</td></tr>
<tr><td class="cat">Automation &amp; Apps</td><td>PowerApps · Power Automate · Chatbots (Dialogflow, custom Python bots) · REST APIs · FastAPI · Flask</td></tr>
<tr><td class="cat">DevOps &amp; IaC</td><td>Terraform · Docker · Kubernetes · GitHub Actions · CI/CD · Helm · ArgoCD</td></tr>
<tr><td class="cat">Certifications</td><td>AWS Cloud Practitioner · Azure Fundamentals (AZ-900) · Google Cloud ACE · FinOps Certified Practitioner · ISTQB CTFL</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Cloud &amp; Data Engineer <span class="meta">&nbsp;|&nbsp; IKEA (Ingka Digital) &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</span></p>
<ul>
<li>Designed and maintained ETL pipelines ingesting data from 20+ SaaS and on-prem sources into BigQuery, processing 5TB+ daily.</li>
<li>Built unified cloud cost dashboards in PowerBI, providing cost visibility across AWS and GCP for 50+ engineering teams.</li>
<li>Developed Python automation for rightsizing recommendations—identified $400K/year savings through underutilized compute analysis.</li>
<li>Created PowerApps-based self-service tooling for teams to tag resources, request budgets, and view cost anomalies.</li>
<li>Built and deployed chatbot (Dialogflow + Cloud Functions) enabling engineers to query cost data and receive alerts conversationally.</li>
<li>Drove FinOps practices: established tagging governance, showback reports, and monthly cost review cadence with stakeholders.</li>
<li>Deployed data pipelines using Airflow + Python on GKE, extending coverage to new data clouds and scopes.</li>
<li>Collaborated with finance stakeholders to build budget forecasting models and chargeback allocation across business units.</li>
</ul>

<p class="role">Cloud &amp; DevOps Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Built and maintained cloud infrastructure (AWS/GCP) supporting 300M+ users globally, optimizing for cost and scalability.</li>
<li>Developed cost monitoring and alerting using AWS Cost Explorer APIs and custom Python scripts for budget threshold breaches.</li>
<li>Automated infrastructure provisioning with Terraform, reducing deployment time by 60% and ensuring cost-efficient resource allocation.</li>
<li>Created PowerBI reports for infrastructure cost metrics and SLA compliance visibility for engineering leadership.</li>
</ul>

<p class="role">Data &amp; Platform Engineer <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Built and operated ETL pipelines (Python) on AWS and GCP for IoT and enterprise data—processing multi-TB datasets from distributed sources.</li>
<li>Designed PowerBI reports surfacing infrastructure cost metrics, resource utilization, and optimization opportunities for leadership.</li>
<li>Implemented infrastructure-as-code (Terraform) for AWS and GCP environments, reducing provisioning time by 70%.</li>
<li>Developed cost monitoring automation using Cloud APIs, alerting on budget threshold breaches and anomalous spend patterns.</li>
<li>Automated CI/CD pipelines (Jenkins, GitHub Actions) for data pipeline deployment across dev/staging/prod environments.</li>
<li>Built Power Automate flows for automated cost anomaly notifications to cost center owners.</li>
<li>Managed cloud infrastructure (EC2, S3, RDS, Lambda, GCE, GCS) for enterprise clients, optimizing for cost and performance.</li>
<li>Developed Python-based automation tools for resource scheduling, auto-scaling policies, and unused resource cleanup—saving 25% monthly spend.</li>
</ul>

<p class="role">Software &amp; Infrastructure Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (Wipro, CSC) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Developed data ingestion scripts pulling from REST APIs, databases, and flat files into centralized data stores.</li>
<li>Created monitoring dashboards (Grafana, CloudWatch) providing real-time visibility into infrastructure health and cost trends.</li>
<li>Built Python-based automation for report generation, data extraction, and infrastructure management tasks.</li>
<li>Supported migration of on-prem workloads to AWS, including cost estimation and architecture recommendations.</li>
<li>Managed enterprise application infrastructure, driving cost optimization through rightsizing and resource consolidation.</li>
</ul>

<h2>KEY PROJECTS</h2>
<ul>
<li><strong>Unified Cost Platform:</strong> End-to-end cost management solution integrating AWS, GCP, and Azure billing data into a single PowerBI workspace with drill-down by team, service, and environment.</li>
<li><strong>FinOps Chatbot:</strong> Dialogflow-based chatbot enabling non-technical stakeholders to query cost trends, get optimization suggestions, and submit budget requests via natural language.</li>
<li><strong>Data Pipeline Scaling:</strong> Extended ingestion framework from 5 to 25+ data sources (Salesforce, ServiceNow, Jira, custom APIs) using modular Python connectors with Airflow orchestration.</li>
<li><strong>PowerApps Cost Portal:</strong> Self-service application for engineering teams to view allocated costs, submit tagging corrections, and track savings initiatives.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
