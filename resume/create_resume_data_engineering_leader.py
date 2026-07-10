"""Generate Data Engineering Leader resume – Inter IKEA Internal Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Data_Engineering_Leader_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Data Engineering Leader with 10+ years of experience designing, building, and operating "
        "data pipelines, data models, and data platform solutions at scale within the IKEA ecosystem. "
        "Combines hands-on technical leadership (Python, SQL, PySpark, BigQuery, cloud-native services) "
        "with people leadership — leading and developing engineering teams, mentoring data engineers, "
        "and fostering a culture of technical excellence and continuous learning. "
        "Deep IKEA domain knowledge with proven ability to collaborate across product teams, business "
        "stakeholders, and the broader IKEA Data Engineering community. "
        "Experienced in establishing best practices for data modelling, data quality, pipeline architecture, "
        "and self-service data platforms. Passionate about data sharing, engineering excellence, "
        "and enabling data-driven decision-making across the organization. "
        "Strong believer in IKEA values — togetherness, simplicity, cost-consciousness, and leading by example."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Technologies")
    skills_data = [
        ("Data Engineering", "Python · SQL · PySpark · Data modelling (star schema, Data Vault) · ETL/ELT pipelines · dbt · Airflow · Data quality frameworks · Metadata management"),
        ("Data Platforms", "BigQuery · Databricks (familiar) · Microsoft Fabric (familiar) · GCP (Cloud Functions, Dataflow, Pub/Sub, GCS) · AWS (Glue, S3, Athena) · Data Lakes · Lakehouse architecture"),
        ("Leadership & People", "Team leadership (direct & indirect) · Mentoring engineers · Code reviews · Knowledge sharing · Community of practice · Performance development · Agile/Scrum"),
        ("Architecture & Standards", "Data pipeline architecture · Semantic layers · Reusable frameworks · Data governance · Version control · Testing · Monitoring · Documentation"),
        ("DevOps & Tooling", "Terraform · Docker · Kubernetes · GitHub Actions · CI/CD · Git · Jira · Confluence · Ingka DevOps Tooling"),
        ("Stakeholder Management", "Business requirement translation · Cross-functional collaboration · Product team alignment · IKEA ways of working"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS
    add_role(doc, "Data Engineering Lead / Team Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Lead, coach, and develop a team of engineers — fostering a culture of technical excellence, collaboration, continuous learning, and knowledge sharing within the IKEA Data Engineering community.",
        "Designed, developed, and maintained enterprise-grade data pipelines (Python, SQL, BigQuery) ingesting data from 20+ SaaS and on-prem sources — processing large-scale datasets for analytics and self-service across the Customer Support domain.",
        "Provided strategic technical leadership for data modelling, pipeline architecture, and data platform design — ensuring scalable, well-governed, and future-proof data solutions aligned with IKEA's data strategy.",
        "Act as a central point of contact within the IKEA data engineering space — driving consistency, best practices, and continuous improvement of methodologies and tools across teams.",
        "Collaborated closely with business stakeholders, product teams, analysts, and architects to translate complex business requirements into performant, scalable data solutions serving 32 markets.",
        "Established and enforced data engineering best practices — data modelling standards, data quality checks, version control (Git), automated testing, monitoring, documentation, and metadata management.",
        "Built reusable frameworks and semantic layers accelerating data product delivery — enabling self-service analytics and data democratization for business users.",
        "Optimised data pipelines for performance, reliability, scalability, and cost-efficiency — reducing processing costs by 40% through query optimization and pipeline restructuring.",
        "Ensured data governance, privacy, and compliance standards (GDPR) — data stored, processed, and accessed securely across all pipelines.",
        "Promoted data sharing culture and engineering excellence — mentoring less experienced co-workers through code reviews, pair programming, and knowledge sharing forums.",
        "Drove adoption of modern data engineering practices (dbt, Airflow, CI/CD for data pipelines) within the team and across adjacent IKEA teams.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Data & Platform Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Designed and maintained data pipelines processing high-volume event data for a globally distributed application (300M+ users).",
        "Built data models and reporting infrastructure enabling data-driven decision-making for product and engineering teams.",
        "Collaborated cross-functionally with product, engineering, and data science teams in an agile environment.",
        "Drove CI/CD automation for data pipeline deployments — reducing deployment time by 50%.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Technical Lead — Data Engineering", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led and developed a team of 8–12 engineers — providing technical direction, mentoring, conducting code reviews, and fostering continuous learning and innovation.",
        "Designed and implemented data models and robust data pipelines (Python, SQL, PySpark) for IoT telemetry, supply chain, and enterprise data across IKEA and LEGO platforms.",
        "Built enterprise-grade ETL/ELT solutions processing multi-TB datasets from distributed sources — ensuring data quality, reliability, and performance at scale.",
        "Collaborated with IKEA business stakeholders, architects, and product teams to translate complex requirements into scalable data solutions.",
        "Established data engineering best practices across teams — data modelling standards, testing frameworks, version control, monitoring, and documentation.",
        "Built reusable data pipeline components and frameworks accelerating delivery across multiple IKEA product teams.",
        "Drove knowledge sharing and community building — regular forums, documentation, and mentoring that built data engineering capability across the organization.",
        "Worked within the IKEA ecosystem (IKEA App, Startcus Spare Parts, Genesys, Verint) — deep understanding of IKEA's data landscape and ways of working.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software & Data Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_4 = [
        "Built data integration pipelines and ETL solutions for enterprise financial and retail domains using Python, SQL, and Java.",
        "Developed data models and automated reporting solutions supporting business analytics and decision-making.",
        "Worked with large datasets across relational databases, flat files, and APIs — building scalable data ingestion frameworks.",
        "Participated in migration from legacy data systems to modern architectures.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── IKEA Ecosystem Experience ─────────────────────────────────────────
    add_heading_block(doc, "IKEA Ecosystem Experience")
    ikea_items = [
        "IKEA Customer Connect (VCS): Built end-to-end data layer — ETL pipelines, BigQuery data models, self-service analytics serving 32 markets within the Customer Support domain.",
        "IKEA Data Engineering Community: Active contributor — driving consistency, knowledge sharing, and best practices across the broader IKEA data engineering landscape.",
        "IKEA Ways of Working: Product team model, Agile/Scrum, Ingka DevOps tooling, cross-functional collaboration with business stakeholders, architects, and analysts.",
        "IKEA App & Connected Products: Data pipelines for IoT telemetry and smart-home product analytics (DIRIGERA/TRÅDFRI ecosystem).",
        "CSSP (Customer Support Staff Planning): Data integration with Verint and Genesys — workforce analytics across global contact centers.",
        "Startcus Spare Parts: Data engineering for logistics and fulfillment tracking systems.",
        "Data Governance & Compliance: GDPR-compliant data handling, metadata management, and security standards within IKEA's governance framework.",
    ]
    for item in ikea_items:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Togetherness: Build strong cross-functional relationships — collaborate with business stakeholders, data scientists, analysts, and product teams to deliver value together.",
        "Leading by Example: Mentor and develop co-workers through code reviews, knowledge sharing, and fostering a culture of engineering excellence and continuous learning.",
        "Simplicity: Design pragmatic, maintainable data solutions — reusable frameworks and clear standards that teams can adopt easily.",
        "Cost-Consciousness: Optimise data pipelines for cost-efficiency — reducing cloud spend while maintaining performance and reliability.",
        "Daring to be Different: Embrace modern data engineering technologies (Lakehouse, semantic layers, dbt) and apply them where they create clear business value.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Data Engineering Leader with 10+ years of experience designing, building, and operating data pipelines, data models, and data platform solutions at scale within the IKEA ecosystem. Combines hands-on technical leadership (Python, SQL, PySpark, BigQuery, cloud-native services) with people leadership — leading and developing engineering teams, mentoring data engineers, and fostering a culture of technical excellence and continuous learning. Deep IKEA domain knowledge with proven ability to collaborate across product teams, business stakeholders, and the broader IKEA Data Engineering community. Experienced in establishing best practices for data modelling, data quality, pipeline architecture, and self-service data platforms. Passionate about data sharing, engineering excellence, and enabling data-driven decision-making across the organization. Strong believer in IKEA values — togetherness, simplicity, cost-consciousness, and leading by example.</p>

<h2>KEY SKILLS &amp; TECHNOLOGIES</h2>
<table>
<tr><td class="cat">Data Engineering</td><td>Python · SQL · PySpark · Data modelling (star schema, Data Vault) · ETL/ELT pipelines · dbt · Airflow · Data quality frameworks · Metadata management</td></tr>
<tr><td class="cat">Data Platforms</td><td>BigQuery · Databricks (familiar) · Microsoft Fabric (familiar) · GCP (Cloud Functions, Dataflow, Pub/Sub, GCS) · AWS (Glue, S3, Athena) · Data Lakes · Lakehouse architecture</td></tr>
<tr><td class="cat">Leadership &amp; People</td><td>Team leadership (direct &amp; indirect) · Mentoring engineers · Code reviews · Knowledge sharing · Community of practice · Performance development · Agile/Scrum</td></tr>
<tr><td class="cat">Architecture &amp; Standards</td><td>Data pipeline architecture · Semantic layers · Reusable frameworks · Data governance · Version control · Testing · Monitoring · Documentation</td></tr>
<tr><td class="cat">DevOps &amp; Tooling</td><td>Terraform · Docker · Kubernetes · GitHub Actions · CI/CD · Git · Jira · Confluence · Ingka DevOps Tooling</td></tr>
<tr><td class="cat">Stakeholder Management</td><td>Business requirement translation · Cross-functional collaboration · Product team alignment · IKEA ways of working</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Data Engineering Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Lead, coach, and develop a team of engineers — fostering a culture of technical excellence, collaboration, continuous learning, and knowledge sharing within the IKEA Data Engineering community.</li>
<li>Designed, developed, and maintained enterprise-grade data pipelines (Python, SQL, BigQuery) ingesting data from 20+ SaaS and on-prem sources — processing large-scale datasets for analytics and self-service across the Customer Support domain.</li>
<li>Provided strategic technical leadership for data modelling, pipeline architecture, and data platform design — ensuring scalable, well-governed, and future-proof data solutions aligned with IKEA's data strategy.</li>
<li>Act as a central point of contact within the IKEA data engineering space — driving consistency, best practices, and continuous improvement of methodologies and tools across teams.</li>
<li>Collaborated closely with business stakeholders, product teams, analysts, and architects to translate complex business requirements into performant, scalable data solutions serving 32 markets.</li>
<li>Established and enforced data engineering best practices — data modelling standards, data quality checks, version control (Git), automated testing, monitoring, documentation, and metadata management.</li>
<li>Built reusable frameworks and semantic layers accelerating data product delivery — enabling self-service analytics and data democratization for business users.</li>
<li>Optimised data pipelines for performance, reliability, scalability, and cost-efficiency — reducing processing costs by 40% through query optimization and pipeline restructuring.</li>
<li>Ensured data governance, privacy, and compliance standards (GDPR) — data stored, processed, and accessed securely across all pipelines.</li>
<li>Promoted data sharing culture and engineering excellence — mentoring less experienced co-workers through code reviews, pair programming, and knowledge sharing forums.</li>
<li>Drove adoption of modern data engineering practices (dbt, Airflow, CI/CD for data pipelines) within the team and across adjacent IKEA teams.</li>
</ul>

<p class="role">Data &amp; Platform Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Designed and maintained data pipelines processing high-volume event data for a globally distributed application (300M+ users).</li>
<li>Built data models and reporting infrastructure enabling data-driven decision-making for product and engineering teams.</li>
<li>Collaborated cross-functionally with product, engineering, and data science teams in an agile environment.</li>
<li>Drove CI/CD automation for data pipeline deployments — reducing deployment time by 50%.</li>
</ul>

<p class="role">Technical Lead — Data Engineering <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led and developed a team of 8–12 engineers — providing technical direction, mentoring, conducting code reviews, and fostering continuous learning and innovation.</li>
<li>Designed and implemented data models and robust data pipelines (Python, SQL, PySpark) for IoT telemetry, supply chain, and enterprise data across IKEA and LEGO platforms.</li>
<li>Built enterprise-grade ETL/ELT solutions processing multi-TB datasets from distributed sources — ensuring data quality, reliability, and performance at scale.</li>
<li>Collaborated with IKEA business stakeholders, architects, and product teams to translate complex requirements into scalable data solutions.</li>
<li>Established data engineering best practices across teams — data modelling standards, testing frameworks, version control, monitoring, and documentation.</li>
<li>Built reusable data pipeline components and frameworks accelerating delivery across multiple IKEA product teams.</li>
<li>Drove knowledge sharing and community building — regular forums, documentation, and mentoring that built data engineering capability across the organization.</li>
<li>Worked within the IKEA ecosystem (IKEA App, Startcus Spare Parts, Genesys, Verint) — deep understanding of IKEA's data landscape and ways of working.</li>
</ul>

<p class="role">Software &amp; Data Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built data integration pipelines and ETL solutions for enterprise financial and retail domains using Python, SQL, and Java.</li>
<li>Developed data models and automated reporting solutions supporting business analytics and decision-making.</li>
<li>Worked with large datasets across relational databases, flat files, and APIs — building scalable data ingestion frameworks.</li>
<li>Participated in migration from legacy data systems to modern architectures.</li>
</ul>

<h2>IKEA ECOSYSTEM EXPERIENCE</h2>
<ul>
<li><strong>IKEA Customer Connect (VCS):</strong> Built end-to-end data layer — ETL pipelines, BigQuery data models, self-service analytics serving 32 markets within the Customer Support domain.</li>
<li><strong>IKEA Data Engineering Community:</strong> Active contributor — driving consistency, knowledge sharing, and best practices across the broader IKEA data engineering landscape.</li>
<li><strong>IKEA Ways of Working:</strong> Product team model, Agile/Scrum, Ingka DevOps tooling, cross-functional collaboration with business stakeholders, architects, and analysts.</li>
<li><strong>IKEA App &amp; Connected Products:</strong> Data pipelines for IoT telemetry and smart-home product analytics (DIRIGERA/TRÅDFRI ecosystem).</li>
<li><strong>CSSP (Customer Support Staff Planning):</strong> Data integration with Verint and Genesys — workforce analytics across global contact centers.</li>
<li><strong>Startcus Spare Parts:</strong> Data engineering for logistics and fulfillment tracking systems.</li>
<li><strong>Data Governance &amp; Compliance:</strong> GDPR-compliant data handling, metadata management, and security standards within IKEA's governance framework.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Togetherness:</strong> Build strong cross-functional relationships — collaborate with business stakeholders, data scientists, analysts, and product teams to deliver value together.</li>
<li><strong>Leading by Example:</strong> Mentor and develop co-workers through code reviews, knowledge sharing, and fostering a culture of engineering excellence and continuous learning.</li>
<li><strong>Simplicity:</strong> Design pragmatic, maintainable data solutions — reusable frameworks and clear standards that teams can adopt easily.</li>
<li><strong>Cost-Consciousness:</strong> Optimise data pipelines for cost-efficiency — reducing cloud spend while maintaining performance and reliability.</li>
<li><strong>Daring to be Different:</strong> Embrace modern data engineering technologies (Lakehouse, semantic layers, dbt) and apply them where they create clear business value.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
