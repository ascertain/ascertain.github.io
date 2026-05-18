from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Test_Specialist_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Test_Specialist_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "manual functional testing",
    "functional testing",
    "micro frontends",
    "micro frontend",
    "end-to-end",
    "E2E",
    "edge cases",
    "test automation",
    "automation framework",
    "test scenarios",
    "test cases",
    "test procedures",
    "quality advocate",
    "defect",
    "risk-based",
    "regression",
    "smoke",
    "exploratory",
    "UAT",
    "user acceptance",
    "BAT",
    "business acceptance",
    "API testing",
    "API automation",
    "API",
    "Selenium",
    "Cypress",
    "Playwright",
    "RestAssured",
    "Postman",
    "Jenkins",
    "GitHub Actions",
    "Azure DevOps",
    "CI/CD",
    "Grafana",
    "Prometheus",
    "ELK",
    "Docker",
    "Kubernetes",
    "cloud",
    "SDLC",
    "STLC",
    "Agile",
    "Waterfall",
    "Scrum",
    "ISTQB",
    "web application",
    "Web application",
    "cross-functional",
    "IKEA",
    "LEGO",
    "Truecaller",
    "Swedish",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Test Specialist  |  Manual, Automation, E2E, UAT, BAT, API & Exploratory Testing  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Test Specialist with 15+ years of hands-on experience in manual functional "
        "testing, test automation, and end-to-end quality assurance for complex Web "
        "applications, including micro frontend architectures. Comprehensive testing "
        "coverage — manual functional, automated E2E, API testing (REST/GraphQL), "
        "UAT (User Acceptance Testing), BAT (Business Acceptance Testing), exploratory "
        "testing, regression, smoke, and risk-based testing. Designs and maintains "
        "automation frameworks (Playwright, Selenium, Cypress) with CI/CD integration "
        "(GitHub Actions, Jenkins, Azure DevOps). Deep understanding of end-to-end "
        "functional flows — dependencies, edge cases, and data requirements across "
        "platform layers. Translates business requirements into structured test "
        "scenarios, test procedures, and test cases. Quality advocate — early lifecycle "
        "feedback. SDLC, STLC, Agile, Waterfall. ISTQB Certified. Currently at IKEA "
        "IT AB (3+ years). Swedish — Basic.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Manual Functional & Exploratory Testing: ",
            "15+ years performing manual functional testing on complex Web applications. "
            "Currently at IKEA — micro frontends, 30+ markets. Analyzes end-to-end "
            "functional flows, identifies dependencies, edge cases, and data requirements. "
            "Exploratory testing — session-based, charter-driven exploration to uncover "
            "defects beyond scripted scenarios. Validates data across platform layers.",
        ),
        (
            "Test Automation & API Testing: ",
            "Designs and maintains automation frameworks — Playwright (primary), Selenium, "
            "Cypress. Automated E2E, regression, and smoke suites. API testing and API "
            "automation (RestAssured, Postman, Karate) — REST and GraphQL endpoint "
            "validation, contract testing, response schema verification. CI/CD pipeline "
            "integration (GitHub Actions, Jenkins). Actively expands test automation suite.",
        ),
        (
            "UAT, BAT & End-to-End Testing: ",
            "Leads and executes UAT (User Acceptance Testing) — validates that business "
            "requirements are met from the end-user perspective. BAT (Business Acceptance "
            "Testing) — verifies business-critical workflows and acceptance criteria before "
            "release. End-to-end testing across full application stack — frontend, APIs, "
            "backend services, databases. Cross-browser and cross-device validation.",
        ),
        (
            "Test Design & Defect Management: ",
            "Translates business requirements into structured test scenarios, test "
            "procedures, and test cases. Risk-based testing, regression, smoke testing. "
            "Identifies, documents, and tracks defects — proactive collaboration with "
            "developers. Quality advocate — early lifecycle feedback on requirements, "
            "designs, and implementations. SDLC, STLC, Agile, Waterfall. ISTQB Certified.",
        ),
        (
            "Tools & Infrastructure: ",
            "CI/CD (GitHub Actions, Jenkins, Azure DevOps). Observability and quality "
            "metrics (Grafana, Prometheus awareness, ELK awareness). Docker, Kubernetes, "
            "cloud (GCP, AWS, Azure). Jira, Confluence, TestRail, Zephyr. Swedish — Basic.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Test Automation: ",
            "Playwright, Selenium, Cypress, Appium, RestAssured, Karate, "
            "Postman, E2E automation, API automation (REST/GraphQL), regression "
            "automation, smoke automation, contract testing, schema validation",
        ),
        (
            "Testing Practices: ",
            "Manual functional testing, exploratory testing (session-based), "
            "UAT (User Acceptance Testing), BAT (Business Acceptance Testing), "
            "end-to-end testing, API testing, risk-based testing, regression, "
            "smoke, edge-case analysis, cross-browser, data validation",
        ),
        (
            "CI/CD & DevOps: ",
            "GitHub Actions, Jenkins, Azure DevOps, Docker, Kubernetes, "
            "Terraform, GitOps, automated test pipelines with quality gates",
        ),
        (
            "Observability & Metrics: ",
            "Grafana, Cloud Monitoring, structured logging, Prometheus awareness, "
            "ELK awareness, quality metrics dashboards",
        ),
        (
            "Cloud & Platforms: ",
            "GCP, AWS, Azure, Docker, Kubernetes, micro frontend architectures, "
            "microservices, Web applications",
        ),
        (
            "Languages & Tools: ",
            "TypeScript, Java, Python, C#, SQL, Jira, Confluence, TestRail, "
            "Zephyr, Git",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Test Specialist / Senior SDET", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Complex Web Application, Micro Frontends, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Manual functional testing on complex Web application with multiple micro "
        "frontends — analyzed end-to-end functional flows, identified dependencies, "
        "edge cases, and data requirements. Exploratory testing — session-based "
        "exploration to uncover defects beyond scripted scenarios. Data validation "
        "across frontend, API, backend, and database layers.",
        "Test automation framework — Playwright for automated E2E, API testing "
        "(RestAssured, Postman — REST endpoint validation, contract testing, schema "
        "verification), regression and smoke automation suites. CI/CD integration "
        "(GitHub Actions). Actively expanded the automation suite.",
        "UAT (User Acceptance Testing) — validated business requirements from end-user "
        "perspective across 30+ markets. BAT (Business Acceptance Testing) — verified "
        "business-critical workflows and acceptance criteria before each release. "
        "Translated requirements into structured test scenarios and test procedures.",
        "Quality advocate — early feedback on requirements, designs, implementations. "
        "Risk-based, regression, smoke, and exploratory testing. Defect tracking and "
        "proactive developer collaboration. Agile (Scrum). Grafana quality metrics.",
        "AI-assisted testing (Claude, Copilot, Gemini) — 30% velocity improvement. "
        "Docker, Kubernetes, GCP cloud-native. Cross-functional collaboration.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior Test Specialist / QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Web & Mobile Platform — 300M+ Users, Global Scale", bold=True, size=10)
    tc_bullets = [
        "Manual functional testing and automated E2E testing on Web and mobile "
        "applications (300M+ users). End-to-end flow analysis, edge-case identification. "
        "API testing (RestAssured, Postman). UAT — validated user-facing flows. "
        "Exploratory testing. Selenium, Appium (iOS/Android). Defect tracking.",
        "Test automation framework and CI/CD. BAT — business-critical workflow "
        "verification before release. Risk-based, regression, smoke testing. "
        "Quality advocate. Agile (Scrum/Kanban). AWS cloud. Mentored engineers.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Test Specialist / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise Web Applications, Multi-Platform", bold=True, size=10)
    lego_bullets = [
        "LEGO: Manual functional testing, automated E2E, and API testing on e-commerce "
        "Web application — end-to-end flow analysis, data validation across layers. "
        "Automation framework (Selenium, RestAssured, Karate, Appium). UAT and BAT "
        "— user acceptance and business acceptance verification. Led 8–10 engineers. "
        "CI/CD (Jenkins). Defect management. Quality advocate.",
        "IKEA (2018–2021): Functional testing across Web applications (IKEA App, "
        "Genesys, Verint/CSSP, Spartacus). Manual functional, exploratory, API testing. "
        "UAT coordination with business stakeholders. Risk-based, regression, smoke "
        "testing. Test automation. Agile (Scrum).",
        "End-to-end testing ownership — translated business requirements into test "
        "scenarios and test procedures. Exploratory testing for edge-case discovery. "
        "Quality advocate — early lifecycle feedback. Quality metrics and reporting.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior Test Specialist / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Financial Services — Enterprise Web Applications", bold=True, size=10)
    fin_bullets = [
        "Manual functional, automated E2E, and API testing on enterprise Web "
        "applications (banking, Finacle). End-to-end flow analysis, data validation, "
        "edge-case and exploratory testing. UAT with business stakeholders. "
        "Automation frameworks (Selenium, TestNG). CI/CD. SQL, Java, C#. SDLC/STLC.",
        "Testing leadership — 15+ engineers. Risk-based, regression, smoke, "
        "exploratory, UAT, and BAT. Cross-functional collaboration. Quality advocate. "
        "Agile and Waterfall. Defect triage and resolution.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Manual functional testing and exploratory testing on complex Web applications "
        "with micro frontends at IKEA (30+ markets) and Truecaller (300M+ users). "
        "End-to-end flow analysis, edge-case discovery, data validation across layers.",
        "Test automation frameworks — Playwright, Selenium, Cypress, Appium. API testing "
        "and API automation (RestAssured, Postman, Karate). CI/CD (GitHub Actions, Jenkins) "
        "with automated E2E, regression, and smoke suites across 4 organizations.",
        "UAT (User Acceptance Testing) — validated business requirements from end-user "
        "perspective across 30+ markets. BAT (Business Acceptance Testing) — verified "
        "business-critical workflows before release. Quality advocate — early feedback.",
        "Risk-based, regression, smoke, and exploratory testing across SDLC/STLC in "
        "Agile and Waterfall. Observability and quality metrics (Grafana).",
        "AI-assisted testing — Claude, Copilot, Gemini — 30% velocity improvement. "
        "Docker, Kubernetes, cloud platforms (GCP, AWS, Azure).",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation (CTFL)\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Test Specialist Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Test Specialist | Manual, Automation, E2E, UAT, BAT, API &amp; Exploratory Testing | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Test Specialist with 15+ years in <span class="hl">manual functional testing</span>, <span class="hl">test automation</span>, and <span class="hl">end-to-end</span> quality assurance for complex <span class="hl">Web application</span>s, including <span class="hl">micro frontend</span> architectures. Comprehensive testing — <span class="hl">manual functional</span>, automated <span class="hl">E2E</span>, <span class="hl">API testing</span> (REST/GraphQL), <span class="hl">UAT</span> (User Acceptance Testing), <span class="hl">BAT</span> (Business Acceptance Testing), <span class="hl">exploratory</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>, <span class="hl">risk-based</span>. <span class="hl">Automation framework</span>s (<span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Cypress</span>) with <span class="hl">CI/CD</span>. <span class="hl">Quality advocate</span>. <span class="hl">SDLC</span>, <span class="hl">STLC</span>, <span class="hl">Agile</span>, <span class="hl">Waterfall</span>. <span class="hl">ISTQB</span> Certified. <span class="hl">Swedish</span> — Basic.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Manual Functional &amp; Exploratory:</b> 15+ years — <span class="hl">manual functional testing</span>, <span class="hl">exploratory</span> (session-based). <span class="hl">IKEA</span> — <span class="hl">micro frontend</span>s, 30+ markets. <span class="hl">End-to-end</span> flows, <span class="hl">edge cases</span>, data validation.<br>
  <b>Test Automation &amp; API Testing:</b> <span class="hl">Automation framework</span>s — <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Cypress</span>. <span class="hl">E2E</span>, <span class="hl">API testing</span> (<span class="hl">RestAssured</span>, <span class="hl">Postman</span>, Karate — REST/GraphQL). <span class="hl">Regression</span>, <span class="hl">smoke</span> suites. <span class="hl">CI/CD</span>.<br>
  <b>UAT, BAT &amp; E2E:</b> <span class="hl">UAT</span> — validates business requirements from end-user perspective. <span class="hl">BAT</span> — business-critical workflow verification before release. <span class="hl">End-to-end</span> across full stack.<br>
  <b>Test Design &amp; Defects:</b> Requirements → <span class="hl">test scenarios</span>/<span class="hl">test procedures</span>/<span class="hl">test cases</span>. <span class="hl">Risk-based</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>. <span class="hl">Defect</span> tracking. <span class="hl">Quality advocate</span>. <span class="hl">SDLC</span>, <span class="hl">STLC</span>, <span class="hl">Agile</span>, <span class="hl">Waterfall</span>. <span class="hl">ISTQB</span>.<br>
  <b>Tools &amp; Infra:</b> <span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>, <span class="hl">Azure DevOps</span>. <span class="hl">Grafana</span>, <span class="hl">Prometheus</span>, <span class="hl">ELK</span> awareness. <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">cloud</span>. <span class="hl">Swedish</span> — Basic.</p>

  <div class="section">Technical Skills</div>
  <p><b>Test Automation:</b> Playwright, Selenium, Cypress, Appium, RestAssured, Karate, Postman, E2E automation, API automation (REST/GraphQL), regression, smoke, contract testing<br>
  <b>Testing Practices:</b> Manual functional testing, exploratory (session-based), UAT (User Acceptance), BAT (Business Acceptance), end-to-end, API testing, risk-based, regression, smoke, edge-case analysis, cross-browser, data validation<br>
  <b>CI/CD &amp; DevOps:</b> GitHub Actions, Jenkins, Azure DevOps, Docker, Kubernetes, Terraform, automated test pipelines<br>
  <b>Observability:</b> Grafana, Cloud Monitoring, Prometheus awareness, ELK awareness, quality metrics<br>
  <b>Cloud:</b> GCP, AWS, Azure, Docker, Kubernetes, micro frontends, microservices, Web apps<br>
  <b>Languages &amp; Tools:</b> TypeScript, Java, Python, C#, SQL, Jira, Confluence, TestRail, Zephyr, Git</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Test Specialist / Senior SDET</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Complex Web Application, Micro Frontends, 30+ Markets</div>
  <ul>
    <li><span class="hl">Manual functional testing</span> on complex <span class="hl">Web application</span> with <span class="hl">micro frontend</span>s — <span class="hl">end-to-end</span> flows, <span class="hl">edge cases</span>, data validation. <span class="hl">Exploratory</span> testing (session-based). 30+ markets.</li>
    <li><span class="hl">Test automation</span> &amp; <span class="hl">automation framework</span> — <span class="hl">Playwright</span> (<span class="hl">E2E</span>, <span class="hl">API testing</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>). <span class="hl">RestAssured</span>, <span class="hl">Postman</span> — REST endpoint validation, contract testing. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>).</li>
    <li><span class="hl">UAT</span> — validated business requirements from end-user perspective (30+ markets). <span class="hl">BAT</span> — business-critical workflow verification before release. <span class="hl">Test scenarios</span>/<span class="hl">test procedures</span>/<span class="hl">test cases</span>.</li>
    <li><span class="hl">Quality advocate</span> — early feedback. <span class="hl">Risk-based</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>, <span class="hl">exploratory</span>. <span class="hl">Defect</span> tracking. <span class="hl">Agile</span> (<span class="hl">Scrum</span>). <span class="hl">Grafana</span>.</li>
    <li>AI-assisted testing (Claude, Copilot, Gemini) — 30% velocity. <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, GCP <span class="hl">cloud</span>-native.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior Test Specialist / QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Web &amp; Mobile Platform — 300M+ Users</div>
  <ul>
    <li><span class="hl">Manual functional</span>, automated <span class="hl">E2E</span>, and <span class="hl">API testing</span> — <span class="hl">Web application</span> &amp; mobile (300M+ users). <span class="hl">End-to-end</span> flows, <span class="hl">edge cases</span>. <span class="hl">UAT</span>. <span class="hl">Exploratory</span>. <span class="hl">Selenium</span>, Appium, <span class="hl">RestAssured</span>. <span class="hl">Defect</span> tracking.</li>
    <li><span class="hl">Automation framework</span> &amp; <span class="hl">CI/CD</span>. <span class="hl">BAT</span> — business-critical workflow verification. <span class="hl">Risk-based</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>. <span class="hl">Quality advocate</span>. <span class="hl">Agile</span>. AWS <span class="hl">cloud</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Test Specialist / Test Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise Web Applications</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Manual functional</span>, automated <span class="hl">E2E</span>, <span class="hl">API testing</span> — e-commerce <span class="hl">Web application</span>. <span class="hl">Automation framework</span> (<span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, Karate, Appium). <span class="hl">UAT</span> &amp; <span class="hl">BAT</span>. 8–10 engineers. <span class="hl">CI/CD</span> (<span class="hl">Jenkins</span>). <span class="hl">Quality advocate</span>.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Functional testing</span>, <span class="hl">exploratory</span>, <span class="hl">API testing</span> across <span class="hl">Web application</span>s (IKEA App, Genesys, Verint, Spartacus). <span class="hl">UAT</span>. <span class="hl">Risk-based</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>. <span class="hl">Agile</span>.</li>
    <li><span class="hl">End-to-end</span> testing ownership — <span class="hl">test scenarios</span>/<span class="hl">test procedures</span>. <span class="hl">Exploratory</span> for edge-case discovery. <span class="hl">Quality advocate</span>. Quality metrics.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior Test Specialist / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Financial Services — Enterprise Web Applications</div>
  <ul>
    <li><span class="hl">Manual functional</span>, automated <span class="hl">E2E</span>, <span class="hl">API testing</span> on enterprise <span class="hl">Web application</span>s (Finacle banking). <span class="hl">End-to-end</span> flows. <span class="hl">UAT</span>. <span class="hl">Exploratory</span>. <span class="hl">Automation framework</span>s (<span class="hl">Selenium</span>, TestNG). <span class="hl">CI/CD</span>. <span class="hl">SDLC</span>, <span class="hl">STLC</span>.</li>
    <li>Leadership — 15+ engineers. <span class="hl">Risk-based</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>, <span class="hl">exploratory</span>, <span class="hl">UAT</span>, <span class="hl">BAT</span>. <span class="hl">Agile</span> &amp; <span class="hl">Waterfall</span>. <span class="hl">Quality advocate</span>. <span class="hl">Defect</span> triage.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Manual functional</span> and <span class="hl">exploratory</span> testing on complex <span class="hl">Web application</span>s with <span class="hl">micro frontend</span>s at <span class="hl">IKEA</span> (30+ markets) and <span class="hl">Truecaller</span> (300M+ users). <span class="hl">End-to-end</span> flows, <span class="hl">edge cases</span>.</li>
    <li><span class="hl">Automation framework</span>s — <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Cypress</span>. <span class="hl">API testing</span> (<span class="hl">RestAssured</span>, <span class="hl">Postman</span>, Karate). <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>). <span class="hl">E2E</span>, <span class="hl">regression</span>, <span class="hl">smoke</span> suites.</li>
    <li><span class="hl">UAT</span> — end-user validation (30+ markets). <span class="hl">BAT</span> — business-critical workflow verification before release. <span class="hl">Quality advocate</span> — early feedback.</li>
    <li><span class="hl">Risk-based</span>, <span class="hl">regression</span>, <span class="hl">smoke</span>, <span class="hl">exploratory</span> across <span class="hl">SDLC</span>/<span class="hl">STLC</span> in <span class="hl">Agile</span> &amp; <span class="hl">Waterfall</span>. <span class="hl">Grafana</span> quality metrics.</li>
    <li>AI-assisted testing — 30% velocity. <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">cloud</span> (GCP, AWS, Azure).</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation (CTFL)<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
