"""
Resume: Senior Software Engineer – Event Intelligence, IKEA (INTERNAL)
Focus: AI/ML/LLM on Observability data, event pipelines, automated incident detection,
       Grafana stack, Splunk ITSI, Python, Kubernetes, GCP, ServiceNow/JIRA integration.
Internal = highlight IKEA ecosystem knowledge.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Senior_SE_Event_Intelligence_Resume"
IKEA_BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = IKEA_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = IKEA_BLUE
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = IKEA_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Senior Software Engineer with 14+ years in IT, currently within the IKEA/Ingka Digital ecosystem, "
        "combining deep platform engineering and data pipeline expertise with a strong DevOps & SRE mindset. "
        "Experienced in building event-driven data pipelines, integrating observability tooling (Grafana stack, "
        "OpenTelemetry), and applying AI/ML techniques to operational data for anomaly detection and automated "
        "incident response. Proficient in Python and cloud-native architectures on GCP and Kubernetes. "
        "Passionate about leveraging AI/LLM as an everyday companion to solve complex engineering challenges "
        "and drive operational intelligence at scale across IKEA's 500+ technology teams."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Event Management & Correlation", "AI/ML/LLM for Observability", "Python & Golang Development",
        "Grafana Stack (Loki/Tempo/Mimir)", "OpenTelemetry & Tracing", "Splunk Enterprise / ITSI",
        "Kubernetes & Cloud-Native (GCP)", "Automated Incident Pipelines", "ServiceNow & JIRA Integration",
        "Data Engineering & Pipelines", "DevOps / SRE / CI-CD", "Slack Bots & Notification Flows",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TECHNICAL SKILLS ──
    add_heading_block(doc, "Technical Skills")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    skills = (
        "Languages: Python, Golang, TypeScript, SQL  •  "
        "AI/ML: Scikit-learn, TensorFlow, LangChain, OpenAI APIs, anomaly detection models  •  "
        "Observability: Grafana, Loki, Tempo, Mimir, OpenTelemetry, Splunk, Datadog  •  "
        "Cloud & Infra: GCP (BigQuery, Cloud Functions, Pub/Sub, GKE), Kubernetes, Docker, Terraform  •  "
        "ITSM: ServiceNow, JIRA, Slack API  •  "
        "CI/CD: GitHub Actions, Jenkins, ArgoCD  •  "
        "Data: BigQuery, Pub/Sub, Apache Beam, event streaming, ETL pipelines"
    )
    r = tp.add_run(skills)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- IKEA / Ingka Digital ---
    role_header(doc, "Team Lead & Senior Engineer – Visual Customer Support (VCS)",
                "Ingka Digital / IKEA", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Architect and lead development of event-driven data pipelines processing high-volume operational data from IKEA's customer contact platforms (VCS, Genesys, Verint, CSSP) using Python, GCP Pub/Sub, BigQuery, and Cloud Functions.")
    bullet(doc, "Build automated event correlation and anomaly detection logic that monitors system health metrics, logs, and traces across 5+ integrated platforms — proactively identifying outages before user impact.")
    bullet(doc, "Implement automated notification and incident creation flows integrating with Slack (bot notifications) and JIRA for cross-team incident collaboration and escalation workflows.")
    bullet(doc, "Design and maintain observability instrumentation using OpenTelemetry and Grafana dashboards for real-time system health visibility across the VCS ecosystem.")
    bullet(doc, "Apply AI/ML techniques (anomaly detection, pattern recognition, LLM-powered log analysis) to operational data to surface actionable insights and drive automated recovery workflows.")
    bullet(doc, "Operate production workloads on GCP (GKE/Kubernetes); manage CI/CD pipelines via GitHub Actions; maintain infrastructure-as-code with Terraform.")
    bullet(doc, "Collaborate closely with IKEA's Observability platform team and Operational Intelligence to align Event Intelligence approaches and share tooling across the organisation.")
    bullet(doc, "Lead a team of engineers in an autonomous, self-organizing agile setup; drive architecture decisions, code reviews, and continuous delivery practices.")
    bullet(doc, "Leverage AI/LLM tools daily as a development companion — code generation, data analysis, incident root-cause exploration — embracing AI-native engineering workflows.")

    # --- Truecaller ---
    role_header(doc, "QA Lead & Platform Engineer",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Built automated testing and monitoring pipelines for microservices on Kubernetes; integrated observability tooling for service health tracking.")
    bullet(doc, "Developed Python-based data validation and anomaly detection scripts for platform metrics; contributed to incident response automation.")
    bullet(doc, "Worked with ServiceNow and JIRA for incident management workflows in a cloud-native SaaS environment.")

    # --- HCLTech ---
    role_header(doc, "Senior Engineer / Technical Lead – Enterprise Platforms",
                "HCLTech (for IKEA & LEGO)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Designed and implemented event-driven integration pipelines connecting 10+ enterprise systems (e-commerce, order management, supply chain) for IKEA and LEGO programmes.")
    bullet(doc, "Built monitoring and alerting solutions correlating application logs, infrastructure metrics, and business events to detect service degradation and trigger automated responses.")
    bullet(doc, "Developed Python tooling for data engineering tasks: ETL pipelines, data quality checks, and automated reporting across large-scale datasets.")
    bullet(doc, "Integrated incident management workflows with ServiceNow and JIRA; automated ticket creation and escalation based on event severity rules.")
    bullet(doc, "Operated services on hybrid infrastructure (on-prem + cloud); gained deep Kubernetes and container orchestration experience.")
    bullet(doc, "Drove DevOps culture: CI/CD pipeline creation (Jenkins, GitHub Actions), infrastructure automation, and shift-left quality practices.")
    bullet(doc, "Led research initiatives exploring ML-based anomaly detection on operational data to improve mean-time-to-detect (MTTD) for critical services.")

    # --- India ---
    role_header(doc, "Software Engineer / Test Engineer",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Developed and tested enterprise web applications (Java, .NET); built automated test frameworks and data validation scripts in Python.")
    bullet(doc, "Gained foundational experience in systems engineering, scripting, database management, and cross-platform integration.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "Google Cloud Associate Cloud Engineer",
        "AWS Certified Cloud Practitioner",
        "ISTQB Certified Tester",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("PGDOM")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – IGNOU  |  ").font.size = Pt(9)
    r2 = p.add_run("B.Tech Information Technology")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – UP Technical University").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Conversational)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    # ── WHY THIS ROLE ──
    add_heading_block(doc, "Why Event Intelligence at IKEA")
    why = (
        "As a current Ingka Digital co-worker, I have first-hand understanding of the scale and complexity of IKEA's "
        "technology landscape — 500+ teams, diverse platforms, and the critical need for proactive operational intelligence. "
        "I'm passionate about applying AI/ML to Observability data and building the automated event pipelines that keep "
        "IKEA's digital products resilient. This role is the perfect intersection of my data engineering skills, "
        "DevOps/SRE mindset, and growing AI/LLM expertise — and I'm excited to help shape the next generation of "
        "Event Intelligence at IKEA."
    )
    pw = doc.add_paragraph(why)
    pw.paragraph_format.space_before = Pt(3)
    for run in pw.runs:
        run.font.size = Pt(9)
        run.italic = True

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#0051BA;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#0051BA;font-size:11pt;border-bottom:1px solid #0051BA;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#0051BA;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    .why{font-style:italic;font-size:9pt;margin-top:4px}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Senior Software Engineer with 14+ years in IT, currently within the IKEA/Ingka Digital ecosystem,
combining deep platform engineering and data pipeline expertise with a strong DevOps &amp; SRE mindset.
Experienced in building event-driven data pipelines, integrating observability tooling (Grafana stack,
OpenTelemetry), and applying AI/ML techniques to operational data for anomaly detection and automated
incident response. Proficient in Python and cloud-native architectures on GCP and Kubernetes.
Passionate about leveraging AI/LLM as an everyday companion to solve complex engineering challenges
and drive operational intelligence at scale across IKEA's 500+ technology teams.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Event Management &amp; Correlation</td><td>✓ AI/ML/LLM for Observability</td><td>✓ Python &amp; Golang Development</td></tr>
<tr><td>✓ Grafana Stack (Loki/Tempo/Mimir)</td><td>✓ OpenTelemetry &amp; Tracing</td><td>✓ Splunk Enterprise / ITSI</td></tr>
<tr><td>✓ Kubernetes &amp; Cloud-Native (GCP)</td><td>✓ Automated Incident Pipelines</td><td>✓ ServiceNow &amp; JIRA Integration</td></tr>
<tr><td>✓ Data Engineering &amp; Pipelines</td><td>✓ DevOps / SRE / CI-CD</td><td>✓ Slack Bots &amp; Notification Flows</td></tr>
</table>

<h2>TECHNICAL SKILLS</h2>
<p class="tools">Languages: Python, Golang, TypeScript, SQL &bull;
AI/ML: Scikit-learn, TensorFlow, LangChain, OpenAI APIs, anomaly detection models &bull;
Observability: Grafana, Loki, Tempo, Mimir, OpenTelemetry, Splunk, Datadog &bull;
Cloud &amp; Infra: GCP (BigQuery, Cloud Functions, Pub/Sub, GKE), Kubernetes, Docker, Terraform &bull;
ITSM: ServiceNow, JIRA, Slack API &bull;
CI/CD: GitHub Actions, Jenkins, ArgoCD &bull;
Data: BigQuery, Pub/Sub, Apache Beam, event streaming, ETL pipelines</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Team Lead &amp; Senior Engineer – Visual Customer Support (VCS) &nbsp;|&nbsp; Ingka Digital / IKEA, Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Architect and lead development of event-driven data pipelines processing high-volume operational data from IKEA's customer contact platforms (VCS, Genesys, Verint, CSSP) using Python, GCP Pub/Sub, BigQuery, and Cloud Functions.</li>
<li>Build automated event correlation and anomaly detection logic that monitors system health metrics, logs, and traces across 5+ integrated platforms — proactively identifying outages before user impact.</li>
<li>Implement automated notification and incident creation flows integrating with Slack (bot notifications) and JIRA for cross-team incident collaboration and escalation workflows.</li>
<li>Design and maintain observability instrumentation using OpenTelemetry and Grafana dashboards for real-time system health visibility across the VCS ecosystem.</li>
<li>Apply AI/ML techniques (anomaly detection, pattern recognition, LLM-powered log analysis) to operational data to surface actionable insights and drive automated recovery workflows.</li>
<li>Operate production workloads on GCP (GKE/Kubernetes); manage CI/CD pipelines via GitHub Actions; maintain infrastructure-as-code with Terraform.</li>
<li>Collaborate closely with IKEA's Observability platform team and Operational Intelligence to align Event Intelligence approaches and share tooling across the organisation.</li>
<li>Lead a team of engineers in an autonomous, self-organizing agile setup; drive architecture decisions, code reviews, and continuous delivery practices.</li>
<li>Leverage AI/LLM tools daily as a development companion — code generation, data analysis, incident root-cause exploration — embracing AI-native engineering workflows.</li>
</ul>

<p class="role">QA Lead &amp; Platform Engineer &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Built automated testing and monitoring pipelines for microservices on Kubernetes; integrated observability tooling for service health tracking.</li>
<li>Developed Python-based data validation and anomaly detection scripts for platform metrics; contributed to incident response automation.</li>
<li>Worked with ServiceNow and JIRA for incident management workflows in a cloud-native SaaS environment.</li>
</ul>

<p class="role">Senior Engineer / Technical Lead – Enterprise Platforms &nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Designed and implemented event-driven integration pipelines connecting 10+ enterprise systems (e-commerce, order management, supply chain) for IKEA and LEGO programmes.</li>
<li>Built monitoring and alerting solutions correlating application logs, infrastructure metrics, and business events to detect service degradation and trigger automated responses.</li>
<li>Developed Python tooling for data engineering tasks: ETL pipelines, data quality checks, and automated reporting across large-scale datasets.</li>
<li>Integrated incident management workflows with ServiceNow and JIRA; automated ticket creation and escalation based on event severity rules.</li>
<li>Operated services on hybrid infrastructure (on-prem + cloud); gained deep Kubernetes and container orchestration experience.</li>
<li>Drove DevOps culture: CI/CD pipeline creation (Jenkins, GitHub Actions), infrastructure automation, and shift-left quality practices.</li>
<li>Led research initiatives exploring ML-based anomaly detection on operational data to improve mean-time-to-detect (MTTD) for critical services.</li>
</ul>

<p class="role">Software Engineer / Test Engineer &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Developed and tested enterprise web applications (Java, .NET); built automated test frameworks and data validation scripts in Python.</li>
<li>Gained foundational experience in systems engineering, scripting, database management, and cross-platform integration.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>Google Cloud Associate Cloud Engineer</li>
<li>AWS Certified Cloud Practitioner</li>
<li>ISTQB Certified Tester</li>
<li>Six Sigma Green Belt</li>
<li>Certified Ethical Hacker (CEH)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>PGDOM</b> – IGNOU &nbsp;|&nbsp; <b>B.Tech Information Technology</b> – UP Technical University</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Conversational) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

<h2>WHY EVENT INTELLIGENCE AT IKEA</h2>
<p class="why">As a current Ingka Digital co-worker, I have first-hand understanding of the scale and complexity of IKEA's
technology landscape — 500+ teams, diverse platforms, and the critical need for proactive operational intelligence.
I'm passionate about applying AI/ML to Observability data and building the automated event pipelines that keep
IKEA's digital products resilient. This role is the perfect intersection of my data engineering skills,
DevOps/SRE mindset, and growing AI/LLM expertise — and I'm excited to help shape the next generation of
Event Intelligence at IKEA.</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
