"""
Generate a tailored resume for Agile Project Manager.
Focus: backlog ownership, agile facilitation, stakeholder coordination,
data-driven decisions, cross-team synchronization, continuous improvement.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Agile_Project_Manager_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Agile_Project_Manager_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "Agile", "Scrum", "Kanban", "backlog", "product backlog",
    "sprint planning", "roadmap", "stakeholder", "product owner",
    "cross-functional", "data-driven", "continuous improvement",
    "prioritization", "dependencies", "CI/CD", "collaboration",
    "facilitation", "retrospective", "planning sessions",
    "Jira", "Confluence", "Azure DevOps",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Agile Project Manager", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Agile Project Manager with 16+ years in software development, combining strong technical "
        "understanding with proven facilitation and delivery skills. Experienced in owning and "
        "maintaining product backlogs, breaking down requirements with development teams, and "
        "driving prioritization aligned with business goals. Skilled at building strong relationships "
        "with product managers, engineering teams, and stakeholders — ensuring clarity on requirements, "
        "scope, and priorities. Track record of coordinating cross-team dependencies, facilitating "
        "planning sessions and roadmap processes, and fostering data-driven decision-making. "
        "Passionate about continuous improvement of agile ways of working and creating conditions "
        "where teams are empowered and deliver effectively. Fluent in Swedish and English."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Backlog Ownership & Prioritization", "Agile Facilitation (Scrum/Kanban)", "Stakeholder & Product Manager Alignment",
        "Cross-Team Coordination & Dependencies", "Sprint & Roadmap Planning", "Data-Driven Decision Making",
        "Continuous Improvement & Ways of Working", "Requirements Breakdown & Scoping", "Open Communication & Collaboration",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Team Lead / Agile Delivery Lead — VCS Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Own and maintain the product backlog for a customer-facing video platform across 30+ markets — ensuring correct prioritization and well-organized work for 3 cross-functional squads.",
        "Break down requirements together with development teams (backend, data engineering, infrastructure), translating business needs into actionable user stories and technical tasks.",
        "Act as central point of contact between product managers, engineering teams, and stakeholders — ensuring alignment on requirements, priority, and scope at all times.",
        "Facilitate sprint planning, backlog refinement, retrospectives, and roadmap sessions — driving predictable and transparent delivery cadence.",
        "Coordinate internal projects and synchronize dependencies across backend (Node.js/TypeScript), data pipeline (Python/BigQuery), and infrastructure (Terraform/GCP) teams.",
        "Champion data-driven decision-making: established delivery metrics (cycle time, throughput, defect rate) used in sprint reviews and roadmap planning.",
        "Foster open communication and collaboration across distributed teams — creating an environment where teams are empowered and blockers are resolved quickly.",
        "Drive continuous improvement of agile ways of working: introduced structured retrospectives, refined Definition of Done, and improved sprint predictability by 35%.",
        "Led scaling from 2,000 to 50,000+ concurrent users — managing backlog prioritization through rapid growth while maintaining quality and team focus.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Release & Delivery Coordinator", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Coordinated release planning and cross-squad dependencies for a 300M+ user platform — synchronizing deliveries across multiple autonomous squads.",
        "Maintained shared backlog visibility and facilitated prioritization discussions between product owners and engineering leads.",
        "Drove process improvements in release workflow: introduced feature flag governance and rollback checklists, reducing release incidents by 40%.",
        "Fostered collaboration between development, QA, and operations teams — ensuring smooth workflow and timely deliveries with clear communication.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Specialist / Agile Delivery Lead", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed product backlogs and sprint planning for IKEA e-commerce (30+ markets) and LEGO digital platform projects — ensuring development teams had clear, prioritized work at all times.",
        "Built strong working relationships with product managers and customer stakeholders — translating requirements into well-scoped user stories and ensuring priority alignment.",
        "Coordinated dependencies across 4+ vendor teams in Denmark, Sweden, and India — synchronizing deliveries, managing blockers, and maintaining milestone accountability.",
        "Facilitated agile ceremonies (sprint planning, reviews, retrospectives) across distributed teams — continuously improving ways of working and team effectiveness.",
        "Drove adoption of data-driven practices: introduced velocity tracking, burn-down reporting, and defect trend analysis to support informed prioritization decisions.",
        "Guided teams through methodology transitions (Waterfall → Agile/Scrum) — coaching team members and establishing sustainable agile practices.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise IT & Banking", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Lead / Project Coordinator", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Coordinated project deliveries for core banking implementations (Finacle CBS) — managing backlog, priorities, and cross-team dependencies in regulated environments.",
        "Broke down complex requirements (data migration, integration, ETL) with development teams and ensured clear scope definition for multi-phase rollouts.",
        "Served as liaison between business analysts, development teams, and stakeholders — maintaining transparency and driving timely delivery across 20+ branches.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Environment ──────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL ENVIRONMENT")
    for label, value in [
        ("Agile & Delivery: ", "Scrum, Kanban, SAFe (exposure), Jira, Confluence, Azure DevOps, Miro"),
        ("Cloud & DevOps: ", "GCP (Cloud Run, BigQuery, Pub/Sub), AWS, Terraform, GitHub Actions, Docker"),
        ("Development: ", "Node.js, TypeScript, Python, REST APIs, Kafka/Pub/Sub, CI/CD pipelines"),
        ("Monitoring & Data: ", "Datadog, Grafana, BigQuery, delivery metrics dashboards"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PG Diploma, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "ITIL v4 Foundation",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "Swedish (Fluent)  •  English (Fluent)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Agile Project Manager</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Agile Project Manager with 16+ years in software development, combining strong technical understanding with proven facilitation and delivery skills. Experienced in owning and maintaining product backlogs, breaking down requirements with development teams, and driving prioritization aligned with business goals. Skilled at building strong relationships with product managers, engineering teams, and stakeholders &mdash; ensuring clarity on requirements, scope, and priorities. Track record of coordinating cross-team dependencies, facilitating planning sessions and roadmap processes, and fostering data-driven decision-making. Passionate about continuous improvement of agile ways of working and creating conditions where teams are empowered and deliver effectively. Fluent in Swedish and English.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Backlog Ownership &amp; Prioritization</td><td>&bull; Agile Facilitation (Scrum/Kanban)</td><td>&bull; Stakeholder &amp; Product Manager Alignment</td></tr>
<tr><td>&bull; Cross-Team Coordination &amp; Dependencies</td><td>&bull; Sprint &amp; Roadmap Planning</td><td>&bull; Data-Driven Decision Making</td></tr>
<tr><td>&bull; Continuous Improvement &amp; Ways of Working</td><td>&bull; Requirements Breakdown &amp; Scoping</td><td>&bull; Open Communication &amp; Collaboration</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Team Lead / Agile Delivery Lead &mdash; VCS Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Own and maintain the product backlog for a customer-facing video platform across 30+ markets &mdash; ensuring correct prioritization and well-organized work for 3 cross-functional squads.</li>
<li>Break down requirements together with development teams (backend, data engineering, infrastructure), translating business needs into actionable user stories and technical tasks.</li>
<li>Act as central point of contact between product managers, engineering teams, and stakeholders &mdash; ensuring alignment on requirements, priority, and scope at all times.</li>
<li>Facilitate sprint planning, backlog refinement, retrospectives, and roadmap sessions &mdash; driving predictable and transparent delivery cadence.</li>
<li>Coordinate internal projects and synchronize dependencies across backend (Node.js/TypeScript), data pipeline (Python/BigQuery), and infrastructure (Terraform/GCP) teams.</li>
<li>Champion data-driven decision-making: established delivery metrics (cycle time, throughput, defect rate) used in sprint reviews and roadmap planning.</li>
<li>Foster open communication and collaboration across distributed teams &mdash; creating an environment where teams are empowered and blockers are resolved quickly.</li>
<li>Drive continuous improvement of agile ways of working: introduced structured retrospectives, refined Definition of Done, and improved sprint predictability by 35%.</li>
<li>Led scaling from 2,000 to 50,000+ concurrent users &mdash; managing backlog prioritization through rapid growth while maintaining quality and team focus.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Release &amp; Delivery Coordinator <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Coordinated release planning and cross-squad dependencies for a 300M+ user platform &mdash; synchronizing deliveries across multiple autonomous squads.</li>
<li>Maintained shared backlog visibility and facilitated prioritization discussions between product owners and engineering leads.</li>
<li>Drove process improvements in release workflow: introduced feature flag governance and rollback checklists, reducing release incidents by 40%.</li>
<li>Fostered collaboration between development, QA, and operations teams &mdash; ensuring smooth workflow and timely deliveries with clear communication.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Technical Specialist / Agile Delivery Lead <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Managed product backlogs and sprint planning for IKEA e-commerce (30+ markets) and LEGO digital platform projects &mdash; ensuring development teams had clear, prioritized work at all times.</li>
<li>Built strong working relationships with product managers and customer stakeholders &mdash; translating requirements into well-scoped user stories and ensuring priority alignment.</li>
<li>Coordinated dependencies across 4+ vendor teams in Denmark, Sweden, and India &mdash; synchronizing deliveries, managing blockers, and maintaining milestone accountability.</li>
<li>Facilitated agile ceremonies (sprint planning, reviews, retrospectives) across distributed teams &mdash; continuously improving ways of working and team effectiveness.</li>
<li>Drove adoption of data-driven practices: introduced velocity tracking, burn-down reporting, and defect trend analysis to support informed prioritization decisions.</li>
<li>Guided teams through methodology transitions (Waterfall &rarr; Agile/Scrum) &mdash; coaching team members and establishing sustainable agile practices.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise IT &amp; Banking</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Coordinated project deliveries for core banking implementations (Finacle CBS) &mdash; managing backlog, priorities, and cross-team dependencies in regulated environments.</li>
<li>Broke down complex requirements (data migration, integration, ETL) with development teams and ensured clear scope definition for multi-phase rollouts.</li>
<li>Served as liaison between business analysts, development teams, and stakeholders &mdash; maintaining transparency and driving timely delivery across 20+ branches.</li>
</ul>

<h2>TECHNICAL ENVIRONMENT</h2>
<p class="tech-line"><b>Agile &amp; Delivery:</b> Scrum, Kanban, SAFe (exposure), Jira, Confluence, Azure DevOps, Miro</p>
<p class="tech-line"><b>Cloud &amp; DevOps:</b> GCP (Cloud Run, BigQuery, Pub/Sub), AWS, Terraform, GitHub Actions, Docker</p>
<p class="tech-line"><b>Development:</b> Node.js, TypeScript, Python, REST APIs, Kafka/Pub/Sub, CI/CD pipelines</p>
<p class="tech-line"><b>Monitoring &amp; Data:</b> Datadog, Grafana, BigQuery, delivery metrics dashboards</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; Six Sigma Green Belt</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">Swedish (Fluent) &bull; English (Fluent) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
