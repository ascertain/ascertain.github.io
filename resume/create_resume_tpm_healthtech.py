from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Technical_Project_Manager_Healthtech_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Technical_Project_Manager_Healthtech_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "integration",
    "API",
    "SaaS",
    "platform",
    "healthcare",
    "healthtech",
    "external partner",
    "stakeholder",
    "cross-functional",
    "data model",
    "system architecture",
    "authentication",
    "scalable",
    "dependencies",
    "bottleneck",
    "process",
    "framework",
    "template",
    "risk",
    "visibility",
    "product mindset",
    "technical fluency",
    "vendor",
    "coordination",
    "delivery",
    "workstream",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Technical Project Manager  |  Integrations, SaaS Platforms & Partner Coordination  |  10+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Technical Project Manager with 10+ years leading integration work with external "
        "partners and enterprise systems across SaaS platforms. Maps external systems against "
        "internal platforms and proposes concrete integration designs. Central point of "
        "coordination between engineering, partnerships, and external stakeholders. Translates "
        "complex technical requirements into clear priorities, plans, and deliverables. "
        "Technically fluent — comfortable discussing APIs, authentication, data models, and "
        "system architecture with engineers. Manages multiple workstreams without losing "
        "clarity or momentum. Identifies bottlenecks, manages dependencies, and keeps projects "
        "moving at high velocity. Builds and improves how teams work — processes, templates, "
        "and frameworks for integrations and technical delivery. Product mindset — focused "
        "on building the right thing. Strong communication — tracks progress, surfaces risks, "
        "and provides clear visibility across stakeholders. Personal project: SimpleRx — "
        "AI-powered health report assistant (Google Play).",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Integration Leadership & Partner Coordination: ",
            "Led integration work with external SaaS partners at IKEA — mapped external "
            "vendor systems against internal platform, proposed integration designs, and "
            "coordinated implementation. Central point between engineering, partnerships, "
            "and external stakeholders. Managed API integrations, event-driven data exchange, "
            "and data pipelines. Ensured high-quality, scalable implementations across 30+ "
            "global markets.",
        ),
        (
            "Technical Fluency — APIs, Data Models, Architecture: ",
            "Hands-on technical background — comfortable discussing APIs (REST, GraphQL), "
            "authentication (OAuth, token-based), data models, and system architecture with "
            "engineers. Worked closely with engineering teams to ensure integration solutions "
            "are robust and aligned with platform architecture. Python, TypeScript, C#, Java. "
            "Cloud-native (GCP, AWS, Docker, Terraform).",
        ),
        (
            "Managing Multiple Workstreams & Dependencies: ",
            "Managed multiple concurrent workstreams across IKEA (30+ markets), Truecaller "
            "(300M+ users), and LEGO — without losing clarity or momentum. Identified "
            "bottlenecks, managed dependencies, and kept projects at high velocity. "
            "Translated complex technical requirements into clear priorities and deliverables.",
        ),
        (
            "Process Building & Continuous Improvement: ",
            "Built and improved how teams work — established processes, templates, and "
            "frameworks for integrations and technical delivery. Led Selenium-to-Playwright "
            "migration (3x faster, 50% CI reduction). Created CoE practices — guidelines, "
            "quality gates, reusable frameworks. Six Sigma Green Belt.",
        ),
        (
            "Healthtech & Product Mindset: ",
            "Co-founder of SimpleRx — an AI-powered health report assistant on Google Play "
            "that translates medical reports into plain language. End-to-end product "
            "ownership: vision, AI/LLM integration, mobile development, deployment. "
            "Privacy-first, health data compliance. Product mindset — building the right "
            "thing, not just delivering on time.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Integration & Partner Management: ",
            "External partner integration. System mapping and integration design. "
            "API integrations (REST, GraphQL, webhooks). Authentication (OAuth, "
            "token-based). Data exchange (event streams, batch, real-time). Partner "
            "coordination and stakeholder alignment",
        ),
        (
            "Technical Project Management: ",
            "Multi-workstream management. Dependency tracking. Bottleneck identification. "
            "Priorities, plans, deliverables. Risk management. Progress visibility "
            "and reporting. Agile (Scrum/SAFe), Waterfall",
        ),
        (
            "Process & Frameworks: ",
            "Process design and improvement. Templates and frameworks for delivery. "
            "CoE practices. Quality gates. Continuous improvement. Six Sigma GB. "
            "ITIL Foundation",
        ),
        (
            "Technical Fluency: ",
            "APIs, data models, system architecture. Python, TypeScript, C#, Java. "
            "GCP, AWS, Docker, Kubernetes, Terraform. CI/CD (GitHub Actions, Jenkins). "
            "Data pipelines (BigQuery, Pub/Sub). Grafana",
        ),
        (
            "Communication & Stakeholders: ",
            "Engineering, partnerships, external stakeholders. Translating technical "
            "complexity into clear decisions. Progress tracking, risk surfacing. "
            "Vendor relationship management. Cross-functional coordination",
        ),
        (
            "Tools: ",
            "Jira, Confluence, MS Project, TestRail. Playwright, Selenium. "
            "Postman, Swagger. Git, GitHub Actions. Grafana dashboards",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Technical Project Lead / Team Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  SaaS Platform — Integrations, External Partners, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Integration leadership — led integration work with external SaaS partner for "
        "IKEA's Visual Collaboration platform. Mapped external vendor systems against "
        "internal platform and proposed concrete integration designs. Managed API "
        "integrations (REST), event-driven data exchange (Pub/Sub), and data pipelines "
        "(BigQuery, Cloud Functions). Ensured high-quality, scalable implementations "
        "across 30+ global markets.",
        "Central coordination point — between engineering, partnerships, and external "
        "stakeholders. Translated complex technical requirements into clear priorities, "
        "plans, and deliverables. Tracked progress, surfaced risks, and communicated "
        "clearly across all stakeholders. Managed vendor relationship — aligned delivery "
        "with business KPIs.",
        "Technical fluency — worked closely with engineers on APIs, data models, "
        "authentication, and system architecture. Ensured integration solutions were "
        "robust and aligned with platform. Python, TypeScript, CI/CD (GitHub Actions), "
        "Terraform, Docker on GCP.",
        "Process building — established processes, templates, and frameworks for "
        "integrations and technical delivery. Built CoE practices — guidelines, quality "
        "gates, reusable frameworks. Led Selenium-to-Playwright migration (3x faster, "
        "50% CI reduction). AI-assisted — 30% velocity improvement.",
        "Team management — led engineers and consultants. Coaching, mentoring, competence "
        "development. Managed multiple workstreams, identified bottlenecks, kept projects "
        "at high velocity. Budget and resource management. Exceptional Performer.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Release & Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  SaaS Platform — 300M+ Users, API-Driven, Distributed", bold=True, size=10)
    tc_bullets = [
        "Technical delivery at scale — managed multiple workstreams for a platform "
        "serving 300M+ users. Identified bottlenecks, managed dependencies, and "
        "maintained velocity. API-driven architecture. Translated technical complexity "
        "into clear decisions. Stakeholder communication and risk visibility. AWS.",
        "Cross-functional coordination — engineering, product, partnerships. Release "
        "automation, CI/CD. Feature flag management and data-driven rollouts. Process "
        "improvement. Agile (Scrum/Kanban).",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "HCLTech — LEGO & IKEA Group, Denmark & Sweden — Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2013 – 2021  |  E-Commerce & Enterprise — Integrations, Partners, Multi-Platform", bold=True, size=10)
    lego_bullets = [
        "LEGO & IKEA (2017–2021): Managed integration work across enterprise platforms — "
        "external partner APIs, third-party system integrations, multi-partner delivery. "
        "Coordinated between engineering, partners, and stakeholders. Multiple concurrent "
        "workstreams. Dependency management. Teams of 8–10 across onshore-offshore. "
        "Translated requirements into plans and deliverables. Agile and Waterfall.",
        "Technical Lead (2013–2017): Hands-on technical delivery — API integrations, "
        "automation frameworks (Selenium, C#), CI/CD. Worked closely with engineers on "
        "system architecture and data flows. Process building — templates, frameworks, "
        "quality gates. Managed dependencies across workstreams. Mentoring.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "Banking & Enterprise — Technical Lead / Consultant", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2013  |  Core Banking (Finacle CBS) — System Integrations, Data Exchange", bold=True, size=10)
    fin_bullets = [
        "Integration-focused delivery — core banking system integrations (Finacle CBS, "
        "biometric hardware, ETL pipelines). Mapped external systems against platform. "
        "Data exchange (batch, API). Managed dependencies across workstreams. "
        "Stakeholder coordination. Regulated environment with compliance requirements.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Personal Project ──
    add_section_heading(document, "Personal Project — Healthtech")
    pp = document.add_paragraph()
    pp.paragraph_format.space_after = Pt(0)
    pp.paragraph_format.line_spacing = 1.0
    add_text(pp, "SimpleRx — AI-Powered Health Report Assistant  ", bold=True, size=10)
    add_text(pp, "(Google Play)  |  Co-Founder & Developer, DQODIFY Pvt. Ltd.", size=10, color=TEXT_MUTED)
    pp2 = document.add_paragraph()
    pp2.paragraph_format.space_after = Pt(0)
    pp2.paragraph_format.line_spacing = 1.0
    add_text(
        pp2,
        "AI-powered healthtech app translating medical reports, lab results, and "
        "prescriptions into plain language (Hindi & English). End-to-end product "
        "ownership: vision, AI/LLM integration, mobile development, Play Store "
        "deployment. Health tracker with trend graphs. Family health profiles. "
        "Privacy-first — encrypted storage, health data compliance. Product mindset.",
        size=10,
    )

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Integration leadership — led external partner integration at IKEA (SaaS vendor "
        "↔ internal platform), mapped systems, proposed designs, coordinated engineering "
        "and stakeholders. Scaled to 30+ markets. Exceptional Performer.",
        "Process building — established integration processes, templates, and frameworks. "
        "CoE practices. Led Selenium-to-Playwright migration (3x faster, 50% CI reduction). "
        "Continuously improved how teams work.",
        "Multi-workstream delivery — IKEA (30+ markets), Truecaller (300M+ users), LEGO. "
        "Identified bottlenecks, managed dependencies, maintained high velocity across "
        "concurrent workstreams.",
        "Healthtech product — co-founded SimpleRx (Google Play), AI-powered health report "
        "assistant. Product mindset. Health data compliance.",
        "Technical fluency across stacks — APIs, data models, system architecture. "
        "Python, TypeScript, C#. GCP, AWS, Docker, Terraform, CI/CD.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Six Sigma Green Belt\n"
        "ITIL Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Technical Project Manager Healthtech Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Technical Project Manager | Integrations, SaaS Platforms &amp; Partner Coordination | 10+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Technical Project Manager with 10+ years leading <span class="hl">integration</span> work with external partners and enterprise systems across <span class="hl">SaaS</span> <span class="hl">platform</span>s. Maps external systems against internal <span class="hl">platform</span>s and proposes concrete <span class="hl">integration</span> designs. Central point of <span class="hl">coordination</span> between engineering, partnerships, and external <span class="hl">stakeholder</span>s. Translates complex technical requirements into clear priorities, plans, and deliverables. Technically fluent — <span class="hl">API</span>s, <span class="hl">authentication</span>, <span class="hl">data model</span>s, <span class="hl">system architecture</span>. Manages multiple <span class="hl">workstream</span>s. Identifies <span class="hl">bottleneck</span>s, manages <span class="hl">dependencies</span>. Builds <span class="hl">process</span>es, <span class="hl">template</span>s, <span class="hl">framework</span>s. <span class="hl">Product mindset</span>. SimpleRx — AI-powered <span class="hl">health</span> report assistant (Google Play).</p>

  <div class="section">How I Match the Role</div>
  <p><b>Integration Leadership &amp; Partner Coordination:</b> Led <span class="hl">integration</span> with external <span class="hl">SaaS</span> partner at IKEA — mapped systems, proposed designs, managed <span class="hl">API</span> <span class="hl">integration</span>s, event-driven data exchange, data pipelines. <span class="hl">Scalable</span> implementations across 30+ markets.<br><br>
  <b>Technical Fluency:</b> <span class="hl">API</span>s (REST, GraphQL), <span class="hl">authentication</span> (OAuth), <span class="hl">data model</span>s, <span class="hl">system architecture</span>. Worked closely with engineers. Python, TypeScript, C#, Java. GCP, AWS, Docker, Terraform.<br><br>
  <b>Multi-Workstream &amp; Dependencies:</b> Multiple concurrent <span class="hl">workstream</span>s — IKEA (30+ markets), Truecaller (300M+), LEGO. <span class="hl">Bottleneck</span>s, <span class="hl">dependencies</span>, high velocity. Clear priorities and deliverables.<br><br>
  <b>Process Building:</b> <span class="hl">Process</span>es, <span class="hl">template</span>s, <span class="hl">framework</span>s for <span class="hl">integration</span>s and <span class="hl">delivery</span>. CoE practices. Selenium-to-Playwright (3x/50%). Six Sigma GB.<br><br>
  <b>Healthtech &amp; Product Mindset:</b> SimpleRx — AI <span class="hl">health</span> report assistant (Google Play). End-to-end product ownership. Health data compliance. <span class="hl">Product mindset</span>.</p>

  <div class="section">Core Competencies</div>
  <p><b>Integration &amp; Partners:</b> External partner integration. System mapping &amp; design. API (REST, GraphQL, webhooks). Authentication (OAuth). Data exchange. Partner coordination<br>
  <b>Technical PM:</b> Multi-workstream. Dependencies. Bottlenecks. Priorities, plans, deliverables. Risk. Visibility &amp; reporting. Agile (Scrum/SAFe), Waterfall<br>
  <b>Process &amp; Frameworks:</b> Process design. Templates. CoE. Quality gates. Continuous improvement. Six Sigma GB. ITIL<br>
  <b>Technical:</b> APIs, data models, system architecture. Python, TypeScript, C#, Java. GCP, AWS, Docker, K8s, Terraform. CI/CD. Data pipelines. Grafana<br>
  <b>Communication:</b> Engineering, partnerships, external stakeholders. Technical → clear decisions. Progress, risks. Vendor management. Cross-functional<br>
  <b>Tools:</b> Jira, Confluence, MS Project, TestRail. Playwright, Selenium. Postman, Swagger. Git, GitHub Actions</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Technical Project Lead / Team Lead</div>
  <div class="job-sub">Mar 2022 – Present | SaaS Platform — Integrations, External Partners, 30+ Markets</div>
  <ul>
    <li><span class="hl">Integration</span> leadership — external <span class="hl">SaaS</span> partner. Mapped systems, proposed <span class="hl">integration</span> designs. <span class="hl">API</span> (REST), event-driven data exchange, data pipelines. <span class="hl">Scalable</span> implementations — 30+ markets.</li>
    <li>Central <span class="hl">coordination</span> — engineering, partnerships, external <span class="hl">stakeholder</span>s. Translated requirements into priorities, plans, deliverables. Progress <span class="hl">visibility</span>, <span class="hl">risk</span>s. <span class="hl">Vendor</span> relationship aligned with KPIs.</li>
    <li>Technical fluency — <span class="hl">API</span>s, <span class="hl">data model</span>s, <span class="hl">authentication</span>, <span class="hl">system architecture</span>. Python, TypeScript, CI/CD, Terraform, Docker, GCP.</li>
    <li><span class="hl">Process</span> building — <span class="hl">template</span>s, <span class="hl">framework</span>s for <span class="hl">integration</span>s and <span class="hl">delivery</span>. CoE. Playwright migration (3x/50%). AI-assisted — 30% velocity. Exceptional Performer.</li>
    <li>Team — engineers + consultants. Coaching, mentoring. Multiple <span class="hl">workstream</span>s, <span class="hl">bottleneck</span>s, <span class="hl">dependencies</span>. Budget &amp; resource management.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Release &amp; Technical Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | SaaS Platform — 300M+ Users, API-Driven, Distributed</div>
  <ul>
    <li>Multiple <span class="hl">workstream</span>s — 300M+ users. <span class="hl">Bottleneck</span>s, <span class="hl">dependencies</span>, velocity. <span class="hl">API</span>-driven. Technical complexity → clear decisions. <span class="hl">Stakeholder</span> <span class="hl">visibility</span>. AWS.</li>
    <li><span class="hl">Cross-functional</span> — engineering, product, partnerships. <span class="hl">Process</span> improvement. CI/CD. Feature flags. Agile.</li>
  </ul>

  <div class="job-title">HCLTech — LEGO &amp; IKEA, Denmark &amp; Sweden — Technical Lead</div>
  <div class="job-sub">2013 – 2021 | E-Commerce &amp; Enterprise — Integrations, Partners, Multi-Platform</div>
  <ul>
    <li>LEGO &amp; IKEA (2017–21): <span class="hl">Integration</span> across enterprise <span class="hl">platform</span>s — external partner <span class="hl">API</span>s, third-party systems. <span class="hl">Coordination</span> (engineering, partners, <span class="hl">stakeholder</span>s). Multiple <span class="hl">workstream</span>s. <span class="hl">Dependencies</span>. Teams of 8–10. Plans &amp; deliverables. Agile &amp; Waterfall.</li>
    <li>Technical Lead (2013–17): <span class="hl">API</span> <span class="hl">integration</span>s, <span class="hl">framework</span>s (Selenium, C#), CI/CD. <span class="hl">System architecture</span>, data flows. <span class="hl">Process</span> building — <span class="hl">template</span>s, <span class="hl">framework</span>s, quality gates. <span class="hl">Dependencies</span> across <span class="hl">workstream</span>s. Mentoring.</li>
  </ul>

  <div class="job-title">Banking &amp; Enterprise — Technical Lead / Consultant</div>
  <div class="job-sub">2008 – 2013 | Core Banking — System Integrations, Data Exchange</div>
  <ul>
    <li><span class="hl">Integration</span> — core banking (Finacle CBS), biometric hardware, ETL. Mapped external systems. Data exchange (batch, <span class="hl">API</span>). <span class="hl">Dependencies</span>. <span class="hl">Stakeholder</span> <span class="hl">coordination</span>. Regulated.</li>
  </ul>

  <div class="section">Personal Project — Healthtech</div>
  <p><b>SimpleRx — AI-Powered Health Report Assistant</b> <span style="color:#4b5563">(Google Play) | Co-Founder &amp; Developer, DQODIFY Pvt. Ltd.</span><br>
  AI <span class="hl">healthtech</span> app — medical reports → plain language. End-to-end product ownership. AI/LLM <span class="hl">integration</span>. Health tracker. Family profiles. Privacy-first, health data compliance. <span class="hl">Product mindset</span>.</p>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Integration</span> leadership — IKEA <span class="hl">SaaS</span> <span class="hl">vendor</span> ↔ <span class="hl">platform</span>. System mapping, designs, 30+ markets. Exceptional Performer.</li>
    <li><span class="hl">Process</span> building — <span class="hl">integration</span> <span class="hl">process</span>es, <span class="hl">template</span>s, <span class="hl">framework</span>s. CoE. Playwright (3x/50%).</li>
    <li>Multi-<span class="hl">workstream</span> — IKEA, Truecaller, LEGO. <span class="hl">Bottleneck</span>s, <span class="hl">dependencies</span>, velocity.</li>
    <li><span class="hl">Healthtech</span> product — SimpleRx (Google Play). AI health reports. <span class="hl">Product mindset</span>.</li>
    <li>Technical fluency — <span class="hl">API</span>s, <span class="hl">data model</span>s, <span class="hl">system architecture</span>. Python, TypeScript, C#. GCP, AWS, CI/CD.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>Six Sigma Green Belt<br>ITIL Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>UiPath RPA Certified</td>
    </tr>
  </table>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
