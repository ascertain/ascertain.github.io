"""
Generate a tailored resume for Software/Systems Engineering Lead.
Focus: system delivery, engineering work structuring, quality/compliance/cybersecurity,
documentation/validation, technical leadership, customer collaboration, risk management.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Systems_Engineering_Lead_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Systems_Engineering_Lead_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "system architecture", "end-to-end delivery", "system-level",
    "engineering work packages", "traceability",
    "cybersecurity", "regulatory", "compliance", "safety",
    "documentation", "validation", "verification",
    "cross-functional", "technical leadership",
    "customer collaboration", "contractual",
    "risk management", "proposals", "estimates", "scope",
    "MS Project", "Jira", "CI/CD",
    "Agile", "V-Model",
    "stakeholder", "milestones",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Software & Systems Engineering Lead", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Software & Systems Engineering Lead with 16+ years of experience driving end-to-end "
        "delivery of complex software and system-level solutions. Proven track record of leading "
        "the planning, execution, and verification of all software deliverables — translating "
        "system architecture into clear, traceable engineering work packages. Strong background "
        "in quality and compliance oversight, ensuring deliverables meet functional, performance, "
        "cybersecurity, and regulatory standards. Experienced in driving documentation, testing, "
        "and validation activities aligned with project milestones. Skilled at leading "
        "cross-functional engineering teams, coordinating technical decisions across disciplines, "
        "and engaging with customers to shape solutions and ensure contractual alignment. "
        "Effective at managing technical risks and contributing to proposals, estimates, and "
        "scope management."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Software System Delivery (E2E)", "System Architecture & Work Structuring", "Quality, Compliance & Cybersecurity",
        "Documentation & Validation (V&V)", "Cross-Functional Technical Leadership", "Customer Collaboration & Contracts",
        "Risk Management & Mitigation", "Proposals, Estimates & Scope Mgmt", "Agile & V-Model Project Delivery",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Engineering Lead / Technical Delivery Manager — Cloud Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Lead end-to-end software system delivery for a customer-facing platform across 30+ markets — owning planning, execution, verification, and go-live of all system-level deliverables.",
        "Translate system architecture into traceable engineering work packages — decomposing requirements into clear development tasks, integration points, and acceptance criteria.",
        "Ensure quality and compliance oversight across all deliverables: functional correctness, performance benchmarks, cybersecurity controls, and regulatory alignment (GDPR, data privacy).",
        "Drive documentation, testing, and validation activities aligned with project milestones — maintaining traceability from requirements through verification and customer acceptance.",
        "Lead cross-functional engineering teams (backend, data, infrastructure, security) — coordinating technical decisions across disciplines and ensuring architectural consistency.",
        "Engage with internal customers (product, business stakeholders, market leads) to shape solutions, resolve technical issues, and ensure contractual/SLA alignment.",
        "Manage technical risks through structured risk registers, impact assessment, and mitigation planning — escalating proactively when thresholds are exceeded.",
        "Contribute to proposals and estimates for new platform capabilities — providing technical scope assessments, effort estimates, and delivery timelines.",
        "Scaled platform from 2,000 to 50,000+ concurrent users while maintaining system integrity, performance SLAs, and security compliance throughout all delivery phases.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Release & Systems Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led planning and verification of software releases for a 300M+ user platform — ensuring all system-level deliverables met quality, performance, and security standards.",
        "Structured engineering work across development squads — coordinating dependencies, integration milestones, and validation activities across microservice boundaries.",
        "Managed technical risk during release cycles — identifying blockers, assessing impact, and coordinating mitigation with cross-functional teams.",
        "Drove documentation and validation practices: release notes, deployment verification checklists, and rollback procedures aligned with compliance requirements.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Specialist / Systems Delivery Lead", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led software system delivery for IKEA e-commerce (30+ markets) and LEGO digital platform — managing end-to-end execution from architecture translation through verification and go-live.",
        "Translated system architectures into structured engineering work packages for multi-partner delivery teams (4+ vendors across Denmark, Sweden, India) — ensuring traceability and accountability.",
        "Ensured quality and compliance across deliverables: functional testing, performance validation, security assessments, and regulatory compliance for global market deployments.",
        "Drove documentation and validation activities: test strategies, verification reports, technical specifications, and milestone-aligned delivery documentation.",
        "Led cross-functional engineering teams across disciplines (frontend, backend, infrastructure, testing) — coordinating technical decisions and resolving architectural conflicts.",
        "Engaged directly with customer technical representatives (IKEA IT, LEGO Digital) to shape solutions, resolve issues, and maintain contractual alignment on scope and timelines.",
        "Managed technical risks within complex multi-system integration projects — maintaining risk registers, conducting impact analysis, and reporting to programme leadership.",
        "Contributed to proposals and estimates for new project phases — providing technical assessments, resource plans, and work breakdown structures for RFP/RFI responses.",
        "Guided teams through V-Model and Agile delivery approaches — ensuring appropriate verification and validation practices regardless of methodology.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise IT & Banking", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Lead / Systems Engineer", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led software system delivery for core banking implementations (Finacle CBS) — managing planning, verification, and validation for regulated financial systems across 20+ branches.",
        "Structured engineering work for data migration, integration, and system deployment projects — translating architecture into executable work packages with clear milestones.",
        "Ensured compliance with banking regulatory standards — driving documentation, audit-trail maintenance, and system validation for production releases.",
        "Coordinated cross-functional teams (development, DBA, infrastructure, business) and engaged with customer IT departments to align on technical scope and contractual obligations.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Environment ──────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL ENVIRONMENT")
    for label, value in [
        ("Systems & Architecture: ", "Cloud-native (GCP, AWS), microservices, event-driven architecture, API design, system decomposition"),
        ("Project & Delivery: ", "MS Project, Jira, Azure DevOps, Confluence, Agile (Scrum/Kanban), V-Model, Waterfall"),
        ("Quality & Compliance: ", "V&V, ISTQB, cybersecurity standards (awareness), GDPR, performance testing, security assessments"),
        ("Engineering: ", "Python, TypeScript, Node.js, Terraform (IaC), Docker, Kubernetes, CI/CD (GitHub Actions, Jenkins)"),
        ("Documentation: ", "Technical specifications, verification reports, risk registers, WBS, proposal documentation"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PG Diploma, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "Certified Ethical Hacker (CEH)",
        "ITIL v4 Foundation",
        "AWS Certified Cloud Practitioner",
        "Six Sigma Green Belt",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Software &amp; Systems Engineering Lead</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Software &amp; Systems Engineering Lead with 16+ years of experience driving end-to-end delivery of complex software and system-level solutions. Proven track record of leading the planning, execution, and verification of all software deliverables &mdash; translating system architecture into clear, traceable engineering work packages. Strong background in quality and compliance oversight, ensuring deliverables meet functional, performance, cybersecurity, and regulatory standards. Experienced in driving documentation, testing, and validation activities aligned with project milestones. Skilled at leading cross-functional engineering teams, coordinating technical decisions across disciplines, and engaging with customers to shape solutions and ensure contractual alignment. Effective at managing technical risks and contributing to proposals, estimates, and scope management.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Software System Delivery (E2E)</td><td>&bull; System Architecture &amp; Work Structuring</td><td>&bull; Quality, Compliance &amp; Cybersecurity</td></tr>
<tr><td>&bull; Documentation &amp; Validation (V&amp;V)</td><td>&bull; Cross-Functional Technical Leadership</td><td>&bull; Customer Collaboration &amp; Contracts</td></tr>
<tr><td>&bull; Risk Management &amp; Mitigation</td><td>&bull; Proposals, Estimates &amp; Scope Mgmt</td><td>&bull; Agile &amp; V-Model Project Delivery</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Engineering Lead / Technical Delivery Manager &mdash; Cloud Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Lead end-to-end software system delivery for a customer-facing platform across 30+ markets &mdash; owning planning, execution, verification, and go-live of all system-level deliverables.</li>
<li>Translate system architecture into traceable engineering work packages &mdash; decomposing requirements into clear development tasks, integration points, and acceptance criteria.</li>
<li>Ensure quality and compliance oversight across all deliverables: functional correctness, performance benchmarks, cybersecurity controls, and regulatory alignment (GDPR, data privacy).</li>
<li>Drive documentation, testing, and validation activities aligned with project milestones &mdash; maintaining traceability from requirements through verification and customer acceptance.</li>
<li>Lead cross-functional engineering teams (backend, data, infrastructure, security) &mdash; coordinating technical decisions across disciplines and ensuring architectural consistency.</li>
<li>Engage with internal customers (product, business stakeholders, market leads) to shape solutions, resolve technical issues, and ensure contractual/SLA alignment.</li>
<li>Manage technical risks through structured risk registers, impact assessment, and mitigation planning &mdash; escalating proactively when thresholds are exceeded.</li>
<li>Contribute to proposals and estimates for new platform capabilities &mdash; providing technical scope assessments, effort estimates, and delivery timelines.</li>
<li>Scaled platform from 2,000 to 50,000+ concurrent users while maintaining system integrity, performance SLAs, and security compliance throughout all delivery phases.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Release &amp; Systems Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Led planning and verification of software releases for a 300M+ user platform &mdash; ensuring all system-level deliverables met quality, performance, and security standards.</li>
<li>Structured engineering work across development squads &mdash; coordinating dependencies, integration milestones, and validation activities across microservice boundaries.</li>
<li>Managed technical risk during release cycles &mdash; identifying blockers, assessing impact, and coordinating mitigation with cross-functional teams.</li>
<li>Drove documentation and validation practices: release notes, deployment verification checklists, and rollback procedures aligned with compliance requirements.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Technical Specialist / Systems Delivery Lead <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Led software system delivery for IKEA e-commerce (30+ markets) and LEGO digital platform &mdash; managing end-to-end execution from architecture translation through verification and go-live.</li>
<li>Translated system architectures into structured engineering work packages for multi-partner delivery teams (4+ vendors across Denmark, Sweden, India) &mdash; ensuring traceability and accountability.</li>
<li>Ensured quality and compliance across deliverables: functional testing, performance validation, security assessments, and regulatory compliance for global market deployments.</li>
<li>Drove documentation and validation activities: test strategies, verification reports, technical specifications, and milestone-aligned delivery documentation.</li>
<li>Led cross-functional engineering teams across disciplines (frontend, backend, infrastructure, testing) &mdash; coordinating technical decisions and resolving architectural conflicts.</li>
<li>Engaged directly with customer technical representatives (IKEA IT, LEGO Digital) to shape solutions, resolve issues, and maintain contractual alignment on scope and timelines.</li>
<li>Managed technical risks within complex multi-system integration projects &mdash; maintaining risk registers, conducting impact analysis, and reporting to programme leadership.</li>
<li>Contributed to proposals and estimates for new project phases &mdash; providing technical assessments, resource plans, and work breakdown structures for RFP/RFI responses.</li>
<li>Guided teams through V-Model and Agile delivery approaches &mdash; ensuring appropriate verification and validation practices regardless of methodology.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise IT &amp; Banking</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Led software system delivery for core banking implementations (Finacle CBS) &mdash; managing planning, verification, and validation for regulated financial systems across 20+ branches.</li>
<li>Structured engineering work for data migration, integration, and system deployment projects &mdash; translating architecture into executable work packages with clear milestones.</li>
<li>Ensured compliance with banking regulatory standards &mdash; driving documentation, audit-trail maintenance, and system validation for production releases.</li>
<li>Coordinated cross-functional teams (development, DBA, infrastructure, business) and engaged with customer IT departments to align on technical scope and contractual obligations.</li>
</ul>

<h2>TECHNICAL ENVIRONMENT</h2>
<p class="tech-line"><b>Systems &amp; Architecture:</b> Cloud-native (GCP, AWS), microservices, event-driven architecture, API design, system decomposition</p>
<p class="tech-line"><b>Project &amp; Delivery:</b> MS Project, Jira, Azure DevOps, Confluence, Agile (Scrum/Kanban), V-Model, Waterfall</p>
<p class="tech-line"><b>Quality &amp; Compliance:</b> V&amp;V, ISTQB, cybersecurity standards (awareness), GDPR, performance testing, security assessments</p>
<p class="tech-line"><b>Engineering:</b> Python, TypeScript, Node.js, Terraform (IaC), Docker, Kubernetes, CI/CD (GitHub Actions, Jenkins)</p>
<p class="tech-line"><b>Documentation:</b> Technical specifications, verification reports, risk registers, WBS, proposal documentation</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; Certified Ethical Hacker (CEH)</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
