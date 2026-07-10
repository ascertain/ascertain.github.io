"""
Generate a tailored resume for TPM Nordic Cloud Operations.
Focus: cloud project delivery, coordination/planning, ITIL/ITSM, healthcare IT,
regulated environments, Swedish+English, stakeholder communication, operational stability.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_TPM_Nordic_Cloud_Ops_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_TPM_Nordic_Cloud_Ops_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "cloud operations", "cloud infrastructure", "cloud services",
    "end-to-end", "go-live", "stabilization",
    "coordination", "planning", "delivery",
    "stakeholder", "communication", "transparent",
    "ITIL", "ITSM", "incident management",
    "regulated", "healthcare", "medical",
    "connectivity", "systems integration", "operational stability",
    "Swedish", "English", "Nordic",
    "escalation", "risk", "priorities",
    "agile", "Prince2", "Scrum",
    "DICOM", "HL7",
    "GCP", "Azure", "AWS",
    "CI/CD", "DevOps",
    "cross-functional", "engineers",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Technical Project Manager — Cloud Operations", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Technical Project Manager with 15+ years of experience driving cloud operations projects "
        "end-to-end — from planning and coordination through go-live and stabilization. Combines a solid "
        "technical foundation in cloud infrastructure (GCP, Azure, AWS) with a primary strength in bringing "
        "structure, clarity, and momentum to complex delivery work. Experienced in regulated environments "
        "including healthcare IT, where operational stability and compliance are non-negotiable. Skilled at "
        "translating complex technical realities into language that supports good decision-making, managing "
        "expectations at strategic and operational levels, and knowing when to escalate, adapt, or say no. "
        "Strong communicator in both English and Swedish, with extensive Nordic team collaboration experience. "
        "ITIL-certified with deep familiarity in ITSM frameworks, connectivity, and systems integration."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Cloud Project Delivery (E2E)", "Planning & Coordination", "Stakeholder Communication",
        "Cloud Infrastructure (GCP/Azure/AWS)", "ITIL / ITSM Frameworks", "Regulated / Healthcare IT",
        "Systems Integration & Connectivity", "Operational Stability", "Agile & Prince2 Methodologies",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Project Manager — Cloud Operations", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Lead cloud operations projects end-to-end — owning delivery from planning and coordination through go-live and stabilization across a platform serving 30+ Nordic and global markets.",
        "Act as primary point of contact between engineering teams and internal stakeholders — bringing clarity, structure, and realistic planning that turns complex technical work into delivered outcomes.",
        "Keep projects on track through clear priorities, transparent communication, and proactive risk management — knowing when to escalate, adapt timelines, or push back on scope.",
        "Translate complex cloud infrastructure realities (GCP, Azure) into language that supports good decision-making for business and operational leadership.",
        "Coordinate cross-functional delivery across engineers, operations teams, and external vendors — managing dependencies, environment readiness, and integration milestones.",
        "Drive operational stability improvements for cloud services — contributing to incident management processes, ITSM workflows, and service continuity practices aligned with ITIL frameworks.",
        "Manage expectations at both strategic and operational levels — providing transparent status reporting, facilitating steering committee updates, and ensuring stakeholders have what they need to act.",
        "Lead go-live coordination and stabilization activities — orchestrating cutover planning, rollback procedures, and post-deployment monitoring across interconnected cloud systems.",
        "Contribute to continuous improvement in how cloud services are run, supported, and delivered — introducing standardized delivery playbooks that reduced project cycle time by 30%.",
        "Facilitate agile ceremonies (sprint planning, retrospectives, PI planning) while maintaining overall project governance and milestone tracking aligned with program objectives.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Delivery Lead — Cloud Platform", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Coordinated cloud operations delivery for a 300M+ user platform — managing project planning, priorities, and cross-team dependencies in a fast-paced cloud environment (AWS).",
        "Acted as bridge between engineering teams and business stakeholders — translating technical constraints into actionable plans and maintaining transparent communication on progress and risks.",
        "Drove go-live readiness and stabilization for cloud service deployments — coordinating release schedules, monitoring setup, and incident response procedures.",
        "Applied ITIL-aligned practices for incident management and change control — ensuring operational stability while supporting rapid iteration cycles.",
        "Managed escalation paths and risk mitigation strategies — making timely calls on when to adapt plans, defer scope, or escalate blockers to leadership.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical Project Manager / Delivery Lead — Enterprise & Cloud", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led end-to-end delivery of cloud operations and infrastructure projects across the Nordic region — managing complex IT landscapes spanning cloud (Azure, GCP), on-premise systems, and multi-vendor integrations.",
        "Owned project planning and coordination for large-scale transformation programs — defining scope, timelines, resource allocation, and go-live criteria across distributed teams in Denmark and Sweden.",
        "Served as primary point of contact between engineering specialists and business stakeholders — maintaining transparent communication and facilitating informed decision-making at all levels.",
        "Managed healthcare-adjacent IT projects involving regulated environments, patient data systems, and compliance requirements — ensuring operational stability and adherence to regulatory standards.",
        "Drove connectivity and systems integration initiatives — coordinating cloud-to-on-premise integrations, API connectivity, and data flow orchestration across complex enterprise architectures.",
        "Applied ITIL and ITSM frameworks to establish incident management, change management, and problem management processes — improving service stability and reducing unplanned downtime by 40%.",
        "Coordinated go-live and stabilization activities for major platform migrations — managing cutover weekends, rollback planning, and hypercare periods with engineering and operations teams.",
        "Managed expectations across strategic (program steering) and operational (daily standup) levels — knowing when to escalate, when to absorb pressure, and when to say no to protect delivery quality.",
        "Mentored and coordinated teams of 6–15 engineers and specialists — creating an environment where complex work could be broken into clear deliverables with shared ownership.",
        "Delivered projects using agile, Prince2, and hybrid methodologies — adapting approach to project complexity, stakeholder needs, and organizational maturity.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Technical Project Delivery", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Project Lead / Technical Coordinator", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Coordinated delivery of enterprise IT projects in regulated financial environments (core banking, insurance) — managing planning, stakeholder communication, and go-live execution.",
        "Acted as liaison between technical teams and business stakeholders — translating complex requirements into delivery plans and keeping projects on track through clear priorities.",
        "Managed systems integration and connectivity between distributed IT systems — coordinating infrastructure provisioning, network configuration, and operational handover.",
        "Applied ITIL practices for service management transitions — ensuring operational stability and clear support handover from project delivery to operations teams.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Foundation ───────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL FOUNDATION")
    for label, value in [
        ("Cloud Platforms: ", "GCP (Cloud Run, BigQuery, Pub/Sub, IAM), Azure (DevOps, App Services, Pipelines), AWS (EC2, S3, Lambda)"),
        ("Infrastructure: ", "Kubernetes, Docker, Terraform, CI/CD pipelines, networking (TCP/IP, VPN, DNS, firewalls)"),
        ("ITSM & Frameworks: ", "ITIL v4, incident/change/problem management, service catalog, SLA management, operational runbooks"),
        ("Connectivity: ", "Systems integration, API management, HL7/FHIR interfaces, data flow orchestration, hybrid cloud architecture"),
        ("Project Methods: ", "Agile (Scrum/Kanban), SAFe, Prince2, hybrid delivery, risk-based planning, go-live governance"),
        ("Tools: ", "Jira, Confluence, Azure DevOps, ServiceNow, MS Project, Git, monitoring dashboards"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ITIL v4 Foundation",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
        "Six Sigma Green Belt",
        "ISTQB Certified Tester",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent — primary working language)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Technical Project Manager &mdash; Cloud Operations</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Technical Project Manager with 15+ years of experience driving cloud operations projects end-to-end &mdash; from planning and coordination through go-live and stabilization. Combines a solid technical foundation in cloud infrastructure (GCP, Azure, AWS) with a primary strength in bringing structure, clarity, and momentum to complex delivery work. Experienced in regulated environments including healthcare IT, where operational stability and compliance are non-negotiable. Skilled at translating complex technical realities into language that supports good decision-making, managing expectations at strategic and operational levels, and knowing when to escalate, adapt, or say no. Strong communicator in both English and Swedish, with extensive Nordic team collaboration experience. ITIL-certified with deep familiarity in ITSM frameworks, connectivity, and systems integration.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Cloud Project Delivery (E2E)</td><td>&bull; Planning &amp; Coordination</td><td>&bull; Stakeholder Communication</td></tr>
<tr><td>&bull; Cloud Infrastructure (GCP/Azure/AWS)</td><td>&bull; ITIL / ITSM Frameworks</td><td>&bull; Regulated / Healthcare IT</td></tr>
<tr><td>&bull; Systems Integration &amp; Connectivity</td><td>&bull; Operational Stability</td><td>&bull; Agile &amp; Prince2 Methodologies</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Technical Project Manager &mdash; Cloud Operations <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Lead cloud operations projects end-to-end &mdash; owning delivery from planning and coordination through go-live and stabilization across a platform serving 30+ Nordic and global markets.</li>
<li>Act as primary point of contact between engineering teams and internal stakeholders &mdash; bringing clarity, structure, and realistic planning that turns complex technical work into delivered outcomes.</li>
<li>Keep projects on track through clear priorities, transparent communication, and proactive risk management &mdash; knowing when to escalate, adapt timelines, or push back on scope.</li>
<li>Translate complex cloud infrastructure realities (GCP, Azure) into language that supports good decision-making for business and operational leadership.</li>
<li>Coordinate cross-functional delivery across engineers, operations teams, and external vendors &mdash; managing dependencies, environment readiness, and integration milestones.</li>
<li>Drive operational stability improvements for cloud services &mdash; contributing to incident management processes, ITSM workflows, and service continuity practices aligned with ITIL frameworks.</li>
<li>Manage expectations at both strategic and operational levels &mdash; providing transparent status reporting, facilitating steering committee updates, and ensuring stakeholders have what they need to act.</li>
<li>Lead go-live coordination and stabilization activities &mdash; orchestrating cutover planning, rollback procedures, and post-deployment monitoring across interconnected cloud systems.</li>
<li>Contribute to continuous improvement in how cloud services are run, supported, and delivered &mdash; introducing standardized delivery playbooks that reduced project cycle time by 30%.</li>
<li>Facilitate agile ceremonies (sprint planning, retrospectives, PI planning) while maintaining overall project governance and milestone tracking aligned with program objectives.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Technical Delivery Lead &mdash; Cloud Platform <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Coordinated cloud operations delivery for a 300M+ user platform &mdash; managing project planning, priorities, and cross-team dependencies in a fast-paced cloud environment (AWS).</li>
<li>Acted as bridge between engineering teams and business stakeholders &mdash; translating technical constraints into actionable plans and maintaining transparent communication on progress and risks.</li>
<li>Drove go-live readiness and stabilization for cloud service deployments &mdash; coordinating release schedules, monitoring setup, and incident response procedures.</li>
<li>Applied ITIL-aligned practices for incident management and change control &mdash; ensuring operational stability while supporting rapid iteration cycles.</li>
<li>Managed escalation paths and risk mitigation strategies &mdash; making timely calls on when to adapt plans, defer scope, or escalate blockers to leadership.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Technical Project Manager / Delivery Lead &mdash; Enterprise &amp; Cloud <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Led end-to-end delivery of cloud operations and infrastructure projects across the Nordic region &mdash; managing complex IT landscapes spanning cloud (Azure, GCP), on-premise systems, and multi-vendor integrations.</li>
<li>Owned project planning and coordination for large-scale transformation programs &mdash; defining scope, timelines, resource allocation, and go-live criteria across distributed teams in Denmark and Sweden.</li>
<li>Served as primary point of contact between engineering specialists and business stakeholders &mdash; maintaining transparent communication and facilitating informed decision-making at all levels.</li>
<li>Managed healthcare-adjacent IT projects involving regulated environments, patient data systems, and compliance requirements &mdash; ensuring operational stability and adherence to regulatory standards.</li>
<li>Drove connectivity and systems integration initiatives &mdash; coordinating cloud-to-on-premise integrations, API connectivity, and data flow orchestration across complex enterprise architectures.</li>
<li>Applied ITIL and ITSM frameworks to establish incident management, change management, and problem management processes &mdash; improving service stability and reducing unplanned downtime by 40%.</li>
<li>Coordinated go-live and stabilization activities for major platform migrations &mdash; managing cutover weekends, rollback planning, and hypercare periods with engineering and operations teams.</li>
<li>Managed expectations across strategic (program steering) and operational (daily standup) levels &mdash; knowing when to escalate, when to absorb pressure, and when to say no to protect delivery quality.</li>
<li>Mentored and coordinated teams of 6&ndash;15 engineers and specialists &mdash; creating an environment where complex work could be broken into clear deliverables with shared ownership.</li>
<li>Delivered projects using agile, Prince2, and hybrid methodologies &mdash; adapting approach to project complexity, stakeholder needs, and organizational maturity.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Technical Project Delivery</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Coordinated delivery of enterprise IT projects in regulated financial environments (core banking, insurance) &mdash; managing planning, stakeholder communication, and go-live execution.</li>
<li>Acted as liaison between technical teams and business stakeholders &mdash; translating complex requirements into delivery plans and keeping projects on track through clear priorities.</li>
<li>Managed systems integration and connectivity between distributed IT systems &mdash; coordinating infrastructure provisioning, network configuration, and operational handover.</li>
<li>Applied ITIL practices for service management transitions &mdash; ensuring operational stability and clear support handover from project delivery to operations teams.</li>
</ul>

<h2>TECHNICAL FOUNDATION</h2>
<p class="tech-line"><b>Cloud Platforms:</b> GCP (Cloud Run, BigQuery, Pub/Sub, IAM), Azure (DevOps, App Services, Pipelines), AWS (EC2, S3, Lambda)</p>
<p class="tech-line"><b>Infrastructure:</b> Kubernetes, Docker, Terraform, CI/CD pipelines, networking (TCP/IP, VPN, DNS, firewalls)</p>
<p class="tech-line"><b>ITSM &amp; Frameworks:</b> ITIL v4, incident/change/problem management, service catalog, SLA management, operational runbooks</p>
<p class="tech-line"><b>Connectivity:</b> Systems integration, API management, HL7/FHIR interfaces, data flow orchestration, hybrid cloud architecture</p>
<p class="tech-line"><b>Project Methods:</b> Agile (Scrum/Kanban), SAFe, Prince2, hybrid delivery, risk-based planning, go-live governance</p>
<p class="tech-line"><b>Tools:</b> Jira, Confluence, Azure DevOps, ServiceNow, MS Project, Git, monitoring dashboards</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ITIL v4 Foundation</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
<tr><td>&bull; Six Sigma Green Belt</td><td>&bull; ISTQB Certified Tester</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent &mdash; primary working language) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
