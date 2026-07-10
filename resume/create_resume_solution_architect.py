"""Generate Solution Architect resume – Inter IKEA SCD Procurement Development."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Solution_Architect_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Solution Architect with 12+ years of experience designing, governing, and delivering enterprise-scale "
        "digital solutions within the IKEA ecosystem. Strong programming and development background combined with "
        "proven ability to translate complex business needs into simple, cost-efficient, and reliable architectures. "
        "Experienced in cloud platforms (GCP, AWS), SaaS integrations, web technologies, and system/data integration "
        "patterns — with a systematic, end-to-end approach to solution design. "
        "Skilled at communicating architecture decisions and transition roadmaps to diverse stakeholders — "
        "from technical teams to digital product leaders and business executives. "
        "Hands-on experience with Agile and DevOps ways of working, global implementation projects, "
        "and large-scale platform transformations. Passionate about technology trends and continuously exploring "
        "how emerging technologies can deliver business value in a modern enterprise landscape. "
        "Track record of collaborating with Enterprise Architects, product teams, and cross-functional stakeholders "
        "to ensure aligned, consistent solution implementations across IKEA. "
        "Entrepreneurial spirit with a deep commitment to IKEA values — simplicity, togetherness, and cost-consciousness."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Solution Architecture", "End-to-end solution design · High-level design & detailed design guidance · Current/future/transition state mapping · Architecture governance · Technology roadmaps · Cost-efficient design · Non-functional requirements"),
        ("Cloud & Platforms", "GCP (BigQuery, Cloud Run, Cloud Functions, Pub/Sub, Dataflow, GCS) · AWS (S3, Glue, Lambda) · Cloud-native patterns · SaaS integration · Platform architecture · Microservices · Serverless"),
        ("Integration & Data", "System integration patterns · Data integration (ETL/ELT, APIs, event-driven) · REST & gRPC APIs · Pub/Sub messaging · Data modelling · Integration middleware · Technical interface design"),
        ("Web & Development", "Python · TypeScript · Node.js · SQL · Java (earlier career) · Web technologies · Software development lifecycle · DevOps & CI/CD · Git · Terraform · Docker · Kubernetes"),
        ("Stakeholder Management", "Communicating architectures to diverse audiences · Influencing senior stakeholders · Cross-organizational alignment · Collaboration with Enterprise Architects · Digital product leader partnership"),
        ("Agile & Delivery", "Agile/Scrum · DevOps · Global implementation projects · Large platform delivery · Iterative development · Self-organizing teams · Continuous improvement"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS
    add_role(doc, "Solution Architect / Technical Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Design end-to-end solution architectures for digital products — translating business needs into robust, cost-efficient, and scalable technical solutions aligned with IKEA's architecture vision and guidelines.",
        "Collaborate with digital product leaders, business analysts, and delivery teams to understand business priorities, success measures, and application portfolio — designing architectures that fulfil business needs and are available when users need support.",
        "Create and communicate current state, future state, and transition state architectures — guiding detailed design and delivery of solutions while ensuring stakeholder alignment across the organization.",
        "Develop and review high-level designs supporting running business — providing guidance on service improvements, service transitions, and technology roadmap evolution.",
        "Collaborate with other Solution Architects and Enterprise Architects — ensuring products adhere to Technology Architecture guidelines and maintaining consistent solution implementations across IKEA.",
        "Proactively identify and advise on technology opportunities to deliver better business results — evaluating trends, new platforms, cost/benefit trade-offs, risks, and dependencies.",
        "Govern solution implementations — seeking understanding and agreement to ensure aligned, consistent architectures across product teams and the broader IKEA ecosystem.",
        "Design and implement system integrations and data integration patterns — connecting 20+ SaaS and on-prem sources via APIs (REST), Pub/Sub messaging, and event-driven architectures (GCP, BigQuery, Cloud Functions).",
        "Communicate solution architectures and transition roadmaps to diverse stakeholders — from engineers to senior leadership — adapting technical complexity to audience needs.",
        "Led a major platform transformation — re-architecting from batch-based processing to event-driven, real-time architecture serving 32 markets (Python, TypeScript, GCP, Terraform, Kubernetes).",
        "Support projects and initiatives taking responsibility for securing IKEA's architecture vision — balancing short-term delivery needs with long-term architectural sustainability.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Platform Engineer / Solution Designer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Designed cloud-native solution architectures for a globally distributed platform (300M+ users) — balancing scalability, reliability, cost, and performance across complex system dependencies.",
        "Worked with product teams to translate business needs into technical solutions — considering end-to-end architecture, integration patterns, and non-functional requirements.",
        "Contributed to technology evaluation and roadmap planning — assessing new cloud services and patterns for adoption in a fast-moving SaaS environment.",
        "Implemented CI/CD automation and Infrastructure as Code — improving deployment reliability and reducing operational complexity.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Technical Lead / Solution Architect", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led solution architecture for global implementation projects — designing end-to-end solutions combining cloud platforms, SaaS products, and custom-built systems for enterprise-scale operations.",
        "Designed system integration architectures connecting diverse platforms — defining technical interfaces, data flows, and integration patterns across complex, multi-vendor environments.",
        "Collaborated with Enterprise Architects to ensure solution designs adhered to Technology Architecture — governing implementations for consistency, scalability, and alignment with strategic direction.",
        "Created current state, future state, and transition roadmaps for platform modernization — guiding teams through legacy migration to cloud-native architectures (GCP, AWS).",
        "Led a team of 8–12 engineers — coaching on solution design, architecture thinking, and systematic end-to-end approaches to complex problems.",
        "Proactively identified technology opportunities and drove them into ways of working — from proof-of-concept (containerization, event-driven patterns) through organization-wide adoption.",
        "Communicated solution architectures and design decisions to senior stakeholders, business leaders, and technical teams — influencing direction across organizational boundaries.",
        "Delivered solutions within the IKEA ecosystem (IKEA App, Customer Support, Spare Parts, IoT) — deep understanding of IKEA's technology landscape and business model.",
        "Worked with SaaS platforms (Genesys, Verint, ServiceNow) — integrating third-party solutions into IKEA's enterprise architecture while maintaining governance and alignment.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_4 = [
        "Built enterprise software solutions across financial and retail domains — gaining strong programming and development background in Java, Python, SQL, and web technologies.",
        "Developed system integrations, APIs, and data pipelines connecting diverse enterprise systems — foundational experience in integration patterns and technical interface design.",
        "Participated in large-scale global implementation projects — delivering solutions within Agile frameworks across distributed teams.",
        "Contributed to platform modernization initiatives — migrating legacy systems to web-based, scalable architectures.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── IKEA Ecosystem Experience ─────────────────────────────────────────
    add_heading_block(doc, "IKEA Ecosystem Experience")
    ikea_items = [
        "IKEA Customer Connect (VCS): Solution architecture for end-to-end digital product — cloud-native platform (GCP), data integration, APIs, event-driven systems serving 32 IKEA markets.",
        "Supply Chain Adjacency: Systems integration with supply chain processes, logistics, and fulfilment systems — understanding of IKEA's integrated value chain and end-to-end approach.",
        "Architecture Governance: Collaborating with Enterprise Architects and Solution Architects — ensuring consistent, aligned implementations adhering to IKEA Technology Architecture.",
        "Platform & SaaS Landscape: Deep experience designing solutions combining Platforms, SaaS (Genesys, Verint, Salesforce), and home-grown systems within IKEA's technology landscape.",
        "Global Implementations: Solution design for products operating across 30+ markets — understanding multi-market deployment, localization, and scalability requirements.",
        "IKEA Ways of Working: Agile/Scrum, DevOps, product team model, cross-functional collaboration with digital product leaders, business analysts, and process developers.",
        "Technology Roadmaps: Driving architecture evolution and transition planning — current state → future state mapping aligned with IKEA's strategic direction.",
    ]
    for item in ikea_items:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Simplicity: Provide simple solutions to complex problems — design architectures that everyone can understand (what, how, and why) while meeting enterprise-grade requirements.",
        "Togetherness: Collaborate across teams — work closely with digital product leaders, business analysts, process developers, delivery teams, and architects to bring ideas into business value.",
        "Cost-Consciousness: Design cost-efficient and reliable digital products — keeping business priorities at the centre while balancing performance, scalability, and sustainability.",
        "Leading by Example: Share knowledge, learn every day, and take pride in contributing to the team's success — coaching engineers and raising the quality bar across the organization.",
        "Entrepreneurial Spirit: Proactively identify technology opportunities — exploring trends, evaluating new platforms, and driving innovation that delivers measurable business results.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Solution Architect with 12+ years of experience designing, governing, and delivering enterprise-scale digital solutions within the IKEA ecosystem. Strong programming and development background combined with proven ability to translate complex business needs into simple, cost-efficient, and reliable architectures. Experienced in cloud platforms (GCP, AWS), SaaS integrations, web technologies, and system/data integration patterns — with a systematic, end-to-end approach to solution design. Skilled at communicating architecture decisions and transition roadmaps to diverse stakeholders — from technical teams to digital product leaders and business executives. Hands-on experience with Agile and DevOps ways of working, global implementation projects, and large-scale platform transformations. Passionate about technology trends and continuously exploring how emerging technologies can deliver business value in a modern enterprise landscape. Track record of collaborating with Enterprise Architects, product teams, and cross-functional stakeholders to ensure aligned, consistent solution implementations across IKEA. Entrepreneurial spirit with a deep commitment to IKEA values — simplicity, togetherness, and cost-consciousness.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Solution Architecture</td><td>End-to-end solution design · High-level design &amp; detailed design guidance · Current/future/transition state mapping · Architecture governance · Technology roadmaps · Cost-efficient design · Non-functional requirements</td></tr>
<tr><td class="cat">Cloud &amp; Platforms</td><td>GCP (BigQuery, Cloud Run, Cloud Functions, Pub/Sub, Dataflow, GCS) · AWS (S3, Glue, Lambda) · Cloud-native patterns · SaaS integration · Platform architecture · Microservices · Serverless</td></tr>
<tr><td class="cat">Integration &amp; Data</td><td>System integration patterns · Data integration (ETL/ELT, APIs, event-driven) · REST &amp; gRPC APIs · Pub/Sub messaging · Data modelling · Integration middleware · Technical interface design</td></tr>
<tr><td class="cat">Web &amp; Development</td><td>Python · TypeScript · Node.js · SQL · Java (earlier career) · Web technologies · Software development lifecycle · DevOps &amp; CI/CD · Git · Terraform · Docker · Kubernetes</td></tr>
<tr><td class="cat">Stakeholder Mgmt</td><td>Communicating architectures to diverse audiences · Influencing senior stakeholders · Cross-organizational alignment · Collaboration with Enterprise Architects · Digital product leader partnership</td></tr>
<tr><td class="cat">Agile &amp; Delivery</td><td>Agile/Scrum · DevOps · Global implementation projects · Large platform delivery · Iterative development · Self-organizing teams · Continuous improvement</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Solution Architect / Technical Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Design end-to-end solution architectures for digital products — translating business needs into robust, cost-efficient, and scalable technical solutions aligned with IKEA's architecture vision and guidelines.</li>
<li>Collaborate with digital product leaders, business analysts, and delivery teams to understand business priorities, success measures, and application portfolio — designing architectures that fulfil business needs and are available when users need support.</li>
<li>Create and communicate current state, future state, and transition state architectures — guiding detailed design and delivery of solutions while ensuring stakeholder alignment across the organization.</li>
<li>Develop and review high-level designs supporting running business — providing guidance on service improvements, service transitions, and technology roadmap evolution.</li>
<li>Collaborate with other Solution Architects and Enterprise Architects — ensuring products adhere to Technology Architecture guidelines and maintaining consistent solution implementations across IKEA.</li>
<li>Proactively identify and advise on technology opportunities to deliver better business results — evaluating trends, new platforms, cost/benefit trade-offs, risks, and dependencies.</li>
<li>Govern solution implementations — seeking understanding and agreement to ensure aligned, consistent architectures across product teams and the broader IKEA ecosystem.</li>
<li>Design and implement system integrations and data integration patterns — connecting 20+ SaaS and on-prem sources via APIs (REST), Pub/Sub messaging, and event-driven architectures (GCP, BigQuery, Cloud Functions).</li>
<li>Communicate solution architectures and transition roadmaps to diverse stakeholders — from engineers to senior leadership — adapting technical complexity to audience needs.</li>
<li>Led a major platform transformation — re-architecting from batch-based processing to event-driven, real-time architecture serving 32 markets (Python, TypeScript, GCP, Terraform, Kubernetes).</li>
<li>Support projects and initiatives taking responsibility for securing IKEA's architecture vision — balancing short-term delivery needs with long-term architectural sustainability.</li>
</ul>

<p class="role">Platform Engineer / Solution Designer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Designed cloud-native solution architectures for a globally distributed platform (300M+ users) — balancing scalability, reliability, cost, and performance across complex system dependencies.</li>
<li>Worked with product teams to translate business needs into technical solutions — considering end-to-end architecture, integration patterns, and non-functional requirements.</li>
<li>Contributed to technology evaluation and roadmap planning — assessing new cloud services and patterns for adoption in a fast-moving SaaS environment.</li>
<li>Implemented CI/CD automation and Infrastructure as Code — improving deployment reliability and reducing operational complexity.</li>
</ul>

<p class="role">Technical Lead / Solution Architect <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led solution architecture for global implementation projects — designing end-to-end solutions combining cloud platforms, SaaS products, and custom-built systems for enterprise-scale operations.</li>
<li>Designed system integration architectures connecting diverse platforms — defining technical interfaces, data flows, and integration patterns across complex, multi-vendor environments.</li>
<li>Collaborated with Enterprise Architects to ensure solution designs adhered to Technology Architecture — governing implementations for consistency, scalability, and alignment with strategic direction.</li>
<li>Created current state, future state, and transition roadmaps for platform modernization — guiding teams through legacy migration to cloud-native architectures (GCP, AWS).</li>
<li>Led a team of 8–12 engineers — coaching on solution design, architecture thinking, and systematic end-to-end approaches to complex problems.</li>
<li>Proactively identified technology opportunities and drove them into ways of working — from proof-of-concept (containerization, event-driven patterns) through organization-wide adoption.</li>
<li>Communicated solution architectures and design decisions to senior stakeholders, business leaders, and technical teams — influencing direction across organizational boundaries.</li>
<li>Delivered solutions within the IKEA ecosystem (IKEA App, Customer Support, Spare Parts, IoT) — deep understanding of IKEA's technology landscape and business model.</li>
<li>Worked with SaaS platforms (Genesys, Verint, ServiceNow) — integrating third-party solutions into IKEA's enterprise architecture while maintaining governance and alignment.</li>
</ul>

<p class="role">Software Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built enterprise software solutions across financial and retail domains — gaining strong programming and development background in Java, Python, SQL, and web technologies.</li>
<li>Developed system integrations, APIs, and data pipelines connecting diverse enterprise systems — foundational experience in integration patterns and technical interface design.</li>
<li>Participated in large-scale global implementation projects — delivering solutions within Agile frameworks across distributed teams.</li>
<li>Contributed to platform modernization initiatives — migrating legacy systems to web-based, scalable architectures.</li>
</ul>

<h2>IKEA ECOSYSTEM EXPERIENCE</h2>
<ul>
<li><strong>IKEA Customer Connect (VCS):</strong> Solution architecture for end-to-end digital product — cloud-native platform (GCP), data integration, APIs, event-driven systems serving 32 IKEA markets.</li>
<li><strong>Supply Chain Adjacency:</strong> Systems integration with supply chain processes, logistics, and fulfilment systems — understanding of IKEA's integrated value chain and end-to-end approach.</li>
<li><strong>Architecture Governance:</strong> Collaborating with Enterprise Architects and Solution Architects — ensuring consistent, aligned implementations adhering to IKEA Technology Architecture.</li>
<li><strong>Platform &amp; SaaS Landscape:</strong> Deep experience designing solutions combining Platforms, SaaS (Genesys, Verint, Salesforce), and home-grown systems within IKEA's technology landscape.</li>
<li><strong>Global Implementations:</strong> Solution design for products operating across 30+ markets — understanding multi-market deployment, localization, and scalability requirements.</li>
<li><strong>IKEA Ways of Working:</strong> Agile/Scrum, DevOps, product team model, cross-functional collaboration with digital product leaders, business analysts, and process developers.</li>
<li><strong>Technology Roadmaps:</strong> Driving architecture evolution and transition planning — current state → future state mapping aligned with IKEA's strategic direction.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Simplicity:</strong> Provide simple solutions to complex problems — design architectures that everyone can understand (what, how, and why) while meeting enterprise-grade requirements.</li>
<li><strong>Togetherness:</strong> Collaborate across teams — work closely with digital product leaders, business analysts, process developers, delivery teams, and architects to bring ideas into business value.</li>
<li><strong>Cost-Consciousness:</strong> Design cost-efficient and reliable digital products — keeping business priorities at the centre while balancing performance, scalability, and sustainability.</li>
<li><strong>Leading by Example:</strong> Share knowledge, learn every day, and take pride in contributing to the team's success — coaching engineers and raising the quality bar across the organization.</li>
<li><strong>Entrepreneurial Spirit:</strong> Proactively identify technology opportunities — exploring trends, evaluating new platforms, and driving innovation that delivers measurable business results.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
