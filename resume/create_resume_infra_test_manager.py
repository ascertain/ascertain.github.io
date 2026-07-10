"""Generate Infrastructure Test Manager resume – Aeven External Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Infrastructure_Test_Manager_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Infrastructure Test Manager with 14+ years of hands-on experience in IT infrastructure testing, "
        "system integration testing, and end-to-end validation across network, identity, datacenter, "
        "endpoint/workplace, and cloud environments. Proven ability to define test scope and acceptance "
        "criteria with customers and project managers, formalize actionable test plans, coordinate complex "
        "test streams, and drive execution through to sign-off. Experienced in validating where "
        "infrastructure changes connect to the application layer. Strong communicator — reporting status, "
        "risks, and decisions clearly to stakeholders at all levels. Pragmatic, structured, and "
        "straightforward approach to testing. Comfortable working in security-sensitive environments."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Infrastructure Testing", "Network testing · Datacenter migrations · Endpoint/workplace validation · Identity & access management · Integration testing · Cloud infrastructure (AWS, GCP, Azure)"),
        ("Test Management", "Test scope definition · Test plan formalization · Acceptance criteria · Entry/exit criteria · Defect management · Retest coordination · Risk-based testing"),
        ("Coordination", "Customer collaboration · Project/programme manager alignment · Environment coordination · Resource & timeline planning · Multi-stream test coordination"),
        ("Automation & Tools", "Playwright · Python (pytest) · Selenium · Postman · Terraform · Ansible · CI/CD pipelines (GitHub Actions, Jenkins) · Infrastructure validation scripts"),
        ("Communication", "Status reporting · Risk escalation · Decision communication · Stakeholder management · Customer-facing professionalism"),
        ("Platforms & Infra", "AWS · GCP · Azure · Docker · Kubernetes · Linux · Windows Server · Active Directory · DNS · DHCP · VPN · Firewalls"),
        ("Certifications", "ISTQB Certified Tester · AWS Cloud Practitioner · Google Cloud ACE · CEH · ITIL v4 · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.5)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "F2F7FC")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - Ingka Digital
    add_role(doc, "Test Manager / Team Lead Acting", "Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Defined and agreed on test scope, test cases, and acceptance criteria with customers, project managers, and programme managers for complex infrastructure and platform deliveries across 32 markets.",
        "Formalized clear and actionable test plans covering infrastructure components (cloud services, networking, identity, integrations) and their connection to application layers.",
        "Coordinated test environments, data, resources, and timelines across multiple parallel delivery streams — ensuring dependencies were tracked and blockers resolved proactively.",
        "Followed up continuously on execution, defects, and retests — driving issue resolution together with development and infrastructure teams through daily triage.",
        "Validated impacts where infrastructure changes (cloud migrations, API gateway changes, network configurations) connected to application-level functionality.",
        "Communicated status, risks, and decisions clearly to stakeholders — providing transparent quality visibility for go/no-go decisions on infrastructure deployments.",
        "Built test automation (Playwright, Python) for infrastructure validation and end-to-end integration testing — reducing manual verification effort by 70%.",
        "Coordinated large-scale customer transitions involving infrastructure, identity, and endpoint changes across multiple environments.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Test & Release Manager", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Managed testing for infrastructure and platform changes supporting 300M+ users — defining test scope with PMs and coordinating across engineering teams.",
        "Validated infrastructure deployments (cloud services, load balancers, CDN, identity services) and their impact on application behaviour.",
        "Tracked defects, retests, and execution progress — communicating status and risks clearly to project leadership.",
        "Built CI/CD quality gates for infrastructure validation — automated smoke tests on every deployment.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Infrastructure & Integration Test Manager", "HCLTech", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Drove hands-on test management for IT infrastructure deliveries — network changes, datacenter migrations, endpoint rollouts, identity configurations, and cloud integrations for enterprise clients.",
        "Worked closely with customers, project managers, and architects to define test scope, acceptance criteria, and test cases for large-scale infrastructure transitions and transformations.",
        "Formalized test plans for complex multi-stream infrastructure projects — covering network, workplace/endpoint, identity (Active Directory, SSO), and datacenter components.",
        "Coordinated test environments, resources, and timelines across parallel workstreams — ensuring infrastructure readiness before customer acceptance.",
        "Validated where infrastructure changes connected to application layers — testing end-to-end workflows across network, identity, cloud services, and business applications.",
        "Followed up continuously on execution, defects, and retests — driving resolution with infrastructure engineers and application teams.",
        "Communicated status, risks, and decisions clearly to customers, programme managers, and steering committees.",
        "Built test automation for infrastructure validation (Python, Ansible, shell scripts) — automated network connectivity checks, DNS validation, and endpoint configuration verification.",
        "Led testing for systems with high availability requirements — ensuring infrastructure changes did not disrupt critical services.",
        "Mentored team members (8–12) on infrastructure testing practices, structured test planning, and stakeholder communication.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Test Engineer / Infrastructure QA", "HCL, Ultimate Digital, Marlabs, TekMindz", "India", "2008 – 2013")
    bullets_4 = [
        "Performed infrastructure and integration testing for enterprise systems — validating network configurations, database migrations, middleware integrations, and endpoint deployments.",
        "Defined test cases and coordinated test execution for multi-tier application and infrastructure deliveries — ensuring cross-component coverage.",
        "Tested infrastructure changes impacting application layers — validating connectivity, authentication flows, load balancing, and failover scenarios.",
        "Coordinated with customers and project managers on acceptance criteria — providing test evidence and supporting customer sign-off.",
        "Built early test automation scripts (Python, shell) for infrastructure validation — network ping checks, port scanning, service health verification.",
        "Tracked defects and testing progress — reporting status to stakeholders and driving resolution with development and infrastructure teams.",
        "Worked in both agile and waterfall delivery models — adapting testing approach to project needs.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent) · Danish (Basic — actively learning)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#1F4E79;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#1F4E79;font-size:10.5pt;border-bottom:1px solid #1F4E79;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#F2F7FC;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Infrastructure Test Manager with 14+ years of hands-on experience in IT infrastructure testing, system integration testing, and end-to-end validation across network, identity, datacenter, endpoint/workplace, and cloud environments. Proven ability to define test scope and acceptance criteria with customers and project managers, formalize actionable test plans, coordinate complex test streams, and drive execution through to sign-off. Experienced in validating where infrastructure changes connect to the application layer. Strong communicator — reporting status, risks, and decisions clearly to stakeholders at all levels. Pragmatic, structured, and straightforward approach to testing. Comfortable working in security-sensitive environments.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Infrastructure Testing</td><td>Network testing · Datacenter migrations · Endpoint/workplace validation · Identity &amp; access management · Integration testing · Cloud infrastructure (AWS, GCP, Azure)</td></tr>
<tr><td class="cat">Test Management</td><td>Test scope definition · Test plan formalization · Acceptance criteria · Entry/exit criteria · Defect management · Retest coordination · Risk-based testing</td></tr>
<tr><td class="cat">Coordination</td><td>Customer collaboration · Project/programme manager alignment · Environment coordination · Resource &amp; timeline planning · Multi-stream test coordination</td></tr>
<tr><td class="cat">Automation &amp; Tools</td><td>Playwright · Python (pytest) · Selenium · Postman · Terraform · Ansible · CI/CD pipelines (GitHub Actions, Jenkins) · Infrastructure validation scripts</td></tr>
<tr><td class="cat">Communication</td><td>Status reporting · Risk escalation · Decision communication · Stakeholder management · Customer-facing professionalism</td></tr>
<tr><td class="cat">Platforms &amp; Infra</td><td>AWS · GCP · Azure · Docker · Kubernetes · Linux · Windows Server · Active Directory · DNS · DHCP · VPN · Firewalls</td></tr>
<tr><td class="cat">Certifications</td><td>ISTQB Certified Tester · AWS Cloud Practitioner · Google Cloud ACE · CEH · ITIL v4 · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Test Manager / Team Lead Acting <span class="meta">&nbsp;|&nbsp; Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Defined and agreed on test scope, test cases, and acceptance criteria with customers, project managers, and programme managers for complex infrastructure and platform deliveries across 32 markets.</li>
<li>Formalized clear and actionable test plans covering infrastructure components (cloud services, networking, identity, integrations) and their connection to application layers.</li>
<li>Coordinated test environments, data, resources, and timelines across multiple parallel delivery streams — ensuring dependencies were tracked and blockers resolved proactively.</li>
<li>Followed up continuously on execution, defects, and retests — driving issue resolution together with development and infrastructure teams through daily triage.</li>
<li>Validated impacts where infrastructure changes (cloud migrations, API gateway changes, network configurations) connected to application-level functionality.</li>
<li>Communicated status, risks, and decisions clearly to stakeholders — providing transparent quality visibility for go/no-go decisions on infrastructure deployments.</li>
<li>Built test automation (Playwright, Python) for infrastructure validation and end-to-end integration testing — reducing manual verification effort by 70%.</li>
<li>Coordinated large-scale customer transitions involving infrastructure, identity, and endpoint changes across multiple environments.</li>
</ul>

<p class="role">Test &amp; Release Manager <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Managed testing for infrastructure and platform changes supporting 300M+ users — defining test scope with PMs and coordinating across engineering teams.</li>
<li>Validated infrastructure deployments (cloud services, load balancers, CDN, identity services) and their impact on application behaviour.</li>
<li>Tracked defects, retests, and execution progress — communicating status and risks clearly to project leadership.</li>
<li>Built CI/CD quality gates for infrastructure validation — automated smoke tests on every deployment.</li>
</ul>

<p class="role">Infrastructure &amp; Integration Test Manager <span class="meta">&nbsp;|&nbsp; HCLTech &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Drove hands-on test management for IT infrastructure deliveries — network changes, datacenter migrations, endpoint rollouts, identity configurations, and cloud integrations for enterprise clients.</li>
<li>Worked closely with customers, project managers, and architects to define test scope, acceptance criteria, and test cases for large-scale infrastructure transitions and transformations.</li>
<li>Formalized test plans for complex multi-stream infrastructure projects — covering network, workplace/endpoint, identity (Active Directory, SSO), and datacenter components.</li>
<li>Coordinated test environments, resources, and timelines across parallel workstreams — ensuring infrastructure readiness before customer acceptance.</li>
<li>Validated where infrastructure changes connected to application layers — testing end-to-end workflows across network, identity, cloud services, and business applications.</li>
<li>Followed up continuously on execution, defects, and retests — driving resolution with infrastructure engineers and application teams.</li>
<li>Communicated status, risks, and decisions clearly to customers, programme managers, and steering committees.</li>
<li>Built test automation for infrastructure validation (Python, Ansible, shell scripts) — automated network connectivity checks, DNS validation, and endpoint configuration verification.</li>
<li>Led testing for systems with high availability requirements — ensuring infrastructure changes did not disrupt critical services.</li>
<li>Mentored team members (8–12) on infrastructure testing practices, structured test planning, and stakeholder communication.</li>
</ul>

<p class="role">Test Engineer / Infrastructure QA <span class="meta">&nbsp;|&nbsp; HCL, Ultimate Digital, Marlabs, TekMindz &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Performed infrastructure and integration testing for enterprise systems — validating network configurations, database migrations, middleware integrations, and endpoint deployments.</li>
<li>Defined test cases and coordinated test execution for multi-tier application and infrastructure deliveries — ensuring cross-component coverage.</li>
<li>Tested infrastructure changes impacting application layers — validating connectivity, authentication flows, load balancing, and failover scenarios.</li>
<li>Coordinated with customers and project managers on acceptance criteria — providing test evidence and supporting customer sign-off.</li>
<li>Built early test automation scripts (Python, shell) for infrastructure validation — network ping checks, port scanning, service health verification.</li>
<li>Tracked defects and testing progress — reporting status to stakeholders and driving resolution with development and infrastructure teams.</li>
<li>Worked in both agile and waterfall delivery models — adapting testing approach to project needs.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent) · Danish (Basic — actively learning)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
