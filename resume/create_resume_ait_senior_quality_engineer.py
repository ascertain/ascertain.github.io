from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_AIT_Senior_Quality_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_AIT_Senior_Quality_Engineer_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "quality planning",
    "quality management",
    "quality assurance",
    "quality engineering",
    "QMS",
    "verification",
    "validation",
    "V&V",
    "FAT",
    "SAT",
    "UAT",
    "compliance",
    "GxP",
    "GAMP5",
    "CSV",
    "Computer System Validation",
    "Data Integrity",
    "ASTM E2500",
    "risk management",
    "risk-based",
    "deviation handling",
    "traceability",
    "DER",
    "Design Review",
    "OV",
    "Operational Verification",
    "SQP",
    "PVP",
    "automation",
    "Automation & IT",
    "cross-disciplinary",
    "cross-functional",
    "stakeholder",
    "project quality",
    "supplier quality",
    "documentation",
    "protocol",
    "IQ/OQ/PQ",
    "lifecycle",
    "SDLC",
    "CI/CD",
    "Agile",
    "Scrum",
    "IKEA",
    "LEGO",
    "Truecaller",
    "Six Sigma",
    "ITIL",
    "ISTQB",
    "CEH",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior Quality Engineer  |  Verification & Validation, Quality Planning, Compliance & QMS  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Quality engineering professional with 15+ years driving verification & "
        "validation, quality planning, and compliance across large-scale, cross-disciplinary "
        "Automation & IT projects. Develops and implements project verification strategies "
        "— FAT/SAT/UAT execution, protocol approval, deviation handling, and traceability. "
        "Ensures QMS adherence and compliance with procedures, specifications, and regulatory "
        "requirements (GxP awareness, Data Integrity, GAMP5, Computer System Validation). "
        "Manages quality activities throughout the project lifecycle — from quality planning "
        "documents (SQP, PVP) through Operational Verification and project closure. "
        "Strong stakeholder management — aligns with QE, Engineering Management, PM, "
        "and cross-functional teams. Risk management and deviation handling. "
        "Structured, result-oriented approach. ISTQB, Six Sigma Green Belt, ITIL, "
        "CEH certified. Engineering degree (M.Tech/B.Tech IT).",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Quality Planning & V&V: ",
            "15+ years developing and implementing verification & validation strategies "
            "for complex Automation & IT systems. Designed test/verification plans covering "
            "FAT, SAT, and UAT execution — protocol development, approval workflows, "
            "deviation handling, and full traceability. Currently at IKEA — built "
            "end-to-end verification frameworks across 30+ market deployments with "
            "structured quality gates and documentation.",
        ),
        (
            "Compliance & QMS: ",
            "Strong compliance background — ensured adherence to procedures, specifications, "
            "and quality standards across regulated environments (banking/finance compliance, "
            "IT security, data integrity). Familiar with GxP principles, GAMP5 concepts, "
            "and Computer System Validation methodologies. Certified Ethical Hacker (CEH) — "
            "data integrity and security awareness. Six Sigma Green Belt — process quality "
            "and continuous improvement.",
        ),
        (
            "Project Quality & Lifecycle Management: ",
            "Drove project quality activities in large, complex, cross-disciplinary projects "
            "at IKEA, Truecaller, and LEGO. Managed quality deliverables and documentation "
            "throughout the full project lifecycle — planning, execution, review, and closure. "
            "Design reviews, quality gate management, and release readiness assessments.",
        ),
        (
            "Stakeholder Management & Coordination: ",
            "Proficient in stakeholder alignment across QE, Engineering Management, PM, "
            "and cross-functional teams. Led design reviews and quality meetings. "
            "Coordinated with suppliers, partners, and external vendors. "
            "Close working relationships at all organizational levels — from engineers "
            "to senior management. Direct, open communication style.",
        ),
        (
            "Risk Management & Deviation Handling: ",
            "Extensive risk-based approach — risk assessment, risk mitigation, and "
            "deviation handling integrated into quality workflows. Root cause analysis. "
            "Ensured traceability from requirements through verification to acceptance. "
            "Managed non-conformances and corrective actions.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Quality & V&V: ",
            "Quality planning (SQP, PVP), verification & validation strategies, "
            "FAT/SAT/UAT, protocol approval, deviation handling, traceability, "
            "Operational Verification, quality gate management",
        ),
        (
            "Compliance & Regulatory: ",
            "QMS adherence, GxP awareness, GAMP5, Computer System Validation (CSV), "
            "ASTM E2500 awareness, Data Integrity, compliance testing, procedure compliance",
        ),
        (
            "Risk & Process: ",
            "Risk management, risk-based testing, root cause analysis, deviation handling, "
            "corrective/preventive actions, Six Sigma Green Belt, continuous improvement",
        ),
        (
            "Automation & IT: ",
            "Test automation (Playwright, Selenium, RestAssured, Appium), CI/CD (GitHub "
            "Actions, Jenkins), cloud platforms (GCP, AWS, Azure), Docker, Kubernetes, "
            "Terraform, Python, Java, TypeScript, SQL",
        ),
        (
            "Project & Documentation: ",
            "Design Engineering Reviews, quality documentation, test plans/protocols, "
            "release management, lifecycle management, Jira, Confluence, TestRail, Zephyr",
        ),
        (
            "Leadership & Stakeholder: ",
            "Cross-disciplinary coordination, stakeholder alignment (QE, PM, Engineering), "
            "supplier quality, team mentoring, Agile (Scrum/SAFe), ITIL",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Senior Quality Engineer / Test Strategy Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Automation & IT Platform, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Quality planning and verification strategy — developed and implemented project "
        "verification plans covering FAT, SAT, and UAT phases for a cross-disciplinary "
        "Automation & IT platform serving 30+ global markets. Ensured compliance with "
        "procedures, specifications, and quality standards. Quality gate management.",
        "Drove project quality activities — managed quality deliverables and documentation "
        "throughout the full project lifecycle. Design reviews, protocol approval, "
        "deviation handling, and traceability from requirements through verification "
        "to acceptance. Maintained quality planning documents.",
        "Risk management and compliance — risk-based approach with structured risk "
        "assessments and mitigation. Deviation handling with root cause analysis and "
        "corrective actions. Data Integrity awareness. Compliance validation. "
        "Application security testing integrated into quality workflows.",
        "Stakeholder alignment — coordinated quality activities across QE, Engineering "
        "Management, PM, and cross-functional teams (BAs, architects, developers). "
        "Supplier/partner coordination. Direct, open communication at all levels. "
        "Mentored engineers in quality practices. Agile (Scrum).",
        "Automation & IT quality — built test automation frameworks (Playwright, API "
        "automation, CI/CD) to ensure verification efficiency and repeatability. "
        "AI-assisted quality practices — 30% velocity improvement. Continuous improvement.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior QA Lead / Quality Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  IT Platform — 300M+ Users, Global Scale", bold=True, size=10)
    tc_bullets = [
        "Verification & validation at scale — developed quality strategies for an IT "
        "platform with 300M+ users. End-to-end V&V covering functional, integration, "
        "performance, security, and UAT. Protocol-driven testing with structured "
        "deviation handling and traceability.",
        "Cross-disciplinary coordination — stakeholder alignment across engineering, "
        "product, and platform teams. Risk management and compliance. Quality "
        "documentation. Release readiness assessments. Agile (Scrum/Kanban).",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Quality Lead / Senior QA Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  Enterprise & E-Commerce — Automation & IT, Multi-Platform", bold=True, size=10)
    lego_bullets = [
        "LEGO: Quality planning and verification strategy for enterprise e-commerce "
        "platform — developed verification plans, FAT/SAT/UAT protocols, and quality "
        "documentation. Led 8–10 engineers. Design reviews. Stakeholder coordination "
        "across engineering, product, and business. Supplier quality oversight. "
        "Risk management. Release management and quality gate approvals.",
        "IKEA (2018–2021): Project quality activities across multiple cross-disciplinary "
        "IT projects (IKEA App, Genesys, Verint). Verification & validation, compliance, "
        "deviation handling, and traceability. Coordinated quality deliverables across "
        "cross-functional teams. Procedure compliance. Agile (Scrum).",
        "Quality lifecycle management — managed quality activities from planning through "
        "execution, review, and closure. Connected quality outcomes with business "
        "objectives. Continuous improvement and process optimization.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior QA Engineer / Quality Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Financial Services — Regulated, Enterprise", bold=True, size=10)
    fin_bullets = [
        "Quality engineering in regulated environments — verification & validation for "
        "enterprise banking systems (Finacle). Compliance with strict regulatory "
        "requirements. Quality planning, protocol development, deviation handling, "
        "traceability. Data Integrity and audit readiness. CI/CD. SQL, Java, C#.",
        "Quality leadership — mentored 15+ engineers. Led cross-functional quality "
        "teams of 10+. Stakeholder coordination across engineering, operations, and "
        "management. Risk management. Process improvement. Six Sigma thinking. "
        "Documentation management. Agile (Scrum/Kanban).",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Verification & validation strategies — developed and implemented V&V frameworks "
        "for Automation & IT platforms at IKEA (30+ markets) and Truecaller (300M+ users). "
        "FAT/SAT/UAT execution with full traceability and quality gate management.",
        "Quality planning and compliance — quality documentation (plans, protocols, reports) "
        "throughout the project lifecycle. QMS adherence. Procedure compliance. "
        "Risk management with structured deviation handling and corrective actions.",
        "Cross-disciplinary stakeholder coordination — aligned quality activities across "
        "QE, Engineering Management, PM, and cross-functional teams at IKEA, LEGO, "
        "and Truecaller. Supplier quality oversight. Design reviews.",
        "Automation & IT quality — built test automation frameworks ensuring verification "
        "efficiency and repeatability. CI/CD integration. AI-assisted quality practices.",
        "Compliance in regulated environments — banking/finance regulatory compliance, "
        "Data Integrity, security (CEH). Six Sigma Green Belt — process quality.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – AIT Senior Quality Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior Quality Engineer | Verification &amp; Validation, Quality Planning, Compliance &amp; QMS | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Quality engineering</span> professional with 15+ years driving <span class="hl">verification</span> &amp; <span class="hl">validation</span>, <span class="hl">quality planning</span>, and <span class="hl">compliance</span> across large-scale, <span class="hl">cross-disciplinary</span> <span class="hl">Automation &amp; IT</span> projects. Develops and implements project <span class="hl">verification</span> strategies — <span class="hl">FAT</span>/<span class="hl">SAT</span>/<span class="hl">UAT</span> execution, <span class="hl">protocol</span> approval, <span class="hl">deviation handling</span>, and <span class="hl">traceability</span>. Ensures <span class="hl">QMS</span> adherence and <span class="hl">compliance</span> with procedures and specifications. <span class="hl">Risk management</span>. <span class="hl">Stakeholder</span> alignment across QE, Engineering Management, and PM. <span class="hl">ISTQB</span>, <span class="hl">Six Sigma</span> Green Belt, <span class="hl">ITIL</span>, <span class="hl">CEH</span>. Engineering degree (M.Tech/B.Tech IT).</p>

  <div class="section">How I Match the Role</div>
  <p><b>Quality Planning &amp; V&amp;V:</b> 15+ years — <span class="hl">verification</span> &amp; <span class="hl">validation</span> strategies for <span class="hl">Automation &amp; IT</span> systems. <span class="hl">FAT</span>/<span class="hl">SAT</span>/<span class="hl">UAT</span>, <span class="hl">protocol</span> approval, <span class="hl">deviation handling</span>, <span class="hl">traceability</span>. <span class="hl">Quality planning</span> <span class="hl">documentation</span>. 30+ markets at <span class="hl">IKEA</span>.<br>
  <b>Compliance &amp; QMS:</b> <span class="hl">Compliance</span> in regulated environments. <span class="hl">GxP</span> awareness, <span class="hl">GAMP5</span> concepts, <span class="hl">CSV</span> methodologies. <span class="hl">Data Integrity</span>. <span class="hl">CEH</span>. <span class="hl">Six Sigma</span> Green Belt.<br>
  <b>Project Quality &amp; Lifecycle:</b> <span class="hl">Project quality</span> in large, complex, <span class="hl">cross-disciplinary</span> projects at <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. Full <span class="hl">lifecycle</span> management — planning, execution, review, closure. Design reviews, quality gates.<br>
  <b>Stakeholder &amp; Risk:</b> <span class="hl">Stakeholder</span> alignment (QE, Engineering, PM, <span class="hl">cross-functional</span> teams). <span class="hl">Supplier quality</span>. <span class="hl">Risk management</span>, <span class="hl">deviation handling</span>, corrective actions. Led design reviews.</p>

  <div class="section">Core Competencies</div>
  <p><b>Quality &amp; V&amp;V:</b> Quality planning (SQP, PVP), V&amp;V strategies, FAT/SAT/UAT, protocol approval, deviation handling, traceability, OV, quality gates<br>
  <b>Compliance:</b> QMS, GxP awareness, GAMP5, CSV, ASTM E2500 awareness, Data Integrity, compliance testing<br>
  <b>Risk &amp; Process:</b> Risk management, risk-based testing, root cause analysis, deviation handling, CAPA, Six Sigma GB, continuous improvement<br>
  <b>Automation &amp; IT:</b> Test automation (Playwright, Selenium, RestAssured, Appium), CI/CD, cloud (GCP, AWS, Azure), Docker, Kubernetes, Python, Java, TypeScript, SQL<br>
  <b>Project &amp; Documentation:</b> Design reviews, quality documentation, test plans/protocols, release management, lifecycle management, Jira, Confluence, TestRail<br>
  <b>Leadership:</b> Cross-disciplinary coordination, stakeholder alignment, supplier quality, mentoring, Agile (Scrum/SAFe), ITIL</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Senior Quality Engineer / Test Strategy Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Automation &amp; IT Platform, 30+ Markets</div>
  <ul>
    <li><span class="hl">Quality planning</span> and <span class="hl">verification</span> strategy — <span class="hl">V&amp;V</span> plans covering <span class="hl">FAT</span>, <span class="hl">SAT</span>, <span class="hl">UAT</span> for <span class="hl">cross-disciplinary</span> <span class="hl">Automation &amp; IT</span> platform (30+ markets). <span class="hl">Compliance</span>. Quality gates.</li>
    <li><span class="hl">Project quality</span> — quality deliverables and <span class="hl">documentation</span> throughout <span class="hl">lifecycle</span>. Design reviews, <span class="hl">protocol</span> approval, <span class="hl">deviation handling</span>, <span class="hl">traceability</span>.</li>
    <li><span class="hl">Risk management</span> and <span class="hl">compliance</span> — <span class="hl">risk-based</span> approach, corrective actions, <span class="hl">Data Integrity</span> awareness. Security testing (<span class="hl">CEH</span>).</li>
    <li><span class="hl">Stakeholder</span> alignment — QE, Engineering, PM, <span class="hl">cross-functional</span> teams. <span class="hl">Supplier quality</span> oversight. Mentored engineers. <span class="hl">Agile</span> (<span class="hl">Scrum</span>).</li>
    <li><span class="hl">Automation</span> &amp; IT quality — test <span class="hl">automation</span> (Playwright, API, <span class="hl">CI/CD</span>). AI-assisted quality (30% velocity improvement). Continuous improvement.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior QA Lead / Quality Engineer</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | IT Platform — 300M+ Users, Global Scale</div>
  <ul>
    <li><span class="hl">V&amp;V</span> at scale — quality strategies for IT platform (300M+ users). Functional, integration, performance, security, <span class="hl">UAT</span>. <span class="hl">Deviation handling</span>, <span class="hl">traceability</span>.</li>
    <li><span class="hl">Cross-disciplinary</span> — <span class="hl">stakeholder</span> alignment, <span class="hl">risk management</span>, <span class="hl">compliance</span>, quality <span class="hl">documentation</span>. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Quality Lead / Senior QA Engineer</div>
  <div class="job-sub">2016 – 2021 | Enterprise &amp; E-Commerce — Automation &amp; IT, Multi-Platform</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Quality planning</span> and <span class="hl">verification</span> — <span class="hl">V&amp;V</span> plans, <span class="hl">FAT</span>/<span class="hl">SAT</span>/<span class="hl">UAT</span> <span class="hl">protocol</span>s, quality <span class="hl">documentation</span>. 8–10 engineers. Design reviews. <span class="hl">Stakeholder</span> coordination. <span class="hl">Supplier quality</span>. <span class="hl">Risk management</span>. Quality gates.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Project quality</span> across <span class="hl">cross-disciplinary</span> IT projects. <span class="hl">V&amp;V</span>, <span class="hl">compliance</span>, <span class="hl">deviation handling</span>, <span class="hl">traceability</span>. <span class="hl">Agile</span>.</li>
    <li>Quality <span class="hl">lifecycle</span> management — planning through closure. <span class="hl">Quality assurance</span> connected to business objectives. Continuous improvement.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior QA Engineer / Quality Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Financial Services — Regulated, Enterprise</div>
  <ul>
    <li><span class="hl">Quality engineering</span> in regulated environments — <span class="hl">V&amp;V</span> for enterprise banking (Finacle). <span class="hl">Compliance</span>. <span class="hl">Quality planning</span>, <span class="hl">protocol</span>s, <span class="hl">deviation handling</span>, <span class="hl">traceability</span>, <span class="hl">Data Integrity</span>, audit readiness.</li>
    <li>Quality leadership — 15+ engineers. <span class="hl">Stakeholder</span> coordination. <span class="hl">Risk management</span>. <span class="hl">Six Sigma</span> thinking. <span class="hl">Documentation</span> management. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">V&amp;V</span> strategies — <span class="hl">Automation &amp; IT</span> platforms at <span class="hl">IKEA</span> (30+ markets) and <span class="hl">Truecaller</span> (300M+ users). <span class="hl">FAT</span>/<span class="hl">SAT</span>/<span class="hl">UAT</span>. Full <span class="hl">traceability</span> and quality gates.</li>
    <li><span class="hl">Quality planning</span> and <span class="hl">compliance</span> — <span class="hl">QMS</span> adherence, <span class="hl">documentation</span>, <span class="hl">risk management</span>, <span class="hl">deviation handling</span>, corrective actions.</li>
    <li><span class="hl">Cross-disciplinary</span> <span class="hl">stakeholder</span> coordination — QE, Engineering, PM at <span class="hl">IKEA</span>, <span class="hl">LEGO</span>, <span class="hl">Truecaller</span>. <span class="hl">Supplier quality</span>. Design reviews.</li>
    <li><span class="hl">Automation</span> &amp; IT quality — test <span class="hl">automation</span> for <span class="hl">verification</span> efficiency. <span class="hl">CI/CD</span>. AI-assisted quality.</li>
    <li><span class="hl">Compliance</span> in regulated environments — banking/finance, <span class="hl">Data Integrity</span>, security (<span class="hl">CEH</span>). <span class="hl">Six Sigma</span> GB.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
