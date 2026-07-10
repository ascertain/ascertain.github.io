"""
Generate a tailored resume for Senior Quality Assurance Engineer.
Focus: system-level testing, hardware-adjacent testing, Python/pytest automation,
Linux, network testing, physical test rigs, exploratory testing, CI pipeline, IoT.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Senior_QA_Engineer_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Senior_QA_Engineer_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "system-level", "end-to-end", "exploratory testing",
    "Python", "pytest", "test automation",
    "Linux", "terminal", "hardware",
    "test rigs", "Raspberry Pi", "IoT",
    "network", "networking", "network traffic",
    "CI pipeline", "CI/CD",
    "defect reports", "root cause",
    "logs", "traces", "troubleshooting",
    "test environments", "physical devices",
    "Jira", "Confluence",
    "agile", "cross-functional",
    "acceptance criteria", "test plans",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Senior Quality Assurance Engineer", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Senior QA Engineer with 10+ years of hands-on experience in system-level, end-to-end, "
        "and exploratory testing across complex technical environments involving software, hardware, "
        "and networking layers. Strong Python automation skills (pytest) with a pragmatic approach — "
        "automating once patterns and test needs are understood, while continuously increasing coverage. "
        "Experienced in preparing and maintaining Linux-based test environments, physical test rigs, "
        "and hardware-adjacent setups including IoT devices and networked systems. Skilled at analyzing "
        "system logs, traces, and network traffic to isolate faults and produce high-quality defect "
        "reports. Collaborative team player who works closely with developers and Product Owners in "
        "agile environments, taking ownership of quality and continuously improving tools, processes, "
        "and test practices."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "System-Level & E2E Testing", "Python Test Automation (pytest)", "Linux & Terminal Proficiency",
        "Hardware-Adjacent / Physical Testing", "Exploratory & Manual Testing", "Log/Trace Analysis & Troubleshooting",
        "CI Pipeline Integration", "Test Rig Setup & Maintenance", "Agile / Cross-Functional Collaboration",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior QA Engineer / SDET — IoT & Cloud Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Design, execute, and maintain system-level, end-to-end, and release test scenarios — balancing exploratory manual testing with growing automation coverage to ensure high product confidence.",
        "Develop and evolve automated test suites in Python (pytest) — maintaining and improving the test automation framework and integrating automated tests into CI pipelines (GitHub Actions).",
        "Prepare, configure, and maintain Linux-based test environments and physical test setups — enabling reliable and repeatable testing across software, hardware, and network layers.",
        "Perform system-level and hardware-based testing: validating interactions between cloud services (GCP), IoT devices, networking components, and external hardware in real-world environments.",
        "Analyze system logs, traces, and network traffic to isolate faults and identify root causes — producing high-quality defect reports with clear reproduction steps and supporting evidence.",
        "Create and maintain test cases, test plans, and documentation in Jira and Confluence — collaborating with developers and Product Owner to clarify requirements, risks, and acceptance criteria.",
        "Continuously improve test processes, tools, and quality practices: introduced structured exploratory testing sessions, extended automation coverage by 60%, and reduced regression time significantly.",
        "Work in a cross-functional agile team (Scrum) — participating in sprint planning, backlog refinement, and demos, taking ownership of quality throughout the development lifecycle.",
        "Build and maintain custom test rigs for hardware-adjacent validation — including device simulators and network configuration for end-to-end scenario execution.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "QA & Release Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Executed system-level and end-to-end testing for a 300M+ user platform — combining exploratory testing with automated regression suites to validate release readiness.",
        "Developed Python-based automated test scripts integrated into CI pipelines — increasing automation coverage and reducing manual regression effort per release cycle.",
        "Analyzed system logs, application traces, and network behavior to troubleshoot production issues — producing detailed defect reports that accelerated developer resolution.",
        "Maintained test environments across Linux-based infrastructure (AWS) — coordinating environment provisioning and ensuring consistent test execution conditions.",
        "Collaborated closely with developers in an agile team — clarifying acceptance criteria, reviewing test coverage, and driving quality improvements.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior Test Engineer / SDET — Enterprise & IoT Platforms", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Designed and executed system-level, end-to-end, and exploratory test scenarios for IKEA IoT/smart-home platforms and LEGO digital systems — validating complex hardware-software-network interactions.",
        "Developed automated test suites in Python (pytest, Selenium) and integrated them into CI/CD pipelines (Jenkins) — gradually increasing automation coverage as test patterns matured.",
        "Prepared and maintained physical test rigs and Linux-based test environments for hardware-adjacent testing — including networked IoT devices, readers, and controller hardware.",
        "Performed hardware-based testing: validated firmware-software interactions, device connectivity, network protocols, and physical device behavior under various operational conditions.",
        "Analyzed system logs, network traffic (Wireshark, tcpdump), and device traces to isolate root causes of failures — producing high-quality defect reports with detailed reproduction steps.",
        "Created and maintained comprehensive test cases, test plans, and test documentation in Jira and Confluence — working closely with developers and Product Owners on acceptance criteria.",
        "Worked in cross-functional agile teams (Scrum/Kanban) across Denmark and Sweden — collaborating with embedded engineers, backend developers, and product roles on quality throughout the lifecycle.",
        "Drove continuous improvement in test practices: introduced exploratory testing charters, risk-based test prioritization, and automated smoke tests that reduced release feedback time by 50%.",
        "Managed and configured test environments including Linux servers, virtual networks, and device simulators — ensuring repeatable and reliable test execution for complex system scenarios.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise Systems Testing", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Engineer / Technical Lead", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Performed system-level testing for core banking platforms — validating end-to-end transaction flows, hardware integrations (biometric readers, ATM devices), and network connectivity.",
        "Developed automation scripts (Python, shell scripting) for regression testing on Linux-based banking servers — improving test coverage and reducing manual effort.",
        "Analyzed system logs and database traces to identify root causes of system failures — producing clear defect reports for development teams in regulated environments.",
        "Configured and maintained test environments across physical hardware, Linux servers, and networked banking terminals for repeatable test execution.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("Automation: ", "Python (pytest, requests, paramiko), Selenium, Playwright, shell scripting, API testing"),
        ("Systems: ", "Linux (Ubuntu, Debian — terminal, systemd, networking), Docker, VM management, system-level debugging"),
        ("Hardware/IoT: ", "Physical test rigs, Raspberry Pi setups, IoT device testing, readers/locks, firmware validation, device simulators"),
        ("Networking: ", "TCP/IP, HTTP/HTTPS, MQTT, Wireshark, tcpdump, network configuration, firewall rules, VPN"),
        ("CI/CD: ", "GitHub Actions, Jenkins, GitLab CI — automated test integration, quality gates, pipeline orchestration"),
        ("Tools: ", "Jira, Confluence, Zephyr, Git, VS Code, Postman, GCP, AWS"),
        ("Methodologies: ", "Agile (Scrum/Kanban), exploratory testing, risk-based testing, ISTQB, TDD, BDD"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
        "ITIL v4 Foundation",
        "Six Sigma Green Belt",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Senior Quality Assurance Engineer</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Senior QA Engineer with 10+ years of hands-on experience in system-level, end-to-end, and exploratory testing across complex technical environments involving software, hardware, and networking layers. Strong Python automation skills (pytest) with a pragmatic approach &mdash; automating once patterns and test needs are understood, while continuously increasing coverage. Experienced in preparing and maintaining Linux-based test environments, physical test rigs, and hardware-adjacent setups including IoT devices and networked systems. Skilled at analyzing system logs, traces, and network traffic to isolate faults and produce high-quality defect reports. Collaborative team player who works closely with developers and Product Owners in agile environments, taking ownership of quality and continuously improving tools, processes, and test practices.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; System-Level &amp; E2E Testing</td><td>&bull; Python Test Automation (pytest)</td><td>&bull; Linux &amp; Terminal Proficiency</td></tr>
<tr><td>&bull; Hardware-Adjacent / Physical Testing</td><td>&bull; Exploratory &amp; Manual Testing</td><td>&bull; Log/Trace Analysis &amp; Troubleshooting</td></tr>
<tr><td>&bull; CI Pipeline Integration</td><td>&bull; Test Rig Setup &amp; Maintenance</td><td>&bull; Agile / Cross-Functional Collaboration</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Senior QA Engineer / SDET &mdash; IoT &amp; Cloud Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Design, execute, and maintain system-level, end-to-end, and release test scenarios &mdash; balancing exploratory manual testing with growing automation coverage to ensure high product confidence.</li>
<li>Develop and evolve automated test suites in Python (pytest) &mdash; maintaining and improving the test automation framework and integrating automated tests into CI pipelines (GitHub Actions).</li>
<li>Prepare, configure, and maintain Linux-based test environments and physical test setups &mdash; enabling reliable and repeatable testing across software, hardware, and network layers.</li>
<li>Perform system-level and hardware-based testing: validating interactions between cloud services (GCP), IoT devices, networking components, and external hardware in real-world environments.</li>
<li>Analyze system logs, traces, and network traffic to isolate faults and identify root causes &mdash; producing high-quality defect reports with clear reproduction steps and supporting evidence.</li>
<li>Create and maintain test cases, test plans, and documentation in Jira and Confluence &mdash; collaborating with developers and Product Owner to clarify requirements, risks, and acceptance criteria.</li>
<li>Continuously improve test processes, tools, and quality practices: introduced structured exploratory testing sessions, extended automation coverage by 60%, and reduced regression time significantly.</li>
<li>Work in a cross-functional agile team (Scrum) &mdash; participating in sprint planning, backlog refinement, and demos, taking ownership of quality throughout the development lifecycle.</li>
<li>Build and maintain custom test rigs for hardware-adjacent validation &mdash; including device simulators and network configuration for end-to-end scenario execution.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">QA &amp; Release Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Executed system-level and end-to-end testing for a 300M+ user platform &mdash; combining exploratory testing with automated regression suites to validate release readiness.</li>
<li>Developed Python-based automated test scripts integrated into CI pipelines &mdash; increasing automation coverage and reducing manual regression effort per release cycle.</li>
<li>Analyzed system logs, application traces, and network behavior to troubleshoot production issues &mdash; producing detailed defect reports that accelerated developer resolution.</li>
<li>Maintained test environments across Linux-based infrastructure (AWS) &mdash; coordinating environment provisioning and ensuring consistent test execution conditions.</li>
<li>Collaborated closely with developers in an agile team &mdash; clarifying acceptance criteria, reviewing test coverage, and driving quality improvements.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Senior Test Engineer / SDET &mdash; Enterprise &amp; IoT Platforms <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Designed and executed system-level, end-to-end, and exploratory test scenarios for IKEA IoT/smart-home platforms and LEGO digital systems &mdash; validating complex hardware-software-network interactions.</li>
<li>Developed automated test suites in Python (pytest, Selenium) and integrated them into CI/CD pipelines (Jenkins) &mdash; gradually increasing automation coverage as test patterns matured.</li>
<li>Prepared and maintained physical test rigs and Linux-based test environments for hardware-adjacent testing &mdash; including networked IoT devices, readers, and controller hardware.</li>
<li>Performed hardware-based testing: validated firmware-software interactions, device connectivity, network protocols, and physical device behavior under various operational conditions.</li>
<li>Analyzed system logs, network traffic (Wireshark, tcpdump), and device traces to isolate root causes of failures &mdash; producing high-quality defect reports with detailed reproduction steps.</li>
<li>Created and maintained comprehensive test cases, test plans, and test documentation in Jira and Confluence &mdash; working closely with developers and Product Owners on acceptance criteria.</li>
<li>Worked in cross-functional agile teams (Scrum/Kanban) across Denmark and Sweden &mdash; collaborating with embedded engineers, backend developers, and product roles on quality throughout the lifecycle.</li>
<li>Drove continuous improvement in test practices: introduced exploratory testing charters, risk-based test prioritization, and automated smoke tests that reduced release feedback time by 50%.</li>
<li>Managed and configured test environments including Linux servers, virtual networks, and device simulators &mdash; ensuring repeatable and reliable test execution for complex system scenarios.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise Systems Testing</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Performed system-level testing for core banking platforms &mdash; validating end-to-end transaction flows, hardware integrations (biometric readers, ATM devices), and network connectivity.</li>
<li>Developed automation scripts (Python, shell scripting) for regression testing on Linux-based banking servers &mdash; improving test coverage and reducing manual effort.</li>
<li>Analyzed system logs and database traces to identify root causes of system failures &mdash; producing clear defect reports for development teams in regulated environments.</li>
<li>Configured and maintained test environments across physical hardware, Linux servers, and networked banking terminals for repeatable test execution.</li>
</ul>

<h2>TECHNICAL SKILLS</h2>
<p class="tech-line"><b>Automation:</b> Python (pytest, requests, paramiko), Selenium, Playwright, shell scripting, API testing</p>
<p class="tech-line"><b>Systems:</b> Linux (Ubuntu, Debian &mdash; terminal, systemd, networking), Docker, VM management, system-level debugging</p>
<p class="tech-line"><b>Hardware/IoT:</b> Physical test rigs, Raspberry Pi setups, IoT device testing, readers/locks, firmware validation, device simulators</p>
<p class="tech-line"><b>Networking:</b> TCP/IP, HTTP/HTTPS, MQTT, Wireshark, tcpdump, network configuration, firewall rules, VPN</p>
<p class="tech-line"><b>CI/CD:</b> GitHub Actions, Jenkins, GitLab CI &mdash; automated test integration, quality gates, pipeline orchestration</p>
<p class="tech-line"><b>Tools:</b> Jira, Confluence, Zephyr, Git, VS Code, Postman, GCP, AWS</p>
<p class="tech-line"><b>Methodologies:</b> Agile (Scrum/Kanban), exploratory testing, risk-based testing, ISTQB, TDD, BDD</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
<tr><td>&bull; ITIL v4 Foundation</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
