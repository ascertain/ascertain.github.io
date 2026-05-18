from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_TPM_Automation_Omni_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_TPM_Automation_Omni_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "technical project manager",
    "project management",
    "program management",
    "automation",
    "test automation",
    "DevOps",
    "CI/CD",
    "cloud",
    "GCP",
    "AWS",
    "omni channel",
    "omni operations",
    "product backlog",
    "roadmap",
    "OKR",
    "agile",
    "scrum",
    "SAFe",
    "stakeholder",
    "cross-functional",
    "distribution",
    "transportation management",
    "quality assurance",
    "quality",
    "testing strategy",
    "product management",
    "scope",
    "schedule",
    "budget",
    "risk",
    "multi-partner",
    "supplier",
    "compliance",
    "regulatory",
    "safety",
    "commissioning",
    "fabrication",
    "FAT",
    "inspection",
    "procurement",
    "resource allocation",
    "client",
    "Denmark",
    "offshore",
    "IKEA",
    "LEGO",
    "Truecaller",
    "Docker",
    "Kubernetes",
    "Terraform",
    "pipeline",
    "Centiro",
    "backlog",
    "benefit realization",
    "user journey",
    "workshop",
    "Six Sigma",
    "ITIL",
    "Grafana",
    "infrastructure as code",
    "scale",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Technical Project Manager  |  Automation Specialist  |  DevOps & Cloud  |  Omni-Channel Scale  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Technical Project Manager and Automation Specialist with 15+ years "
        "leading complex, multi-partner projects from development through testing, "
        "delivery, and commissioning. Combines deep DevOps and cloud expertise "
        "(GCP, AWS, Docker, Kubernetes, Terraform, CI/CD) with strong project "
        "management discipline — scope, schedule, budget, risk. Proven track record "
        "scaling Omni-Channel platforms at IKEA (30+ markets), driving product "
        "roadmaps, backlog prioritization, and benefit realization. Leads agile "
        "product teams (Scrum/SAFe), owns quality assurance and testing strategy, "
        "and collaborates closely with distribution centres, partners, and "
        "cross-functional stakeholders. Interfaces between development, engineering, "
        "and execution. Conducts FATs, inspections, and quality readiness reviews. "
        "Manages suppliers, partners, and subcontractors. Engineering degree "
        "(B.Tech IT), PG Diploma in Operations Management. Six Sigma Green Belt, "
        "ITIL. Experience working in Denmark (LEGO, Copenhagen). Willing to travel.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Technical Project & Program Management (8+ Years): ",
            "Led projects through development, testing, delivery, and commissioning "
            "across IKEA, Truecaller, and LEGO. Managed project scope, schedule, budget, "
            "and risk. Oversaw complex, multi-partner projects including budget monitoring "
            "and audit-ready reporting. Owned product roadmaps and OKRs. Aligned priorities "
            "across multiple stakeholders. Delivered projects on time and within budget "
            "across global, cross-functional teams. 8+ years in project/program leadership.",
        ),
        (
            "Automation Specialist — DevOps & Cloud: ",
            "Built and scaled test automation frameworks (Playwright, Selenium, Cypress, "
            "Appium) running thousands of tests in CI/CD pipelines. Infrastructure as code "
            "with Terraform, Docker, Kubernetes on GCP and AWS. GitHub Actions pipelines. "
            "Grafana monitoring dashboards. Automated data pipelines (BigQuery, Cloud "
            "Functions). AI-assisted automation delivering 30% velocity gain. Enabled "
            "quality assurance and testing strategy across agile product teams.",
        ),
        (
            "Omni-Channel Scale & Product Leadership: ",
            "Scaled IKEA's VCS platform across 30+ global markets — an Omni-Channel "
            "operations platform spanning in-store, online, and distribution centre "
            "touchpoints. Defined and prioritized product backlog to optimize benefit "
            "realization. Collaborated closely with distribution centres to identify "
            "needs, pain points, and opportunities. Managed dependencies across product "
            "teams. Challenged current ways of working and designed future global processes. "
            "Worked with external partners (vendor collaboration similar to Centiro model).",
        ),
        (
            "Supplier, Partner & Stakeholder Management: ",
            "Managed suppliers and partners contributing to projects — vendor selection, "
            "scope definition, performance monitoring. Primary project interface towards "
            "clients and stakeholders. Coordinated with procurement, engineering, "
            "operations, and external partners. Fostered best practices across product "
            "management discipline. Strong leadership and organizational skills.",
        ),
        (
            "Compliance, Quality & Agile Delivery: ",
            "Ensured compliance with regulatory requirements and quality processes. "
            "Quality assurance and testing strategy — communicated quality standards "
            "and product acceptance criteria. Conducted FATs and inspections. "
            "Led agile product teams (Scrum/SAFe), ensuring delivery of outcomes "
            "meets objectives and customer needs. Six Sigma Green Belt. ISTQB certified.",
        ),
        (
            "Denmark Experience & International Team: ",
            "Experience working in Denmark — LEGO Group e-commerce platform (Copenhagen, "
            "2016–2018). International, collaborative working environment across IKEA "
            "(Sweden), Truecaller (Sweden), LEGO (Denmark). Willing to travel. "
            "Danish — Conversational.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Project & Product Management: ",
            "Scope, schedule, budget, risk management. Product roadmap and OKR ownership. "
            "Backlog prioritization and benefit realization. Multi-partner coordination. "
            "Budget monitoring and audit-ready reporting. Resource allocation. "
            "Stakeholder alignment. Business case support",
        ),
        (
            "Automation & DevOps: ",
            "Test automation (Playwright, Selenium, Cypress, Appium). CI/CD pipelines "
            "(GitHub Actions, Jenkins). Infrastructure as code (Terraform). "
            "Docker, Kubernetes. Automated data pipelines. AI-assisted automation. "
            "Quality assurance and testing strategy",
        ),
        (
            "Cloud & Infrastructure: ",
            "GCP (Cloud Run, BigQuery, Cloud Functions, Secret Manager, Artifact Registry). "
            "AWS (EC2, S3, Lambda). Monitoring (Grafana, Cloud Monitoring). "
            "Infrastructure as code. Scalable platform architecture",
        ),
        (
            "Omni-Channel & Distribution: ",
            "Omni-channel operations platforms. Distribution centre collaboration. "
            "User journey mapping and workshop facilitation. Global process design. "
            "Transportation management. Dependency management across product teams. "
            "Partner collaboration (vendor integration)",
        ),
        (
            "Compliance & Quality: ",
            "Regulatory compliance, safety standards awareness, quality processes. "
            "Quality gate management. Six Sigma Green Belt. ITIL Foundation. "
            "ISTQB. Risk identification, assessment, and mitigation. FATs and inspections",
        ),
        (
            "Tools & Practices: ",
            "Jira, Confluence, MS Project, TestRail. Agile (Scrum/SAFe), Kanban. "
            "Grafana dashboards. Git, GitHub Actions. Python, Java, TypeScript, C#. "
            "BigQuery, Terraform, Docker",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(
        cp1,
        "IKEA IT AB, Malmö — Technical Project Manager / Automation Lead / Senior SDET",
        bold=True, size=10, color=SECTION_COLOR,
    )
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(
        tp1,
        "Mar 2022 – Present  |  VCS Omni-Channel Platform — 30+ Global Markets, Multi-Partner",
        bold=True, size=10,
    )
    ikea_bullets = [
        "Technical project leadership — led projects through development, testing, and "
        "commissioning for IKEA's Omni-Channel video communication platform serving 30+ "
        "markets. Managed project scope, schedule, budget, and risk. Owned product roadmap "
        "and OKRs. Aligned priorities across multiple stakeholders including distribution "
        "centres and retail operations. Multi-partner delivery coordination.",

        "Automation and DevOps — built and scaled test automation frameworks (Playwright, "
        "API testing) integrated into CI/CD pipelines (GitHub Actions). Infrastructure as "
        "code with Terraform, Docker, and Kubernetes on GCP. Automated data pipelines "
        "(BigQuery, Cloud Functions). Grafana monitoring dashboards. AI-assisted automation "
        "delivering 30% velocity improvement.",

        "Omni-channel scale — collaborated closely with distribution centres and retail "
        "teams to identify needs, opportunities, and pain points. Defined and prioritized "
        "product backlog to optimize benefit realization. Managed dependencies across "
        "product teams. Challenged current ways of working and designed future global "
        "processes and capabilities.",

        "Supplier and partner coordination — managed external vendors and partners. "
        "Defined scopes of work for subcontractors. Primary technical and project "
        "interface towards clients and stakeholders. Cross-functional collaboration "
        "with engineering, procurement, and operations.",

        "Quality assurance and compliance — enabled quality assurance and testing strategy. "
        "Communicated quality standards and product acceptance criteria. Conducted FATs, "
        "inspections, and quality readiness reviews. Risk identification, assessment, "
        "and mitigation. Audit-ready reporting.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(
        cp2,
        "Truecaller, Stockholm — Project Lead / Automation Lead / Senior QA Lead",
        bold=True, size=10, color=SECTION_COLOR,
    )
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(
        tp2,
        "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users, Global Scale",
        bold=True, size=10,
    )
    tc_bullets = [
        "Project delivery at global scale — managed project scope, schedule, and risk "
        "for a platform serving 300M+ users. Coordinated delivery activities across "
        "multiple teams and partners. Release management and commissioning. Stakeholder "
        "interface. AWS cloud infrastructure. Automation frameworks for API and mobile.",

        "Cross-functional coordination and product leadership — engineering, product, "
        "and operations teams. Risk management. Supplier coordination. Resource allocation. "
        "Backlog prioritization. Agile (Scrum/Kanban). International working environment.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(
        cp3,
        "LEGO Group & IKEA (via HCLTech) — Project Lead / Automation Lead / Technical Lead",
        bold=True, size=10, color=SECTION_COLOR,
    )
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(
        tp3,
        "2016 – 2021  |  E-Commerce & Omni-Channel — Denmark & Sweden, Multi-Partner",
        bold=True, size=10,
    )
    lego_bullets = [
        "LEGO (Denmark): Led complex, multi-partner project delivery for e-commerce "
        "platform — scope, schedule, budget, and risk management. Managed suppliers "
        "and partners. Built automation frameworks (Selenium, BDD). Coordinated "
        "development, testing, and delivery. Led 8–10 engineers. FATs, inspections, "
        "quality reviews. Client and stakeholder interface. Denmark experience "
        "(Copenhagen, 2016–2018).",

        "IKEA (2018–2021): Project leadership across omni-channel technology projects "
        "(IKEA App, Genesys, Verint/CSSP, Spartacus). Multi-partner coordination. "
        "Product backlog ownership. Compliance with quality processes. Scope definition "
        "for subcontractors. Automation strategy (Selenium, Appium, REST Assured). "
        "Resource allocation. Risk management. Agile (Scrum). Collaborated with "
        "distribution and retail operations.",

        "DevOps and cloud — CI/CD pipelines (Jenkins, GitHub Actions). Docker containerization. "
        "Cloud infrastructure. Technical interface between development, engineering, "
        "and execution. Supported procurement in defining scopes of work. Budget "
        "monitoring. Audit-ready reporting.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(
        cp4,
        "HCLTech, Samin TekMindz & Banking — Project Lead / Automation Lead",
        bold=True, size=10, color=SECTION_COLOR,
    )
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(
        tp4,
        "2008 – 2016  |  Banking & Enterprise — Regulated, Multi-Partner Delivery",
        bold=True, size=10,
    )
    fin_bullets = [
        "Project delivery in regulated environments — managed scope, schedule, budget, "
        "and risk for enterprise banking projects (Finacle). Multi-partner coordination. "
        "Compliance with strict regulatory requirements and quality processes. FATs, "
        "inspections, and quality gate management. Built automation frameworks (Selenium, "
        "UFT). Audit-ready documentation.",

        "Leadership and resource management — led teams of 10–15 engineers. Resource "
        "allocation across internal teams and subcontractors. Supplier management. "
        "Client interface. DevOps adoption (CI/CD, build automation). Cross-functional "
        "coordination. Agile and Waterfall. Six Sigma thinking.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Omni-channel scale — scaled IKEA's platform across 30+ global markets, "
        "collaborating with distribution centres and retail operations. Defined product "
        "roadmaps and OKRs. Backlog optimization for benefit realization.",

        "Automation and DevOps transformation — built CI/CD pipelines and automation "
        "frameworks from scratch at IKEA, Truecaller, and LEGO. Infrastructure as code "
        "(Terraform, Docker, Kubernetes). Cloud-native on GCP and AWS. 30% velocity "
        "improvement through AI-assisted automation.",

        "Multi-partner project delivery — IKEA (30+ markets), Truecaller (300M+ users), "
        "LEGO (Denmark). Managed scope, schedule, budget, and risk. On-time, within-budget "
        "delivery. Supplier and stakeholder management.",

        "Quality and compliance — FATs, inspections, quality readiness reviews. "
        "Testing strategy ownership. Regulatory compliance. Risk management. "
        "Six Sigma Green Belt. Audit-ready reporting.",

        "Denmark experience — LEGO Group (Copenhagen, 2016–2018). International, "
        "collaborative working environment. Danish — Conversational. Willing to travel.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "Six Sigma Green Belt\n"
        "ITIL Foundation\n"
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Danish — Conversational  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – TPM Automation Omni Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Technical Project Manager | Automation Specialist | DevOps &amp; Cloud | Omni-Channel Scale | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Technical <span class="hl">Project Manager</span> and <span class="hl">Automation</span> Specialist with 15+ years leading complex, <span class="hl">multi-partner</span> projects from development through <span class="hl">testing</span>, delivery, and <span class="hl">commissioning</span>. Combines deep <span class="hl">DevOps</span> and <span class="hl">cloud</span> expertise (<span class="hl">GCP</span>, <span class="hl">AWS</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">Terraform</span>, <span class="hl">CI/CD</span>) with strong project management — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. Proven track record scaling <span class="hl">Omni-Channel</span> platforms at <span class="hl">IKEA</span> (30+ markets). Drives product <span class="hl">roadmap</span>s, <span class="hl">backlog</span> prioritization, and <span class="hl">benefit realization</span>. Leads <span class="hl">agile</span> product teams (<span class="hl">Scrum</span>/<span class="hl">SAFe</span>). <span class="hl">Quality assurance</span> and <span class="hl">testing strategy</span> ownership. Collaborates with <span class="hl">distribution</span> centres, partners, <span class="hl">cross-functional</span> <span class="hl">stakeholder</span>s. Engineering degree, <span class="hl">Six Sigma</span> GB, <span class="hl">ITIL</span>. <span class="hl">Denmark</span> experience (<span class="hl">LEGO</span>).</p>

  <div class="section">How I Match the Role</div>
  <p><b>Technical Project &amp; Program Management (8+ Years):</b> Led projects through development, <span class="hl">testing</span>, delivery, <span class="hl">commissioning</span> at <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. <span class="hl">Scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. Owned product <span class="hl">roadmap</span>s and <span class="hl">OKR</span>s. <span class="hl">Multi-partner</span>. Audit-ready reporting.</p>
  <p><b>Automation Specialist — DevOps &amp; Cloud:</b> Built <span class="hl">test automation</span> frameworks (Playwright, Selenium, Cypress, Appium). <span class="hl">CI/CD</span> <span class="hl">pipeline</span>s. <span class="hl">Infrastructure as code</span> (<span class="hl">Terraform</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). <span class="hl">GCP</span>/<span class="hl">AWS</span>. <span class="hl">Grafana</span> monitoring. AI-assisted <span class="hl">automation</span> — 30% velocity gain.</p>
  <p><b>Omni-Channel Scale &amp; Product Leadership:</b> Scaled <span class="hl">IKEA</span>'s platform across 30+ markets — <span class="hl">Omni-Channel</span> <span class="hl">operations</span>. Defined <span class="hl">product backlog</span>. Collaborated with <span class="hl">distribution</span> centres. <span class="hl">Benefit realization</span>. Managed dependencies across product teams. Designed future global processes. Partner collaboration.</p>
  <p><b>Supplier &amp; Stakeholder Management:</b> <span class="hl">Supplier</span>/partner management. <span class="hl">Scope</span> definition. <span class="hl">Client</span> interface. <span class="hl">Procurement</span> support. Fostered best practices across <span class="hl">product management</span>. <span class="hl">Cross-functional</span> collaboration.</p>
  <p><b>Compliance, Quality &amp; Agile Delivery:</b> <span class="hl">Regulatory</span> <span class="hl">compliance</span>. <span class="hl">Quality assurance</span> and <span class="hl">testing strategy</span>. <span class="hl">FAT</span>s, <span class="hl">inspection</span>s. Led <span class="hl">agile</span> product teams. <span class="hl">Six Sigma</span> GB. ISTQB.</p>
  <p><b>Denmark &amp; International:</b> <span class="hl">LEGO</span> (<span class="hl">Denmark</span>, Copenhagen 2016–18). <span class="hl">IKEA</span> (Sweden). <span class="hl">Truecaller</span> (Sweden). Danish — Conversational. Willing to travel.</p>

  <div class="section">Core Competencies</div>
  <p><b>Project &amp; Product:</b> Scope, schedule, budget, risk. Roadmap &amp; OKR ownership. Backlog prioritization. Benefit realization. Multi-partner. Audit-ready reporting. Resource allocation. Business case support<br>
  <b>Automation &amp; DevOps:</b> Test automation (Playwright, Selenium, Cypress, Appium). CI/CD (GitHub Actions, Jenkins). Terraform. Docker, Kubernetes. AI-assisted automation. Testing strategy<br>
  <b>Cloud:</b> GCP (Cloud Run, BigQuery, Cloud Functions). AWS (EC2, S3, Lambda). Grafana monitoring. Infrastructure as code. Scalable architecture<br>
  <b>Omni-Channel:</b> Omni-channel operations. Distribution centre collaboration. User journey mapping. Workshop facilitation. Global process design. Transportation management. Partner integration<br>
  <b>Compliance &amp; Quality:</b> Regulatory compliance, safety. Quality gates. Six Sigma GB, ITIL, ISTQB. Risk management. FATs &amp; inspections<br>
  <b>Tools:</b> Jira, Confluence, MS Project, TestRail. Agile (Scrum/SAFe), Kanban. Grafana. Git, GitHub Actions. Python, Java, TypeScript, C#</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Technical Project Manager / Automation Lead / Senior SDET</div>
  <div class="job-sub">Mar 2022 – Present | VCS Omni-Channel Platform — 30+ Global Markets, Multi-Partner</div>
  <ul>
    <li>Technical <span class="hl">project</span> leadership — development, <span class="hl">testing</span>, <span class="hl">commissioning</span> for <span class="hl">Omni-Channel</span> platform (30+ markets). <span class="hl">Scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Roadmap</span> &amp; <span class="hl">OKR</span>s. <span class="hl">Stakeholder</span> alignment. <span class="hl">Distribution</span> centre collaboration.</li>
    <li><span class="hl">Automation</span> &amp; <span class="hl">DevOps</span> — <span class="hl">test automation</span> (Playwright, API). <span class="hl">CI/CD</span> (GitHub Actions). <span class="hl">Terraform</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span> on <span class="hl">GCP</span>. Automated data <span class="hl">pipeline</span>s (BigQuery). <span class="hl">Grafana</span> monitoring. AI-assisted — 30% velocity gain.</li>
    <li><span class="hl">Omni-channel</span> <span class="hl">scale</span> — <span class="hl">product backlog</span> &amp; <span class="hl">benefit realization</span>. <span class="hl">Distribution</span> centre needs &amp; pain points. Dependency management. Designed future global processes.</li>
    <li><span class="hl">Supplier</span>/partner coordination. Scopes of work. <span class="hl">Client</span>/<span class="hl">stakeholder</span> interface. <span class="hl">Cross-functional</span> (<span class="hl">engineering</span>, <span class="hl">procurement</span>, operations).</li>
    <li><span class="hl">Quality assurance</span> — <span class="hl">testing strategy</span>. <span class="hl">FAT</span>s, <span class="hl">inspection</span>s, readiness reviews. <span class="hl">Risk</span> management. <span class="hl">Compliance</span>. Audit-ready reporting.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Project Lead / Automation Lead / Senior QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users, Global Scale</div>
  <ul>
    <li><span class="hl">Project</span> delivery at global <span class="hl">scale</span> — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">risk</span> for 300M+ users. <span class="hl">Multi-partner</span> coordination. Release/<span class="hl">commissioning</span>. <span class="hl">Stakeholder</span> interface. <span class="hl">AWS</span>. <span class="hl">Automation</span> frameworks.</li>
    <li><span class="hl">Cross-functional</span> &amp; product leadership — <span class="hl">backlog</span> prioritization. <span class="hl">Risk</span> management. <span class="hl">Supplier</span> coordination. <span class="hl">Resource allocation</span>. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Project Lead / Automation Lead / Technical Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Omni-Channel — Denmark &amp; Sweden, Multi-Partner</div>
  <ul>
    <li><span class="hl">LEGO</span> (<span class="hl">Denmark</span>): <span class="hl">Multi-partner</span> <span class="hl">project</span> delivery — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Supplier</span>/partner management. <span class="hl">Automation</span> (Selenium, BDD). <span class="hl">FAT</span>s, <span class="hl">inspection</span>s. <span class="hl">Client</span> interface. Copenhagen 2016–18.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Omni-channel</span> projects (App, Genesys, Spartacus). <span class="hl">Product backlog</span>. <span class="hl">Automation</span> (Selenium, Appium). <span class="hl">Distribution</span>/retail collaboration. <span class="hl">Multi-partner</span>. <span class="hl">Agile</span>.</li>
    <li><span class="hl">DevOps</span> &amp; <span class="hl">cloud</span> — <span class="hl">CI/CD</span> (Jenkins, GitHub Actions). <span class="hl">Docker</span>. <span class="hl">Procurement</span> support. <span class="hl">Budget</span> monitoring. <span class="hl">Cross-functional</span>.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Project Lead / Automation Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Enterprise — Regulated, Multi-Partner Delivery</div>
  <ul>
    <li><span class="hl">Project</span> delivery — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Multi-partner</span>. <span class="hl">Regulatory</span> <span class="hl">compliance</span>/<span class="hl">quality</span>. <span class="hl">FAT</span>s, <span class="hl">quality</span> gates. <span class="hl">Automation</span> (Selenium, UFT). Audit-ready.</li>
    <li>Led 10–15 engineers. <span class="hl">Resource allocation</span>. <span class="hl">Supplier</span> management. <span class="hl">Client</span> interface. <span class="hl">DevOps</span> adoption (<span class="hl">CI/CD</span>). <span class="hl">Cross-functional</span>. <span class="hl">Six Sigma</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Omni-channel</span> <span class="hl">scale</span> — <span class="hl">IKEA</span> 30+ markets. <span class="hl">Distribution</span> centre collaboration. <span class="hl">Roadmap</span>/<span class="hl">OKR</span>s. <span class="hl">Backlog</span> optimization. <span class="hl">Benefit realization</span>.</li>
    <li><span class="hl">Automation</span> &amp; <span class="hl">DevOps</span> transformation — <span class="hl">CI/CD</span>, <span class="hl">Terraform</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>. <span class="hl">Cloud</span>-native (<span class="hl">GCP</span>/<span class="hl">AWS</span>). 30% velocity via AI-assisted <span class="hl">automation</span>.</li>
    <li><span class="hl">Multi-partner</span> delivery — <span class="hl">IKEA</span>, <span class="hl">Truecaller</span> (300M+), <span class="hl">LEGO</span> (<span class="hl">Denmark</span>). On-time, within <span class="hl">budget</span>. <span class="hl">Supplier</span>/<span class="hl">stakeholder</span> management.</li>
    <li><span class="hl">Quality</span>/<span class="hl">compliance</span> — <span class="hl">FAT</span>s, <span class="hl">inspection</span>s. <span class="hl">Testing strategy</span>. <span class="hl">Risk</span> management. <span class="hl">Six Sigma</span> GB. Audit-ready.</li>
    <li><span class="hl">Denmark</span> — <span class="hl">LEGO</span> (Copenhagen 2016–18). International. Danish — Conversational. Willing to travel.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>Six Sigma Green Belt<br>ITIL Foundation<br>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Danish — Conversational | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
