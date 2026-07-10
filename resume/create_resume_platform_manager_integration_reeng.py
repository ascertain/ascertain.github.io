"""
Resume: Platform Manager – Integration Reengineering, Inter IKEA (INTERNAL)
Focus: Modernisation of legacy integration, Common Technology Foundation adoption,
       platform-as-a-product, decommissioning/uplifting, people leadership, DevOps,
       cross-unit coordination, developer experience.
NOTE: User has 8 years in IKEA ecosystem, can code but not a hardcore developer.
      Frame as a technology leader who bridges platform engineering with transformation leadership.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Platform_Manager_Integration_Reengineering_Resume"
IKEA_BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = IKEA_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = IKEA_BLUE
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = IKEA_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Technology transformation leader with 8+ years in the IKEA ecosystem, experienced in driving integration "
        "platform modernisation — from decommissioning legacy landscapes to adopting cloud-native, event-driven "
        "architectures. Combines practical platform engineering knowledge with strong people leadership and a "
        "platform-as-a-product mindset where success is measured by the velocity and autonomy of consuming teams. "
        "Proven at coordinating cross-unit dependencies, aligning roadmaps with enterprise strategy, and communicating "
        "complex technical transitions in simple business language. Energised by change, motivated by learning, and "
        "grounded in IKEA values — building trust, collaborating across boundaries, and influencing without formal authority."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Integration Modernisation & Reeng.", "Platform-as-a-Product Leadership", "Legacy Decommissioning Strategy",
        "People Leadership & Team Growth", "Cross-Unit Dependency Coordination", "DevOps & Developer Experience",
        "API-First & Event-Driven Design", "Roadmap & Priority Communication", "Budget & Cost Optimisation",
        "IKEA Ecosystem (8+ years)", "Agile / Scaled Delivery / OKRs", "Cloud, Hybrid & On-Prem Platforms",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TECHNICAL & FRAMEWORK KNOWLEDGE ──
    add_heading_block(doc, "Technical & Framework Knowledge")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    skills = (
        "Integration Patterns: Event-driven architecture, API gateways, strangler-fig (legacy-to-cloud), data replication, Pub/Sub  •  "
        "Cloud & Hybrid: GCP (Cloud Run, Pub/Sub, BigQuery, GKE), hybrid cloud/on-prem environments  •  "
        "DevOps & Automation: GitHub Actions, Docker, Kubernetes, Terraform, CI/CD pipeline design  •  "
        "API Management: OpenAPI, API-first design, versioning, contract governance  •  "
        "Development: TypeScript, Python, Node.js; understanding of Java ecosystem  •  "
        "Frameworks: Agile (Scrum, SAFe awareness), DevOps, ITIL awareness, TOGAF familiarity  •  "
        "Metrics & Governance: OKRs, platform adoption metrics, delivery velocity KPIs  •  "
        "IKEA Context: Franchise model, Inter IKEA / Ingka operating model, Common Technology Foundation"
    )
    r = tp.add_run(skills)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- Ingka Digital / IKEA ---
    role_header(doc, "Team Lead & Platform Engineer – Customer Connect Integration",
                "Ingka Digital / IKEA", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Lead the full platform lifecycle for integration services connecting 5+ IKEA customer platforms — from strategy and architecture through to delivery, operations, and continuous improvement.")
    bullet(doc, "Drive integration modernisation: decommissioning legacy point-to-point connections and uplifting to event-driven architecture (GCP Pub/Sub) and API-first patterns aligned with Common Technology Foundation.")
    bullet(doc, "Apply platform-as-a-product thinking: measure platform success by the velocity and autonomy of consuming product teams; provide self-service tooling, documentation, and adoption guidance.")
    bullet(doc, "Coordinate cross-unit dependencies: align integration roadmaps with enterprise strategies, manage shared timelines, and communicate priorities clearly to technology networks and business stakeholders.")
    bullet(doc, "Guide and empower product teams on their modernisation journey — providing architectural guidance, migration playbooks, and hands-on support during strangler-pattern transitions.")
    bullet(doc, "Foster DevOps practices and automation: CI/CD pipelines (GitHub Actions), infrastructure-as-code (Terraform), containerised deployments — improving delivery speed and reliability.")
    bullet(doc, "Lead and grow a high-performing engineering team: competence development, succession planning, innovation time, and creating conditions where curiosity and collaboration thrive.")
    bullet(doc, "Own budget planning and cost monitoring; optimise cloud and platform spend while delivering on modernisation commitments and operational SLAs.")
    bullet(doc, "Drive developer experience improvements: better onboarding, clearer API documentation, self-service environment provisioning — reducing friction for teams adopting the platform.")
    bullet(doc, "Communicate platform roadmaps and transition plans to architects, product leaders, and core business in simple, accessible language — building trust across organisational boundaries.")

    # --- Truecaller ---
    role_header(doc, "Platform Engineer",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Contributed to platform modernisation in a cloud-native SaaS environment; experienced event-driven architectures and distributed system operations at high scale.")
    bullet(doc, "Gained perspective on how platform teams empower product engineering through self-service, automation, and clear developer experience.")

    # --- HCLTech for IKEA ---
    role_header(doc, "Senior Engineer & Integration Lead – IKEA Enterprise Platforms",
                "HCLTech (for IKEA)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Spent 8 years in IKEA's integration landscape across multiple domains (supply chain, e-commerce, order management, customer platforms) — building deep understanding of the legacy landscape and its modernisation opportunities.")
    bullet(doc, "Led integration reengineering initiatives: assessed legacy systems for decommissioning, designed migration paths to modern patterns (API-first, event-based), and executed transitions with minimal business disruption.")
    bullet(doc, "Managed a team of 8–12 engineers (onshore & offshore): people development, skills assessment, hiring, succession planning — growing capability aligned with future technology needs.")
    bullet(doc, "Designed and delivered integration architectures connecting 10+ enterprise systems using strangler-fig patterns to incrementally modernise while maintaining operational continuity.")
    bullet(doc, "Coordinated dependencies across multiple product and platform teams; aligned activity plans, managed shared environments, and ensured consistent approaches across IKEA's integration ecosystem.")
    bullet(doc, "Owned service management for integration platforms: incident management, SLA governance, operational reporting — applying ITIL practices pragmatically.")
    bullet(doc, "Drove Agile delivery using OKRs and metrics to track value; championed continuous improvement, test-and-learn approaches, and iterative delivery.")
    bullet(doc, "Built trusted stakeholder relationships across Inter IKEA and Ingka — influencing decisions through collaboration rather than formal authority, aligned with IKEA values.")
    bullet(doc, "Introduced new integration technologies when appropriate; evaluated emerging platforms, ran pilots at small scale, and scaled successes into production adoption.")

    # --- India ---
    role_header(doc, "Software Engineer",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Developed enterprise applications and integration components (Java, .NET); gained foundational experience in service-oriented architecture and software engineering at scale.")
    bullet(doc, "Worked across full software lifecycle in globally distributed delivery models; built foundation in systematic problem-solving and cross-team collaboration.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "Google Cloud Associate Cloud Engineer",
        "AWS Certified Cloud Practitioner",
        "ISTQB Certified Tester",
        "Six Sigma Green Belt (Process Improvement)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("PGDOM")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – IGNOU  |  ").font.size = Pt(9)
    r2 = p.add_run("B.Tech Information Technology")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – UP Technical University").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Conversational)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    # ── WHY THIS ROLE ──
    add_heading_block(doc, "Why This Role")
    why = (
        "After 8+ years living and breathing IKEA's integration landscape, I've seen first-hand where legacy friction "
        "slows teams down and where modernisation unlocks genuine velocity. I'm energised by the opportunity to lead "
        "integration reengineering at scale — defining the path forward for our legacy landscape while empowering "
        "product teams to adopt the Common Technology Foundation with confidence. I believe platform success is measured "
        "by the autonomy it creates for others, and I'm passionate about building the conditions where teams can move "
        "faster, integrate simpler, and focus on delivering value rather than fighting infrastructure. This role brings "
        "together my integration depth, people leadership, and drive for modernisation — grounded in the IKEA values "
        "of togetherness, simplicity, and always finding a better way."
    )
    pw = doc.add_paragraph(why)
    pw.paragraph_format.space_before = Pt(3)
    for run in pw.runs:
        run.font.size = Pt(9)
        run.italic = True

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#0051BA;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#0051BA;font-size:11pt;border-bottom:1px solid #0051BA;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#0051BA;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    .why{font-style:italic;font-size:9pt;margin-top:4px}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Technology transformation leader with 8+ years in the IKEA ecosystem, experienced in driving integration
platform modernisation — from decommissioning legacy landscapes to adopting cloud-native, event-driven
architectures. Combines practical platform engineering knowledge with strong people leadership and a
platform-as-a-product mindset where success is measured by the velocity and autonomy of consuming teams.
Proven at coordinating cross-unit dependencies, aligning roadmaps with enterprise strategy, and communicating
complex technical transitions in simple business language. Energised by change, motivated by learning, and
grounded in IKEA values — building trust, collaborating across boundaries, and influencing without formal authority.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Integration Modernisation &amp; Reeng.</td><td>✓ Platform-as-a-Product Leadership</td><td>✓ Legacy Decommissioning Strategy</td></tr>
<tr><td>✓ People Leadership &amp; Team Growth</td><td>✓ Cross-Unit Dependency Coordination</td><td>✓ DevOps &amp; Developer Experience</td></tr>
<tr><td>✓ API-First &amp; Event-Driven Design</td><td>✓ Roadmap &amp; Priority Communication</td><td>✓ Budget &amp; Cost Optimisation</td></tr>
<tr><td>✓ IKEA Ecosystem (8+ years)</td><td>✓ Agile / Scaled Delivery / OKRs</td><td>✓ Cloud, Hybrid &amp; On-Prem Platforms</td></tr>
</table>

<h2>TECHNICAL &amp; FRAMEWORK KNOWLEDGE</h2>
<p class="tools">Integration Patterns: Event-driven architecture, API gateways, strangler-fig (legacy-to-cloud), data replication, Pub/Sub &bull;
Cloud &amp; Hybrid: GCP (Cloud Run, Pub/Sub, BigQuery, GKE), hybrid cloud/on-prem environments &bull;
DevOps &amp; Automation: GitHub Actions, Docker, Kubernetes, Terraform, CI/CD pipeline design &bull;
API Management: OpenAPI, API-first design, versioning, contract governance &bull;
Development: TypeScript, Python, Node.js; understanding of Java ecosystem &bull;
Frameworks: Agile (Scrum, SAFe awareness), DevOps, ITIL awareness, TOGAF familiarity &bull;
Metrics &amp; Governance: OKRs, platform adoption metrics, delivery velocity KPIs &bull;
IKEA Context: Franchise model, Inter IKEA / Ingka operating model, Common Technology Foundation</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Team Lead &amp; Platform Engineer – Customer Connect Integration &nbsp;|&nbsp; Ingka Digital / IKEA, Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Lead the full platform lifecycle for integration services connecting 5+ IKEA customer platforms — from strategy and architecture through to delivery, operations, and continuous improvement.</li>
<li>Drive integration modernisation: decommissioning legacy point-to-point connections and uplifting to event-driven architecture (GCP Pub/Sub) and API-first patterns aligned with Common Technology Foundation.</li>
<li>Apply platform-as-a-product thinking: measure platform success by the velocity and autonomy of consuming product teams; provide self-service tooling, documentation, and adoption guidance.</li>
<li>Coordinate cross-unit dependencies: align integration roadmaps with enterprise strategies, manage shared timelines, and communicate priorities clearly to technology networks and business stakeholders.</li>
<li>Guide and empower product teams on their modernisation journey — providing architectural guidance, migration playbooks, and hands-on support during strangler-pattern transitions.</li>
<li>Foster DevOps practices and automation: CI/CD pipelines (GitHub Actions), infrastructure-as-code (Terraform), containerised deployments — improving delivery speed and reliability.</li>
<li>Lead and grow a high-performing engineering team: competence development, succession planning, innovation time, and creating conditions where curiosity and collaboration thrive.</li>
<li>Own budget planning and cost monitoring; optimise cloud and platform spend while delivering on modernisation commitments and operational SLAs.</li>
<li>Drive developer experience improvements: better onboarding, clearer API documentation, self-service environment provisioning — reducing friction for teams adopting the platform.</li>
<li>Communicate platform roadmaps and transition plans to architects, product leaders, and core business in simple, accessible language — building trust across organisational boundaries.</li>
</ul>

<p class="role">Platform Engineer &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Contributed to platform modernisation in a cloud-native SaaS environment; experienced event-driven architectures and distributed system operations at high scale.</li>
<li>Gained perspective on how platform teams empower product engineering through self-service, automation, and clear developer experience.</li>
</ul>

<p class="role">Senior Engineer &amp; Integration Lead – IKEA Enterprise Platforms &nbsp;|&nbsp; HCLTech (for IKEA), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Spent 8 years in IKEA's integration landscape across multiple domains (supply chain, e-commerce, order management, customer platforms) — building deep understanding of the legacy landscape and its modernisation opportunities.</li>
<li>Led integration reengineering initiatives: assessed legacy systems for decommissioning, designed migration paths to modern patterns (API-first, event-based), and executed transitions with minimal business disruption.</li>
<li>Managed a team of 8–12 engineers (onshore &amp; offshore): people development, skills assessment, hiring, succession planning — growing capability aligned with future technology needs.</li>
<li>Designed and delivered integration architectures connecting 10+ enterprise systems using strangler-fig patterns to incrementally modernise while maintaining operational continuity.</li>
<li>Coordinated dependencies across multiple product and platform teams; aligned activity plans, managed shared environments, and ensured consistent approaches across IKEA's integration ecosystem.</li>
<li>Owned service management for integration platforms: incident management, SLA governance, operational reporting — applying ITIL practices pragmatically.</li>
<li>Drove Agile delivery using OKRs and metrics to track value; championed continuous improvement, test-and-learn approaches, and iterative delivery.</li>
<li>Built trusted stakeholder relationships across Inter IKEA and Ingka — influencing decisions through collaboration rather than formal authority, aligned with IKEA values.</li>
<li>Introduced new integration technologies when appropriate; evaluated emerging platforms, ran pilots at small scale, and scaled successes into production adoption.</li>
</ul>

<p class="role">Software Engineer &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Developed enterprise applications and integration components (Java, .NET); gained foundational experience in service-oriented architecture and software engineering at scale.</li>
<li>Worked across full software lifecycle in globally distributed delivery models; built foundation in systematic problem-solving and cross-team collaboration.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>Google Cloud Associate Cloud Engineer</li>
<li>AWS Certified Cloud Practitioner</li>
<li>ISTQB Certified Tester</li>
<li>Six Sigma Green Belt (Process Improvement)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>PGDOM</b> – IGNOU &nbsp;|&nbsp; <b>B.Tech Information Technology</b> – UP Technical University</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Conversational) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

<h2>WHY THIS ROLE</h2>
<p class="why">After 8+ years living and breathing IKEA's integration landscape, I've seen first-hand where legacy friction
slows teams down and where modernisation unlocks genuine velocity. I'm energised by the opportunity to lead
integration reengineering at scale — defining the path forward for our legacy landscape while empowering
product teams to adopt the Common Technology Foundation with confidence. I believe platform success is measured
by the autonomy it creates for others, and I'm passionate about building the conditions where teams can move
faster, integrate simpler, and focus on delivering value rather than fighting infrastructure. This role brings
together my integration depth, people leadership, and drive for modernisation — grounded in the IKEA values
of togetherness, simplicity, and always finding a better way.</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
