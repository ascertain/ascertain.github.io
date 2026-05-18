from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Media_SDET_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Media_SDET_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "WebRTC",
    "SIP",
    "media",
    "real-time communication",
    "test automation framework",
    "automation framework",
    "test automation",
    "shift-left",
    "quality strategy",
    "C#",
    "Java",
    "TypeScript",
    "GCP",
    "AWS",
    "cloud",
    "VDI",
    "Windows 10",
    "Windows 11",
    "CI/CD",
    "Git",
    "code review",
    "mentor",
    "cross-functional",
    "cross-team",
    "integration",
    "scalable",
    "end-to-end",
    "E2E",
    "API",
    "regression",
    "performance",
    "Playwright",
    "Selenium",
    "Docker",
    "Kubernetes",
    "Terraform",
    "GitHub Actions",
    "Jenkins",
    "Agile",
    "Scrum",
    "IKEA",
    "LEGO",
    "Truecaller",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Media SDET / Technical Lead  |  WebRTC, SIP, Test Automation Frameworks, C#  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Media SDET / Technical Lead with 15+ years and hands-on experience with "
        "real-time communication systems (WebRTC, SIP). Owns and drives quality "
        "strategy for media backend and frontend components. Designs, implements, "
        "and maintains robust, scalable test automation frameworks (C#, Java, "
        "TypeScript). Validates and troubleshoots WebRTC/SIP communication flows "
        "— audio, video, screen sharing, call routing, codec negotiation. "
        "Shift-left testing advocate — collaborates with developers, PMs, and QA "
        "engineers to embed testing early. Supports automated testing across VDI, "
        "Cloud (GCP/AWS), and Windows 10/11 environments. CI/CD pipelines, code "
        "reviews, mentors SDETs. Cross-functional coordination across multiple "
        "teams and products. Currently at IKEA IT AB (3+ years) building media "
        "communication platform with WebRTC/SIP. Also Truecaller and LEGO.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "WebRTC & SIP — Media Quality Ownership: ",
            "Hands-on experience with WebRTC and SIP real-time communication systems "
            "at IKEA — validates and troubleshoots audio/video calls, screen sharing, "
            "call routing, codec negotiation, SRTP, TURN/STUN, ICE connectivity, and "
            "SIP trunk integration. Owns quality strategy for media backend and "
            "frontend components. End-to-end media flow validation across browsers "
            "and environments.",
        ),
        (
            "Test Automation Frameworks — C#, Java, TypeScript: ",
            "Designs, implements, and maintains robust, scalable test automation "
            "frameworks. Strong programming: C#, Java, TypeScript. Currently at "
            "IKEA — Playwright (TypeScript) E2E and API automation framework for "
            "media platform. Previous: Selenium/RestAssured (Java, C#) automation "
            "frameworks at LEGO and banking. Framework design patterns (POM, BDD). "
            "Code reviews and mentors other SDETs.",
        ),
        (
            "Shift-Left & Quality Strategy: ",
            "Drives shift-left testing practices — embeds quality early in the SDLC. "
            "Collaborates with developers, product managers, and QA engineers. "
            "Supports the Media Test Strategy. Contributes to continuous improvement "
            "of test processes, tools, and methodologies. Quality advocate from "
            "requirements to production.",
        ),
        (
            "Cloud, VDI & Windows Environments: ",
            "Cloud environments — GCP (primary, ACE certified), AWS (CP). Supports "
            "automated testing across Cloud, VDI, and Windows 10/11 setups. Docker, "
            "Kubernetes, Terraform. CI/CD pipelines (GitHub Actions, Jenkins). "
            "Version control (Git). Modern testing tools.",
        ),
        (
            "Leadership & Cross-Team Coordination: ",
            "Demonstrated leadership in driving quality initiatives across teams. "
            "Coordinates with multiple teams and products to ensure seamless "
            "integration and quality delivery. Mentors SDETs. Global, cross-functional "
            "teams at IKEA, Truecaller, and LEGO. Fast-paced, self-driven, proactive.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Languages: ",
            "C# (.NET), Java, TypeScript, Python, SQL, Bash",
        ),
        (
            "Media Technologies: ",
            "WebRTC (audio/video/screen sharing, SRTP, TURN/STUN, ICE), SIP "
            "(trunking, call routing, codec negotiation), real-time communication "
            "protocol validation",
        ),
        (
            "Test Automation: ",
            "Playwright, Selenium, Appium, RestAssured, Karate, Postman, E2E, "
            "API automation, regression, smoke, performance (JMeter, K6), "
            "framework design (POM, BDD)",
        ),
        (
            "Cloud & Infrastructure: ",
            "GCP (Cloud Run, BigQuery, Pub/Sub), AWS, Docker, Kubernetes, "
            "Terraform, VDI environments, Windows 10/11",
        ),
        (
            "CI/CD & Tools: ",
            "GitHub Actions, Jenkins, Git, Jira, Confluence, TestRail, Zephyr, "
            "Grafana, structured logging",
        ),
        (
            "Practices: ",
            "Shift-left testing, quality strategy, test process improvement, "
            "code reviews, mentoring, Agile (Scrum/SAFe), cross-team coordination",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Media SDET / Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — WebRTC/SIP Media Communication Platform, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "WebRTC & SIP quality ownership — validates and troubleshoots real-time "
        "communication systems: audio/video calls, screen sharing, call routing, "
        "codec negotiation, SRTP, TURN/STUN, ICE connectivity, SIP trunk "
        "integration. End-to-end media flow validation across browsers and "
        "environments. Owns quality strategy for media backend and frontend.",
        "Test automation framework — designed, implemented, and maintains a robust, "
        "scalable automation framework (Playwright/TypeScript) for E2E and API "
        "testing of media platform. Regression and smoke automation suites. "
        "Framework design patterns (POM). Code reviews. Shift-left testing — "
        "embedded quality early in SDLC with developers and PMs.",
        "Cloud, VDI & Windows environments — automated testing across GCP Cloud, "
        "VDI setups, and Windows 10/11. CI/CD pipelines (GitHub Actions, Docker, "
        "Kubernetes, Terraform). Version control (Git). Modern testing tools. "
        "Continuous improvement of test processes and tools.",
        "Cross-team coordination — coordinates with multiple teams and products "
        "(media backend, frontend, infrastructure) to ensure seamless integration "
        "and quality delivery. Mentors SDETs. Quality initiatives across teams. "
        "Agile (Scrum). Grafana observability and quality metrics.",
        "AI-assisted testing (Claude, Copilot, Gemini) — 30% velocity improvement. "
        "Performance testing for media services. Supports Media Test Strategy "
        "evolution. Self-driven, proactive. 30+ global markets.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior SDET / QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users, VoIP/Calling", bold=True, size=10)
    tc_bullets = [
        "Communication platform quality — tested VoIP/calling features, real-time "
        "communication flows at scale (300M+ users). End-to-end validation of call "
        "flows, API testing (RestAssured), mobile testing (Appium — iOS/Android). "
        "Test automation framework design. Shift-left practices with developers.",
        "Cross-team coordination and quality leadership. CI/CD. AWS cloud. "
        "Code reviews. Mentored engineers. Performance testing. Regression and "
        "smoke automation. Agile (Scrum/Kanban). Fast-paced global environment.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Communication Platforms — Multi-Channel", bold=True, size=10)
    lego_bullets = [
        "LEGO: Designed and maintained scalable test automation frameworks — "
        "Selenium, RestAssured, Karate (Java, C#). E2E and API automation for "
        "e-commerce platform. Framework design patterns (POM, BDD). Led 8–10 "
        "SDETs. Code reviews. CI/CD (Jenkins). Cross-team coordination. "
        "Shift-left practices. Quality strategy.",
        "IKEA (2018–2021): Communication and media platforms — Genesys (contact "
        "center/VoIP), Verint/CSSP, IKEA App. Validated real-time communication "
        "flows. Test automation. API testing. Regression, smoke, performance. "
        "Cross-functional coordination. Agile (Scrum). Windows environments.",
        "Quality leadership — drove quality initiatives across teams. Mentored "
        "SDETs. Shift-left testing. Continuous improvement of test processes and "
        "tools. Seamless integration and quality delivery.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Enterprise — Automation Frameworks, C#, Java", bold=True, size=10)
    fin_bullets = [
        "Test automation frameworks — designed and maintained automation frameworks "
        "(Selenium, TestNG, C#, Java) for enterprise banking applications (Finacle). "
        "E2E, API, regression, smoke, performance testing. CI/CD. Windows "
        "environments. Code reviews. SQL. Shift-left practices.",
        "Quality leadership — mentored 15+ engineers. Drove quality initiatives. "
        "Cross-team coordination. Framework design patterns. Continuous improvement. "
        "Agile and Waterfall. Fast-paced, global environment.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "WebRTC/SIP media quality ownership at IKEA — end-to-end validation of "
        "real-time communication systems (audio/video, screen sharing, call routing, "
        "SIP trunking) serving 30+ global markets. Media Test Strategy support.",
        "Scalable test automation frameworks — designed and maintained across 4 "
        "organizations (C#, Java, TypeScript). Playwright, Selenium, RestAssured. "
        "Framework patterns (POM, BDD). CI/CD integration (GitHub Actions, Jenkins).",
        "Shift-left testing advocate — embedded quality early in SDLC. Collaborated "
        "with developers, PMs, QA engineers. Continuous improvement of test "
        "processes, tools, and methodologies.",
        "Cross-team coordination — multiple teams and products at IKEA, Truecaller, "
        "LEGO. Seamless integration and quality delivery. Mentored SDETs. Code "
        "reviews. Quality initiatives across teams.",
        "Cloud, VDI & Windows — automated testing across GCP, AWS, VDI, and "
        "Windows 10/11 environments. Docker, Kubernetes, Terraform. Performance "
        "testing for media services. AI-assisted testing (30% velocity).",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Media SDET Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Media SDET / Technical Lead | WebRTC, SIP, Test Automation Frameworks, C# | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Media</span> SDET / Technical Lead with 15+ years. Hands-on <span class="hl">WebRTC</span> and <span class="hl">SIP</span> <span class="hl">real-time communication</span> systems. Owns <span class="hl">quality strategy</span> for <span class="hl">media</span> backend and frontend. Designs, implements, maintains robust, <span class="hl">scalable</span> <span class="hl">test automation framework</span>s (<span class="hl">C#</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>). <span class="hl">Shift-left</span> advocate. <span class="hl">Cloud</span> (<span class="hl">GCP</span>/<span class="hl">AWS</span>), <span class="hl">VDI</span>, <span class="hl">Windows 10</span>/11. <span class="hl">CI/CD</span>, <span class="hl">Git</span>, <span class="hl">code review</span>s, <span class="hl">mentor</span>s SDETs. <span class="hl">Cross-functional</span> coordination. At <span class="hl">IKEA</span> (3+ years). Also <span class="hl">Truecaller</span> and <span class="hl">LEGO</span>.</p>

  <div class="section">How I Match the Role</div>
  <p><b>WebRTC &amp; SIP — Media Quality:</b> <span class="hl">WebRTC</span>/<span class="hl">SIP</span> at <span class="hl">IKEA</span> — audio/video, screen sharing, call routing, codec, SRTP, TURN/STUN, ICE, <span class="hl">SIP</span> trunking. Owns <span class="hl">quality strategy</span> for <span class="hl">media</span> backend/frontend. <span class="hl">End-to-end</span> media validation.<br>
  <b>Test Automation Frameworks — C#, Java, TS:</b> Designs/maintains <span class="hl">scalable</span> <span class="hl">automation framework</span>s. <span class="hl">C#</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>. <span class="hl">Playwright</span>, <span class="hl">Selenium</span>. POM, BDD. <span class="hl">Code review</span>s. <span class="hl">Mentor</span>s SDETs.<br>
  <b>Shift-Left &amp; Quality Strategy:</b> <span class="hl">Shift-left</span> — embeds quality early. Collaborates with devs, PMs, QA. Media Test Strategy. Continuous improvement of processes/tools.<br>
  <b>Cloud, VDI &amp; Windows:</b> <span class="hl">GCP</span> (ACE certified), <span class="hl">AWS</span> (CP). <span class="hl">VDI</span>, <span class="hl">Windows 10</span>/11. <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">Terraform</span>. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>). <span class="hl">Git</span>.<br>
  <b>Leadership &amp; Cross-Team:</b> Quality initiatives across teams. <span class="hl">Cross-functional</span> coordination at <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. <span class="hl">Mentor</span>s SDETs. Seamless <span class="hl">integration</span> and quality delivery.</p>

  <div class="section">Technical Skills</div>
  <p><b>Languages:</b> C# (.NET), Java, TypeScript, Python, SQL, Bash<br>
  <b>Media:</b> WebRTC (audio/video/screen sharing, SRTP, TURN/STUN, ICE), SIP (trunking, call routing, codec), real-time communication<br>
  <b>Test Automation:</b> Playwright, Selenium, Appium, RestAssured, Karate, Postman, E2E, API, regression, smoke, performance (JMeter, K6), POM, BDD<br>
  <b>Cloud &amp; Infra:</b> GCP (Cloud Run, BigQuery, Pub/Sub), AWS, Docker, Kubernetes, Terraform, VDI, Windows 10/11<br>
  <b>CI/CD &amp; Tools:</b> GitHub Actions, Jenkins, Git, Jira, Confluence, TestRail, Zephyr, Grafana<br>
  <b>Practices:</b> Shift-left, quality strategy, code reviews, mentoring, Agile (Scrum/SAFe), cross-team coordination</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Media SDET / Technical Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — WebRTC/SIP Media Communication Platform, 30+ Markets</div>
  <ul>
    <li><span class="hl">WebRTC</span> &amp; <span class="hl">SIP</span> <span class="hl">quality strategy</span> — validates <span class="hl">real-time communication</span>: audio/video, screen sharing, call routing, codec, SRTP, TURN/STUN, ICE, <span class="hl">SIP</span> trunking. <span class="hl">End-to-end</span> <span class="hl">media</span> flow validation. 30+ markets.</li>
    <li><span class="hl">Test automation framework</span> — <span class="hl">Playwright</span>/<span class="hl">TypeScript</span>, <span class="hl">E2E</span> &amp; <span class="hl">API</span>. <span class="hl">Scalable</span>, POM. <span class="hl">Code review</span>s. <span class="hl">Shift-left</span> — devs &amp; PMs. <span class="hl">Regression</span>, smoke suites.</li>
    <li><span class="hl">Cloud</span>, <span class="hl">VDI</span> &amp; <span class="hl">Windows 10</span>/11 — <span class="hl">GCP</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">Terraform</span>. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>). <span class="hl">Git</span>. Test process improvement.</li>
    <li><span class="hl">Cross-team</span> coordination — <span class="hl">media</span> backend, frontend, infra. <span class="hl">Mentor</span>s SDETs. Quality initiatives. <span class="hl">Agile</span> (<span class="hl">Scrum</span>). Grafana.</li>
    <li>AI-assisted testing — 30% velocity. <span class="hl">Performance</span> testing for <span class="hl">media</span> services. Media Test Strategy evolution.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior SDET / QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users, VoIP/Calling</div>
  <ul>
    <li>Communication platform — VoIP/calling, <span class="hl">real-time communication</span> (300M+ users). <span class="hl">End-to-end</span> call flow validation. <span class="hl">API</span> testing. <span class="hl">Test automation framework</span>. <span class="hl">Shift-left</span>. <span class="hl">Selenium</span>, Appium, RestAssured.</li>
    <li><span class="hl">Cross-team</span> coordination. <span class="hl">CI/CD</span>. <span class="hl">AWS</span> <span class="hl">cloud</span>. <span class="hl">Code review</span>s. <span class="hl">Mentor</span>ed engineers. <span class="hl">Performance</span>. <span class="hl">Regression</span>/smoke. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Senior SDET / Test Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Communication Platforms</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Scalable</span> <span class="hl">test automation framework</span>s — <span class="hl">Selenium</span>, RestAssured, Karate (<span class="hl">Java</span>, <span class="hl">C#</span>). <span class="hl">E2E</span>, <span class="hl">API</span>, POM, BDD. 8–10 SDETs. <span class="hl">Code review</span>s. <span class="hl">CI/CD</span> (<span class="hl">Jenkins</span>). <span class="hl">Shift-left</span>. <span class="hl">Quality strategy</span>.</li>
    <li><span class="hl">IKEA</span> (2018–21): Communication/<span class="hl">media</span> — Genesys (VoIP), Verint, IKEA App. <span class="hl">Real-time communication</span> flows. <span class="hl">Test automation</span>. <span class="hl">API</span>. <span class="hl">Regression</span>, smoke, <span class="hl">performance</span>. <span class="hl">Windows</span>. <span class="hl">Agile</span>.</li>
    <li>Quality leadership — quality initiatives across teams. <span class="hl">Mentor</span>ed SDETs. <span class="hl">Shift-left</span>. Continuous improvement. Seamless <span class="hl">integration</span> and quality delivery.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior SDET / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Enterprise — Automation Frameworks, C#, Java</div>
  <ul>
    <li><span class="hl">Test automation framework</span>s — <span class="hl">Selenium</span>, TestNG (<span class="hl">C#</span>, <span class="hl">Java</span>). Enterprise banking (Finacle). <span class="hl">E2E</span>, <span class="hl">API</span>, <span class="hl">regression</span>, smoke, <span class="hl">performance</span>. <span class="hl">CI/CD</span>. <span class="hl">Windows</span>. <span class="hl">Code review</span>s. <span class="hl">Shift-left</span>.</li>
    <li>Quality leadership — 15+ engineers. Quality initiatives. <span class="hl">Cross-team</span> coordination. Continuous improvement. <span class="hl">Agile</span> &amp; Waterfall.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">WebRTC</span>/<span class="hl">SIP</span> <span class="hl">media</span> quality at <span class="hl">IKEA</span> — <span class="hl">real-time communication</span> validation (audio/video, screen sharing, call routing, <span class="hl">SIP</span> trunking) for 30+ markets.</li>
    <li><span class="hl">Scalable</span> <span class="hl">test automation framework</span>s across 4 orgs (<span class="hl">C#</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>). <span class="hl">Playwright</span>, <span class="hl">Selenium</span>. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>).</li>
    <li><span class="hl">Shift-left</span> — quality early in SDLC. <span class="hl">Code review</span>s. <span class="hl">Mentor</span>ed SDETs. Continuous improvement of test processes and tools.</li>
    <li><span class="hl">Cross-team</span> coordination — <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. Seamless <span class="hl">integration</span> and quality delivery. Quality initiatives.</li>
    <li><span class="hl">Cloud</span>, <span class="hl">VDI</span>, <span class="hl">Windows 10</span>/11 — <span class="hl">GCP</span>, <span class="hl">AWS</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>. <span class="hl">Performance</span> testing. AI-assisted (30% velocity).</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
