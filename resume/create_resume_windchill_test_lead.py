"""
Generate a tailored resume for Windchill Test Lead.
Focus: PTC Windchill testing, PLM, test strategy, UAT, integrations (ERP/CAD),
defect management, quality metrics, lifecycle/workflows/access control.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Windchill_Test_Lead_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Windchill_Test_Lead_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "Windchill", "PTC Windchill", "PLM", "PDMLink", "MPMLink", "ProjectLink",
    "test strategy", "test plan", "UAT", "user acceptance testing",
    "ERP", "CAD", "Creo", "SAP", "Inventor",
    "regression", "integration testing", "performance testing",
    "defect management", "triage", "root-cause analysis",
    "lifecycle", "workflows", "access control", "Change Management",
    "BOM", "Document Management", "quality metrics",
    "ISTQB", "Azure DevOps", "Jira", "HP ALM",
    "cloud", "migration", "upgrades",
    "entry/exit criteria", "test coverage", "stakeholder",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Test Lead — PTC Windchill / PLM", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Test Lead with 10+ years of experience in enterprise PLM/Windchill testing across "
        "implementations, upgrades, migrations, and cloud deployments. Proven ability to define "
        "test strategies, create comprehensive test plans, manage entry/exit criteria, and drive "
        "end-to-end quality assurance for PTC Windchill environments including PDMLink, MPMLink, "
        "and ProjectLink modules. Deep understanding of PLM business processes — Change Management, "
        "BOMs, Document Management, lifecycle states, workflows, and access control. Experienced "
        "in coordinating system, integration, regression, performance, and UAT testing across "
        "Windchill and its integrations with ERP (SAP), CAD (Creo, Inventor), and other enterprise "
        "systems. Strong communicator who reports quality metrics, manages risks, and ensures "
        "stakeholder alignment throughout the testing lifecycle."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "PTC Windchill Testing (PDMLink/MPMLink)", "Test Strategy & Plan Ownership", "UAT Coordination & Business Signoff",
        "PLM Lifecycle & Workflow Validation", "Defect Triage & Root-Cause Analysis", "ERP/CAD Integration Testing",
        "Quality Metrics & Stakeholder Reporting", "Regression & Performance Testing", "Cloud Migration & Upgrade Testing",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Lead / SDET Lead — Enterprise Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Lead end-to-end testing activities for a large-scale enterprise platform deployed across 30+ markets — defining test strategy, test plans, schedules, and entry/exit criteria.",
        "Manage system, integration, regression, and UAT testing across multiple modules and integrations with ERP systems, third-party services, and internal enterprise tools.",
        "Review and validate business and technical requirements to ensure comprehensive test coverage — working closely with product managers and architects.",
        "Design, review, and maintain test cases, test scenarios, and test data aligned with lifecycle configurations, workflows, and access control policies.",
        "Lead defect management including triage, prioritisation, root-cause analysis, and retesting — driving resolution with development teams.",
        "Communicate test progress, risks, and quality metrics to stakeholders through structured dashboards and weekly status reports.",
        "Support UAT by guiding business users, managing test cycles, and ensuring signoff readiness — coordinating across multiple Nordic markets.",
        "Drive continuous improvement in testing practices: introduced automated regression suites, risk-based testing, and structured retrospectives.",
        "Coordinate testing across cloud infrastructure (GCP), data pipelines, and API layers — validating end-to-end data flows and system integrations.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Release & Quality Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Defined test strategies for release readiness across a 300M+ user platform — managing regression testing, performance validation, and quality gate enforcement.",
        "Led defect triage and root-cause analysis during release cycles — coordinating with development squads to ensure timely resolution and retesting.",
        "Drove quality metrics adoption (defect leakage, test coverage, escape rate) to support data-driven release decisions.",
        "Coordinated integration testing across microservices and third-party system dependencies (AWS infrastructure).",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech (IKEA & LEGO) ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior Test Analyst / Test Lead — Enterprise PLM & E-Commerce", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led end-to-end testing for PTC Windchill implementations and upgrades at LEGO Group — validating PDMLink configurations, lifecycle states, workflows, access control, and BOM structures.",
        "Defined and owned test strategies and test plans for Windchill module testing including Change Management, Document Management, and revision control processes.",
        "Coordinated integration testing between Windchill and enterprise systems — validating data flows to/from ERP (SAP), CAD tools (Creo, Inventor), and downstream manufacturing systems.",
        "Managed system, regression, and UAT testing for Windchill cloud migration projects — ensuring zero data loss and functional parity post-migration.",
        "Supported UAT by guiding engineering and business users through test cycles, managing test data preparation, and ensuring stakeholder signoff readiness.",
        "Led defect management using HP ALM and Jira — running triage sessions, tracking root-cause analysis, and reporting quality metrics to programme leadership.",
        "Designed and maintained comprehensive test case libraries for PLM workflows: part creation, change notices (ECN/ECR), promotion requests, and access policy validation.",
        "Drove test automation initiatives for regression coverage on Windchill — reducing manual regression effort by 40% across upgrade cycles.",
        "Tested Windchill integrations with ProjectLink for project scheduling and resource allocation workflows.",
        "Operated in complex engineering environments with strict regulatory and quality standards — applying ISTQB methodologies throughout.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise IT & Banking", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Analyst / Technical Lead", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led testing for core banking system implementations (Finacle CBS) — managing test plans, integration testing, data migration validation, and UAT coordination in regulated environments.",
        "Designed test strategies for ERP-adjacent systems: billing, ETL pipelines, and third-party integrations — ensuring end-to-end data integrity across enterprise platforms.",
        "Managed defect lifecycle using HP ALM — triage, prioritisation, and stakeholder reporting for multi-phase rollouts across 20+ branches.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Tools & Technologies ───────────────────────────────────────────────
    add_section_heading(doc, "TOOLS & TECHNOLOGIES")
    for label, value in [
        ("PLM/PDM: ", "PTC Windchill (PDMLink, MPMLink, ProjectLink), lifecycle/workflow configuration, access control, Change Management, BOMs"),
        ("Test Management: ", "Azure DevOps, Jira, HP ALM (Quality Center), Confluence, Zephyr"),
        ("Integration: ", "SAP (ERP), Creo, Inventor (CAD), REST APIs, middleware/ETL, enterprise service bus"),
        ("Automation: ", "Selenium, Playwright, Python, API test automation, CI/CD (GitHub Actions, Jenkins)"),
        ("Methodologies: ", "Agile (Scrum/Kanban), Waterfall, V-Model, ISTQB, risk-based testing"),
        ("Cloud: ", "GCP, AWS — cloud migration testing, infrastructure validation"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PG Diploma, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "ITIL v4 Foundation",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Test Lead &mdash; PTC Windchill / PLM</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Test Lead with 10+ years of experience in enterprise PLM/Windchill testing across implementations, upgrades, migrations, and cloud deployments. Proven ability to define test strategies, create comprehensive test plans, manage entry/exit criteria, and drive end-to-end quality assurance for PTC Windchill environments including PDMLink, MPMLink, and ProjectLink modules. Deep understanding of PLM business processes &mdash; Change Management, BOMs, Document Management, lifecycle states, workflows, and access control. Experienced in coordinating system, integration, regression, performance, and UAT testing across Windchill and its integrations with ERP (SAP), CAD (Creo, Inventor), and other enterprise systems. Strong communicator who reports quality metrics, manages risks, and ensures stakeholder alignment throughout the testing lifecycle.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; PTC Windchill Testing (PDMLink/MPMLink)</td><td>&bull; Test Strategy &amp; Plan Ownership</td><td>&bull; UAT Coordination &amp; Business Signoff</td></tr>
<tr><td>&bull; PLM Lifecycle &amp; Workflow Validation</td><td>&bull; Defect Triage &amp; Root-Cause Analysis</td><td>&bull; ERP/CAD Integration Testing</td></tr>
<tr><td>&bull; Quality Metrics &amp; Stakeholder Reporting</td><td>&bull; Regression &amp; Performance Testing</td><td>&bull; Cloud Migration &amp; Upgrade Testing</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Test Lead / SDET Lead &mdash; Enterprise Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Lead end-to-end testing activities for a large-scale enterprise platform deployed across 30+ markets &mdash; defining test strategy, test plans, schedules, and entry/exit criteria.</li>
<li>Manage system, integration, regression, and UAT testing across multiple modules and integrations with ERP systems, third-party services, and internal enterprise tools.</li>
<li>Review and validate business and technical requirements to ensure comprehensive test coverage &mdash; working closely with product managers and architects.</li>
<li>Design, review, and maintain test cases, test scenarios, and test data aligned with lifecycle configurations, workflows, and access control policies.</li>
<li>Lead defect management including triage, prioritisation, root-cause analysis, and retesting &mdash; driving resolution with development teams.</li>
<li>Communicate test progress, risks, and quality metrics to stakeholders through structured dashboards and weekly status reports.</li>
<li>Support UAT by guiding business users, managing test cycles, and ensuring signoff readiness &mdash; coordinating across multiple Nordic markets.</li>
<li>Drive continuous improvement in testing practices: introduced automated regression suites, risk-based testing, and structured retrospectives.</li>
<li>Coordinate testing across cloud infrastructure (GCP), data pipelines, and API layers &mdash; validating end-to-end data flows and system integrations.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Release &amp; Quality Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Defined test strategies for release readiness across a 300M+ user platform &mdash; managing regression testing, performance validation, and quality gate enforcement.</li>
<li>Led defect triage and root-cause analysis during release cycles &mdash; coordinating with development squads to ensure timely resolution and retesting.</li>
<li>Drove quality metrics adoption (defect leakage, test coverage, escape rate) to support data-driven release decisions.</li>
<li>Coordinated integration testing across microservices and third-party system dependencies (AWS infrastructure).</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Senior Test Analyst / Test Lead &mdash; Enterprise PLM &amp; E-Commerce <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Led end-to-end testing for PTC Windchill implementations and upgrades at LEGO Group &mdash; validating PDMLink configurations, lifecycle states, workflows, access control, and BOM structures.</li>
<li>Defined and owned test strategies and test plans for Windchill module testing including Change Management, Document Management, and revision control processes.</li>
<li>Coordinated integration testing between Windchill and enterprise systems &mdash; validating data flows to/from ERP (SAP), CAD tools (Creo, Inventor), and downstream manufacturing systems.</li>
<li>Managed system, regression, and UAT testing for Windchill cloud migration projects &mdash; ensuring zero data loss and functional parity post-migration.</li>
<li>Supported UAT by guiding engineering and business users through test cycles, managing test data preparation, and ensuring stakeholder signoff readiness.</li>
<li>Led defect management using HP ALM and Jira &mdash; running triage sessions, tracking root-cause analysis, and reporting quality metrics to programme leadership.</li>
<li>Designed and maintained comprehensive test case libraries for PLM workflows: part creation, change notices (ECN/ECR), promotion requests, and access policy validation.</li>
<li>Drove test automation initiatives for regression coverage on Windchill &mdash; reducing manual regression effort by 40% across upgrade cycles.</li>
<li>Tested Windchill integrations with ProjectLink for project scheduling and resource allocation workflows.</li>
<li>Operated in complex engineering environments with strict regulatory and quality standards &mdash; applying ISTQB methodologies throughout.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise IT &amp; Banking</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Led testing for core banking system implementations (Finacle CBS) &mdash; managing test plans, integration testing, data migration validation, and UAT coordination in regulated environments.</li>
<li>Designed test strategies for ERP-adjacent systems: billing, ETL pipelines, and third-party integrations &mdash; ensuring end-to-end data integrity across enterprise platforms.</li>
<li>Managed defect lifecycle using HP ALM &mdash; triage, prioritisation, and stakeholder reporting for multi-phase rollouts across 20+ branches.</li>
</ul>

<h2>TOOLS &amp; TECHNOLOGIES</h2>
<p class="tech-line"><b>PLM/PDM:</b> PTC Windchill (PDMLink, MPMLink, ProjectLink), lifecycle/workflow configuration, access control, Change Management, BOMs</p>
<p class="tech-line"><b>Test Management:</b> Azure DevOps, Jira, HP ALM (Quality Center), Confluence, Zephyr</p>
<p class="tech-line"><b>Integration:</b> SAP (ERP), Creo, Inventor (CAD), REST APIs, middleware/ETL, enterprise service bus</p>
<p class="tech-line"><b>Automation:</b> Selenium, Playwright, Python, API test automation, CI/CD (GitHub Actions, Jenkins)</p>
<p class="tech-line"><b>Methodologies:</b> Agile (Scrum/Kanban), Waterfall, V-Model, ISTQB, risk-based testing</p>
<p class="tech-line"><b>Cloud:</b> GCP, AWS &mdash; cloud migration testing, infrastructure validation</p>

<h2>EDUCATION</h2>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; Six Sigma Green Belt</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
