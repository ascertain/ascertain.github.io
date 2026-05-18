from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_System_Test_Architect_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_System_Test_Architect_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "system-level test",
    "system test",
    "test architecture",
    "test architect",
    "hardware/software",
    "HW/SW",
    "requirements analysis",
    "test specification",
    "coverage definition",
    "test coverage",
    "cross-domain",
    "hardware",
    "software",
    "verification",
    "validation",
    "test lead",
    "SoC",
    "Linux",
    "Android",
    "ISTQB",
    "FuSA",
    "functional safety",
    "safety standards",
    "test strategy",
    "test plan",
    "test automation",
    "automation framework",
    "integration testing",
    "regression",
    "E2E",
    "end-to-end",
    "API",
    "CI/CD",
    "embedded",
    "firmware",
    "protocol",
    "WebRTC",
    "SIP",
    "Playwright",
    "Selenium",
    "Docker",
    "Kubernetes",
    "GCP",
    "AWS",
    "Agile",
    "Scrum",
    "IKEA",
    "LEGO",
    "Truecaller",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "System Test Architect  |  HW/SW System-Level Testing, Test Architecture & Test Lead  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "System Test Architect with 15+ years in system-level testing and test "
        "architecture for combined hardware/software systems. Extensive experience "
        "in requirements analysis, test specification authoring, and coverage "
        "definition. Hands-on background in both hardware and software verification, "
        "with the ability to operate effectively across domain boundaries — from "
        "embedded/firmware-level protocols (WebRTC, SIP, real-time media) to "
        "application-layer software and cloud-based services. Proven test lead "
        "with demonstrated ownership of test strategy, test plans, and quality "
        "delivery for complex, cross-domain systems. Experience with complex SoC-based "
        "systems and Linux environments. ISTQB Certified. Functional safety (FuSA) "
        "awareness. Currently at IKEA IT AB (3+ years) — system-level testing for "
        "a WebRTC/SIP communication platform spanning hardware endpoints, software "
        "stack, and cloud infrastructure across 30+ markets.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "System-Level Test & Test Architecture: ",
            "15+ years designing test architecture for combined hardware/software "
            "systems. Requirements analysis — traces system requirements to test "
            "specifications and defines coverage criteria (functional, structural, "
            "boundary, risk-based). Authors test specifications for system-level, "
            "integration, and end-to-end testing. Currently at IKEA — test "
            "architecture for WebRTC/SIP communication platform spanning hardware "
            "endpoints (phones, headsets, codecs), embedded protocols, software "
            "stack, and cloud services (GCP) serving 30+ markets.",
        ),
        (
            "Cross-Domain HW/SW Verification: ",
            "Hands-on background in both hardware and software verification — "
            "operates effectively across domain boundaries. Hardware: endpoint "
            "device testing (audio/video hardware, codecs, SIP phones), protocol "
            "validation (WebRTC, SIP, SRTP, TURN/STUN, ICE), firmware interaction. "
            "Software: application-layer testing, API verification, E2E automation, "
            "cloud-service validation. Integration testing across HW/SW boundaries.",
        ),
        (
            "Test Lead Experience: ",
            "Proven test lead across 4 organizations — IKEA, Truecaller, LEGO, "
            "and banking/enterprise. Defines test strategy, authors test plans, "
            "manages test execution, and drives quality delivery. Led teams of "
            "8–15+ engineers. Cross-functional coordination with developers, "
            "architects, product managers, and hardware engineers. Mentors SDETs.",
        ),
        (
            "SoC, Linux & Safety Awareness: ",
            "Experience with Linux-based systems — cloud infrastructure (GCP, AWS, "
            "Docker, Kubernetes), server-side Linux environments, CI/CD on Linux "
            "runners. Understanding of complex SoC architectures through embedded "
            "protocol validation (WebRTC/SIP codec negotiation, media processing). "
            "ISTQB Certified. Functional safety (FuSA) awareness — safety-critical "
            "thinking from regulated environments (banking compliance, CEH security).",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "System Test & Architecture: ",
            "System-level test design, test architecture, requirements analysis, "
            "test specification authoring, coverage definition, traceability, "
            "risk-based testing, integration testing, E2E testing",
        ),
        (
            "HW/SW Verification: ",
            "Hardware endpoint testing (audio/video, codecs, SIP phones), embedded "
            "protocol validation (WebRTC, SIP, SRTP, TURN/STUN, ICE), firmware "
            "interaction, HW/SW integration testing, cross-domain verification",
        ),
        (
            "Test Automation: ",
            "Playwright, Selenium, Appium, RestAssured, Karate, Postman, E2E, "
            "API automation, regression, smoke, performance (JMeter, K6), "
            "framework design (POM, BDD)",
        ),
        (
            "Platforms & Infrastructure: ",
            "Linux (server, CI/CD runners), GCP, AWS, Docker, Kubernetes, "
            "Terraform, Windows 10/11, SoC architecture awareness",
        ),
        (
            "Languages & CI/CD: ",
            "TypeScript, Java, C#, Python, SQL, Bash | GitHub Actions, Jenkins, "
            "Git, Jira, Confluence, TestRail, Zephyr, Grafana",
        ),
        (
            "Practices & Standards: ",
            "ISTQB, FuSA awareness, shift-left testing, test strategy/planning, "
            "Agile (Scrum/SAFe), cross-domain coordination, mentoring, code reviews",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — System Test Architect / Senior SDET", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — WebRTC/SIP HW/SW Communication Platform, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "System-level test architecture for combined hardware/software communication "
        "platform — WebRTC/SIP spanning hardware endpoints (phones, headsets, codecs), "
        "embedded protocols, software stack, and cloud services (GCP). Requirements "
        "analysis — traced system requirements to test specifications. Coverage "
        "definition across functional, integration, and E2E levels.",
        "Cross-domain HW/SW verification — validated hardware endpoint interaction "
        "(audio/video, codecs, SIP phones) alongside software-layer testing (API, "
        "E2E, cloud services). Protocol verification: WebRTC, SIP, SRTP, TURN/STUN, "
        "ICE connectivity, codec negotiation, call routing. Integration testing "
        "across hardware/software boundaries.",
        "Test automation framework — designed and maintained scalable automation "
        "(Playwright/TypeScript) for E2E and API testing. Regression and smoke "
        "suites. CI/CD (GitHub Actions, Docker, Kubernetes, Terraform). Linux-based "
        "CI/CD runners. Test specification authoring and coverage tracking.",
        "Test lead — owned test strategy and quality delivery for 30+ global "
        "markets. Cross-functional coordination with developers, architects, "
        "hardware engineers, and product managers. Mentored SDETs. Agile (Scrum). "
        "Quality metrics (Grafana). AI-assisted testing — 30% velocity improvement.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users, VoIP/Calling", bold=True, size=10)
    tc_bullets = [
        "System-level testing for communication platform (300M+ users) — VoIP/calling "
        "across hardware (mobile devices, iOS/Android) and software layers. "
        "Requirements analysis and test specification for call flow verification. "
        "API testing (RestAssured). Cross-domain — mobile hardware, application "
        "software, cloud services (AWS). Test automation framework design.",
        "Test lead — test strategy, quality delivery. Cross-team coordination. "
        "Regression, performance testing. CI/CD. Mentored engineers. Agile.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Communication — Multi-Platform, HW/SW Integration", bold=True, size=10)
    lego_bullets = [
        "LEGO: System-level test architecture for e-commerce platform — requirements "
        "analysis, test specification, coverage definition across web, mobile, and "
        "backend. Automation frameworks (Selenium, RestAssured, Karate — Java, C#). "
        "Integration testing across platform layers. Led 8–10 SDETs. CI/CD (Jenkins). "
        "Test strategy ownership. Cross-domain coordination.",
        "IKEA (2018–2021): System testing for communication platforms — Genesys "
        "(contact center/VoIP hardware + software), Verint/CSSP, IKEA App. "
        "Cross-domain HW/SW verification: telephony hardware, SIP integration, "
        "application software. Integration testing. Test specifications. "
        "Regression, API testing. Agile (Scrum).",
        "Test lead — drove test strategy and quality delivery across multiple "
        "projects. Mentored engineers. Cross-functional coordination with "
        "developers, architects, and hardware teams. Continuous improvement.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Enterprise — System Testing, Regulated Environments", bold=True, size=10)
    fin_bullets = [
        "System-level testing for enterprise banking platforms (Finacle) — "
        "requirements analysis, test specification authoring, coverage definition. "
        "Integration testing across hardware (ATMs, POS terminals) and software "
        "layers. Automation frameworks (Selenium, TestNG — C#, Java). "
        "CI/CD. SQL. Compliance in regulated environments.",
        "Test lead — 15+ engineers. Test strategy and quality delivery. "
        "Cross-domain coordination (hardware, software, network). Mentored "
        "engineers. Agile and Waterfall. Continuous improvement.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "System-level test architecture for combined HW/SW platform at IKEA — "
        "WebRTC/SIP spanning hardware endpoints, embedded protocols, software "
        "stack, and cloud (GCP). Requirements traceability. Coverage definition. "
        "30+ global markets.",
        "Cross-domain HW/SW verification — hardware endpoint testing (audio/video, "
        "codecs, SIP phones, ATMs, mobile devices) alongside software/cloud "
        "verification. Integration testing across boundaries at IKEA, Truecaller, "
        "LEGO, and banking.",
        "Test lead across 4 organizations — test strategy, test plans, quality "
        "delivery. Led teams of 8–15+ SDETs. Cross-functional coordination. "
        "Mentoring. Test specification authoring.",
        "Scalable test automation frameworks (Playwright, Selenium, RestAssured — "
        "TypeScript, Java, C#). CI/CD (GitHub Actions, Jenkins). Linux-based "
        "environments. AI-assisted testing — 30% velocity improvement.",
        "ISTQB Certified. FuSA awareness from regulated environments (banking "
        "compliance, CEH security). Shift-left testing practices.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – System Test Architect Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>System Test Architect | HW/SW System-Level Testing, Test Architecture &amp; Test Lead | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">System Test Architect</span> with 15+ years in <span class="hl">system-level test</span>ing and <span class="hl">test architecture</span> for combined <span class="hl">hardware/software</span> systems. <span class="hl">Requirements analysis</span>, <span class="hl">test specification</span> authoring, <span class="hl">coverage definition</span>. Hands-on <span class="hl">hardware</span> and <span class="hl">software</span> <span class="hl">verification</span> — operates across domain boundaries: <span class="hl">embedded</span> <span class="hl">protocol</span>s (<span class="hl">WebRTC</span>, <span class="hl">SIP</span>), application-layer software, <span class="hl">cloud</span> services. Proven <span class="hl">test lead</span>. <span class="hl">SoC</span>-based <span class="hl">Linux</span> experience. <span class="hl">ISTQB</span> Certified. <span class="hl">FuSA</span> awareness. At <span class="hl">IKEA</span> (3+ years) — <span class="hl">system-level test</span>ing for <span class="hl">WebRTC</span>/<span class="hl">SIP</span> platform (HW endpoints + SW + <span class="hl">cloud</span>, 30+ markets).</p>

  <div class="section">How I Match the Role</div>
  <p><b>System-Level Test &amp; Test Architecture:</b> 15+ years — <span class="hl">test architecture</span> for <span class="hl">HW/SW</span> systems. <span class="hl">Requirements analysis</span>, <span class="hl">test specification</span>, <span class="hl">coverage definition</span>. At <span class="hl">IKEA</span> — <span class="hl">WebRTC</span>/<span class="hl">SIP</span> platform: <span class="hl">hardware</span> endpoints, <span class="hl">embedded</span> <span class="hl">protocol</span>s, <span class="hl">software</span>, <span class="hl">cloud</span> (<span class="hl">GCP</span>). 30+ markets.<br>
  <b>Cross-Domain HW/SW Verification:</b> <span class="hl">Hardware</span>: endpoint devices, codecs, <span class="hl">SIP</span> phones, ATMs, mobile. <span class="hl">Software</span>: <span class="hl">API</span>, <span class="hl">E2E</span>, <span class="hl">cloud</span>-service <span class="hl">validation</span>. <span class="hl">Protocol</span>: <span class="hl">WebRTC</span>, <span class="hl">SIP</span>, SRTP, TURN/STUN, ICE. <span class="hl">Integration testing</span> across <span class="hl">HW/SW</span> boundaries.<br>
  <b>Test Lead:</b> <span class="hl">Test lead</span> across 4 orgs — <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>, banking. <span class="hl">Test strategy</span>, test plans, quality delivery. 8–15+ engineers. Cross-functional. Mentors SDETs.<br>
  <b>SoC, Linux &amp; Safety:</b> <span class="hl">Linux</span>-based systems (<span class="hl">cloud</span>, <span class="hl">CI/CD</span>, server). <span class="hl">SoC</span> architecture awareness (<span class="hl">embedded</span> <span class="hl">protocol</span>/<span class="hl">codec</span> <span class="hl">validation</span>). <span class="hl">ISTQB</span> Certified. <span class="hl">FuSA</span> awareness (regulated environments, <span class="hl">CEH</span>).</p>

  <div class="section">Technical Skills</div>
  <p><b>System Test:</b> System-level test design, test architecture, requirements analysis, test specification, coverage definition, traceability, risk-based, integration, E2E<br>
  <b>HW/SW:</b> Hardware endpoint testing (audio/video, codecs, SIP phones), embedded protocol validation (WebRTC, SIP, SRTP, TURN/STUN, ICE), firmware interaction, HW/SW integration<br>
  <b>Test Automation:</b> Playwright, Selenium, Appium, RestAssured, Karate, Postman, E2E, API, regression, smoke, performance (JMeter, K6), POM, BDD<br>
  <b>Platforms:</b> Linux, GCP, AWS, Docker, Kubernetes, Terraform, Windows 10/11, SoC awareness<br>
  <b>Languages &amp; CI/CD:</b> TypeScript, Java, C#, Python, SQL, Bash | GitHub Actions, Jenkins, Git, Jira, Confluence, TestRail, Zephyr, Grafana<br>
  <b>Standards:</b> ISTQB, FuSA awareness, shift-left, test strategy/planning, Agile (Scrum/SAFe), cross-domain coordination</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — System Test Architect / Senior SDET</div>
  <div class="job-sub">Mar 2022 – Present | VCS — WebRTC/SIP HW/SW Communication Platform, 30+ Markets</div>
  <ul>
    <li><span class="hl">System-level test</span> <span class="hl">architecture</span> for <span class="hl">HW/SW</span> platform — <span class="hl">WebRTC</span>/<span class="hl">SIP</span>: <span class="hl">hardware</span> endpoints, <span class="hl">embedded</span> <span class="hl">protocol</span>s, <span class="hl">software</span>, <span class="hl">cloud</span> (<span class="hl">GCP</span>). <span class="hl">Requirements analysis</span>, <span class="hl">test specification</span>, <span class="hl">coverage definition</span>. 30+ markets.</li>
    <li><span class="hl">Cross-domain</span> <span class="hl">HW/SW</span> <span class="hl">verification</span> — <span class="hl">hardware</span> endpoints (audio/video, codecs, <span class="hl">SIP</span> phones) + <span class="hl">software</span> (<span class="hl">API</span>, <span class="hl">E2E</span>, <span class="hl">cloud</span>). <span class="hl">Protocol</span>: <span class="hl">WebRTC</span>, <span class="hl">SIP</span>, SRTP, TURN/STUN, ICE. <span class="hl">Integration testing</span> across boundaries.</li>
    <li><span class="hl">Test automation</span> — <span class="hl">Playwright</span>/TypeScript, <span class="hl">E2E</span> &amp; <span class="hl">API</span>. <span class="hl">Regression</span>, smoke. <span class="hl">CI/CD</span> (GitHub Actions, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, Terraform). <span class="hl">Linux</span> CI runners.</li>
    <li><span class="hl">Test lead</span> — <span class="hl">test strategy</span>, quality delivery. Cross-functional (devs, architects, HW engineers, PMs). Mentored SDETs. <span class="hl">Agile</span>. AI-assisted — 30% velocity.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior SDET / Test Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users, VoIP/Calling</div>
  <ul>
    <li><span class="hl">System-level test</span>ing — VoIP/calling across <span class="hl">hardware</span> (mobile devices) and <span class="hl">software</span>. <span class="hl">Requirements analysis</span>, <span class="hl">test specification</span>. <span class="hl">Cross-domain</span> — mobile HW, app SW, <span class="hl">cloud</span> (<span class="hl">AWS</span>). <span class="hl">API</span> testing. <span class="hl">Test automation</span>.</li>
    <li><span class="hl">Test lead</span> — <span class="hl">test strategy</span>, quality delivery. Cross-team. <span class="hl">Regression</span>, performance. <span class="hl">CI/CD</span>. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Senior SDET / Test Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Communication — Multi-Platform, HW/SW Integration</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">System test</span> <span class="hl">architecture</span> — <span class="hl">requirements analysis</span>, <span class="hl">test specification</span>, <span class="hl">coverage definition</span>. <span class="hl">Automation framework</span>s (<span class="hl">Selenium</span>, RestAssured, Karate — <span class="hl">Java</span>, C#). <span class="hl">Integration testing</span>. 8–10 SDETs. <span class="hl">CI/CD</span> (<span class="hl">Jenkins</span>). <span class="hl">Test lead</span>.</li>
    <li><span class="hl">IKEA</span> (2018–21): Communication — Genesys (VoIP <span class="hl">HW/SW</span>), Verint, IKEA App. <span class="hl">Cross-domain</span> <span class="hl">HW/SW</span>: telephony <span class="hl">hardware</span>, <span class="hl">SIP</span>, <span class="hl">software</span>. <span class="hl">Integration testing</span>. <span class="hl">Test specification</span>s. <span class="hl">Agile</span>.</li>
    <li><span class="hl">Test lead</span> — <span class="hl">test strategy</span>, quality delivery. Mentored engineers. <span class="hl">Cross-domain</span> coordination.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior SDET / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Enterprise — System Testing, Regulated</div>
  <ul>
    <li><span class="hl">System-level test</span>ing — enterprise banking (Finacle). <span class="hl">Requirements analysis</span>, <span class="hl">test specification</span>, <span class="hl">coverage definition</span>. <span class="hl">Integration testing</span> across <span class="hl">hardware</span> (ATMs, POS) and <span class="hl">software</span>. <span class="hl">Automation</span> (<span class="hl">Selenium</span>, C#, Java). <span class="hl">CI/CD</span>. Compliance.</li>
    <li><span class="hl">Test lead</span> — 15+ engineers. <span class="hl">Test strategy</span>. <span class="hl">Cross-domain</span> (<span class="hl">hardware</span>, <span class="hl">software</span>, network). <span class="hl">Agile</span> &amp; Waterfall.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">System test</span> <span class="hl">architecture</span> for <span class="hl">HW/SW</span> platform at <span class="hl">IKEA</span> — <span class="hl">WebRTC</span>/<span class="hl">SIP</span> (<span class="hl">hardware</span> endpoints, <span class="hl">embedded</span> <span class="hl">protocol</span>s, <span class="hl">software</span>, <span class="hl">cloud</span>). <span class="hl">Requirements</span> traceability, <span class="hl">coverage definition</span>. 30+ markets.</li>
    <li><span class="hl">Cross-domain</span> <span class="hl">HW/SW</span> <span class="hl">verification</span> — <span class="hl">hardware</span> endpoints, <span class="hl">embedded</span> <span class="hl">protocol</span>s, <span class="hl">software</span>/<span class="hl">cloud</span>. <span class="hl">Integration testing</span> at <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>, banking.</li>
    <li><span class="hl">Test lead</span> across 4 orgs — <span class="hl">test strategy</span>, plans, quality delivery. 8–15+ SDETs. Cross-functional. <span class="hl">Test specification</span> authoring.</li>
    <li><span class="hl">Test automation</span> (<span class="hl">Playwright</span>, <span class="hl">Selenium</span>, RestAssured — TS, Java, C#). <span class="hl">CI/CD</span>. <span class="hl">Linux</span>. AI-assisted (30% velocity).</li>
    <li><span class="hl">ISTQB</span> Certified. <span class="hl">FuSA</span> awareness (regulated banking, <span class="hl">CEH</span>). Shift-left testing.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
