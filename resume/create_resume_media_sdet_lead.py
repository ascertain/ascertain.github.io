from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Media_SDET_Lead_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Media_SDET_Lead_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "SDET",
    "Media Test Strategy",
    "test strategy",
    "quality strategy",
    "test automation",
    "automation frameworks",
    "WebRTC",
    "SIP",
    "media technologies",
    "real-time communication",
    "C#",
    "Java",
    "TypeScript",
    "C++",
    "Python",
    "shift-left",
    "CI/CD",
    "GCP",
    "AWS",
    "cloud",
    "VDI",
    "Windows",
    "Git",
    "Playwright",
    "Selenium",
    "Appium",
    "RestAssured",
    "Karate",
    "code review",
    "mentor",
    "leadership",
    "cross-functional",
    "cross-team",
    "integration",
    "scalable",
    "reliable",
    "Agile",
    "Scrum",
    "Docker",
    "Kubernetes",
    "Terraform",
    "GitHub Actions",
    "Jenkins",
    "IKEA",
    "LEGO",
    "Truecaller",
    "media",
    "quality",
    "ownership",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior SDET / Technical Lead  |  Media Technologies, Test Automation & Quality Leadership  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior SDET and technical leader with 15+ years driving quality strategy and "
        "building scalable test automation frameworks for media and real-time communication "
        "systems. Hands-on experience with media technologies (WebRTC, SIP) — validating "
        "and troubleshooting real-time audio/video pipelines across backend and frontend "
        "components. Strong programming skills in C#, Java, TypeScript, Python, and C++. "
        "Designs robust automation frameworks, drives shift-left testing, and supports "
        "testing across cloud (GCP/AWS), VDI, and Windows 10/11 environments. Performs "
        "code reviews, mentors SDETs, and coordinates across multiple teams and products "
        "to ensure seamless integration and quality delivery. Passionate about media "
        "technologies, automation, and continuous improvement. Experience at IKEA, "
        "Truecaller, and LEGO. Agile/Scrum practitioner.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Media Technologies & Real-Time Communication: ",
            "Hands-on experience with WebRTC and SIP — validating, testing, and "
            "troubleshooting real-time communication systems across media backend and "
            "frontend. Built a WebRTC-based Video Collaboration Solution (VCS) at IKEA "
            "serving 30+ markets with real-time audio/video, signaling, and TURN/STUN.",
        ),
        (
            "Test Automation Frameworks: ",
            "Designed, implemented, and maintained robust, scalable test automation "
            "frameworks (Playwright, Selenium, RestAssured, Karate, Appium). CI/CD "
            "integration (GitHub Actions, Jenkins, Docker). Shift-left testing practices.",
        ),
        (
            "Programming & Quality Leadership: ",
            "Strong programming in C#, Java, TypeScript, Python, C++. Performs code "
            "reviews, mentors SDETs, and drives quality initiatives across teams. "
            "Coordinates with multiple teams/products for seamless integration.",
        ),
        (
            "Cloud, VDI & Windows: ",
            "Automated testing across GCP, AWS cloud environments, VDI setups, and "
            "Windows 10/11. Kubernetes, Docker, Terraform for infrastructure. "
            "Cross-environment test execution and validation.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Media & Real-Time: ",
            "WebRTC, SIP, TURN/STUN, real-time audio/video, signaling, "
            "media backend/frontend testing, codec validation",
        ),
        (
            "Languages: ",
            "C#, Java, TypeScript, Python, C++, JavaScript, SQL",
        ),
        (
            "Test Automation: ",
            "Playwright, Selenium, Appium (iOS/Android), RestAssured, Karate, "
            "TestNG, Jest, Postman, E2E, API, regression, contract testing",
        ),
        (
            "CI/CD & Cloud: ",
            "GitHub Actions, Jenkins, Docker, Kubernetes, GCP, AWS, Azure, "
            "Terraform, Git, VDI, Windows 10/11",
        ),
        (
            "AI-Driven Testing: ",
            "Claude, Copilot, Gemini — AI test generation, code review, "
            "exploratory testing agents, workflow automation",
        ),
        (
            "Tools & Frameworks: ",
            "TestRail, Grafana, JMeter, Maven, .NET, Jira, Confluence",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Senior SDET / Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — WebRTC-Based Video Collaboration Solution (SaaS, Media)", bold=True, size=10)
    ikea_bullets = [
        "Owned quality strategy for a WebRTC-based real-time communication platform — "
        "media backend and frontend components. Validated WebRTC signaling, TURN/STUN, "
        "SIP interoperability, audio/video quality, and codec handling across 30+ markets.",
        "Designed, implemented, and maintained scalable test automation frameworks — "
        "Playwright E2E, API testing (RestAssured, Karate), regression suites. "
        "Integrated into CI/CD pipelines (GitHub Actions, Docker, Terraform).",
        "Drove shift-left testing practices — collaborated with developers, product "
        "managers, and QA engineers. Performed code reviews, mentored SDETs, ensured "
        "adherence to testing best practices.",
        "Automated testing across cloud (GCP) and Windows environments. Kubernetes, "
        "Docker, infrastructure-as-code. Continuous improvement of test processes, "
        "tools, and methodologies.",
        "Coordinated with multiple teams and products to ensure seamless integration "
        "and quality delivery. AI-driven testing — Claude, Copilot, Gemini daily for "
        "test generation and workflow automation. 30% velocity improvement.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior SDET / QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  SaaS — 300M+ Users, Real-Time Communication & Mobile", bold=True, size=10)
    tc_bullets = [
        "Test automation for a real-time communication platform (VoIP/WebRTC-adjacent) "
        "with 300M+ users — Selenium, Appium (iOS/Android), RestAssured, Karate, TestNG. "
        "E2E, regression, API, and media testing.",
        "Mentored SDETs and developers on testing best practices. Code reviews. "
        "Coordinated cross-team integration testing. CI/CD pipelines. Performance "
        "testing (JMeter). Agile (Scrum/Kanban).",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Web, Mobile, Multi-Platform", bold=True, size=10)
    lego_bullets = [
        "LEGO: Designed automation frameworks — Selenium, RestAssured, Karate, TestNG, "
        "Appium (iOS/Android). E2E, regression, API, integration testing. Led 8–10 "
        "engineers. Code reviews, mentored SDETs. Maven, SQL, TestRail.",
        "IKEA App, Genesys (cloud contact center / SIP), Verint (2018–2021): Test "
        "automation across media-adjacent products — SIP-based contact center, voice "
        "quality validation, cross-team coordination. CI/CD, shift-left, Windows environments.",
        "Quality leadership across cross-functional, global teams. Drove shift-left "
        "testing practices and continuous improvement. Agile (Scrum).",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech & Government — Full-Stack, Multi-Platform", bold=True, size=10)
    fin_bullets = [
        "Built automation frameworks from scratch — Selenium, RestAssured, TestNG. "
        "E2E, regression, API, integration testing across enterprise platforms "
        "(Finacle). CI/CD, Git, Windows environments, SQL.",
        "Mentored 15+ engineers. Led teams of 10+. Code reviews. Quality "
        "leadership, shift-left, Agile (Scrum/Kanban).",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "WebRTC/SIP media quality ownership — built and owned quality strategy for a "
        "real-time communication platform (IKEA VCS) serving 30+ markets. Validated "
        "signaling, TURN/STUN, codec, audio/video quality.",
        "Scalable automation frameworks — Playwright, Selenium, Appium, RestAssured, "
        "Karate across web, mobile, API, and media layers.",
        "CI/CD and shift-left — GitHub Actions, Jenkins, Docker, Kubernetes. Automated "
        "testing across GCP, AWS, VDI, and Windows 10/11 environments.",
        "Quality leadership — mentored SDETs and developers, performed code reviews, "
        "coordinated cross-team integration across multiple products. Drove continuous "
        "improvement of test processes and methodologies.",
        "AI-driven testing — Claude, Copilot, Gemini daily. AI test generation, "
        "exploratory agents. 30% velocity improvement. Team rollout.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Media SDET / Technical Lead Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior SDET / Technical Lead | Media Technologies, Test Automation &amp; Quality Leadership | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Senior <span class="hl">SDET</span> and technical leader with 15+ years driving <span class="hl">quality strategy</span> and building <span class="hl">scalable</span> <span class="hl">test automation</span> frameworks for <span class="hl">media</span> and <span class="hl">real-time communication</span> systems. Hands-on with <span class="hl">media technologies</span> (<span class="hl">WebRTC</span>, <span class="hl">SIP</span>) — validating and troubleshooting real-time audio/video. Strong programming: <span class="hl">C#</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>, <span class="hl">Python</span>, <span class="hl">C++</span>. <span class="hl">Shift-left</span> testing. <span class="hl">Code review</span>s, <span class="hl">mentor</span>s SDETs. <span class="hl">Cloud</span> (<span class="hl">GCP</span>/<span class="hl">AWS</span>), <span class="hl">VDI</span>, <span class="hl">Windows</span> 10/11. Coordinates <span class="hl">cross-team</span> <span class="hl">integration</span>. <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. <span class="hl">Agile</span>/<span class="hl">Scrum</span>.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Media Technologies &amp; Real-Time Communication:</b> Hands-on <span class="hl">WebRTC</span> and <span class="hl">SIP</span> — built and tested a WebRTC-based Video Collaboration Solution (IKEA VCS) serving 30+ markets. Signaling, TURN/STUN, audio/video quality, codec validation.<br>
  <b>Test Automation Frameworks:</b> Designed, implemented, maintained robust <span class="hl">scalable</span> frameworks — Playwright, Selenium, RestAssured, Karate, Appium. <span class="hl">CI/CD</span> (GitHub Actions, Jenkins, Docker). <span class="hl">Shift-left</span>.<br>
  <b>Programming &amp; Quality Leadership:</b> <span class="hl">C#</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>, <span class="hl">Python</span>, <span class="hl">C++</span>. <span class="hl">Code review</span>s, <span class="hl">mentor</span>s SDETs, drives <span class="hl">quality</span> initiatives. <span class="hl">Cross-team</span> coordination.<br>
  <b>Cloud, VDI &amp; Windows:</b> Testing across <span class="hl">GCP</span>, <span class="hl">AWS</span>, <span class="hl">VDI</span>, and <span class="hl">Windows</span> 10/11. <span class="hl">Kubernetes</span>, <span class="hl">Docker</span>, <span class="hl">Terraform</span>.</p>

  <div class="section">Technical Skills</div>
  <p><b>Media &amp; Real-Time:</b> WebRTC, SIP, TURN/STUN, real-time audio/video, signaling, codec validation<br>
  <b>Languages:</b> C#, Java, TypeScript, Python, C++, JavaScript, SQL<br>
  <b>Test Automation:</b> Playwright, Selenium, Appium (iOS/Android), RestAssured, Karate, TestNG, Jest, Postman<br>
  <b>CI/CD &amp; Cloud:</b> GitHub Actions, Jenkins, Docker, Kubernetes, GCP, AWS, Azure, Terraform, Git, VDI, Windows 10/11<br>
  <b>AI-Driven Testing:</b> Claude, Copilot, Gemini — test generation, code review, exploratory agents<br>
  <b>Tools:</b> TestRail, Grafana, JMeter, Maven, .NET, Jira, Confluence</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Senior SDET / Technical Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — WebRTC-Based Video Collaboration Solution (SaaS, Media)</div>
  <ul>
    <li>Owned <span class="hl">quality strategy</span> for <span class="hl">WebRTC</span>-based <span class="hl">real-time communication</span> platform — <span class="hl">media</span> backend &amp; frontend. Validated <span class="hl">WebRTC</span> signaling, TURN/STUN, <span class="hl">SIP</span> interop, audio/video quality. 30+ markets.</li>
    <li><span class="hl">Scalable</span> <span class="hl">test automation</span> frameworks — <span class="hl">Playwright</span> E2E, <span class="hl">API</span> (<span class="hl">RestAssured</span>, <span class="hl">Karate</span>), regression. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, <span class="hl">Terraform</span>).</li>
    <li><span class="hl">Shift-left</span>. <span class="hl">Code review</span>s. <span class="hl">Mentor</span>ed SDETs. <span class="hl">Cloud</span> (<span class="hl">GCP</span>), <span class="hl">Windows</span>. <span class="hl">Cross-team</span> <span class="hl">integration</span>. AI-driven testing (Claude, Copilot, Gemini). 30% velocity.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior SDET / QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | SaaS — 300M+ Users, Real-Time Communication &amp; Mobile</div>
  <ul>
    <li><span class="hl">Test automation</span> for <span class="hl">real-time communication</span> platform (VoIP/<span class="hl">WebRTC</span>-adjacent) — <span class="hl">Selenium</span>, <span class="hl">Appium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span>. E2E, regression, <span class="hl">API</span>, <span class="hl">media</span> testing.</li>
    <li><span class="hl">Mentor</span>ed SDETs. <span class="hl">Code review</span>s. <span class="hl">Cross-team</span> <span class="hl">integration</span>. <span class="hl">CI/CD</span>. JMeter performance. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — SDET / Test Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Web, Mobile, Multi-Platform</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Automation frameworks</span> — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span>, <span class="hl">Appium</span>. E2E, regression, <span class="hl">API</span>, <span class="hl">integration</span>. 8–10 engineers. <span class="hl">Code review</span>s, <span class="hl">mentor</span>ed SDETs.</li>
    <li><span class="hl">IKEA</span> Genesys (<span class="hl">SIP</span>-based cloud contact center) &amp; Verint (2018–21): <span class="hl">Media</span>-adjacent testing — <span class="hl">SIP</span>, voice <span class="hl">quality</span>. <span class="hl">CI/CD</span>, <span class="hl">shift-left</span>, <span class="hl">Windows</span>.</li>
    <li><span class="hl">Quality</span> <span class="hl">leadership</span> across <span class="hl">cross-functional</span> global teams. Continuous improvement. <span class="hl">Agile</span> (<span class="hl">Scrum</span>).</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior SDET / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech &amp; Government — Full-Stack, Multi-Platform</div>
  <ul>
    <li><span class="hl">Automation frameworks</span> — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, TestNG. E2E, regression, <span class="hl">API</span>, <span class="hl">integration</span>. <span class="hl">CI/CD</span>, <span class="hl">Git</span>, <span class="hl">Windows</span>.</li>
    <li><span class="hl">Mentor</span>ed 15+ engineers. <span class="hl">Code review</span>s. <span class="hl">Quality</span> <span class="hl">leadership</span>. <span class="hl">Shift-left</span>. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">WebRTC</span>/<span class="hl">SIP</span> <span class="hl">media</span> <span class="hl">quality</span> <span class="hl">ownership</span> — <span class="hl">real-time communication</span> platform (IKEA VCS), 30+ markets. Signaling, TURN/STUN, codec, audio/video.</li>
    <li><span class="hl">Scalable</span> <span class="hl">automation frameworks</span> — <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Appium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span> across web, mobile, <span class="hl">API</span>, <span class="hl">media</span>.</li>
    <li><span class="hl">CI/CD</span> &amp; <span class="hl">shift-left</span> — <span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>. <span class="hl">GCP</span>, <span class="hl">AWS</span>, <span class="hl">VDI</span>, <span class="hl">Windows</span> 10/11.</li>
    <li><span class="hl">Quality</span> <span class="hl">leadership</span> — <span class="hl">mentor</span>ed SDETs, <span class="hl">code review</span>s, <span class="hl">cross-team</span> coordination, continuous improvement.</li>
    <li>AI-driven testing — Claude, Copilot, Gemini. 30% velocity. Team rollout.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
