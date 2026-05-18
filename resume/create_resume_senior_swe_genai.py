from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Senior_SWE_GenAI_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Senior_SWE_GenAI_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "Python",
    "FastAPI",
    "RAG",
    "Retrieval-Augmented Generation",
    "GenAI",
    "Generative AI",
    "LangChain",
    "LangGraph",
    "AutoGen",
    "multi-agent",
    "agent-based",
    "prompt engineering",
    "prompt design",
    "microservice",
    "microservices",
    "OpenAPI",
    "REST API",
    "API",
    "CI/CD",
    "Terraform",
    "Infrastructure as Code",
    "IaC",
    "AWS",
    "Azure",
    "GCP",
    "cloud-native",
    "Docker",
    "Kubernetes",
    "observability",
    "scalable",
    "resilient",
    "production-grade",
    "backend",
    "full-stack",
    "AI compliance",
    "responsible AI",
    "guardrails",
    "data privacy",
    "open-source",
    "IKEA",
    "Truecaller",
    "LEGO",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior Software Engineer  |  Python, GenAI/RAG, Cloud-Native Microservices  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior Software Engineer with 15+ years building and operating production-grade "
        "backend and full-stack systems. Advanced Python proficiency — designs and develops "
        "high-performance services using FastAPI, with hands-on Generative AI experience: "
        "RAG pipelines, GenAI agent-based workflows, multi-agent orchestration "
        "(LangChain/LangGraph), and advanced prompt engineering. Architects microservice-based "
        "systems with OpenAPI/REST APIs that are reliable, observable, secure, and scalable. "
        "Strong CI/CD and Infrastructure as Code practice (Terraform, GitHub Actions, Docker, "
        "Kubernetes). Cloud-native on GCP, AWS, and Azure — builds cost-aware, resilient "
        "services. AI compliance and responsible AI awareness — guardrails, data privacy, "
        "auditability. Open-source contributor (GitHub). Deep technical ownership with "
        "continuous improvement mindset. Currently at IKEA IT AB (3+ years).",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Python & High-Performance Backend: ",
            "Advanced Python — builds production-grade services with FastAPI. "
            "Designs reliable, observable, secure APIs. Understands trade-offs in "
            "scaling real-world systems. Also proficient in TypeScript/Node.js, "
            "Java, C#. Currently at IKEA — built Python data-ingestion pipelines, "
            "FastAPI services, and backend microservices serving 30+ markets.",
        ),
        (
            "Generative AI / RAG / Multi-Agent: ",
            "Hands-on GenAI implementation — RAG (Retrieval-Augmented Generation) "
            "pipelines, GenAI plugin and agent-based workflows. Experience with "
            "multi-agent frameworks (LangChain/LangGraph) for agent coordination, "
            "handoffs, and stateful orchestration. Advanced prompt engineering — "
            "prompt design techniques, patterns, and trade-offs. AI-assisted "
            "development (Claude, Copilot, Gemini) — 30% velocity improvement.",
        ),
        (
            "Microservices & OpenAPI/REST: ",
            "Designs and implements microservice-based architectures. Defines, "
            "builds, and consumes OpenAPI/REST APIs. Event-driven patterns "
            "(Pub/Sub, EventArc). API gateway management. Contract testing. "
            "Service observability (Grafana, Cloud Monitoring, structured logging).",
        ),
        (
            "CI/CD & Infrastructure as Code: ",
            "Establishes and maintains CI/CD pipelines with strong quality gates "
            "(GitHub Actions, Jenkins). Infrastructure as Code (Terraform — GCP, "
            "AWS). Docker, Kubernetes. Automated testing in pipelines (unit, "
            "integration, E2E). GitOps. Environment management (dev/staging/prod).",
        ),
        (
            "Cloud-Native & Security: ",
            "GCP (primary — Associate Cloud Engineer certified), AWS (Cloud "
            "Practitioner), Azure. Scalable, secure, cost-aware, resilient "
            "cloud-native services. Data privacy awareness. AI compliance and "
            "responsible AI — guardrails, auditability, risk classification. "
            "Security background (Certified Ethical Hacker).",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Languages: ",
            "Python (advanced), TypeScript, Java, C#, SQL, Bash",
        ),
        (
            "Frameworks & Libraries: ",
            "FastAPI, LangChain, LangGraph, Flask, Node.js/Express, "
            "Playwright, Selenium, RestAssured, pytest, vitest",
        ),
        (
            "GenAI & AI/ML: ",
            "RAG pipelines, multi-agent orchestration, GenAI plugins, "
            "prompt engineering, embeddings, vector stores, LLM integration "
            "(OpenAI, Claude, Gemini), responsible AI, AI compliance",
        ),
        (
            "Cloud & Infrastructure: ",
            "GCP (Cloud Run, BigQuery, Pub/Sub, Cloud Functions, Secret Manager), "
            "AWS, Azure, Docker, Kubernetes, Terraform (IaC), GitHub Actions, Jenkins",
        ),
        (
            "Architecture & APIs: ",
            "Microservices, OpenAPI/REST, event-driven (Pub/Sub, EventArc), "
            "API gateway, contract testing, observability (Grafana, Cloud Monitoring), "
            "structured logging, scalable system design",
        ),
        (
            "Practices: ",
            "CI/CD with quality gates, GitOps, TDD, code review, "
            "Agile (Scrum/SAFe), open-source contribution, security (CEH)",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Senior Software Engineer / SDET", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Cloud-Native Microservices Platform, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Python backend engineering — designed and developed production-grade Python "
        "services (FastAPI, Flask) for data-ingestion pipelines, API services, and "
        "automation workflows. High-performance, reliable, observable APIs serving "
        "30+ markets. OpenAPI/REST API design and contract testing.",
        "GenAI and RAG integration — implemented RAG pipelines and agent-based "
        "workflows using LangChain/LangGraph for intelligent test generation and "
        "quality automation. Advanced prompt engineering (Claude, Copilot, Gemini) — "
        "30% development velocity improvement. AI-assisted code review and analysis.",
        "Microservice architecture — designed and operated cloud-native microservices "
        "(Docker, Kubernetes, GCP Cloud Run). Event-driven patterns (Pub/Sub, EventArc). "
        "Observability (Grafana, Cloud Monitoring, structured logging). Scalable, "
        "resilient, cost-aware system design.",
        "CI/CD and Infrastructure as Code — established CI/CD pipelines (GitHub Actions) "
        "with strong quality gates (automated unit, integration, E2E tests). Terraform "
        "for GCP infrastructure. Docker containerization. GitOps. Environment management "
        "(dev/staging/prod). Deep technical ownership and continuous improvement.",
        "Cloud-native on GCP — Cloud Run, BigQuery, Pub/Sub, Cloud Functions, Secret "
        "Manager, Artifact Registry. Security-aware development. Data pipeline "
        "engineering (Python, BigQuery, GCS). Full-stack contributions (TypeScript).",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior Software Engineer / QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Backend Platform — 300M+ Users, Global Scale", bold=True, size=10)
    tc_bullets = [
        "Production-grade backend systems at scale — API development and testing for "
        "a platform serving 300M+ users. REST API design, microservice integration, "
        "performance optimization. Python and Java. Observability and monitoring. "
        "AWS cloud-native. CI/CD pipelines with quality gates.",
        "Technical ownership — end-to-end responsibility for service reliability, "
        "scalability, and security. Cross-team coordination. Agile (Scrum/Kanban). "
        "Mentored engineers in backend engineering practices.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Senior Engineer / Tech Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Microservices, Cloud, APIs", bold=True, size=10)
    lego_bullets = [
        "LEGO: Backend and API engineering for cloud e-commerce platform — designed "
        "microservice-based architecture, OpenAPI/REST APIs, event-driven integrations. "
        "Java, Python, TypeScript. CI/CD pipelines (Jenkins). Docker. Led 8–10 engineers. "
        "Production operations, observability, release management.",
        "IKEA (2018–2021): Full-stack engineering across multiple IKEA projects (IKEA App, "
        "Genesys, Verint/CSSP, Spartacus). API development, microservice integration, "
        "cloud-native services. CI/CD. Security-aware development. Cross-functional "
        "team coordination. Agile (Scrum).",
        "Technical ownership — continuous improvement of backend services, API reliability, "
        "and deployment pipelines. Scalable system design. Open-source tooling adoption.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior Engineer / Tech Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Financial Services — Enterprise Backend, APIs", bold=True, size=10)
    fin_bullets = [
        "Enterprise backend engineering — built and operated production-grade services "
        "for banking platforms (Finacle). API development (REST, SOAP), database "
        "engineering (SQL, Oracle), enterprise integration patterns. Java, C#, Python. "
        "CI/CD. Data privacy and compliance in regulated environments.",
        "Engineering leadership — mentored 15+ engineers. Led teams of 10+. Technical "
        "ownership of backend services. Continuous improvement. Agile (Scrum/Kanban). "
        "Cross-functional coordination.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── GenAI & Innovation ──
    add_section_heading(document, "GenAI & Innovation")
    genai_items = [
        "RAG pipelines — implemented Retrieval-Augmented Generation using LangChain, "
        "vector stores, and embeddings for intelligent document retrieval and quality "
        "automation. Integrated with OpenAI and Claude APIs.",
        "Multi-agent orchestration — built agent-based workflows using LangGraph for "
        "coordinating agents, handling handoffs, and stateful orchestration. GenAI "
        "plugins for code generation, test generation, and analysis.",
        "Prompt engineering — advanced prompt design techniques (chain-of-thought, "
        "few-shot, system prompts). Deep understanding of patterns and trade-offs. "
        "Applied across Claude, Copilot, and Gemini for 30% velocity improvement.",
        "AI compliance and responsible AI — guardrails implementation, data privacy "
        "awareness, auditability, risk classification for AI systems. Ethical AI "
        "practices in production workflows.",
        "Open-source — contributor on GitHub (ascertain.github.io). Internal framework "
        "development and tooling. Innovation-driven continuous improvement.",
    ]
    for item in genai_items:
        add_highlighted_bullet(document, item)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ISTQB Certified Tester Foundation\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent (Professional)  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Senior SWE GenAI Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior Software Engineer | Python, GenAI/RAG, Cloud-Native Microservices | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Senior Software Engineer with 15+ years building and operating <span class="hl">production-grade</span> <span class="hl">backend</span> and <span class="hl">full-stack</span> systems. Advanced <span class="hl">Python</span> proficiency — <span class="hl">FastAPI</span>, <span class="hl">GenAI</span>: <span class="hl">RAG</span> pipelines, <span class="hl">agent-based</span> workflows, <span class="hl">multi-agent</span> orchestration (<span class="hl">LangChain</span>/<span class="hl">LangGraph</span>), <span class="hl">prompt engineering</span>. <span class="hl">Microservice</span> architecture, <span class="hl">OpenAPI</span>/<span class="hl">REST API</span>s — reliable, <span class="hl">observable</span>, secure, <span class="hl">scalable</span>. <span class="hl">CI/CD</span> with quality gates, <span class="hl">Terraform</span> (<span class="hl">IaC</span>). <span class="hl">Cloud-native</span> on <span class="hl">GCP</span>, <span class="hl">AWS</span>, <span class="hl">Azure</span>. <span class="hl">AI compliance</span>, <span class="hl">responsible AI</span>, <span class="hl">data privacy</span>. <span class="hl">Open-source</span> contributor. Currently at <span class="hl">IKEA</span> IT AB (3+ years).</p>

  <div class="section">How I Match the Role</div>
  <p><b>Python &amp; Backend:</b> Advanced <span class="hl">Python</span> — <span class="hl">production-grade</span> <span class="hl">FastAPI</span> services. Reliable, <span class="hl">observable</span>, secure <span class="hl">API</span>s. Scaling trade-offs. Also TypeScript, Java, C#. <span class="hl">IKEA</span> — <span class="hl">Python</span> data pipelines, <span class="hl">FastAPI</span>, <span class="hl">backend</span> <span class="hl">microservices</span> (30+ markets).<br>
  <b>GenAI / RAG / Multi-Agent:</b> <span class="hl">RAG</span> pipelines, <span class="hl">GenAI</span> <span class="hl">agent-based</span> workflows, <span class="hl">multi-agent</span> frameworks (<span class="hl">LangChain</span>/<span class="hl">LangGraph</span>) — agent coordination, handoffs, stateful orchestration. Advanced <span class="hl">prompt engineering</span>. AI-assisted dev (Claude, Copilot, Gemini) — 30% velocity gain.<br>
  <b>Microservices &amp; APIs:</b> <span class="hl">Microservice</span> architecture, <span class="hl">OpenAPI</span>/<span class="hl">REST API</span>s, event-driven (Pub/Sub). <span class="hl">Observability</span> (Grafana, Cloud Monitoring). Contract testing.<br>
  <b>CI/CD &amp; IaC:</b> <span class="hl">CI/CD</span> with quality gates (GitHub Actions, Jenkins). <span class="hl">Terraform</span> (<span class="hl">Infrastructure as Code</span>). <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>. GitOps.<br>
  <b>Cloud &amp; AI Compliance:</b> <span class="hl">GCP</span> (ACE certified), <span class="hl">AWS</span>, <span class="hl">Azure</span>. <span class="hl">Cloud-native</span>, <span class="hl">scalable</span>, <span class="hl">resilient</span>. <span class="hl">AI compliance</span>, <span class="hl">responsible AI</span>, <span class="hl">guardrails</span>, <span class="hl">data privacy</span>. Security (<span class="hl">CEH</span>).</p>

  <div class="section">Technical Skills</div>
  <p><b>Languages:</b> Python (advanced), TypeScript, Java, C#, SQL, Bash<br>
  <b>Frameworks:</b> FastAPI, LangChain, LangGraph, Flask, Node.js/Express, Playwright, pytest, vitest<br>
  <b>GenAI:</b> RAG pipelines, multi-agent orchestration, GenAI plugins, prompt engineering, embeddings, vector stores, LLM integration (OpenAI, Claude, Gemini), responsible AI<br>
  <b>Cloud &amp; Infra:</b> GCP (Cloud Run, BigQuery, Pub/Sub, Functions), AWS, Azure, Docker, Kubernetes, Terraform, GitHub Actions, Jenkins<br>
  <b>Architecture:</b> Microservices, OpenAPI/REST, event-driven, API gateway, observability (Grafana), structured logging, scalable design<br>
  <b>Practices:</b> CI/CD with quality gates, GitOps, TDD, code review, Agile (Scrum/SAFe), open-source, security (CEH)</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Senior Software Engineer / SDET</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Cloud-Native Microservices Platform, 30+ Markets</div>
  <ul>
    <li><span class="hl">Python</span> <span class="hl">backend</span> — <span class="hl">production-grade</span> <span class="hl">FastAPI</span> services, data-ingestion pipelines, <span class="hl">API</span> services. <span class="hl">OpenAPI</span>/<span class="hl">REST API</span> design. <span class="hl">Observable</span>, reliable, <span class="hl">scalable</span>. 30+ markets.</li>
    <li><span class="hl">GenAI</span> &amp; <span class="hl">RAG</span> — <span class="hl">RAG</span> pipelines, <span class="hl">agent-based</span> workflows (<span class="hl">LangChain</span>/<span class="hl">LangGraph</span>). <span class="hl">Prompt engineering</span> (Claude, Copilot, Gemini) — 30% velocity. AI-assisted code review.</li>
    <li><span class="hl">Microservice</span> architecture — <span class="hl">cloud-native</span> (<span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">GCP</span> Cloud Run). Event-driven (Pub/Sub). <span class="hl">Observability</span> (Grafana). <span class="hl">Scalable</span>, <span class="hl">resilient</span> design.</li>
    <li><span class="hl">CI/CD</span> &amp; <span class="hl">IaC</span> — GitHub Actions with quality gates. <span class="hl">Terraform</span> for <span class="hl">GCP</span>. <span class="hl">Docker</span>. GitOps. Deep technical ownership. Continuous improvement.</li>
    <li><span class="hl">Cloud-native</span> on <span class="hl">GCP</span> — Cloud Run, BigQuery, Pub/Sub, Functions, Secret Manager. Security-aware. <span class="hl">Full-stack</span> (TypeScript). Data pipelines (<span class="hl">Python</span>, BigQuery).</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior Software Engineer / QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Backend Platform — 300M+ Users, Global Scale</div>
  <ul>
    <li><span class="hl">Production-grade</span> <span class="hl">backend</span> at scale — <span class="hl">API</span> development, <span class="hl">REST API</span> design, <span class="hl">microservice</span> integration, performance optimization. <span class="hl">Python</span>, Java. <span class="hl">Observability</span>. <span class="hl">AWS</span> <span class="hl">cloud-native</span>. <span class="hl">CI/CD</span>.</li>
    <li>Technical ownership — service reliability, <span class="hl">scalable</span>, secure. Cross-team. <span class="hl">Agile</span>. Mentored engineers.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Senior Engineer / Tech Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Microservices, Cloud, APIs</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Backend</span> &amp; <span class="hl">API</span> engineering — <span class="hl">microservice</span> architecture, <span class="hl">OpenAPI</span>/<span class="hl">REST API</span>s, event-driven. Java, <span class="hl">Python</span>, TypeScript. <span class="hl">CI/CD</span> (Jenkins). <span class="hl">Docker</span>. 8–10 engineers. <span class="hl">Observability</span>, release management.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Full-stack</span> engineering — <span class="hl">API</span> dev, <span class="hl">microservice</span> integration, <span class="hl">cloud-native</span>. <span class="hl">CI/CD</span>. Security-aware. <span class="hl">Agile</span>.</li>
    <li>Technical ownership — continuous improvement, <span class="hl">API</span> reliability, deployment pipelines. <span class="hl">Scalable</span> design. <span class="hl">Open-source</span> tooling.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior Engineer / Tech Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Financial Services — Enterprise Backend, APIs</div>
  <ul>
    <li>Enterprise <span class="hl">backend</span> — <span class="hl">production-grade</span> services (Finacle banking). <span class="hl">API</span> dev (REST, SOAP), SQL, enterprise integration. Java, C#, <span class="hl">Python</span>. <span class="hl">CI/CD</span>. <span class="hl">Data privacy</span> and compliance.</li>
    <li>Engineering leadership — 15+ engineers. Technical ownership. Continuous improvement. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="section">GenAI &amp; Innovation</div>
  <ul>
    <li><span class="hl">RAG</span> pipelines — <span class="hl">Retrieval-Augmented Generation</span> using <span class="hl">LangChain</span>, vector stores, embeddings. OpenAI &amp; Claude <span class="hl">API</span> integration.</li>
    <li><span class="hl">Multi-agent</span> orchestration — <span class="hl">agent-based</span> workflows (<span class="hl">LangGraph</span>), agent coordination, handoffs, stateful orchestration. <span class="hl">GenAI</span> plugins.</li>
    <li><span class="hl">Prompt engineering</span> — chain-of-thought, few-shot, system prompts. Patterns &amp; trade-offs. Claude, Copilot, Gemini. 30% velocity.</li>
    <li><span class="hl">AI compliance</span> &amp; <span class="hl">responsible AI</span> — <span class="hl">guardrails</span>, <span class="hl">data privacy</span>, auditability, risk classification.</li>
    <li><span class="hl">Open-source</span> — GitHub contributor. Internal framework development. Innovation-driven.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>Google Cloud – Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ISTQB Certified Tester Foundation<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent (Professional) | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
