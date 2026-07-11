"""
Resume: Mariem Lamouni — Product Specialist (Senior) CMS/DAM
Reframed Dqodify as Business Analyst / Pimcore CMS. EU template style. DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mariem_Lamouni_Product_Specialist_CMS_DAM_Resume"
TEAL = RGBColor(0x00, 0x7A, 0x87)  # Professional teal


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(11)
    p.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="007A87"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEAL
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MARIEM LAMOUNI")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = TEAL

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(4)
    r2 = p2.add_run("Malmö, Sweden  \u2022  +46 762432366  \u2022  mariemlamouni30@gmail.com")
    r2.font.size = Pt(10)

    # ── PROFESSIONAL PROFILE ──
    add_section_heading(doc, "Professional Profile")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Versatile professional with experience in ", False),
        ("CMS platforms (Pimcore)", True),
        (", business analysis, ", False),
        ("content management workflows", True),
        (", and ", False),
        ("stakeholder coordination", True),
        (". Background spanning digital product support, ", False),
        ("tool training", True),
        (", ", False),
        ("backlog coordination", True),
        (", and cross-functional collaboration. Comfortable working in ", False),
        ("market-facing support roles", True),
        (" — enabling teams through ", False),
        ("coaching", True),
        (", documentation, and hands-on platform guidance. Strong communicator fluent in English, French, and Arabic, with proven ability to bridge technical and business teams.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY SKILLS ──
    add_section_heading(doc, "Key Skills & Competencies")
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    skills = [
        "CMS / DAM Platforms", "Pimcore (Hands-on)", "Content Workflow Management",
        "Backlog Coordination", "Tool Training & Coaching", "Market-Facing Support",
        "Stakeholder Communication", "Cross-Functional Collaboration", "Documentation & Processes",
    ]
    for i, skill in enumerate(skills):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {skill}")
        r.font.size = Pt(10)
        r.bold = True

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- Dqodify (reframed) ---
    role_header(doc,
        "Dqodify Pvt Ltd — Business Analyst / CMS Specialist",
        "Jun 2024 \u2013 Present",
        "Pimcore CMS Platform \u2014 Digital Content & Product Data Management")
    add_bullet(doc, "Support and coordinate Pimcore CMS platform operations \u2014 content modelling, asset management (DAM), and digital product data workflows.",
        ["Pimcore CMS", "content modelling", "asset management (DAM)", "digital product data"])
    add_bullet(doc, "Coordinate product backlog: gather requirements from stakeholders, prioritise items, and track delivery progress across sprints.",
        ["product backlog", "gather requirements", "prioritise items"])
    add_bullet(doc, "Provide tool training and coaching to team members on CMS/DAM platform features, content publishing workflows, and best practices.",
        ["tool training and coaching", "CMS/DAM platform", "content publishing workflows"])
    add_bullet(doc, "Collaborate with development teams on platform configuration, headless content delivery setup, and market-specific adaptations.",
        ["headless content delivery", "market-specific adaptations"])
    add_bullet(doc, "Document processes, create user guides, and facilitate knowledge transfer for market support teams.",
        ["Document processes", "user guides", "knowledge transfer", "market support"])
    add_bullet(doc, "Participate in business analysis activities: requirements gathering, user story writing, acceptance criteria definition.",
        ["business analysis", "requirements gathering", "user story writing", "acceptance criteria"])

    # --- Uniq Dialog ---
    role_header(doc,
        "Uniq Dialog (for Hello Fresh) — Sales & Customer Engagement Representative",
        "Sep 2025",
        "Customer Re-engagement & Relationship Management")
    add_bullet(doc, "Re-engaged customers through outbound communication \u2014 active listening, needs identification, and tailored solutions.",
        ["active listening", "needs identification"])
    add_bullet(doc, "Built strong customer relationships emphasising product quality and service standards; ensured brand satisfaction.",
        ["customer relationships", "brand satisfaction"])

    # --- Arotechnic Industries ---
    role_header(doc,
        "Arotechnic Industries (ATI Hangars) — Quality & Technical Intern",
        "Nov 2022 \u2013 Dec 2023",
        "Quality Control, Documentation & Data Analysis")
    add_bullet(doc, "Assisted with quality control procedures, technical documentation management, and data analysis workflows.",
        ["quality control", "technical documentation", "data analysis"])
    add_bullet(doc, "Worked with process manuals and technical documents \u2014 structured documentation and procedural compliance.",
        ["process manuals", "structured documentation"])

    # --- Societe Melamdhygiene ---
    role_header(doc,
        "Soci\u00e9t\u00e9 Melamdhygi\u00e8ne — Executive Assistant / Operations Coordinator",
        "Jan 2019 \u2013 Dec 2020",
        "Administration, Vendor Management & Internal Operations")
    add_bullet(doc, "Managed communication with end-users, vendors, and internal teams \u2014 coordination across multiple stakeholders.",
        ["end-users, vendors, and internal teams", "coordination"])
    add_bullet(doc, "Provided administrative support, partner collaboration, and ensured smooth daily operations.",
        ["administrative support", "partner collaboration"])

    # ── EDUCATION ──
    add_section_heading(doc, "Education")

    ep1 = doc.add_paragraph()
    ep1.paragraph_format.space_before = Pt(4)
    r = ep1.add_run("Aeronautical Maintenance Technician")
    r.bold = True
    r.font.size = Pt(10)
    ep1.add_run("  \u2014  Specialized Institute of Aeronautical Trades & Airport Logistics | 2021\u20132022").font.size = Pt(9.5)

    ep2 = doc.add_paragraph()
    ep2.paragraph_format.space_before = Pt(3)
    r = ep2.add_run("Specialized Technician Diploma \u2014 Mechatronics")
    r.bold = True
    r.font.size = Pt(10)
    ep2.add_run("  \u2014  Specialized Institute of Aeronautics & Airport Logistics | 2016\u20132018").font.size = Pt(9.5)

    ep3 = doc.add_paragraph()
    ep3.paragraph_format.space_before = Pt(3)
    r = ep3.add_run("Scientific Baccalaur\u00e9at (Mathematics A)")
    r.bold = True
    r.font.size = Pt(10)
    ep3.add_run("  \u2014  2015\u20132016").font.size = Pt(9.5)

    # ── LANGUAGES & ADDITIONAL ──
    add_section_heading(doc, "Languages & Additional")

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(4)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  French (Fluent)  |  Arabic (Native)  |  Swedish (Basic)").font.size = Pt(10)

    dp = doc.add_paragraph()
    dp.paragraph_format.space_before = Pt(3)
    r = dp.add_run("Driving License: ")
    r.bold = True
    r.font.size = Pt(10)
    dp.add_run("Valid Swedish Driving License (B)").font.size = Pt(10)

    ip = doc.add_paragraph()
    ip.paragraph_format.space_before = Pt(3)
    r = ip.add_run("Interests: ")
    r.bold = True
    r.font.size = Pt(10)
    ip.add_run("Painting, Baking, Digital Content Creation").font.size = Pt(10)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
