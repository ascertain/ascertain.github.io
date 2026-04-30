from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Quinyx_Senior_QA_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Quinyx_Senior_QA_Engineer_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "15+ years",
    "Senior QA Engineer",
    "end-to-end",
    "test strategy",
    "automation",
    "AI-powered",
    "AI tool",
    "AI-assisted",
    "Claude",
    "Gemini",
    "Copilot",
    "Selenium",
    "RestAssured",
    "Appium",
    "TestNG",
    "Playwright",
    "K6",
    "performance testing",
    "load testing",
    "Java",
    "Spring Boot",
    "React",
    "full stack",
    "full-stack",
    "CI/CD",
    "GitHub Actions",
    "Docker",
    "Kubernetes",
    "K8s",
    "AWS",
    "Maven",
    "SQL",
    "TestRail",
    "regression",
    "exploratory",
    "acceptance",
    "system testing",
    "manual testing",
    "mentor",
    "coaching",
    "Agile",
    "Scrum",
    "cross-functional",
    "autonomous",
    "ownership",
    "quality",
    "SaaS",
    "WFM",
    "collaboration",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior QA Engineer | Full-Stack Quality, Automation & AI-Powered QA | 15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior QA Engineer with 15+ years of hands-on test engineering experience, taking "
        "end-to-end ownership of quality across full-stack products (Java/Spring Boot + React). "
        "Shapes test strategy, drives automation at scale (Selenium, RestAssured, Appium, TestNG, "
        "Playwright), and champions AI-powered QA workflows using Claude, Gemini, and Copilot "
        "daily. Experienced in performance testing (K6), CI/CD integration (GitHub Actions, Docker, "
        "Kubernetes, AWS), and mentoring QA engineers and developers on modern testing practices. "
        "Daily AI tool user who has rolled out AI-assisted workflows to teams. Proactive, "
        "collaborative, passionate about quality — not just checking boxes. Thrives in autonomous, "
        "cross-functional teams with Agile ceremonies. SaaS product background across workforce "
        "management, e-commerce, and fintech.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "End-to-End Quality Ownership: ",
            "Own quality across full-stack products — from test strategy and test case creation "
            "through manual testing (regression, exploratory, acceptance, system testing) to "
            "automated suites and CI/CD integration. Not a 'write tests and file bugs' approach.",
        ),
        (
            "Automation at Scale: ",
            "Built and maintained automated test suites with Selenium, RestAssured, Appium, "
            "TestNG, and Playwright. Integrated testing into CI/CD pipelines (GitHub Actions, "
            "Docker, Kubernetes, AWS). Strong Java and Maven skills.",
        ),
        (
            "AI-Powered QA Workflows: ",
            "Daily AI tool user — Claude, Gemini, Copilot integrated into daily QA workflows. "
            "Rolled out AI-assisted test generation, code reviews, and automation to teams. "
            "Champion AI adoption across engineering.",
        ),
        (
            "Performance Testing: ",
            "Experience with performance, load, and stress testing using K6. Established "
            "performance testing practices and benchmarks in production-like environments.",
        ),
        (
            "Mentoring & Collaboration: ",
            "Natural mentor — raised the bar for QA engineers and developers on testing best "
            "practices. Collaborated with Staff Engineers and Directors on automation strategy. "
            "Active in Agile ceremonies (Scrum/Kanban) and QA community.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Automation: ",
            "Selenium, RestAssured, Appium, TestNG, Playwright, API testing, E2E testing, "
            "regression suites, test case design",
        ),
        (
            "Full Stack: ",
            "Java, Spring Boot, React, Maven, SQL databases, OOP, TestRail",
        ),
        (
            "CI/CD & Cloud: ",
            "GitHub Actions, Docker, Kubernetes (K8s), AWS, Terraform, CI/CD pipeline design",
        ),
        (
            "Performance: ",
            "K6 (load, stress, soak testing), performance benchmarking, observability",
        ),
        (
            "AI Tools: ",
            "Claude, Gemini, GitHub Copilot — daily user, AI-assisted test generation, "
            "code reviews, workflow automation, team rollout experience",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Senior QA Engineer / Engineering Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Video Collaboration Solution (SaaS, Full-Stack)", bold=True, size=10)
    ikea_bullets = [
        "End-to-end quality ownership across a full-stack product (TypeScript/Node.js + React) "
        "— shaped test strategy, drove automation, and integrated testing into CI/CD pipelines "
        "(GitHub Actions, Docker, Terraform) serving 30+ markets.",
        "Built and maintained automated test suites — Playwright E2E, API testing (RestAssured "
        "patterns), regression, exploratory, acceptance, and system testing. Actively "
        "participated in manual testing and test case creation.",
        "Championed AI-powered QA workflows — rolled out Claude, Gemini, and Copilot for "
        "AI-assisted test generation, code reviews, and automation. Daily AI tool user who "
        "helped the entire team adopt AI workflows, improving velocity by 30%.",
        "Expanded performance testing practice — load and stress testing for API endpoints, "
        "performance benchmarking, and observability. Set up monitoring (iLert) for automated "
        "error detection and rapid triage.",
        "Mentored QA engineers and developers on testing best practices — raised the bar for "
        "quality culture. Collaborated with engineering leadership on automation strategy and "
        "roadmap. Active in Agile ceremonies (Scrum) and QA community.",
        "Worked across Kubernetes (K8s) and cloud-native infrastructure (GCP Cloud Run, Docker). "
        "SQL databases, Maven-like build tooling, and TestRail for test management.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  SaaS — 300M+ Users, Multi-Platform", bold=True, size=10)
    tc_bullets = [
        "End-to-end quality ownership at a fast-growing SaaS startup (300M+ users) — test "
        "strategy, automation (Selenium, Appium, RestAssured, TestNG), manual testing "
        "(regression, exploratory, acceptance), and CI/CD pipeline integration.",
        "Mentored QA engineers and developers on modern testing practices. Collaborated with "
        "engineering directors on automation strategy. Proactive, autonomous, cross-functional "
        "team environment with Agile (Scrum/Kanban).",
        "Performance and load testing for high-traffic APIs. Championed early AI-assisted "
        "testing workflows, managed test infrastructure, and drove quality culture.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Test Lead / Senior QA Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Java/React Full-Stack", bold=True, size=10)
    lego_bullets = [
        "LEGO: Senior QA for Java-based e-commerce platform — Selenium, RestAssured, TestNG, "
        "Appium automation suites. Managed 8–10 engineers. Manual testing (regression, "
        "exploratory, acceptance, system testing). Maven, SQL, TestRail.",
        "IKEA App, Genesys, Verint (2018–2021): End-to-end quality across multiple Java + "
        "React products — test strategy, CI/CD integration, performance testing, and "
        "mentoring developers on testing best practices.",
        "Drove automation strategy with engineering directors. Active in Agile ceremonies. "
        "Built autonomous, cross-functional quality culture — exactly the Quinyx values "
        "of passion, trust, collaboration, quality, and innovation.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior QA Engineer / Team Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech & Government — Java Full-Stack", bold=True, size=10)
    fin_bullets = [
        "Hands-on QA/test engineering across Java full-stack banking platforms (Finacle) — "
        "Selenium, TestNG, RestAssured, manual testing, regression, system testing, SQL "
        "databases. Built automation frameworks from scratch.",
        "Mentored 15+ QA engineers and developers. Led teams of 10+ engineers with "
        "autonomous ownership. Drove CI/CD adoption and Agile practices (Scrum/Kanban).",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "End-to-end quality ownership across SaaS products — shaped test strategy, drove "
        "automation (Selenium, RestAssured, Appium, TestNG, Playwright), integrated into "
        "CI/CD (GitHub Actions, Docker, K8s).",
        "AI-powered QA champion — rolled out Claude, Gemini, Copilot to teams. Daily AI user. "
        "30% velocity improvement through AI-assisted test generation and workflows.",
        "Performance testing practice — K6 patterns, load/stress testing, benchmarking for "
        "high-traffic APIs serving 300M+ users (Truecaller) and 30+ markets (IKEA).",
        "Natural mentor — raised the bar for QA engineers and developers across every "
        "organization. Built quality culture where testing is everyone's responsibility.",
        "Full-stack Java + React experience — automation across the entire stack, from API "
        "(RestAssured) through UI (Selenium/Playwright), in SaaS and enterprise environments.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "Certified Ethical Hacker (CEH)\n"
        "AWS Cloud Practitioner\nITIL Foundation\n"
        "Six Sigma Green Belt\nUiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Quinyx Senior QA Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior QA Engineer | Full-Stack Quality, Automation &amp; AI-Powered QA | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Senior QA Engineer</span> with <span class="hl">15+ years</span> of hands-on test engineering, taking <span class="hl">end-to-end</span> <span class="hl">ownership</span> of <span class="hl">quality</span> across <span class="hl">full-stack</span> products (<span class="hl">Java</span>/<span class="hl">Spring Boot</span> + <span class="hl">React</span>). Shapes <span class="hl">test strategy</span>, drives <span class="hl">automation</span> at scale (<span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Appium</span>, <span class="hl">TestNG</span>, <span class="hl">Playwright</span>), and champions <span class="hl">AI-powered</span> QA workflows using <span class="hl">Claude</span>, <span class="hl">Gemini</span>, and <span class="hl">Copilot</span> daily. <span class="hl">K6</span> <span class="hl">performance testing</span>, <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, <span class="hl">K8s</span>, <span class="hl">AWS</span>). Natural <span class="hl">mentor</span> who raises the bar. Proactive, <span class="hl">collaborative</span>, passionate about <span class="hl">quality</span>. Thrives in <span class="hl">autonomous</span>, <span class="hl">cross-functional</span> teams. <span class="hl">SaaS</span> background.</p>

  <div class="section">How I Match the Role</div>
  <p><b>End-to-End Quality Ownership:</b> Full-stack (Java + React) — test strategy, manual testing (regression, exploratory, acceptance, system), automated suites, CI/CD.<br>
  <b>Automation at Scale:</b> Selenium, RestAssured, Appium, TestNG, Playwright. CI/CD (GitHub Actions, Docker, K8s, AWS). Java, Maven.<br>
  <b>AI-Powered QA:</b> Daily Claude, Gemini, Copilot user. Rolled out AI-assisted test generation and workflows to teams. 30% velocity improvement.<br>
  <b>Performance Testing:</b> K6 patterns, load/stress testing, benchmarking for high-traffic APIs.<br>
  <b>Mentoring &amp; Collaboration:</b> Natural mentor — QA engineers and developers. Automation strategy with Staff Engineers/Directors. Agile ceremonies, QA community.</p>

  <div class="section">Technical Skills</div>
  <p><b>Automation:</b> Selenium, RestAssured, Appium, TestNG, Playwright, E2E, API testing, regression suites<br>
  <b>Full Stack:</b> Java, Spring Boot, React, Maven, SQL databases, TestRail<br>
  <b>CI/CD &amp; Cloud:</b> GitHub Actions, Docker, Kubernetes (K8s), AWS, Terraform<br>
  <b>Performance:</b> K6 (load, stress, soak), benchmarking, observability<br>
  <b>AI Tools:</b> Claude, Gemini, Copilot — daily user, AI-assisted test generation, team rollout</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Senior QA Engineer / Engineering Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Video Collaboration Solution (SaaS, Full-Stack)</div>
  <ul>
    <li><span class="hl">End-to-end</span> <span class="hl">quality</span> <span class="hl">ownership</span> — <span class="hl">test strategy</span>, <span class="hl">automation</span> (<span class="hl">Playwright</span> E2E, API testing), <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, Terraform). 30+ markets.</li>
    <li>Built automated suites — <span class="hl">regression</span>, <span class="hl">exploratory</span>, <span class="hl">acceptance</span>, <span class="hl">system testing</span>. Active <span class="hl">manual testing</span> and test case creation.</li>
    <li>Championed <span class="hl">AI-powered</span> QA — rolled out <span class="hl">Claude</span>, <span class="hl">Gemini</span>, <span class="hl">Copilot</span> to team. <span class="hl">AI-assisted</span> test generation, code reviews. 30% velocity improvement.</li>
    <li><span class="hl">Performance testing</span> — load/stress for APIs. Monitoring (iLert), <span class="hl">observability</span>, automated error detection.</li>
    <li><span class="hl">Mentor</span>ed QA engineers and developers. Collaborated with engineering leadership on <span class="hl">automation</span> strategy. <span class="hl">Agile</span> (<span class="hl">Scrum</span>). <span class="hl">K8s</span>, <span class="hl">SQL</span>, <span class="hl">Maven</span>, <span class="hl">TestRail</span>.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | SaaS — 300M+ Users, Multi-Platform</div>
  <ul>
    <li><span class="hl">End-to-end</span> <span class="hl">quality</span> at <span class="hl">SaaS</span> startup (300M+ users) — <span class="hl">Selenium</span>, <span class="hl">Appium</span>, <span class="hl">RestAssured</span>, <span class="hl">TestNG</span>. <span class="hl">Manual testing</span>, <span class="hl">CI/CD</span>.</li>
    <li><span class="hl">Mentor</span>ed engineers. <span class="hl">Automation</span> strategy with directors. <span class="hl">Performance</span>/<span class="hl">load testing</span>. <span class="hl">Autonomous</span>, <span class="hl">cross-functional</span>, <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Test Lead / Senior QA Engineer</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Java/React Full-Stack</div>
  <ul>
    <li>LEGO: <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">TestNG</span>, <span class="hl">Appium</span>. 8–10 engineers. <span class="hl">Regression</span>, <span class="hl">exploratory</span>, <span class="hl">acceptance</span>, <span class="hl">system testing</span>. <span class="hl">Maven</span>, <span class="hl">SQL</span>, <span class="hl">TestRail</span>.</li>
    <li>IKEA (2018–21): <span class="hl">End-to-end</span> <span class="hl">quality</span> across <span class="hl">Java</span> + <span class="hl">React</span> products. <span class="hl">CI/CD</span>, <span class="hl">performance testing</span>, <span class="hl">mentor</span>ing developers.</li>
    <li><span class="hl">Automation</span> strategy with directors. <span class="hl">Agile</span> (<span class="hl">Scrum</span>). <span class="hl">Autonomous</span>, <span class="hl">cross-functional</span> <span class="hl">quality</span> <span class="hl">culture</span>.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior QA Engineer / Team Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech &amp; Government — Java Full-Stack</div>
  <ul>
    <li><span class="hl">Java</span> <span class="hl">full-stack</span> banking (Finacle) — <span class="hl">Selenium</span>, <span class="hl">TestNG</span>, <span class="hl">RestAssured</span>, <span class="hl">SQL</span>. Built <span class="hl">automation</span> frameworks from scratch.</li>
    <li><span class="hl">Mentor</span>ed 15+ engineers. 10+ team. <span class="hl">Autonomous</span> <span class="hl">ownership</span>. <span class="hl">CI/CD</span>, <span class="hl">Agile</span> (<span class="hl">Scrum</span>).</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">End-to-end</span> <span class="hl">quality</span> across <span class="hl">SaaS</span> — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Appium</span>, <span class="hl">TestNG</span>, <span class="hl">Playwright</span>. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, <span class="hl">K8s</span>).</li>
    <li><span class="hl">AI-powered</span> QA champion — <span class="hl">Claude</span>, <span class="hl">Gemini</span>, <span class="hl">Copilot</span>. Daily user. 30% velocity. Team rollout.</li>
    <li><span class="hl">Performance testing</span> — <span class="hl">K6</span> patterns, high-traffic APIs (300M+ users, 30+ markets).</li>
    <li>Natural <span class="hl">mentor</span> — raised the bar across every organization. <span class="hl">Quality</span> <span class="hl">culture</span>.</li>
    <li><span class="hl">Full-stack</span> <span class="hl">Java</span> + <span class="hl">React</span> — API (<span class="hl">RestAssured</span>) through UI (<span class="hl">Selenium</span>/<span class="hl">Playwright</span>) in <span class="hl">SaaS</span>.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>Certified Ethical Hacker (CEH)<br>AWS Cloud Practitioner<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
