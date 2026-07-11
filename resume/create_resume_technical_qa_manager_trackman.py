"""
Resume: Technical QA Manager — Trackman / Unity Applications / Sports Technology
Focus: QA strategy, test automation (Unity), game development, team mentorship, HW/SW integration.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Technical_QA_Manager_Trackman_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Technical QA Manager with ", False),
        ("16+ years", True),
        (" driving ", False),
        ("quality strategy", True),
        (", ", False),
        ("test automation frameworks", True),
        (", and ", False),
        ("team leadership", True),
        (" for complex applications combining ", False),
        ("hardware/software integration", True),
        (", real-time systems, and interactive user experiences. Strong background in ", False),
        ("system-level testing", True),
        (", ", False),
        ("CI/CD integration", True),
        (", and ", False),
        ("production debugging", True),
        (". Proven ability to define and evolve ", False),
        ("comprehensive QA strategies", True),
        (" across the full product lifecycle \u2014 from concept through release and post-launch evaluation. Experienced in ", False),
        ("mentoring QA teams", True),
        (", championing quality at every development stage, and driving ", False),
        ("ROI-driven test automation", True),
        (" that supports fast iteration cycles. Passionate about ", False),
        ("sports technology", True),
        (" and interactive applications where quality directly shapes user experience.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "QA Strategy & Leadership", "Test Automation Architecture", "Hardware/Software Integration",
        "System-Level Testing", "CI/CD Pipeline Integration", "Team Mentorship & Growth",
        "Root Cause Analysis", "Production Debugging", "Agile/Scrum Methodologies",
        "Quality Framework Ownership", "Risk Identification & Mitigation", "Continuous Improvement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills_data = [
        ("QA & Testing:", " Test strategy design, system-level testing, HW/SW integration testing, simulation validation, gameplay testing"),
        ("Automation:", " Python, Pytest, Playwright, Selenium, Appium, Robot Framework, custom test frameworks for interactive applications"),
        ("CI/CD:", " GitHub Actions, Jenkins, Azure DevOps, Docker, automated test pipelines, build verification"),
        ("Game/App Tech:", " Unity (exposure), real-time application testing, device integration, graphics/rendering validation"),
        ("Debugging:", " Production environment analysis, log analysis, defect reproduction, performance profiling"),
        ("Tools:", " JIRA, XRAY, TestRail, Confluence, Git/GitHub, defect tracking & triage workflows"),
        ("Practices:", " Agile/Scrum, Kanban, risk-based testing, exploratory testing, regression automation, ROI-driven automation"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Technical QA Manager / Team Lead (Acting)",
        "Mar 2022 \u2013 Present",
        "Interactive Applications Platform \u2014 Real-Time Systems, Device Integration, 30+ Markets | Agile/Scrum")
    add_bullet(doc, "Define, implement, and continuously evolve overall QA strategy for interactive applications combining hardware devices, real-time software, and user-facing experiences.",
        ["QA strategy", "interactive applications", "hardware devices", "real-time software"])
    add_bullet(doc, "Lead the test automation strategy \u2014 scalable, efficient approach aligned with product goals; continuously assess automation opportunities and advocate for ROI-driven automation.",
        ["test automation strategy", "scalable", "ROI-driven automation"])
    add_bullet(doc, "Design, maintain, and evolve both manual and automated test frameworks supporting fast iteration cycles, simulation complexities, and multi-configuration testing.",
        ["manual and automated test frameworks", "fast iteration cycles", "simulation complexities"])
    add_bullet(doc, "Work closely with product owners, developers, designers, and domain experts to secure high quality, realism, and consistency throughout the product lifecycle.",
        ["product owners, developers, designers", "quality, realism, and consistency", "product lifecycle"])
    add_bullet(doc, "Champion quality at every stage \u2014 from concept discussions through implementation, testing, release, and post-launch evaluation.",
        ["Champion quality at every stage", "concept", "release", "post-launch evaluation"])
    add_bullet(doc, "Mentor, guide, and lead a team of QA engineers \u2014 fostering technical growth, ownership, and continuous improvement within the QA discipline.",
        ["Mentor, guide, and lead", "technical growth", "ownership", "continuous improvement"])
    add_bullet(doc, "Drive root-cause analysis and bug triage processes; promote continuous enhancement of testing practices, tools, and quality standards.",
        ["root-cause analysis", "bug triage", "continuous enhancement"])
    add_bullet(doc, "Identify and mitigate risks related to software updates and customer-specific configurations; debug production environments and troubleshoot issues.",
        ["mitigate risks", "software updates", "customer-specific configurations", "debug production"])
    add_bullet(doc, "System-level testing across hardware/software integration \u2014 device communication, sensor data validation, real-time feedback loops.",
        ["System-level testing", "hardware/software integration", "sensor data", "real-time feedback"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Automation Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | CI/CD Integration & Release Quality")
    add_bullet(doc, "Integrated automated test suites into CI/CD pipelines (Jenkins) \u2014 build verification, quality gates, fast feedback for development teams.",
        ["automated test suites", "CI/CD pipelines", "build verification", "quality gates"])
    add_bullet(doc, "Analysed and debugged production environments; root-cause analysis for release issues; drove testing practice improvements.",
        ["debugged production environments", "root-cause analysis", "testing practice improvements"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 QA Lead / Technical Specialist",
        "2013 \u2013 2021",
        "Interactive Digital Platforms & Applications \u2014 E-Commerce, Mobile, Device Integration | Team Leadership")
    add_bullet(doc, "Defined and implemented comprehensive QA strategies for complex interactive applications \u2014 multi-platform (web, mobile, devices), fast release cycles.",
        ["QA strategies", "interactive applications", "multi-platform", "fast release cycles"])
    add_bullet(doc, "Led test automation framework development (Python, Java, Selenium) \u2014 scalable architecture supporting simulation, integration, and regression testing.",
        ["test automation framework", "scalable architecture", "simulation", "regression testing"])
    add_bullet(doc, "Mentored and led team of 8\u201312 QA engineers; fostered culture of ownership, technical curiosity, and continuous improvement.",
        ["Mentored and led", "8\u201312 QA engineers", "ownership", "continuous improvement"])
    add_bullet(doc, "System-level testing for hardware/software integrated products \u2014 device communication protocols, sensor validation, real-time data flows.",
        ["System-level testing", "hardware/software", "device communication protocols", "real-time data flows"])
    add_bullet(doc, "CI/CD integration (Jenkins, GitHub Actions, Amazon EC2); automated build verification; quality gates preventing defect leakage to production.",
        ["CI/CD integration", "automated build verification", "quality gates"])
    add_bullet(doc, "Bug triage and root-cause analysis across complex distributed systems; drove continuous improvement of testing tools and standards.",
        ["Bug triage", "root-cause analysis", "continuous improvement"])
    add_bullet(doc, "Collaborated with product owners, developers, and domain experts ensuring quality from concept through post-launch; risk identification and mitigation.",
        ["product owners, developers", "concept through post-launch", "risk identification"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 QA Engineer / Consultant",
        "2008 \u2013 2013",
        "Core Banking & Device Integration \u2014 Regulated Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "System-level testing for Finacle Core Banking \u2014 hardware/software integration (biometric devices), transaction simulation, production debugging.",
        ["System-level testing", "hardware/software integration", "biometric devices", "production debugging"])
    add_bullet(doc, "Built automated test frameworks (Python, Java); CI integration; experience in complex application testing where reliability directly impacts users.",
        ["automated test frameworks", "CI integration", "reliability directly impacts users"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "Six Sigma Green Belt", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
