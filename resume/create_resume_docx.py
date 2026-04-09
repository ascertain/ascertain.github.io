from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Funda_QA_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Funda_QA_Engineer_Resume.doc"

ACCENT = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10.5, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    add_text(paragraph, title.upper(), bold=True, size=11, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.06
    tokens = [
        "15+ years",
        "test automation",
        "Playwright",
        "continuous deployment",
        "continuous delivery",
        "C#",
        ".NET Core",
        "Scrum",
        "Agile",
        "Azure DevOps",
        "Kubernetes",
        "Git",
        "CI/CD",
        "unit testing",
        "integration tests",
        "quality assurance",
        "Selenium",
        "xunit",
        "Jira",
        "whole team",
        "defects",
        "collaboration",
        "mentoring",
    ]
    for token in tokens:
        if token in text:
            before, after = text.split(token, 1)
            add_text(paragraph, before, size=10.5)
            add_text(paragraph, token, bold=True, size=10.5, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10.5)
            return
    add_text(paragraph, text, size=10.5)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=20, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_after = Pt(2)
    add_text(role, "QA Engineer | Test Automation | Continuous Deployment | Agile & Scrum", bold=True, size=11.5)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=10, color=TEXT_MUTED)
    add_text(contact, "\nLinkedIn: linkedin.com/in/md-kashif | GitHub: github.com/ascertain", size=10, color=TEXT_MUTED)
    set_paragraph_bottom_border(contact, color="7F8FA6", size="10")

    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(4)
    profile.paragraph_format.line_spacing = 1.08
    add_text(
        profile,
        "Quality-driven QA Engineer with 15+ years of experience in test automation, quality assurance, and continuous delivery across high-traffic digital products. Strong background in designing, creating, and maintaining automated test suites using Playwright, Selenium, C#, and .NET Core, with deep experience in unit testing (xunit) and integration testing. Proven ability to enable teams to ship faster through continuous deployment practices, whole-team quality ownership, and streamlined CI/CD pipelines. Experienced working in multidisciplinary Scrum teams alongside developers, product owners, and UX designers, promoting a collaborative whole-team approach to testing. Brings practical expertise in Azure DevOps, Kubernetes, Git, GitHub, and Jira, with a developer background that strengthens code-level quality contribution and automation design.",
        size=10.5,
    )

    add_section_heading(document, "Technical & Functional Strengths")
    skill_lines = [
        ("Test Automation & Frameworks: ", "Playwright, Selenium, C#, .NET Core integration tests, xunit, unit testing, API testing, end-to-end automation, test suite design and maintenance"),
        ("Quality Assurance & Testing: ", "defect identification, bug investigation, exploratory testing, regression testing, quality standards, specification validation, test strategy"),
        ("CI/CD & DevOps: ", "continuous deployment, continuous delivery, Azure DevOps, Kubernetes, Git, GitHub, CI/CD pipelines, release automation, deployment confidence"),
        ("Agile & Collaboration: ", "Scrum, Agile testing, whole-team quality approach, developer collaboration, product owner alignment, cross-functional teamwork, Jira"),
        ("Developer Background: ", "C#, .NET Core, JavaScript, TypeScript, Python, code reviews, version control, programming-driven test design, shift-left quality"),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        add_text(paragraph, label, bold=True, size=10.5)
        add_text(paragraph, value, size=10.5)

    add_section_heading(document, "Professional Experience")
    jobs = [
        (
            "IKEA IT AB, Malmö, Sweden",
            "Software Engineer / Senior QA & Test Automation Contributor | Mar 2022 – Present",
            [
                "Drive test automation strategy and execution across complex digital products, designing and maintaining automated test suites that enable continuous deployment with multiple releases per day.",
                "Build and maintain Playwright-based end-to-end tests and .NET Core integration tests, replacing legacy Selenium suites to improve speed, reliability, and developer confidence.",
                "Promote a whole-team approach to quality by embedding testing knowledge across the Scrum team, collaborating closely with developers, product owners, and UX designers.",
                "Strengthen CI/CD pipelines through Azure DevOps, ensuring automated tests run reliably on every commit and supporting continuous delivery with fast feedback loops.",
                "Conduct thorough testing including exploratory, regression, and specification-based testing to identify defects early and prevent issues from reaching production.",
                "Contribute to code quality through code reviews, unit testing with xunit, and practical guidance that helps developers write more testable and reliable code.",
                "Work across Kubernetes-hosted environments, ensuring test infrastructure scales with deployment frequency and supports high-traffic production workloads.",
                "Recognized as an Exceptional Performer through sustained ownership, quality leadership, and dependable collaboration across cross-functional initiatives.",
            ],
        ),
        (
            "Truecaller, Sweden",
            "Release & Automation Engineer | Sep 2021 – Feb 2022",
            [
                "Supported release quality and automation in a fast-moving product environment through structured validation, defect investigation, and stronger readiness checks before rollout.",
                "Designed and maintained automated tests to streamline the testing process and enable faster, more confident releases under continuous change.",
                "Collaborated closely with developers and stakeholders to investigate bugs, improve test coverage, and strengthen confidence in release quality.",
                "Operated effectively in a dynamic Agile environment with rapid deployment cycles requiring practical judgment and disciplined execution.",
            ],
        ),
        (
            "HCLTech / Enterprise Programs / Automation Experience",
            "Automation Lead / SDET / QA Engineer | Dec 2013 – Aug 2021",
            [
                "Led test automation initiatives across enterprise programs, building and maintaining automated test suites using Selenium, C#, and integration testing frameworks.",
                "Worked across web applications, APIs, and integrated platforms to improve end-to-end test coverage, defect detection, and release confidence.",
                "Partnered with development teams to promote shift-left testing, improve unit test adoption, and embed quality practices into CI/CD pipelines.",
                "Mentored team members on automation best practices, test design patterns, and collaborative quality approaches across Scrum teams.",
            ],
        ),
    ]
    for company, title, bullets in jobs:
        company_paragraph = document.add_paragraph()
        company_paragraph.paragraph_format.space_before = Pt(4)
        company_paragraph.paragraph_format.space_after = Pt(0)
        add_text(company_paragraph, company, bold=True, size=11, color=SECTION_COLOR)

        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_after = Pt(0)
        add_text(title_paragraph, title, bold=True, size=10.5)

        for bullet in bullets:
            add_highlighted_bullet(document, bullet)

    add_section_heading(document, "Selected Achievements")
    achievements = [
        "Built 15+ years of experience across test automation, quality assurance, and continuous delivery in high-traffic digital product environments.",
        "Migrated test automation from Selenium to Playwright, improving test reliability, execution speed, and developer adoption across multiple teams.",
        "Enabled continuous deployment practices with 10+ releases per day through robust CI/CD pipelines, fast automated feedback, and reliable test infrastructure.",
        "Strengthened whole-team quality ownership through mentoring, knowledge sharing, and embedding testing practices into Scrum team workflows.",
        "Improved defect detection rates through better test coverage, exploratory testing discipline, and earlier quality feedback in the development cycle.",
        "Built trust across multidisciplinary teams through collaboration, practical quality contribution, and dependable execution in Agile delivery environments.",
    ]
    for achievement in achievements:
        add_highlighted_bullet(document, achievement)

    add_section_heading(document, "Education, Certifications & Languages")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        paragraph = cell.paragraphs[0]
        add_text(paragraph, title, bold=True, size=10.5, color=SECTION_COLOR)

    education = table.cell(1, 0).paragraphs[0]
    add_text(education, "B.Tech in Information Technology\nPG Diploma in Operations", size=10.2)
    certifications = table.cell(1, 1).paragraphs[0]
    add_text(certifications, "ISTQB | ITIL | AWS | GCP | Ethical Hacking | Six Sigma Green Belt", size=10.2)

    add_section_heading(document, "Languages")
    language_paragraph = document.add_paragraph()
    language_paragraph.paragraph_format.space_after = Pt(0)
    add_text(language_paragraph, "English — Fluent | Swedish — Basic", size=10.5)

    document.save(DOCX_PATH)


def build_doc():
    html = """
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif Funda QA Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 32px; color: #1f2937; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 4px; font-size: 24pt; }
    h2 { text-align: center; font-size: 12pt; margin-top: 0; margin-bottom: 8px; }
    .contact { text-align: center; color: #4b5563; border-bottom: 2px solid #1f4e79; padding-bottom: 10px; margin-bottom: 14px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 11pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 3px; margin-top: 14px; margin-bottom: 8px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 8px; }
    .job-sub { font-weight: 700; margin-bottom: 4px; }
    ul { margin-top: 2px; margin-bottom: 6px; }
    li { margin-bottom: 4px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 6px; }
    td { border: 1px solid #b8cce4; padding: 7px; vertical-align: top; }
    .tag { background: #d9eaf7; font-weight: 700; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>QA Engineer | Test Automation | Continuous Deployment | Agile &amp; Scrum</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com<br>LinkedIn: linkedin.com/in/md-kashif | GitHub: github.com/ascertain</div>

  <div class="section">Profile</div>
  <p>Quality-driven QA Engineer with <span class="hl">15+ years</span> of experience in <span class="hl">test automation</span>, <span class="hl">quality assurance</span>, and <span class="hl">continuous delivery</span> across high-traffic digital products. Strong background in designing, creating, and maintaining automated test suites using <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">C#</span>, and <span class="hl">.NET Core</span>, with deep experience in <span class="hl">unit testing</span> (<span class="hl">xunit</span>) and <span class="hl">integration tests</span>. Proven ability to enable teams to ship faster through <span class="hl">continuous deployment</span> practices, <span class="hl">whole team</span> quality ownership, and streamlined <span class="hl">CI/CD</span> pipelines. Experienced working in multidisciplinary <span class="hl">Scrum</span> teams alongside developers, product owners, and UX designers, promoting a collaborative whole-team approach to testing. Brings practical expertise in <span class="hl">Azure DevOps</span>, <span class="hl">Kubernetes</span>, <span class="hl">Git</span>, GitHub, and <span class="hl">Jira</span>, with a developer background that strengthens code-level quality contribution and automation design.</p>

  <div class="section">Technical &amp; Functional Strengths</div>
  <p><b>Test Automation &amp; Frameworks:</b> <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">C#</span>, <span class="hl">.NET Core</span> <span class="hl">integration tests</span>, <span class="hl">xunit</span>, <span class="hl">unit testing</span>, API testing, end-to-end automation, test suite design and maintenance<br>
  <b>Quality Assurance &amp; Testing:</b> <span class="hl">defects</span> identification, bug investigation, exploratory testing, regression testing, quality standards, specification validation, test strategy<br>
  <b>CI/CD &amp; DevOps:</b> <span class="hl">continuous deployment</span>, <span class="hl">continuous delivery</span>, <span class="hl">Azure DevOps</span>, <span class="hl">Kubernetes</span>, <span class="hl">Git</span>, GitHub, <span class="hl">CI/CD</span> pipelines, release automation, deployment confidence<br>
  <b>Agile &amp; Collaboration:</b> <span class="hl">Scrum</span>, <span class="hl">Agile</span> testing, <span class="hl">whole team</span> quality approach, developer <span class="hl">collaboration</span>, product owner alignment, cross-functional teamwork, <span class="hl">Jira</span><br>
  <b>Developer Background:</b> <span class="hl">C#</span>, <span class="hl">.NET Core</span>, JavaScript, TypeScript, Python, code reviews, version control, programming-driven test design, shift-left quality</p>

  <div class="section">Professional Experience</div>
  <div class="job-title">IKEA IT AB, Malmö, Sweden</div>
  <div class="job-sub">Software Engineer / Senior QA &amp; Test Automation Contributor | Mar 2022 – Present</div>
  <ul>
    <li>Drive <span class="hl">test automation</span> strategy and execution across complex digital products, designing and maintaining automated test suites that enable <span class="hl">continuous deployment</span> with multiple releases per day.</li>
    <li>Build and maintain <span class="hl">Playwright</span>-based end-to-end tests and <span class="hl">.NET Core</span> <span class="hl">integration tests</span>, replacing legacy <span class="hl">Selenium</span> suites to improve speed, reliability, and developer confidence.</li>
    <li>Promote a <span class="hl">whole team</span> approach to quality by embedding testing knowledge across the <span class="hl">Scrum</span> team, collaborating closely with developers, product owners, and UX designers.</li>
    <li>Strengthen <span class="hl">CI/CD</span> pipelines through <span class="hl">Azure DevOps</span>, ensuring automated tests run reliably on every commit and supporting <span class="hl">continuous delivery</span> with fast feedback loops.</li>
    <li>Conduct thorough testing including exploratory, regression, and specification-based testing to identify <span class="hl">defects</span> early and prevent issues from reaching production.</li>
    <li>Contribute to code quality through code reviews, <span class="hl">unit testing</span> with <span class="hl">xunit</span>, and practical guidance that helps developers write more testable and reliable code.</li>
    <li>Work across <span class="hl">Kubernetes</span>-hosted environments, ensuring test infrastructure scales with deployment frequency and supports high-traffic production workloads.</li>
    <li>Recognized as an Exceptional Performer through sustained ownership, quality leadership, and dependable <span class="hl">collaboration</span> across cross-functional initiatives.</li>
  </ul>

  <div class="job-title">Truecaller, Sweden</div>
  <div class="job-sub">Release &amp; Automation Engineer | Sep 2021 – Feb 2022</div>
  <ul>
    <li>Supported release quality and automation in a fast-moving product environment through structured validation, defect investigation, and stronger readiness checks before rollout.</li>
    <li>Designed and maintained automated tests to streamline the testing process and enable faster, more confident releases under continuous change.</li>
    <li>Collaborated closely with developers and stakeholders to investigate bugs, improve test coverage, and strengthen confidence in release quality.</li>
    <li>Operated effectively in a dynamic <span class="hl">Agile</span> environment with rapid deployment cycles requiring practical judgment and disciplined execution.</li>
  </ul>

  <div class="job-title">HCLTech / Enterprise Programs / Automation Experience</div>
  <div class="job-sub">Automation Lead / SDET / QA Engineer | Dec 2013 – Aug 2021</div>
  <ul>
    <li>Led <span class="hl">test automation</span> initiatives across enterprise programs, building and maintaining automated test suites using <span class="hl">Selenium</span>, <span class="hl">C#</span>, and integration testing frameworks.</li>
    <li>Worked across web applications, APIs, and integrated platforms to improve end-to-end test coverage, defect detection, and release confidence.</li>
    <li>Partnered with development teams to promote shift-left testing, improve <span class="hl">unit testing</span> adoption, and embed quality practices into <span class="hl">CI/CD</span> pipelines.</li>
    <li>Mentored team members on automation best practices, test design patterns, and collaborative quality approaches across <span class="hl">Scrum</span> teams.</li>
  </ul>

  <div class="section">Selected Achievements</div>
  <ul>
    <li>Built <span class="hl">15+ years</span> of experience across <span class="hl">test automation</span>, <span class="hl">quality assurance</span>, and <span class="hl">continuous delivery</span> in high-traffic digital product environments.</li>
    <li>Migrated <span class="hl">test automation</span> from <span class="hl">Selenium</span> to <span class="hl">Playwright</span>, improving test reliability, execution speed, and developer adoption across multiple teams.</li>
    <li>Enabled <span class="hl">continuous deployment</span> practices with 10+ releases per day through robust <span class="hl">CI/CD</span> pipelines, fast automated feedback, and reliable test infrastructure.</li>
    <li>Strengthened <span class="hl">whole team</span> quality ownership through <span class="hl">mentoring</span>, knowledge sharing, and embedding testing practices into <span class="hl">Scrum</span> team workflows.</li>
    <li>Improved defect detection rates through better test coverage, exploratory testing discipline, and earlier quality feedback in the development cycle.</li>
    <li>Built trust across multidisciplinary teams through <span class="hl">collaboration</span>, practical quality contribution, and dependable execution in <span class="hl">Agile</span> delivery environments.</li>
  </ul>

  <div class="section">Education, Certifications &amp; Languages</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr><td>B.Tech in Information Technology<br>PG Diploma in Operations</td><td>ISTQB | ITIL | AWS | GCP | Ethical Hacking | Six Sigma Green Belt</td></tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
