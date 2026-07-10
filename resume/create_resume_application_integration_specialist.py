"""Generate Application & Integration Specialist resume – IKEA Kuwait."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Application_Integration_Specialist_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Application & Integration Specialist with 12+ years of experience delivering enterprise-grade "
        "technical solutions — specializing in system integrations, API development, data pipelines, "
        "ERP connectivity, and process automation within the IKEA ecosystem. "
        "Hands-on expertise building and maintaining integrations between ERP systems and connected business platforms "
        "(POS, CRM, e-commerce, HR, finance) using APIs (REST/SOAP), middleware, and event-driven architectures. "
        "Strong background in application support, system configuration, performance optimization, "
        "and automation workflows (Power Platform, Azure Logic Apps, scripting). "
        "Proficient with data formats (JSON, XML), SQL databases, and real-time data flow patterns "
        "ensuring seamless, accurate data exchange across enterprise systems. "
        "Excellent communicator who partners closely with business stakeholders to understand needs "
        "and translate requirements into sustainable, user-friendly technical solutions. "
        "Passionate about simplifying ways of working, reducing manual administration, "
        "and enabling business growth through digital solutions. "
        "Deeply aligned with IKEA values — togetherness, diversity, equality, and creating a better everyday life for the many people. "
        "Eager to relocate to Kuwait — strong personal connection to the MENA region (spouse from Morocco, homeowner in Casablanca)."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Technologies")
    skills_data = [
        ("Integration & APIs", "REST & SOAP APIs · Middleware (Azure Logic Apps, integration hubs) · Event-driven integration · Pub/Sub messaging · Real-time data flow · System-to-system connectivity · Integration patterns · Data pipelines"),
        ("ERP & Enterprise Apps", "Microsoft Dynamics (D365, Business Central, AX, NAV familiarity) · ERP integration · System configuration · Upgrades & enhancements · Performance optimization · Application support"),
        ("Automation", "Power Automate · Power Platform · Azure Logic Apps · RPA · Python scripting · Shell scripting · Workflow automation · Process optimization · Reducing manual administration"),
        ("Data & Databases", "SQL (BigQuery, SQL Server) · JSON · XML · Data modelling · ETL/ELT pipelines · Data integrity · Data governance · Real-time data synchronization"),
        ("Cloud & DevOps", "GCP (Cloud Functions, Pub/Sub, BigQuery, GCS, Dataflow) · Azure (Logic Apps, Functions) · AWS · Terraform · Docker · CI/CD · Git · Monitoring & observability"),
        ("Programming", "Python · TypeScript · Node.js · SQL · Shell scripting · Java (earlier career)"),
        ("Stakeholder & Communication", "Business requirement translation · Cross-functional collaboration · Technical documentation · User support · Troubleshooting · Continuous improvement"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS
    add_role(doc, "Integration & Application Lead / Team Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Design, develop, and maintain integrations between enterprise applications and connected business systems — building APIs, middleware, and data pipelines ensuring seamless, accurate, real-time data flow across 20+ platforms (ERP, CRM, POS, HR, finance).",
        "Build and support RESTful APIs and event-driven integrations (Pub/Sub, webhooks) connecting ERP systems with e-commerce, workforce management, and analytics platforms — processing data for 32 IKEA markets.",
        "Identify opportunities to automate business processes and reduce manual administration — developing automation workflows using scripting (Python), cloud functions, and orchestration tools to improve operational efficiency.",
        "Manage system configurations, enhancements, and performance optimization — securing reliable operations, positive user experiences, and continuous improvement of enterprise applications.",
        "Secure data integrity, system security, governance, and compliance — maintaining documentation, supporting audits, monitoring system performance, and ensuring GDPR-compliant data handling.",
        "Partner closely with stakeholders across functions (business, operations, IT, finance) to understand business needs — translating requirements into sustainable, user-friendly technical solutions.",
        "Troubleshoot application, integration, and system issues — providing timely support and driving root cause analysis to prevent recurrence and improve system stability.",
        "Led a major platform integration project — re-architecting from batch-based data exchange to real-time, event-driven integrations (GCP, Python, Pub/Sub, Cloud Functions, BigQuery).",
        "Develop and maintain data pipelines processing JSON/XML data from diverse sources — ensuring accurate transformation, validation, and loading into enterprise databases.",
        "Contribute to continuous improvement initiatives — evaluating new integration patterns, automation tools, and platform enhancements to simplify ways of working.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Platform & Integration Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Designed and maintained system integrations for a globally distributed platform (300M+ users) — building APIs and data pipelines ensuring real-time data flow across distributed services.",
        "Implemented automation workflows reducing manual operations — scripting deployment pipelines and monitoring processes to improve operational efficiency by 50%.",
        "Supported enterprise applications and platform infrastructure — managing configurations, troubleshooting issues, and driving performance optimization.",
        "Collaborated with cross-functional teams to translate business requirements into technical integration solutions — ensuring seamless data exchange between systems.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Technical Lead — Integration & Application Engineering", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led design, development, and maintenance of enterprise integrations connecting ERP, CRM, POS, e-commerce, and finance systems — building APIs (REST/SOAP), middleware, and data pipelines for IKEA and LEGO platforms.",
        "Managed ERP system connectivity and application support — system configurations, upgrades, enhancements, and performance optimization ensuring reliable operations and user satisfaction.",
        "Built and supported APIs and middleware solutions enabling seamless data flow between enterprise applications — handling JSON, XML, and database (SQL) transformations at scale.",
        "Developed automation solutions reducing manual business processes — scripting (Python, Shell), workflow automation, and scheduled jobs improving operational efficiency across teams.",
        "Partnered with business stakeholders across functions to understand needs — translating complex requirements into sustainable integration and automation solutions.",
        "Led a team of 8–12 engineers — providing technical direction, troubleshooting support, and coaching on integration patterns, API design, and automation best practices.",
        "Secured data integrity and system governance — maintaining documentation, supporting compliance requirements, and monitoring system performance across integrated platforms.",
        "Drove continuous improvement of integration architecture — identifying opportunities for consolidation, standardization, and modernization of legacy integrations.",
        "Managed integrations with SaaS platforms (Genesys, Verint, ServiceNow) — configuring connectors, APIs, and data synchronization for enterprise business systems.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software & Integration Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_4 = [
        "Built enterprise integrations and APIs connecting financial, retail, and HR systems — developing RESTful services, data pipelines, and ETL solutions using Java, Python, and SQL.",
        "Supported enterprise applications — managing system configurations, troubleshooting issues, and implementing enhancements for ERP and business platforms.",
        "Developed automation scripts reducing manual data processing — improving accuracy and operational efficiency through scheduled workflows and batch processing.",
        "Worked with diverse data formats (JSON, XML, flat files) and databases (SQL Server, Oracle) — building reliable data transformation and loading pipelines.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── IKEA Experience ───────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Experience & Knowledge")
    ikea_items = [
        "IKEA Customer Connect: Built end-to-end integration layer — APIs, data pipelines, event-driven integrations connecting 20+ business systems across 32 IKEA markets.",
        "Enterprise Application Management: Experience with IKEA's application landscape — ERP connectivity, CRM integration, POS data flows, and workforce management system integration.",
        "Automation & Efficiency: Delivered automation solutions simplifying IKEA ways of working — reducing manual administration and enabling business growth through digital tools.",
        "Data Integrity & Governance: Ensuring compliance, security, and data quality across all integrations — supporting audits, maintaining documentation, and monitoring system health.",
        "IKEA Values: Deep alignment with IKEA culture — togetherness, diversity, equality, simplicity, cost-consciousness, and creating a better everyday life for the many people.",
        "Stakeholder Collaboration: Partnering with business functions (operations, finance, HR, retail) to understand needs and deliver sustainable technical solutions.",
    ]
    for item in ikea_items:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # ─── Motivation / Personal ─────────────────────────────────────────────
    add_heading_block(doc, "Motivation & Relocation")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Highly motivated to relocate to Kuwait and join IKEA Kuwait. Strong personal connection to the MENA region — "
        "my wife is from Morocco, and we own a home in Casablanca. I am genuinely excited about the opportunity to bring "
        "my IKEA experience, integration expertise, and passion for simplifying ways of working to support IKEA Kuwait's "
        "business growth and digital transformation. Ready for immediate relocation."
    )
    r.font.size = Pt(10)

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Application &amp; Integration Specialist with 12+ years of experience delivering enterprise-grade technical solutions — specializing in system integrations, API development, data pipelines, ERP connectivity, and process automation within the IKEA ecosystem. Hands-on expertise building and maintaining integrations between ERP systems and connected business platforms (POS, CRM, e-commerce, HR, finance) using APIs (REST/SOAP), middleware, and event-driven architectures. Strong background in application support, system configuration, performance optimization, and automation workflows (Power Platform, Azure Logic Apps, scripting). Proficient with data formats (JSON, XML), SQL databases, and real-time data flow patterns ensuring seamless, accurate data exchange across enterprise systems. Excellent communicator who partners closely with business stakeholders to understand needs and translate requirements into sustainable, user-friendly technical solutions. Passionate about simplifying ways of working, reducing manual administration, and enabling business growth through digital solutions. Deeply aligned with IKEA values — togetherness, diversity, equality, and creating a better everyday life for the many people. Eager to relocate to Kuwait — strong personal connection to the MENA region (spouse from Morocco, homeowner in Casablanca).</p>

<h2>KEY SKILLS &amp; TECHNOLOGIES</h2>
<table>
<tr><td class="cat">Integration &amp; APIs</td><td>REST &amp; SOAP APIs · Middleware (Azure Logic Apps, integration hubs) · Event-driven integration · Pub/Sub messaging · Real-time data flow · System-to-system connectivity · Integration patterns · Data pipelines</td></tr>
<tr><td class="cat">ERP &amp; Enterprise Apps</td><td>Microsoft Dynamics (D365, Business Central, AX, NAV familiarity) · ERP integration · System configuration · Upgrades &amp; enhancements · Performance optimization · Application support</td></tr>
<tr><td class="cat">Automation</td><td>Power Automate · Power Platform · Azure Logic Apps · RPA · Python scripting · Shell scripting · Workflow automation · Process optimization · Reducing manual administration</td></tr>
<tr><td class="cat">Data &amp; Databases</td><td>SQL (BigQuery, SQL Server) · JSON · XML · Data modelling · ETL/ELT pipelines · Data integrity · Data governance · Real-time data synchronization</td></tr>
<tr><td class="cat">Cloud &amp; DevOps</td><td>GCP (Cloud Functions, Pub/Sub, BigQuery, GCS, Dataflow) · Azure (Logic Apps, Functions) · AWS · Terraform · Docker · CI/CD · Git · Monitoring &amp; observability</td></tr>
<tr><td class="cat">Programming</td><td>Python · TypeScript · Node.js · SQL · Shell scripting · Java (earlier career)</td></tr>
<tr><td class="cat">Stakeholder &amp; Comms</td><td>Business requirement translation · Cross-functional collaboration · Technical documentation · User support · Troubleshooting · Continuous improvement</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Integration &amp; Application Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Design, develop, and maintain integrations between enterprise applications and connected business systems — building APIs, middleware, and data pipelines ensuring seamless, accurate, real-time data flow across 20+ platforms (ERP, CRM, POS, HR, finance).</li>
<li>Build and support RESTful APIs and event-driven integrations (Pub/Sub, webhooks) connecting ERP systems with e-commerce, workforce management, and analytics platforms — processing data for 32 IKEA markets.</li>
<li>Identify opportunities to automate business processes and reduce manual administration — developing automation workflows using scripting (Python), cloud functions, and orchestration tools to improve operational efficiency.</li>
<li>Manage system configurations, enhancements, and performance optimization — securing reliable operations, positive user experiences, and continuous improvement of enterprise applications.</li>
<li>Secure data integrity, system security, governance, and compliance — maintaining documentation, supporting audits, monitoring system performance, and ensuring GDPR-compliant data handling.</li>
<li>Partner closely with stakeholders across functions (business, operations, IT, finance) to understand business needs — translating requirements into sustainable, user-friendly technical solutions.</li>
<li>Troubleshoot application, integration, and system issues — providing timely support and driving root cause analysis to prevent recurrence and improve system stability.</li>
<li>Led a major platform integration project — re-architecting from batch-based data exchange to real-time, event-driven integrations (GCP, Python, Pub/Sub, Cloud Functions, BigQuery).</li>
<li>Develop and maintain data pipelines processing JSON/XML data from diverse sources — ensuring accurate transformation, validation, and loading into enterprise databases.</li>
<li>Contribute to continuous improvement initiatives — evaluating new integration patterns, automation tools, and platform enhancements to simplify ways of working.</li>
</ul>

<p class="role">Platform &amp; Integration Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Designed and maintained system integrations for a globally distributed platform (300M+ users) — building APIs and data pipelines ensuring real-time data flow across distributed services.</li>
<li>Implemented automation workflows reducing manual operations — scripting deployment pipelines and monitoring processes to improve operational efficiency by 50%.</li>
<li>Supported enterprise applications and platform infrastructure — managing configurations, troubleshooting issues, and driving performance optimization.</li>
<li>Collaborated with cross-functional teams to translate business requirements into technical integration solutions — ensuring seamless data exchange between systems.</li>
</ul>

<p class="role">Technical Lead — Integration &amp; Application Engineering <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led design, development, and maintenance of enterprise integrations connecting ERP, CRM, POS, e-commerce, and finance systems — building APIs (REST/SOAP), middleware, and data pipelines for IKEA and LEGO platforms.</li>
<li>Managed ERP system connectivity and application support — system configurations, upgrades, enhancements, and performance optimization ensuring reliable operations and user satisfaction.</li>
<li>Built and supported APIs and middleware solutions enabling seamless data flow between enterprise applications — handling JSON, XML, and database (SQL) transformations at scale.</li>
<li>Developed automation solutions reducing manual business processes — scripting (Python, Shell), workflow automation, and scheduled jobs improving operational efficiency across teams.</li>
<li>Partnered with business stakeholders across functions to understand needs — translating complex requirements into sustainable integration and automation solutions.</li>
<li>Led a team of 8–12 engineers — providing technical direction, troubleshooting support, and coaching on integration patterns, API design, and automation best practices.</li>
<li>Secured data integrity and system governance — maintaining documentation, supporting compliance requirements, and monitoring system performance across integrated platforms.</li>
<li>Drove continuous improvement of integration architecture — identifying opportunities for consolidation, standardization, and modernization of legacy integrations.</li>
<li>Managed integrations with SaaS platforms (Genesys, Verint, ServiceNow) — configuring connectors, APIs, and data synchronization for enterprise business systems.</li>
</ul>

<p class="role">Software &amp; Integration Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built enterprise integrations and APIs connecting financial, retail, and HR systems — developing RESTful services, data pipelines, and ETL solutions using Java, Python, and SQL.</li>
<li>Supported enterprise applications — managing system configurations, troubleshooting issues, and implementing enhancements for ERP and business platforms.</li>
<li>Developed automation scripts reducing manual data processing — improving accuracy and operational efficiency through scheduled workflows and batch processing.</li>
<li>Worked with diverse data formats (JSON, XML, flat files) and databases (SQL Server, Oracle) — building reliable data transformation and loading pipelines.</li>
</ul>

<h2>IKEA EXPERIENCE &amp; KNOWLEDGE</h2>
<ul>
<li><strong>IKEA Customer Connect:</strong> Built end-to-end integration layer — APIs, data pipelines, event-driven integrations connecting 20+ business systems across 32 IKEA markets.</li>
<li><strong>Enterprise Application Management:</strong> Experience with IKEA's application landscape — ERP connectivity, CRM integration, POS data flows, and workforce management system integration.</li>
<li><strong>Automation &amp; Efficiency:</strong> Delivered automation solutions simplifying IKEA ways of working — reducing manual administration and enabling business growth through digital tools.</li>
<li><strong>Data Integrity &amp; Governance:</strong> Ensuring compliance, security, and data quality across all integrations — supporting audits, maintaining documentation, and monitoring system health.</li>
<li><strong>IKEA Values:</strong> Deep alignment with IKEA culture — togetherness, diversity, equality, simplicity, cost-consciousness, and creating a better everyday life for the many people.</li>
<li><strong>Stakeholder Collaboration:</strong> Partnering with business functions (operations, finance, HR, retail) to understand needs and deliver sustainable technical solutions.</li>
</ul>

<h2>MOTIVATION &amp; RELOCATION</h2>
<p>Highly motivated to relocate to Kuwait and join IKEA Kuwait. Strong personal connection to the MENA region — my wife is from Morocco, and we own a home in Casablanca. I am genuinely excited about the opportunity to bring my IKEA experience, integration expertise, and passion for simplifying ways of working to support IKEA Kuwait's business growth and digital transformation. Ready for immediate relocation.</p>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
