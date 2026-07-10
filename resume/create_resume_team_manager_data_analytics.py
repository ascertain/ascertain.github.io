"""Generate Team Manager Data & Analytics resume – IKEA Internal Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Team_Manager_Data_Analytics_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Data & Analytics Leader with 10+ years of experience building analytics products at scale, "
        "designing data pipelines, and leading cross-functional teams within IKEA. "
        "Proven track record in Power BI reporting, ETL pipeline development, semantic layer design, "
        "and enabling self-service data democratization across business domains. "
        "Deep understanding of Data Warehousing, Lakehouse architectures, and large-scale data processing. "
        "Experienced people manager — mentoring data analysts, driving collaboration and knowledge sharing, "
        "and translating complex business needs into actionable data-driven solutions. "
        "Strong IKEA domain knowledge across Customer Support systems with a passion for continuous improvement, "
        "togetherness, and promoting a data-driven culture across the organization."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Data & Analytics", "Power BI (DAX, reports, dashboards, paginated reports) · Data Warehousing · Data Vault · Lakehouse · Data Modeling · Semantic Layers · Data Science fundamentals"),
        ("ETL & Pipelines", "Python · SQL · BigQuery · Airflow · dbt · ETL/ELT pipelines · REST API integrations · Data Lakes · Large-scale dataset processing"),
        ("Self-Service & Democratization", "Scalable semantic layer design · Self-service analytics enablement · Data governance · Data quality frameworks · Business glossaries"),
        ("Cloud Platforms", "GCP (BigQuery, Cloud Functions, Cloud Run, Pub/Sub, Dataflow) · AWS (S3, Glue, Athena, CloudWatch) · Terraform · Docker"),
        ("Leadership & People", "Team management (direct & indirect) · Mentoring Data Analysts · Coaching · Knowledge sharing · Agile/Scrum · Stakeholder management"),
        ("Tools & Practices", "Jira · Confluence · Ingka DevOps Tooling · GitHub Actions · CI/CD · Grafana · Cloud Monitoring · Looker Studio"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS Team Lead Acting
    add_role(doc, "Team Lead Acting — Data & Analytics, Visual Customer Support (VCS)", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2023 – Present")
    bullets_1 = [
        "Lead and develop a team of data analysts and engineers, providing both technical and people leadership — coaching on analytics methods, tools, and best practices.",
        "Designed and built scalable Power BI dashboards and reports for VCS operations, enabling self-service analytics for business stakeholders across 32 markets.",
        "Developed the VCS Data Layer (ETL pipeline) — ingesting data from multiple SaaS and on-prem sources into BigQuery, processing large-scale datasets for cross-organizational analytics. (Ref: https://datacatalog.ingka.com/data/128/)",
        "Built semantic layers and data models enabling data democratization — business users can independently explore and derive insights without engineering dependency.",
        "Act as point of contact and escalation for data analytics within the Customer Support domain — driving consistency in how analytics capabilities are developed and used.",
        "Drove collaboration and knowledge sharing across data analyst community — established regular forums, documentation standards, and reusable templates.",
        "Collaborated with business stakeholders to understand needs and devised data-driven solutions — translating requirements into Power BI reports, automated alerts, and actionable KPI dashboards.",
        "Provided technical leadership and support to data analysts across the organisation — mentoring on DAX, data modeling, Power BI best practices, and SQL optimization.",
        "Managed vendor relationships and consultant coordination — ensured delivery alignment and knowledge transfer to internal team members.",
        "Drove continuous improvement of analytics methods, tools, and ways of working — introduced dbt for transformation layer and established data quality checks.",
        "Active interaction with Business Owner — delivered bi-weekly progress updates and ensured alignment between data delivery and business priorities.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - IKEA CSSP
    add_role(doc, "Data & Analytics Engineer — CSSP (Customer Support Staff Planning)", "IKEA, Ingka Digital", "Malmö, Sweden", "2022 – 2023")
    bullets_2 = [
        "Built analytics products at scale using Power BI — designed dashboards for workforce scheduling, forecasting, and real-time adherence monitoring used by planning teams worldwide.",
        "Designed data warehouse architecture integrating Verint and Genesys platforms — enabling unified workforce analytics across global contact centers.",
        "Developed ETL pipelines (Python, SQL, BigQuery) processing large complex datasets to derive actionable workforce optimization insights.",
        "Created scalable semantic layers enabling self-service reporting for business planners — reducing time-to-insight from days to minutes.",
        "Collaborated with Verint and Genesys specialists to design integration architectures, navigating complex stakeholder environments.",
        "Mentored junior data analysts on Power BI development, DAX patterns, and data modeling best practices.",
        "Helped identify, measure, and automate data points to improve overall analytics process and support data-driven decision-making.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - Truecaller
    add_role(doc, "Data & Release Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_3 = [
        "Built analytics dashboards and reporting solutions for a globally distributed application serving 300M+ users.",
        "Developed data pipelines for release metrics, quality KPIs, and operational insights — enabling data-driven release decisions.",
        "Collaborated cross-functionally with product, engineering, and QA teams in an agile environment.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - HCLTech
    add_role(doc, "Technical Lead — Data & Analytics", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_4 = [
        "Built analytics products at scale using Power BI and Grafana — designed dashboards for IoT device telemetry, customer support metrics, and operational KPIs across IKEA and LEGO.",
        "Led a team of 8–12 engineers and analysts — providing technical leadership, mentoring, conducting knowledge sharing sessions, and driving continuous improvement.",
        "Designed and implemented data warehouse solutions and ETL pipelines processing large-scale IoT and enterprise datasets (Python, SQL, cloud-native services).",
        "Enabled data democratization by building self-service analytics layers — business teams could independently access and analyze data without engineering support.",
        "Navigated complex stakeholder environments across IKEA and LEGO — translating business needs into scalable analytical solutions.",
        "Drove collaboration and consistency in analytics practices across distributed teams spanning multiple countries and domains.",
        "Worked on Startcus Spare Parts, IKEA App, Genesys, and Verint systems — building reporting and analytics integrations.",
        "Established best practices for data quality, documentation, and reusable analytics components.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # Role 5 - Earlier
    add_role(doc, "Software & Data Engineer", "Earlier Career (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_5 = [
        "Built data integration pipelines and automated reporting solutions for financial and retail domains.",
        "Developed analytics dashboards and business intelligence reports for enterprise decision-making.",
        "Worked with large datasets, SQL databases, and ETL processes to support business analytics needs.",
    ]
    for b in bullets_5:
        bullet(doc, b)

    # ─── Key Analytics Projects ────────────────────────────────────────────
    add_heading_block(doc, "Key Data & Analytics Projects")
    projects = [
        "VCS Analytics Platform: End-to-end data layer for Visual Customer Support — ETL pipelines, BigQuery warehouse, Power BI dashboards enabling self-service analytics across 32 markets.",
        "CSSP Workforce Analytics: Power BI solution integrating Verint and Genesys data for real-time scheduling, forecasting, and adherence monitoring across global contact centers.",
        "Customer Support KPI Dashboards: Scalable Power BI reports providing actionable insights for business partners across IKEA Customer Connect domain.",
        "Data Democratization Initiative: Designed semantic layers and self-service frameworks enabling non-technical stakeholders to independently derive insights from complex datasets.",
        "IoT Telemetry Analytics: Large-scale data processing and visualization for IKEA smart-home products — device health, usage patterns, and quality metrics.",
    ]
    for proj in projects:
        bullet(doc, proj, bold_prefix=proj.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Togetherness: Thrive in cross-functional collaboration — actively drive knowledge sharing, community building, and collective problem-solving across data analyst teams.",
        "Simplicity: Advocate for simple, accessible analytics — design self-service solutions that empower business users to make data-driven decisions independently.",
        "Cost-Consciousness: Optimize data infrastructure costs, build reusable analytics components, and maximize value from existing tools and platforms.",
        "Leading by Example: Mentor and develop team members, model best practices, and foster a culture of continuous learning and improvement.",
        "Constant Improvement: Continuously evolve analytics methods, tools, and ways of working — promoting a data-driven culture across the organization.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Data &amp; Analytics Leader with 10+ years of experience building analytics products at scale, designing data pipelines, and leading cross-functional teams within IKEA. Proven track record in Power BI reporting, ETL pipeline development, semantic layer design, and enabling self-service data democratization across business domains. Deep understanding of Data Warehousing, Lakehouse architectures, and large-scale data processing. Experienced people manager — mentoring data analysts, driving collaboration and knowledge sharing, and translating complex business needs into actionable data-driven solutions. Strong IKEA domain knowledge across Customer Support systems with a passion for continuous improvement, togetherness, and promoting a data-driven culture across the organization.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Data &amp; Analytics</td><td>Power BI (DAX, reports, dashboards, paginated reports) · Data Warehousing · Data Vault · Lakehouse · Data Modeling · Semantic Layers · Data Science fundamentals</td></tr>
<tr><td class="cat">ETL &amp; Pipelines</td><td>Python · SQL · BigQuery · Airflow · dbt · ETL/ELT pipelines · REST API integrations · Data Lakes · Large-scale dataset processing</td></tr>
<tr><td class="cat">Self-Service &amp; Democratization</td><td>Scalable semantic layer design · Self-service analytics enablement · Data governance · Data quality frameworks · Business glossaries</td></tr>
<tr><td class="cat">Cloud Platforms</td><td>GCP (BigQuery, Cloud Functions, Cloud Run, Pub/Sub, Dataflow) · AWS (S3, Glue, Athena, CloudWatch) · Terraform · Docker</td></tr>
<tr><td class="cat">Leadership &amp; People</td><td>Team management (direct &amp; indirect) · Mentoring Data Analysts · Coaching · Knowledge sharing · Agile/Scrum · Stakeholder management</td></tr>
<tr><td class="cat">Tools &amp; Practices</td><td>Jira · Confluence · Ingka DevOps Tooling · GitHub Actions · CI/CD · Grafana · Cloud Monitoring · Looker Studio</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Team Lead Acting — Data &amp; Analytics, Visual Customer Support (VCS) <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2023 – Present</span></p>
<ul>
<li>Lead and develop a team of data analysts and engineers, providing both technical and people leadership — coaching on analytics methods, tools, and best practices.</li>
<li>Designed and built scalable Power BI dashboards and reports for VCS operations, enabling self-service analytics for business stakeholders across 32 markets.</li>
<li>Developed the VCS Data Layer (ETL pipeline) — ingesting data from multiple SaaS and on-prem sources into BigQuery, processing large-scale datasets for cross-organizational analytics. (Ref: <a href="https://datacatalog.ingka.com/data/128/">Data Catalog</a>)</li>
<li>Built semantic layers and data models enabling data democratization — business users can independently explore and derive insights without engineering dependency.</li>
<li>Act as point of contact and escalation for data analytics within the Customer Support domain — driving consistency in how analytics capabilities are developed and used.</li>
<li>Drove collaboration and knowledge sharing across data analyst community — established regular forums, documentation standards, and reusable templates.</li>
<li>Collaborated with business stakeholders to understand needs and devised data-driven solutions — translating requirements into Power BI reports, automated alerts, and actionable KPI dashboards.</li>
<li>Provided technical leadership and support to data analysts across the organisation — mentoring on DAX, data modeling, Power BI best practices, and SQL optimization.</li>
<li>Managed vendor relationships and consultant coordination — ensured delivery alignment and knowledge transfer to internal team members.</li>
<li>Drove continuous improvement of analytics methods, tools, and ways of working — introduced dbt for transformation layer and established data quality checks.</li>
<li>Active interaction with Business Owner — delivered bi-weekly progress updates and ensured alignment between data delivery and business priorities.</li>
</ul>

<p class="role">Data &amp; Analytics Engineer — CSSP (Customer Support Staff Planning) <span class="meta">&nbsp;|&nbsp; IKEA, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – 2023</span></p>
<ul>
<li>Built analytics products at scale using Power BI — designed dashboards for workforce scheduling, forecasting, and real-time adherence monitoring used by planning teams worldwide.</li>
<li>Designed data warehouse architecture integrating Verint and Genesys platforms — enabling unified workforce analytics across global contact centers.</li>
<li>Developed ETL pipelines (Python, SQL, BigQuery) processing large complex datasets to derive actionable workforce optimization insights.</li>
<li>Created scalable semantic layers enabling self-service reporting for business planners — reducing time-to-insight from days to minutes.</li>
<li>Collaborated with Verint and Genesys specialists to design integration architectures, navigating complex stakeholder environments.</li>
<li>Mentored junior data analysts on Power BI development, DAX patterns, and data modeling best practices.</li>
<li>Helped identify, measure, and automate data points to improve overall analytics process and support data-driven decision-making.</li>
</ul>

<p class="role">Data &amp; Release Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Built analytics dashboards and reporting solutions for a globally distributed application serving 300M+ users.</li>
<li>Developed data pipelines for release metrics, quality KPIs, and operational insights — enabling data-driven release decisions.</li>
<li>Collaborated cross-functionally with product, engineering, and QA teams in an agile environment.</li>
</ul>

<p class="role">Technical Lead — Data &amp; Analytics <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Built analytics products at scale using Power BI and Grafana — designed dashboards for IoT device telemetry, customer support metrics, and operational KPIs across IKEA and LEGO.</li>
<li>Led a team of 8–12 engineers and analysts — providing technical leadership, mentoring, conducting knowledge sharing sessions, and driving continuous improvement.</li>
<li>Designed and implemented data warehouse solutions and ETL pipelines processing large-scale IoT and enterprise datasets (Python, SQL, cloud-native services).</li>
<li>Enabled data democratization by building self-service analytics layers — business teams could independently access and analyze data without engineering support.</li>
<li>Navigated complex stakeholder environments across IKEA and LEGO — translating business needs into scalable analytical solutions.</li>
<li>Drove collaboration and consistency in analytics practices across distributed teams spanning multiple countries and domains.</li>
<li>Worked on Startcus Spare Parts, IKEA App, Genesys, and Verint systems — building reporting and analytics integrations.</li>
<li>Established best practices for data quality, documentation, and reusable analytics components.</li>
</ul>

<p class="role">Software &amp; Data Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built data integration pipelines and automated reporting solutions for financial and retail domains.</li>
<li>Developed analytics dashboards and business intelligence reports for enterprise decision-making.</li>
<li>Worked with large datasets, SQL databases, and ETL processes to support business analytics needs.</li>
</ul>

<h2>KEY DATA &amp; ANALYTICS PROJECTS</h2>
<ul>
<li><strong>VCS Analytics Platform:</strong> End-to-end data layer for Visual Customer Support — ETL pipelines, BigQuery warehouse, Power BI dashboards enabling self-service analytics across 32 markets.</li>
<li><strong>CSSP Workforce Analytics:</strong> Power BI solution integrating Verint and Genesys data for real-time scheduling, forecasting, and adherence monitoring across global contact centers.</li>
<li><strong>Customer Support KPI Dashboards:</strong> Scalable Power BI reports providing actionable insights for business partners across IKEA Customer Connect domain.</li>
<li><strong>Data Democratization Initiative:</strong> Designed semantic layers and self-service frameworks enabling non-technical stakeholders to independently derive insights from complex datasets.</li>
<li><strong>IoT Telemetry Analytics:</strong> Large-scale data processing and visualization for IKEA smart-home products — device health, usage patterns, and quality metrics.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Togetherness:</strong> Thrive in cross-functional collaboration — actively drive knowledge sharing, community building, and collective problem-solving across data analyst teams.</li>
<li><strong>Simplicity:</strong> Advocate for simple, accessible analytics — design self-service solutions that empower business users to make data-driven decisions independently.</li>
<li><strong>Cost-Consciousness:</strong> Optimize data infrastructure costs, build reusable analytics components, and maximize value from existing tools and platforms.</li>
<li><strong>Leading by Example:</strong> Mentor and develop team members, model best practices, and foster a culture of continuous learning and improvement.</li>
<li><strong>Constant Improvement:</strong> Continuously evolve analytics methods, tools, and ways of working — promoting a data-driven culture across the organization.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
