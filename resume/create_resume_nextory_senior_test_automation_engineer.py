from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Nextory_Senior_Test_Automation_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Nextory_Senior_Test_Automation_Engineer_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "15+ years",
    "Senior Test Automation Engineer",
    "test automation",
    "automated tests",
    "end-to-end",
    "E2E",
    "regression",
    "API",
    "contract testing",
    "shift-left",
    "AI-driven",
    "AI-powered",
    "AI-assisted",
    "AI tool",
    "test generation",
    "exploratory testing",
    "testing frameworks",
    "CI/CD",
    "pipelines",
    "Playwright",
    "Cypress",
    "Maestro",
    "Selenium",
    "Appium",
    "RestAssured",
    "mobile",
    "iOS",
    "Android",
    "coaching",
    "mentor",
    "technical guidance",
    "quality ownership",
    "ownership",
    "quality",
    "scalable",
    "reliable",
    "short cycle times",
    "engineering mindset",
    "collaborative",
    "continuous improvement",
    "fast-moving",
    "Claude",
    "Copilot",
    "Gemini",
    "IKEA",
    "LEGO",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior Test Automation Engineer | Testing Frameworks, CI/CD & AI-Driven QA | 15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior Test Automation Engineer with 15+ years of hands-on experience designing and "
        "implementing systems that prevent bugs — not just find them. Strong engineering mindset "
        "focused on building robust testing frameworks, scalable CI/CD pipelines, and enabling "
        "teams to move fast with confidence. Drives shift-left testing to build quality into the "
        "development process from the start. Actively exploring and implementing AI-driven "
        "approaches for test generation, maintenance, and exploratory testing (Claude, Copilot, "
        "Gemini — daily user). Coaches engineers on test automation and creates environments "
        "where quality is whole-team ownership. Experience across web (Playwright, Cypress, "
        "Selenium) and mobile (Appium, iOS, Android). Thrives in fast-moving, collaborative "
        "environments with short cycle times and continuous improvement.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Testing Frameworks & Automation: ",
            "Designed and built robust testing frameworks from scratch — end-to-end, regression, "
            "API (RestAssured), and contract testing. Playwright, Cypress, Selenium for web; "
            "Appium for mobile (iOS/Android). Coaching others in test automation is core to how I work.",
        ),
        (
            "CI/CD & Short Cycle Times: ",
            "Built seamless CI/CD pipelines (GitHub Actions, Docker, Kubernetes) enabling "
            "efficient, scalable, reliable test processes with short cycle times. Tests run on "
            "every PR — fast feedback, high confidence.",
        ),
        (
            "AI-Driven Testing: ",
            "Daily AI tool user (Claude, Copilot, Gemini). Implementing AI-driven approaches "
            "for test generation, test maintenance, and exploratory testing agents. Rolled out "
            "AI-assisted workflows to teams — 30% velocity improvement.",
        ),
        (
            "Shift-Left & Quality Ownership: ",
            "Drives shift-left testing — quality built into the process from the start. Creates "
            "environments where the whole team takes ownership of quality together. Technical "
            "guidance to developers on testing best practices.",
        ),
        (
            "Mobile Experience: ",
            "Testing on mobile platforms (iOS and Android) using Appium. Familiar with mobile "
            "app testing workflows, device farms, and cross-platform regression.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Test Automation: ",
            "Playwright, Cypress, Selenium, Appium, RestAssured, E2E, regression, API, "
            "contract testing, test framework design",
        ),
        (
            "CI/CD & Infrastructure: ",
            "GitHub Actions, Docker, Kubernetes, Terraform, pipeline design, scalable test processes",
        ),
        (
            "Mobile: ",
            "iOS, Android, Appium, cross-platform testing, mobile app automation",
        ),
        (
            "AI-Driven Testing: ",
            "Claude, Copilot, Gemini — AI test generation, maintenance, exploratory testing "
            "agents, team rollout",
        ),
        (
            "Languages & Frameworks: ",
            "TypeScript, Java, Node.js, React, Spring Boot, Maven, SQL",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Senior Test Automation Engineer / Engineering Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Video Collaboration Solution (SaaS, Full-Stack)", bold=True, size=10)
    ikea_bullets = [
        "Designed and implemented robust testing frameworks — Playwright E2E, API testing "
        "(RestAssured patterns), regression suites, and contract testing. Built systems that "
        "prevent bugs, not just find them. Serving 30+ markets.",
        "Built seamless CI/CD pipelines (GitHub Actions, Docker, Terraform) enabling short "
        "cycle times — automated tests on every PR, fast feedback, scalable and reliable "
        "test processes. Shift-left testing built into the development process from the start.",
        "Implementing AI-driven approaches for testing — Claude, Copilot, Gemini for test "
        "generation, test maintenance, and exploratory testing. Daily AI tool user. Rolled "
        "out AI-assisted workflows to the entire team, improving velocity by 30%.",
        "Coaching engineers on test automation and quality ownership — created an environment "
        "where quality is something the whole team takes ownership of together. Technical "
        "guidance to developers on testing best practices and shift-left thinking.",
        "Mobile testing experience — cross-platform testing, responsive UI automation for "
        "web and mobile views. Kubernetes, Docker, cloud-native infrastructure.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior QA Lead / Test Automation Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  SaaS — 300M+ Users, Mobile-First (iOS & Android)", bold=True, size=10)
    tc_bullets = [
        "Test automation across mobile platforms (iOS and Android) using Appium, plus web "
        "(Selenium) and API (RestAssured). End-to-end, regression, and exploratory testing "
        "at a fast-moving SaaS company with 300M+ users.",
        "Built scalable, reliable test processes with short cycle times. Coached engineers "
        "on test automation. CI/CD pipeline integration. Collaborative, ownership-driven "
        "environment with continuous improvement.",
        "Technical guidance to engineering teams — shift-left testing, quality ownership, "
        "and modern testing practices in a fast-moving tech environment.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Test Lead / Senior Test Automation Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Web & Mobile", bold=True, size=10)
    lego_bullets = [
        "LEGO: Built testing frameworks — Selenium, RestAssured, Appium (iOS/Android). "
        "End-to-end, regression, API, and contract testing. Managed 8–10 engineers. "
        "Coaching on test automation and quality ownership.",
        "IKEA App & Integrations (2018–2021): Test automation across web and mobile "
        "platforms. CI/CD pipelines, shift-left testing. Scalable, reliable test processes "
        "enabling short cycle times across multiple product teams.",
        "Technical guidance to developers — drove quality ownership as a whole-team "
        "responsibility. Collaborative environment with continuous improvement. "
        "Fast-moving, engineering-mindset culture.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior Test Automation Engineer / Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech & Government — Java Full-Stack", bold=True, size=10)
    fin_bullets = [
        "Designed testing frameworks from scratch — Selenium, RestAssured, end-to-end, "
        "regression, and API test automation across Java full-stack banking platforms. "
        "Scalable, reliable test processes with CI/CD integration.",
        "Coached 15+ engineers on test automation. Led teams of 10+. Drove quality "
        "ownership, shift-left practices, and continuous improvement.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Built robust testing frameworks from scratch — Playwright, Selenium, Appium, "
        "RestAssured — enabling scalable, reliable test processes with short cycle times.",
        "AI-driven testing pioneer — Claude, Copilot, Gemini for test generation, "
        "maintenance, and exploratory testing agents. 30% velocity improvement. Team rollout.",
        "Seamless CI/CD pipelines (GitHub Actions, Docker, K8s) — automated tests on every "
        "PR, shift-left testing, fast feedback loops across SaaS products.",
        "Mobile testing across iOS and Android (Appium) — 300M+ user platform (Truecaller) "
        "and enterprise apps (IKEA, LEGO).",
        "Coaching and technical guidance — created whole-team quality ownership culture "
        "across every organization. Natural mentor who raises the bar.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "Certified Ethical Hacker (CEH)\n"
        "AWS Cloud Practitioner\nITIL Foundation\n"
        "Six Sigma Green Belt\nUiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Nextory Senior Test Automation Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior Test Automation Engineer | Testing Frameworks, CI/CD &amp; AI-Driven QA | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Senior Test Automation Engineer</span> with <span class="hl">15+ years</span> designing and implementing systems that prevent bugs. Strong <span class="hl">engineering mindset</span> — builds robust <span class="hl">testing frameworks</span>, <span class="hl">scalable</span> <span class="hl">CI/CD</span> <span class="hl">pipelines</span>, and enables teams to move fast with confidence. Drives <span class="hl">shift-left</span> testing. <span class="hl">AI-driven</span> approaches for <span class="hl">test generation</span>, maintenance, and <span class="hl">exploratory testing</span> agents (<span class="hl">Claude</span>, <span class="hl">Copilot</span>, <span class="hl">Gemini</span> — daily user). <span class="hl">Coaching</span> engineers on <span class="hl">test automation</span>. Web (<span class="hl">Playwright</span>, <span class="hl">Cypress</span>, <span class="hl">Selenium</span>) and <span class="hl">mobile</span> (<span class="hl">Appium</span>, <span class="hl">iOS</span>, <span class="hl">Android</span>). <span class="hl">Fast-moving</span>, <span class="hl">collaborative</span> environments with <span class="hl">short cycle times</span> and <span class="hl">continuous improvement</span>.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Testing Frameworks &amp; Automation:</b> Playwright, Cypress, Selenium, Appium (iOS/Android), RestAssured. E2E, regression, API, contract testing. Coaching others in test automation.<br>
  <b>CI/CD &amp; Short Cycle Times:</b> GitHub Actions, Docker, K8s. Seamless pipelines, scalable reliable test processes, fast feedback.<br>
  <b>AI-Driven Testing:</b> Claude, Copilot, Gemini daily. Test generation, maintenance, exploratory testing agents. 30% velocity improvement. Team rollout.<br>
  <b>Shift-Left &amp; Quality Ownership:</b> Quality built into the process from the start. Whole-team ownership. Technical guidance.<br>
  <b>Mobile:</b> iOS and Android testing (Appium). Mobile app automation, cross-platform regression.</p>

  <div class="section">Technical Skills</div>
  <p><b>Test Automation:</b> Playwright, Cypress, Selenium, Appium, RestAssured, E2E, regression, API, contract testing<br>
  <b>CI/CD:</b> GitHub Actions, Docker, Kubernetes, Terraform, pipeline design<br>
  <b>Mobile:</b> iOS, Android, Appium, cross-platform testing<br>
  <b>AI-Driven Testing:</b> Claude, Copilot, Gemini — test generation, maintenance, exploratory agents<br>
  <b>Languages:</b> TypeScript, Java, Node.js, React, Spring Boot, Maven, SQL</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Senior Test Automation Engineer / Engineering Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — SaaS, Full-Stack</div>
  <ul>
    <li>Designed robust <span class="hl">testing frameworks</span> — <span class="hl">Playwright</span> <span class="hl">E2E</span>, <span class="hl">API</span> testing, <span class="hl">regression</span>, <span class="hl">contract testing</span>. Systems that prevent bugs. 30+ markets.</li>
    <li>Built seamless <span class="hl">CI/CD</span> <span class="hl">pipelines</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>) — <span class="hl">short cycle times</span>, <span class="hl">scalable</span> <span class="hl">reliable</span> test processes. <span class="hl">Shift-left</span> testing.</li>
    <li><span class="hl">AI-driven</span> testing — <span class="hl">Claude</span>, <span class="hl">Copilot</span>, <span class="hl">Gemini</span> for <span class="hl">test generation</span>, maintenance, <span class="hl">exploratory testing</span>. 30% velocity. Team rollout.</li>
    <li><span class="hl">Coaching</span> engineers on <span class="hl">test automation</span> and <span class="hl">quality ownership</span>. <span class="hl">Technical guidance</span>. Whole-team quality.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior QA Lead / Test Automation Engineer</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | SaaS — 300M+ Users, Mobile-First (iOS &amp; Android)</div>
  <ul>
    <li><span class="hl">Test automation</span> across <span class="hl">mobile</span> (<span class="hl">iOS</span>, <span class="hl">Android</span> — <span class="hl">Appium</span>), web (<span class="hl">Selenium</span>), <span class="hl">API</span> (<span class="hl">RestAssured</span>). <span class="hl">E2E</span>, <span class="hl">regression</span>, <span class="hl">exploratory testing</span>.</li>
    <li><span class="hl">Scalable</span> <span class="hl">reliable</span> test processes, <span class="hl">short cycle times</span>. <span class="hl">Coaching</span> engineers. <span class="hl">CI/CD</span>. <span class="hl">Fast-moving</span>, <span class="hl">collaborative</span>, <span class="hl">continuous improvement</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Test Lead / Senior Test Automation Engineer</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Web &amp; Mobile</div>
  <ul>
    <li>LEGO: <span class="hl">Testing frameworks</span> — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Appium</span> (<span class="hl">iOS</span>/<span class="hl">Android</span>). <span class="hl">E2E</span>, <span class="hl">regression</span>, <span class="hl">API</span>, <span class="hl">contract testing</span>. 8–10 engineers.</li>
    <li>IKEA (2018–21): <span class="hl">Test automation</span> across web &amp; <span class="hl">mobile</span>. <span class="hl">CI/CD</span>, <span class="hl">shift-left</span>, <span class="hl">short cycle times</span>. <span class="hl">Coaching</span> developers.</li>
    <li><span class="hl">Quality ownership</span> as whole-team responsibility. <span class="hl">Collaborative</span>, <span class="hl">continuous improvement</span>, <span class="hl">engineering mindset</span>.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior Test Automation Engineer / Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech &amp; Government — Java Full-Stack</div>
  <ul>
    <li><span class="hl">Testing frameworks</span> from scratch — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">E2E</span>, <span class="hl">regression</span>, <span class="hl">API</span>. <span class="hl">Scalable</span> <span class="hl">reliable</span> test processes. <span class="hl">CI/CD</span>.</li>
    <li><span class="hl">Coaching</span> 15+ engineers. <span class="hl">Quality ownership</span>. <span class="hl">Shift-left</span>. <span class="hl">Continuous improvement</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Robust <span class="hl">testing frameworks</span> — <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Appium</span>, <span class="hl">RestAssured</span>. <span class="hl">Scalable</span> <span class="hl">reliable</span> processes, <span class="hl">short cycle times</span>.</li>
    <li><span class="hl">AI-driven</span> testing pioneer — <span class="hl">Claude</span>, <span class="hl">Copilot</span>, <span class="hl">Gemini</span>. <span class="hl">Test generation</span>, <span class="hl">exploratory testing</span> agents. 30% velocity. Team rollout.</li>
    <li>Seamless <span class="hl">CI/CD</span> <span class="hl">pipelines</span> — <span class="hl">shift-left</span>, <span class="hl">automated tests</span> on every PR across SaaS products.</li>
    <li><span class="hl">Mobile</span> testing (<span class="hl">iOS</span>/<span class="hl">Android</span>) — 300M+ user platform (Truecaller) and enterprise apps.</li>
    <li><span class="hl">Coaching</span> and <span class="hl">technical guidance</span> — whole-team <span class="hl">quality ownership</span> culture across every organization.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>Certified Ethical Hacker (CEH)<br>AWS Cloud Practitioner<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
