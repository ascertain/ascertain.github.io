"""Generate Engineering & Quality Leader resume – Inter IKEA Purchasing."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Engineering_Quality_Leader_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Engineering & Quality Leader with 12+ years of experience driving engineering excellence, quality assurance, "
        "process control, and continuous improvement in global, matrix organizations within the IKEA ecosystem. "
        "Proven leadership skills in influencing and making things happen across cross-functional teams, suppliers, "
        "and stakeholders — leading business agendas at suppliers and creating preconditions for quality ownership. "
        "Strong engineering background (B.Tech IT) combined with deep knowledge of quality management, process improvement (Six Sigma), "
        "cost development, and production optimization. "
        "Analytical and holistic thinker who leads with the customer, supplier, and total IKEA in mind — "
        "translating strategy into operational/tactical execution while delivering results through involving, leading, "
        "developing, and inspiring people in remote environments. "
        "Experienced in product development support, compliance assurance, customer feedback analysis, and data-driven "
        "decision-making. Skilled at building and maintaining relationships with stakeholders across Range Areas, "
        "Supply Areas, and the broader IKEA organization. "
        "Committed to IKEA values, Democratic Design, and Fact-based Business Leadership — acting as coach, mentor, "
        "and ambassador for engineering and quality priorities."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Engineering & Quality", "Quality assurance & control · Process control · Continuous improvement · Compliance management · Engineering development · Production improvement · Automation preconditions · Root cause analysis"),
        ("Leadership", "Matrix organization leadership · Remote team leadership · Coaching & mentoring · Influencing without authority · Supplier business agenda · Cross-functional collaboration · People development"),
        ("Process & Strategy", "Category strategy execution · Cost development · Business process optimization · Six Sigma (Green Belt) · PDCA · Data-driven decision making · KPI management · Risk assessment"),
        ("Product Development", "News/Moves/Improves support · Democratic Design · Engineering dimension · Customer satisfaction · Quality standards · Material transformation · Supplier development"),
        ("IKEA Knowledge", "IKEA Strategic Landscape · Purchasing & Quality Strategies · IKEA Business Model · One IKEA Business Plan · PDP · DPOP · Supplier Development Process · IKEA Purchasing Manual"),
        ("Technical & Analytical", "Data analysis (Python, SQL, BigQuery) · Quality metrics & reporting · Process automation · Cloud platforms (GCP) · Systems integration · Performance monitoring"),
        ("Certifications", "Six Sigma Green Belt · ISTQB Certified · Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS
    add_role(doc, "Engineering & Quality Lead / Team Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Lead engineering and quality assurance across the product area — securing performance, compliance, process control, and continuous improvement in alignment with IKEA's quality and engineering strategy.",
        "Drive quality development work from product development through delivery — creating preconditions for automation, establishing quality gates, and ensuring customer expectations are met across 32 markets.",
        "Lead and support engineering teams in quality improvement projects — analyzing customer feedback, developing action plans, and driving measurable improvements in customer satisfaction and product reliability.",
        "Secure engineering and quality priorities are known and integrated into category/product plans — co-creating strategy with management and ensuring execution aligns with business goals.",
        "Support competence development for engineering and quality within business teams and at supplier partners — planning development needs in alignment with business development managers.",
        "Lead and support implementation of quality assurance practices across the value chain — creating preconditions for teams and suppliers to take full responsibility for quality performance.",
        "Build and maintain relationships with stakeholders across Range Areas, Supply Areas, and adjacent IKEA organizations — securing alignment, implementation of actions, and achievement of goals.",
        "Drive data-driven quality management — building analytics and reporting capabilities (Python, SQL, BigQuery) that enable fact-based business leadership and proactive quality interventions.",
        "Contribute to global engineering and quality priorities — co-creating ways of working across the engineering matrix and identifying cross-category sharing and synergy opportunities.",
        "Act as ambassador for IKEA values and role model for Fact-based Business Leadership — leading with the customer, supplier, and total IKEA in mind.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Quality & Platform Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Drove engineering quality and process improvement for a global platform (300M+ users) — establishing quality assurance practices, performance monitoring, and continuous improvement cycles.",
        "Implemented automated quality checks and process control mechanisms — reducing defect rates and improving production reliability through data-driven analysis.",
        "Collaborated cross-functionally with product and engineering teams — ensuring quality was embedded from the start of development, not as an afterthought.",
        "Contributed to cost development and production improvement — optimizing deployment pipelines and reducing operational overhead by 50%.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Engineering & Quality Lead / Technical Lead", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led engineering and quality agenda at supplier level — driving quality development, compliance assurance, and continuous improvement across IKEA and LEGO product teams in a matrix organization.",
        "Proven experience leading business agenda at suppliers — creating preconditions for suppliers to take full responsibility for quality performance through process establishment, competence development, and governance.",
        "Led a team of 8–12 engineers — developing and inspiring people in remote environments, acting as coach and mentor, and delivering results through involvement and empowerment.",
        "Drove process control and continuous improvement initiatives — implementing Six Sigma methodologies, root cause analysis, and structured problem-solving to reduce defects and improve production quality.",
        "Supported product development projects (equivalent to News/Moves/Improves) in the engineering and quality dimension — ensuring quality requirements were addressed from design through delivery.",
        "Led customer feedback analysis and developed improvement projects — translating customer satisfaction data into actionable engineering and quality improvement plans.",
        "Built and maintained relationships with stakeholders across multiple IKEA domains — securing alignment, coordinating cross-functional initiatives, and driving shared quality goals.",
        "Supported material and technology transformation — deploying solutions aligned with strategic direction and securing that engineering priorities were reflected in action plans.",
        "Contributed to global engineering and quality priorities — co-creating ways of working, identifying synergy opportunities, and sharing best practices across categories and teams.",
        "Communicated clearly in English (verbal and written) — translating complex engineering/quality topics into simple, actionable language for diverse stakeholder audiences.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software & Quality Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_4 = [
        "Built engineering and quality foundations across enterprise software projects — developing quality assurance processes, compliance checks, and production improvement practices.",
        "Led quality assurance initiatives in global implementation projects — ensuring product quality through systematic testing, process control, and continuous improvement.",
        "Contributed to cost development through process automation and production optimization — reducing rework and improving delivery efficiency.",
        "Worked in cross-functional teams delivering to business requirements — developing analytical skills and systematic approaches to end-to-end quality.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── IKEA Ecosystem Experience ─────────────────────────────────────────
    add_heading_block(doc, "IKEA Ecosystem Experience")
    ikea_items = [
        "IKEA Customer Connect (VCS): Engineering & quality leadership for digital product area — quality assurance, process control, continuous improvement, and customer satisfaction across 32 markets.",
        "Supplier Collaboration: Extensive experience working with IKEA's supplier ecosystem — leading quality agendas, competence development, and creating preconditions for supplier quality ownership.",
        "IKEA Ways of Working: Product team model, matrix organization leadership, cross-functional collaboration, Agile/DevOps, fact-based business leadership.",
        "Quality Strategy: Driving IKEA's quality agenda — connecting engineering and quality priorities to business strategy, customer satisfaction, and Democratic Design principles.",
        "Cross-Category Collaboration: Contributing to global engineering and quality initiatives — identifying synergies, sharing best practices, and co-creating ways of working across the engineering matrix.",
        "Stakeholder Management: Building relationships across Range Areas, Supply Areas, and retail organization — securing alignment and coordinated execution of engineering and quality goals.",
        "Data-Driven Quality: Leveraging analytics and reporting (Python, SQL, BigQuery) for fact-based quality management — proactive identification of trends, risks, and improvement opportunities.",
    ]
    for item in ikea_items:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Customer Focus: Lead with the customer in mind — understanding how everyday engineering and quality work adds customer value by delivering to Democratic Design principles.",
        "Togetherness: Build and maintain strong relationships across Range Areas, Supply Areas, suppliers, and the broader IKEA organization — achieving results through collaboration and shared goals.",
        "Leading by Example: Act as ambassador for IKEA values and role model for Fact-based Business Leadership — coaching, mentoring, and inspiring people to perform and develop.",
        "Cost-Consciousness: Drive cost development and production improvement — balancing quality, cost, and customer satisfaction in every decision.",
        "Simplicity: Communicate clearly and in a simple way — making engineering and quality priorities accessible and actionable for all stakeholders.",
        "Continuous Improvement: Never stop improving — applying analytical thinking, Six Sigma, and systematic approaches to drive better quality at lower cost.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Engineering &amp; Quality Leader with 12+ years of experience driving engineering excellence, quality assurance, process control, and continuous improvement in global, matrix organizations within the IKEA ecosystem. Proven leadership skills in influencing and making things happen across cross-functional teams, suppliers, and stakeholders — leading business agendas at suppliers and creating preconditions for quality ownership. Strong engineering background (B.Tech IT) combined with deep knowledge of quality management, process improvement (Six Sigma), cost development, and production optimization. Analytical and holistic thinker who leads with the customer, supplier, and total IKEA in mind — translating strategy into operational/tactical execution while delivering results through involving, leading, developing, and inspiring people in remote environments. Experienced in product development support, compliance assurance, customer feedback analysis, and data-driven decision-making. Skilled at building and maintaining relationships with stakeholders across Range Areas, Supply Areas, and the broader IKEA organization. Committed to IKEA values, Democratic Design, and Fact-based Business Leadership — acting as coach, mentor, and ambassador for engineering and quality priorities.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Engineering &amp; Quality</td><td>Quality assurance &amp; control · Process control · Continuous improvement · Compliance management · Engineering development · Production improvement · Automation preconditions · Root cause analysis</td></tr>
<tr><td class="cat">Leadership</td><td>Matrix organization leadership · Remote team leadership · Coaching &amp; mentoring · Influencing without authority · Supplier business agenda · Cross-functional collaboration · People development</td></tr>
<tr><td class="cat">Process &amp; Strategy</td><td>Category strategy execution · Cost development · Business process optimization · Six Sigma (Green Belt) · PDCA · Data-driven decision making · KPI management · Risk assessment</td></tr>
<tr><td class="cat">Product Development</td><td>News/Moves/Improves support · Democratic Design · Engineering dimension · Customer satisfaction · Quality standards · Material transformation · Supplier development</td></tr>
<tr><td class="cat">IKEA Knowledge</td><td>IKEA Strategic Landscape · Purchasing &amp; Quality Strategies · IKEA Business Model · One IKEA Business Plan · PDP · DPOP · Supplier Development Process · IKEA Purchasing Manual</td></tr>
<tr><td class="cat">Technical &amp; Analytical</td><td>Data analysis (Python, SQL, BigQuery) · Quality metrics &amp; reporting · Process automation · Cloud platforms (GCP) · Systems integration · Performance monitoring</td></tr>
<tr><td class="cat">Certifications</td><td>Six Sigma Green Belt · ISTQB Certified · Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Engineering &amp; Quality Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Lead engineering and quality assurance across the product area — securing performance, compliance, process control, and continuous improvement in alignment with IKEA's quality and engineering strategy.</li>
<li>Drive quality development work from product development through delivery — creating preconditions for automation, establishing quality gates, and ensuring customer expectations are met across 32 markets.</li>
<li>Lead and support engineering teams in quality improvement projects — analyzing customer feedback, developing action plans, and driving measurable improvements in customer satisfaction and product reliability.</li>
<li>Secure engineering and quality priorities are known and integrated into category/product plans — co-creating strategy with management and ensuring execution aligns with business goals.</li>
<li>Support competence development for engineering and quality within business teams and at supplier partners — planning development needs in alignment with business development managers.</li>
<li>Lead and support implementation of quality assurance practices across the value chain — creating preconditions for teams and suppliers to take full responsibility for quality performance.</li>
<li>Build and maintain relationships with stakeholders across Range Areas, Supply Areas, and adjacent IKEA organizations — securing alignment, implementation of actions, and achievement of goals.</li>
<li>Drive data-driven quality management — building analytics and reporting capabilities (Python, SQL, BigQuery) that enable fact-based business leadership and proactive quality interventions.</li>
<li>Contribute to global engineering and quality priorities — co-creating ways of working across the engineering matrix and identifying cross-category sharing and synergy opportunities.</li>
<li>Act as ambassador for IKEA values and role model for Fact-based Business Leadership — leading with the customer, supplier, and total IKEA in mind.</li>
</ul>

<p class="role">Quality &amp; Platform Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Drove engineering quality and process improvement for a global platform (300M+ users) — establishing quality assurance practices, performance monitoring, and continuous improvement cycles.</li>
<li>Implemented automated quality checks and process control mechanisms — reducing defect rates and improving production reliability through data-driven analysis.</li>
<li>Collaborated cross-functionally with product and engineering teams — ensuring quality was embedded from the start of development, not as an afterthought.</li>
<li>Contributed to cost development and production improvement — optimizing deployment pipelines and reducing operational overhead by 50%.</li>
</ul>

<p class="role">Engineering &amp; Quality Lead / Technical Lead <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led engineering and quality agenda at supplier level — driving quality development, compliance assurance, and continuous improvement across IKEA and LEGO product teams in a matrix organization.</li>
<li>Proven experience leading business agenda at suppliers — creating preconditions for suppliers to take full responsibility for quality performance through process establishment, competence development, and governance.</li>
<li>Led a team of 8–12 engineers — developing and inspiring people in remote environments, acting as coach and mentor, and delivering results through involvement and empowerment.</li>
<li>Drove process control and continuous improvement initiatives — implementing Six Sigma methodologies, root cause analysis, and structured problem-solving to reduce defects and improve production quality.</li>
<li>Supported product development projects (equivalent to News/Moves/Improves) in the engineering and quality dimension — ensuring quality requirements were addressed from design through delivery.</li>
<li>Led customer feedback analysis and developed improvement projects — translating customer satisfaction data into actionable engineering and quality improvement plans.</li>
<li>Built and maintained relationships with stakeholders across multiple IKEA domains — securing alignment, coordinating cross-functional initiatives, and driving shared quality goals.</li>
<li>Supported material and technology transformation — deploying solutions aligned with strategic direction and securing that engineering priorities were reflected in action plans.</li>
<li>Contributed to global engineering and quality priorities — co-creating ways of working, identifying synergy opportunities, and sharing best practices across categories and teams.</li>
<li>Communicated clearly in English (verbal and written) — translating complex engineering/quality topics into simple, actionable language for diverse stakeholder audiences.</li>
</ul>

<p class="role">Software &amp; Quality Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built engineering and quality foundations across enterprise software projects — developing quality assurance processes, compliance checks, and production improvement practices.</li>
<li>Led quality assurance initiatives in global implementation projects — ensuring product quality through systematic testing, process control, and continuous improvement.</li>
<li>Contributed to cost development through process automation and production optimization — reducing rework and improving delivery efficiency.</li>
<li>Worked in cross-functional teams delivering to business requirements — developing analytical skills and systematic approaches to end-to-end quality.</li>
</ul>

<h2>IKEA ECOSYSTEM EXPERIENCE</h2>
<ul>
<li><strong>IKEA Customer Connect (VCS):</strong> Engineering &amp; quality leadership for digital product area — quality assurance, process control, continuous improvement, and customer satisfaction across 32 markets.</li>
<li><strong>Supplier Collaboration:</strong> Extensive experience working with IKEA's supplier ecosystem — leading quality agendas, competence development, and creating preconditions for supplier quality ownership.</li>
<li><strong>IKEA Ways of Working:</strong> Product team model, matrix organization leadership, cross-functional collaboration, Agile/DevOps, fact-based business leadership.</li>
<li><strong>Quality Strategy:</strong> Driving IKEA's quality agenda — connecting engineering and quality priorities to business strategy, customer satisfaction, and Democratic Design principles.</li>
<li><strong>Cross-Category Collaboration:</strong> Contributing to global engineering and quality initiatives — identifying synergies, sharing best practices, and co-creating ways of working across the engineering matrix.</li>
<li><strong>Stakeholder Management:</strong> Building relationships across Range Areas, Supply Areas, and retail organization — securing alignment and coordinated execution of engineering and quality goals.</li>
<li><strong>Data-Driven Quality:</strong> Leveraging analytics and reporting (Python, SQL, BigQuery) for fact-based quality management — proactive identification of trends, risks, and improvement opportunities.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Customer Focus:</strong> Lead with the customer in mind — understanding how everyday engineering and quality work adds customer value by delivering to Democratic Design principles.</li>
<li><strong>Togetherness:</strong> Build and maintain strong relationships across Range Areas, Supply Areas, suppliers, and the broader IKEA organization — achieving results through collaboration and shared goals.</li>
<li><strong>Leading by Example:</strong> Act as ambassador for IKEA values and role model for Fact-based Business Leadership — coaching, mentoring, and inspiring people to perform and develop.</li>
<li><strong>Cost-Consciousness:</strong> Drive cost development and production improvement — balancing quality, cost, and customer satisfaction in every decision.</li>
<li><strong>Simplicity:</strong> Communicate clearly and in a simple way — making engineering and quality priorities accessible and actionable for all stakeholders.</li>
<li><strong>Continuous Improvement:</strong> Never stop improving — applying analytical thinking, Six Sigma, and systematic approaches to drive better quality at lower cost.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
