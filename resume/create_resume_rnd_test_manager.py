"""
Generate a tailored resume for R&D Testing Technology Manager.
Focus: test methods development, team leadership, test planning, technical competence,
high-speed separators/manufacturing adjacent, structured & systematic approach.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_RnD_Test_Manager_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_RnD_Test_Manager_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "test methods", "test routines", "test planning",
    "team competence", "people development", "leadership",
    "product development", "R&D", "innovation",
    "structured", "systematic", "continuous improvement",
    "manufacturing", "testing technology",
    "cross-functional", "stakeholder",
    "quality", "performance testing",
    "automation", "CI/CD",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "R&D Testing Technology Manager", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Engineering manager with 5+ years of people leadership experience and a strong technical "
        "background in testing and product development. Proven ability to develop test methods and "
        "routines, build team competence, and drive efficient test planning within complex R&D "
        "environments. Constructive, structured, and systematic in approach — with a track record "
        "of leading cross-functional collaboration to deliver high-quality products. Passionate "
        "about leading and developing individuals while maintaining hands-on technical interest in "
        "testing technology. Proactive leader who inspires teams, manages stakeholder relationships "
        "effectively, and drives continuous improvement in testing processes. Fluent in Swedish "
        "and English."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Test Methods & Routines Development", "People Leadership & Team Development", "Test Planning & Project Coordination",
        "Structured & Systematic Approach", "Cross-Functional Collaboration", "Continuous Improvement & Innovation",
        "Technical Competence in Testing", "Stakeholder & Relationship Management", "Efficient Product Development",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test & Engineering Manager — Platform Development", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Manage and lead a testing technology team responsible for test methods, test routines, and quality assurance across a customer-facing platform deployed in 30+ markets.",
        "Develop and grow team competence through structured coaching, individual development plans, regular 1-on-1s, and skills gap analysis — building a high-performing testing organization.",
        "Define and evolve test methods and routines: system testing, integration testing, performance testing, regression automation, and exploratory testing approaches.",
        "Drive test planning in collaboration with R&D and product teams — ensuring testing activities are aligned with project milestones and product development cadence.",
        "Lead cross-functional collaboration between development, infrastructure, and product teams — coordinating technical decisions and ensuring testing supports efficient delivery.",
        "Maintain a constructive, structured, and systematic approach to test strategy evolution — introducing risk-based testing, improved coverage models, and data-driven quality decisions.",
        "Inspire and motivate team members through positive leadership — fostering an innovative environment where individuals take ownership and grow professionally.",
        "Manage stakeholder relationships at both operational and strategic levels — communicating test progress, quality metrics, and improvement initiatives to leadership.",
        "Drive continuous improvement in testing technology: introduced automated test pipelines (CI/CD), performance benchmarking tools, and real-time quality dashboards.",
        "Recognized as Exceptional Performer for 2023 — acknowledged for leadership impact and testing innovation.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test & Release Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Developed and maintained test routines for release readiness across a 300M+ user platform — ensuring systematic quality verification before each production deployment.",
        "Contributed to test planning and coordination across development squads — aligning testing timelines with product development sprints and release milestones.",
        "Drove improvements in test methods: introduced structured regression approaches, performance validation gates, and automated quality checks in CI/CD pipelines.",
        "Collaborated cross-functionally with engineering and operations teams — managing relationships and ensuring testing supported efficient product delivery.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Lead / Technical Specialist — Product Testing", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led testing teams (5–12 members) across IKEA and LEGO projects — responsible for team competence development, workload distribution, and performance management.",
        "Developed test methods and routines for enterprise product platforms: system testing, integration testing, performance/load testing, and manufacturing-adjacent hardware-software integration validation.",
        "Drove test planning for complex product development projects — collaborating with R&D, architects, and product owners to align test activities with delivery milestones.",
        "Built a structured and systematic testing practice: defined test strategies, entry/exit criteria, coverage models, and reusable test frameworks adopted across multiple projects.",
        "Led cross-functional coordination between development teams, product management, and customer representatives across Denmark, Sweden, and India.",
        "Mentored and developed team members — conducting technical training, introducing new testing technologies, and supporting career growth within the testing discipline.",
        "Managed testing for product releases including hardware-software integration scenarios — validating system behavior under real-world operational conditions.",
        "Drove continuous improvement in testing processes: reduced regression cycles by 40% through automation, improved defect escape rates, and introduced quality metrics dashboards.",
        "Maintained strong stakeholder relationships — communicating test results and quality status constructively to programme leadership and customer representatives.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise IT & Banking", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Test Lead / Technical Lead", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Led testing teams for core banking system implementations — developing test methods, managing team workload, and ensuring systematic quality assurance across multi-phase deployments.",
        "Defined test routines for integration, data migration, and system validation — establishing structured approaches adopted across 20+ branch rollouts.",
        "Collaborated with cross-functional teams (development, operations, business) to align test planning with product development timelines and customer expectations.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Environment ──────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL ENVIRONMENT")
    for label, value in [
        ("Testing Technology: ", "Test automation frameworks, performance/load testing, system integration testing, hardware-software validation, CI/CD quality gates"),
        ("Tools: ", "Jira, Azure DevOps, Confluence, Playwright, Selenium, Pytest, JMeter, Grafana, Datadog"),
        ("Engineering: ", "Python, TypeScript, Node.js, Docker, Terraform, GCP, AWS, GitHub Actions, Jenkins"),
        ("Methodologies: ", "Agile (Scrum/Kanban), V-Model, ISTQB, risk-based testing, exploratory testing, TDD"),
        ("Leadership: ", "Team development, 1-on-1 coaching, competence planning, performance management, stakeholder communication"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("M.Tech, Computer Science", "JNTU, India"),
        ("B.Tech, Information Technology", "JNTU, India"),
        ("PG Diploma, Operations Management", "IGNOU, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Six Sigma Green Belt",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "ITIL v4 Foundation",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "Swedish (Fluent)  •  English (Fluent)  •  Danish (Conversational)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">R&amp;D Testing Technology Manager</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Engineering manager with 5+ years of people leadership experience and a strong technical background in testing and product development. Proven ability to develop test methods and routines, build team competence, and drive efficient test planning within complex R&amp;D environments. Constructive, structured, and systematic in approach &mdash; with a track record of leading cross-functional collaboration to deliver high-quality products. Passionate about leading and developing individuals while maintaining hands-on technical interest in testing technology. Proactive leader who inspires teams, manages stakeholder relationships effectively, and drives continuous improvement in testing processes. Fluent in Swedish and English.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Test Methods &amp; Routines Development</td><td>&bull; People Leadership &amp; Team Development</td><td>&bull; Test Planning &amp; Project Coordination</td></tr>
<tr><td>&bull; Structured &amp; Systematic Approach</td><td>&bull; Cross-Functional Collaboration</td><td>&bull; Continuous Improvement &amp; Innovation</td></tr>
<tr><td>&bull; Technical Competence in Testing</td><td>&bull; Stakeholder &amp; Relationship Management</td><td>&bull; Efficient Product Development</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Test &amp; Engineering Manager &mdash; Platform Development <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Manage and lead a testing technology team responsible for test methods, test routines, and quality assurance across a customer-facing platform deployed in 30+ markets.</li>
<li>Develop and grow team competence through structured coaching, individual development plans, regular 1-on-1s, and skills gap analysis &mdash; building a high-performing testing organization.</li>
<li>Define and evolve test methods and routines: system testing, integration testing, performance testing, regression automation, and exploratory testing approaches.</li>
<li>Drive test planning in collaboration with R&amp;D and product teams &mdash; ensuring testing activities are aligned with project milestones and product development cadence.</li>
<li>Lead cross-functional collaboration between development, infrastructure, and product teams &mdash; coordinating technical decisions and ensuring testing supports efficient delivery.</li>
<li>Maintain a constructive, structured, and systematic approach to test strategy evolution &mdash; introducing risk-based testing, improved coverage models, and data-driven quality decisions.</li>
<li>Inspire and motivate team members through positive leadership &mdash; fostering an innovative environment where individuals take ownership and grow professionally.</li>
<li>Manage stakeholder relationships at both operational and strategic levels &mdash; communicating test progress, quality metrics, and improvement initiatives to leadership.</li>
<li>Drive continuous improvement in testing technology: introduced automated test pipelines (CI/CD), performance benchmarking tools, and real-time quality dashboards.</li>
<li>Recognized as Exceptional Performer for 2023 &mdash; acknowledged for leadership impact and testing innovation.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Test &amp; Release Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Developed and maintained test routines for release readiness across a 300M+ user platform &mdash; ensuring systematic quality verification before each production deployment.</li>
<li>Contributed to test planning and coordination across development squads &mdash; aligning testing timelines with product development sprints and release milestones.</li>
<li>Drove improvements in test methods: introduced structured regression approaches, performance validation gates, and automated quality checks in CI/CD pipelines.</li>
<li>Collaborated cross-functionally with engineering and operations teams &mdash; managing relationships and ensuring testing supported efficient product delivery.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Test Lead / Technical Specialist &mdash; Product Testing <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Led testing teams (5&ndash;12 members) across IKEA and LEGO projects &mdash; responsible for team competence development, workload distribution, and performance management.</li>
<li>Developed test methods and routines for enterprise product platforms: system testing, integration testing, performance/load testing, and manufacturing-adjacent hardware-software integration validation.</li>
<li>Drove test planning for complex product development projects &mdash; collaborating with R&amp;D, architects, and product owners to align test activities with delivery milestones.</li>
<li>Built a structured and systematic testing practice: defined test strategies, entry/exit criteria, coverage models, and reusable test frameworks adopted across multiple projects.</li>
<li>Led cross-functional coordination between development teams, product management, and customer representatives across Denmark, Sweden, and India.</li>
<li>Mentored and developed team members &mdash; conducting technical training, introducing new testing technologies, and supporting career growth within the testing discipline.</li>
<li>Managed testing for product releases including hardware-software integration scenarios &mdash; validating system behavior under real-world operational conditions.</li>
<li>Drove continuous improvement in testing processes: reduced regression cycles by 40% through automation, improved defect escape rates, and introduced quality metrics dashboards.</li>
<li>Maintained strong stakeholder relationships &mdash; communicating test results and quality status constructively to programme leadership and customer representatives.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise IT &amp; Banking</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Led testing teams for core banking system implementations &mdash; developing test methods, managing team workload, and ensuring systematic quality assurance across multi-phase deployments.</li>
<li>Defined test routines for integration, data migration, and system validation &mdash; establishing structured approaches adopted across 20+ branch rollouts.</li>
<li>Collaborated with cross-functional teams (development, operations, business) to align test planning with product development timelines and customer expectations.</li>
</ul>

<h2>TECHNICAL ENVIRONMENT</h2>
<p class="tech-line"><b>Testing Technology:</b> Test automation frameworks, performance/load testing, system integration testing, hardware-software validation, CI/CD quality gates</p>
<p class="tech-line"><b>Tools:</b> Jira, Azure DevOps, Confluence, Playwright, Selenium, Pytest, JMeter, Grafana, Datadog</p>
<p class="tech-line"><b>Engineering:</b> Python, TypeScript, Node.js, Docker, Terraform, GCP, AWS, GitHub Actions, Jenkins</p>
<p class="tech-line"><b>Methodologies:</b> Agile (Scrum/Kanban), V-Model, ISTQB, risk-based testing, exploratory testing, TDD</p>
<p class="tech-line"><b>Leadership:</b> Team development, 1-on-1 coaching, competence planning, performance management, stakeholder communication</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Six Sigma Green Belt</td></tr>
<tr><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td><td>&bull; ITIL v4 Foundation</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">Swedish (Fluent) &bull; English (Fluent) &bull; Danish (Conversational) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
