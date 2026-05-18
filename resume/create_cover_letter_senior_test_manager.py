from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Senior_Test_Manager_Cover_Letter.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Senior_Test_Manager_Cover_Letter.doc"

SECTION_COLOR = RGBColor(31, 78, 121)
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)


def add_text(paragraph, text, *, bold=False, size=11, color=TEXT_DARK):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Header ──
    header = document.add_paragraph()
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "Mohammad Kashif", bold=True, size=14, color=SECTION_COLOR)

    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(4)
    add_text(contact, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif", size=10, color=TEXT_MUTED)

    # ── Date ──
    date_p = document.add_paragraph()
    date_p.paragraph_format.space_before = Pt(12)
    date_p.paragraph_format.space_after = Pt(12)
    add_text(date_p, "May 8, 2026", size=11)

    # ── Recipient ──
    to_p = document.add_paragraph()
    to_p.paragraph_format.space_after = Pt(0)
    add_text(to_p, "Hiring Manager", size=11)
    to_p2 = document.add_paragraph()
    to_p2.paragraph_format.space_after = Pt(12)
    add_text(to_p2, "via Incluso — NGO2B Program, Stockholm", size=11)

    # ── Subject ──
    subj = document.add_paragraph()
    subj.paragraph_format.space_after = Pt(12)
    add_text(subj, "Re: Senior Test Manager — NGO2B (Next Generation Offer to Billing)", bold=True, size=11)

    # ── Body ──
    paragraphs = [
        (
            "Dear Hiring Manager,"
        ),
        (
            "I am writing to express my strong interest in the Senior Test Manager role "
            "for the NGO2B program. With 15+ years of experience leading test strategy "
            "and execution across large-scale IT transformation programs — including "
            "integration-heavy solutions, SaaS platforms, and financial systems — I am "
            "confident I can deliver the quality leadership this business-critical "
            "transformation requires."
        ),
        (
            "At IKEA IT AB, I currently own the end-to-end test strategy for a SaaS-integrated "
            "platform serving 30+ global markets. This includes managing testing across API "
            "integrations, event-driven data exchange, and downstream systems — aligning "
            "closely with the integration layer, SaaS vendor coordination, and distributed "
            "architecture patterns central to NGO2B. I define test strategy, drive planning "
            "and execution, manage defects, and ensure go-live readiness while coordinating "
            "across development teams, architects, business stakeholders, and the external "
            "SaaS provider."
        ),
        (
            "What makes me particularly well-suited for this role is my financial sector "
            "background. I spent 5+ years testing core banking systems (Finacle CBS) — "
            "billing, invoicing, accounts receivable, subledger reconciliation, and data "
            "migration between legacy and new platforms. I understand the complexity of "
            "financial data flows, batch processing, and regulatory compliance — directly "
            "relevant to replacing on-prem billing solutions with a modern SaaS architecture."
        ),
        (
            "I bring a structured, pragmatic approach to test management: clear strategy, "
            "risk-based planning, strong stakeholder communication, and a focus on delivering "
            "quality in complex environments. I am experienced in managing multiple "
            "stakeholders across business and technology, securing alignment between internal "
            "teams and external vendors, and providing transparent reporting to program leadership."
        ),
        (
            "I am fluent in both Swedish and English, available to work 100% on-site in "
            "Stockholm/Solna, and can start within 2–4 weeks. I am fully aware of and "
            "comfortable with the 12-month contract structure."
        ),
        (
            "I would welcome the opportunity to discuss how my experience in SaaS "
            "transformation, integration testing, and financial systems can contribute "
            "to the success of the NGO2B program."
        ),
        (
            "Kind regards,"
        ),
    ]

    for text in paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        add_text(p, text, size=11)

    # ── Signature ──
    sig = document.add_paragraph()
    sig.paragraph_format.space_before = Pt(4)
    add_text(sig, "Mohammad Kashif", bold=True, size=11, color=SECTION_COLOR)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Cover Letter – Senior Test Manager</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 48px; color: #1f2937; font-size: 11pt; line-height: 1.5; }
    .name { color: #1f4e79; font-size: 14pt; font-weight: 700; margin-bottom: 2px; }
    .contact { color: #4b5563; font-size: 10pt; margin-bottom: 20px; }
    .subject { font-weight: 700; margin-bottom: 16px; }
    .signature { color: #1f4e79; font-weight: 700; margin-top: 8px; }
    p { margin-bottom: 12px; }
  </style>
</head>
<body>
  <div class="name">Mohammad Kashif</div>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | linkedin.com/in/md-kashif</div>

  <p>May 8, 2026</p>

  <p>Hiring Manager<br>via Incluso — NGO2B Program, Stockholm</p>

  <p class="subject">Re: Senior Test Manager — NGO2B (Next Generation Offer to Billing)</p>

  <p>Dear Hiring Manager,</p>

  <p>I am writing to express my strong interest in the Senior Test Manager role for the NGO2B program. With 15+ years of experience leading test strategy and execution across large-scale IT transformation programs — including integration-heavy solutions, SaaS platforms, and financial systems — I am confident I can deliver the quality leadership this business-critical transformation requires.</p>

  <p>At IKEA IT AB, I currently own the end-to-end test strategy for a SaaS-integrated platform serving 30+ global markets. This includes managing testing across API integrations, event-driven data exchange, and downstream systems — aligning closely with the integration layer, SaaS vendor coordination, and distributed architecture patterns central to NGO2B. I define test strategy, drive planning and execution, manage defects, and ensure go-live readiness while coordinating across development teams, architects, business stakeholders, and the external SaaS provider.</p>

  <p>What makes me particularly well-suited for this role is my financial sector background. I spent 5+ years testing core banking systems (Finacle CBS) — billing, invoicing, accounts receivable, subledger reconciliation, and data migration between legacy and new platforms. I understand the complexity of financial data flows, batch processing, and regulatory compliance — directly relevant to replacing on-prem billing solutions with a modern SaaS architecture.</p>

  <p>I bring a structured, pragmatic approach to test management: clear strategy, risk-based planning, strong stakeholder communication, and a focus on delivering quality in complex environments. I am experienced in managing multiple stakeholders across business and technology, securing alignment between internal teams and external vendors, and providing transparent reporting to program leadership.</p>

  <p>I am fluent in both Swedish and English, available to work 100% on-site in Stockholm/Solna, and can start within 2–4 weeks. I am fully aware of and comfortable with the 12-month contract structure.</p>

  <p>I would welcome the opportunity to discuss how my experience in SaaS transformation, integration testing, and financial systems can contribute to the success of the NGO2B program.</p>

  <p>Kind regards,</p>

  <p class="signature">Mohammad Kashif</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
