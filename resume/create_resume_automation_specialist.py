from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Automation_Specialist_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Automation_Specialist_Resume.doc"

ACCENT = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10.5, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    add_text(paragraph, title.upper(), bold=True, size=11, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "15+ years",
    "Agentic AI",
    "AI-driven",
    "AI/ML",
    "automation",
    "GCP",
    "Google Cloud",
    "Terraform",
    "CI/CD",
    "DevOps",
    "Playwright",
    "data pipeline",
    "data layer",
    "BigQuery",
    "CoBrowser",
    "Real-Time Transcription",
    "Visual Collaboration",
    "business growth",
    "plug and play",
    "sales conversion",
    "Kubernetes",
    "Docker",
    "Cloud Run",
    "GitHub Actions",
    "product ownership",
    "cross-functional",
    "stakeholder",
    "strategic",
    "LangChain",
    "scalable",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.06
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10.5)
            add_text(paragraph, matched, bold=True, size=10.5, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10.5)
            return
    add_text(paragraph, text, size=10.5)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=20, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_after = Pt(2)
    add_text(
        role,
        "Automation Specialist | AI & Agentic Workflows | DevOps & GCP Cloud | Product Delivery",
        bold=True,
        size=11.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(4)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=10, color=TEXT_MUTED)
    add_text(
        contact,
        "\nLinkedIn: linkedin.com/in/md-kashif | GitHub: github.com/ascertain | Blog: ascertain.github.io",
        size=10,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="7F8FA6", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(4)
    profile.paragraph_format.line_spacing = 1.08
    add_text(
        profile,
        "Results-driven Automation Specialist with 15+ years of experience spanning test automation, "
        "AI-driven quality engineering, DevOps, and cloud-native delivery on GCP. Proven track record of "
        "owning end-to-end product delivery — from strategic vision and stakeholder alignment to hands-on "
        "automation, data pipeline design, and production rollout. Led the Visual Collaboration Solution "
        "(VCS) programme at IKEA, enabling business growth through Agentic AI workflows, CoBrowser, "
        "Real-Time Transcription, and Call Quality Assessment capabilities for remote planning and complex "
        "resolution scenarios. Architected a flexible, plug-and-play Business API model and a scalable data "
        "layer for sales conversion analytics and proactive follow-up identification. Experienced in building, "
        "testing, releasing, and rolling out consumer-scale apps (IKEA, LEGO, Truecaller) to global markets. "
        "Combines deep technical understanding with business acumen to drive measurable outcomes in fast-paced, "
        "cross-functional environments.",
        size=10.5,
    )

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "AI & Agentic Workflows: ",
            "Agentic AI design, LangChain, AI-assisted quality gates, intelligent test orchestration, "
            "conversational AI integration, prompt engineering, AI-driven analytics",
        ),
        (
            "Automation & Quality: ",
            "Playwright, Selenium, end-to-end automation, API testing, performance testing, "
            "test strategy, shift-left quality, exploratory testing, regression suites",
        ),
        (
            "DevOps & Cloud (GCP): ",
            "Google Cloud Platform, Cloud Run, BigQuery, Pub/Sub, Eventarc, Cloud Functions, "
            "Terraform, Docker, Kubernetes, CI/CD pipelines, GitHub Actions, infrastructure as code",
        ),
        (
            "Product & Delivery: ",
            "Product ownership, roadmap planning, stakeholder management, Agile/Scrum, "
            "cross-functional leadership, release management, market rollout, KPI-driven delivery",
        ),
        (
            "Data & Analytics: ",
            "Data pipeline design, sales conversion analysis, BigQuery analytics, "
            "proactive follow-up identification, business intelligence, data-driven decision-making",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        add_text(paragraph, label, bold=True, size=10.5)
        add_text(paragraph, value, size=10.5)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")
    jobs = [
        (
            "IKEA IT AB, Malmö, Sweden",
            "Automation Specialist / Product Owner – Visual Collaboration Solution | Mar 2022 – Present",
            [
                "Own the Visual Collaboration Solution (VCS) end-to-end — driving strategic direction, "
                "feature delivery, and business growth through AI-powered capabilities including CoBrowser, "
                "Real-Time Transcription, and Call Quality Assessments for IKEA's remote planning and "
                "complex resolution calls.",
                "Designed and delivered an Agentic AI workflow layer that automates quality assessments, "
                "intelligently routes customer interactions, and surfaces actionable insights to business "
                "stakeholders for continuous improvement.",
                "Architected a flexible, plug-and-play Business API model enabling any product team to "
                "integrate collaboration features with minimal effort — reducing onboarding time and "
                "accelerating feature adoption across business units.",
                "Built a scalable data layer on GCP (BigQuery, Pub/Sub, Cloud Functions, Eventarc) to "
                "process interaction files, analyze sales conversion rates, and identify high-potential "
                "calls for future follow-up — directly contributing to revenue growth.",
                "Led automation strategy across the platform using Playwright and CI/CD pipelines on "
                "GitHub Actions, enabling continuous deployment with multiple daily releases and near-zero "
                "regression risk.",
                "Managed DevOps and cloud infrastructure on GCP using Terraform, Docker, Kubernetes, and "
                "Cloud Run — ensuring scalable, cost-efficient, and reliable production environments.",
                "Collaborated with cross-functional stakeholders (product, design, engineering, business) "
                "to align technical delivery with business outcomes, KPIs, and market requirements.",
                "Recognized as an Exceptional Performer for sustained ownership, leadership, and measurable "
                "business impact across the VCS programme.",
            ],
        ),
        (
            "Truecaller, Sweden",
            "Release & Automation Engineer | Sep 2021 – Feb 2022",
            [
                "Owned release automation and quality gates for the Truecaller consumer app — serving "
                "300M+ active users across global markets.",
                "Built and maintained automated test suites and release validation pipelines to ensure "
                "rapid, confident rollouts in a high-velocity Agile environment.",
                "Collaborated with product and engineering teams to investigate production issues, improve "
                "test coverage, and strengthen release readiness processes.",
                "Contributed to the build, test, and market rollout workflow for the Truecaller app across "
                "Android and iOS platforms.",
            ],
        ),
        (
            "LEGO Group (via HCLTech), Denmark",
            "Senior Automation Engineer | 2019 – 2021",
            [
                "Led test automation for LEGO's digital commerce and consumer-facing applications — building "
                "robust Selenium/Playwright suites across web and mobile platforms.",
                "Supported the build, test, release, and market rollout lifecycle for the LEGO App, ensuring "
                "quality at scale for millions of users worldwide.",
                "Drove shift-left testing practices and CI/CD integration, embedding automation into every "
                "stage of the development pipeline.",
                "Mentored team members on automation architecture, test design patterns, and DevOps best "
                "practices across Scrum teams.",
            ],
        ),
        (
            "HCLTech / Enterprise Programs",
            "Automation Lead / SDET | Dec 2013 – 2019",
            [
                "Led automation initiatives across enterprise programmes for global clients — designing "
                "scalable test frameworks, CI/CD pipelines, and quality dashboards.",
                "Built end-to-end automation for web, API, and integrated platforms using Selenium, C#, "
                "and Python — improving release confidence and reducing manual effort by 70%.",
                "Partnered with development and operations teams to implement DevOps practices, "
                "infrastructure automation, and continuous delivery pipelines.",
                "Mentored and trained 15+ engineers on automation strategy, tooling, and collaborative "
                "quality practices across distributed teams.",
            ],
        ),
    ]
    for company, title, bullets in jobs:
        cp = document.add_paragraph()
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(0)
        add_text(cp, company, bold=True, size=11, color=SECTION_COLOR)

        tp = document.add_paragraph()
        tp.paragraph_format.space_after = Pt(0)
        add_text(tp, title, bold=True, size=10.5)

        for bullet in bullets:
            add_highlighted_bullet(document, bullet)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Owned the Visual Collaboration Solution at IKEA, enabling CoBrowser, Real-Time Transcription, "
        "and AI-driven Call Quality Assessments that increased remote planning adoption by 40%.",
        "Designed a plug-and-play Business API architecture allowing 5+ product teams to integrate "
        "collaboration capabilities with zero custom development.",
        "Built a GCP data pipeline processing 100K+ interaction records, delivering sales conversion "
        "analytics and identifying 15% more high-potential follow-up opportunities.",
        "Implemented Agentic AI workflows that automated quality scoring and intelligent call routing, "
        "reducing manual assessment effort by 60%.",
        "Delivered consumer-scale apps (IKEA, LEGO, Truecaller) through full build-test-release-rollout "
        "cycles to global markets serving 300M+ users.",
        "Established automation-first culture and CI/CD excellence — achieving 10+ daily releases with "
        "near-zero regression across GCP-hosted microservices.",
    ]
    for achievement in achievements:
        add_highlighted_bullet(document, achievement)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10.5, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "B.Tech in Information Technology\nPG Diploma in Operations Management", size=10.2)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester\nITIL Foundation\nAWS Cloud Practitioner\n"
        "Google Cloud Professional\nCertified Ethical Hacker\nSix Sigma Green Belt",
        size=10.2,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic  |  Hindi / Urdu — Native", size=10.5)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Automation Specialist Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 32px; color: #1f2937; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 4px; font-size: 24pt; }
    h2 { text-align: center; font-size: 12pt; margin-top: 0; margin-bottom: 8px; }
    .contact { text-align: center; color: #4b5563; border-bottom: 2px solid #1f4e79; padding-bottom: 10px; margin-bottom: 14px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 11pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 3px; margin-top: 14px; margin-bottom: 8px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 8px; }
    .job-sub { font-weight: 700; margin-bottom: 4px; }
    ul { margin-top: 2px; margin-bottom: 6px; }
    li { margin-bottom: 4px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 6px; }
    td { border: 1px solid #b8cce4; padding: 7px; vertical-align: top; }
    .tag { background: #d9eaf7; font-weight: 700; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Automation Specialist | AI &amp; Agentic Workflows | DevOps &amp; GCP Cloud | Product Delivery</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com<br>LinkedIn: linkedin.com/in/md-kashif | GitHub: github.com/ascertain | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Results-driven Automation Specialist with <span class="hl">15+ years</span> of experience spanning test <span class="hl">automation</span>, <span class="hl">AI-driven</span> quality engineering, <span class="hl">DevOps</span>, and cloud-native delivery on <span class="hl">GCP</span>. Proven track record of owning end-to-end product delivery — from <span class="hl">strategic</span> vision and <span class="hl">stakeholder</span> alignment to hands-on automation, <span class="hl">data pipeline</span> design, and production rollout. Led the <span class="hl">Visual Collaboration</span> Solution (VCS) programme at IKEA, enabling <span class="hl">business growth</span> through <span class="hl">Agentic AI</span> workflows, <span class="hl">CoBrowser</span>, <span class="hl">Real-Time Transcription</span>, and Call Quality Assessment capabilities for remote planning and complex resolution scenarios. Architected a flexible, <span class="hl">plug and play</span> Business API model and a <span class="hl">scalable</span> <span class="hl">data layer</span> for <span class="hl">sales conversion</span> analytics and proactive follow-up identification. Experienced in building, testing, releasing, and rolling out consumer-scale apps (IKEA, LEGO, Truecaller) to global markets.</p>

  <div class="section">Core Competencies</div>
  <p><b>AI &amp; Agentic Workflows:</b> <span class="hl">Agentic AI</span> design, <span class="hl">LangChain</span>, AI-assisted quality gates, intelligent test orchestration, conversational AI integration, prompt engineering, <span class="hl">AI-driven</span> analytics<br>
  <b>Automation &amp; Quality:</b> <span class="hl">Playwright</span>, Selenium, end-to-end automation, API testing, performance testing, test strategy, shift-left quality<br>
  <b>DevOps &amp; Cloud (GCP):</b> <span class="hl">Google Cloud</span> Platform, <span class="hl">Cloud Run</span>, <span class="hl">BigQuery</span>, Pub/Sub, Eventarc, Cloud Functions, <span class="hl">Terraform</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">CI/CD</span> pipelines, <span class="hl">GitHub Actions</span>, infrastructure as code<br>
  <b>Product &amp; Delivery:</b> <span class="hl">Product ownership</span>, roadmap planning, <span class="hl">stakeholder</span> management, Agile/Scrum, <span class="hl">cross-functional</span> leadership, release management, market rollout, KPI-driven delivery<br>
  <b>Data &amp; Analytics:</b> <span class="hl">Data pipeline</span> design, <span class="hl">sales conversion</span> analysis, <span class="hl">BigQuery</span> analytics, proactive follow-up identification, business intelligence, data-driven decision-making</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö, Sweden</div>
  <div class="job-sub">Automation Specialist / Product Owner – Visual Collaboration Solution | Mar 2022 – Present</div>
  <ul>
    <li>Own the <span class="hl">Visual Collaboration</span> Solution (VCS) end-to-end — driving <span class="hl">strategic</span> direction, feature delivery, and <span class="hl">business growth</span> through AI-powered capabilities including <span class="hl">CoBrowser</span>, <span class="hl">Real-Time Transcription</span>, and Call Quality Assessments for IKEA's remote planning and complex resolution calls.</li>
    <li>Designed and delivered an <span class="hl">Agentic AI</span> workflow layer that automates quality assessments, intelligently routes customer interactions, and surfaces actionable insights to business <span class="hl">stakeholder</span>s.</li>
    <li>Architected a flexible, <span class="hl">plug and play</span> Business API model enabling any product team to integrate collaboration features with minimal effort — reducing onboarding time and accelerating feature adoption.</li>
    <li>Built a <span class="hl">scalable</span> <span class="hl">data layer</span> on <span class="hl">GCP</span> (<span class="hl">BigQuery</span>, Pub/Sub, Cloud Functions, Eventarc) to process interaction files, analyze <span class="hl">sales conversion</span> rates, and identify high-potential calls for future follow-up.</li>
    <li>Led <span class="hl">automation</span> strategy across the platform using <span class="hl">Playwright</span> and <span class="hl">CI/CD</span> pipelines on <span class="hl">GitHub Actions</span>, enabling continuous deployment with multiple daily releases.</li>
    <li>Managed <span class="hl">DevOps</span> and cloud infrastructure on <span class="hl">GCP</span> using <span class="hl">Terraform</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, and <span class="hl">Cloud Run</span>.</li>
    <li>Collaborated with <span class="hl">cross-functional</span> <span class="hl">stakeholder</span>s (product, design, engineering, business) to align technical delivery with business outcomes and KPIs.</li>
    <li>Recognized as an Exceptional Performer for sustained ownership, leadership, and measurable business impact.</li>
  </ul>

  <div class="job-title">Truecaller, Sweden</div>
  <div class="job-sub">Release &amp; Automation Engineer | Sep 2021 – Feb 2022</div>
  <ul>
    <li>Owned release <span class="hl">automation</span> and quality gates for the Truecaller consumer app — serving 300M+ active users across global markets.</li>
    <li>Built and maintained automated test suites and release validation pipelines for rapid, confident rollouts.</li>
    <li>Contributed to the build, test, and market rollout workflow for the Truecaller app across Android and iOS platforms.</li>
  </ul>

  <div class="job-title">LEGO Group (via HCLTech), Denmark</div>
  <div class="job-sub">Senior Automation Engineer | 2019 – 2021</div>
  <ul>
    <li>Led test <span class="hl">automation</span> for LEGO's digital commerce and consumer-facing applications.</li>
    <li>Supported the build, test, release, and market rollout lifecycle for the LEGO App serving millions of users worldwide.</li>
    <li>Drove shift-left testing practices and <span class="hl">CI/CD</span> integration across Scrum teams.</li>
  </ul>

  <div class="job-title">HCLTech / Enterprise Programs</div>
  <div class="job-sub">Automation Lead / SDET | Dec 2013 – 2019</div>
  <ul>
    <li>Led <span class="hl">automation</span> initiatives across enterprise programmes — designing <span class="hl">scalable</span> test frameworks, <span class="hl">CI/CD</span> pipelines, and quality dashboards.</li>
    <li>Partnered with development and operations teams to implement <span class="hl">DevOps</span> practices and continuous delivery pipelines.</li>
    <li>Mentored 15+ engineers on <span class="hl">automation</span> strategy and collaborative quality practices.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Owned the <span class="hl">Visual Collaboration</span> Solution at IKEA, enabling <span class="hl">CoBrowser</span>, <span class="hl">Real-Time Transcription</span>, and <span class="hl">AI-driven</span> Call Quality Assessments — increasing remote planning adoption by 40%.</li>
    <li>Designed a <span class="hl">plug and play</span> Business API architecture allowing 5+ product teams to integrate collaboration capabilities with zero custom development.</li>
    <li>Built a <span class="hl">GCP</span> <span class="hl">data pipeline</span> processing 100K+ interaction records, delivering <span class="hl">sales conversion</span> analytics and identifying 15% more high-potential follow-up opportunities.</li>
    <li>Implemented <span class="hl">Agentic AI</span> workflows that automated quality scoring and intelligent call routing, reducing manual assessment effort by 60%.</li>
    <li>Delivered consumer-scale apps (IKEA, LEGO, Truecaller) through full build-test-release-rollout cycles to global markets serving 300M+ users.</li>
    <li>Established <span class="hl">automation</span>-first culture and <span class="hl">CI/CD</span> excellence — achieving 10+ daily releases with near-zero regression across <span class="hl">GCP</span>-hosted microservices.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester<br>ITIL Foundation<br>AWS Cloud Practitioner<br>Google Cloud Professional<br>Certified Ethical Hacker<br>Six Sigma Green Belt</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic  |  Hindi / Urdu — Native</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
