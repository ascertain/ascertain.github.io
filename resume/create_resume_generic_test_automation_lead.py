from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_SDET_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_SDET_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "SDET",
    "Test Automation",
    "Test Lead",
    "Playwright",
    "Selenium",
    "Cypress",
    "Appium",
    "RestAssured",
    "Karate",
    "TestNG",
    "Jest",
    "Vitest",
    "JMeter",
    "Postman",
    "Grafana",
    "APIM",
    "CI/CD",
    "GitHub Actions",
    "Jenkins",
    "Docker",
    "Kubernetes",
    "AWS",
    "GCP",
    "Azure",
    "Terraform",
    "AI",
    "Claude",
    "Copilot",
    "Gemini",
    "Java",
    "C#",
    "TypeScript",
    "Python",
    "API",
    "E2E",
    "regression",
    "performance",
    "load testing",
    "shift-left",
    "Agile",
    "Scrum",
    "SaaS",
    "mobile",
    "iOS",
    "Android",
    "SQL",
    "Maven",
    "IKEA",
    "LEGO",
    "Truecaller",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "SDET  |  Full-Stack Test Automation & AI-Driven Quality  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "SDET with 15+ years building robust testing frameworks and driving quality across "
        "full-stack products. Deep automation skills across web (Playwright, Selenium, Cypress), "
        "mobile (Appium — iOS/Android), and API (RestAssured, Karate, Postman). Integrates "
        "testing into CI/CD pipelines (GitHub Actions, Jenkins, Docker, Kubernetes). Daily AI "
        "tool user — Claude, Copilot, Gemini for AI-driven test generation, exploratory testing, "
        "and workflow automation. Polyglot engineer — Java, C#, Python, TypeScript, React. "
        "Performance testing with JMeter. APIM testing and API gateway validation. Shift-left "
        "advocate who coaches teams and builds quality culture. Experience across SaaS, "
        "e-commerce, mobile, fintech, and enterprise at companies including IKEA, Truecaller, "
        "and LEGO. Agile/Scrum practitioner. Collaborative, proactive, and takes ownership.",
        size=10,
    )

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Test Automation: ",
            "Playwright, Selenium, Cypress, Appium (iOS/Android), RestAssured, "
            "Karate, TestNG, Jest, Vitest, Postman, E2E, API, regression, contract testing",
        ),
        (
            "Languages: ",
            "Java, C#, Python, TypeScript, JavaScript, SQL",
        ),
        (
            "CI/CD & Cloud: ",
            "GitHub Actions, Jenkins, Docker, Kubernetes (K8s), AWS, GCP, Azure, "
            "Terraform, pipeline design, infrastructure-as-code",
        ),
        (
            "Performance & Observability: ",
            "JMeter (load, stress, soak, spike), Grafana, monitoring, benchmarking",
        ),
        (
            "API & Integration: ",
            "APIM, API gateway testing, RestAssured, Karate, Postman, contract testing",
        ),
        (
            "AI-Driven Testing: ",
            "Claude, GitHub Copilot, Gemini — AI test generation, code review, "
            "exploratory testing agents, workflow automation",
        ),
        (
            "Frameworks & Tools: ",
            "Spring Boot, React, .NET, Maven, TestRail, Git, Jira, Confluence",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — SDET / Test Automation Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — SaaS Platform, Full-Stack, Multi-Market", bold=True, size=10)
    ikea_bullets = [
        "Designed and built test automation frameworks — Playwright E2E, API testing "
        "(RestAssured, Karate), regression suites, contract testing. Integrated into "
        "CI/CD pipelines (GitHub Actions, Docker, Terraform) serving 30+ markets.",
        "AI-driven testing — daily Claude, Copilot, Gemini user. AI-assisted test "
        "generation, code reviews, exploratory testing agents. Rolled out AI workflows "
        "to the team — 30% velocity improvement.",
        "Performance testing — JMeter load/stress/spike testing for APIs. APIM and "
        "API gateway validation. Monitoring (iLert, Grafana), observability, automated "
        "error detection. Shift-left testing built into the development process.",
        "Coached engineers on test automation best practices. Built quality culture "
        "where testing is everyone's ownership. Kubernetes, Docker, cloud-native.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — SDET / Senior QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  SaaS — 300M+ Users, Mobile-First", bold=True, size=10)
    tc_bullets = [
        "Test automation across web (Selenium), mobile (Appium — iOS/Android), and "
        "API (RestAssured, Karate, TestNG). E2E, regression, exploratory, and "
        "performance testing at a fast-moving SaaS company with 300M+ users.",
        "Performance and load testing (JMeter) for high-traffic APIs. Coached "
        "engineers on modern testing practices. CI/CD pipeline integration. "
        "Agile (Scrum/Kanban), autonomous, collaborative team environment.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Web & Mobile, Multi-Platform", bold=True, size=10)
    lego_bullets = [
        "LEGO: Built automation frameworks — Selenium, RestAssured, Karate, TestNG, "
        "Appium (iOS/Android). E2E, regression, API testing. Led 8–10 engineers. "
        "Maven, SQL, TestRail.",
        "IKEA App, Genesys, Verint (2018–2021): Test automation across multiple "
        "full-stack products. CI/CD pipelines, shift-left testing, JMeter performance "
        "testing. Coached developers on quality ownership.",
        "Agile (Scrum), cross-functional teams, continuous improvement. Drove quality "
        "culture across multiple product teams.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech & Government — Full-Stack, Multi-Platform", bold=True, size=10)
    fin_bullets = [
        "Built automation frameworks from scratch — Selenium, RestAssured, TestNG. "
        "E2E, regression, API, and system testing across enterprise banking platforms "
        "(Finacle). JMeter performance testing. SQL, Maven, CI/CD integration.",
        "Coached 15+ engineers on test automation. Led teams of 10+. Shift-left "
        "practices, Agile (Scrum/Kanban), quality ownership.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Built robust test automation frameworks (Playwright, Selenium, Cypress, Appium, "
        "RestAssured, Karate) across web, mobile (iOS/Android), and API layers.",
        "AI-driven testing pioneer — Claude, Copilot, Gemini daily. AI test generation, "
        "exploratory testing agents. 30% velocity improvement. Team rollout.",
        "Seamless CI/CD pipelines (GitHub Actions, Jenkins, Docker, Kubernetes) — "
        "shift-left testing, automated tests on every PR, short cycle times.",
        "Performance testing (JMeter) — load/stress/spike for high-traffic APIs serving "
        "300M+ users (Truecaller) and 30+ markets (IKEA). APIM and API gateway testing.",
        "Coached engineers across every organization — built whole-team quality culture. "
        "Collaborative, proactive, ownership-driven.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Test Automation Lead / SDET Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>SDET | Full-Stack Test Automation &amp; AI-Driven Quality | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">SDET</span> with 15+ years building robust testing frameworks and driving quality across full-stack products. Deep automation across web (<span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Cypress</span>), mobile (<span class="hl">Appium</span> — <span class="hl">iOS</span>/<span class="hl">Android</span>), and <span class="hl">API</span> (<span class="hl">RestAssured</span>, <span class="hl">Karate</span>, <span class="hl">Postman</span>). <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). Daily <span class="hl">AI</span> user — <span class="hl">Claude</span>, <span class="hl">Copilot</span>, <span class="hl">Gemini</span>. Polyglot: <span class="hl">Java</span>, <span class="hl">C#</span>, <span class="hl">Python</span>, <span class="hl">TypeScript</span>. Performance (<span class="hl">JMeter</span>). <span class="hl">APIM</span> &amp; API gateway testing. Shift-left. SaaS, e-commerce, mobile, fintech at <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. <span class="hl">Agile</span>/<span class="hl">Scrum</span>.</p>

  <div class="section">Technical Skills</div>
  <p><b>Test Automation:</b> Playwright, Selenium, Cypress, Appium (iOS/Android), RestAssured, Karate, TestNG, Jest, Vitest, Postman, E2E, API, regression, contract testing<br>
  <b>Languages:</b> Java, C#, Python, TypeScript, JavaScript, SQL<br>
  <b>CI/CD &amp; Cloud:</b> GitHub Actions, Jenkins, Docker, Kubernetes, AWS, GCP, Azure, Terraform<br>
  <b>Performance:</b> JMeter (load, stress, soak, spike), Grafana, monitoring, benchmarking<br>
  <b>API &amp; Integration:</b> APIM, API gateway testing, RestAssured, Karate, Postman, contract testing<br>
  <b>AI-Driven Testing:</b> Claude, Copilot, Gemini — test generation, code review, exploratory agents<br>
  <b>Frameworks &amp; Tools:</b> Spring Boot, React, .NET, Maven, TestRail, Git, Jira, Confluence</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — SDET / Test Automation Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — SaaS Platform, Full-Stack, Multi-Market</div>
  <ul>
    <li>Built <span class="hl">test automation</span> frameworks — <span class="hl">Playwright</span> E2E, <span class="hl">API</span> testing (<span class="hl">RestAssured</span>, <span class="hl">Karate</span>), <span class="hl">regression</span>, contract testing. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, <span class="hl">Terraform</span>). 30+ markets.</li>
    <li><span class="hl">AI</span>-driven testing — <span class="hl">Claude</span>, <span class="hl">Copilot</span>, <span class="hl">Gemini</span> daily. AI test generation, exploratory agents. 30% velocity. Team rollout.</li>
    <li><span class="hl">Performance</span> testing — <span class="hl">JMeter</span>, <span class="hl">Grafana</span>, monitoring. <span class="hl">APIM</span> &amp; API gateway. <span class="hl">Shift-left</span>. <span class="hl">Kubernetes</span>, cloud-native.</li>
    <li>Coached engineers on <span class="hl">test automation</span>. Quality culture. <span class="hl">Agile</span> (<span class="hl">Scrum</span>).</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — SDET / Senior QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | SaaS — 300M+ Users, Mobile-First</div>
  <ul>
    <li><span class="hl">Test automation</span> — <span class="hl">Selenium</span>, <span class="hl">Appium</span> (<span class="hl">iOS</span>/<span class="hl">Android</span>), <span class="hl">RestAssured</span>, <span class="hl">Karate</span>, <span class="hl">TestNG</span>. E2E, regression, exploratory. 300M+ users.</li>
    <li><span class="hl">JMeter</span> <span class="hl">performance</span>/<span class="hl">load testing</span> for high-traffic APIs. Coached engineers. <span class="hl">CI/CD</span>. <span class="hl">Agile</span>, collaborative, autonomous.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — SDET / Test Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Web &amp; Mobile, Multi-Platform</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span>, <span class="hl">TestNG</span>, <span class="hl">Appium</span> (<span class="hl">iOS</span>/<span class="hl">Android</span>). E2E, regression, <span class="hl">API</span>. 8–10 engineers. <span class="hl">Maven</span>, <span class="hl">SQL</span>, TestRail.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Test automation</span> across multiple full-stack products. <span class="hl">CI/CD</span>, <span class="hl">shift-left</span>, <span class="hl">JMeter</span> <span class="hl">performance</span> testing. Coached developers.</li>
    <li><span class="hl">Agile</span> (<span class="hl">Scrum</span>), cross-functional teams, continuous improvement.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior SDET / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech &amp; Government — Full-Stack, Multi-Platform</div>
  <ul>
    <li>Automation frameworks — <span class="hl">Selenium</span>, <span class="hl">RestAssured</span>, <span class="hl">TestNG</span>. E2E, regression, <span class="hl">API</span>. <span class="hl">JMeter</span> <span class="hl">performance</span>. <span class="hl">SQL</span>, <span class="hl">Maven</span>, <span class="hl">CI/CD</span>.</li>
    <li>Coached 15+ engineers. <span class="hl">Shift-left</span>. <span class="hl">Agile</span> (<span class="hl">Scrum</span>). Quality ownership.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Robust <span class="hl">test automation</span> frameworks — <span class="hl">Playwright</span>, <span class="hl">Selenium</span>, <span class="hl">Cypress</span>, <span class="hl">Appium</span>, <span class="hl">RestAssured</span>, <span class="hl">Karate</span> across web, <span class="hl">mobile</span>, <span class="hl">API</span>.</li>
    <li><span class="hl">AI</span>-driven testing — <span class="hl">Claude</span>, <span class="hl">Copilot</span>, <span class="hl">Gemini</span>. Test generation, exploratory agents. 30% velocity. Team rollout.</li>
    <li><span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>) — <span class="hl">shift-left</span>, automated on every PR.</li>
    <li><span class="hl">Performance</span> (<span class="hl">JMeter</span>) — 300M+ users (<span class="hl">Truecaller</span>), 30+ markets (<span class="hl">IKEA</span>). <span class="hl">APIM</span> &amp; API gateway testing.</li>
    <li>Coached engineers everywhere — whole-team quality culture.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
