"""
Resume: IT Service Management Manager — ITSM, Team Leadership, DevOps
Focus: ITIL processes, Incident/Problem/Change management, team leadership, stakeholder alignment.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_IT_Service_Management_Manager_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("IT Service Management leader with ", False),
        ("16+ years", True),
        (" in IT operations, service delivery, and ", False),
        ("team leadership", True),
        (" within complex, global, cross-functional environments. Proven experience owning and evolving ", False),
        ("ITSM processes", True),
        (" (Incident, Problem, Change, Request, Monitoring & Alert Management). Track record of ", False),
        ("leading and developing teams", True),
        (", driving ", False),
        ("service management reporting", True),
        (" (availability, performance, KPIs), and ", False),
        ("optimising service delivery", True),
        (" within budget constraints. Strong ", False),
        ("stakeholder management", True),
        (" skills, balancing ", False),
        ("operational execution with strategic thinking", True),
        (". Experienced in implementing ", False),
        ("agile and DevOps", True),
        (" ways of working and driving ", False),
        ("continuous improvement", True),
        (" across enterprise-scale IT environments.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "ITIL / ITSM Processes", "Incident & Problem Management", "Change & Request Management",
        "Team Leadership & Coaching", "Service Performance & KPIs", "Monitoring & Alert Management",
        "Capacity & Resource Planning", "Agile & DevOps Practices", "Stakeholder Communication",
        "Operational Excellence", "Risk & Bottleneck Identification", "Strategic Thinking (6\u201318mo)",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical & Domain Skills")
    skills_data = [
        ("ITSM:", " Incident, Problem, Change, Request, Monitoring & Alert Management, SLA/SLO management"),
        ("Tools:", " JIRA Service Management, JIRA Operations, ServiceNow, Confluence, PagerDuty, Opsgenie"),
        ("Monitoring:", " Grafana, Datadog, GCP Cloud Monitoring, alerting pipelines, availability dashboards"),
        ("DevOps & CI/CD:", " GitHub Actions, Jenkins, Docker, Kubernetes, GitOps, automated deployments"),
        ("Cloud & Infra:", " GCP (Cloud Run, GKE, Pub/Sub, BigQuery), AWS, Terraform, microservices"),
        ("Reporting:", " Service performance metrics, KPI dashboards, capacity planning, budget tracking"),
        ("Practices:", " ITIL v4, Agile/Scrum, Kanban, DevOps, blameless post-mortems, OKRs"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Team Lead (Acting) / Service & Delivery Manager",
        "Mar 2022 \u2013 Present",
        "Global Customer Connect Platform \u2014 30+ Markets | Enterprise-Scale IT Operations | Agile & DevOps")
    add_bullet(doc, "Own and evolve IT Service Management processes (Incident, Problem, Change, Monitoring & Alert Management) for a platform serving 30+ global markets.",
        ["IT Service Management processes", "Incident, Problem, Change, Monitoring & Alert Management", "30+ global markets"])
    add_bullet(doc, "Lead and develop two functional areas: service management operations and 2nd-line application support \u2014 building high-performing, accountable teams.",
        ["Lead and develop", "service management", "2nd-line application support", "high-performing"])
    add_bullet(doc, "Drive service management reporting: availability dashboards, performance metrics, KPIs (MTTR, MTTD, SLA compliance) \u2014 translating data into actionable insights.",
        ["service management reporting", "availability", "performance metrics", "KPIs", "actionable insights"])
    add_bullet(doc, "Ensure effective capacity and resource planning across teams; manage deviations proactively; balance delivery within budget and operational constraints.",
        ["capacity and resource planning", "manage deviations", "budget and operational constraints"])
    add_bullet(doc, "Identify risks, bottlenecks, and improvement opportunities \u2014 structured problem-solving and root cause analysis to optimise service delivery.",
        ["risks, bottlenecks", "improvement opportunities", "optimise service delivery"])
    add_bullet(doc, "Contribute to mid-term strategy (6\u201318 months): roadmap alignment, technology modernisation, and service maturity evolution.",
        ["mid-term strategy", "roadmap alignment", "service maturity"])
    add_bullet(doc, "Implement and strengthen agile and DevOps ways of working \u2014 CI/CD pipelines, automated monitoring, infrastructure as code, blameless post-mortems.",
        ["agile and DevOps", "CI/CD pipelines", "automated monitoring", "blameless post-mortems"])
    add_bullet(doc, "Coach and support team members in their growth; drive recruitment, competence development, and retention; create conditions for high delivery efficiency.",
        ["Coach and support", "recruitment", "competence development", "retention", "delivery efficiency"])
    add_bullet(doc, "Collaborate with stakeholders across IT, Product, and business functions \u2014 aligning execution with strategic priorities in a fast-changing global environment.",
        ["stakeholders across IT, Product, and business", "strategic priorities", "global environment"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Operations Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Service Reliability & Release Operations")
    add_bullet(doc, "Managed release operations and service reliability for high-scale platform \u2014 incident response, deployment pipelines, monitoring and alerting.",
        ["release operations", "service reliability", "incident response", "monitoring and alerting"])
    add_bullet(doc, "Drove improvements in change management processes; ensured service performance metrics met SLA targets across environments.",
        ["change management", "service performance metrics", "SLA targets"])
    add_bullet(doc, "Collaborated cross-functionally to identify and resolve production issues; structured post-incident reviews driving process improvements.",
        ["cross-functionally", "production issues", "post-incident reviews"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 Technical Lead / Service Delivery Lead",
        "2013 \u2013 2021",
        "Enterprise IT Operations \u2014 Global E-Commerce & Digital Platforms | Multi-Geography Delivery")
    add_bullet(doc, "Led IT service delivery across complex, global, cross-functional programmes \u2014 managed distributed teams (8\u201312) spanning multiple geographies.",
        ["IT service delivery", "global, cross-functional", "distributed teams", "multiple geographies"])
    add_bullet(doc, "Owned incident and problem management processes; established monitoring, alerting, and escalation frameworks for enterprise platforms.",
        ["incident and problem management", "monitoring, alerting, and escalation"])
    add_bullet(doc, "Drove change management and continuous improvement \u2014 evolving ways of working from manual operations to automated, DevOps-driven service delivery.",
        ["change management", "continuous improvement", "DevOps-driven service delivery"])
    add_bullet(doc, "Service management reporting: defined KPIs, tracked availability and performance, delivered executive dashboards for stakeholder communication.",
        ["Service management reporting", "KPIs", "availability and performance", "stakeholder communication"])
    add_bullet(doc, "Resource planning, capacity management, and budget-conscious optimisation across multiple concurrent projects and service lines.",
        ["Resource planning", "capacity management", "budget-conscious optimisation"])
    add_bullet(doc, "Strengthened collaboration across teams and geographies; coached engineers on operational best practices and service ownership.",
        ["collaboration across teams and geographies", "coached", "service ownership"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 IT Operations / Consultant",
        "2008 \u2013 2013",
        "Core Banking Systems \u2014 Enterprise-Scale IT (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "IT operations and service management for Finacle Core Banking \u2014 incident management, change control, and compliance in regulated environments.",
        ["IT operations", "incident management", "change control", "regulated environments"])
    add_bullet(doc, "Built monitoring and alert automation; service performance tracking for mission-critical banking transaction systems.",
        ["monitoring and alert automation", "service performance", "mission-critical"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "Six Sigma Green Belt", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
