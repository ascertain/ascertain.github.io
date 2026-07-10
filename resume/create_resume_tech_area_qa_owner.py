"""
Generate a tailored resume for Tech Area QA Owner (Game Development).
Focus: technical QA ownership, system-level testing (performance, stability, reliability),
QA strategy, cross-functional leadership through influence, risk management, developer
enablement, internal/external QA coordination, hands-on + strategic balance.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Tech_Area_QA_Owner_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_Tech_Area_QA_Owner_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "technical QA", "QA strategy", "QA ownership",
    "stability", "performance", "reliability",
    "systemic risk", "risk management", "risk assessment",
    "technical quality", "quality outcomes",
    "developer enablement", "critical friend",
    "end-to-end", "systems-heavy",
    "cross-functional", "influence", "collaboration",
    "internal and external", "external partners",
    "test planning", "coverage", "test strategy",
    "reporting", "triaging", "tracking",
    "technical debt", "continuous improvement",
    "hands-on", "strategic",
    "platform", "compatibility",
    "CI/CD", "automation",
    "Python", "pytest",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Tech Area QA Owner", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Technical QA leader working in IT since 2008, with deep experience owning quality outcomes end-to-end "
        "across complex, systems-heavy platforms. Proven ability to define and drive technical QA strategy — "
        "covering stability, performance, reliability, and systemic risk — while embedding quality into planning, "
        "design, and delivery through collaboration and influence rather than authority alone. Experienced in "
        "coordinating internal and external QA contributors as one cohesive team, aligning priorities and standards "
        "across engineers, producers, and partners. Acts as a critical friend to development teams — supporting and "
        "challenging engineers to build effective testing approaches and integrate quality into their daily work. "
        "Comfortable operating both hands-on (investigating issues, validating assumptions) and strategically "
        "(shaping long-term QA approaches, managing technical debt vs. delivery constraints). Strong communicator "
        "who makes complex technical quality topics understandable and actionable."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Technical QA Ownership (E2E)", "QA Strategy & Test Planning", "Stability & Performance Testing",
        "Risk Assessment & Triaging", "Developer Enablement & Coaching", "Internal/External QA Coordination",
        "Leadership Through Influence", "Continuous Improvement", "Hands-On + Strategic Balance",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Technical QA Lead / Quality Owner — IoT & Cloud Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Own technical QA outcomes within the platform area — acting as primary point of accountability for stability, performance, reliability, and systemic risk across a complex, distributed system serving 30+ markets.",
        "Define and drive area-level technical QA strategy, test planning, and coverage across development, release, and live phases — including performance, stability, compatibility, and platform considerations.",
        "Lead technical quality through collaboration and influence — partnering with Tech Directors, engineers, and production to embed QA into design, implementation, and decision-making from early stages.",
        "Coordinate internal and external QA contributors as one cohesive team — ensuring alignment on technical priorities, standards, and ways of working across internal engineers and external partners.",
        "Act as a critical friend to development teams — supporting and challenging engineers in building effective testing approaches, integrating quality into their daily work, and driving improvement assertively.",
        "Drive visibility and decision-making through clear technical QA reporting, effective triaging, risk assessment, and tracking of systemic issues and trends — enabling leadership to act on quality data.",
        "Balance technical debt, delivery constraints, and user-facing risk — guiding where QA effort will have the highest impact and making pragmatic trade-off recommendations.",
        "Stay close to the product and systems through hands-on involvement — investigating issues, validating assumptions, and grounding strategic decisions in real outcomes.",
        "Champion developer enablement practices — building shared testing standards, reusable automation frameworks (Python/pytest), and quality feedback loops within CI/CD pipelines.",
        "Create structure and drive continuous improvement across teams and systems — evolving QA processes, metrics, dashboards, and reporting to make quality transparent and actionable.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "QA Lead / Technical Quality Owner", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Owned technical quality outcomes for a 300M+ user platform — taking end-to-end accountability for stability, performance, and release reliability.",
        "Defined QA strategy and test coverage across development and release cycles — balancing automated regression, performance validation, and manual exploratory testing.",
        "Led through influence — aligning engineering, product, and QA contributors around clear technical priorities and quality standards without direct authority.",
        "Drove technical QA reporting and risk assessment — making quality status, systemic issues, and release risks transparent and actionable for leadership.",
        "Operated both hands-on and strategically — investigating production issues while shaping testing approaches and continuous improvement initiatives.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior QA Lead / Technical Quality Owner — IoT & Enterprise Platforms", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Owned technical QA outcomes across complex, systems-heavy platforms (IKEA IoT/smart-home, LEGO digital) — accountable for stability, performance, reliability, and systemic risk across multiple releases.",
        "Defined and drove technical QA strategy, test planning, and coverage — spanning unit, integration, system, and E2E levels across development, release, and live operations.",
        "Coordinated internal QA team and external QA partners (co-dev, offshore) as one cohesive team — aligning on technical priorities, standards, and ways of working across locations.",
        "Acted as critical friend to development teams — giving constructive feedback to technical engineers, assertively driving improvement in testing practices and code quality.",
        "Led through influence across cross-functional teams (firmware, mobile, cloud, security, production) — embedding quality into design and implementation decisions.",
        "Drove visibility through technical QA reporting, triaging, and tracking of systemic issues — presenting risk assessments and release readiness recommendations to steering committees.",
        "Managed stability and performance testing for complex distributed systems — validating behavior under load, across platforms, and under adverse conditions (network, hardware faults).",
        "Balanced technical debt, delivery constraints, and user-facing risk — guiding QA investment toward highest-impact areas and making pragmatic scope recommendations.",
        "Built automation frameworks (Python, pytest, Selenium) and integrated them into CI/CD pipelines — championing developer enablement and shift-left testing practices.",
        "Drove continuous improvement — evolving QA processes, reducing flakiness, implementing regression strategies, and improving ways of working across teams and systems.",
        "Stayed hands-on when needed — investigating complex issues across distributed systems, validating assumptions, and debugging systemic problems to ground decisions in reality.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — Technical QA & Systems Testing", bold=True, size=Pt(10))
    add_text(p, "  |  India  |  2008 – 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Owned technical quality for enterprise systems in regulated financial environments — validating stability, reliability, and performance across complex multi-tier architectures.",
        "Coordinated QA activities across internal and external teams — aligning testing priorities, coverage, and standards across distributed contributors.",
        "Built testing automation and tooling (Python, shell) — enabling developer self-service testing and continuous quality feedback.",
        "Drove continuous improvement in QA processes — creating structure, improving ways of working, and establishing quality metrics and reporting.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("QA Strategy: ", "Test planning, coverage analysis, risk assessment, exit criteria, release readiness, QA metrics & dashboards"),
        ("Testing Types: ", "Performance, stability, reliability, compatibility, regression, system-level, E2E, exploratory"),
        ("Automation: ", "Python (pytest, requests, custom frameworks), Selenium, Playwright, shell scripting, API testing"),
        ("CI/CD: ", "GitHub Actions, Jenkins, GitLab CI — automated quality gates, pipeline integration, shift-left practices"),
        ("Platforms: ", "GCP, Azure, AWS, Linux, Docker, Kubernetes, distributed systems"),
        ("Tools: ", "Jira, Confluence, Grafana dashboards, monitoring/alerting, log analysis, Git"),
        ("Leadership: ", "Influence-based leadership, mentoring, cross-functional coordination, internal/external team alignment"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("Post Graduate Diploma in Operation and Management", "IGNOU, India"),
        ("B.Tech, Information Technology", "UP Technical University, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
        "ITIL v4 Foundation",
        "Six Sigma Green Belt",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent — written and spoken)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Tech Area QA Owner</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Technical QA leader working in IT since 2008, with deep experience owning quality outcomes end-to-end across complex, systems-heavy platforms. Proven ability to define and drive technical QA strategy &mdash; covering stability, performance, reliability, and systemic risk &mdash; while embedding quality into planning, design, and delivery through collaboration and influence rather than authority alone. Experienced in coordinating internal and external QA contributors as one cohesive team, aligning priorities and standards across engineers, producers, and partners. Acts as a critical friend to development teams &mdash; supporting and challenging engineers to build effective testing approaches and integrate quality into their daily work. Comfortable operating both hands-on (investigating issues, validating assumptions) and strategically (shaping long-term QA approaches, managing technical debt vs. delivery constraints). Strong communicator who makes complex technical quality topics understandable and actionable.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Technical QA Ownership (E2E)</td><td>&bull; QA Strategy &amp; Test Planning</td><td>&bull; Stability &amp; Performance Testing</td></tr>
<tr><td>&bull; Risk Assessment &amp; Triaging</td><td>&bull; Developer Enablement &amp; Coaching</td><td>&bull; Internal/External QA Coordination</td></tr>
<tr><td>&bull; Leadership Through Influence</td><td>&bull; Continuous Improvement</td><td>&bull; Hands-On + Strategic Balance</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Technical QA Lead / Quality Owner &mdash; IoT &amp; Cloud Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Own technical QA outcomes within the platform area &mdash; acting as primary point of accountability for stability, performance, reliability, and systemic risk across a complex, distributed system serving 30+ markets.</li>
<li>Define and drive area-level technical QA strategy, test planning, and coverage across development, release, and live phases &mdash; including performance, stability, compatibility, and platform considerations.</li>
<li>Lead technical quality through collaboration and influence &mdash; partnering with Tech Directors, engineers, and production to embed QA into design, implementation, and decision-making from early stages.</li>
<li>Coordinate internal and external QA contributors as one cohesive team &mdash; ensuring alignment on technical priorities, standards, and ways of working across internal engineers and external partners.</li>
<li>Act as a critical friend to development teams &mdash; supporting and challenging engineers in building effective testing approaches, integrating quality into their daily work, and driving improvement assertively.</li>
<li>Drive visibility and decision-making through clear technical QA reporting, effective triaging, risk assessment, and tracking of systemic issues and trends &mdash; enabling leadership to act on quality data.</li>
<li>Balance technical debt, delivery constraints, and user-facing risk &mdash; guiding where QA effort will have the highest impact and making pragmatic trade-off recommendations.</li>
<li>Stay close to the product and systems through hands-on involvement &mdash; investigating issues, validating assumptions, and grounding strategic decisions in real outcomes.</li>
<li>Champion developer enablement practices &mdash; building shared testing standards, reusable automation frameworks (Python/pytest), and quality feedback loops within CI/CD pipelines.</li>
<li>Create structure and drive continuous improvement across teams and systems &mdash; evolving QA processes, metrics, dashboards, and reporting to make quality transparent and actionable.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">QA Lead / Technical Quality Owner <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Owned technical quality outcomes for a 300M+ user platform &mdash; taking end-to-end accountability for stability, performance, and release reliability.</li>
<li>Defined QA strategy and test coverage across development and release cycles &mdash; balancing automated regression, performance validation, and manual exploratory testing.</li>
<li>Led through influence &mdash; aligning engineering, product, and QA contributors around clear technical priorities and quality standards without direct authority.</li>
<li>Drove technical QA reporting and risk assessment &mdash; making quality status, systemic issues, and release risks transparent and actionable for leadership.</li>
<li>Operated both hands-on and strategically &mdash; investigating production issues while shaping testing approaches and continuous improvement initiatives.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Senior QA Lead / Technical Quality Owner &mdash; IoT &amp; Enterprise Platforms <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Owned technical QA outcomes across complex, systems-heavy platforms (IKEA IoT/smart-home, LEGO digital) &mdash; accountable for stability, performance, reliability, and systemic risk across multiple releases.</li>
<li>Defined and drove technical QA strategy, test planning, and coverage &mdash; spanning unit, integration, system, and E2E levels across development, release, and live operations.</li>
<li>Coordinated internal QA team and external QA partners (co-dev, offshore) as one cohesive team &mdash; aligning on technical priorities, standards, and ways of working across locations.</li>
<li>Acted as critical friend to development teams &mdash; giving constructive feedback to technical engineers, assertively driving improvement in testing practices and code quality.</li>
<li>Led through influence across cross-functional teams (firmware, mobile, cloud, security, production) &mdash; embedding quality into design and implementation decisions.</li>
<li>Drove visibility through technical QA reporting, triaging, and tracking of systemic issues &mdash; presenting risk assessments and release readiness recommendations to steering committees.</li>
<li>Managed stability and performance testing for complex distributed systems &mdash; validating behavior under load, across platforms, and under adverse conditions (network, hardware faults).</li>
<li>Balanced technical debt, delivery constraints, and user-facing risk &mdash; guiding QA investment toward highest-impact areas and making pragmatic scope recommendations.</li>
<li>Built automation frameworks (Python, pytest, Selenium) and integrated them into CI/CD pipelines &mdash; championing developer enablement and shift-left testing practices.</li>
<li>Drove continuous improvement &mdash; evolving QA processes, reducing flakiness, implementing regression strategies, and improving ways of working across teams and systems.</li>
<li>Stayed hands-on when needed &mdash; investigating complex issues across distributed systems, validating assumptions, and debugging systemic problems to ground decisions in reality.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Technical QA &amp; Systems Testing</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;2008 &ndash; 2013</span></p>
<ul>
<li>Owned technical quality for enterprise systems in regulated financial environments &mdash; validating stability, reliability, and performance across complex multi-tier architectures.</li>
<li>Coordinated QA activities across internal and external teams &mdash; aligning testing priorities, coverage, and standards across distributed contributors.</li>
<li>Built testing automation and tooling (Python, shell) &mdash; enabling developer self-service testing and continuous quality feedback.</li>
<li>Drove continuous improvement in QA processes &mdash; creating structure, improving ways of working, and establishing quality metrics and reporting.</li>
</ul>

<h2>TECHNICAL SKILLS</h2>
<p class="tech-line"><b>QA Strategy:</b> Test planning, coverage analysis, risk assessment, exit criteria, release readiness, QA metrics &amp; dashboards</p>
<p class="tech-line"><b>Testing Types:</b> Performance, stability, reliability, compatibility, regression, system-level, E2E, exploratory</p>
<p class="tech-line"><b>Automation:</b> Python (pytest, requests, custom frameworks), Selenium, Playwright, shell scripting, API testing</p>
<p class="tech-line"><b>CI/CD:</b> GitHub Actions, Jenkins, GitLab CI &mdash; automated quality gates, pipeline integration, shift-left practices</p>
<p class="tech-line"><b>Platforms:</b> GCP, Azure, AWS, Linux, Docker, Kubernetes, distributed systems</p>
<p class="tech-line"><b>Tools:</b> Jira, Confluence, Grafana dashboards, monitoring/alerting, log analysis, Git</p>
<p class="tech-line"><b>Leadership:</b> Influence-based leadership, mentoring, cross-functional coordination, internal/external team alignment</p>

<h2>EDUCATION</h2>
<p><b>Post Graduate Diploma in Operation and Management</b> &mdash; IGNOU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; UP Technical University, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
<tr><td>&bull; ITIL v4 Foundation</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent &mdash; written and spoken)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
