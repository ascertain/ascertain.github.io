"""
Resume: eCommerce Testing & Optimisation Specialist — Pandora
Focus: A/B testing, experimentation, CRO, customer journey analysis, web analytics,
       global-local coordination, stakeholder education, data-driven optimisation.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_eCom_Testing_Pandora_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Data-driven ", False),
        ("eCommerce testing & optimisation specialist", True),
        (" with ", False),
        ("16+ years", True),
        (" in digital quality, ", False),
        ("experimentation", True),
        (", and ", False),
        ("customer journey analysis", True),
        (" across global platforms. Proven track record of driving ", False),
        ("A/B testing programmes", True),
        (", ", False),
        ("conversion rate optimisation (CRO)", True),
        (", and ", False),
        ("data-backed improvements", True),
        (" within large-scale ", False),
        ("eCommerce ecosystems", True),
        (" (IKEA, LEGO). Skilled at ", False),
        ("collaborating with UX, product, and eCommerce teams", True),
        (" to translate ", False),
        ("web analytics insights", True),
        (" into actionable recommendations. Experienced in ", False),
        ("global-local coordination", True),
        (" across 30+ markets, ", False),
        ("stakeholder education", True),
        (", and ", False),
        ("process standardisation", True),
        (". Passionate about uncovering ", False),
        ("what drives customer behaviour", True),
        (" and building a ", False),
        ("culture of continuous improvement", True),
        (" through experimentation.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "A/B Testing & Experimentation", "Conversion Rate Optimisation (CRO)", "Customer Journey Analysis",
        "Web Analytics & Reporting", "Global-Local Test Coordination", "Stakeholder Education & Influence",
        "eCommerce Performance Insights", "UX & Product Collaboration", "Process Standardisation",
        "Data-Driven Decision Making", "Cross-Market Knowledge Sharing", "Continuous Improvement Culture",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills_data = [
        ("Analytics & Insights:", " Google Analytics, Adobe Analytics (exposure), Contentsquare (exposure), Monetate (exposure), BigQuery, Looker Studio, custom dashboards"),
        ("Experimentation:", " A/B test design & analysis, multivariate testing, feature flagging, hypothesis formulation, statistical significance evaluation"),
        ("eCommerce Platforms:", " IKEA.com ecosystem, Sitecore (LEGO multi-site), Pimcore (Campaign Factory), Salesforce Commerce, CMS-driven storefronts"),
        ("Web Technologies:", " HTML, CSS, JavaScript, REST APIs, responsive web testing, cross-browser/device validation"),
        ("Automation & CI/CD:", " Python, Playwright, Selenium, GitHub Actions, Jenkins, Docker, automated regression pipelines"),
        ("Collaboration:", " JIRA, Confluence, Miro, Figma (review), Slack, stakeholder workshops, cross-functional facilitation"),
        ("Practices:", " Agile/Scrum, Kanban, data storytelling, post-test reporting, metric definition, funnel analysis"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA / Ingka Digital ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 eCommerce Testing & Optimisation Lead (Acting)",
        "Mar 2022 \u2013 Present",
        "Customer-Facing Digital Platforms \u2014 30+ Markets | eCommerce, UX, Data-Driven Experimentation")

    add_bullet(doc,
        "Drive the global-local testing roadmap for IKEA\u2019s digital customer experience platforms, aligning experimentation initiatives with both regional market priorities and global eCommerce strategies.",
        ["global-local testing roadmap", "experimentation initiatives", "regional market priorities", "global eCommerce strategies"])

    add_bullet(doc,
        "Design and execute A/B tests and experimentation programmes to optimise customer journeys \u2014 hypothesis formulation, test setup, statistical analysis, and post-test reporting with clear actionable insights.",
        ["A/B tests", "experimentation programmes", "customer journeys", "hypothesis formulation", "statistical analysis", "post-test reporting", "actionable insights"])

    add_bullet(doc,
        "Conduct in-depth analysis of customer journeys and feature performance using analytics tools (BigQuery, Looker Studio, custom dashboards) \u2014 identifying optimisation opportunities and translating data into recommendations for product and UX teams.",
        ["in-depth analysis of customer journeys", "analytics tools", "optimisation opportunities", "recommendations for product and UX teams"])

    add_bullet(doc,
        "Collaborate closely with UX, product, and eCommerce teams to ensure site improvements are rooted in user insights and testing outcomes; help prioritise impactful changes based on data.",
        ["UX, product, and eCommerce teams", "user insights and testing outcomes", "prioritise impactful changes"])

    add_bullet(doc,
        "Educate and influence stakeholders across the organisation on A/B testing best practices, the value of experimentation, and data-driven decision making \u2014 building a culture of continuous improvement.",
        ["Educate and influence stakeholders", "A/B testing best practices", "experimentation", "data-driven decision making", "culture of continuous improvement"])

    add_bullet(doc,
        "Drive process standardisation for testing execution across markets \u2014 governance frameworks, consistent methodologies, and knowledge-sharing rituals ensuring efficient and repeatable experimentation.",
        ["process standardisation", "governance frameworks", "consistent methodologies", "knowledge-sharing", "efficient and repeatable experimentation"])

    add_bullet(doc,
        "Share best practices and testing insights across 30+ markets, facilitating cross-market learning and fostering a test-and-learn mindset organisation-wide.",
        ["best practices", "testing insights across 30+ markets", "cross-market learning", "test-and-learn mindset"])

    add_bullet(doc,
        "Define and track key metrics (conversion rate, task completion, customer satisfaction) to measure the impact of experiments and digital experience changes.",
        ["key metrics", "conversion rate", "task completion", "customer satisfaction", "impact of experiments"])

    add_bullet(doc,
        "Lead and mentor a team of 10+ engineers; orchestrate testing workstreams across a complex global structure while remaining hands-on with analysis and reporting.",
        ["Lead and mentor", "10+", "testing workstreams", "complex global structure", "hands-on with analysis"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Experimentation Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Data-Driven Release Quality & Feature Optimisation")

    add_bullet(doc,
        "Supported feature experimentation and release validation for a 300M+ user platform \u2014 analysing feature performance data to inform go/no-go decisions and optimise user experience.",
        ["feature experimentation", "release validation", "analysing feature performance data", "optimise user experience"])

    add_bullet(doc,
        "Integrated automated quality checks into CI/CD pipelines; drove data-backed improvements to release processes and testing practices.",
        ["automated quality checks", "CI/CD pipelines", "data-backed improvements", "testing practices"])

    # --- HCLTech / LEGO ---
    role_header(doc,
        "HCLTech \u2014 LEGO Group, Denmark \u2014 eCommerce QA & Optimisation Lead",
        "2017 \u2013 2021",
        "eCommerce & Digital Campaign Platforms \u2014 Sitecore, Pimcore, Multi-Market | Web Performance & CRO")

    add_bullet(doc,
        "Led testing and optimisation for LEGO\u2019s global eCommerce and digital campaign platforms (Sitecore multi-site: DUPLO, Friends, Minecraft, Ninjago; Campaign Factory on Pimcore) \u2014 ensuring seamless customer experiences across markets.",
        ["testing and optimisation", "global eCommerce", "digital campaign platforms", "Sitecore", "Pimcore", "customer experiences across markets"])

    add_bullet(doc,
        "Analysed eCommerce performance data to identify conversion bottlenecks and optimisation opportunities; provided actionable insights to product and UX stakeholders.",
        ["eCommerce performance data", "conversion bottlenecks", "optimisation opportunities", "actionable insights", "product and UX stakeholders"])

    add_bullet(doc,
        "Drove A/B testing and experimentation for website features and campaign pages \u2014 test design, execution, analysis, and post-test reporting to guide iterative improvements.",
        ["A/B testing and experimentation", "test design, execution, analysis", "post-test reporting", "iterative improvements"])

    add_bullet(doc,
        "Built and maintained automated regression frameworks (Selenium, Python, Jenkins, Amazon EC2) supporting fast release cycles and consistent quality across markets.",
        ["automated regression frameworks", "fast release cycles", "consistent quality across markets"])

    add_bullet(doc,
        "Collaborated cross-functionally with UX, product, and development teams; educated stakeholders on testing value and best practices.",
        ["cross-functionally", "UX, product, and development teams", "educated stakeholders", "best practices"])

    add_bullet(doc,
        "Managed and mentored a team of 8\u201312 engineers; coordinated testing workstreams across global and local priorities.",
        ["8\u201312 engineers", "testing workstreams", "global and local priorities"])

    # --- HCLTech / IKEA ---
    role_header(doc,
        "HCLTech \u2014 IKEA, Sweden \u2014 Digital QA Specialist",
        "2013 \u2013 2017",
        "IKEA Digital Ecosystem \u2014 eCommerce, Mobile, Customer-Facing Platforms")

    add_bullet(doc,
        "Tested and validated customer-facing digital experiences across IKEA\u2019s eCommerce ecosystem \u2014 web, mobile, cross-browser, and multi-market rollouts.",
        ["customer-facing digital experiences", "eCommerce ecosystem", "cross-browser", "multi-market rollouts"])

    add_bullet(doc,
        "Performed customer journey analysis and usability validation; worked closely with UX and product teams to ensure improvements were data-informed.",
        ["customer journey analysis", "usability validation", "UX and product teams", "data-informed"])

    add_bullet(doc,
        "Developed automated test suites (Python, Selenium) for regression and feature validation; integrated tests into CI/CD pipelines for rapid feedback.",
        ["automated test suites", "regression and feature validation", "CI/CD pipelines", "rapid feedback"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 QA Analyst / Consultant",
        "2008 \u2013 2013",
        "Core Banking & Enterprise Applications (HCL, Marlabs, TekMindz, India)")

    add_bullet(doc,
        "Functional and regression testing for Finacle Core Banking and enterprise applications; experience with complex data analysis and quality assurance in regulated environments.",
        ["Functional and regression testing", "complex data analysis", "quality assurance", "regulated environments"])

    add_bullet(doc,
        "Built early automation frameworks (Python, Java); gained foundational skills in analytical thinking, data-driven quality assessment, and process improvement.",
        ["automation frameworks", "analytical thinking", "data-driven quality assessment", "process improvement"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "Six Sigma Green Belt", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
