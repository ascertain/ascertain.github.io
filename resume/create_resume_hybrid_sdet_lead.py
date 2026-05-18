from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_SDET_Lead_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_SDET_Lead_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "SDET",
    "automation",
    "Playwright",
    "Selenium",
    "CI/CD",
    "DevOps",
    "leadership",
    "mentor",
    "coaching",
    "CoE",
    "Center of Excellence",
    "stakeholder",
    "vendor",
    "RFP",
    "RFI",
    "framework",
    "testing strategy",
    "quality",
    "security",
    "CEH",
    "cloud",
    "GCP",
    "AWS",
    "Terraform",
    "Docker",
    "Kubernetes",
    "Python",
    "TypeScript",
    "C#",
    "Agile",
    "SAFe",
    "Six Sigma",
    "ISTQB",
    "AI",
    "data pipeline",
    "Grafana",
    "project delivery",
    "budget",
    "risk",
    "team",
    "cross-functional",
    "line management",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(
                paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW
            )
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "SDET Lead / Technical Lead  —  Automation, DevOps, Project Delivery & Team Leadership  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(
        contact,
        "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com",
        size=9.5,
        color=TEXT_MUTED,
    )
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Hands-on SDET Lead and Technical Lead with 15+ years combining deep automation "
        "engineering with project delivery and people leadership. Codes, architects, and "
        "delivers — owns end-to-end quality across the full SDLC. Builds and scales automation "
        "frameworks (Playwright, Selenium, Cypress, Appium) and CI/CD pipelines. Leads and "
        "mentors engineering teams (8–15 engineers), coaches competence development, and drives "
        "hiring. Establishes Center of Excellence (CoE) practices — defines testing strategy, "
        "automation guidelines, quality gates, and reusable frameworks. Manages project scope, "
        "schedule, budget, and risk. Connects with business stakeholders, manages vendor "
        "relationships, and leads RFP/RFI technical responses. Designs solutions and frameworks "
        "for complex platforms. Security-aware — Certified Ethical Hacker (CEH). Cloud-native "
        "delivery on GCP and AWS (Terraform, Docker, Kubernetes, CI/CD). Shipped products to "
        "300M+ users at IKEA, Truecaller, and LEGO. Line management experience.",
        size=10,
    )

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    comp_lines = [
        (
            "Automation & SDET: ",
            "Playwright, Selenium, Cypress, Appium. API testing (REST, GraphQL). "
            "E2E, integration, regression, performance. TDD/BDD. Framework design "
            "and architecture. CI/CD integration. AI-assisted test generation",
        ),
        (
            "Project & Delivery Management: ",
            "Scope, schedule, budget, risk management. Roadmap & OKRs. Backlog "
            "prioritization. Multi-partner coordination. Resource allocation. "
            "Agile (Scrum/SAFe), Waterfall. Audit-ready reporting",
        ),
        (
            "Leadership & CoE: ",
            "Line management (8–15 engineers). Coaching, mentoring, competence "
            "development. Hiring & talent retention. Center of Excellence — testing "
            "strategy, automation guidelines, quality gates, reusable frameworks. "
            "Knowledge sharing & evangelization",
        ),
        (
            "Stakeholder, Vendor & Business: ",
            "Business stakeholder alignment. Vendor management & managed service "
            "relationships. RFP/RFI technical input & POC delivery. Solution "
            "design. Cross-functional collaboration (engineering, product, business)",
        ),
        (
            "DevOps & Cloud: ",
            "GCP (Cloud Run, BigQuery, Cloud Functions), AWS. Terraform, Docker, "
            "Kubernetes. GitHub Actions, Jenkins. Infrastructure as Code. "
            "Grafana observability. Data pipelines",
        ),
        (
            "Security & Compliance: ",
            "CEH (Certified Ethical Hacker). Security testing. Regulatory "
            "compliance awareness. Quality gate management. Six Sigma Green Belt. "
            "ISTQB. ITIL Foundation",
        ),
        (
            "Programming: ",
            "Python, TypeScript, C#, Java, Node.js. REST APIs. SQL. Git. "
            "Jira, Confluence, TestRail, MS Project",
        ),
    ]
    for label, value in comp_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(
        cp1,
        "IKEA IT AB, Malmö — Team Lead / SDET Lead",
        bold=True,
        size=10,
        color=SECTION_COLOR,
    )
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(
        tp1,
        "Mar 2022 – Present  |  VCS Platform — Omni-Channel, 30+ Global Markets",
        bold=True,
        size=10,
    )
    ikea_bullets = [
        "End-to-end ownership of IKEA's Visual Collaboration Solution (VCS) — scaled "
        "RCMP from 2K to 50K usage across 30+ markets, turning the RCMP sales business "
        "green. Full solution ownership: automation, DevOps, quality, and delivery.",
        "Line management — responsible for engineering team plus consultants. Coaching, "
        "competence development, continuous follow-up. Hiring and talent retention. "
        "Mentoring engineers on automation best practices, code quality, and career growth.",
        "CoE & testing strategy — established automation guidelines, quality gates, and "
        "reusable frameworks. Defined testing strategy and acceptance criteria. Owned "
        "roadmap, OKRs, and backlog prioritization for benefit realization.",
        "Business stakeholder & vendor management — bridged IKEA business leadership and "
        "SaaS platform vendor. Managed service relationship, aligned delivery with business "
        "KPIs. Cross-functional collaboration with POs, PMs, and stakeholders.",
        "Automation & DevOps — Playwright, API testing, CI/CD (GitHub Actions). Terraform, "
        "Docker on GCP. Data pipelines (BigQuery, Cloud Functions). Grafana observability. "
        "AI-assisted — 30% velocity improvement. Led Selenium-to-Playwright migration "
        "(3x faster execution, 50% CI pipeline reduction).",
        "Budget & resource management — planning, forecasting, follow-up. Capacity "
        "planning. Prioritization. Risk identification and mitigation. Audit-ready reporting.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(
        cp2,
        "Truecaller, Stockholm — Release & Automation Engineer",
        bold=True,
        size=10,
        color=SECTION_COLOR,
    )
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(
        tp2,
        "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users",
        bold=True,
        size=10,
    )
    tc_bullets = [
        "Project leadership at global scale — owned scope, schedule, risk, and release "
        "readiness for 300M+ users. Go/no-go decisions. Stakeholder interface. Release "
        "automation — built tools and workflows to streamline release processes. CI pipelines, "
        "build systems, deployment automation. AWS cloud infrastructure.",
        "Cross-functional coordination — backlog prioritization, supplier coordination, "
        "resource allocation. Feature flag management and data-driven rollouts across "
        "global markets. Onboarded engineers. Agile (Scrum/Kanban).",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(
        cp3,
        "HCLTech — LEGO & IKEA Group, Denmark & Sweden — Technical Specialist / SDET Lead",
        bold=True,
        size=10,
        color=SECTION_COLOR,
    )
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(
        tp3,
        "2013 – 2021  |  E-Commerce, Mobile & Enterprise — Multi-Partner Delivery",
        bold=True,
        size=10,
    )
    lego_bullets = [
        "LEGO & IKEA (2017–2021): Led test scope, timelines, risk, and delivery across "
        "iOS, Android, and web platforms. Multi-partner coordination. Designed and executed "
        "automation strategy across mobile and web. Built test automation for beta OS releases. "
        "Release readiness, quality gate decisions. Onshore-offshore team management.",
        "SDET Lead (2013–2017): Full SDLC ownership — requirements through production "
        "monitoring. Designed and scaled Selenium frameworks (C#, NUnit). CI execution. "
        "RFP/RFI technical input and POC delivery — built working automation POCs to "
        "demonstrate feasibility and win new business. Sprint and quality reporting for "
        "leadership visibility.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(
        cp4,
        "Banking & Enterprise — SDET / Consultant",
        bold=True,
        size=10,
        color=SECTION_COLOR,
    )
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(
        tp4,
        "2008 – 2013  |  Finacle CBS, Core Banking — Regulated Environments",
        bold=True,
        size=10,
    )
    fin_bullets = [
        "Led testing workstreams for core banking systems (Finacle CBS, biometric "
        "authentication). Test automation — automated scripts from functional specifications. "
        "ETL automation with Pentaho. Data validation (SQL, UNIX). Regulatory compliance "
        "and banking security standards. Client interface and defect coordination.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Scaled IKEA RCMP from 2K to 50K usage across 30+ markets — end-to-end solution "
        "ownership, directly turning the RCMP sales business green. Recognized as "
        "Exceptional Performer.",
        "Led Selenium-to-Playwright migration — 3x faster test execution, 50% CI pipeline "
        "reduction, improved developer adoption across multiple teams.",
        "CoE & automation evangelization — established testing strategy, automation "
        "guidelines, quality gates, and reusable frameworks. Mentored engineers, drove "
        "competence development, and built high-performing teams.",
        "RFP/RFI & solution design — led technical input for RFP/RFI responses. Built "
        "working automation POCs to demonstrate feasibility and win new engagements.",
        "Business–vendor bridge — managed SaaS vendor relationship at IKEA, aligning "
        "vendor roadmap with IKEA's omni-channel growth strategy and business KPIs.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Personal Project ──
    add_section_heading(document, "Personal Project")
    pp = document.add_paragraph()
    pp.paragraph_format.space_after = Pt(0)
    pp.paragraph_format.line_spacing = 1.0
    add_text(pp, "SimpleRx — AI-Powered Health Report Assistant  ", bold=True, size=10)
    add_text(
        pp,
        "(Google Play)  |  Co-Founder & Developer, DQODIFY Pvt. Ltd.",
        size=10,
        color=TEXT_MUTED,
    )
    pp2 = document.add_paragraph()
    pp2.paragraph_format.space_after = Pt(0)
    pp2.paragraph_format.line_spacing = 1.0
    add_text(
        pp2,
        "AI app translating medical reports into plain language (Hindi & English). "
        "End-to-end ownership: product vision, AI/LLM integration, mobile development, "
        "Play Store deployment. Privacy-first — encrypted storage, user-controlled data deletion.",
        size=10,
    )

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(
        table.rows[0].cells, ["Education", "Certifications"]
    ):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(
        edu,
        "B.Tech in Information Technology\nPG Diploma in Operations Management",
        size=9.5,
    )
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Certified Ethical Hacker (CEH)\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Six Sigma Green Belt\n"
        "ITIL Foundation\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – SDET Lead / Technical Lead Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>SDET Lead / Technical Lead — Automation, DevOps, Project Delivery &amp; Team Leadership | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Hands-on <span class="hl">SDET</span> Lead and Technical Lead with 15+ years combining deep <span class="hl">automation</span> engineering with <span class="hl">project delivery</span> and people <span class="hl">leadership</span>. Codes, architects, and delivers — owns end-to-end quality across the full SDLC. Builds and scales <span class="hl">automation</span> <span class="hl">framework</span>s (<span class="hl">Playwright</span>, <span class="hl">Selenium</span>, Cypress, Appium) and <span class="hl">CI/CD</span> pipelines. Leads and <span class="hl">mentor</span>s engineering <span class="hl">team</span>s (8–15 engineers), coaches competence development, and drives hiring. Establishes <span class="hl">Center of Excellence</span> (<span class="hl">CoE</span>) practices — defines <span class="hl">testing strategy</span>, <span class="hl">automation</span> guidelines, <span class="hl">quality</span> gates, and reusable <span class="hl">framework</span>s. Manages project scope, schedule, <span class="hl">budget</span>, and <span class="hl">risk</span>. Connects with business <span class="hl">stakeholder</span>s, manages <span class="hl">vendor</span> relationships, and leads <span class="hl">RFP</span>/<span class="hl">RFI</span> technical responses. Designs solutions and <span class="hl">framework</span>s. <span class="hl">Security</span>-aware — <span class="hl">CEH</span>. <span class="hl">Cloud</span>-native (<span class="hl">GCP</span>, <span class="hl">AWS</span>). Products shipped to 300M+ users at IKEA, Truecaller, LEGO. <span class="hl">Line management</span> experience.</p>

  <div class="section">Core Competencies</div>
  <p><b>Automation &amp; SDET:</b> Playwright, Selenium, Cypress, Appium. API testing (REST, GraphQL). E2E, integration, regression, performance. TDD/BDD. Framework design. CI/CD integration. AI-assisted test generation<br>
  <b>Project &amp; Delivery:</b> Scope, schedule, budget, risk. Roadmap &amp; OKRs. Backlog prioritization. Multi-partner. Resource allocation. Agile (Scrum/SAFe), Waterfall. Audit-ready reporting<br>
  <b>Leadership &amp; CoE:</b> Line management (8–15 engineers). Coaching, mentoring, competence development. Hiring. Center of Excellence — testing strategy, guidelines, quality gates, reusable frameworks<br>
  <b>Stakeholder, Vendor &amp; Business:</b> Business stakeholder alignment. Vendor management. RFP/RFI &amp; POC delivery. Solution design. Cross-functional collaboration<br>
  <b>DevOps &amp; Cloud:</b> GCP (Cloud Run, BigQuery), AWS. Terraform, Docker, Kubernetes. GitHub Actions, Jenkins. IaC. Grafana. Data pipelines<br>
  <b>Security &amp; Compliance:</b> CEH, security testing, regulatory compliance. Quality gates. Six Sigma GB. ISTQB. ITIL<br>
  <b>Programming:</b> Python, TypeScript, C#, Java, Node.js. REST APIs. SQL. Git. Jira, Confluence, TestRail</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Team Lead / SDET Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS Platform — Omni-Channel, 30+ Global Markets</div>
  <ul>
    <li>End-to-end ownership of IKEA's VCS — scaled RCMP from 2K to 50K across 30+ markets, turning RCMP sales business green. Full solution ownership: <span class="hl">automation</span>, <span class="hl">DevOps</span>, <span class="hl">quality</span>, and delivery.</li>
    <li><span class="hl">Line management</span> — engineering <span class="hl">team</span> + consultants. <span class="hl">Coaching</span>, competence development, hiring, talent retention. <span class="hl">Mentor</span>ing on <span class="hl">automation</span> best practices, code <span class="hl">quality</span>, and career growth.</li>
    <li><span class="hl">CoE</span> &amp; <span class="hl">testing strategy</span> — established <span class="hl">automation</span> guidelines, <span class="hl">quality</span> gates, reusable <span class="hl">framework</span>s. Owned roadmap, OKRs, backlog prioritization.</li>
    <li>Business <span class="hl">stakeholder</span> &amp; <span class="hl">vendor</span> management — bridged IKEA business <span class="hl">leadership</span> and SaaS <span class="hl">vendor</span>. Aligned delivery with business KPIs. <span class="hl">Cross-functional</span> collaboration.</li>
    <li><span class="hl">Automation</span> &amp; <span class="hl">DevOps</span> — <span class="hl">Playwright</span>, API testing, <span class="hl">CI/CD</span> (GitHub Actions). <span class="hl">Terraform</span>, <span class="hl">Docker</span> on <span class="hl">GCP</span>. Data pipelines (BigQuery). <span class="hl">Grafana</span>. <span class="hl">AI</span>-assisted — 30% velocity improvement. <span class="hl">Selenium</span>-to-<span class="hl">Playwright</span> migration (3x faster, 50% CI reduction).</li>
    <li><span class="hl">Budget</span> &amp; resource management — planning, forecasting, capacity, <span class="hl">risk</span>. Audit-ready reporting.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Release &amp; Automation Engineer</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users</div>
  <ul>
    <li><span class="hl">Project</span> <span class="hl">leadership</span> — scope, schedule, <span class="hl">risk</span>, release readiness for 300M+ users. Go/no-go decisions. <span class="hl">Stakeholder</span> interface. Release <span class="hl">automation</span> — tools, workflows, <span class="hl">CI/CD</span> pipelines. <span class="hl">AWS</span>.</li>
    <li><span class="hl">Cross-functional</span> — backlog prioritization, supplier coordination, resource allocation. Feature flag management. Onboarded engineers. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">HCLTech — LEGO &amp; IKEA Group, Denmark &amp; Sweden — Technical Specialist / SDET Lead</div>
  <div class="job-sub">2013 – 2021 | E-Commerce, Mobile &amp; Enterprise — Multi-Partner Delivery</div>
  <ul>
    <li>LEGO &amp; IKEA (2017–21): Led test scope, timelines, <span class="hl">risk</span>, delivery across iOS, Android, web. Multi-partner. <span class="hl">Automation</span> strategy (mobile &amp; web). <span class="hl">Quality</span> gates. Onshore-offshore <span class="hl">team</span> management.</li>
    <li><span class="hl">SDET</span> Lead (2013–17): Full SDLC. <span class="hl">Selenium</span> <span class="hl">framework</span>s (<span class="hl">C#</span>, NUnit). <span class="hl">CI/CD</span>. <span class="hl">RFP</span>/<span class="hl">RFI</span> technical input &amp; POC delivery — built <span class="hl">automation</span> POCs to win new business.</li>
  </ul>

  <div class="job-title">Banking &amp; Enterprise — SDET / Consultant</div>
  <div class="job-sub">2008 – 2013 | Finacle CBS, Core Banking — Regulated Environments</div>
  <ul>
    <li>Core banking systems (Finacle CBS, biometric auth). Test <span class="hl">automation</span>. ETL <span class="hl">automation</span> (Pentaho). Data validation (SQL, UNIX). Regulatory compliance and banking <span class="hl">security</span>. Client interface.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Scaled IKEA RCMP 2K → 50K across 30+ markets — end-to-end ownership. Exceptional Performer recognition.</li>
    <li><span class="hl">Selenium</span>-to-<span class="hl">Playwright</span> migration — 3x faster execution, 50% CI reduction, improved adoption.</li>
    <li><span class="hl">CoE</span> &amp; <span class="hl">automation</span> evangelization — <span class="hl">testing strategy</span>, guidelines, <span class="hl">quality</span> gates, reusable <span class="hl">framework</span>s. <span class="hl">Mentor</span>ed engineers.</li>
    <li><span class="hl">RFP</span>/<span class="hl">RFI</span> &amp; solution design — led technical responses. Built <span class="hl">automation</span> POCs to win engagements.</li>
    <li>Business–<span class="hl">vendor</span> bridge — managed SaaS <span class="hl">vendor</span> relationship, aligned roadmap with business KPIs.</li>
  </ul>

  <div class="section">Personal Project</div>
  <p><b>SimpleRx — AI-Powered Health Report Assistant</b> <span style="color:#4b5563">(Google Play) | Co-Founder &amp; Developer, DQODIFY Pvt. Ltd.</span><br>
  AI app translating medical reports into plain language (Hindi &amp; English). End-to-end ownership: product vision, AI/LLM integration, mobile dev, Play Store deployment. Privacy-first — encrypted storage, user-controlled deletion.</p>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>Certified Ethical Hacker (CEH)<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Six Sigma Green Belt<br>ITIL Foundation<br>UiPath RPA Certified</td>
    </tr>
  </table>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
