"""
Resume: Engineering Manager — SaaS/B2B, Team Leadership, Cloud Architecture
Focus: People leadership, coaching, production reliability, ambiguity navigation, cross-functional alignment.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Engineering_Manager_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Engineering leader with ", False),
        ("16+ years", True),
        (" in software development, now focused on ", False),
        ("growing high-performing teams", True),
        (" through ", False),
        ("coaching, mentoring, and setting clear direction", True),
        (". Experienced in forming the ", False),
        ("Product-Design-Engineering trio", True),
        (" that enables autonomous, empowered teams. Strong ", False),
        ("technical background", True),
        (" with hands-on capability — able to contribute to engineering work, ", False),
        ("architectural decisions", True),
        (", and ", False),
        ("AI-assisted development", True),
        (" practices. Proven track record ensuring ", False),
        ("predictable, sustainable delivery", True),
        (", maintaining ", False),
        ("engineering and security standards", True),
        (", and cultivating a culture of ", False),
        ("psychological safety", True),
        (", ", False),
        ("inclusion", True),
        (", and ", False),
        ("continuous improvement", True),
        (". Experienced in ", False),
        ("SaaS/B2B platforms", True),
        (" with cloud-native, event-driven architecture.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── LEADERSHIP COMPETENCIES ──
    add_section_heading(doc, "Leadership Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Coaching & Mentoring", "Hiring & Retaining Talent", "Psychological Safety & Inclusion",
        "Predictable Delivery Pace", "AI-Assisted Development", "Delivery Forecasting & Planning",
        "Product-Design-EM Trio Model", "Engineering & Security Standards", "Continuous Improvement",
        "Hands-On Technical Leadership", "Cross-Team Collaboration", "Roadmap & Capacity Planning",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL BACKGROUND ──
    add_section_heading(doc, "Technical Background")
    skills_data = [
        ("Architecture:", " Cloud-native (GCP, AWS), event-driven systems, microservices, SaaS platforms"),
        ("Languages:", " Python, TypeScript/JavaScript, Java, SQL, Bash"),
        ("DevOps & CI/CD:", " GitHub Actions, Jenkins, Azure DevOps, Docker, Kubernetes, GitOps"),
        ("Platforms:", " GCP (Cloud Run, GKE, Pub/Sub, BigQuery), AWS, Terraform"),
        ("Communication:", " REST API, WebSocket, SIP/RTP (exposure), real-time messaging"),
        ("AI & Tooling:", " AI-assisted development workflows, Copilot, automated quality tooling"),
        ("Practices:", " Agile/Scrum, Kanban, OKRs, incident management, SLOs/SLIs"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Engineering Manager / Team Lead (Acting)",
        "Mar 2022 \u2013 Present",
        "SaaS Customer Connect Platform \u2014 Omni-Channel, 30+ Markets | Cloud-Native, Event-Driven Architecture")
    add_bullet(doc, "Lead, develop and grow a high-performing engineering team through coaching, mentoring, and setting clear direction \u2014 connecting daily work to strategic product objectives.",
        ["Lead, develop and grow", "coaching, mentoring", "clear direction", "strategic product objectives"])
    add_bullet(doc, "Support each team member to identify, grow and develop engineering capabilities and career through regular 1:1s, feedback, and personalised development plans.",
        ["grow and develop engineering capabilities", "career", "development plans"])
    add_bullet(doc, "Form the Product-Design-Engineering trio \u2014 working closely with product manager and designer to enable an autonomous and empowered team.",
        ["Product-Design-Engineering trio", "autonomous and empowered team"])
    add_bullet(doc, "Create a culture of psychological safety and inclusion \u2014 monitor team health, empower members to raise ideas and concerns, drive improvements in how the team works together.",
        ["psychological safety and inclusion", "team health", "raise ideas and concerns"])
    add_bullet(doc, "Ensure delivery pipeline flows at a predictable and sustainable pace \u2014 delivery forecasting, capacity planning, informed estimates, and associated reporting.",
        ["predictable and sustainable pace", "delivery forecasting", "capacity planning"])
    add_bullet(doc, "Drive continuous improvement: review, develop and implement best-practice ways of working; remove blockers proactively; ensure operational hygiene.",
        ["continuous improvement", "best-practice ways of working", "remove blockers"])
    add_bullet(doc, "Actively explore and adopt AI-assisted development tools \u2014 enabling the team to work smarter, reduce repetitive effort, and push boundaries of delivery.",
        ["AI-assisted development tools", "work smarter", "push boundaries"])
    add_bullet(doc, "Advocate and enforce engineering and security standards; facilitate compliant engineering approach fostering sustainably excellent engineering.",
        ["engineering and security standards", "sustainably excellent engineering"])
    add_bullet(doc, "Attract, recruit and retain capable individuals \u2014 building on and complementing team diversity; contribute to development roadmaps and plans.",
        ["recruit and retain", "team diversity", "development roadmaps"])
    add_bullet(doc, "Collaborate with peer Engineering Managers to shape broader engineering strategy, culture, and cross-team knowledge sharing.",
        ["Engineering Managers", "engineering strategy", "cross-team knowledge sharing"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Automation Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "B2B/B2C Communication Platform \u2014 300M+ Users | Telephony & Real-Time Messaging")
    add_bullet(doc, "Owned release pipelines for production systems where reliability and quality directly impacted 300M+ users \u2014 operational mindset in action.",
        ["production systems", "reliability and quality", "operational mindset"])
    add_bullet(doc, "Collaborated cross-functionally with engineers, product, and ops to ensure smooth deployments; structured problem-solving for release blockers.",
        ["cross-functionally", "structured problem-solving"])
    add_bullet(doc, "Exposure to telephony-related infrastructure (SIP/RTP, real-time communication) within a high-scale SaaS platform.",
        ["telephony-related", "SIP/RTP", "SaaS platform"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 Technical Lead / SDET Lead",
        "2013 \u2013 2021",
        "SaaS/Enterprise E-Commerce & Digital Platforms \u2014 Multi-Team Delivery, Distributed Engineering")
    add_bullet(doc, "Led and mentored distributed engineering teams (8\u201312 people) through periods of growth, reorganisation, and programme changes \u2014 supporting individual career development.",
        ["Led and mentored", "growth, reorganisation", "career development"])
    add_bullet(doc, "Established ways of working, rituals, and decision-making frameworks that helped teams deliver at a predictable, sustainable pace.",
        ["ways of working", "predictable, sustainable pace"])
    add_bullet(doc, "Drove continuous improvement: reviewed processes, implemented best practices, removed unnecessary overhead \u2014 engineering-first transformation.",
        ["continuous improvement", "best practices", "engineering-first transformation"])
    add_bullet(doc, "Operated production-grade SaaS platforms where reliability, quality, engineering standards, and long-term maintainability were non-negotiable.",
        ["production-grade SaaS", "reliability", "engineering standards"])
    add_bullet(doc, "Worked with product and business stakeholders to align on roadmaps, capacity planning, and delivery forecasting across programmes.",
        ["product and business stakeholders", "roadmaps", "capacity planning", "delivery forecasting"])
    add_bullet(doc, "Hands-on technical contribution: architectural decisions, CI/CD pipelines, cloud infrastructure, troubleshooting production issues.",
        ["Hands-on technical", "architectural decisions", "troubleshooting production issues"])

    # --- India ---
    role_header(doc,
        "Banking & Enterprise \u2014 Senior Engineer / Consultant",
        "2008 \u2013 2013",
        "Core Banking Systems \u2014 Regulated B2B Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "Developed and maintained production banking software (Finacle CBS) where reliability, compliance, and quality were critical.",
        ["production banking software", "reliability, compliance, and quality"])
    add_bullet(doc, "Early leadership experience: coordinated testing across teams, mentored junior engineers, established team practices.",
        ["leadership experience", "coordinated", "mentored", "established team practices"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "CEH", "Six Sigma Green Belt"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
