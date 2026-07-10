"""Generate QA Manager resume – Inter IKEA Internal Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_QA_Manager_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Quality Assurance Manager with 12+ years of experience in software quality assurance, "
        "requirement management, verification & validation, and leading QA teams in digital and software environments. "
        "Deep IKEA domain expertise with proven ability to set quality standards, drive competence development, "
        "and ensure high-quality outcomes across digital products and platforms. "
        "Experienced in information security and data privacy practices, test automation, and balancing "
        "business needs with quality standards and delivery outcomes. "
        "Strong people leader — coaching, mentoring, and developing QA engineers while driving collaboration "
        "across product, platform, and technology areas. Passionate about fact-based decision-making, "
        "continuous improvement, and embedding quality into every stage of the development lifecycle. "
        "ISTQB certified with experience working in large, complex organisations with shared standards and ways of working."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Quality Assurance", "QA Strategy · Requirement Management · Verification & Validation · Quality Gates · Go/No-Go Decisions · Defect Management · Release Quality"),
        ("People Leadership", "Team leadership (direct & indirect) · Coaching & mentoring · Performance management · Competence development · Recruitment · Succession planning"),
        ("Test Automation", "Playwright · Vitest · Jest · Python (pytest) · CI/CD integration · Test automation frameworks · TDD · Shift-left testing"),
        ("Security & Privacy", "Information security principles · Data privacy (GDPR) · Security testing practices · CEH certified · Secure development lifecycle"),
        ("Standards & Governance", "Setting QA standards & guidelines · Cross-organizational commonality · Best practices · Process improvement · Six Sigma"),
        ("Tools & Platforms", "Jira · Confluence · Azure DevOps · TestRail · Zephyr · GitHub Actions · Terraform · GCP · AWS · Docker · Kubernetes"),
        ("Certifications", "ISTQB Certified Tester · Certified Ethical Hacker (CEH) · AWS Cloud Practitioner · Google Cloud ACE · Six Sigma Green Belt · ITIL v4"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.5)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS Team Lead Acting
    add_role(doc, "QA & Team Lead Acting — Visual Customer Support (VCS)", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2023 – Present")
    bullets_1 = [
        "Lead and develop the QA competence within the VCS team — coaching, supporting, and motivating co-workers on quality assurance practices, requirement management, and verification & validation.",
        "Responsible for driving competence development — identifying skill gaps, securing upskilling, and ensuring the team meets evolving business quality needs across 32 markets.",
        "Set QA standards and guidelines for the VCS platform — establishing commonality in testing practices, requirement traceability, and quality gates across the product lifecycle.",
        "Drive test automation strategy using Playwright (UI + accessibility testing) — achieving 90%+ coverage on critical user flows and reducing manual regression effort by 70%.",
        "Ensure information security and data privacy compliance — validating that video support solutions meet GDPR requirements and IKEA's security standards before release.",
        "Manage requirement management end-to-end — ensuring business requirements are traceable through design, implementation, and verification across all platform components.",
        "Perform verification & validation across the full ecosystem — coordinating testing of APIs, frontend, integrations, and vendor solutions as an integrated system.",
        "Recommend Quality go/no-go decisions based on risk assessment, defect analysis, and test evidence — advising product leadership on release readiness.",
        "Collaborate with Product and Platform Area Managers to support knowledge sharing and ensure quality outcomes across the Customer Support ecosystem.",
        "Act as ambassador of IKEA values — modelling fact-based business leadership, transparency, and a quality-first culture within the team and across stakeholders.",
        "Manage vendor quality accountability — setting quality expectations, reviewing deliverables, and ensuring vendor solutions meet IKEA standards.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - IKEA CSSP
    add_role(doc, "QA Engineer / Quality Lead — CSSP (Customer Support Staff Planning)", "IKEA, Ingka Digital", "Malmö, Sweden", "2022 – 2023")
    bullets_2 = [
        "Led quality assurance practices for the CSSP system integrating with Verint and Genesys platforms — ensuring requirement management and V&V across global contact center tools.",
        "Defined and implemented QA standards for the product area — establishing test strategies, quality gates, and release criteria aligned with IKEA's shared ways of working.",
        "Drove test automation using GitHub Actions CI/CD pipelines — automated testing, security scanning, and continuous quality validation on every deployment.",
        "Collaborated with multiple stakeholders across product, platform, and technology areas — balancing business needs, quality standards, and delivery timelines.",
        "Ensured data privacy compliance for workforce planning data — validating GDPR requirements in data handling across Verint and Genesys integrations.",
        "Mentored team members on QA best practices, requirement traceability, and test-driven development approaches.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - Truecaller
    add_role(doc, "Quality & Release Manager", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_3 = [
        "Drove quality assurance strategy and verification for a globally released mobile application (300M+ users) — ensuring release quality across multiple platforms.",
        "Managed requirement validation and quality go/no-go decisions for production releases — assessing defect severity, test coverage, and risk exposure.",
        "Built CI/CD automation integrating quality gates — reducing release cycle time by 50% while maintaining quality standards.",
        "Collaborated cross-functionally with product, engineering, and QA teams in a large, complex organisation with shared standards.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - HCLTech
    add_role(doc, "QA Manager / Test Lead", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_4 = [
        "Led and developed QA teams (8–12 members) — responsible for coaching, performance management, competence development, and succession planning within the quality assurance function.",
        "Set QA standards and guidelines across multiple IKEA and LEGO digital products — securing commonality and uniformity in testing practices across cross-functional teams.",
        "Drove requirement management and verification & validation for IoT/connected products — ensuring full traceability from business requirements through implementation to test evidence.",
        "Built and scaled test automation frameworks (Python, Selenium, Playwright) — improving quality and efficiency through automation across regression, integration, and security testing.",
        "Ensured information security and data privacy compliance for connected products — validating firmware, cloud, and app security as part of the overall quality strategy (CEH certified).",
        "Collaborated with Product and Platform Area Managers across IKEA and LEGO — supporting knowledge sharing, driving quality outcomes, and balancing business needs with quality standards.",
        "Drove competence agenda — identified skill gaps, created upskilling programs, and built QA capability across Product Development, Engineering, and Quality teams.",
        "Worked across product, platform, and technology areas in a large, complex organisation with shared standards and ways of working.",
        "Acted as ambassador of IKEA values — role modelling fact-based leadership, togetherness, and continuous improvement within the QA community.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # Role 5 - Earlier
    add_role(doc, "QA Engineer / Software Engineer", "Earlier Career (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_5 = [
        "Worked as an engineer with QA tasks — hands-on testing, automation development, and quality assurance for enterprise financial and retail systems.",
        "Built test automation and tooling (Python, shell scripting) — establishing continuous integration practices and driving systematic quality improvement.",
        "Managed requirement traceability and defect lifecycle in regulated environments.",
        "Drove quality practices across distributed teams — coordinating test execution, reporting, and process improvement.",
    ]
    for b in bullets_5:
        bullet(doc, b)

    # ─── Key QA Achievements ──────────────────────────────────────────────
    add_heading_block(doc, "Key Quality Assurance Achievements")
    achievements = [
        "QA Competence Hub: Built and led QA competence across multiple IKEA domains — established standards, mentoring programs, and knowledge-sharing forums driving quality consistency.",
        "Test Automation at Scale: Drove automation strategy (Playwright, Python, CI/CD) achieving 90%+ coverage on critical flows and reducing manual effort by 70%.",
        "Cross-Platform Quality: Ensured quality across 32 markets for VCS — coordinating V&V across APIs, frontend, vendor integrations, and security as an integrated system.",
        "Security & Privacy: Embedded information security and GDPR compliance into QA processes — ensuring digital products meet data privacy requirements before release.",
        "Standards & Uniformity: Set QA standards and guidelines used across multiple product teams — securing commonality in requirement management, testing, and release quality.",
    ]
    for a in achievements:
        bullet(doc, a, bold_prefix=a.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Togetherness: Drive collaboration across product, platform, and technology areas — fostering knowledge sharing and collective ownership of quality outcomes.",
        "Leading by Example: Coach and develop co-workers, model fact-based business leadership, and set the standard for quality practices across the organization.",
        "Simplicity: Establish clear, practical QA standards and guidelines that teams can adopt easily — reducing complexity while ensuring high quality.",
        "Cost-Consciousness: Drive automation to improve quality and efficiency — reducing manual effort and enabling teams to deliver more with existing resources.",
        "Constant Improvement: Continuously identify competence gaps, drive upskilling, and evolve quality practices to meet changing business needs.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:20%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Quality Assurance Manager with 12+ years of experience in software quality assurance, requirement management, verification &amp; validation, and leading QA teams in digital and software environments. Deep IKEA domain expertise with proven ability to set quality standards, drive competence development, and ensure high-quality outcomes across digital products and platforms. Experienced in information security and data privacy practices, test automation, and balancing business needs with quality standards and delivery outcomes. Strong people leader — coaching, mentoring, and developing QA engineers while driving collaboration across product, platform, and technology areas. Passionate about fact-based decision-making, continuous improvement, and embedding quality into every stage of the development lifecycle. ISTQB certified with experience working in large, complex organisations with shared standards and ways of working.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Quality Assurance</td><td>QA Strategy · Requirement Management · Verification &amp; Validation · Quality Gates · Go/No-Go Decisions · Defect Management · Release Quality</td></tr>
<tr><td class="cat">People Leadership</td><td>Team leadership (direct &amp; indirect) · Coaching &amp; mentoring · Performance management · Competence development · Recruitment · Succession planning</td></tr>
<tr><td class="cat">Test Automation</td><td>Playwright · Vitest · Jest · Python (pytest) · CI/CD integration · Test automation frameworks · TDD · Shift-left testing</td></tr>
<tr><td class="cat">Security &amp; Privacy</td><td>Information security principles · Data privacy (GDPR) · Security testing practices · CEH certified · Secure development lifecycle</td></tr>
<tr><td class="cat">Standards &amp; Governance</td><td>Setting QA standards &amp; guidelines · Cross-organizational commonality · Best practices · Process improvement · Six Sigma</td></tr>
<tr><td class="cat">Tools &amp; Platforms</td><td>Jira · Confluence · Azure DevOps · TestRail · Zephyr · GitHub Actions · Terraform · GCP · AWS · Docker · Kubernetes</td></tr>
<tr><td class="cat">Certifications</td><td>ISTQB Certified Tester · Certified Ethical Hacker (CEH) · AWS Cloud Practitioner · Google Cloud ACE · Six Sigma Green Belt · ITIL v4</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">QA &amp; Team Lead Acting — Visual Customer Support (VCS) <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2023 – Present</span></p>
<ul>
<li>Lead and develop the QA competence within the VCS team — coaching, supporting, and motivating co-workers on quality assurance practices, requirement management, and verification &amp; validation.</li>
<li>Responsible for driving competence development — identifying skill gaps, securing upskilling, and ensuring the team meets evolving business quality needs across 32 markets.</li>
<li>Set QA standards and guidelines for the VCS platform — establishing commonality in testing practices, requirement traceability, and quality gates across the product lifecycle.</li>
<li>Drive test automation strategy using Playwright (UI + accessibility testing) — achieving 90%+ coverage on critical user flows and reducing manual regression effort by 70%.</li>
<li>Ensure information security and data privacy compliance — validating that video support solutions meet GDPR requirements and IKEA's security standards before release.</li>
<li>Manage requirement management end-to-end — ensuring business requirements are traceable through design, implementation, and verification across all platform components.</li>
<li>Perform verification &amp; validation across the full ecosystem — coordinating testing of APIs, frontend, integrations, and vendor solutions as an integrated system.</li>
<li>Recommend Quality go/no-go decisions based on risk assessment, defect analysis, and test evidence — advising product leadership on release readiness.</li>
<li>Collaborate with Product and Platform Area Managers to support knowledge sharing and ensure quality outcomes across the Customer Support ecosystem.</li>
<li>Act as ambassador of IKEA values — modelling fact-based business leadership, transparency, and a quality-first culture within the team and across stakeholders.</li>
<li>Manage vendor quality accountability — setting quality expectations, reviewing deliverables, and ensuring vendor solutions meet IKEA standards.</li>
</ul>

<p class="role">QA Engineer / Quality Lead — CSSP (Customer Support Staff Planning) <span class="meta">&nbsp;|&nbsp; IKEA, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – 2023</span></p>
<ul>
<li>Led quality assurance practices for the CSSP system integrating with Verint and Genesys platforms — ensuring requirement management and V&amp;V across global contact center tools.</li>
<li>Defined and implemented QA standards for the product area — establishing test strategies, quality gates, and release criteria aligned with IKEA's shared ways of working.</li>
<li>Drove test automation using GitHub Actions CI/CD pipelines — automated testing, security scanning, and continuous quality validation on every deployment.</li>
<li>Collaborated with multiple stakeholders across product, platform, and technology areas — balancing business needs, quality standards, and delivery timelines.</li>
<li>Ensured data privacy compliance for workforce planning data — validating GDPR requirements in data handling across Verint and Genesys integrations.</li>
<li>Mentored team members on QA best practices, requirement traceability, and test-driven development approaches.</li>
</ul>

<p class="role">Quality &amp; Release Manager <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Drove quality assurance strategy and verification for a globally released mobile application (300M+ users) — ensuring release quality across multiple platforms.</li>
<li>Managed requirement validation and quality go/no-go decisions for production releases — assessing defect severity, test coverage, and risk exposure.</li>
<li>Built CI/CD automation integrating quality gates — reducing release cycle time by 50% while maintaining quality standards.</li>
<li>Collaborated cross-functionally with product, engineering, and QA teams in a large, complex organisation with shared standards.</li>
</ul>

<p class="role">QA Manager / Test Lead <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led and developed QA teams (8–12 members) — responsible for coaching, performance management, competence development, and succession planning within the quality assurance function.</li>
<li>Set QA standards and guidelines across multiple IKEA and LEGO digital products — securing commonality and uniformity in testing practices across cross-functional teams.</li>
<li>Drove requirement management and verification &amp; validation for IoT/connected products — ensuring full traceability from business requirements through implementation to test evidence.</li>
<li>Built and scaled test automation frameworks (Python, Selenium, Playwright) — improving quality and efficiency through automation across regression, integration, and security testing.</li>
<li>Ensured information security and data privacy compliance for connected products — validating firmware, cloud, and app security as part of the overall quality strategy (CEH certified).</li>
<li>Collaborated with Product and Platform Area Managers across IKEA and LEGO — supporting knowledge sharing, driving quality outcomes, and balancing business needs with quality standards.</li>
<li>Drove competence agenda — identified skill gaps, created upskilling programs, and built QA capability across Product Development, Engineering, and Quality teams.</li>
<li>Worked across product, platform, and technology areas in a large, complex organisation with shared standards and ways of working.</li>
<li>Acted as ambassador of IKEA values — role modelling fact-based leadership, togetherness, and continuous improvement within the QA community.</li>
</ul>

<p class="role">QA Engineer / Software Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Worked as an engineer with QA tasks — hands-on testing, automation development, and quality assurance for enterprise financial and retail systems.</li>
<li>Built test automation and tooling (Python, shell scripting) — establishing continuous integration practices and driving systematic quality improvement.</li>
<li>Managed requirement traceability and defect lifecycle in regulated environments.</li>
<li>Drove quality practices across distributed teams — coordinating test execution, reporting, and process improvement.</li>
</ul>

<h2>KEY QUALITY ASSURANCE ACHIEVEMENTS</h2>
<ul>
<li><strong>QA Competence Hub:</strong> Built and led QA competence across multiple IKEA domains — established standards, mentoring programs, and knowledge-sharing forums driving quality consistency.</li>
<li><strong>Test Automation at Scale:</strong> Drove automation strategy (Playwright, Python, CI/CD) achieving 90%+ coverage on critical flows and reducing manual effort by 70%.</li>
<li><strong>Cross-Platform Quality:</strong> Ensured quality across 32 markets for VCS — coordinating V&amp;V across APIs, frontend, vendor integrations, and security as an integrated system.</li>
<li><strong>Security &amp; Privacy:</strong> Embedded information security and GDPR compliance into QA processes — ensuring digital products meet data privacy requirements before release.</li>
<li><strong>Standards &amp; Uniformity:</strong> Set QA standards and guidelines used across multiple product teams — securing commonality in requirement management, testing, and release quality.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Togetherness:</strong> Drive collaboration across product, platform, and technology areas — fostering knowledge sharing and collective ownership of quality outcomes.</li>
<li><strong>Leading by Example:</strong> Coach and develop co-workers, model fact-based business leadership, and set the standard for quality practices across the organization.</li>
<li><strong>Simplicity:</strong> Establish clear, practical QA standards and guidelines that teams can adopt easily — reducing complexity while ensuring high quality.</li>
<li><strong>Cost-Consciousness:</strong> Drive automation to improve quality and efficiency — reducing manual effort and enabling teams to deliver more with existing resources.</li>
<li><strong>Constant Improvement:</strong> Continuously identify competence gaps, drive upskilling, and evolve quality practices to meet changing business needs.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
