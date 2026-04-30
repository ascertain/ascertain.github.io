from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Schneider_SW_Quality_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Schneider_SW_Quality_Engineer_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "SW Quality",
    "software quality",
    "quality assurance",
    "offer quality",
    "project quality",
    "8D",
    "5 Why",
    "Ishikawa",
    "problem solving",
    "customer escalation",
    "lessons learned",
    "Voice of Customer",
    "Critical to Quality",
    "SQI",
    "defect",
    "verification",
    "validation",
    "OLM",
    "Offer Lifecycle",
    "milestone",
    "Go/No go",
    "Early Warning",
    "risk",
    "cyber security",
    "IEC62443",
    "Agile",
    "Scrum",
    "PI planning",
    "Release Management",
    "CI/CD",
    "cross-functional",
    "cross-team",
    "C#",
    "C++",
    ".NET",
    "Java",
    "TypeScript",
    "Python",
    "SQL",
    "customer satisfaction",
    "continuous improvement",
    "field quality",
    "IKEA",
    "LEGO",
    "Truecaller",
    "Six Sigma",
    "ISTQB",
    "leadership",
    "coaching",
    "mentor",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Software Quality Engineer  |  Offer & Project Quality, V&V, Problem Solving  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Software Quality Engineer with 15+ years of enterprise software development and "
        "testing experience, driving offer quality and project quality across SaaS, IoT, "
        "and enterprise products. Deploys SW Quality Fundamentals — verification and "
        "validation, defect analysis (SQI/dSQI), 8D customer escalation management, "
        "root cause analysis (5 Why, Ishikawa), and lessons learned processes. Ensures "
        "quality is built into every project following Offer Lifecycle Management (OLM), "
        "captures Voice of Customer, defines Critical to Quality requirements, and prepares "
        "milestone Go/No go decisions. Strong programming background (C#, C++, Java, "
        "TypeScript, Python, .NET, SQL). Cyber security awareness (CEH certified, threat "
        "modeling). Leads quality initiatives cross-functionally without direct reports — "
        "coaches teams on quality tools and methods. Agile practitioner with PI planning "
        "and Release Management experience. Six Sigma Green Belt. ISTQB certified.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Offer Quality: ",
            "Deployed SW Quality Fundamentals across products serving 30+ markets (IKEA) "
            "and 300M+ users (Truecaller). Managed customer escalations using 8D method. "
            "Regular follow-up and analysis of SW performance through defects, SQI, dSQI. "
            "Drove lessons learned and implemented best practices based on field performance. "
            "Supported R&D and Tech Support on quality aspects.",
        ),
        (
            "Project Quality: ",
            "Ensured offer quality built into every project — Voice of Customer collection, "
            "Critical to Quality requirements, Offer Quality Goals, risk identification and "
            "mitigation. Prepared quality summaries for milestone reviews with Go/No go "
            "recommendations. Defined Early Warning Systems before product launches. "
            "Coached project teams on quality tools and methods for robustness and "
            "customer satisfaction targets.",
        ),
        (
            "Problem Solving & Continuous Improvement: ",
            "Applied 8D, 5 Why, Ishikawa, and statistical methods for root cause analysis "
            "and corrective actions. Drove Issue Review Board escalations cross-functionally. "
            "Captured field quality experience and fed it back into offer development for "
            "continuous improvement.",
        ),
        (
            "Programming & Technical Breadth: ",
            "C#, C++, Java, TypeScript, Python, .NET, SQL, JavaScript, HTML, XML. "
            "Test automation frameworks (Playwright, Selenium, RestAssured). CI/CD pipelines. "
            "Cyber security awareness (CEH certified, IEC62443 familiarity). Broad technical "
            "understanding across SW and embedded systems.",
        ),
        (
            "Cross-Functional Leadership: ",
            "Leads using quality authority without direct reports — builds cross-team "
            "relationships, mobilizes cross-functional activities. Coaches and mentors "
            "engineers on quality processes. Strong influencing, customer-focused mindset. "
            "Agile (Scrum, PI planning), Release Management.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Quality Methods: ",
            "8D, 5 Why, Ishikawa, root cause analysis, FMEA, SQI/dSQI, "
            "defect analysis, V&V, risk assessment, lessons learned, OLM",
        ),
        (
            "Languages: ",
            "C#, C++, Java, TypeScript, Python, .NET, JavaScript, SQL, "
            "HTML, XML, Go (familiarity)",
        ),
        (
            "Test Automation: ",
            "Playwright, Selenium, Appium, RestAssured, Karate, TestNG, "
            "Jest, Postman, E2E, API, regression, contract testing",
        ),
        (
            "CI/CD & Tools: ",
            "GitHub Actions, Jenkins, Docker, Kubernetes, Terraform, "
            "Git, TestRail, Jira, Confluence, Grafana",
        ),
        (
            "Security: ",
            "Certified Ethical Hacker (CEH), cyber security fundamentals, "
            "threat modeling, IEC62443 awareness",
        ),
        (
            "Process: ",
            "Agile/Scrum, PI planning, Release Management, Six Sigma "
            "Green Belt, ISTQB, ITIL, OLM, milestone reviews",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Software Quality Engineer / Engineering Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — SaaS Platform, Enterprise Software, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Owned offer quality for enterprise SaaS product — deployed SW Quality Fundamentals, "
        "managed customer escalations (8D method), drove defect analysis (SQI tracking), "
        "and implemented lessons learned from field performance across 30+ markets.",
        "Ensured project quality built into every release — Voice of Customer captured, "
        "Critical to Quality requirements defined, risk identification and mitigation "
        "plans maintained. Prepared quality summaries for milestone reviews with Go/No go "
        "recommendations. Defined Early Warning Systems before launches.",
        "Built scalable test automation frameworks (Playwright, RestAssured, Karate) and "
        "integrated into CI/CD pipelines (GitHub Actions, Docker, Terraform). Verification "
        "and validation across the full stack. Shift-left testing in Agile (Scrum, PI planning).",
        "Led quality initiatives cross-functionally without direct reports — coached "
        "engineers on problem solving methods (8D, 5 Why, Ishikawa), quality tools, and "
        "robustness targets. Drove Issue Review Board escalations. Strong influencing.",
        "Cyber security awareness — threat modeling, security testing. Release Management "
        "in Agile. Continuous improvement of quality processes and methodologies.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior QA Lead / Software Quality Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  SaaS — 300M+ Users, Mobile & Enterprise", bold=True, size=10)
    tc_bullets = [
        "Software quality assurance for SaaS platform with 300M+ users — customer "
        "escalation management, defect analysis, field quality monitoring, root cause "
        "analysis. Test automation (Selenium, Appium, RestAssured). V&V across mobile "
        "and backend. Customer-focused quality mindset.",
        "Cross-functional quality leadership — coached engineers on testing best "
        "practices and problem solving. CI/CD. Release Management in Agile "
        "(Scrum/Kanban). Continuous improvement of quality processes.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Quality Lead / Senior QA Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — In-House & Brand-Labeled SW", bold=True, size=10)
    lego_bullets = [
        "LEGO: Offer quality for e-commerce platform — SW Quality Fundamentals, "
        "defect tracking, customer escalation handling, lessons learned. Built "
        "automation frameworks (Selenium, RestAssured, Karate, Appium). Led 8–10 "
        "engineers. V&V, risk assessment, milestone reviews.",
        "IKEA App, Genesys, Verint (2018–2021): Project quality across multiple "
        "products — Voice of Customer, quality goals, risk mitigation. Cross-functional "
        "coordination, coached teams on quality tools and methods. Release Management "
        "in Agile. Problem solving (8D, root cause).",
        "Quality leadership without direct reports — built cross-team relationships, "
        "mobilized cross-functional activities. Strong influencing. Continuous "
        "improvement. Customer satisfaction focus.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior QA Engineer / Quality Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech & Government — Enterprise Software", bold=True, size=10)
    fin_bullets = [
        "Enterprise software quality assurance — V&V, defect management, customer "
        "escalations, root cause analysis (8D, 5 Why). Built automation frameworks "
        "(Selenium, RestAssured, TestNG). CI/CD. SQL, Java, C#.",
        "Coached 15+ engineers on quality processes and problem solving methods. Led "
        "teams of 10+. Quality leadership across cross-functional teams. Agile "
        "(Scrum/Kanban). Continuous improvement. Customer-focused mindset.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Offer quality ownership — SW Quality Fundamentals, 8D customer escalations, "
        "defect analysis (SQI/dSQI), lessons learned across SaaS products serving "
        "300M+ users (Truecaller) and 30+ markets (IKEA).",
        "Project quality built into every release — Voice of Customer, Critical to "
        "Quality, risk identification, milestone Go/No go, Early Warning Systems.",
        "Problem solving champion — 8D, 5 Why, Ishikawa, root cause analysis. Drove "
        "Issue Review Board escalations cross-functionally. Field quality feedback "
        "loop for continuous improvement.",
        "Cross-functional quality leadership — led without direct reports, built "
        "cross-team relationships, coached teams on quality tools and methods. "
        "Strong influencing and customer satisfaction focus.",
        "Technical breadth — C#, C++, Java, TypeScript, Python, .NET, SQL. Test "
        "automation, CI/CD, cyber security (CEH). Six Sigma Green Belt. ISTQB.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(
        edu,
        "M.Tech / B.Tech in Information Technology\n"
        "PG Diploma in Operations Management",
        size=9.5,
    )
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Six Sigma Green Belt\n"
        "Certified Ethical Hacker (CEH)\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "ITIL Foundation\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Schneider Electric SW Quality Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Software Quality Engineer | Offer &amp; Project Quality, V&amp;V, Problem Solving | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Software Quality</span> Engineer with 15+ years of enterprise software development and testing. Drives <span class="hl">offer quality</span> and <span class="hl">project quality</span> — SW Quality Fundamentals, <span class="hl">verification</span> &amp; <span class="hl">validation</span>, <span class="hl">defect</span> analysis (<span class="hl">SQI</span>/dSQI), <span class="hl">8D</span> <span class="hl">customer escalation</span> management, root cause analysis (<span class="hl">5 Why</span>, <span class="hl">Ishikawa</span>), <span class="hl">lessons learned</span>. Ensures quality built into every project (<span class="hl">OLM</span>), captures <span class="hl">Voice of Customer</span>, defines <span class="hl">Critical to Quality</span>, prepares <span class="hl">milestone</span> <span class="hl">Go/No go</span>. Programming: <span class="hl">C#</span>, <span class="hl">C++</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>, <span class="hl">Python</span>, <span class="hl">.NET</span>, <span class="hl">SQL</span>. <span class="hl">Cyber security</span> (CEH, IEC62443 awareness). <span class="hl">Cross-functional</span> <span class="hl">leadership</span> without direct reports. <span class="hl">Agile</span>, <span class="hl">PI planning</span>, <span class="hl">Release Management</span>. <span class="hl">Six Sigma</span> Green Belt. <span class="hl">ISTQB</span>.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Offer Quality:</b> SW Quality Fundamentals. <span class="hl">8D</span> <span class="hl">customer escalation</span>s. <span class="hl">Defect</span> analysis (<span class="hl">SQI</span>/dSQI). <span class="hl">Lessons learned</span> from <span class="hl">field quality</span>. Supported R&amp;D &amp; Tech Support. 30+ markets (<span class="hl">IKEA</span>), 300M+ users (<span class="hl">Truecaller</span>).<br>
  <b>Project Quality:</b> <span class="hl">Voice of Customer</span>, <span class="hl">Critical to Quality</span>, <span class="hl">risk</span> identification &amp; mitigation, <span class="hl">milestone</span> <span class="hl">Go/No go</span>, <span class="hl">Early Warning</span> Systems, <span class="hl">OLM</span>. Coached teams on robustness &amp; <span class="hl">customer satisfaction</span> targets.<br>
  <b>Problem Solving:</b> <span class="hl">8D</span>, <span class="hl">5 Why</span>, <span class="hl">Ishikawa</span>, root cause, Issue Review Board. <span class="hl">Continuous improvement</span> — <span class="hl">field quality</span> feedback loop.<br>
  <b>Programming:</b> <span class="hl">C#</span>, <span class="hl">C++</span>, <span class="hl">Java</span>, <span class="hl">TypeScript</span>, <span class="hl">Python</span>, <span class="hl">.NET</span>, <span class="hl">SQL</span>. Test automation. <span class="hl">CI/CD</span>. <span class="hl">Cyber security</span> (CEH, IEC62443).<br>
  <b>Cross-Functional Leadership:</b> Leads without direct reports. <span class="hl">Cross-team</span> relationships. <span class="hl">Coaching</span>. <span class="hl">Agile</span> (<span class="hl">Scrum</span>, <span class="hl">PI planning</span>), <span class="hl">Release Management</span>. <span class="hl">Six Sigma</span>.</p>

  <div class="section">Technical Skills</div>
  <p><b>Quality Methods:</b> 8D, 5 Why, Ishikawa, FMEA, SQI/dSQI, V&amp;V, risk assessment, lessons learned, OLM<br>
  <b>Languages:</b> C#, C++, Java, TypeScript, Python, .NET, JavaScript, SQL, HTML, XML, Go (familiarity)<br>
  <b>Test Automation:</b> Playwright, Selenium, Appium, RestAssured, Karate, TestNG, Jest, Postman<br>
  <b>CI/CD &amp; Tools:</b> GitHub Actions, Jenkins, Docker, Kubernetes, Terraform, Git, TestRail, Jira, Confluence<br>
  <b>Security:</b> CEH, cyber security fundamentals, threat modeling, IEC62443 awareness<br>
  <b>Process:</b> Agile/Scrum, PI planning, Release Management, Six Sigma Green Belt, ISTQB, ITIL, OLM</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Software Quality Engineer / Engineering Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS — SaaS Platform, Enterprise Software, 30+ Markets</div>
  <ul>
    <li><span class="hl">Offer quality</span> <span class="hl">ownership</span> — <span class="hl">SW Quality</span> Fundamentals, <span class="hl">8D</span> <span class="hl">customer escalation</span>s, <span class="hl">defect</span> analysis (<span class="hl">SQI</span>), <span class="hl">lessons learned</span> from <span class="hl">field quality</span>. 30+ markets.</li>
    <li><span class="hl">Project quality</span> — <span class="hl">Voice of Customer</span>, <span class="hl">Critical to Quality</span>, <span class="hl">risk</span> mitigation, <span class="hl">milestone</span> <span class="hl">Go/No go</span>, <span class="hl">Early Warning</span> Systems. <span class="hl">OLM</span>.</li>
    <li>Test automation (Playwright, RestAssured, Karate). <span class="hl">CI/CD</span> (GitHub Actions, Docker). <span class="hl">V&amp;V</span>. <span class="hl">Shift-left</span> in <span class="hl">Agile</span> (<span class="hl">Scrum</span>, <span class="hl">PI planning</span>).</li>
    <li><span class="hl">Cross-functional</span> <span class="hl">leadership</span> — <span class="hl">coaching</span> on <span class="hl">problem solving</span> (<span class="hl">8D</span>, <span class="hl">5 Why</span>, <span class="hl">Ishikawa</span>). <span class="hl">Cyber security</span>. <span class="hl">Release Management</span>. <span class="hl">Continuous improvement</span>.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior QA Lead / Software Quality Engineer</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | SaaS — 300M+ Users, Mobile &amp; Enterprise</div>
  <ul>
    <li><span class="hl">Software quality assurance</span> — <span class="hl">customer escalation</span>s, <span class="hl">defect</span> analysis, <span class="hl">field quality</span>, root cause. Test automation (Selenium, Appium, RestAssured). <span class="hl">V&amp;V</span>.</li>
    <li><span class="hl">Cross-functional</span> <span class="hl">leadership</span>. <span class="hl">Coaching</span>. <span class="hl">CI/CD</span>. <span class="hl">Release Management</span>. <span class="hl">Agile</span>. <span class="hl">Continuous improvement</span>. <span class="hl">Customer satisfaction</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Quality Lead / Senior QA Engineer</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — In-House &amp; Brand-Labeled SW</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Offer quality</span> — <span class="hl">SW Quality</span> Fundamentals, <span class="hl">defect</span> tracking, <span class="hl">customer escalation</span>s, <span class="hl">lessons learned</span>. Automation (Selenium, RestAssured, Karate, Appium). 8–10 engineers. <span class="hl">V&amp;V</span>, <span class="hl">risk</span>, <span class="hl">milestone</span> reviews.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Project quality</span> — <span class="hl">Voice of Customer</span>, quality goals, <span class="hl">risk</span> mitigation. <span class="hl">Cross-functional</span> coordination. <span class="hl">Release Management</span>. <span class="hl">Problem solving</span> (<span class="hl">8D</span>). <span class="hl">Customer satisfaction</span>.</li>
    <li><span class="hl">Leadership</span> without direct reports — <span class="hl">cross-team</span> relationships. <span class="hl">Coaching</span>. <span class="hl">Continuous improvement</span>. <span class="hl">Agile</span> (<span class="hl">Scrum</span>).</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior QA Engineer / Quality Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech &amp; Government — Enterprise Software</div>
  <ul>
    <li><span class="hl">Software quality assurance</span> — <span class="hl">V&amp;V</span>, <span class="hl">defect</span> management, <span class="hl">customer escalation</span>s, root cause (<span class="hl">8D</span>, <span class="hl">5 Why</span>). Automation (Selenium, RestAssured). <span class="hl">CI/CD</span>. <span class="hl">SQL</span>, <span class="hl">Java</span>, <span class="hl">C#</span>.</li>
    <li><span class="hl">Coaching</span> 15+ engineers. <span class="hl">Leadership</span>. <span class="hl">Cross-functional</span>. <span class="hl">Agile</span>. <span class="hl">Continuous improvement</span>. <span class="hl">Customer satisfaction</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Offer quality</span> — <span class="hl">8D</span>, <span class="hl">SQI</span>/dSQI, <span class="hl">lessons learned</span> across 300M+ users (<span class="hl">Truecaller</span>) and 30+ markets (<span class="hl">IKEA</span>).</li>
    <li><span class="hl">Project quality</span> — <span class="hl">Voice of Customer</span>, <span class="hl">Critical to Quality</span>, <span class="hl">risk</span>, <span class="hl">milestone</span> <span class="hl">Go/No go</span>, <span class="hl">Early Warning</span> Systems.</li>
    <li><span class="hl">Problem solving</span> — <span class="hl">8D</span>, <span class="hl">5 Why</span>, <span class="hl">Ishikawa</span>. <span class="hl">Field quality</span> feedback. <span class="hl">Continuous improvement</span>.</li>
    <li><span class="hl">Cross-functional</span> <span class="hl">leadership</span> — no direct reports, <span class="hl">cross-team</span>, <span class="hl">coaching</span>, <span class="hl">customer satisfaction</span>.</li>
    <li>Technical breadth — <span class="hl">C#</span>, <span class="hl">C++</span>, <span class="hl">Java</span>, <span class="hl">Python</span>, <span class="hl">.NET</span>. <span class="hl">Cyber security</span> (CEH). <span class="hl">Six Sigma</span>. <span class="hl">ISTQB</span>.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>Six Sigma Green Belt<br>Certified Ethical Hacker (CEH)<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>ITIL Foundation<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
