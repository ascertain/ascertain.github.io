"""
Resume: Hybrid — IT Service & Test Engineering Manager
Combines ITSM processes, service delivery, monitoring AND test engineering, verification, quality.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_IT_Service_Test_Manager_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("IT Service & Test Engineering Manager with ", False),
        ("16+ years", True),
        (" leading teams across ", False),
        ("IT service management", True),
        (", ", False),
        ("test engineering", True),
        (", and ", False),
        ("service delivery", True),
        (" in complex global environments. Experienced in owning ", False),
        ("ITSM processes", True),
        (" (Incident, Problem, Change, Monitoring) while simultaneously driving ", False),
        ("test method development", True),
        (", ", False),
        ("verification strategies", True),
        (", and ", False),
        ("quality systems", True),
        (". Proven ", False),
        ("people leader", True),
        (" who builds high-performing teams, manages ", False),
        ("capacity and resource planning", True),
        (", and aligns execution with strategic priorities. Strong in ", False),
        ("stakeholder management", True),
        (", ", False),
        ("agile & DevOps", True),
        (" ways of working, and ", False),
        ("service performance reporting", True),
        (". Established a ", False),
        ("Centre of Excellence", True),
        (" for quality solutions serving teams horizontally.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "ITIL / ITSM Processes", "Test Engineering & Verification", "Team Leadership & Coaching",
        "Incident & Problem Management", "Test Method Development", "Capacity & Resource Planning",
        "Monitoring & Alert Management", "Service Performance & KPIs", "Agile & DevOps Practices",
        "Stakeholder Management", "Centre of Excellence (QA)", "Continuous Improvement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical & Domain Skills")
    skills_data = [
        ("ITSM:", " Incident, Problem, Change, Request, Monitoring & Alert Mgmt, SLA/SLO management"),
        ("Test Engineering:", " Test method design & validation, verification strategies, V-model, risk-based testing"),
        ("Tools:", " JIRA Service Management, XRAY, TestRail, Confluence, PagerDuty, Opsgenie"),
        ("Automation:", " Python, Pytest, Playwright, Selenium, Robot Framework, custom test infrastructure"),
        ("Monitoring & Reporting:", " Grafana, dashboards, alerting pipelines, KPI tracking, availability metrics"),
        ("CI/CD & DevOps:", " GitHub Actions, Jenkins, Docker, Kubernetes, GitOps, automated quality gates"),
        ("Cloud:", " GCP (Cloud Run, GKE, Pub/Sub, BigQuery), AWS, Terraform, microservices"),
        ("Practices:", " ITIL v4, Agile/Scrum, Kanban, DevOps, Six Sigma, blameless post-mortems"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Team Lead (Acting) / IT Service & Test Engineering Manager",
        "Mar 2022 \u2013 Present",
        "Global Customer Connect Platform \u2014 30+ Markets | IT Operations, Testing & Service Delivery | Agile & DevOps")
    add_bullet(doc, "Own and evolve ITSM processes (Incident, Problem, Change, Monitoring & Alert Management) while leading test engineering for integrated HW/SW platform.",
        ["ITSM processes", "Incident, Problem, Change, Monitoring", "test engineering"])
    add_bullet(doc, "Lead and develop two teams: service management operations and test engineering \u2014 building an engaging, safe, high-performing work environment.",
        ["Lead and develop two teams", "service management", "test engineering", "high-performing"])
    add_bullet(doc, "Drive service management and test reporting: availability, performance KPIs, test coverage, defect trends, release readiness \u2014 actionable insights for stakeholders.",
        ["service management and test reporting", "KPIs", "test coverage", "defect trends", "release readiness"])
    add_bullet(doc, "Define and implement new test methods and test setups for innovation projects \u2014 generating robust data supporting sound design decisions and verification strategies.",
        ["new test methods", "innovation projects", "robust data", "verification strategies"])
    add_bullet(doc, "Ensure effective capacity and resource planning across teams; identify risks, bottlenecks, and improvement opportunities proactively.",
        ["capacity and resource planning", "risks, bottlenecks", "improvement opportunities"])
    add_bullet(doc, "Established a Test Engineering Centre of Excellence \u2014 designing quality solutions and frameworks shared horizontally across teams while owning one product\u2019s service and verification.",
        ["Centre of Excellence", "quality solutions", "horizontally across teams"])
    add_bullet(doc, "Implement and strengthen agile and DevOps ways of working \u2014 CI/CD pipelines, automated monitoring, quality gates, blameless post-mortems.",
        ["agile and DevOps", "CI/CD pipelines", "automated monitoring", "quality gates"])
    add_bullet(doc, "Coach and support team members: recruitment, competence development, goal-setting, career growth \u2014 translate operational insights into strategic context.",
        ["Coach and support", "recruitment", "competence development", "career growth"])
    add_bullet(doc, "Collaborate with stakeholders across IT, Product, R&D, and business \u2014 aligning execution with strategic priorities in a fast-changing global environment.",
        ["stakeholders across IT, Product, R&D", "strategic priorities", "global environment"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Operations Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Release Operations, Service Reliability & Test Automation")
    add_bullet(doc, "Managed release operations and service reliability \u2014 incident response, change management, deployment pipelines, monitoring and alerting.",
        ["release operations", "service reliability", "incident response", "change management"])
    add_bullet(doc, "Developed test methods for release verification; ensured quality gates and service performance metrics met SLA targets.",
        ["test methods", "release verification", "quality gates", "SLA targets"])
    add_bullet(doc, "Collaborated cross-functionally to resolve production issues; structured post-incident reviews driving process and testing improvements.",
        ["cross-functionally", "production issues", "post-incident reviews"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 Service & Test Delivery Lead",
        "2013 \u2013 2021",
        "Enterprise IT \u2014 Service Management, Test Engineering & Multi-Geography Delivery")
    add_bullet(doc, "Led IT service delivery and test engineering across global programmes \u2014 managed distributed teams (8\u201312) spanning multiple geographies and functions.",
        ["IT service delivery", "test engineering", "distributed teams", "multiple geographies"])
    add_bullet(doc, "Owned incident, problem, and change management processes; established monitoring, alerting, and escalation frameworks for enterprise platforms.",
        ["incident, problem, and change management", "monitoring, alerting, and escalation"])
    add_bullet(doc, "Defined and implemented new test methods for complex integrated systems \u2014 verification strategies from concept through design verification (V-model, risk-based).",
        ["test methods", "verification strategies", "V-model, risk-based"])
    add_bullet(doc, "Built test infrastructure and automated setups (Python, Java, Selenium); service performance dashboards; KPI-driven reporting for stakeholders.",
        ["test infrastructure", "service performance dashboards", "KPI-driven reporting"])
    add_bullet(doc, "Drove continuous improvement: evolved from manual operations to automated, DevOps-driven service delivery and testing \u2014 ways of working transformation.",
        ["continuous improvement", "DevOps-driven", "ways of working transformation"])
    add_bullet(doc, "Resource planning, capacity management, and budget optimisation across concurrent projects; structured prioritisation and stakeholder communication.",
        ["Resource planning", "capacity management", "budget optimisation", "stakeholder communication"])
    add_bullet(doc, "Coached engineers on operational and testing best practices; strengthened collaboration across teams and geographies.",
        ["Coached", "collaboration across teams and geographies"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 IT Operations & Test Engineer / Consultant",
        "2008 \u2013 2013",
        "Core Banking & Biometric Devices \u2014 Enterprise-Scale, Regulated (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "IT operations and test engineering for Finacle Core Banking \u2014 incident management, change control, verification, and compliance in regulated environment.",
        ["IT operations", "test engineering", "incident management", "change control", "regulated"])
    add_bullet(doc, "Device integration testing (Morpho BAS 2FA) \u2014 HW/SW verification, test method development, monitoring and alert automation for mission-critical systems.",
        ["Device integration testing", "HW/SW verification", "test method development", "monitoring"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Six Sigma Green Belt", "Google Cloud ACE", "AWS Cloud Practitioner", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
