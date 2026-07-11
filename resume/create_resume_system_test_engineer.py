"""
Resume: System Test Engineer / Test Lead — Embedded Linux, HW/SW Validation
Focus: Exploratory testing, embedded systems, Python automation, CI pipelines, debugging.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_System_Test_Engineer_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("System Test professional with ", False),
        ("16+ years", True),
        (" of hands-on experience in ", False),
        ("system-level test strategy", True),
        (", ", False),
        ("exploratory testing", True),
        (", and ", False),
        ("deep debugging across hardware and software layers", True),
        (". Strong ", False),
        ("Python", True),
        (" developer building ", False),
        ("automated test frameworks", True),
        (" integrated into ", False),
        ("CI pipelines", True),
        (". Proven ", False),
        ("test leadership", True),
        (" across ", False),
        ("embedded systems", True),
        (", ", False),
        ("communication protocol testing", True),
        (", and cross-disciplinary platforms combining HW, SW, and integrations. Structured ", False),
        ("problem-solver", True),
        (" with a quality-driven mindset and passion for ", False),
        ("investigative testing", True),
        (" and ", False),
        ("continuous improvement", True),
        (".", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "System Testing & Test Leadership", "Exploratory Testing", "Embedded Systems (Linux)",
        "Python Test Automation", "Debugging & Troubleshooting", "Communication Protocol Testing",
        "CI Pipeline Integration", "Jira / XRAY Test Management", "Cross-Functional Collaboration",
        "Problem Solving (Systematic)", "HW/SW Integration Validation", "Continuous Improvement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills_data = [
        ("Languages:", " Python (primary), Java, TypeScript/JavaScript, Bash, SQL"),
        ("Test Automation:", " Pytest, Playwright, Selenium, Appium, Robot Framework, custom HW test harnesses"),
        ("CI/CD:", " GitHub Actions, Jenkins, Azure DevOps, Docker, Kubernetes"),
        ("Embedded/HW:", " Embedded Linux validation, HW/SW integration, device interface testing, system-level debugging"),
        ("Protocols:", " REST API, WebSocket, MQTT, serial communication, network protocol validation"),
        ("Test Management:", " Jira, XRAY, TestRail, Confluence, defect tracking & reporting"),
        ("Cloud & Infra:", " GCP (Cloud Run, Pub/Sub), Docker, Kubernetes, Amazon EC2"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB, Malm\u00f6 \u2014 Team Lead (Acting) / SDET",
        "Mar 2022 \u2013 Present",
        "IKEA App & Customer Connect VCS Platform \u2014 Embedded Devices, Software, Data Pipelines | 30+ Markets")
    add_bullet(doc, "Design and execute system-level test strategies across integrated platforms combining hardware devices, embedded software, and cloud services.",
        ["system-level test strategies", "hardware devices", "embedded software", "cloud services"])
    add_bullet(doc, "Lead system testing activities \u2014 planning, coordination, and hands-on execution ensuring thorough end-to-end validation.",
        ["Lead system testing", "planning, coordination", "hands-on execution", "end-to-end validation"])
    add_bullet(doc, "Perform exploratory testing and deep debugging across HW/SW layers \u2014 uncovering hidden defects in complex system interactions.",
        ["exploratory testing", "deep debugging across HW/SW layers", "hidden defects"])
    add_bullet(doc, "Develop and maintain automated test frameworks in Python; integrate tests into CI pipelines (GitHub Actions) for continuous validation.",
        ["automated test frameworks in Python", "CI pipelines", "continuous validation"])
    add_bullet(doc, "Test and validate communication protocols (REST API, WebSocket, device interfaces) and system integrations across distributed components.",
        ["communication protocols", "REST API", "WebSocket", "system integrations"])
    add_bullet(doc, "Drive structured problem-solving and continuous improvement of testing processes \u2014 root cause analysis, defect trend identification.",
        ["structured problem-solving", "continuous improvement", "root cause analysis"])
    add_bullet(doc, "Manage test cases, defects, and reporting using Jira/XRAY; quality telemetry dashboards (Grafana) for data-driven decisions.",
        ["Jira/XRAY", "quality telemetry", "data-driven decisions"])
    add_bullet(doc, "Collaborate cross-functionally with developers, product owners, and ops teams to identify, analyse, and resolve defects.",
        ["cross-functionally", "identify, analyse, and resolve defects"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Automation Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Microservices")
    add_bullet(doc, "Managed release pipelines and CI infrastructure (Jenkins) \u2014 automated test integration ensuring system stability across deployments.",
        ["release pipelines", "CI infrastructure", "system stability"])
    add_bullet(doc, "Performed system-level validation of communication platform components; debugging across distributed microservices.",
        ["system-level validation", "debugging", "distributed microservices"])
    add_bullet(doc, "Drove defect analysis and structured problem-solving to improve release quality and reduce regression risks.",
        ["defect analysis", "structured problem-solving", "release quality"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 Technical Specialist / SDET Lead",
        "2013 \u2013 2021",
        "E-Commerce, Mobile & Enterprise \u2014 Cross-Disciplinary Systems (HW + SW + Integrations)")
    add_bullet(doc, "Led system testing activities for cross-disciplinary platforms \u2014 combining hardware, embedded software, web applications, and third-party integrations.",
        ["system testing", "cross-disciplinary platforms", "hardware, embedded software", "integrations"])
    add_bullet(doc, "Designed system-level test strategies with strong focus on exploratory testing \u2014 investigating complex behaviours across multi-component systems.",
        ["system-level test strategies", "exploratory testing", "complex behaviours"])
    add_bullet(doc, "Built automated test frameworks (Python, Java, Selenium, Cucumber) and integrated into CI pipelines (Jenkins, Amazon EC2).",
        ["automated test frameworks", "Python", "Java", "CI pipelines", "Jenkins"])
    add_bullet(doc, "Hands-on debugging and troubleshooting across HW/SW boundaries \u2014 device communication, system integration failures, protocol mismatches.",
        ["debugging and troubleshooting", "HW/SW boundaries", "device communication", "protocol mismatches"])
    add_bullet(doc, "Validated communication protocols and system integrations across 10+ interconnected platforms \u2014 40% regression cycle reduction.",
        ["communication protocols", "system integrations", "40% regression cycle reduction"])
    add_bullet(doc, "Managed distributed teams (8\u201312 engineers); mentored on testing practices, framework adoption, and structured problem-solving.",
        ["distributed teams", "mentored", "structured problem-solving"])
    add_bullet(doc, "Improved test frameworks, processes, and automation maturity across programmes \u2014 manual-first to automation-first transformation.",
        ["test frameworks, processes, and automation maturity", "automation-first transformation"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 SDET / Consultant",
        "2008 \u2013 2013",
        "Finacle CBS, Core Banking \u2014 Regulated Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "System testing for Finacle Core Banking \u2014 transaction processing, account management, regulatory compliance validation.",
        ["System testing", "Finacle Core Banking", "regulatory compliance"])
    add_bullet(doc, "Morpho BAS 2FA: Hardware/software integration testing for biometric authentication devices \u2014 protocol validation, device communication.",
        ["Hardware/software integration testing", "biometric authentication", "protocol validation", "device communication"])
    add_bullet(doc, "Built Python and Java test automation frameworks; deep debugging of complex transaction flows across system boundaries.",
        ["Python", "Java", "test automation frameworks", "deep debugging"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Google Cloud ACE", "AWS Cloud Practitioner", "CEH", "Six Sigma Green Belt"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
