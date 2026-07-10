"""
Resume: Test Manager – WHSmith (UK, EXTERNAL)
Focus: Test strategy & leadership, release management, environment/data management,
       supplier management, non-functional testing, CI/CD, stakeholder management,
       offshore team management, continuous improvement.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Test_Manager_WHSmith_Resume"
BRAND = RGBColor(0x1B, 0x3A, 0x5C)  # Deep navy


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="1B3A5C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Experienced Test Manager with 14+ years leading end-to-end testing capabilities across retail, "
        "e-commerce, and enterprise platforms. Proven expertise in defining test strategies, managing release processes, "
        "coordinating test environments, and governing non-functional testing (performance, security, resilience). "
        "Track record of managing distributed onshore/offshore QA teams and external test suppliers while maintaining "
        "cost efficiency and quality KPIs. Skilled in stakeholder management at all levels, driving continuous improvement "
        "through automation adoption, process standardisation, and tool optimisation. ISTQB certified with hands-on experience "
        "across Agile, Scrum, and Waterfall delivery methodologies."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Test Strategy & Planning", "Release Management", "Environment & Data Management",
        "Team Leadership & KPIs", "Supplier & Vendor Management", "Non-Functional Testing (Perf/Sec)",
        "Stakeholder Management", "CI/CD & Automation Frameworks", "Defect & Risk Governance",
        "Offshore/Distributed Teams", "Test Tool Selection & Mgmt", "Continuous Improvement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TOOLS & TECHNOLOGIES ──
    add_heading_block(doc, "Tools & Technologies")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    tools_text = (
        "Jira  •  Azure DevOps  •  TestRail  •  HP ALM / Quality Centre  •  "
        "Selenium  •  Cypress  •  Postman  •  JMeter  •  k6  •  "
        "Jenkins  •  GitHub Actions  •  Azure Pipelines  •  "
        "AWS  •  GCP  •  Azure  •  Docker  •  Kubernetes"
    )
    r = tp.add_run(tools_text)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- Ingka Digital / IKEA ---
    role_header(doc, "Test Manager / Team Lead – Customer Connect Platforms",
                "Ingka Digital (IKEA)", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Define and own the test strategy across 5+ integrated systems; establish repeatable processes for SIT, UAT, regression, and business acceptance testing.")
    bullet(doc, "Lead release management activities: coordinate test readiness, deployment scheduling, and Go/No-Go governance across multiple product groups.")
    bullet(doc, "Accountable for test environment provisioning and data management; ensure environments are stable, fit for purpose, and aligned with release calendars.")
    bullet(doc, "Manage distributed testing resources (onshore & offshore); define KPIs, track capacity, and drive quality standards across the testing capability.")
    bullet(doc, "Own non-functional testing decisions: identify when performance, resilience, and security testing is required; coordinate execution with specialist teams.")
    bullet(doc, "Evaluate, select, and manage testing tools; drive adoption of automation frameworks to improve test coverage and reduce regression cycle times by 40%.")
    bullet(doc, "Build strong stakeholder relationships with technology partners, business SMEs, and project teams; deliver concise project-level test reporting to steering committees.")
    bullet(doc, "Drive continuous improvement: introduced shift-left testing practices, automated smoke suites in CI/CD pipelines, and standardised defect triage processes.")
    bullet(doc, "Identify upstream/downstream dependencies, risks, and issues; ensure testing activities are accurately reflected in project plans and RAID logs.")

    # --- Truecaller ---
    role_header(doc, "QA Lead – Platform Testing",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Developed test strategies for microservices architecture; coordinated API, integration, and E2E testing across multiple technology stacks.")
    bullet(doc, "Managed test environments for cloud-native SaaS platform; ensured test data integrity and environment stability for concurrent release trains.")
    bullet(doc, "Established structured release readiness checkpoints and defect management workflows using Jira and automation reporting dashboards.")

    # --- HCLTech ---
    role_header(doc, "Test Manager – Enterprise Programmes",
                "HCLTech (for IKEA & LEGO)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Led the testing capability for large-scale retail and e-commerce programmes spanning 10+ interconnected systems (order mgmt, supply chain, CRM, POS).")
    bullet(doc, "Managed a team of 8–12 Test Analysts across onshore and offshore locations; accountable for capacity planning, capability development, and performance reviews.")
    bullet(doc, "Owned test supplier relationships: managed external vendor teams, negotiated capacity, tracked delivery against SLAs, and controlled testing costs.")
    bullet(doc, "Defined and executed non-functional test strategies including performance testing (JMeter), security scanning, and API contract testing.")
    bullet(doc, "Governed release management processes: coordinated multi-team deployments, environment bookings, and deployment verification testing.")
    bullet(doc, "Managed test tool ecosystem including HP ALM/Quality Centre, Jira, and introduced TestRail for improved test case management and traceability.")
    bullet(doc, "Facilitated Agile ceremonies (Scrum of Scrums, sprint demos, retrospectives) and Waterfall gate reviews; adapted testing approach to delivery methodology.")
    bullet(doc, "Drove continuous improvement initiatives: automated regression suites, reduced defect leakage to production by 35%, standardised test process documentation.")
    bullet(doc, "Collaborated with offshore delivery centres (India) on test execution, knowledge transfer, and follow-the-sun testing models.")

    # --- India ---
    role_header(doc, "Senior Test Engineer / Test Lead",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Progressed from Test Engineer to Test Lead; owned test planning, execution, and defect lifecycle for web and enterprise applications.")
    bullet(doc, "Gained experience across diverse technologies (Java, .NET, Oracle, web services) and testing tools (QC, Selenium, LoadRunner).")
    bullet(doc, "Built foundation in SDLC/STLC processes, stakeholder communication, and cross-functional team collaboration.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "ISTQB Certified Tester – Foundation Level",
        "AWS Certified Cloud Practitioner",
        "Google Cloud Associate Cloud Engineer",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("PGDOM")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – IGNOU  |  ").font.size = Pt(9)
    r2 = p.add_run("B.Tech Information Technology")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – UP Technical University").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Basic)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#1B3A5C;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#1B3A5C;font-size:11pt;border-bottom:1px solid #1B3A5C;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#1B3A5C;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Experienced Test Manager with 14+ years leading end-to-end testing capabilities across retail,
e-commerce, and enterprise platforms. Proven expertise in defining test strategies, managing release processes,
coordinating test environments, and governing non-functional testing (performance, security, resilience).
Track record of managing distributed onshore/offshore QA teams and external test suppliers while maintaining
cost efficiency and quality KPIs. Skilled in stakeholder management at all levels, driving continuous improvement
through automation adoption, process standardisation, and tool optimisation. ISTQB certified with hands-on experience
across Agile, Scrum, and Waterfall delivery methodologies.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Test Strategy &amp; Planning</td><td>✓ Release Management</td><td>✓ Environment &amp; Data Management</td></tr>
<tr><td>✓ Team Leadership &amp; KPIs</td><td>✓ Supplier &amp; Vendor Management</td><td>✓ Non-Functional Testing (Perf/Sec)</td></tr>
<tr><td>✓ Stakeholder Management</td><td>✓ CI/CD &amp; Automation Frameworks</td><td>✓ Defect &amp; Risk Governance</td></tr>
<tr><td>✓ Offshore/Distributed Teams</td><td>✓ Test Tool Selection &amp; Mgmt</td><td>✓ Continuous Improvement</td></tr>
</table>

<h2>TOOLS &amp; TECHNOLOGIES</h2>
<p class="tools">Jira &bull; Azure DevOps &bull; TestRail &bull; HP ALM / Quality Centre &bull;
Selenium &bull; Cypress &bull; Postman &bull; JMeter &bull; k6 &bull;
Jenkins &bull; GitHub Actions &bull; Azure Pipelines &bull;
AWS &bull; GCP &bull; Azure &bull; Docker &bull; Kubernetes</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Test Manager / Team Lead – Customer Connect Platforms &nbsp;|&nbsp; Ingka Digital (IKEA), Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Define and own the test strategy across 5+ integrated systems; establish repeatable processes for SIT, UAT, regression, and business acceptance testing.</li>
<li>Lead release management activities: coordinate test readiness, deployment scheduling, and Go/No-Go governance across multiple product groups.</li>
<li>Accountable for test environment provisioning and data management; ensure environments are stable, fit for purpose, and aligned with release calendars.</li>
<li>Manage distributed testing resources (onshore &amp; offshore); define KPIs, track capacity, and drive quality standards across the testing capability.</li>
<li>Own non-functional testing decisions: identify when performance, resilience, and security testing is required; coordinate execution with specialist teams.</li>
<li>Evaluate, select, and manage testing tools; drive adoption of automation frameworks to improve test coverage and reduce regression cycle times by 40%.</li>
<li>Build strong stakeholder relationships with technology partners, business SMEs, and project teams; deliver concise project-level test reporting to steering committees.</li>
<li>Drive continuous improvement: introduced shift-left testing practices, automated smoke suites in CI/CD pipelines, and standardised defect triage processes.</li>
<li>Identify upstream/downstream dependencies, risks, and issues; ensure testing activities are accurately reflected in project plans and RAID logs.</li>
</ul>

<p class="role">QA Lead – Platform Testing &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Developed test strategies for microservices architecture; coordinated API, integration, and E2E testing across multiple technology stacks.</li>
<li>Managed test environments for cloud-native SaaS platform; ensured test data integrity and environment stability for concurrent release trains.</li>
<li>Established structured release readiness checkpoints and defect management workflows using Jira and automation reporting dashboards.</li>
</ul>

<p class="role">Test Manager – Enterprise Programmes &nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Led the testing capability for large-scale retail and e-commerce programmes spanning 10+ interconnected systems (order mgmt, supply chain, CRM, POS).</li>
<li>Managed a team of 8–12 Test Analysts across onshore and offshore locations; accountable for capacity planning, capability development, and performance reviews.</li>
<li>Owned test supplier relationships: managed external vendor teams, negotiated capacity, tracked delivery against SLAs, and controlled testing costs.</li>
<li>Defined and executed non-functional test strategies including performance testing (JMeter), security scanning, and API contract testing.</li>
<li>Governed release management processes: coordinated multi-team deployments, environment bookings, and deployment verification testing.</li>
<li>Managed test tool ecosystem including HP ALM/Quality Centre, Jira, and introduced TestRail for improved test case management and traceability.</li>
<li>Facilitated Agile ceremonies (Scrum of Scrums, sprint demos, retrospectives) and Waterfall gate reviews; adapted testing approach to delivery methodology.</li>
<li>Drove continuous improvement initiatives: automated regression suites, reduced defect leakage to production by 35%, standardised test process documentation.</li>
<li>Collaborated with offshore delivery centres (India) on test execution, knowledge transfer, and follow-the-sun testing models.</li>
</ul>

<p class="role">Senior Test Engineer / Test Lead &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Progressed from Test Engineer to Test Lead; owned test planning, execution, and defect lifecycle for web and enterprise applications.</li>
<li>Gained experience across diverse technologies (Java, .NET, Oracle, web services) and testing tools (QC, Selenium, LoadRunner).</li>
<li>Built foundation in SDLC/STLC processes, stakeholder communication, and cross-functional team collaboration.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>ISTQB Certified Tester – Foundation Level</li>
<li>AWS Certified Cloud Practitioner</li>
<li>Google Cloud Associate Cloud Engineer</li>
<li>Six Sigma Green Belt</li>
<li>Certified Ethical Hacker (CEH)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>PGDOM</b> – IGNOU &nbsp;|&nbsp; <b>B.Tech Information Technology</b> – UP Technical University</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Basic) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
