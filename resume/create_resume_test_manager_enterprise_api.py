"""
Resume: Test Manager – Enterprise Applications & API Testing (EXTERNAL)
Focus: Test strategy & leadership, API validation, Agile QA, test governance,
       sourcing/procurement domain, SQL data validation, cross-team coordination,
       UX-driven platforms, defect management, test metrics & dashboards.
EXTERNAL = show diverse company experience.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Test_Manager_Enterprise_API_Resume"
BRAND = RGBColor(0x1A, 0x3C, 0x5E)  # Professional dark navy


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="1A3C5E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Test Manager with 14+ years of experience leading end-to-end testing strategy and execution across "
        "enterprise applications, web services, and sourcing/procurement platforms. Deep expertise in Agile QA "
        "leadership, API validation (Postman, SoapUI, REST), and test governance — with a strong focus on aligning "
        "testing with real user journeys and UX outcomes. Proven track record managing functional, integration, "
        "regression, and UAT testing across multiple Agile teams, driving risk-based testing practices, and delivering "
        "clear release readiness reporting through metrics and dashboards. Skilled in SQL-based data validation, "
        "cross-team coordination, defect management, and mentoring teams on quality best practices."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Test Strategy & Governance", "API / Web Services Testing", "Agile QA Leadership",
        "Functional / Integration / UAT", "SQL Data Validation", "Risk-Based Testing",
        "Test Metrics & Dashboards", "Cross-Team Coordination", "Defect Management (Jira/GitHub)",
        "Sourcing / Procurement Domain", "CI/CD & Automation (Exposure)", "Team Mentoring & Best Practices",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TOOLS & TECHNOLOGIES ──
    add_heading_block(doc, "Tools & Technologies")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    skills = (
        "API Testing: Postman, SoapUI, REST Assured, k6  •  "
        "Test Management: qTest, TestRail, HP ALM / Quality Centre, Zephyr  •  "
        "Defect Tracking: Jira, Bugzilla, GitHub Issues  •  "
        "SQL: PostgreSQL, BigQuery, Oracle — backend data validation & query writing  •  "
        "SCM: Git, GitHub, Bitbucket  •  "
        "CI/CD: Jenkins, GitHub Actions, Azure DevOps  •  "
        "Automation: Selenium, Cypress, Cucumber/Gherkin, Pytest, Vitest  •  "
        "Monitoring: Grafana dashboards, test metrics reporting  •  "
        "Collaboration: Confluence, Miro, Slack"
    )
    r = tp.add_run(skills)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- Ingka Digital / IKEA ---
    role_header(doc, "Test Manager / Team Lead – Enterprise Digital Platforms",
                "Ingka Digital (IKEA)", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Lead end-to-end test strategy, planning, and execution across multiple Agile teams for 5+ integrated enterprise applications and web services supporting the customer journey.")
    bullet(doc, "Oversee functional, integration, regression, and UAT testing cycles; define quality gates and release readiness criteria for each deployment.")
    bullet(doc, "Manage API testing and validation using Postman and REST frameworks across microservices; ensure contract compliance and backward compatibility for all web service endpoints.")
    bullet(doc, "Drive test case design, maintenance, and traceability using TestRail and Jira; maintain living test documentation aligned with evolving business requirements.")
    bullet(doc, "Define and track test metrics, dashboards, and release readiness reports — providing stakeholders with clear visibility into quality status, defect trends, and automation coverage.")
    bullet(doc, "Implement risk-based testing practices: prioritise test efforts based on business impact, change risk, and historical defect data to optimise coverage within sprint timelines.")
    bullet(doc, "Perform SQL-based data validation across integrated systems (BigQuery, PostgreSQL); verify data integrity, transformation accuracy, and cross-system consistency.")
    bullet(doc, "Collaborate closely with product owners, business analysts, and UX teams to align testing with real user journeys and customer experience requirements.")
    bullet(doc, "Coordinate defect management using Jira and GitHub; lead defect triage sessions, drive resolution timelines, and ensure clear root-cause documentation.")
    bullet(doc, "Mentor team members on Agile QA best practices, test automation adoption, and quality-first mindset; conduct knowledge-sharing sessions and pairing workshops.")

    # --- Truecaller ---
    role_header(doc, "QA Lead – Platform Testing",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Led API test strategy for microservices-based SaaS platform; built Postman collections and automated API validation suites for contract and integration testing.")
    bullet(doc, "Defined test metrics and quality dashboards; tracked defect density, pass/fail rates, and release readiness across sprint cycles.")
    bullet(doc, "Validated backend data using SQL queries; ensured data consistency across distributed services and databases.")
    bullet(doc, "Coordinated testing across multiple Agile feature teams; managed cross-dependency defect triage and integration test scheduling.")

    # --- HCLTech ---
    role_header(doc, "Test Manager / Senior QA Lead – Enterprise Sourcing & E-Commerce Platforms",
                "HCLTech (for IKEA & LEGO)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Led test strategy and execution for large-scale enterprise applications in sourcing, procurement, and e-commerce domains — spanning 10+ integrated systems and multiple Agile release trains.")
    bullet(doc, "Managed API and web services testing using Postman, SoapUI, and REST frameworks; validated complex integrations between sourcing platforms, order management, and supply chain systems.")
    bullet(doc, "Oversaw functional, integration, regression, and UAT testing for enterprise-wide releases; established quality governance practices and release readiness criteria.")
    bullet(doc, "Drove test case design and traceability using HP ALM / Quality Centre and Jira; maintained comprehensive test repositories linked to business requirements and user stories.")
    bullet(doc, "Performed extensive SQL data validation (Oracle, PostgreSQL) across integrated backend systems; verified data flows, transformations, and business rule accuracy.")
    bullet(doc, "Defined and tracked test metrics and dashboards for programme-level reporting: defect trends, test progress, automation ROI, and quality gate compliance.")
    bullet(doc, "Implemented risk-based testing approaches; prioritised test coverage based on business criticality, change impact, and historical defect patterns.")
    bullet(doc, "Coordinated defect management using Jira and Bugzilla; led cross-team defect triage, escalation workflows, and root-cause analysis.")
    bullet(doc, "Collaborated with product, business, and UX teams to map test scenarios to real user journeys; ensured testing aligned with customer experience goals.")
    bullet(doc, "Managed a team of 8–12 test analysts (onshore & offshore); mentored team members on Agile QA practices, promoted continuous improvement culture.")
    bullet(doc, "Supported CI/CD pipeline integration for automated test suites (Jenkins, GitHub Actions); drove adoption of automation frameworks to reduce manual regression effort.")

    # --- India ---
    role_header(doc, "QA Engineer / Test Lead",
                "Multiple Companies (HCL, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullet(doc, "Progressed from QA Engineer to Test Lead; owned test planning, execution, and defect lifecycle for web applications and enterprise systems.")
    bullet(doc, "Built foundation in API testing (SOAP/REST), SQL validation, and test management tools; worked across Agile and Waterfall delivery environments.")
    bullet(doc, "Gained experience across diverse technology stacks (Java, .NET, Oracle, web services) in digital and enterprise application testing.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "ISTQB Certified Tester – Foundation Level",
        "Google Cloud Associate Cloud Engineer",
        "AWS Certified Cloud Practitioner",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – UP Technical University  |  ").font.size = Pt(9)
    r2 = p.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – IGNOU").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Basic)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#1A3C5E;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#1A3C5E;font-size:11pt;border-bottom:1px solid #1A3C5E;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#1A3C5E;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Test Manager with 14+ years of experience leading end-to-end testing strategy and execution across
enterprise applications, web services, and sourcing/procurement platforms. Deep expertise in Agile QA
leadership, API validation (Postman, SoapUI, REST), and test governance — with a strong focus on aligning
testing with real user journeys and UX outcomes. Proven track record managing functional, integration,
regression, and UAT testing across multiple Agile teams, driving risk-based testing practices, and delivering
clear release readiness reporting through metrics and dashboards. Skilled in SQL-based data validation,
cross-team coordination, defect management, and mentoring teams on quality best practices.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Test Strategy &amp; Governance</td><td>✓ API / Web Services Testing</td><td>✓ Agile QA Leadership</td></tr>
<tr><td>✓ Functional / Integration / UAT</td><td>✓ SQL Data Validation</td><td>✓ Risk-Based Testing</td></tr>
<tr><td>✓ Test Metrics &amp; Dashboards</td><td>✓ Cross-Team Coordination</td><td>✓ Defect Management (Jira/GitHub)</td></tr>
<tr><td>✓ Sourcing / Procurement Domain</td><td>✓ CI/CD &amp; Automation (Exposure)</td><td>✓ Team Mentoring &amp; Best Practices</td></tr>
</table>

<h2>TOOLS &amp; TECHNOLOGIES</h2>
<p class="tools">API Testing: Postman, SoapUI, REST Assured, k6 &bull;
Test Management: qTest, TestRail, HP ALM / Quality Centre, Zephyr &bull;
Defect Tracking: Jira, Bugzilla, GitHub Issues &bull;
SQL: PostgreSQL, BigQuery, Oracle — backend data validation &amp; query writing &bull;
SCM: Git, GitHub, Bitbucket &bull;
CI/CD: Jenkins, GitHub Actions, Azure DevOps &bull;
Automation: Selenium, Cypress, Cucumber/Gherkin, Pytest, Vitest &bull;
Monitoring: Grafana dashboards, test metrics reporting &bull;
Collaboration: Confluence, Miro, Slack</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Test Manager / Team Lead – Enterprise Digital Platforms &nbsp;|&nbsp; Ingka Digital (IKEA), Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Lead end-to-end test strategy, planning, and execution across multiple Agile teams for 5+ integrated enterprise applications and web services supporting the customer journey.</li>
<li>Oversee functional, integration, regression, and UAT testing cycles; define quality gates and release readiness criteria for each deployment.</li>
<li>Manage API testing and validation using Postman and REST frameworks across microservices; ensure contract compliance and backward compatibility for all web service endpoints.</li>
<li>Drive test case design, maintenance, and traceability using TestRail and Jira; maintain living test documentation aligned with evolving business requirements.</li>
<li>Define and track test metrics, dashboards, and release readiness reports — providing stakeholders with clear visibility into quality status, defect trends, and automation coverage.</li>
<li>Implement risk-based testing practices: prioritise test efforts based on business impact, change risk, and historical defect data to optimise coverage within sprint timelines.</li>
<li>Perform SQL-based data validation across integrated systems (BigQuery, PostgreSQL); verify data integrity, transformation accuracy, and cross-system consistency.</li>
<li>Collaborate closely with product owners, business analysts, and UX teams to align testing with real user journeys and customer experience requirements.</li>
<li>Coordinate defect management using Jira and GitHub; lead defect triage sessions, drive resolution timelines, and ensure clear root-cause documentation.</li>
<li>Mentor team members on Agile QA best practices, test automation adoption, and quality-first mindset; conduct knowledge-sharing sessions and pairing workshops.</li>
</ul>

<p class="role">QA Lead – Platform Testing &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Led API test strategy for microservices-based SaaS platform; built Postman collections and automated API validation suites for contract and integration testing.</li>
<li>Defined test metrics and quality dashboards; tracked defect density, pass/fail rates, and release readiness across sprint cycles.</li>
<li>Validated backend data using SQL queries; ensured data consistency across distributed services and databases.</li>
<li>Coordinated testing across multiple Agile feature teams; managed cross-dependency defect triage and integration test scheduling.</li>
</ul>

<p class="role">Test Manager / Senior QA Lead – Enterprise Sourcing &amp; E-Commerce Platforms &nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Led test strategy and execution for large-scale enterprise applications in sourcing, procurement, and e-commerce domains — spanning 10+ integrated systems and multiple Agile release trains.</li>
<li>Managed API and web services testing using Postman, SoapUI, and REST frameworks; validated complex integrations between sourcing platforms, order management, and supply chain systems.</li>
<li>Oversaw functional, integration, regression, and UAT testing for enterprise-wide releases; established quality governance practices and release readiness criteria.</li>
<li>Drove test case design and traceability using HP ALM / Quality Centre and Jira; maintained comprehensive test repositories linked to business requirements and user stories.</li>
<li>Performed extensive SQL data validation (Oracle, PostgreSQL) across integrated backend systems; verified data flows, transformations, and business rule accuracy.</li>
<li>Defined and tracked test metrics and dashboards for programme-level reporting: defect trends, test progress, automation ROI, and quality gate compliance.</li>
<li>Implemented risk-based testing approaches; prioritised test coverage based on business criticality, change impact, and historical defect patterns.</li>
<li>Coordinated defect management using Jira and Bugzilla; led cross-team defect triage, escalation workflows, and root-cause analysis.</li>
<li>Collaborated with product, business, and UX teams to map test scenarios to real user journeys; ensured testing aligned with customer experience goals.</li>
<li>Managed a team of 8–12 test analysts (onshore &amp; offshore); mentored team members on Agile QA practices, promoted continuous improvement culture.</li>
<li>Supported CI/CD pipeline integration for automated test suites (Jenkins, GitHub Actions); drove adoption of automation frameworks to reduce manual regression effort.</li>
</ul>

<p class="role">QA Engineer / Test Lead &nbsp;|&nbsp; Multiple Companies (HCL, Marlabs, TekMindz), India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li>Progressed from QA Engineer to Test Lead; owned test planning, execution, and defect lifecycle for web applications and enterprise systems.</li>
<li>Built foundation in API testing (SOAP/REST), SQL validation, and test management tools; worked across Agile and Waterfall delivery environments.</li>
<li>Gained experience across diverse technology stacks (Java, .NET, Oracle, web services) in digital and enterprise application testing.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>ISTQB Certified Tester – Foundation Level</li>
<li>Google Cloud Associate Cloud Engineer</li>
<li>AWS Certified Cloud Practitioner</li>
<li>Six Sigma Green Belt</li>
<li>Certified Ethical Hacker (CEH)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>B.Tech Information Technology</b> – UP Technical University &nbsp;|&nbsp; <b>PGDOM</b> – IGNOU</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Basic) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
