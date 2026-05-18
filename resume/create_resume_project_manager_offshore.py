from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Project_Manager_Offshore_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Project_Manager_Offshore_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "project management",
    "program management",
    "project delivery",
    "scope",
    "schedule",
    "budget",
    "risk management",
    "risk",
    "multi-partner",
    "supplier",
    "stakeholder",
    "compliance",
    "regulatory",
    "safety",
    "quality",
    "FAT",
    "inspection",
    "commissioning",
    "fabrication",
    "testing",
    "procurement",
    "engineering",
    "resource allocation",
    "cross-functional",
    "client",
    "Denmark",
    "offshore",
    "maritime",
    "IKEA",
    "LEGO",
    "Truecaller",
    "Maersk",
    "Six Sigma",
    "ITIL",
    "Agile",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Project Manager  |  Technical Delivery, Multi-Partner Coordination, Compliance & Quality  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Technical Project Manager with 15+ years leading complex, multi-partner "
        "projects from development through testing, delivery, and commissioning. "
        "Manages project scope, schedule, budget, and risk to ensure timely and "
        "successful delivery. Strong technical background — interfaces effectively "
        "between development, engineering, and execution. Coordinates fabrication, "
        "testing, and delivery activities across distributed locations. Ensures "
        "compliance with safety standards, regulatory requirements, and quality "
        "processes. Acts as primary project interface towards clients and key "
        "stakeholders. Conducts FATs, inspections, and quality readiness reviews. "
        "Manages suppliers, partners, and subcontractors. Plans resource allocation "
        "in coordination with internal teams and partners. Engineering degree "
        "(B.Tech IT), PG Diploma in Operations Management. Six Sigma Green Belt, "
        "ITIL. Experience working in Denmark (LEGO, Copenhagen). Willing to travel.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Project & Program Management (8+ Years): ",
            "Led projects through development, testing, delivery, and commissioning "
            "across IKEA, Truecaller, and LEGO. Managed project scope, schedule, budget, "
            "and risk. Oversaw complex, multi-partner projects including budget monitoring "
            "and audit-ready reporting. Delivered projects on time and within budget across "
            "global, cross-functional teams. 8+ years in project/program leadership roles.",
        ),
        (
            "Technical Background — Dev, Engineering & Execution: ",
            "Solid technical background interfacing between development, engineering, "
            "and execution. Led technical delivery teams of 8–15 engineers. Hands-on "
            "engineering experience (software, systems, infrastructure). Understands "
            "technical trade-offs and translates between engineering and business. "
            "B.Tech in IT. PG Diploma in Operations Management.",
        ),
        (
            "Supplier, Partner & Stakeholder Management: ",
            "Managed suppliers and partners contributing to projects — vendor selection, "
            "scope definition, performance monitoring. Primary project interface towards "
            "clients and stakeholders. Coordinated with procurement, engineering, "
            "operations, and external partners. Strong leadership, organizational, "
            "and communication skills at all levels.",
        ),
        (
            "Compliance, Safety & Quality (FATs, Inspections): ",
            "Ensured compliance with regulatory requirements and quality processes across "
            "regulated environments (banking, enterprise, global platforms). Conducted "
            "FATs and inspections to monitor progress, quality, and readiness. Quality "
            "gate management. Risk identification, assessment, and mitigation. "
            "Six Sigma Green Belt — process quality and continuous improvement.",
        ),
        (
            "Denmark Experience & International Team: ",
            "Experience working in Denmark — LEGO Group e-commerce platform (Copenhagen, "
            "2016–2018). International, collaborative working environment across IKEA "
            "(Sweden), Truecaller (Sweden), LEGO (Denmark), and global teams. "
            "Willing to travel to support project activities. Danish — Conversational.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Project Management: ",
            "Scope, schedule, budget, and risk management. Multi-partner project "
            "coordination. Budget monitoring and audit-ready reporting. Resource "
            "allocation. Project planning (Gantt, milestones, critical path)",
        ),
        (
            "Delivery & Execution: ",
            "Development, fabrication, testing, commissioning. FATs, inspections, "
            "quality readiness. Delivery coordination across distributed locations. "
            "Release/deployment management. Go-live readiness",
        ),
        (
            "Compliance & Quality: ",
            "Regulatory compliance, safety standards awareness, quality processes. "
            "Quality gate management. Six Sigma Green Belt. ITIL Foundation. "
            "Risk identification, assessment, and mitigation. ISTQB certified",
        ),
        (
            "Stakeholder & Partner Management: ",
            "Client interface, stakeholder management at all levels. Supplier and "
            "partner management. Procurement support (scope of work definition). "
            "Cross-functional collaboration (engineering, procurement, operations)",
        ),
        (
            "Technical Background: ",
            "Software/systems engineering, cloud infrastructure (GCP, AWS), CI/CD, "
            "Docker, Kubernetes, Terraform. Python, Java, TypeScript, C#. Data "
            "pipelines. Interface between development, engineering, and execution",
        ),
        (
            "Tools & Practices: ",
            "Jira, Confluence, MS Project, TestRail. Agile (Scrum/SAFe), Waterfall. "
            "Grafana (dashboards, reporting). Git, GitHub Actions",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Technical Project Lead / Senior SDET", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS Platform — Multi-Partner, 30+ Global Markets", bold=True, size=10)
    ikea_bullets = [
        "Project delivery leadership — led projects through development, testing, and "
        "commissioning for a global platform serving 30+ markets. Managed project scope, "
        "schedule, and risk. Coordinated multi-partner delivery across distributed teams "
        "and locations. Budget-aware planning and reporting.",
        "Supplier and partner coordination — managed external vendors and partners "
        "contributing to the project. Defined scopes of work. Primary technical interface "
        "towards stakeholders. Coordinated with engineering, operations, and external "
        "partners for seamless integration and delivery.",
        "Compliance and quality — ensured compliance with regulatory requirements and "
        "internal quality processes. Conducted FATs, inspections, and quality readiness "
        "reviews. Risk identification, assessment, and mitigation. Quality gate management. "
        "Audit-ready documentation and reporting (Grafana dashboards).",
        "Resource allocation and team management — planned resource allocation across "
        "internal teams and partners. Led cross-functional coordination across development, "
        "engineering, and operations. Mentored engineers. Agile (Scrum/SAFe).",
        "Technical execution — oversaw infrastructure delivery (GCP cloud, Docker, "
        "Kubernetes, Terraform, CI/CD). Interfaced between development, engineering, "
        "and execution. AI-assisted productivity improvement (30% velocity gain).",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Project Lead / Senior QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users, Global Delivery", bold=True, size=10)
    tc_bullets = [
        "Project delivery at global scale — managed project scope, schedule, and risk "
        "for a platform serving 300M+ users. Coordinated delivery activities across "
        "multiple teams and partners. Release management and commissioning. Technical "
        "interface towards stakeholders. AWS cloud infrastructure.",
        "Cross-functional coordination — engineering, product, and operations teams. "
        "Risk management. Supplier coordination. Resource planning. Agile (Scrum/Kanban). "
        "International, fast-paced working environment.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Project Lead / Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Denmark & Sweden, Multi-Partner", bold=True, size=10)
    lego_bullets = [
        "LEGO (Denmark): Led complex, multi-partner project delivery for e-commerce "
        "platform — scope, schedule, budget, and risk management. Managed suppliers "
        "and partners. Coordinated fabrication/development, testing, and delivery. "
        "Led 8–10 engineers. FATs, inspections, quality reviews. Client and stakeholder "
        "interface. Release management. Denmark experience (Copenhagen, 2016–2018).",
        "IKEA (2018–2021): Project leadership across multiple technology projects "
        "(IKEA App, Genesys, Verint/CSSP, Spartacus). Multi-partner coordination. "
        "Compliance with quality processes. Scope definition for subcontractors. "
        "Resource allocation. Risk management. Agile (Scrum).",
        "Technical interface — bridged development, engineering, and execution. "
        "Supported procurement in defining scopes of work. Budget monitoring. "
        "Audit-ready reporting. Cross-functional collaboration.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Project Lead / Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking & Enterprise — Regulated, Multi-Partner Delivery", bold=True, size=10)
    fin_bullets = [
        "Project delivery in regulated environments — managed scope, schedule, budget, "
        "and risk for enterprise banking projects (Finacle). Multi-partner coordination. "
        "Compliance with strict regulatory requirements and quality processes. FATs, "
        "inspections, and quality gate management. Audit-ready documentation.",
        "Leadership and resource management — led teams of 10–15 engineers. Resource "
        "allocation across internal teams and subcontractors. Supplier management. "
        "Client interface. Cross-functional coordination (engineering, operations, "
        "procurement). Agile and Waterfall. Six Sigma thinking.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Project delivery across global platforms — IKEA (30+ markets), Truecaller "
        "(300M+ users), LEGO (Denmark). Managed scope, schedule, budget, and risk. "
        "Multi-partner coordination. On-time, within-budget delivery.",
        "Supplier and stakeholder management — managed vendors, partners, subcontractors. "
        "Primary client/stakeholder interface. Procurement support (scope of work). "
        "Cross-functional collaboration with engineering, operations, and partners.",
        "Compliance and quality — FATs, inspections, quality readiness reviews. "
        "Regulatory compliance in regulated environments. Risk management. "
        "Quality gate management. Audit-ready reporting.",
        "Denmark experience — LEGO Group (Copenhagen, 2016–2018). International, "
        "collaborative, trust-based working environment. Willing to travel.",
        "Technical background — interfaces between development, engineering, and "
        "execution. Cloud infrastructure (GCP, AWS). Six Sigma Green Belt. ITIL.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "Six Sigma Green Belt\n"
        "ITIL Foundation\n"
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Danish — Conversational  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Project Manager Offshore Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Project Manager | Technical Delivery, Multi-Partner Coordination, Compliance &amp; Quality | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Technical <span class="hl">Project Manager</span> with 15+ years leading complex, <span class="hl">multi-partner</span> projects from development through <span class="hl">testing</span>, delivery, and <span class="hl">commissioning</span>. Manages <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, and <span class="hl">risk</span>. Technical background — interfaces between development, <span class="hl">engineering</span>, and execution. Coordinates <span class="hl">fabrication</span>, <span class="hl">testing</span>, and delivery across distributed locations. <span class="hl">Compliance</span> with <span class="hl">safety</span> standards, <span class="hl">regulatory</span> requirements, <span class="hl">quality</span> processes. <span class="hl">Client</span> and <span class="hl">stakeholder</span> interface. <span class="hl">FAT</span>s, <span class="hl">inspection</span>s. <span class="hl">Supplier</span> and partner management. <span class="hl">Resource allocation</span>. B.Tech IT, PG Diploma Operations Management. <span class="hl">Six Sigma</span> GB, <span class="hl">ITIL</span>. <span class="hl">Denmark</span> experience (<span class="hl">LEGO</span>). Willing to travel.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Project &amp; Program Management (8+ Years):</b> Led projects through development, <span class="hl">testing</span>, delivery, <span class="hl">commissioning</span> at <span class="hl">IKEA</span>, <span class="hl">Truecaller</span>, <span class="hl">LEGO</span>. <span class="hl">Scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Multi-partner</span>. Audit-ready reporting. On-time delivery.<br>
  <b>Technical Background:</b> Solid technical — interfaces between development, <span class="hl">engineering</span>, execution. Led 8–15 engineers. B.Tech IT. PG Diploma Operations Management.<br>
  <b>Supplier &amp; Stakeholder:</b> <span class="hl">Supplier</span>/partner management. <span class="hl">Scope</span> of work definition. <span class="hl">Client</span> interface. <span class="hl">Procurement</span> support. <span class="hl">Cross-functional</span> (engineering, <span class="hl">procurement</span>, operations).<br>
  <b>Compliance, Quality &amp; FATs:</b> <span class="hl">Regulatory</span> <span class="hl">compliance</span>. <span class="hl">Quality</span> processes. <span class="hl">FAT</span>s, <span class="hl">inspection</span>s, readiness reviews. <span class="hl">Risk</span> identification/mitigation. <span class="hl">Six Sigma</span> GB.<br>
  <b>Denmark &amp; International:</b> <span class="hl">LEGO</span> (<span class="hl">Denmark</span>, Copenhagen 2016–18). International teams — <span class="hl">IKEA</span> (Sweden), <span class="hl">Truecaller</span> (Sweden). Danish — Conversational. Willing to travel.</p>

  <div class="section">Core Competencies</div>
  <p><b>Project Management:</b> Scope, schedule, budget, risk. Multi-partner coordination. Budget monitoring, audit-ready reporting. Resource allocation. Project planning<br>
  <b>Delivery &amp; Execution:</b> Development, fabrication, testing, commissioning. FATs, inspections, quality readiness. Distributed delivery. Release management<br>
  <b>Compliance &amp; Quality:</b> Regulatory compliance, safety awareness, quality processes. Quality gates. Six Sigma GB, ITIL, ISTQB. Risk management<br>
  <b>Stakeholder &amp; Partners:</b> Client interface, stakeholder management. Supplier/partner management. Procurement support. Cross-functional collaboration<br>
  <b>Technical:</b> Software/systems engineering, cloud (GCP, AWS), CI/CD, Docker, Kubernetes, Terraform. Python, Java, TypeScript, C#<br>
  <b>Tools:</b> Jira, Confluence, MS Project, TestRail, Grafana (dashboards), Git, GitHub Actions. Agile (Scrum/SAFe), Waterfall</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Technical Project Lead / Senior SDET</div>
  <div class="job-sub">Mar 2022 – Present | VCS Platform — Multi-Partner, 30+ Global Markets</div>
  <ul>
    <li><span class="hl">Project delivery</span> — development, <span class="hl">testing</span>, <span class="hl">commissioning</span> for 30+ markets. <span class="hl">Scope</span>, <span class="hl">schedule</span>, <span class="hl">risk</span>. <span class="hl">Multi-partner</span> delivery across distributed teams. <span class="hl">Budget</span>-aware planning.</li>
    <li><span class="hl">Supplier</span>/partner coordination — managed vendors. Defined scopes of work. <span class="hl">Client</span>/<span class="hl">stakeholder</span> interface. <span class="hl">Engineering</span>, operations, partner <span class="hl">integration</span>.</li>
    <li><span class="hl">Compliance</span> and <span class="hl">quality</span> — <span class="hl">FAT</span>s, <span class="hl">inspection</span>s, readiness reviews. <span class="hl">Risk</span> management. <span class="hl">Quality</span> gates. Audit-ready reporting (Grafana).</li>
    <li><span class="hl">Resource allocation</span> — internal teams and partners. <span class="hl">Cross-functional</span> coordination. Mentored engineers. <span class="hl">Agile</span> (Scrum/SAFe).</li>
    <li>Technical execution — cloud (GCP, Docker, Kubernetes, Terraform, CI/CD). 30% productivity improvement (AI-assisted).</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Project Lead / Senior QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users, Global Delivery</div>
  <ul>
    <li><span class="hl">Project delivery</span> at global scale — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">risk</span> for 300M+ users. Delivery coordination, <span class="hl">multi-partner</span>. Release/<span class="hl">commissioning</span>. <span class="hl">Client</span> interface. AWS.</li>
    <li><span class="hl">Cross-functional</span> — <span class="hl">engineering</span>, product, operations. <span class="hl">Risk</span> management. <span class="hl">Supplier</span> coordination. <span class="hl">Resource allocation</span>. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Project Lead / Technical Lead</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Denmark &amp; Sweden, Multi-Partner</div>
  <ul>
    <li><span class="hl">LEGO</span> (<span class="hl">Denmark</span>): <span class="hl">Multi-partner</span> <span class="hl">project delivery</span> — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Supplier</span>/partner management. Development, <span class="hl">testing</span>, delivery. 8–10 engineers. <span class="hl">FAT</span>s, <span class="hl">inspection</span>s. <span class="hl">Client</span>/<span class="hl">stakeholder</span> interface. Copenhagen 2016–18.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Multi-partner</span> coordination. <span class="hl">Compliance</span>/<span class="hl">quality</span>. <span class="hl">Scope</span> definition for subcontractors. <span class="hl">Resource allocation</span>. <span class="hl">Risk</span> management. <span class="hl">Agile</span>.</li>
    <li>Technical interface — development, <span class="hl">engineering</span>, execution. <span class="hl">Procurement</span> support. <span class="hl">Budget</span> monitoring. Audit-ready reporting. <span class="hl">Cross-functional</span>.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Project Lead / Technical Lead</div>
  <div class="job-sub">2008 – 2016 | Banking &amp; Enterprise — Regulated, Multi-Partner Delivery</div>
  <ul>
    <li><span class="hl">Project delivery</span> in regulated environments — <span class="hl">scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Multi-partner</span>. <span class="hl">Regulatory</span> <span class="hl">compliance</span>/<span class="hl">quality</span>. <span class="hl">FAT</span>s, <span class="hl">inspection</span>s, <span class="hl">quality</span> gates. Audit-ready.</li>
    <li>Led 10–15 engineers. <span class="hl">Resource allocation</span>. <span class="hl">Supplier</span> management. <span class="hl">Client</span> interface. <span class="hl">Cross-functional</span>. <span class="hl">Agile</span> &amp; Waterfall. <span class="hl">Six Sigma</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Project delivery</span> — <span class="hl">IKEA</span> (30+ markets), <span class="hl">Truecaller</span> (300M+ users), <span class="hl">LEGO</span> (<span class="hl">Denmark</span>). <span class="hl">Scope</span>, <span class="hl">schedule</span>, <span class="hl">budget</span>, <span class="hl">risk</span>. <span class="hl">Multi-partner</span>. On-time delivery.</li>
    <li><span class="hl">Supplier</span>/<span class="hl">stakeholder</span> management — vendors, partners, subcontractors. <span class="hl">Client</span> interface. <span class="hl">Procurement</span> (scope of work). <span class="hl">Cross-functional</span>.</li>
    <li><span class="hl">Compliance</span>/<span class="hl">quality</span> — <span class="hl">FAT</span>s, <span class="hl">inspection</span>s, <span class="hl">quality</span> readiness. <span class="hl">Regulatory</span> <span class="hl">compliance</span>. <span class="hl">Risk</span> management. Audit-ready.</li>
    <li><span class="hl">Denmark</span> — <span class="hl">LEGO</span> (Copenhagen 2016–18). International, collaborative. Danish — Conversational.</li>
    <li>Technical — development/<span class="hl">engineering</span>/execution interface. Cloud (GCP, AWS). <span class="hl">Six Sigma</span> GB. <span class="hl">ITIL</span>.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>Six Sigma Green Belt<br>ITIL Foundation<br>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent | Danish — Conversational | Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
