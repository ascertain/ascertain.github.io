from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Test_Automation_Engineer_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Test_Automation_Engineer_Resume.doc"

ACCENT = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "15+ years",
    "banking",
    "fintech",
    "financial",
    "Playwright",
    "TypeScript",
    "test automation",
    "automation strategy",
    "automation framework",
    "continuous delivery",
    "CI/CD",
    "API testing",
    "E2E",
    "UI testing",
    "performance testing",
    "Scrum",
    "Scrum Master",
    "agile",
    "cross-functional",
    "exploratory testing",
    "web application",
    "web portal",
    "Selenium",
    "Appium",
    "shift-left",
    "quality assurance",
    "test strategy",
    "test plan",
    "mentoring",
    "code review",
    "GCP",
    "Docker",
    "Git",
    "core banking",
    "Finacle",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Senior Test Automation Engineer | Playwright & TypeScript | Banking & Fintech",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif",
        size=9.5,
        color=TEXT_MUTED,
    )
    add_text(
        contact,
        "\nPortfolio: ascertain.github.io/kashif/  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="7F8FA6", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Senior Test Automation Engineer with 15+ years of experience in banking, fintech, and "
        "enterprise web applications. Expert in Playwright (TypeScript) with a proven track record "
        "of defining and implementing test automation strategies — including API, E2E UI, and "
        "performance testing — within continuous delivery pipelines. Strong background in the "
        "financial sector, having delivered test automation for core banking (Finacle), payment "
        "systems, and business web portals. Collaborative team player thriving in agile environments "
        "with cross-functional teams, driving shift-left quality practices and engaging in "
        "exploratory testing. Experienced Scrum practitioner with interest in the Scrum Master role. "
        "Author of technical blog posts on Playwright migration and modern test frameworks.",
        size=10,
    )

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Test Automation: ",
            "Playwright (TypeScript), Selenium, Appium, Karate — E2E UI testing, API testing, "
            "performance testing, test automation framework design & architecture",
        ),
        (
            "Programming & Web: ",
            "TypeScript, C#, Java, Python — modern web development, framework architecture, "
            "code reviews, design patterns",
        ),
        (
            "Continuous Delivery: ",
            "CI/CD pipelines, GitHub Actions, Docker, Terraform, Git, continuous testing, "
            "shift-left quality integration, automated regression",
        ),
        (
            "Quality & Strategy: ",
            "Test automation strategy, test plans, quality assurance, exploratory testing, "
            "TDD/BDD, UAT, BAT, test coverage analysis, defect triage",
        ),
        (
            "Banking & Domain: ",
            "Core banking (Finacle), payment systems, financial web portals, e-commerce, "
            "government payment platforms, POS integration, transaction validation",
        ),
        (
            "Agile & Leadership: ",
            "Scrum/Scrum Master experience, agile delivery, cross-functional collaboration, "
            "team mentoring, test process improvement, stakeholder management",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")
    jobs = [
        (
            "IKEA IT AB, Malmö, Sweden",
            "Senior Test Automation Engineer / Software Engineer | Mar 2022 – Present",
            [
                "Define and implement test automation strategy for an innovative business web portal "
                "(VCS) — covering API testing, E2E UI testing, and performance testing within "
                "continuous delivery pipelines using Playwright (TypeScript) and GitHub Actions.",
                "Architect and maintain a scalable test automation framework — driving shift-left "
                "quality practices and integrating automated testing across the full CI/CD pipeline.",
                "Collaborate closely with cross-functional team members (developers, product managers, "
                "QA engineers) to drive and include the team in test plans and agile delivery.",
                "Lead exploratory testing efforts and define test automation strategies for new features "
                "across web, API, and real-time communication components.",
                "Serve as Scrum practitioner — facilitating agile ceremonies, driving sprint quality "
                "goals, and mentoring team members on testing best practices and code reviews.",
                "Built quality analytics on GCP (BigQuery, Pub/Sub) for test metrics and coverage "
                "tracking. Recognized as Exceptional Performer.",
            ],
        ),
        (
            "Truecaller, Sweden",
            "Release & Automation Engineer | Sep 2021 – Feb 2022",
            [
                "Led test automation for Android and iOS apps serving 300M+ users — defining test "
                "automation strategies across API, UI, and performance dimensions with CI/CD integration.",
                "Designed test plans for release readiness in agile sprints. Built automated suites "
                "using Appium and Selenium within continuous delivery pipelines.",
                "Collaborated cross-functionally on quality assurance, defect triage, and exploratory "
                "testing for release coordination.",
            ],
        ),
        (
            "LEGO and IKEA Group (via HCLTech), Denmark & Sweden",
            "Senior Test Automation Engineer / Test Lead | 2016 – 2021",
            [
                "Led test automation framework architecture for LEGO e-commerce web portal and IKEA "
                "App — Selenium, Appium, and Karate (TypeScript/Java) with full CI/CD integration.",
                "Defined and implemented test automation strategies — API testing, E2E UI testing, and "
                "performance testing — driving shift-left quality and agile delivery practices.",
                "Test lead managing 8+ engineers in Scrum teams — mentoring on automation best practices, "
                "code reviews, test plans, and cross-functional collaboration across global teams.",
            ],
        ),
        (
            "HCLTech / Enterprise Programs",
            "Automation Lead / SDET | Dec 2013 – 2014",
            [
                "Led test automation framework design across enterprise banking and fintech programmes — "
                "scalable frameworks with CI/CD pipelines, reducing manual effort by 70%.",
                "Mentored 15+ engineers on test automation architecture, agile practices, and "
                "quality-driven development across distributed teams.",
            ],
        ),
        (
            "Samin TekMindz India Pvt. Ltd. — Banking & Government",
            "Test Automation Engineer | 2011 – 2013",
            [
                "Designed test automation for the AMA government payment web portal (Ghana) — "
                "a complex banking/POS integration validating real-time financial transactions.",
                "Implemented test automation strategies across API and UI layers. Drove quality "
                "assurance and test process improvements for financial platforms.",
            ],
        ),
        (
            "Earlier Career — Core Banking & Financial Services",
            "Test Automation Engineer / QA | 2008 – 2011",
            [
                "Built test automation frameworks for core banking solutions (Finacle) — validating "
                "real-time transaction systems, ATM networks, POS terminals, and banking web portals.",
                "Delivered API and UI test automation for banking and financial applications. Led "
                "UAT/BAT cycles ensuring quality assurance across payment and transaction workflows.",
            ],
        ),
    ]
    for company, title, bullets in jobs:
        cp = document.add_paragraph()
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after = Pt(0)
        add_text(cp, company, bold=True, size=10, color=SECTION_COLOR)

        tp = document.add_paragraph()
        tp.paragraph_format.space_after = Pt(0)
        add_text(tp, title, bold=True, size=10)

        for bullet in bullets:
            add_highlighted_bullet(document, bullet)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Migrated test automation from Selenium to Playwright (TypeScript) — 3x reliability, 50% "
        "faster CI/CD pipelines. Authored blog post on Playwright migration: "
        "ascertain.github.io/2026/04/08/sunset-for-selenium-and-rise-for-playwright",
        "Defined and implemented test automation strategies across banking, fintech, and e-commerce "
        "web portals — API, E2E UI, and performance testing in continuous delivery.",
        "Delivered test automation for core banking (Finacle) — validating real-time financial "
        "transactions, ATM networks, POS systems, and payment platforms.",
        "Led and mentored teams of 8-15+ engineers in Scrum/agile environments — driving quality "
        "assurance culture and cross-functional collaboration.",
        "Drove shift-left testing adoption across engineering — reducing defect escape rate by 60% "
        "through early quality integration in agile sprints.",
    ]
    for achievement in achievements:
        add_highlighted_bullet(document, achievement)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\nGoogle Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\nITIL Foundation\nCertified Ethical Hacker (CEH)\n"
        "Six Sigma Green Belt\nUiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Senior Test Automation Engineer Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #1f4e79; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Senior Test Automation Engineer | Playwright &amp; TypeScript | Banking &amp; Fintech</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif<br>Portfolio: ascertain.github.io/kashif/  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Senior Test Automation Engineer with <span class="hl">15+ years</span> of experience in <span class="hl">banking</span>, <span class="hl">fintech</span>, and enterprise <span class="hl">web application</span>s. Expert in <span class="hl">Playwright</span> (<span class="hl">TypeScript</span>) with a proven track record of defining and implementing <span class="hl">test automation</span> strategies — including <span class="hl">API testing</span>, <span class="hl">E2E</span> <span class="hl">UI testing</span>, and <span class="hl">performance testing</span> — within <span class="hl">continuous delivery</span> pipelines. Strong background in the <span class="hl">financial</span> sector: <span class="hl">core banking</span> (<span class="hl">Finacle</span>), payment systems, and business <span class="hl">web portal</span>s. Collaborative team player in <span class="hl">agile</span> environments with <span class="hl">cross-functional</span> teams, driving <span class="hl">shift-left</span> quality and <span class="hl">exploratory testing</span>. <span class="hl">Scrum</span> practitioner with interest in <span class="hl">Scrum Master</span> role. Author of blog posts on <span class="hl">Playwright</span> migration.</p>

  <div class="section">Core Competencies</div>
  <p><b>Test Automation:</b> Playwright (TypeScript), Selenium, Appium, Karate — E2E UI, API, performance testing, framework design &amp; architecture<br>
  <b>Programming &amp; Web:</b> TypeScript, C#, Java, Python — modern web development, framework architecture, code reviews<br>
  <b>Continuous Delivery:</b> CI/CD pipelines, GitHub Actions, Docker, Terraform, Git, continuous testing, shift-left quality<br>
  <b>Quality &amp; Strategy:</b> Test automation strategy, test plans, quality assurance, exploratory testing, TDD/BDD, UAT, BAT<br>
  <b>Banking &amp; Domain:</b> Core banking (Finacle), payment systems, financial web portals, e-commerce, POS integration<br>
  <b>Agile &amp; Leadership:</b> Scrum/Scrum Master, agile delivery, cross-functional collaboration, mentoring, test process improvement</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö, Sweden</div>
  <div class="job-sub">Senior Test Automation Engineer / Software Engineer | Mar 2022 – Present</div>
  <ul>
    <li>Define and implement <span class="hl">test automation</span> strategy for a business <span class="hl">web portal</span> — <span class="hl">API testing</span>, <span class="hl">E2E</span> <span class="hl">UI testing</span>, and <span class="hl">performance testing</span> within <span class="hl">continuous delivery</span> pipelines using <span class="hl">Playwright</span> (<span class="hl">TypeScript</span>) and GitHub Actions.</li>
    <li>Architect and maintain a scalable <span class="hl">automation framework</span> — driving <span class="hl">shift-left</span> quality and integrating automated testing across <span class="hl">CI/CD</span>.</li>
    <li>Collaborate with <span class="hl">cross-functional</span> team members to drive <span class="hl">test plan</span>s and <span class="hl">agile</span> delivery. Lead <span class="hl">exploratory testing</span> for new features.</li>
    <li><span class="hl">Scrum</span> practitioner — facilitating ceremonies, driving sprint quality, <span class="hl">mentoring</span> on best practices and <span class="hl">code review</span>s.</li>
    <li>Built quality analytics on <span class="hl">GCP</span> (BigQuery, Pub/Sub) for test metrics. Recognized as Exceptional Performer.</li>
  </ul>

  <div class="job-title">Truecaller, Sweden</div>
  <div class="job-sub">Release &amp; Automation Engineer | Sep 2021 – Feb 2022</div>
  <ul>
    <li>Led <span class="hl">test automation</span> for apps serving 300M+ users — defining strategies across <span class="hl">API testing</span>, UI, and performance with <span class="hl">CI/CD</span>.</li>
    <li>Designed <span class="hl">test plan</span>s for release readiness in <span class="hl">agile</span> sprints. <span class="hl">Cross-functional</span> <span class="hl">quality assurance</span> and <span class="hl">exploratory testing</span>.</li>
  </ul>

  <div class="job-title">LEGO and IKEA Group (via HCLTech), Denmark &amp; Sweden</div>
  <div class="job-sub">Senior Test Automation Engineer / Test Lead | 2016 – 2021</div>
  <ul>
    <li>Led <span class="hl">automation framework</span> architecture for e-commerce <span class="hl">web portal</span> — <span class="hl">Selenium</span>, <span class="hl">Appium</span>, Karate with <span class="hl">CI/CD</span>. Defined <span class="hl">test automation</span> strategies for <span class="hl">API testing</span>, <span class="hl">E2E</span> UI, and performance.</li>
    <li>Test lead in <span class="hl">Scrum</span> teams managing 8+ engineers — <span class="hl">mentoring</span>, <span class="hl">code review</span>s, <span class="hl">test plan</span>s, <span class="hl">agile</span> delivery.</li>
  </ul>

  <div class="job-title">HCLTech / Enterprise Programs</div>
  <div class="job-sub">Automation Lead / SDET | Dec 2013 – 2014</div>
  <ul>
    <li>Led <span class="hl">automation framework</span> design for enterprise <span class="hl">banking</span> and <span class="hl">fintech</span> programmes — <span class="hl">CI/CD</span> pipelines, 70% manual effort reduction. Mentored 15+ engineers.</li>
  </ul>

  <div class="job-title">Samin TekMindz India Pvt. Ltd. — Banking &amp; Government</div>
  <div class="job-sub">Test Automation Engineer | 2011 – 2013</div>
  <ul>
    <li>Designed <span class="hl">test automation</span> for government payment <span class="hl">web portal</span> (AMA, Ghana) — <span class="hl">banking</span>/POS integration validating real-time <span class="hl">financial</span> transactions.</li>
    <li>Implemented <span class="hl">automation strategy</span> across API and UI layers. Drove <span class="hl">quality assurance</span> for <span class="hl">financial</span> platforms.</li>
  </ul>

  <div class="job-title">Earlier Career — Core Banking &amp; Financial Services</div>
  <div class="job-sub">Test Automation Engineer / QA | 2008 – 2011</div>
  <ul>
    <li>Built <span class="hl">automation framework</span>s for <span class="hl">core banking</span> (<span class="hl">Finacle</span>) — real-time transactions, ATM networks, POS terminals, <span class="hl">banking</span> <span class="hl">web portal</span>s.</li>
    <li>Delivered API and UI <span class="hl">test automation</span> for <span class="hl">banking</span> applications. Led UAT/BAT ensuring <span class="hl">quality assurance</span> across payment workflows.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li>Migrated <span class="hl">Selenium</span> to <span class="hl">Playwright</span> (<span class="hl">TypeScript</span>) — 3x reliability, 50% faster <span class="hl">CI/CD</span>. Blog: ascertain.github.io/2026/04/08/sunset-for-selenium-and-rise-for-playwright</li>
    <li>Defined <span class="hl">test automation</span> strategies across <span class="hl">banking</span>, <span class="hl">fintech</span>, and e-commerce — <span class="hl">API testing</span>, <span class="hl">E2E</span> UI, performance in <span class="hl">continuous delivery</span>.</li>
    <li>Delivered <span class="hl">test automation</span> for <span class="hl">core banking</span> (<span class="hl">Finacle</span>) — real-time <span class="hl">financial</span> transactions, ATM networks, POS, payment platforms.</li>
    <li>Led teams of 8-15+ engineers in <span class="hl">Scrum</span>/<span class="hl">agile</span> — driving <span class="hl">quality assurance</span> culture and <span class="hl">cross-functional</span> collaboration.</li>
    <li>Drove <span class="hl">shift-left</span> testing — 60% reduction in defect escape rate through early quality in <span class="hl">agile</span> sprints.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>ITIL Foundation<br>Certified Ethical Hacker<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
