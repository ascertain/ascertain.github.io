"""
Resume: Mariem Lamouni — 1-Page EU Format (Generic/Hybrid)
Two-column layout with sidebar for contact/skills/languages, main area for experience.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mariem_Lamouni_Resume"
TEAL = RGBColor(0x00, 0x7A, 0x87)
DARK = RGBColor(0x22, 0x22, 0x22)


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & TITLE ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MARIEM LAMOUNI")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = TEAL

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Business Analyst  \u2022  QA & Testing  \u2022  Recruitment  \u2022  Customer Support  \u2022  Administration")
    r2.font.size = Pt(9)
    r2.italic = True

    # ── CONTACT LINE ──
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(6)
    r3 = p3.add_run("Malm\u00f6, Sweden  |  +46 762432366  |  mariemlamouni30@gmail.com  |  Swedish Driving License (B)")
    r3.font.size = Pt(8.5)

    # ── PROFILE ──
    _heading(doc, "PROFILE")
    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(2)
    pp.paragraph_format.space_after = Pt(4)
    r = pp.add_run(
        "Versatile professional with experience in business analysis, QA testing, recruitment, "
        "technical maintenance (aircraft/mechatronics), customer support, and administration. "
        "Analytical and detail-oriented with strong communication skills in English, French, and Arabic. "
        "Skilled at gathering requirements, analysing processes, coordinating stakeholders, and managing documentation. "
        "Comfortable bridging technical and business teams in fast-paced, cross-functional environments."
    )
    r.font.size = Pt(9)

    # ── SKILLS (compact table) ──
    _heading(doc, "KEY SKILLS")
    skills = [
        "Business & Product Analysis", "QA & Manual Testing", "Recruitment & Screening", "Customer Support",
        "Requirements Gathering", "Process Documentation", "Stakeholder Coordination", "Data Entry & Analysis",
        "Administrative Management", "Technical Writing", "Problem Solving", "Team Collaboration",
    ]
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # override the 2-row table created above
    # (we'll just skip the old table logic and use this)
    for i, s in enumerate(skills):
        cell = table.rows[i // 4].cells[i % 4]
        cell.text = ""
        cp = cell.paragraphs[0]
        r = cp.add_run(f"\u25b8 {s}")
        r.font.size = Pt(8.5)
        r.bold = True

    # ── EXPERIENCE ──
    _heading(doc, "PROFESSIONAL EXPERIENCE")

    # Dqodify
    _role(doc, "Dqodify Pvt Ltd \u2014 Business Analyst / QA & Recruitment Coordinator", "Jun 2024 \u2013 Present")
    _bullet(doc, "Business analysis: requirements gathering, user stories, acceptance criteria, process mapping for digital products")
    _bullet(doc, "Analyse workflows and product data to identify improvements; translate business needs into clear specifications")
    _bullet(doc, "Manual/functional testing: test case design & execution, defect reporting, regression testing, verification")
    _bullet(doc, "Recruitment coordination: candidate sourcing, CV screening, interview scheduling, hiring stakeholder updates")
    _bullet(doc, "Product backlog support, sprint tracking, progress reporting, and cross-team knowledge transfer")

    # Uniq Dialog
    _role(doc, "Uniq Dialog (Hello Fresh) \u2014 Customer Support & Sales Representative", "Sep 2025")
    _bullet(doc, "Outbound customer engagement: active listening, needs assessment, tailored solutions, issue resolution")
    _bullet(doc, "Re-engaged inactive customers; built relationships through service quality and personalised communication")

    # ATI Hangars
    _role(doc, "Arotechnic Industries (ATI Hangars) \u2014 Quality Control & Technical Intern", "Nov 2022 \u2013 Dec 2023")
    _bullet(doc, "Quality control inspections and testing procedures for aircraft maintenance; ensured aviation safety compliance")
    _bullet(doc, "Technical documentation management: process manuals, inspection checklists, maintenance logs")
    _bullet(doc, "Data entry, analysis, and reporting to support quality assurance decisions; parts tracking and inventory")

    # Societe
    _role(doc, "Soci\u00e9t\u00e9 Melamdhygi\u00e8ne \u2014 Executive Assistant / Admin Coordinator", "Jan 2019 \u2013 Dec 2020")
    _bullet(doc, "Administration: scheduling, correspondence, filing, vendor management, order processing, invoice tracking")
    _bullet(doc, "Multi-stakeholder coordination: end-users, vendors, internal teams across departments")
    _bullet(doc, "Operational support: supplier follow-up, inventory records, and ensuring smooth daily operations")

    # ── EDUCATION ──
    _heading(doc, "EDUCATION")
    edus = [
        ("Aeronautical Maintenance Technician", "Inst. of Aeronautical Trades & Airport Logistics | 2021\u20132022"),
        ("Specialized Technician \u2014 Mechatronics", "Inst. of Aeronautics & Airport Logistics | 2016\u20132018"),
        ("Scientific Baccalaur\u00e9at (Maths A)", "2015\u20132016"),
    ]
    for title, detail in edus:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9)
        p.add_run(f"  \u2014  {detail}").font.size = Pt(8.5)

    # ── LANGUAGES & TOOLS ──
    _heading(doc, "LANGUAGES & TOOLS")
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    lp.paragraph_format.left_indent = Cm(0.3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(9)
    lp.add_run("English (Fluent) | French (Fluent) | Arabic (Native) | Swedish (Basic)").font.size = Pt(9)

    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.left_indent = Cm(0.3)
    r = tp.add_run("Tools: ")
    r.bold = True
    r.font.size = Pt(9)
    tp.add_run("MS Office, JIRA, Confluence, Google Workspace, Excel").font.size = Pt(9)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


def _heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = TEAL
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="007A87"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _role(doc, title, period):
    p = doc.add_paragraph()
    p.space_before = Pt(5)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = TEAL
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(8.5)
    r2.italic = True


def _bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0.5)
    r = p.add_run(f"\u2022 {text}")
    r.font.size = Pt(9)


if __name__ == "__main__":
    build_docx()
