"""
Resume: QA Lead – Digital Banking (EXTERNAL)
Focus: QA strategy at program level, cross-team coordination, risk-based testing,
       stakeholder management, end-to-end testing, automation, banking/fintech domain.
NOTE: User has early-career banking experience (Finacle Core Banking, Morpho BAS 2FA integration).
      Highlight banking domain knowledge from early career + current enterprise QA leadership.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_QA_Lead_Digital_Banking_Resume"
BRAND = RGBColor(0x00, 0x2B, 0x5C)  # Banking navy


def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="002B5C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def bullet(doc, text, bold_prefix=None, indent=Cm(0.5), size=Pt(9.5)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    if bold_prefix:
        r = p.add_run(f"• {bold_prefix}: ")
        r.bold = True
        r.font.size = size
        r2 = p.add_run(text)
        r2.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND
    r2 = p.add_run(f"  |  {company}, {location}")
    r2.font.size = Pt(9.5)
    r3 = p.add_run(f"  |  {period}")
    r3.font.size = Pt(9)
    r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME & CONTACT ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9)

    # ── PROFESSIONAL SUMMARY ──
    add_heading_block(doc, "Professional Summary")
    summary = (
        "Quality Assurance Lead with 14+ years of experience orchestrating QA efforts across programs, projects, "
        "and multi-team product initiatives — including early-career banking/fintech domain experience (Finacle Core "
        "Banking, Morpho BAS 2FA integration). Proven ability to define and lead QA strategy at programme level, "
        "coordinate cross-team testing, and serve as the central QA point of contact for product, engineering, and "
        "business stakeholders. Strong in risk-based testing, end-to-end quality planning across customer journeys, "
        "and enabling teams to adopt automation and modern testing practices. Experienced leading without line "
        "authority — creating alignment, facilitating collaboration forums, and driving decisions through influence "
        "and transparency. Brings structured reporting, clear KPI tracking, and a continuous improvement mindset "
        "to every programme."
    )
    ps = doc.add_paragraph(summary)
    ps.paragraph_format.space_before = Pt(4)
    for run in ps.runs:
        run.font.size = Pt(9.5)

    # ── KEY COMPETENCIES ──
    add_heading_block(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Programme-Level QA Strategy", "Cross-Team Coordination & Alignment", "Risk-Based Testing & Coverage",
        "Stakeholder Management", "End-to-End Test Planning", "QA KPIs & Metrics Reporting",
        "Automation Enablement (Playwright)", "Agile / DevOps / CI/CD", "Defect Triage & Root Cause Analysis",
        "Banking / Fintech Domain", "Distributed Teams Leadership", "Continuous Process Improvement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"✓ {comp}")
        r.font.size = Pt(9)

    # ── TOOLS & KNOWLEDGE ──
    add_heading_block(doc, "Tools & Knowledge")
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(3)
    skills = (
        "QA Strategy: Test lifecycle management, programme-level planning, risk-based prioritisation, quality gates  •  "
        "Automation: Playwright, Cypress, Selenium, Pytest, Vitest; CI/CD integration (GitHub Actions, Jenkins)  •  "
        "Test Management: Jira, TestRail, HP ALM, Zephyr  •  "
        "Architecture Understanding: Microservices, event-driven, distributed systems, API integrations  •  "
        "Cloud & DevOps: GCP, AWS, Docker, Kubernetes, CI/CD pipelines  •  "
        "Reporting: Quality dashboards, defect trend analysis, KPI tracking (Grafana, Jira dashboards)  •  "
        "Banking Domain: Finacle Core Banking, Morpho BAS (biometric/2FA), banking integration services  •  "
        "Methodologies: Agile (Scrum, SAFe), DevOps, ISTQB test processes"
    )
    r = tp.add_run(skills)
    r.font.size = Pt(9)

    # ── PROFESSIONAL EXPERIENCE ──
    add_heading_block(doc, "Professional Experience")

    # --- Ingka Digital / IKEA ---
    role_header(doc, "QA Lead / Team Lead – Multi-Team Digital Platforms",
                "Ingka Digital (IKEA)", "Malmö, Sweden", "Mar 2022 – Present")
    bullet(doc, "Define and lead QA strategy at programme level across 5+ integrated digital platforms; plan and coordinate QA activities spanning multiple Agile teams and release trains.")
    bullet(doc, "Serve as the primary QA point of contact for product owners, engineering leads, and business stakeholders — providing transparency on quality status, risks, and mitigation actions.")
    bullet(doc, "Drive end-to-end test planning across systems and customer journeys; align testing scope, priorities, and approaches across concurrent initiatives and releases.")
    bullet(doc, "Oversee cross-team dependencies and ensure coordinated test execution; facilitate QA collaboration forums where independent teams align on shared quality goals.")
    bullet(doc, "Implement risk-based testing methodology: assess business impact, change risk, and complexity to ensure sufficient coverage at all levels (unit, integration, system, E2E).")
    bullet(doc, "Establish and track QA KPIs and performance metrics: defect density, test coverage, automation percentage, cycle time — presented in clear dashboards for stakeholder visibility.")
    bullet(doc, "Support adoption of automation and modern testing practices: champion Playwright and Cypress adoption, integrate automated suites into CI/CD pipelines (GitHub Actions).")
    bullet(doc, "Lead defect triage sessions and root cause analysis across teams; drive resolution priorities and ensure quality outcomes improve with each release.")
    bullet(doc, "Continuously improve QA processes, standards, and tooling; introduce structured test design techniques, standardise reporting, and streamline environment management.")
    bullet(doc, "Lead QA efforts across multiple teams without direct line responsibility — creating alignment through influence, facilitation, and transparent communication.")

    # --- Truecaller ---
    role_header(doc, "QA Lead – Platform Quality",
                "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullet(doc, "Coordinated QA activities across multiple feature teams in a fast-paced Agile/DevOps environment; served as central QA contact for cross-team alignment.")
    bullet(doc, "Defined risk-based test approaches for microservices platform; ensured E2E coverage across distributed systems and API integrations.")
    bullet(doc, "Drove automation adoption and CI/CD quality gates; established structured reporting on quality status and release readiness.")

    # --- HCLTech ---
    role_header(doc, "QA Lead / Test Manager – Enterprise Programmes",
                "HCLTech (for IKEA & LEGO)", "Denmark / Sweden", "2013 – 2021")
    bullet(doc, "Led QA strategy and coordination across large-scale programmes spanning 10+ integrated systems and multiple Agile teams — acting as central QA orchestrator for enterprise-wide releases.")
    bullet(doc, "Planned and coordinated QA activities across projects, teams, and releases; aligned testing scope and priorities with product, engineering, and business stakeholders.")
    bullet(doc, "Drove end-to-end test planning across complex system integrations and customer journeys — from frontend UX flows through backend services and data layers.")
    bullet(doc, "Managed cross-team dependencies: coordinated shared test environments, integration test schedules, and release calendars across distributed teams (onshore & offshore).")
    bullet(doc, "Implemented risk-based testing at programme level; defined coverage models based on business criticality, integration complexity, and change velocity.")
    bullet(doc, "Established QA KPIs and programme-level dashboards: defect trends, test progress, automation ROI, quality gate compliance — enabling data-driven release decisions.")
    bullet(doc, "Managed a distributed team of 8–12 test engineers; drove competence development, facilitated collaboration forums, and promoted modern QA practices.")
    bullet(doc, "Supported automation strategy adoption: guided teams on framework selection, integration with CI/CD pipelines (Jenkins, GitHub Actions), and sustainable test maintenance.")
    bullet(doc, "Led defect triage and root cause analysis sessions across teams; improved defect resolution time and reduced production escape rate by 35%.")

    # --- India (BANKING HIGHLIGHTED) ---
    role_header(doc, "QA Engineer / Test Lead – Banking & Enterprise Systems",
                "HCL Technologies / Marlabs / TekMindz", "India", "2008 – 2013")
    bullet(doc, "Finacle Core Banking Implementation", "Banking Domain")
    bullet(doc, "Led QA for Finacle Core Banking implementation projects — end-to-end testing of banking modules including accounts, transactions, loans, and customer onboarding workflows.")
    bullet(doc, "Morpho BAS 2FA Integration", "Banking Domain")
    bullet(doc, "Tested Morpho BAS (Biometric Authentication Service) second-factor authentication integration with multiple banking applications — validating security flows, biometric device integration, and transaction authentication.")
    bullet(doc, "Coordinated testing across banking system integrations: core banking ↔ channels (internet banking, mobile) ↔ authentication services ↔ payment gateways.")
    bullet(doc, "Developed comprehensive test strategies for banking regulatory compliance, data security, and transaction integrity across interconnected financial systems.")
    bullet(doc, "Gained deep understanding of banking domain: account lifecycle, transaction processing, authentication flows, and financial data validation using SQL.")
    bullet(doc, "Progressed from QA Engineer to Test Lead; built foundational expertise in test planning, defect lifecycle management, and cross-team coordination in complex programme environments.")

    # ── CERTIFICATIONS ──
    add_heading_block(doc, "Certifications")
    certs = [
        "ISTQB Certified Tester – Foundation Level",
        "Google Cloud Associate Cloud Engineer",
        "AWS Certified Cloud Practitioner",
        "Six Sigma Green Belt",
        "Certified Ethical Hacker (CEH)",
    ]
    for c in certs:
        bullet(doc, c, indent=Cm(0.3), size=Pt(9))

    # ── EDUCATION & LANGUAGES ──
    add_heading_block(doc, "Education & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(9.5)
    p.add_run(" – UP Technical University  |  ").font.size = Pt(9)
    r2 = p.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(9.5)
    p.add_run(" – IGNOU").font.size = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r3 = p2.add_run("Languages: ")
    r3.bold = True
    r3.font.size = Pt(9.5)
    p2.add_run("English (Fluent)  |  Swedish (Conversational)  |  Hindi/Urdu (Native)").font.size = Pt(9)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.5cm;color:#222}
    h1{text-align:center;color:#002B5C;font-size:18pt;margin-bottom:2px}
    .contact{text-align:center;font-size:9pt;margin-bottom:10px}
    h2{color:#002B5C;font-size:11pt;border-bottom:1px solid #002B5C;padding-bottom:2px;margin-top:12px}
    .role{font-weight:bold;color:#002B5C;font-size:10pt;margin-top:8px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:9.5pt;margin-bottom:2px}
    li b{color:#002B5C}
    .summary{font-size:9.5pt;margin-top:4px}
    .comp-table{width:100%;font-size:9pt;margin-top:4px}
    .comp-table td{padding:2px 6px}
    .tools{font-size:9pt;margin-top:4px}
    .certs li,.edu{font-size:9pt}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p class="summary">Quality Assurance Lead with 14+ years of experience orchestrating QA efforts across programs, projects,
and multi-team product initiatives — including early-career banking/fintech domain experience (Finacle Core
Banking, Morpho BAS 2FA integration). Proven ability to define and lead QA strategy at programme level,
coordinate cross-team testing, and serve as the central QA point of contact for product, engineering, and
business stakeholders. Strong in risk-based testing, end-to-end quality planning across customer journeys,
and enabling teams to adopt automation and modern testing practices. Experienced leading without line
authority — creating alignment, facilitating collaboration forums, and driving decisions through influence
and transparency. Brings structured reporting, clear KPI tracking, and a continuous improvement mindset
to every programme.</p>

<h2>KEY COMPETENCIES</h2>
<table class="comp-table">
<tr><td>✓ Programme-Level QA Strategy</td><td>✓ Cross-Team Coordination &amp; Alignment</td><td>✓ Risk-Based Testing &amp; Coverage</td></tr>
<tr><td>✓ Stakeholder Management</td><td>✓ End-to-End Test Planning</td><td>✓ QA KPIs &amp; Metrics Reporting</td></tr>
<tr><td>✓ Automation Enablement (Playwright)</td><td>✓ Agile / DevOps / CI/CD</td><td>✓ Defect Triage &amp; Root Cause Analysis</td></tr>
<tr><td>✓ Banking / Fintech Domain</td><td>✓ Distributed Teams Leadership</td><td>✓ Continuous Process Improvement</td></tr>
</table>

<h2>TOOLS &amp; KNOWLEDGE</h2>
<p class="tools">QA Strategy: Test lifecycle management, programme-level planning, risk-based prioritisation, quality gates &bull;
Automation: Playwright, Cypress, Selenium, Pytest, Vitest; CI/CD integration (GitHub Actions, Jenkins) &bull;
Test Management: Jira, TestRail, HP ALM, Zephyr &bull;
Architecture Understanding: Microservices, event-driven, distributed systems, API integrations &bull;
Cloud &amp; DevOps: GCP, AWS, Docker, Kubernetes, CI/CD pipelines &bull;
Reporting: Quality dashboards, defect trend analysis, KPI tracking (Grafana, Jira dashboards) &bull;
Banking Domain: Finacle Core Banking, Morpho BAS (biometric/2FA), banking integration services &bull;
Methodologies: Agile (Scrum, SAFe), DevOps, ISTQB test processes</p>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">QA Lead / Team Lead – Multi-Team Digital Platforms &nbsp;|&nbsp; Ingka Digital (IKEA), Malmö, Sweden &nbsp;|&nbsp; Mar 2022 – Present</p>
<ul>
<li>Define and lead QA strategy at programme level across 5+ integrated digital platforms; plan and coordinate QA activities spanning multiple Agile teams and release trains.</li>
<li>Serve as the primary QA point of contact for product owners, engineering leads, and business stakeholders — providing transparency on quality status, risks, and mitigation actions.</li>
<li>Drive end-to-end test planning across systems and customer journeys; align testing scope, priorities, and approaches across concurrent initiatives and releases.</li>
<li>Oversee cross-team dependencies and ensure coordinated test execution; facilitate QA collaboration forums where independent teams align on shared quality goals.</li>
<li>Implement risk-based testing methodology: assess business impact, change risk, and complexity to ensure sufficient coverage at all levels (unit, integration, system, E2E).</li>
<li>Establish and track QA KPIs and performance metrics: defect density, test coverage, automation percentage, cycle time — presented in clear dashboards for stakeholder visibility.</li>
<li>Support adoption of automation and modern testing practices: champion Playwright and Cypress adoption, integrate automated suites into CI/CD pipelines (GitHub Actions).</li>
<li>Lead defect triage sessions and root cause analysis across teams; drive resolution priorities and ensure quality outcomes improve with each release.</li>
<li>Continuously improve QA processes, standards, and tooling; introduce structured test design techniques, standardise reporting, and streamline environment management.</li>
<li>Lead QA efforts across multiple teams without direct line responsibility — creating alignment through influence, facilitation, and transparent communication.</li>
</ul>

<p class="role">QA Lead – Platform Quality &nbsp;|&nbsp; Truecaller, Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</p>
<ul>
<li>Coordinated QA activities across multiple feature teams in a fast-paced Agile/DevOps environment; served as central QA contact for cross-team alignment.</li>
<li>Defined risk-based test approaches for microservices platform; ensured E2E coverage across distributed systems and API integrations.</li>
<li>Drove automation adoption and CI/CD quality gates; established structured reporting on quality status and release readiness.</li>
</ul>

<p class="role">QA Lead / Test Manager – Enterprise Programmes &nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO), Denmark / Sweden &nbsp;|&nbsp; 2013 – 2021</p>
<ul>
<li>Led QA strategy and coordination across large-scale programmes spanning 10+ integrated systems and multiple Agile teams — acting as central QA orchestrator for enterprise-wide releases.</li>
<li>Planned and coordinated QA activities across projects, teams, and releases; aligned testing scope and priorities with product, engineering, and business stakeholders.</li>
<li>Drove end-to-end test planning across complex system integrations and customer journeys — from frontend UX flows through backend services and data layers.</li>
<li>Managed cross-team dependencies: coordinated shared test environments, integration test schedules, and release calendars across distributed teams (onshore &amp; offshore).</li>
<li>Implemented risk-based testing at programme level; defined coverage models based on business criticality, integration complexity, and change velocity.</li>
<li>Established QA KPIs and programme-level dashboards: defect trends, test progress, automation ROI, quality gate compliance — enabling data-driven release decisions.</li>
<li>Managed a distributed team of 8–12 test engineers; drove competence development, facilitated collaboration forums, and promoted modern QA practices.</li>
<li>Supported automation strategy adoption: guided teams on framework selection, integration with CI/CD pipelines (Jenkins, GitHub Actions), and sustainable test maintenance.</li>
<li>Led defect triage and root cause analysis sessions across teams; improved defect resolution time and reduced production escape rate by 35%.</li>
</ul>

<p class="role">QA Engineer / Test Lead – Banking &amp; Enterprise Systems &nbsp;|&nbsp; HCL Technologies / Marlabs / TekMindz, India &nbsp;|&nbsp; 2008 – 2013</p>
<ul>
<li><b>Banking Domain – Finacle Core Banking Implementation:</b> Led QA for Finacle Core Banking implementation projects — end-to-end testing of banking modules including accounts, transactions, loans, and customer onboarding workflows.</li>
<li><b>Banking Domain – Morpho BAS 2FA Integration:</b> Tested Morpho BAS (Biometric Authentication Service) second-factor authentication integration with multiple banking applications — validating security flows, biometric device integration, and transaction authentication.</li>
<li>Coordinated testing across banking system integrations: core banking ↔ channels (internet banking, mobile) ↔ authentication services ↔ payment gateways.</li>
<li>Developed comprehensive test strategies for banking regulatory compliance, data security, and transaction integrity across interconnected financial systems.</li>
<li>Gained deep understanding of banking domain: account lifecycle, transaction processing, authentication flows, and financial data validation using SQL.</li>
<li>Progressed from QA Engineer to Test Lead; built foundational expertise in test planning, defect lifecycle management, and cross-team coordination in complex programme environments.</li>
</ul>

<h2>CERTIFICATIONS</h2>
<ul class="certs">
<li>ISTQB Certified Tester – Foundation Level</li>
<li>Google Cloud Associate Cloud Engineer</li>
<li>AWS Certified Cloud Practitioner</li>
<li>Six Sigma Green Belt</li>
<li>Certified Ethical Hacker (CEH)</li>
</ul>

<h2>EDUCATION &amp; LANGUAGES</h2>
<p class="edu"><b>B.Tech Information Technology</b> – UP Technical University &nbsp;|&nbsp; <b>PGDOM</b> – IGNOU</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Swedish (Conversational) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
