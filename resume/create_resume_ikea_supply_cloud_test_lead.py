from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_IKEA_Supply_Cloud_Test_Lead_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_IKEA_Supply_Cloud_Test_Lead_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "cloud-native",
    "cloud transformation",
    "SaaS",
    "M3 Cloud",
    "Infor M3",
    "test strategy",
    "testing strategy",
    "testing framework",
    "test automation",
    "automation framework",
    "end-to-end",
    "E2E",
    "unit",
    "integration",
    "system",
    "performance",
    "security",
    "user acceptance",
    "UAT",
    "risk-based",
    "Playwright",
    "UIPath",
    "UiPath",
    "Zephyr",
    "API automation",
    "API",
    "containerized",
    "Docker",
    "Kubernetes",
    "CI/CD",
    "GitHub Actions",
    "cloud infrastructure",
    "GCP",
    "AWS",
    "Azure",
    "Terraform",
    "release management",
    "enterprise integration",
    "cross-functional",
    "cross-stakeholder",
    "stakeholder",
    "business outcomes",
    "compliance",
    "application security",
    "cloud security",
    "technical debt",
    "Agile",
    "Scrum",
    "IKEA",
    "LEGO",
    "Truecaller",
    "Supply",
    "Food",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Cloud Test Strategy Lead  |  Cloud-Native Testing, SaaS & Enterprise Integration  |  15+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Testing leadership professional with 15+ years and a proven track record in "
        "large-scale cloud transformation projects. Designs and implements cloud-native "
        "test strategies for SaaS architecture validation across all deployment phases "
        "(development, staging, production). Orchestrates multi-phase testing — unit, "
        "integration, system, performance, security, and user acceptance testing. Drives "
        "test automation using modern frameworks (Playwright, UiPath, Zephyr) and "
        "containerized testing environments (Docker, Kubernetes). Strong understanding "
        "of application security, cloud security architectures, and risk-based testing. "
        "Connects testing strategy with business outcomes. Coordinates cross-stakeholder "
        "testing across business, development, platform, and partner teams. Currently "
        "at IKEA (3+ years, Malmö) — deeply familiar with IKEA culture, values, and "
        "ways of working. Experience also at Truecaller and LEGO.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Cloud-Native Test Strategy: ",
            "Designed and implemented end-to-end testing frameworks for cloud SaaS "
            "architecture validation across development, staging, and production. "
            "Currently at IKEA — built cloud-native test automation with Playwright, "
            "API automation, CI/CD pipelines (GitHub Actions, Docker, Kubernetes, "
            "Terraform) serving 30+ markets across multiple deployment phases.",
        ),
        (
            "Multi-Phase Testing Orchestration: ",
            "Established comprehensive testing coverage across unit, integration, system, "
            "performance, security, and user acceptance testing phases. Ensured seamless "
            "progression through release management gates. Risk-based testing — prioritized "
            "business-critical flows while managing technical debt reduction.",
        ),
        (
            "Security & Compliance: ",
            "Application security and cloud security awareness — Certified Ethical Hacker "
            "(CEH), threat modeling, security testing integrated into CI/CD. Compliance "
            "testing experience. Understands security architectures and best practices "
            "for cloud-native environments.",
        ),
        (
            "Cross-Stakeholder Collaboration: ",
            "Coordinated testing activities across business stakeholders, product "
            "development teams (BAs, solution architects, developers, testers), and "
            "platform teams. Ensured unified quality delivery and stakeholder alignment. "
            "Experienced working with external partners and vendor relationships. "
            "Global, cross-functional teams across IKEA, Truecaller, and LEGO.",
        ),
        (
            "IKEA Insider: ",
            "3+ years at IKEA IT AB in Malmö — deeply understand IKEA culture, "
            "values, and ways of working. Familiar with IKEA Supply, cross-business "
            "coordination, and enterprise-scale delivery across markets and channels.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Technical Skills ──
    add_section_heading(document, "Technical Skills")
    skill_lines = [
        (
            "Test Automation: ",
            "Playwright, UiPath, Selenium, Appium, RestAssured, Karate, Zephyr, "
            "Postman, E2E, API automation, regression, contract testing",
        ),
        (
            "Cloud & Infrastructure: ",
            "GCP, AWS, Azure, Docker, Kubernetes, Terraform, containerized "
            "testing environments, cloud-native architecture, SaaS validation",
        ),
        (
            "CI/CD & Release Management: ",
            "GitHub Actions, Jenkins, pipeline design, multi-environment "
            "deployment (dev/staging/prod), release management, Git",
        ),
        (
            "Security & Compliance: ",
            "Application security, cloud security, CEH certified, threat "
            "modeling, compliance testing, IEC62443 awareness",
        ),
        (
            "Testing Coverage: ",
            "Unit, integration, system, performance (JMeter, K6), security, "
            "UAT, risk-based testing, technical debt management",
        ),
        (
            "Languages & Tools: ",
            "Java, C#, TypeScript, Python, SQL, TestRail, Jira, Confluence, "
            "Grafana, Maven, .NET",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Test Strategy Lead / Senior SDET", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Cloud-Native SaaS Platform, 30+ Markets", bold=True, size=10)
    ikea_bullets = [
        "Cloud-native test strategy — designed and implemented end-to-end testing "
        "frameworks for SaaS architecture validation across development, staging, and "
        "production. Multi-phase testing: unit, integration, system, performance, "
        "security, and UAT with seamless progression through release gates.",
        "Drove cloud-native test automation — Playwright E2E, API automation (RestAssured, "
        "Karate), containerized testing environments (Docker, Kubernetes). CI/CD pipelines "
        "(GitHub Actions, Terraform) serving 30+ markets across all deployment phases.",
        "Risk-based testing — prioritized business-critical flows, managed technical debt "
        "reduction. Application security and cloud security testing integrated into CI/CD. "
        "Compliance validation. Performance testing for APIs and services.",
        "Cross-stakeholder collaboration — coordinated testing across business stakeholders, "
        "product development teams (BAs, architects, developers, testers), and platform "
        "teams. Unified quality delivery and stakeholder alignment. Release management.",
        "Connected testing strategy with business outcomes — ensured quality delivery "
        "aligned with IKEA Supply and market needs. Mentored engineers. AI-driven testing "
        "(Claude, Copilot, Gemini) — 30% velocity improvement. Agile (Scrum).",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior QA Lead / Test Strategy Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Cloud SaaS — 300M+ Users, Global Scale", bold=True, size=10)
    tc_bullets = [
        "Cloud transformation testing at scale — test strategy for SaaS platform with "
        "300M+ users. End-to-end testing across unit, integration, performance, security, "
        "and UAT phases. Selenium, Appium (iOS/Android), RestAssured, API automation.",
        "Cross-stakeholder coordination — development, platform, and business teams. "
        "Risk-based testing for business-critical flows. Release management. Cloud "
        "infrastructure (AWS). CI/CD. Agile (Scrum/Kanban). Mentored engineers.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Test Lead / Senior SDET", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise — Cloud, SaaS, Multi-Platform", bold=True, size=10)
    lego_bullets = [
        "LEGO: Testing strategy for cloud e-commerce platform — end-to-end testing "
        "frameworks (Selenium, RestAssured, Karate, Appium). Unit, integration, system, "
        "performance, UAT. Led 8–10 engineers. Enterprise integration patterns. "
        "Release management. Cross-functional stakeholder coordination.",
        "IKEA App, Genesys, Verint (2018–2021): Cloud transformation testing — SaaS "
        "architecture validation, multi-phase testing, CI/CD. Cross-stakeholder "
        "collaboration across business and development teams. Risk-based testing. "
        "Security testing. Agile (Scrum).",
        "Connected testing strategy with business outcomes — ensured quality delivery "
        "across global markets and sales channels. Drove continuous improvement.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior SDET / Test Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech & Government — Enterprise, Multi-Platform", bold=True, size=10)
    fin_bullets = [
        "Enterprise testing strategy — automation frameworks (Selenium, RestAssured, "
        "TestNG). End-to-end testing across unit, integration, system, performance, "
        "UAT. Enterprise integration patterns (Finacle banking). CI/CD. Release "
        "management. SQL, Java, C#.",
        "Testing leadership — mentored 15+ engineers. Led teams of 10+. Cross-functional "
        "stakeholder coordination. Connected testing with business outcomes. "
        "Risk-based testing. Agile (Scrum/Kanban). Compliance testing.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Cloud-native test strategy — SaaS architecture validation across dev/staging/prod "
        "at IKEA (30+ markets) and Truecaller (300M+ users). Multi-phase testing orchestration.",
        "Test automation with modern frameworks — Playwright, UiPath, Selenium, Appium, "
        "RestAssured, Karate, Zephyr. Containerized testing (Docker, K8s). API automation.",
        "Risk-based testing and security — prioritized business-critical flows, managed "
        "technical debt. Application security, cloud security (CEH). Compliance testing.",
        "Cross-stakeholder coordination — business, development, platform, and partner "
        "teams. Connected testing strategy with business outcomes. Release management.",
        "IKEA insider (3+ years, Malmö) — deeply familiar with IKEA culture, values, "
        "cross-business coordination, and enterprise-scale delivery.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – IKEA Supply Cloud Test Strategy Lead Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Cloud Test Strategy Lead | Cloud-Native Testing, SaaS &amp; Enterprise Integration | 15+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p>Testing leadership with 15+ years and a proven track record in large-scale <span class="hl">cloud transformation</span> projects. Designs <span class="hl">cloud-native</span> <span class="hl">test strategy</span> for <span class="hl">SaaS</span> architecture validation across all deployment phases (dev, staging, prod). Multi-phase testing: <span class="hl">unit</span>, <span class="hl">integration</span>, <span class="hl">system</span>, <span class="hl">performance</span>, <span class="hl">security</span>, <span class="hl">UAT</span>. Modern <span class="hl">automation framework</span>s (<span class="hl">Playwright</span>, <span class="hl">UiPath</span>, <span class="hl">Zephyr</span>) and <span class="hl">containerized</span> testing (<span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). <span class="hl">Application security</span>, <span class="hl">cloud security</span>, <span class="hl">risk-based</span> testing. Connects testing with <span class="hl">business outcomes</span>. <span class="hl">Cross-stakeholder</span> coordination. Currently at <span class="hl">IKEA</span> (3+ years, Malmö). Also <span class="hl">Truecaller</span> and <span class="hl">LEGO</span>.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Cloud-Native Test Strategy:</b> <span class="hl">End-to-end</span> <span class="hl">testing framework</span>s for <span class="hl">SaaS</span> validation across dev/staging/prod. <span class="hl">Playwright</span>, <span class="hl">API automation</span>, <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Docker</span>, <span class="hl">Kubernetes</span>, <span class="hl">Terraform</span>). 30+ markets at <span class="hl">IKEA</span>.<br>
  <b>Multi-Phase Testing:</b> <span class="hl">Unit</span>, <span class="hl">integration</span>, <span class="hl">system</span>, <span class="hl">performance</span>, <span class="hl">security</span>, <span class="hl">UAT</span>. <span class="hl">Risk-based</span> — business-critical flows. <span class="hl">Technical debt</span> reduction. <span class="hl">Release management</span>.<br>
  <b>Security &amp; Compliance:</b> <span class="hl">Application security</span>, <span class="hl">cloud security</span> (CEH). <span class="hl">Compliance</span> testing. Threat modeling.<br>
  <b>Cross-Stakeholder:</b> Business <span class="hl">stakeholder</span>s, dev teams (BAs, architects, devs, testers), platform teams, external partners. Unified quality delivery.<br>
  <b>IKEA Insider:</b> 3+ years at IKEA IT AB, Malmö. IKEA culture, values, <span class="hl">Supply</span>, cross-business coordination.</p>

  <div class="section">Technical Skills</div>
  <p><b>Test Automation:</b> Playwright, UiPath, Selenium, Appium, RestAssured, Karate, Zephyr, Postman, E2E, API automation<br>
  <b>Cloud &amp; Infra:</b> GCP, AWS, Azure, Docker, Kubernetes, Terraform, containerized testing, SaaS validation<br>
  <b>CI/CD &amp; Release:</b> GitHub Actions, Jenkins, multi-environment deployment, release management, Git<br>
  <b>Security:</b> Application security, cloud security, CEH, threat modeling, compliance testing<br>
  <b>Testing Coverage:</b> Unit, integration, system, performance (JMeter, K6), security, UAT, risk-based<br>
  <b>Languages &amp; Tools:</b> Java, C#, TypeScript, Python, SQL, TestRail, Jira, Confluence, Grafana</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Test Strategy Lead / Senior SDET</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Cloud-Native SaaS Platform, 30+ Markets</div>
  <ul>
    <li><span class="hl">Cloud-native</span> <span class="hl">test strategy</span> — <span class="hl">end-to-end</span> <span class="hl">testing framework</span>s for <span class="hl">SaaS</span> validation across dev/staging/prod. Multi-phase: <span class="hl">unit</span>, <span class="hl">integration</span>, <span class="hl">system</span>, <span class="hl">performance</span>, <span class="hl">security</span>, <span class="hl">UAT</span>. 30+ markets.</li>
    <li><span class="hl">Test automation</span> — <span class="hl">Playwright</span>, <span class="hl">API automation</span> (RestAssured, Karate), <span class="hl">containerized</span> testing (<span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Terraform</span>).</li>
    <li><span class="hl">Risk-based</span> testing — business-critical flows, <span class="hl">technical debt</span>. <span class="hl">Application security</span>, <span class="hl">cloud security</span>. <span class="hl">Release management</span>. AI-driven testing (30% velocity).</li>
    <li><span class="hl">Cross-stakeholder</span> — business, dev, platform teams. <span class="hl">Business outcomes</span> alignment. Mentored engineers. <span class="hl">Agile</span> (<span class="hl">Scrum</span>).</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior QA Lead / Test Strategy Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Cloud SaaS — 300M+ Users, Global Scale</div>
  <ul>
    <li><span class="hl">Cloud transformation</span> testing — <span class="hl">SaaS</span> platform, 300M+ users. <span class="hl">End-to-end</span> across <span class="hl">unit</span>, <span class="hl">integration</span>, <span class="hl">performance</span>, <span class="hl">security</span>, <span class="hl">UAT</span>. Selenium, Appium, RestAssured, <span class="hl">API automation</span>.</li>
    <li><span class="hl">Cross-stakeholder</span>. <span class="hl">Risk-based</span>. <span class="hl">Release management</span>. <span class="hl">Cloud infrastructure</span> (<span class="hl">AWS</span>). <span class="hl">CI/CD</span>. <span class="hl">Agile</span>. Mentored engineers.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Test Lead / Senior SDET</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise — Cloud, SaaS, Multi-Platform</div>
  <ul>
    <li><span class="hl">LEGO</span>: <span class="hl">Testing strategy</span> for cloud e-commerce — <span class="hl">end-to-end</span> <span class="hl">testing framework</span>s (Selenium, RestAssured, Karate, Appium). <span class="hl">Unit</span>, <span class="hl">integration</span>, <span class="hl">system</span>, <span class="hl">performance</span>, <span class="hl">UAT</span>. 8–10 engineers. <span class="hl">Enterprise integration</span>. <span class="hl">Release management</span>.</li>
    <li><span class="hl">IKEA</span> (2018–21): <span class="hl">Cloud transformation</span> — <span class="hl">SaaS</span> validation, multi-phase testing, <span class="hl">CI/CD</span>. <span class="hl">Cross-stakeholder</span>. <span class="hl">Risk-based</span>. <span class="hl">Security</span> testing. <span class="hl">Agile</span>.</li>
    <li>Connected <span class="hl">testing strategy</span> with <span class="hl">business outcomes</span> across global markets. Continuous improvement.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior SDET / Test Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech &amp; Government — Enterprise, Multi-Platform</div>
  <ul>
    <li><span class="hl">Enterprise integration</span> testing — <span class="hl">automation framework</span>s (Selenium, RestAssured). <span class="hl">End-to-end</span>: <span class="hl">unit</span>, <span class="hl">integration</span>, <span class="hl">system</span>, <span class="hl">performance</span>, <span class="hl">UAT</span>. <span class="hl">CI/CD</span>. <span class="hl">Release management</span>. <span class="hl">Compliance</span>.</li>
    <li>Testing <span class="hl">leadership</span> — 15+ engineers. <span class="hl">Cross-functional</span> <span class="hl">stakeholder</span> coordination. <span class="hl">Business outcomes</span>. <span class="hl">Risk-based</span>. <span class="hl">Agile</span>.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Cloud-native</span> <span class="hl">test strategy</span> — <span class="hl">SaaS</span> validation across dev/staging/prod at <span class="hl">IKEA</span> (30+ markets) and <span class="hl">Truecaller</span> (300M+ users).</li>
    <li>Modern <span class="hl">automation framework</span>s — <span class="hl">Playwright</span>, <span class="hl">UiPath</span>, <span class="hl">Zephyr</span>, <span class="hl">containerized</span> (<span class="hl">Docker</span>, <span class="hl">Kubernetes</span>). <span class="hl">API automation</span>.</li>
    <li><span class="hl">Risk-based</span> testing &amp; <span class="hl">security</span> — <span class="hl">application security</span>, <span class="hl">cloud security</span> (CEH). <span class="hl">Compliance</span>. <span class="hl">Technical debt</span> management.</li>
    <li><span class="hl">Cross-stakeholder</span> — business, dev, platform, partner teams. <span class="hl">Business outcomes</span>. <span class="hl">Release management</span>.</li>
    <li><span class="hl">IKEA</span> insider (3+ years, Malmö) — culture, values, <span class="hl">Supply</span>, enterprise-scale delivery.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA Certified</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
