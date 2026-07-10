"""
Generate 3 LinkedIn-style profiles as a formatted DOCX document.
1. Project & Technical Lead/Manager
2. Test (Manual & Automation) Lead/Manager
3. DevOps & Cloud Engineer
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_3_Profiles.docx")

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(14); p.space_after = Pt(4)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(12), color=RGBColor(0x1F,0x47,0x88))
    return p

def add_field(doc, label, value, value_bold=False, value_size=Pt(10)):
    p = doc.add_paragraph(); p.space_after = Pt(4)
    add_text(p, label, bold=True, size=Pt(10))
    add_text(p, value, bold=value_bold, size=value_size)
    return p

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(1.0); section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5); section.right_margin = Cm(1.5)

    # Title
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "MOHAMMAD KASHIF — PROFESSIONAL PROFILES", bold=True, size=Pt(14), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(10)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ═══════════════════════════════════════════════════════════════════════
    # PROFILE 1
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "PROFILE 1 — PROJECT & TECHNICAL LEAD / MANAGER")

    add_field(doc, "Headline: ", "Project & Technical Lead | Cloud & Enterprise Delivery | Nordic IT Operations | Working in IT since 2008")

    p = doc.add_paragraph(); p.space_before = Pt(6); p.space_after = Pt(2)
    add_text(p, "About:", bold=True, size=Pt(10))

    about1 = (
        "Project and Technical Lead with a broad background spanning cloud operations, enterprise platforms, "
        "and cross-functional delivery across Nordic and global organizations. Since 2008, I have led complex "
        "IT projects end-to-end — from planning, coordination, and stakeholder alignment through go-live and "
        "stabilization. My strength lies in bringing structure, clarity, and momentum to technical delivery work, "
        "translating complex realities into actionable plans, and keeping teams aligned toward outcomes.\n\n"
        "I have managed delivery across cloud platforms (GCP, Azure, AWS), enterprise integrations, and regulated "
        "environments — coordinating engineers, vendors, and business stakeholders to ship high-quality solutions "
        "on time. Experienced in agile, SAFe, ITIL, and hybrid methodologies. Strong communicator in English and Swedish."
    )
    p = doc.add_paragraph(); p.space_after = Pt(4)
    add_text(p, about1, size=Pt(10))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Core Areas: ", bold=True, size=Pt(9.5))
    add_text(p, "Cloud project delivery, stakeholder management, cross-functional coordination, risk & escalation management, go-live governance, ITSM/ITIL, continuous improvement, agile/SAFe/Prince2.", size=Pt(9.5))

    # ═══════════════════════════════════════════════════════════════════════
    # PROFILE 2
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "PROFILE 2 — TEST (MANUAL & AUTOMATION) LEAD / MANAGER")

    add_field(doc, "Headline: ", "Test Lead & Manager | Manual & Automation Testing | Quality Strategy & Delivery | Working in IT since 2008")

    p = doc.add_paragraph(); p.space_before = Pt(6); p.space_after = Pt(2)
    add_text(p, "About:", bold=True, size=Pt(10))

    about2 = (
        "Test Lead and Manager with comprehensive experience across manual testing, test automation, and quality "
        "strategy — working in IT since 2008 across enterprise, cloud, IoT, and financial domains. I lead test "
        "planning, estimation, coordination, execution, and reporting for complex programs involving multi-tier "
        "architectures, cloud platforms, and regulated environments.\n\n"
        "I build and grow test teams, define test strategies that balance exploratory and automated approaches, "
        "and drive continuous improvement in test processes and tooling. Proficient in Python (pytest), Selenium, "
        "Playwright, and CI/CD pipeline integration. Experienced in managing testing across Azure, GCP, and hybrid "
        "environments, with domain expertise in financial services, retail, and healthcare-adjacent systems."
    )
    p = doc.add_paragraph(); p.space_after = Pt(4)
    add_text(p, about2, size=Pt(10))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Core Areas: ", bold=True, size=Pt(9.5))
    add_text(p, "Test management (planning, estimation, coordination, reporting, execution), manual & exploratory testing, test automation (Python/pytest/Selenium/Playwright), CI/CD integration, defect management, stakeholder communication, agile/SAFe, ISTQB.", size=Pt(9.5))

    # ═══════════════════════════════════════════════════════════════════════
    # PROFILE 3
    # ═══════════════════════════════════════════════════════════════════════
    add_section_heading(doc, "PROFILE 3 — DEVOPS & CLOUD ENGINEER")

    add_field(doc, "Headline: ", "DevOps & Cloud Engineer | GCP, Azure, AWS | CI/CD & Infrastructure Automation | Working in IT since 2008")

    p = doc.add_paragraph(); p.space_before = Pt(6); p.space_after = Pt(2)
    add_text(p, "About:", bold=True, size=Pt(10))

    about3 = (
        "DevOps and Cloud Engineer focused on building, automating, and operating scalable cloud infrastructure "
        "— working in IT since 2008, with dedicated cloud and DevOps focus since the mid-2010s. I design and "
        "maintain CI/CD pipelines, infrastructure-as-code deployments, and cloud-native architectures across "
        "GCP, Azure, and AWS.\n\n"
        "Hands-on with Terraform, Kubernetes, Docker, GitHub Actions, Jenkins, and cloud-native services. "
        "Experienced in building reliable, observable systems — from automated deployments and monitoring to "
        "incident response and operational stability. I work closely with development teams to streamline "
        "delivery, reduce toil, and ensure production systems are secure, scalable, and resilient."
    )
    p = doc.add_paragraph(); p.space_after = Pt(4)
    add_text(p, about3, size=Pt(10))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Core Areas: ", bold=True, size=Pt(9.5))
    add_text(p, "Cloud infrastructure (GCP, Azure, AWS), CI/CD pipelines (GitHub Actions, Jenkins), Infrastructure as Code (Terraform), Kubernetes & Docker, system reliability, monitoring & observability, Linux, networking, security & compliance, Python/shell scripting.", size=Pt(9.5))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

if __name__ == "__main__":
    build_docx()
