from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Release_Manager_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Release_Manager_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "release management",
    "release plan",
    "release",
    "branching",
    "version",
    "configuration management",
    "Git",
    "CI/CD",
    "deployment",
    "stakeholder",
    "PO",
    "project manager",
    "systems engineer",
    "development team",
    "conflict",
    "risk",
    "documentation",
    "guidelines",
    "software development",
    "fast-paced",
    "initiative",
    "organized",
    "communication",
    "GitHub Actions",
    "Jenkins",
    "Terraform",
    "Docker",
    "pipeline",
    "feature flag",
    "go/no-go",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Release Manager  |  Release Planning, Branching Strategy & Configuration Management  |  10+ Years",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Release Manager with 10+ years of software development experience and a strong "
        "background in release management, branching strategies, and configuration management. "
        "Interface between development plans and customer project needs. Maintains clear "
        "communication with stakeholders — POs, development teams, project managers, and "
        "systems engineers. Monitors and reports status and changes of release plans. "
        "Resolves conflicts within release planning. Detects and manages risks in the "
        "delivery process. Ensures releases are done in accordance with technical and "
        "administrative routines and guidelines. Handles product release documentation. "
        "Solid understanding of branching models (Git Flow, trunk-based) and version management. "
        "Highly organized, efficient, and able to work in a fast-paced environment. "
        "Takes own initiative — works both independently and in team. Built and maintained "
        "CI/CD pipelines (GitHub Actions, Jenkins). Agile (Scrum/SAFe) and Waterfall.",
        size=10,
    )

    # ── Core Competencies ──
    add_section_heading(document, "Core Competencies")
    skill_lines = [
        (
            "Release Management: ",
            "Release planning, scheduling, coordination. Release plan monitoring "
            "and reporting. Conflict resolution within release plans. Go/no-go "
            "decisions. Release documentation. Technical and administrative "
            "guidelines compliance. Feature flag management and controlled rollouts",
        ),
        (
            "Branching & Version Management: ",
            "Git Flow, trunk-based development, release branching, hotfix branching. "
            "Version management and tagging. Merge conflict resolution. Branch "
            "protection rules. Configuration management. Repository management",
        ),
        (
            "CI/CD & DevOps: ",
            "GitHub Actions, Jenkins. Build, test, and deployment pipelines. "
            "Terraform, Docker, Kubernetes. Infrastructure as Code. GCP (Cloud Run, "
            "BigQuery), AWS. Automated release workflows. Environment management",
        ),
        (
            "Risk & Stakeholder Management: ",
            "Risk detection and management in delivery. Stakeholder communication "
            "(PO, dev teams, project managers, systems engineers). Status reporting. "
            "Cross-functional coordination. Conflict resolution",
        ),
        (
            "Software Development: ",
            "Python, TypeScript, C#, Java, Node.js. REST APIs. SQL. Git. "
            "Test automation (Playwright, Selenium). Agile (Scrum/SAFe), Waterfall. "
            "Full SDLC understanding",
        ),
        (
            "Tools: ",
            "Git, GitHub, GitHub Actions, Jenkins. Jira, Confluence. Docker, "
            "Terraform. Grafana (dashboards, reporting). TestRail. MS Project",
        ),
    ]
    for label, value in skill_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Team Lead / Release & Delivery Lead", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS Platform — Release Management, CI/CD, 30+ Global Markets", bold=True, size=10)
    ikea_bullets = [
        "Release management — interface between development plans and customer project "
        "needs across 30+ global markets. Monitored and reported status and changes of "
        "release plans to stakeholders (POs, development teams, project managers). "
        "Resolved conflicts within release planning. Go/no-go decisions. Ensured releases "
        "in accordance with technical and administrative routines and guidelines.",
        "Branching strategy & version management — defined and maintained branching "
        "strategy (Git Flow, feature branches, release branches, hotfix branches). "
        "Version tagging and release documentation. Branch protection rules. Merge "
        "conflict resolution. Configuration management across environments.",
        "Risk management — detected and managed risks in the delivery process. "
        "Communicated risks and mitigation plans to stakeholders. Tracked dependencies "
        "across teams and releases. Maintained clear visibility of release status.",
        "CI/CD & automation — built and maintained CI/CD pipelines (GitHub Actions). "
        "Automated release workflows. Terraform, Docker on GCP (Cloud Run, BigQuery). "
        "Feature flag management for controlled rollouts. AI-assisted — 30% velocity "
        "improvement. Led Selenium-to-Playwright migration (3x faster, 50% CI reduction).",
        "Team leadership — managed engineers and consultants. Coaching, mentoring. "
        "Coordinated with external SaaS vendor on release alignment. Budget and resource "
        "management. Fast-paced, initiative-driven. Recognized as Exceptional Performer.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Release & Automation Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Communication Platform — 300M+ Users, Release-Driven Delivery", bold=True, size=10)
    tc_bullets = [
        "Release management at global scale — owned release planning, readiness, and "
        "coordination for 300M+ users. Go/no-go decisions. Release documentation. "
        "Monitored release plans and reported status to stakeholders. Resolved conflicts "
        "across teams. Risk detection and management. Branching and version management.",
        "Release automation — built tools and automated workflows to streamline release "
        "processes. CI pipelines, build systems, deployment automation. Feature flag "
        "management — controlled rollouts, risk analysis, data-driven decisions. "
        "Cross-functional coordination (engineering, product, operations). AWS. Agile.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "HCLTech — LEGO & IKEA Group, Denmark & Sweden — Technical Lead", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2013 – 2021  |  E-Commerce & Enterprise — Release Coordination, Multi-Partner", bold=True, size=10)
    lego_bullets = [
        "LEGO & IKEA (2017–2021): Release coordination across multi-partner delivery. "
        "Interface between development plans and project needs. Managed release planning, "
        "scheduling, and status reporting. Resolved conflicts across teams and partners. "
        "Risk management. Branching strategy and version management (Git). Release "
        "documentation and guidelines. Led teams of 8–10 across onshore-offshore. "
        "Quality gate decisions. Agile and Waterfall.",
        "Technical Lead (2013–2017): Software development and CI/CD. Built automation "
        "frameworks (Selenium, C#, NUnit). Branching models and version management. "
        "Configuration management. Release documentation. Full SDLC. Worked both "
        "independently and in team. Fast-paced, initiative-driven. Mentoring.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "Banking & Enterprise — Technical Lead / Consultant", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2013  |  Core Banking (Finacle CBS) — Releases, Migrations, Configuration", bold=True, size=10)
    fin_bullets = [
        "Release and configuration management — managed software releases for core "
        "banking platform (Finacle CBS). Release planning, documentation, and guidelines "
        "compliance. Post go-live stabilization. Configuration management across "
        "environments. Data migration releases. Stakeholder coordination (development "
        "teams, project managers, business). Risk detection. Fast-paced, regulated.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "Release management ownership — IKEA (30+ markets), Truecaller (300M+ users), "
        "LEGO, banking. Interface between development and customer needs. Release planning, "
        "status reporting, conflict resolution, risk management, documentation.",
        "Branching strategy & CI/CD — defined branching models (Git Flow, trunk-based). "
        "Built CI/CD pipelines (GitHub Actions, Jenkins). Automated release workflows. "
        "Feature flag management for controlled rollouts. 3x faster execution, 50% CI reduction.",
        "Configuration management — managed configurations across environments, branches, "
        "and releases. Version tagging. Release documentation. Technical guidelines.",
        "Stakeholder communication — POs, development teams, project managers, systems "
        "engineers, external vendors. Clear status reporting, risk surfacing, conflict "
        "resolution. Organized, efficient, initiative-driven.",
        "Scaled IKEA platform from 2K to 50K usage across 30+ markets — end-to-end "
        "release and delivery ownership. Recognized as Exceptional Performer.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "ISTQB Certified Tester Foundation\n"
        "ITIL Foundation\n"
        "Six Sigma Green Belt\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\n"
        "Certified Ethical Hacker (CEH)\n"
        "UiPath RPA Certified",
        size=9.5,
    )

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Release Manager Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Release Manager | Release Planning, Branching Strategy &amp; Configuration Management | 10+ Years</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com | LinkedIn: linkedin.com/in/md-kashif | Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Release</span> Manager with 10+ years of <span class="hl">software development</span> experience and strong background in <span class="hl">release management</span>, <span class="hl">branching</span> strategies, and <span class="hl">configuration management</span>. Interface between development plans and customer project needs. Clear <span class="hl">communication</span> with <span class="hl">stakeholder</span>s — <span class="hl">PO</span>s, <span class="hl">development team</span>s, <span class="hl">project manager</span>s, <span class="hl">systems engineer</span>s. Monitors/reports <span class="hl">release plan</span> status. Resolves <span class="hl">conflict</span>s. Detects/manages <span class="hl">risk</span>s. Ensures <span class="hl">release</span>s follow technical/administrative <span class="hl">guidelines</span>. <span class="hl">Release</span> <span class="hl">documentation</span>. <span class="hl">Branching</span> models (<span class="hl">Git</span> Flow, trunk-based), <span class="hl">version</span> management. Highly <span class="hl">organized</span>, efficient, <span class="hl">fast-paced</span>. Own <span class="hl">initiative</span>. <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>). Agile &amp; Waterfall.</p>

  <div class="section">Core Competencies</div>
  <p><b>Release Management:</b> Release planning, scheduling, coordination. Monitoring &amp; reporting. Conflict resolution. Go/no-go. Documentation. Guidelines compliance. Feature flags<br>
  <b>Branching &amp; Version:</b> Git Flow, trunk-based, release/hotfix branching. Version tagging. Merge conflicts. Branch protection. Configuration management<br>
  <b>CI/CD &amp; DevOps:</b> GitHub Actions, Jenkins. Build/test/deploy pipelines. Terraform, Docker, K8s. GCP, AWS. IaC. Automated release workflows<br>
  <b>Risk &amp; Stakeholders:</b> Risk detection/management. Stakeholder communication (PO, dev, PM, systems). Status reporting. Cross-functional. Conflict resolution<br>
  <b>Software Development:</b> Python, TypeScript, C#, Java. REST APIs. SQL. Git. Playwright, Selenium. Agile (Scrum/SAFe), Waterfall. Full SDLC<br>
  <b>Tools:</b> Git, GitHub, GitHub Actions, Jenkins. Jira, Confluence. Docker, Terraform. Grafana. TestRail. MS Project</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Team Lead / Release &amp; Delivery Lead</div>
  <div class="job-sub">Mar 2022 – Present | VCS Platform — Release Management, CI/CD, 30+ Global Markets</div>
  <ul>
    <li><span class="hl">Release management</span> — interface between development plans and customer needs (30+ markets). Monitored/reported <span class="hl">release plan</span> status to <span class="hl">stakeholder</span>s (POs, dev teams, PMs). Resolved <span class="hl">conflict</span>s. <span class="hl">Go/no-go</span>. <span class="hl">Release</span>s per technical/<span class="hl">guidelines</span>.</li>
    <li><span class="hl">Branching</span> strategy &amp; <span class="hl">version</span> management — <span class="hl">Git</span> Flow, feature/release/hotfix branches. <span class="hl">Version</span> tagging. <span class="hl">Release</span> <span class="hl">documentation</span>. Branch protection. <span class="hl">Configuration management</span>.</li>
    <li><span class="hl">Risk</span> — detected/managed <span class="hl">risk</span>s in delivery. <span class="hl">Communication</span> of mitigation plans. Dependency tracking. <span class="hl">Release</span> status visibility.</li>
    <li><span class="hl">CI/CD</span> — <span class="hl">GitHub Actions</span> <span class="hl">pipeline</span>s. Automated <span class="hl">release</span> workflows. <span class="hl">Terraform</span>, <span class="hl">Docker</span>, GCP. <span class="hl">Feature flag</span> rollouts. 30% velocity. Playwright migration (3x/50%).</li>
    <li>Team — engineers + consultants. Coaching. Vendor <span class="hl">release</span> alignment. Budget. <span class="hl">Fast-paced</span>, <span class="hl">initiative</span>-driven. Exceptional Performer.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Release &amp; Automation Engineer</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Communication Platform — 300M+ Users, Release-Driven Delivery</div>
  <ul>
    <li><span class="hl">Release management</span> at scale — <span class="hl">release plan</span>ning, readiness, coordination (300M+ users). <span class="hl">Go/no-go</span>. <span class="hl">Release</span> <span class="hl">documentation</span>. Status reporting. <span class="hl">Conflict</span> resolution. <span class="hl">Risk</span> detection. <span class="hl">Branching</span>/<span class="hl">version</span> management.</li>
    <li><span class="hl">Release</span> automation — tools, workflows, <span class="hl">CI/CD</span> <span class="hl">pipeline</span>s, <span class="hl">deployment</span> automation. <span class="hl">Feature flag</span> management — controlled rollouts, <span class="hl">risk</span> analysis. Cross-functional. AWS. Agile.</li>
  </ul>

  <div class="job-title">HCLTech — LEGO &amp; IKEA, Denmark &amp; Sweden — Technical Lead</div>
  <div class="job-sub">2013 – 2021 | E-Commerce &amp; Enterprise — Release Coordination, Multi-Partner</div>
  <ul>
    <li>LEGO &amp; IKEA (2017–21): <span class="hl">Release</span> coordination across multi-partner delivery. Interface between development and project needs. <span class="hl">Release plan</span>ning, scheduling, status. <span class="hl">Conflict</span> resolution. <span class="hl">Risk</span> management. <span class="hl">Branching</span>, <span class="hl">version</span> management (<span class="hl">Git</span>). <span class="hl">Release</span> <span class="hl">documentation</span>/<span class="hl">guidelines</span>. Teams 8–10. Agile/Waterfall.</li>
    <li>Technical Lead (2013–17): <span class="hl">Software development</span>, <span class="hl">CI/CD</span>. Automation frameworks (Selenium, C#). <span class="hl">Branching</span> models, <span class="hl">version</span> management. <span class="hl">Configuration management</span>. <span class="hl">Release</span> <span class="hl">documentation</span>. <span class="hl">Initiative</span>-driven. Mentoring.</li>
  </ul>

  <div class="job-title">Banking &amp; Enterprise — Technical Lead / Consultant</div>
  <div class="job-sub">2008 – 2013 | Core Banking (Finacle CBS) — Releases, Migrations, Configuration</div>
  <ul>
    <li><span class="hl">Release</span> &amp; <span class="hl">configuration management</span> — core banking (Finacle CBS). <span class="hl">Release plan</span>ning, <span class="hl">documentation</span>, <span class="hl">guidelines</span>. Post <span class="hl">go/no-go</span> stabilization. <span class="hl">Configuration management</span>. Data migration <span class="hl">release</span>s. <span class="hl">Stakeholder</span> coordination. <span class="hl">Risk</span>. <span class="hl">Fast-paced</span>, regulated.</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">Release management</span> — IKEA (30+ markets), Truecaller (300M+), LEGO, banking. Interface dev ↔ customer. <span class="hl">Release plan</span>ning, status, <span class="hl">conflict</span> resolution, <span class="hl">risk</span>, <span class="hl">documentation</span>.</li>
    <li><span class="hl">Branching</span> &amp; <span class="hl">CI/CD</span> — <span class="hl">Git</span> Flow, trunk-based. <span class="hl">GitHub Actions</span>, <span class="hl">Jenkins</span>. Automated <span class="hl">release</span> workflows. <span class="hl">Feature flag</span>s. 3x faster, 50% CI reduction.</li>
    <li><span class="hl">Configuration management</span> — environments, branches, <span class="hl">release</span>s. <span class="hl">Version</span> tagging. <span class="hl">Release</span> <span class="hl">documentation</span>. Technical <span class="hl">guidelines</span>.</li>
    <li>Scaled IKEA platform 2K → 50K across 30+ markets. Exceptional Performer.</li>
    <li><span class="hl">Stakeholder</span> <span class="hl">communication</span> — POs, dev teams, PMs, <span class="hl">systems engineer</span>s, vendors. <span class="hl">Organized</span>, efficient, <span class="hl">initiative</span>-driven.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>ISTQB Certified Tester Foundation<br>ITIL Foundation<br>Six Sigma Green Belt<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>Certified Ethical Hacker (CEH)<br>UiPath RPA Certified</td>
    </tr>
  </table>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
