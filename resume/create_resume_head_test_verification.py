"""
Resume: Head of Test & Verification — Software & Electronics (Ambu / Medical Device)
Focus: Test strategy & vision, verification maturity, automation, regulated environment, cross-site leadership.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Head_Test_Verification_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Test & Verification leader with ", False),
        ("16+ years", True),
        (" setting ", False),
        ("vision and strategy for verification functions", True),
        (" and building setups that enable ", False),
        ("earlier verification", True),
        (", ", False),
        ("stronger automation", True),
        (", and ", False),
        ("greater release confidence", True),
        (". Proven experience driving ", False),
        ("verification maturity", True),
        (" \u2014 bringing together requirements, development, verification, and ", False),
        ("design control", True),
        (" to support a ", False),
        ("regular, predictable release cadence", True),
        (" while maintaining ", False),
        ("quality and regulatory compliance", True),
        (". Strong in ", False),
        ("test automation", True),
        (", ", False),
        ("CI/CD pipelines", True),
        (", and ", False),
        ("HW/SW integration testing", True),
        (". Experienced ", False),
        ("cross-site leader", True),
        (" who builds ", False),
        ("scalable organisations", True),
        (" with clear roles, capabilities, and effective collaboration. Background in ", False),
        ("regulated environments", True),
        (" (banking compliance, device integration) with structured approach to ", False),
        ("QMS and standards alignment", True),
        (".", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Verification Vision & Strategy", "Test Automation & Frameworks", "Release Confidence & Cadence",
        "Shift-Left Verification", "CI/CD Pipeline Integration", "HW/SW Integration Testing",
        "Cross-Site Team Leadership", "Organisational Scaling", "Regulated Environment (QMS)",
        "Risk-Based Decision Making", "Stakeholder Alignment", "Continuous Improvement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical & Verification Skills")
    skills_data = [
        ("Verification:", " Test strategy design, shift-left verification, design control, V-model, risk-based test planning"),
        ("Automation:", " Python, Pytest, Playwright, Selenium, Appium, Robot Framework, automated regression suites"),
        ("CI/CD:", " GitHub Actions, Jenkins, Azure DevOps, Docker, Kubernetes, automated test pipelines"),
        ("HW/SW:", " Hardware/software integration testing, system-level verification, device communication, sensor validation"),
        ("Quality Systems:", " QMS alignment, regulatory compliance documentation, ISTQB, Six Sigma, standards adherence"),
        ("Reporting:", " Grafana, release readiness dashboards, test coverage metrics, defect trend analysis, KPIs"),
        ("Practices:", " Agile/Scrum, Kanban, SAFe, release planning, cross-functional collaboration"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Head of Test & Verification / Team Lead (Acting)",
        "Mar 2022 \u2013 Present",
        "Platform SW & Device Integration \u2014 30+ Markets | Cross-Site Collaboration | Regulated Quality Standards")
    add_bullet(doc, "Set vision and strategy for Test & Verification across the platform \u2014 shaped the operating model that supports earlier verification, stronger automation, and predictable release cadence.",
        ["vision and strategy for Test & Verification", "operating model", "earlier verification", "predictable release cadence"])
    add_bullet(doc, "Drive shift-left verification: brought requirements, development, verification, and design control together earlier in the process \u2014 reducing late surprises and improving release confidence.",
        ["shift-left verification", "requirements, development, verification", "design control", "release confidence"])
    add_bullet(doc, "Increase test automation across software, HW/SW integration, and regression testing \u2014 scalable automated frameworks integrated into CI/CD pipelines (GitHub Actions).",
        ["test automation", "HW/SW integration", "regression testing", "CI/CD pipelines"])
    add_bullet(doc, "Lead and develop a team of 10+ test and verification professionals; define clear roles, build capabilities, and foster effective cross-site collaboration.",
        ["Lead and develop", "test and verification professionals", "cross-site collaboration"])
    add_bullet(doc, "Established Centre of Excellence for verification \u2014 designing scalable practices shared horizontally across teams while owning one product\u2019s end-to-end verification.",
        ["Centre of Excellence", "scalable practices", "end-to-end verification"])
    add_bullet(doc, "Build verification processes and documentation aligned with quality management system (QMS), regulatory standards, and compliance requirements.",
        ["verification processes", "quality management system (QMS)", "regulatory standards", "compliance"])
    add_bullet(doc, "Collaborate broadly: software, systems engineering, QA, product management, and operations \u2014 ensuring verification is aligned with product risks, architecture, and release plans.",
        ["Collaborate broadly", "product risks", "architecture", "release plans"])
    add_bullet(doc, "Drive release readiness through better planning, visibility, and risk-based decision-making \u2014 dashboards, metrics, and stakeholder reporting.",
        ["release readiness", "risk-based decision-making", "dashboards"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Verification Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Release Confidence & Automated Verification")
    add_bullet(doc, "Improved release confidence and cadence through automated verification gates, deployment validation, and release stability monitoring.",
        ["release confidence", "cadence", "automated verification gates"])
    add_bullet(doc, "Identified verification gaps and drove process improvements \u2014 earlier detection of issues, reduced late-stage surprises.",
        ["verification gaps", "earlier detection", "late-stage surprises"])
    add_bullet(doc, "Collaborated across development, QA, and operations to align verification with release plans and product risks.",
        ["development, QA, and operations", "release plans", "product risks"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 Verification Lead / SDET Lead",
        "2013 \u2013 2021",
        "Software & Electronics Platforms \u2014 Cross-Site Delivery | Organisation Building | Verification Maturity")
    add_bullet(doc, "Set verification strategy and built a mature Test & Verification function across multi-site delivery (Denmark, Sweden, offshore) \u2014 clear roles, interfaces, and ways of working.",
        ["verification strategy", "Test & Verification function", "multi-site delivery", "roles, interfaces, and ways of working"])
    add_bullet(doc, "Drove significant verification maturity step: shifted from late-stage manual testing to earlier, automated, risk-based verification across 10+ integrated systems.",
        ["verification maturity", "shifted", "earlier, automated, risk-based verification"])
    add_bullet(doc, "Built and scaled automated test frameworks (Python, Java, Selenium) integrated into CI/CD pipelines (Jenkins, EC2) \u2014 40% regression cycle reduction.",
        ["automated test frameworks", "CI/CD pipelines", "40% regression cycle reduction"])
    add_bullet(doc, "Led team of 8\u201312 skilled individual contributors across sites; shaped organisation, capabilities, and leadership capacity for growth.",
        ["8\u201312 skilled individual contributors", "shaped organisation", "leadership capacity"])
    add_bullet(doc, "Aligned verification with development, quality, and release stakeholders \u2014 translated strategic ambitions into practical, executable plans.",
        ["Aligned verification", "development, quality, and release stakeholders", "strategic ambitions into practical"])
    add_bullet(doc, "Ensured verification processes and documentation met quality system requirements and compliance standards across regulated and enterprise environments.",
        ["verification processes", "quality system requirements", "compliance standards"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 Verification Engineer / Consultant",
        "2008 \u2013 2013",
        "Core Banking & Device Integration \u2014 Regulated Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "System-level verification for Finacle Core Banking \u2014 regulated environment with rigorous compliance, documentation, and quality standards.",
        ["System-level verification", "regulated environment", "compliance"])
    add_bullet(doc, "HW/SW integration verification for biometric authentication devices (Morpho BAS) \u2014 device communication, sensor validation, security flows.",
        ["HW/SW integration verification", "biometric authentication", "device communication"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Six Sigma Green Belt", "Google Cloud ACE", "AWS Cloud Practitioner", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
