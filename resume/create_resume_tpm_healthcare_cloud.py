"""
Generate a tailored resume for Technical Project Manager – Healthcare Cloud Operations (Nordic).
Combines two JDs: Nordic Cloud Operations TPM + Customer Operations TPM (patient care solutions).
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy, os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_TPM_Healthcare_Cloud_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_TPM_Healthcare_Cloud_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading)

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.color.rgb = color
    run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(7)
    p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

from docx.enum.text import WD_COLOR_INDEX

HIGHLIGHT_TOKENS = [
    "Technical Project Manager", "Cloud Operations", "Healthcare",
    "Nordic", "stakeholder communication", "project delivery",
    "cloud services", "implementation", "patient care",
    "Azure", "GCP", "CI/CD", "Terraform", "Kubernetes",
    "planning", "coordination", "go-live", "stabilization",
    "ITIL", "Agile", "Scrum", "Kanban",
    "escalation", "risk management", "SLA",
    "cross-functional", "engineering teams", "migration",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower()
            rl = remaining.lower()
            idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)])
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                remaining = remaining[len(token):]
                matched = True
                break
            elif idx > 0:
                run = p.add_run(remaining[:idx])
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                remaining = remaining[idx:]
                matched = True
                break
        if not matched:
            run = p.add_run(remaining)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)
    add_text(p, "Technical Project Manager — Healthcare Cloud Operations", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Technical Project Manager with 16+ years in IT, spanning cloud operations, "
        "healthcare-adjacent platforms, and large-scale enterprise delivery across the Nordics. "
        "Proven track record of leading end-to-end implementation projects — from planning and "
        "coordination through go-live and stabilization — ensuring cloud services remain stable, "
        "secure, and performant. Skilled at translating complex technical realities into actionable "
        "plans and clear stakeholder communication. Experienced bridge between engineering teams "
        "and business/customer representatives, managing expectations at both strategic and "
        "operational levels. Known for bringing structure, clarity, and momentum to complex "
        "cross-functional programs."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Cloud Project Delivery (GCP, Azure)", "End-to-End Implementation Leadership", "Stakeholder & Customer Communication",
        "Nordic Multi-Market Operations", "Technical-to-Business Translation", "Risk Management & Escalation",
        "Agile / Scrum / Kanban Delivery", "Go-Live Planning & Stabilization", "CI/CD & Infrastructure as Code",
    ]
    for i, comp in enumerate(competencies):
        row = i // 3
        col = i % 3
        cell = table.rows[row].cells[col]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Team Lead / Technical Project Manager — Cloud Platform (VCS)", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    bullets_ikea = [
        "Lead cloud operations projects end-to-end for a customer-facing video platform deployed across 30+ Nordic and global markets — planning, coordination, go-live, and stabilization.",
        "Act as the primary point of contact between engineering teams, product owners, and internal stakeholders — translating complex technical needs into clear priorities and decisions.",
        "Manage cloud infrastructure on GCP (Cloud Run, BigQuery, Pub/Sub, IAM) with Terraform IaC, ensuring 99.9% uptime SLA for healthcare-grade service reliability.",
        "Drive CI/CD pipeline strategy (GitHub Actions, Cloud Build) reducing deployment lead time from days to hours while maintaining stability.",
        "Coordinate cross-functional delivery across backend (Node.js/TypeScript), data engineering (Python), and infrastructure teams — removing blockers and keeping projects on track.",
        "Manage expectations at strategic and operational levels — presenting roadmaps to leadership while handling day-to-day engineering decisions and escalations.",
        "Led platform migration and scaling from 2,000 to 50,000+ concurrent users, coordinating with external SaaS vendors, security teams, and Nordic market leads.",
        "Established project delivery practices: sprint planning, risk registers, milestone tracking, and transparent status communication to all stakeholders.",
    ]
    for b in bullets_ikea:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Release & Cloud Operations Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    bullets_tc = [
        "Managed release coordination and cloud deployment pipelines for a 300M+ user platform on AWS — ensuring stable, zero-downtime rollouts.",
        "Served as bridge between development squads and operations — coordinating release schedules, feature flags, and rollback strategies.",
        "Maintained cloud service stability through monitoring (Datadog, PagerDuty), incident coordination, and post-mortem-driven improvements.",
        "Communicated release status and technical risks to engineering leadership, enabling informed go/no-go decisions.",
    ]
    for b in bullets_tc:
        add_highlighted_bullet(doc, b)

    # --- HCLTech (IKEA & LEGO) ---
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Technical Specialist / Delivery Lead — Implementation Projects", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    bullets_hcl = [
        "Led technical implementation projects for IKEA (e-commerce, 30+ markets) and LEGO (digital platform) — managing end-to-end delivery from requirements through go-live and post-launch support.",
        "Acted as clear point of contact between customer IT teams, offshore engineering resources, and onsite stakeholders — ensuring alignment and timely delivery.",
        "Coordinated multi-partner implementation programs involving 4+ vendor teams across Denmark, Sweden, and India — maintaining milestone accountability.",
        "Translated complex technical needs (API integrations, platform migrations, test automation frameworks) into business-friendly project plans and executive updates.",
        "Drove continuous improvement in delivery processes: introduced structured sprint reviews, automated quality gates, and risk-based planning across projects.",
        "Managed stabilization and hypercare phases post-go-live — monitoring production health, coordinating defect resolution, and ensuring SLA compliance.",
    ]
    for b in bullets_hcl:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    add_text(p, "Earlier Career — Enterprise IT & Banking", bold=True, size=Pt(10))
    add_text(p, "  |  India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Technical Lead / Project Coordinator", bold=True, size=Pt(10))
    add_text(p, "    Jan 2008 – May 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    bullets_early = [
        "Coordinated implementation projects for core banking systems (Finacle CBS) — managing deliverables, timelines, and customer expectations for regulated environments.",
        "Led data migration and integration projects involving ETL pipelines, biometric systems, and third-party service providers — ensuring zero-downtime cutovers.",
        "Served as technical liaison between business analysts, development teams, and bank IT departments for multi-phase rollouts across 20+ branches.",
    ]
    for b in bullets_early:
        add_highlighted_bullet(doc, b)

    # ─── Technical Environment ──────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL ENVIRONMENT")
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Cloud: ", bold=True, size=Pt(9))
    add_text(p, "GCP (Cloud Run, BigQuery, Pub/Sub, IAM, Secret Manager), AWS (EC2, S3, CloudWatch)", size=Pt(9))
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "IaC & CI/CD: ", bold=True, size=Pt(9))
    add_text(p, "Terraform, GitHub Actions, Cloud Build, Docker, Kubernetes", size=Pt(9))
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Monitoring: ", bold=True, size=Pt(9))
    add_text(p, "Datadog, PagerDuty, GCP Operations Suite, Grafana", size=Pt(9))
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Delivery: ", bold=True, size=Pt(9))
    add_text(p, "Jira, Confluence, Azure DevOps, Agile/Scrum/Kanban, ITIL", size=Pt(9))
    p = doc.add_paragraph()
    p.space_after = Pt(2)
    add_text(p, "Integration: ", bold=True, size=Pt(9))
    add_text(p, "REST APIs, Kafka/Pub/Sub, Microsoft Graph API, Node.js, Python", size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    p = doc.add_paragraph()
    add_text(p, "M.Tech, Computer Science", bold=True, size=Pt(10))
    add_text(p, "  —  JNTU, India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph()
    add_text(p, "B.Tech, Information Technology", bold=True, size=Pt(10))
    add_text(p, "  —  JNTU, India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))
    p = doc.add_paragraph()
    add_text(p, "PG Diploma, Operations Management", bold=True, size=Pt(10))
    add_text(p, "  —  IGNOU, India", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ITIL v4 Foundation",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "ISTQB Certified Tester — Foundation Level",
        "Certified Ethical Hacker (CEH)",
        "Six Sigma Green Belt",
    ]
    # two-column layout
    table = doc.add_table(rows=3, cols=2)
    set_table_borders(table)
    for i, cert in enumerate(certs):
        row = i // 2
        col = i % 2
        cell = table.rows[row].cells[col]
        cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)  •  Swedish (Conversational – B1)  •  Danish (Basic)  •  Hindi / Urdu (Native)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Technical Project Manager &mdash; Healthcare Cloud Operations</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Technical Project Manager with 16+ years in IT, spanning cloud operations, healthcare-adjacent platforms, and large-scale enterprise delivery across the Nordics. Proven track record of leading end-to-end implementation projects &mdash; from planning and coordination through go-live and stabilization &mdash; ensuring cloud services remain stable, secure, and performant. Skilled at translating complex technical realities into actionable plans and clear stakeholder communication. Experienced bridge between engineering teams and business/customer representatives, managing expectations at both strategic and operational levels.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Cloud Project Delivery (GCP, Azure)</td><td>&bull; End-to-End Implementation Leadership</td><td>&bull; Stakeholder &amp; Customer Communication</td></tr>
<tr><td>&bull; Nordic Multi-Market Operations</td><td>&bull; Technical-to-Business Translation</td><td>&bull; Risk Management &amp; Escalation</td></tr>
<tr><td>&bull; Agile / Scrum / Kanban Delivery</td><td>&bull; Go-Live Planning &amp; Stabilization</td><td>&bull; CI/CD &amp; Infrastructure as Code</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Team Lead / Technical Project Manager &mdash; Cloud Platform (VCS) <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Lead cloud operations projects end-to-end for a customer-facing video platform deployed across 30+ Nordic and global markets &mdash; planning, coordination, go-live, and stabilization.</li>
<li>Act as the primary point of contact between engineering teams, product owners, and internal stakeholders &mdash; translating complex technical needs into clear priorities and decisions.</li>
<li>Manage cloud infrastructure on GCP (Cloud Run, BigQuery, Pub/Sub, IAM) with Terraform IaC, ensuring 99.9% uptime SLA for healthcare-grade service reliability.</li>
<li>Drive CI/CD pipeline strategy (GitHub Actions, Cloud Build) reducing deployment lead time from days to hours while maintaining stability.</li>
<li>Coordinate cross-functional delivery across backend, data engineering, and infrastructure teams &mdash; removing blockers and keeping projects on track.</li>
<li>Manage expectations at strategic and operational levels &mdash; presenting roadmaps to leadership while handling day-to-day engineering decisions and escalations.</li>
<li>Led platform migration and scaling from 2,000 to 50,000+ concurrent users, coordinating with external SaaS vendors, security teams, and Nordic market leads.</li>
<li>Established project delivery practices: sprint planning, risk registers, milestone tracking, and transparent status communication.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Release &amp; Cloud Operations Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Managed release coordination and cloud deployment pipelines for a 300M+ user platform on AWS &mdash; ensuring stable, zero-downtime rollouts.</li>
<li>Served as bridge between development squads and operations &mdash; coordinating release schedules, feature flags, and rollback strategies.</li>
<li>Maintained cloud service stability through monitoring (Datadog, PagerDuty), incident coordination, and post-mortem-driven improvements.</li>
<li>Communicated release status and technical risks to engineering leadership, enabling informed go/no-go decisions.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Technical Specialist / Delivery Lead &mdash; Implementation Projects <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Led technical implementation projects for IKEA (e-commerce, 30+ markets) and LEGO (digital platform) &mdash; managing end-to-end delivery from requirements through go-live and post-launch support.</li>
<li>Acted as clear point of contact between customer IT teams, offshore engineering resources, and onsite stakeholders &mdash; ensuring alignment and timely delivery.</li>
<li>Coordinated multi-partner implementation programs involving 4+ vendor teams across Denmark, Sweden, and India &mdash; maintaining milestone accountability.</li>
<li>Translated complex technical needs (API integrations, platform migrations, test automation frameworks) into business-friendly project plans and executive updates.</li>
<li>Drove continuous improvement in delivery processes: introduced structured sprint reviews, automated quality gates, and risk-based planning.</li>
<li>Managed stabilization and hypercare phases post-go-live &mdash; monitoring production health, coordinating defect resolution, and ensuring SLA compliance.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; Enterprise IT &amp; Banking</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;Jan 2008 &ndash; May 2013</span></p>
<ul>
<li>Coordinated implementation projects for core banking systems (Finacle CBS) &mdash; managing deliverables, timelines, and customer expectations for regulated environments.</li>
<li>Led data migration and integration projects involving ETL pipelines, biometric systems, and third-party service providers &mdash; ensuring zero-downtime cutovers.</li>
<li>Served as technical liaison between business analysts, development teams, and bank IT departments for multi-phase rollouts across 20+ branches.</li>
</ul>

<h2>TECHNICAL ENVIRONMENT</h2>
<p class="tech-line"><b>Cloud:</b> GCP (Cloud Run, BigQuery, Pub/Sub, IAM, Secret Manager), AWS (EC2, S3, CloudWatch)</p>
<p class="tech-line"><b>IaC &amp; CI/CD:</b> Terraform, GitHub Actions, Cloud Build, Docker, Kubernetes</p>
<p class="tech-line"><b>Monitoring:</b> Datadog, PagerDuty, GCP Operations Suite, Grafana</p>
<p class="tech-line"><b>Delivery:</b> Jira, Confluence, Azure DevOps, Agile/Scrum/Kanban, ITIL</p>
<p class="tech-line"><b>Integration:</b> REST APIs, Kafka/Pub/Sub, Microsoft Graph API, Node.js, Python</p>

<h2>EDUCATION</h2>
<p><b>M.Tech, Computer Science</b> &mdash; JNTU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; JNTU, India</p>
<p><b>PG Diploma, Operations Management</b> &mdash; IGNOU, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ITIL v4 Foundation</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td></tr>
<tr><td>&bull; Certified Ethical Hacker (CEH)</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent) &bull; Swedish (Conversational &ndash; B1) &bull; Danish (Basic) &bull; Hindi / Urdu (Native)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
