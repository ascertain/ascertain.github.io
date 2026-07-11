"""
Resume: Team Manager Test Engineering — Atos Medical, Hörby
Focus: Test method development, innovation projects, verification strategy, people leadership, medical device R&D.
DOCX only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Team_Manager_Test_Engineering_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    if bold_parts:
        remaining = f"\u2022 {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx + len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx + len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"\u2022 {text}")
        r.font.size = size


def role_header(doc, title, period, context=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True
    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malm\u00f6, Sweden  \u2022  +46 702624230  \u2022  mo.kashif@gmail.com  \u2022  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    parts = [
        ("Test Engineering leader with ", False),
        ("16+ years", True),
        (" combining ", False),
        ("strong technical contribution", True),
        (" with ", False),
        ("people leadership", True),
        (". Experienced in ", False),
        ("test method development", True),
        (", ", False),
        ("verification strategy", True),
        (", and guiding teams through ", False),
        ("innovation projects from early concept to design verification", True),
        (". Hands-on approach to ", False),
        ("test design", True),
        (", ", False),
        ("problem-solving", True),
        (", and ", False),
        ("data-driven evaluation", True),
        (" \u2014 staying close to the technical work while leading and developing the team. Proven ability to ", False),
        ("create clarity around priorities", True),
        (", ", False),
        ("manage stakeholders across R&D and product", True),
        (", and foster a collaborative environment where ", False),
        ("technical curiosity and continuous improvement", True),
        (" thrive. Background in ", False),
        ("regulated environments", True),
        (" (banking compliance, device integration) with structured approach to ", False),
        ("processes, standards, and quality management", True),
        (".", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Test Method Development", "Verification Strategy & Design", "People Leadership & Coaching",
        "Innovation Project Support", "Hands-On Technical Contribution", "Prioritisation & Resource Planning",
        "Stakeholder Management (R&D/Product)", "Continuous Improvement", "Regulated Environment Experience",
        "Data-Driven Decision Making", "Process & Standards (QMS)", "Team Development & Engagement",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"\u25b8 {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills & Methods")
    skills_data = [
        ("Test Engineering:", " Test method design & validation, test setup development, data analysis, statistical methods"),
        ("Verification:", " Design verification strategies, V-model, traceability, acceptance criteria, risk-based test planning"),
        ("Automation & Tools:", " Python (test scripting, data analysis), Pytest, Playwright, Selenium, Robot Framework"),
        ("Data & Analysis:", " Statistical analysis, data visualisation (Grafana, dashboards), BigQuery, measurement uncertainty"),
        ("CI/CD & Infrastructure:", " GitHub Actions, Jenkins, Docker, automated test pipelines, test environment management"),
        ("Quality Systems:", " QMS documentation, process standards, ISTQB methodology, Six Sigma (Green Belt)"),
        ("Domain:", " Device integration testing, HW/SW verification, communication protocol validation, system-level testing"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB (Ingka Digital), Malm\u00f6 \u2014 Team Lead (Acting) / Test Engineering Lead",
        "Mar 2022 \u2013 Present",
        "Omni-Channel Platform \u2014 Device Integration, SW/HW Systems, Innovation Projects | 30+ Markets")
    add_bullet(doc, "Lead a team of test engineers and consultants \u2014 support and develop team members, create clarity around priorities and deliverables, foster technical curiosity.",
        ["Lead a team of test engineers and consultants", "support and develop", "clarity around priorities", "technical curiosity"])
    add_bullet(doc, "Define, develop, and implement new test methods and test setups for innovation projects \u2014 ensuring approaches generate robust data and support sound design decisions.",
        ["new test methods and test setups", "innovation projects", "robust data", "sound design decisions"])
    add_bullet(doc, "Guide team in test strategy, method development and validation, design verification, and documentation across integrated HW/SW systems.",
        ["test strategy", "method development and validation", "design verification", "documentation"])
    add_bullet(doc, "Prioritise resources across multiple projects; support stakeholders (R&D, Product, Quality) with technical input and verification planning.",
        ["Prioritise resources", "stakeholders", "technical input", "verification planning"])
    add_bullet(doc, "Drive processes, standards, and quality management within test engineering \u2014 structured documentation, traceability, and compliance.",
        ["processes, standards, and quality management", "traceability", "compliance"])
    add_bullet(doc, "Stay close to the technical work: hands-on contribution to test design, problem-solving, data analysis, and evaluation alongside the team.",
        ["hands-on contribution", "test design", "problem-solving", "data analysis"])
    add_bullet(doc, "Foster collaborative environment with continuous improvement \u2014 blameless retrospectives, learning reviews, iterating on test practices.",
        ["collaborative environment", "continuous improvement", "iterating on test practices"])
    add_bullet(doc, "Established a Test Engineering Centre of Excellence \u2014 designing quality solutions, frameworks, and best practices shared horizontally across teams while retaining sole ownership of one product\u2019s verification.",
        ["Centre of Excellence", "designing quality solutions", "horizontally across teams", "sole ownership of one product"])
    add_bullet(doc, "Build verification strategies for complex systems combining hardware devices, software, and communication protocols.",
        ["verification strategies", "hardware devices", "communication protocols"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm \u2014 Release & Test Automation Engineer",
        "Sep 2021 \u2013 Feb 2022",
        "Communication Platform \u2014 300M+ Users | Production Verification & Release Quality")
    add_bullet(doc, "Developed and implemented test methods for release verification \u2014 ensuring robust quality gates before production deployment.",
        ["test methods for release verification", "quality gates"])
    add_bullet(doc, "Collaborated across R&D and operations to strengthen testing capabilities and align test strategies with business needs.",
        ["Collaborated across R&D", "testing capabilities", "align test strategies with business needs"])

    # --- HCLTech ---
    role_header(doc,
        "HCLTech \u2014 IKEA & LEGO Group, Denmark & Sweden \u2014 Senior Test Engineer / SDET Lead",
        "2013 \u2013 2021",
        "E-Commerce, Digital & Device Platforms \u2014 Innovation Projects, Method Development, Team Leadership")
    add_bullet(doc, "Led team of 8\u201312 test engineers and consultants across innovation and product development projects \u2014 prioritised resources, set direction, developed individuals.",
        ["Led team", "test engineers and consultants", "prioritised resources", "developed individuals"])
    add_bullet(doc, "Defined and implemented new test methods for complex integrated systems \u2014 scientifically sound, practical to implement, aligned with project needs.",
        ["new test methods", "scientifically sound", "practical to implement", "aligned with project needs"])
    add_bullet(doc, "Developed verification strategies from early concept through design verification \u2014 V-model, traceability, risk-based test planning.",
        ["verification strategies", "early concept through design verification", "risk-based test planning"])
    add_bullet(doc, "Built test infrastructure and automated test setups: Python, Java, Selenium \u2014 generating robust, repeatable data for design decisions.",
        ["test infrastructure", "automated test setups", "robust, repeatable data"])
    add_bullet(doc, "Managed stakeholders across R&D, product management, and quality \u2014 structured prioritisation, technical input, and clear communication.",
        ["stakeholders across R&D, product management, and quality", "structured prioritisation"])
    add_bullet(doc, "Drove processes and standards within test engineering; documentation aligned with quality system requirements.",
        ["processes and standards", "quality system requirements"])
    add_bullet(doc, "Acted as senior technical sparring partner in projects where thoughtful test design and hands-on execution were both essential.",
        ["senior technical sparring partner", "thoughtful test design", "hands-on execution"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise \u2014 Test Engineer / Consultant",
        "2008 \u2013 2013",
        "Core Banking & Biometric Devices \u2014 Regulated Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "Test engineering for Finacle Core Banking \u2014 regulated environment requiring rigorous verification, documentation, and compliance with quality standards.",
        ["regulated environment", "rigorous verification", "compliance with quality standards"])
    add_bullet(doc, "Morpho BAS 2FA: Device integration testing for biometric authentication hardware \u2014 test method development for HW/SW interfaces.",
        ["Device integration testing", "biometric authentication hardware", "test method development", "HW/SW interfaces"])
    add_bullet(doc, "Built automated test frameworks (Python, Java); data analysis and validation tools supporting sound, evidence-based decisions.",
        ["automated test frameworks", "data analysis", "evidence-based decisions"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = ["ISTQB Certified Tester", "Six Sigma Green Belt", "Google Cloud ACE", "AWS Cloud Practitioner", "CEH"]
    for i, cert in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  \u2022  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" \u2013 UP Technical University (Engineering Degree)  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" \u2013 IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)  |  Swedish (Basic)").font.size = Pt(9.5)

    cp2 = doc.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    r = cp2.add_run("Citizenship: ")
    r.bold = True
    r.font.size = Pt(10)
    cp2.add_run("Swedish").font.size = Pt(9.5)

    lp2 = doc.add_paragraph()
    lp2.paragraph_format.space_before = Pt(3)
    r = lp2.add_run("Location: ")
    r.bold = True
    r.font.size = Pt(10)
    lp2.add_run("Malm\u00f6, Sweden (available for on-site H\u00f6rby 4 days/week)").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")


if __name__ == "__main__":
    build_docx()
