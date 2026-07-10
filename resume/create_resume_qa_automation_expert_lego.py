"""
Resume: QA Automation Expert – LEGO Creative Play Lab (EXTERNAL)
2-page, beautified, LEGO experience highlighted, larger fonts, JD keywords bold.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_QA_Automation_Expert_LEGO_Resume"
BLUE = RGBColor(0x00, 0x51, 0xBA)
DARK = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0051BA"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, bold_parts=None, size=Pt(10)):
    """Add a bullet point. bold_parts is a list of substrings to bold."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.first_line_indent = Cm(-0.35)

    if bold_parts:
        remaining = f"• {text}"
        for bp in sorted(bold_parts, key=lambda x: remaining.find(x)):
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = size
            rb = p.add_run(remaining[idx:idx+len(bp)])
            rb.bold = True
            rb.font.size = size
            remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = size
    else:
        r = p.add_run(f"• {text}")
        r.font.size = size


def role_header(doc, title, period, context):
    p = doc.add_paragraph()
    p.space_before = Pt(7)
    p.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    # Period on same line, right-ish
    r2 = p.add_run(f"  |  {period}")
    r2.font.size = Pt(9.5)
    r2.italic = True

    if context:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r3 = p2.add_run(context)
        r3.font.size = Pt(9.5)
        r3.italic = True


# ─── DOCX ───────────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    # ── NAME ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    # ── CONTACT ──
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    r2 = p2.add_run("Malmö, Sweden  •  +46 702624230  •  mo.kashif@gmail.com  •  linkedin.com/in/md-kashif")
    r2.font.size = Pt(9.5)

    # ── PROFESSIONAL SUMMARY ──
    add_section_heading(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(3)
    sp.paragraph_format.space_after = Pt(2)
    # Build summary with bold keywords
    parts = [
        ("Quality Engineering professional with ", False),
        ("14+ years", True),
        (" driving ", False),
        ("test automation strategy", True),
        (", ", False),
        ("framework architecture", True),
        (", and ", False),
        ("CI/CD integration", True),
        (" across complex multi-platform systems — including ", False),
        ("embedded hardware-software integration", True),
        (", data pipelines, and customer-facing applications. Strong ", False),
        ("Python", True),
        (" developer with expertise in ", False),
        ("scaling test strategies", True),
        (" for increasing ", False),
        ("product complexity", True),
        (" and ", False),
        ("variant expansion", True),
        (". Direct ", False),
        ("LEGO platform experience", True),
        (" (Franchises, Campaign Factory, ", False),
        ("Sitecore", True),
        (", ", False),
        ("Pimcore", True),
        (") — deep familiarity with the product ecosystem, SKU complexity, and engineering culture.", False),
    ]
    for text, bold in parts:
        r = sp.add_run(text)
        r.font.size = Pt(10)
        if bold:
            r.bold = True

    # ── KEY COMPETENCIES ──
    add_section_heading(doc, "Key Competencies")
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    competencies = [
        "Test Automation Architecture", "Python (Core Language)", "Embedded / HW-SW Integration",
        "CI/CD Pipeline Design", "Scalable Test Frameworks", "Quality Metrics & KPIs",
        "Defect Trend Analysis & RCA", "Product Variant / SKU Testing", "Mentoring & Coaching",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"▸ {comp}")
        r.font.size = Pt(9.5)
        r.bold = True

    # ── TECHNICAL SKILLS ──
    add_section_heading(doc, "Technical Skills")
    skills_data = [
        ("Languages:", " Python (primary), TypeScript/JavaScript, SQL, Bash"),
        ("Test Frameworks:", " Pytest, Playwright, Selenium, Appium, Robot Framework"),
        ("CI/CD:", " GitHub Actions, Jenkins, Azure DevOps, Amazon EC2"),
        ("Embedded/HW:", " Hardware-software integration, system-level test harnesses, device testing"),
        ("CMS/PIM:", " Sitecore (multi-site), Pimcore (PIM/DAM)"),
        ("Cloud & Infra:", " GCP (Cloud Run, GKE, Pub/Sub), Docker, Kubernetes"),
        ("Tools:", " Jira, TestRail, Git/GitHub, SVN, Confluence"),
    ]
    for label, value in skills_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9.5)
        rv = p.add_run(value)
        rv.font.size = Pt(9.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_heading(doc, "Professional Experience")

    # --- IKEA ---
    role_header(doc,
        "IKEA IT AB, Malmö — Team Lead (Acting) / SDET",
        "Mar 2022 – Present",
        "IKEA App & Customer Connect VCS Platform — Omni-Channel, 30+ Global Markets")
    add_bullet(doc, "Drive technical QA architecture across 5+ integrated platforms (software, hardware, data pipelines) — scalable test automation strategy.",
        ["technical QA architecture", "scalable test automation strategy"])
    add_bullet(doc, "Implement automation using Python, Playwright, Selenium, Appium — 80%+ automated regression coverage across web and mobile.",
        ["Python", "Playwright", "Selenium", "Appium", "80%+ automated regression"])
    add_bullet(doc, "Manage test complexity for product variants and multi-market configurations (30+ markets); parameterised suites that scale without increasing cycle time.",
        ["product variants", "multi-market configurations", "scale"])
    add_bullet(doc, "CI/CD quality gates via GitHub Actions — security scanning, deployment verification, fast feedback loops.",
        ["CI/CD quality gates", "GitHub Actions", "fast feedback loops"])
    add_bullet(doc, "Operationalise quality metrics & KPIs (defect density, coverage, MTTD); Grafana dashboards for data-driven decisions.",
        ["quality metrics & KPIs", "data-driven decisions"])
    add_bullet(doc, "Mentor and coach engineers on automation practices; cross-functional alignment on quality standards.",
        ["Mentor and coach", "cross-functional alignment"])

    # --- Truecaller ---
    role_header(doc,
        "Truecaller, Stockholm — Release & Automation Engineer",
        "Sep 2021 – Feb 2022",
        "Communication Platform — 300M+ Users")
    add_bullet(doc, "Managed release pipelines and CI/CD infrastructure for microservices platform — reliable multi-environment deployments.",
        ["release pipelines", "CI/CD infrastructure"])
    add_bullet(doc, "Integrated automated test suites into Jenkins pipelines; quality gates with fast feedback within development workflows.",
        ["automated test suites", "Jenkins", "quality gates"])

    # --- HCLTech LEGO & IKEA ---
    role_header(doc,
        "HCLTech — LEGO & IKEA Group, Denmark & Sweden — Technical Specialist / SDET Lead",
        "2013 – 2021",
        "E-Commerce, Mobile & Enterprise — Multi-Partner Delivery")

    # ★ LEGO HIGHLIGHT BOX ★
    p_lego = doc.add_paragraph()
    p_lego.space_before = Pt(5)
    p_lego.space_after = Pt(2)
    rl = p_lego.add_run("★  LEGO DIRECT EXPERIENCE  ★")
    rl.bold = True
    rl.font.size = Pt(10.5)
    rl.font.color.rgb = BLUE

    add_bullet(doc, "LEGO Franchises (Multi-Site Sitecore): Led QA & test automation for DUPLO, Friends, Minecraft, Ninjago digital platforms and companion mobile apps. Sitecore CMS — content delivery, personalisation, multi-site architecture validation.",
        ["LEGO Franchises", "Sitecore", "DUPLO", "Friends", "Minecraft", "test automation", "multi-site architecture"])
    add_bullet(doc, "Risk management, test strategy (automation + manual), system test case design/execution/reporting across full lifecycle.",
        ["Risk management", "test strategy", "automation + manual"])
    add_bullet(doc, "LEGO Campaign Factory (Pimcore): Test design & automation for Pimcore-based digital campaign platform — product data models, DAM, asset management, e-commerce integration.",
        ["LEGO Campaign Factory", "Pimcore", "automation", "e-commerce integration"])
    add_bullet(doc, "Comprehensive test deliverables: E2E Test Strategy, Plans, Estimates, Calendar, Cases, Results — full accountability for test quality.",
        ["E2E Test Strategy", "full accountability"])
    add_bullet(doc, "LEGO Campaign Factory CI: Continuous Integration value initiative — Jenkins, Tortoise SVN, Amazon EC2 — automated build-test-deploy cycles.",
        ["Continuous Integration", "Jenkins", "Amazon EC2", "automated build-test-deploy"])
    add_bullet(doc, "Designed scalable test strategies for increasing product SKUs and franchise variant configurations across age groups and product lines.",
        ["scalable test strategies", "product SKUs", "variant configurations"])
    add_bullet(doc, "Validated hardware-software integration across LEGO connected products; system-level test harnesses for device-to-cloud data flow.",
        ["hardware-software integration", "connected products", "device-to-cloud"])

    # IKEA sub-section
    p_ikea = doc.add_paragraph()
    p_ikea.space_before = Pt(4)
    p_ikea.space_after = Pt(1)
    ri = p_ikea.add_run("IKEA & Cross-Programme:")
    ri.bold = True
    ri.font.size = Pt(10)
    ri.font.color.rgb = BLUE

    add_bullet(doc, "End-to-end test automation (Selenium, Python, Cucumber) across 10+ systems — 40% regression cycle reduction.",
        ["End-to-end test automation", "40% regression cycle reduction"])
    add_bullet(doc, "Test transformation: manual → automation-first. Managed 8–12 engineers (onshore & offshore); mentored on automation and coding standards.",
        ["Test transformation", "automation-first", "mentored"])

    # --- India Banking ---
    role_header(doc,
        "Banking & Enterprise — SDET / Consultant",
        "2008 – 2013",
        "Finacle CBS, Core Banking — Regulated Environments (HCL, Marlabs, TekMindz, India)")
    add_bullet(doc, "Finacle Core Banking QA — accounts, transactions, loans in regulated environments. Morpho BAS 2FA biometric integration testing.",
        ["Finacle Core Banking", "biometric integration"])
    add_bullet(doc, "Built automated test frameworks (Selenium, Python, Java); data validation tools and test scripting.",
        ["automated test frameworks", "Python"])

    # ── CERTIFICATIONS & EDUCATION ──
    add_section_heading(doc, "Certifications & Education")
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(3)
    certs = [("ISTQB", " Certified Tester"), ("Google Cloud ACE", ""), ("AWS Cloud Practitioner", ""),
             ("CEH", ""), ("Six Sigma Green Belt", "")]
    for i, (cert, suffix) in enumerate(certs):
        r = cp.add_run(cert)
        r.bold = True
        r.font.size = Pt(9.5)
        if suffix:
            r2 = cp.add_run(suffix)
            r2.font.size = Pt(9.5)
        if i < len(certs) - 1:
            cp.add_run("  •  ").font.size = Pt(9.5)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(3)
    r = ep.add_run("B.Tech Information Technology")
    r.bold = True
    r.font.size = Pt(10)
    ep.add_run(" – UP Technical University  |  ").font.size = Pt(9.5)
    r2 = ep.add_run("PGDOM")
    r2.bold = True
    r2.font.size = Pt(10)
    ep.add_run(" – IGNOU").font.size = Pt(9.5)

    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(3)
    r = lp.add_run("Languages: ")
    r.bold = True
    r.font.size = Pt(10)
    lp.add_run("English (Fluent)  |  Hindi/Urdu (Native)").font.size = Pt(9.5)

    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out


# ─── DOC (HTML) ─────────────────────────────────────────────────────────────
def build_doc():
    css = """
    body{font-family:Calibri,sans-serif;font-size:10pt;margin:1cm 1.2cm;color:#1a1a1a;line-height:1.35}
    h1{text-align:center;color:#0051BA;font-size:18pt;margin-bottom:2px;letter-spacing:1px}
    .contact{text-align:center;font-size:9.5pt;margin-bottom:10px;color:#444}
    h2{color:#0051BA;font-size:11pt;border-bottom:2px solid #0051BA;padding-bottom:2px;margin-top:12px;margin-bottom:4px;text-transform:uppercase;letter-spacing:0.5px}
    .role-line{margin-top:8px;margin-bottom:0}
    .role-title{font-weight:bold;color:#0051BA;font-size:10.5pt}
    .role-period{font-style:italic;font-size:9.5pt;color:#555}
    .context{font-style:italic;font-size:9.5pt;margin-top:0;margin-bottom:3px;color:#555}
    .lego-header{font-weight:bold;color:#0051BA;font-size:10.5pt;margin-top:6px;margin-bottom:2px;
                  background:#E8F0FE;padding:3px 8px;border-left:4px solid #0051BA;display:inline-block}
    .sub{font-weight:bold;color:#0051BA;font-size:10pt;margin-top:5px;margin-bottom:2px}
    ul{margin:2px 0 4px 18px;padding:0}
    li{font-size:10pt;margin-bottom:2px;line-height:1.3}
    .summary{font-size:10pt;margin-top:4px;line-height:1.4}
    .comp-table{width:100%;font-size:9.5pt;margin-top:4px;border-collapse:collapse}
    .comp-table td{padding:2px 8px;font-weight:bold}
    .skills-row{font-size:9.5pt;margin:1px 0 1px 8px}
    .skills-row b{color:#0051BA}
    .certs{font-size:9.5pt;margin-top:3px}
    .edu{font-size:9.5pt;margin-top:3px}
    """

    body = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;•&nbsp; +46 702624230 &nbsp;•&nbsp; mo.kashif@gmail.com &nbsp;•&nbsp; linkedin.com/in/md-kashif</p>

<h2>Professional Summary</h2>
<p class="summary">Quality Engineering professional with <b>14+ years</b> driving <b>test automation strategy</b>,
<b>framework architecture</b>, and <b>CI/CD integration</b> across complex multi-platform systems — including
<b>embedded hardware-software integration</b>, data pipelines, and customer-facing applications.
Strong <b>Python</b> developer with expertise in <b>scaling test strategies</b> for increasing
<b>product complexity</b> and <b>variant expansion</b>. Direct <b>LEGO platform experience</b>
(Franchises, Campaign Factory, <b>Sitecore</b>, <b>Pimcore</b>) — deep familiarity with the product ecosystem,
SKU complexity, and engineering culture.</p>

<h2>Key Competencies</h2>
<table class="comp-table">
<tr><td>▸ Test Automation Architecture</td><td>▸ Python (Core Language)</td><td>▸ Embedded / HW-SW Integration</td></tr>
<tr><td>▸ CI/CD Pipeline Design</td><td>▸ Scalable Test Frameworks</td><td>▸ Quality Metrics &amp; KPIs</td></tr>
<tr><td>▸ Defect Trend Analysis &amp; RCA</td><td>▸ Product Variant / SKU Testing</td><td>▸ Mentoring &amp; Coaching</td></tr>
</table>

<h2>Technical Skills</h2>
<p class="skills-row"><b>Languages:</b> Python (primary), TypeScript/JavaScript, SQL, Bash</p>
<p class="skills-row"><b>Test Frameworks:</b> Pytest, <b>Playwright</b>, <b>Selenium</b>, <b>Appium</b>, Robot Framework</p>
<p class="skills-row"><b>CI/CD:</b> GitHub Actions, <b>Jenkins</b>, Azure DevOps, <b>Amazon EC2</b></p>
<p class="skills-row"><b>Embedded/HW:</b> Hardware-software integration, system-level test harnesses, device testing</p>
<p class="skills-row"><b>CMS/PIM:</b> <b>Sitecore</b> (multi-site), <b>Pimcore</b> (PIM/DAM)</p>
<p class="skills-row"><b>Cloud &amp; Infra:</b> GCP (Cloud Run, GKE, Pub/Sub), Docker, Kubernetes</p>
<p class="skills-row"><b>Tools:</b> Jira, TestRail, Git/GitHub, SVN, Confluence</p>

<h2>Professional Experience</h2>

<p class="role-line"><span class="role-title">IKEA IT AB, Malmö — Team Lead (Acting) / SDET</span> <span class="role-period">&nbsp;|&nbsp; Mar 2022 – Present</span></p>
<p class="context">IKEA App &amp; Customer Connect VCS Platform — Omni-Channel, 30+ Global Markets</p>
<ul>
<li>Drive <b>technical QA architecture</b> across 5+ integrated platforms (software, hardware, data pipelines) — <b>scalable test automation strategy</b>.</li>
<li>Implement automation using <b>Python</b>, <b>Playwright</b>, <b>Selenium</b>, <b>Appium</b> — <b>80%+ automated regression</b> across web and mobile.</li>
<li>Manage test complexity for <b>product variants</b> and multi-market configurations (30+ markets); <b>parameterised suites that scale</b>.</li>
<li><b>CI/CD quality gates</b> via <b>GitHub Actions</b> — security scanning, deployment verification, <b>fast feedback loops</b>.</li>
<li>Operationalise <b>quality metrics &amp; KPIs</b> (defect density, coverage, MTTD); <b>Grafana dashboards</b> for data-driven decisions.</li>
<li><b>Mentor and coach</b> engineers on automation practices; <b>cross-functional alignment</b> on quality standards.</li>
</ul>

<p class="role-line"><span class="role-title">Truecaller, Stockholm — Release &amp; Automation Engineer</span> <span class="role-period">&nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<p class="context">Communication Platform — 300M+ Users</p>
<ul>
<li>Managed <b>release pipelines</b> and <b>CI/CD infrastructure</b> for microservices platform — reliable multi-environment deployments.</li>
<li>Integrated <b>automated test suites</b> into <b>Jenkins</b> pipelines; <b>quality gates</b> with fast feedback.</li>
</ul>

<p class="role-line"><span class="role-title">HCLTech — LEGO &amp; IKEA Group, Denmark &amp; Sweden — Technical Specialist / SDET Lead</span> <span class="role-period">&nbsp;|&nbsp; 2013 – 2021</span></p>
<p class="context">E-Commerce, Mobile &amp; Enterprise — Multi-Partner Delivery</p>

<p class="lego-header">★&nbsp; LEGO DIRECT EXPERIENCE &nbsp;★</p>
<ul>
<li><b>LEGO Franchises (Multi-Site Sitecore):</b> Led QA &amp; <b>test automation</b> for <b>DUPLO</b>, <b>Friends</b>, <b>Minecraft</b>, Ninjago digital platforms and companion mobile apps. <b>Sitecore CMS</b> — content delivery, personalisation, <b>multi-site architecture</b> validation.</li>
<li><b>Risk management</b>, <b>test strategy</b> (automation + manual), system test case design/execution/reporting across full lifecycle.</li>
<li><b>LEGO Campaign Factory (Pimcore):</b> Test design &amp; <b>automation</b> for <b>Pimcore</b>-based digital campaign platform — product data models, DAM, <b>asset management</b>, <b>e-commerce integration</b>.</li>
<li>Comprehensive <b>test deliverables</b>: <b>E2E Test Strategy</b>, Plans, Estimates, Calendar, Cases, Results — full accountability for test quality.</li>
<li><b>LEGO Campaign Factory CI:</b> <b>Continuous Integration</b> value initiative — <b>Jenkins</b>, Tortoise SVN, <b>Amazon EC2</b> — <b>automated build-test-deploy</b> cycles.</li>
<li>Designed <b>scalable test strategies</b> for increasing <b>product SKUs</b> and franchise <b>variant configurations</b> across age groups and product lines.</li>
<li>Validated <b>hardware-software integration</b> across LEGO <b>connected products</b>; <b>system-level test harnesses</b> for <b>device-to-cloud</b> data flow.</li>
</ul>

<p class="sub">IKEA &amp; Cross-Programme:</p>
<ul>
<li><b>End-to-end test automation</b> (Selenium, Python, Cucumber) across 10+ systems — <b>40% regression cycle reduction</b>.</li>
<li><b>Test transformation</b>: manual → <b>automation-first</b>. Managed 8–12 engineers (onshore &amp; offshore); <b>mentored</b> on automation and coding standards.</li>
</ul>

<p class="role-line"><span class="role-title">Banking &amp; Enterprise — SDET / Consultant</span> <span class="role-period">&nbsp;|&nbsp; 2008 – 2013</span></p>
<p class="context">Finacle CBS, Core Banking — Regulated Environments (HCL, Marlabs, TekMindz, India)</p>
<ul>
<li><b>Finacle Core Banking</b> QA — accounts, transactions, loans in regulated environments. <b>Morpho BAS 2FA</b> biometric integration testing.</li>
<li>Built <b>automated test frameworks</b> (Selenium, <b>Python</b>, Java); data validation tools and test scripting.</li>
</ul>

<h2>Certifications &amp; Education</h2>
<p class="certs"><b>ISTQB</b> Certified Tester &nbsp;•&nbsp; <b>Google Cloud ACE</b> &nbsp;•&nbsp; <b>AWS Cloud Practitioner</b> &nbsp;•&nbsp; <b>CEH</b> &nbsp;•&nbsp; <b>Six Sigma Green Belt</b></p>
<p class="edu"><b>B.Tech Information Technology</b> – UP Technical University &nbsp;|&nbsp; <b>PGDOM</b> – IGNOU</p>
<p class="edu"><b>Languages:</b> English (Fluent) &nbsp;|&nbsp; Hindi/Urdu (Native)</p>

</body></html>"""

    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(body, encoding="utf-8")
    print(f"DOC saved: {out}")


if __name__ == "__main__":
    build_docx()
    build_doc()
