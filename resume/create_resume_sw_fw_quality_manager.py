"""
Generate a tailored resume for SW/FW Quality Manager (LEGO-style role).
Focus: SW/FW verification & validation, quality strategy, embedded systems,
test management, defect triaging, go/no-go decisions, ISTQB, cross-functional,
stakeholder management, cybersecurity awareness, project management.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_SW_FW_Quality_Manager_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_SW_FW_Quality_Manager_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "quality strategy", "quality assurance", "quality requirements",
    "verification and validation", "V&V",
    "firmware", "FW", "embedded", "SW",
    "go/no-go", "release",
    "defect triaging", "defect management",
    "risk", "mitigation",
    "stakeholder", "cross-functional",
    "ISTQB", "test management",
    "Jira", "JAMA", "Confluence",
    "cybersecurity", "security",
    "project management",
    "end-to-end", "supply chain",
    "capability building", "resource allocation",
    "objectives", "targets", "KPI",
    "CI/CD", "automation", "Python", "pytest",
    "structured", "systematic", "self-driven",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "SW / FW Quality Assurance Manager", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "SW/FW Quality Assurance Manager with 15+ years of proven experience in Software Test Management, "
        "Software Engineering, and embedded SW/FW quality assurance within cross-functional, global environments. "
        "Builds and drives quality strategies for SW/FW verification and validation — securing line-of-sight to "
        "product and customer requirements, setting quality objectives and targets, and ensuring full ecosystem "
        "qualification across all elements. Experienced in defect triaging, Quality go/no-go recommendations, and "
        "connecting project leadership on trade-off decisions, priorities, capability building, and resource allocation. "
        "ISTQB certified with strong end-to-end supply chain understanding from a quality angle. Engages effectively "
        "with both deep tech specialists and business leaders — translating technical topics into quality approaches "
        "and business decisions. Knowledge in cybersecurity areas. Structured, systematic, and self-driven with "
        "excellent stakeholder management and communication skills."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Quality Strategy (SW/FW V&V)", "Defect Triaging & Go/No-Go", "Embedded SW/FW Quality",
        "Test Management (ISTQB)", "Risk Mitigation & Escalation", "Stakeholder Management",
        "Cross-Functional Leadership", "Cybersecurity Awareness", "Capability Building & Coaching",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "SW/FW Quality Manager / Test Lead — IoT & Connected Products", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Build and drive the quality strategy for SW/FW verification and validation of connected IoT products — securing line-of-sight to product and customer requirements across the full development lifecycle.",
        "Set objectives and targets for quality requirements and performance — defining measurable KPIs for release quality, defect density, test coverage, and ecosystem stability.",
        "Ensure the full SW/FW ecosystem is qualified across all elements — coordinating verification of firmware, cloud services, mobile apps, and hardware integrations as an integrated system.",
        "Drive and secure objective goals for release of firmware and applications — managing release criteria, quality gates, and ensuring readiness across all platform components.",
        "Perform defect triaging across severity levels — prioritizing, escalating, and tracking resolution of critical defects across firmware, software, and integration layers.",
        "Recommend Quality go/no-go decisions based on comprehensive risk assessment, defect analysis, and test evidence — advising project and sponsor leadership on release readiness.",
        "Connect and advise project leadership teams on trade-off decisions, priorities, capability building, and resource allocation related to quality assurance in development projects.",
        "Report, escalate, and mitigate risks and critical project matters related to quality and technical feasibility — maintaining transparent visibility for steering committees.",
        "Engage and build Quality Assurance capabilities with colleagues across Product Development, Engineering, and Quality teams — coaching engineers, driving best practices, and fostering a quality-first culture.",
        "Drive end-to-end supply chain quality understanding — ensuring quality is embedded from requirement through design, development, manufacturing, and field operation.",
        "Apply cybersecurity quality practices — ensuring firmware and application security is validated as part of the overall quality strategy.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Quality & Release Manager", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Drove quality strategy and verification for a globally released mobile application (300M+ users) — ensuring release quality across multiple platforms and firmware dependencies.",
        "Recommended Quality go/no-go decisions for production releases — assessing defect severity, test coverage, and risk exposure to advise leadership on release readiness.",
        "Performed defect triaging and priority management — coordinating resolution across engineering teams and tracking critical issues to closure.",
        "Reported and escalated risks related to quality and technical feasibility — providing transparent status to project leadership and driving mitigation actions.",
        "Built QA capability within the engineering organization — coaching developers on testing practices and embedding quality into the development workflow.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "SW/FW Quality Lead / Test Manager — IoT & Consumer Products", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Built and executed quality strategies for SW/FW verification and validation across IKEA IoT/smart-home products and LEGO connected digital products — ensuring product and customer requirements were met end-to-end.",
        "Set quality objectives, targets, and performance criteria for firmware releases, app updates, and cloud service deployments — tracking against KPIs and driving continuous improvement.",
        "Ensured full SW/FW ecosystem qualification — coordinating verification across embedded firmware, mobile applications, cloud backend, and hardware integration layers as a complete system.",
        "Drove objective release goals for firmware and apps — establishing release criteria, quality gates, and go/no-go decision frameworks used by project leadership.",
        "Performed defect triaging across the full product stack — prioritizing by business impact, coordinating cross-team resolution, and tracking systemic patterns for root-cause action.",
        "Recommended Quality go/no-go decisions to project sponsors and leadership — presenting risk assessments, test evidence, and quality metrics to support informed release decisions.",
        "Connected and supervised project leadership teams on trade-off decisions, priorities, and resource allocation — balancing quality, schedule, and scope across interdisciplinary development activities.",
        "Reported, escalated, and mitigated risks related to quality and technical feasibility — maintaining risk registers, driving mitigation plans, and ensuring critical matters reached appropriate leadership.",
        "Engaged and built QA capabilities across Product Development, Engineering, and Quality teams — mentoring engineers (8–12 team members), establishing testing standards, and fostering quality ownership.",
        "Managed interdisciplinary development activities — coordinating quality across hardware, embedded firmware, mobile, cloud, and security teams in a global, cross-functional environment.",
        "Applied cybersecurity quality practices — ensuring security testing was integrated into the verification strategy for connected products and firmware updates.",
        "Drove end-to-end supply chain quality understanding — from component-level firmware through manufacturing validation to field operation and post-release monitoring.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — SW Quality & Test Management", bold=True, size=Pt(10))
    add_text(p, "  |  India  |  2008 – 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Managed SW quality strategy and verification for enterprise financial systems — setting quality objectives, performing defect triaging, and recommending go/no-go decisions for production releases.",
        "Coordinated cross-functional quality activities across distributed teams — managing priorities, resource allocation, and capability building within test engineering organizations.",
        "Reported and escalated quality risks to project leadership — providing structured risk assessments and driving mitigation actions in regulated environments.",
        "Built test automation and tooling (Python, shell scripting) — establishing continuous integration practices and driving systematic quality improvement.",
        "Engaged with both technical specialists and business stakeholders — translating quality findings into business-relevant decisions and trade-off recommendations.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS & TOOLS")
    for label, value in [
        ("Quality & V&V: ", "Quality strategy, V&V planning, release criteria, go/no-go frameworks, KPIs, ecosystem qualification"),
        ("Test Management: ", "Jira, Confluence, Azure DevOps, Zephyr, TestRail, HP ALM — test planning, tracking, reporting"),
        ("Automation: ", "Python (pytest, requests, custom frameworks), Selenium, Playwright, CI/CD integration (GitHub Actions, Jenkins)"),
        ("Embedded/FW: ", "Firmware verification, hardware-software integration testing, IoT device validation, connectivity protocols (BLE, NFC, MQTT)"),
        ("Cybersecurity: ", "Security testing practices, vulnerability assessment awareness, CEH certified, secure development lifecycle"),
        ("Project Mgmt: ", "Agile (Scrum/Kanban), SAFe, cross-functional coordination, risk management, resource planning, trade-off facilitation"),
        ("Platforms: ", "GCP, Azure, AWS, Linux, Docker, Kubernetes, distributed systems"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("Post Graduate Diploma in Operation and Management", "IGNOU, India"),
        ("B.Tech, Information Technology", "UP Technical University, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
        "ITIL v4 Foundation",
        "Six Sigma Green Belt",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent — written and verbal)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">SW / FW Quality Assurance Manager</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>SW/FW Quality Assurance Manager with 15+ years of proven experience in Software Test Management, Software Engineering, and embedded SW/FW quality assurance within cross-functional, global environments. Builds and drives quality strategies for SW/FW verification and validation &mdash; securing line-of-sight to product and customer requirements, setting quality objectives and targets, and ensuring full ecosystem qualification across all elements. Experienced in defect triaging, Quality go/no-go recommendations, and connecting project leadership on trade-off decisions, priorities, capability building, and resource allocation. ISTQB certified with strong end-to-end supply chain understanding from a quality angle. Engages effectively with both deep tech specialists and business leaders &mdash; translating technical topics into quality approaches and business decisions. Knowledge in cybersecurity areas. Structured, systematic, and self-driven with excellent stakeholder management and communication skills.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Quality Strategy (SW/FW V&amp;V)</td><td>&bull; Defect Triaging &amp; Go/No-Go</td><td>&bull; Embedded SW/FW Quality</td></tr>
<tr><td>&bull; Test Management (ISTQB)</td><td>&bull; Risk Mitigation &amp; Escalation</td><td>&bull; Stakeholder Management</td></tr>
<tr><td>&bull; Cross-Functional Leadership</td><td>&bull; Cybersecurity Awareness</td><td>&bull; Capability Building &amp; Coaching</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">SW/FW Quality Manager / Test Lead &mdash; IoT &amp; Connected Products <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Build and drive the quality strategy for SW/FW verification and validation of connected IoT products &mdash; securing line-of-sight to product and customer requirements across the full development lifecycle.</li>
<li>Set objectives and targets for quality requirements and performance &mdash; defining measurable KPIs for release quality, defect density, test coverage, and ecosystem stability.</li>
<li>Ensure the full SW/FW ecosystem is qualified across all elements &mdash; coordinating verification of firmware, cloud services, mobile apps, and hardware integrations as an integrated system.</li>
<li>Drive and secure objective goals for release of firmware and applications &mdash; managing release criteria, quality gates, and ensuring readiness across all platform components.</li>
<li>Perform defect triaging across severity levels &mdash; prioritizing, escalating, and tracking resolution of critical defects across firmware, software, and integration layers.</li>
<li>Recommend Quality go/no-go decisions based on comprehensive risk assessment, defect analysis, and test evidence &mdash; advising project and sponsor leadership on release readiness.</li>
<li>Connect and advise project leadership teams on trade-off decisions, priorities, capability building, and resource allocation related to quality assurance in development projects.</li>
<li>Report, escalate, and mitigate risks and critical project matters related to quality and technical feasibility &mdash; maintaining transparent visibility for steering committees.</li>
<li>Engage and build Quality Assurance capabilities with colleagues across Product Development, Engineering, and Quality teams &mdash; coaching engineers, driving best practices, and fostering a quality-first culture.</li>
<li>Drive end-to-end supply chain quality understanding &mdash; ensuring quality is embedded from requirement through design, development, manufacturing, and field operation.</li>
<li>Apply cybersecurity quality practices &mdash; ensuring firmware and application security is validated as part of the overall quality strategy.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">Quality &amp; Release Manager <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Drove quality strategy and verification for a globally released mobile application (300M+ users) &mdash; ensuring release quality across multiple platforms and firmware dependencies.</li>
<li>Recommended Quality go/no-go decisions for production releases &mdash; assessing defect severity, test coverage, and risk exposure to advise leadership on release readiness.</li>
<li>Performed defect triaging and priority management &mdash; coordinating resolution across engineering teams and tracking critical issues to closure.</li>
<li>Reported and escalated risks related to quality and technical feasibility &mdash; providing transparent status to project leadership and driving mitigation actions.</li>
<li>Built QA capability within the engineering organization &mdash; coaching developers on testing practices and embedding quality into the development workflow.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">SW/FW Quality Lead / Test Manager &mdash; IoT &amp; Consumer Products <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Built and executed quality strategies for SW/FW verification and validation across IKEA IoT/smart-home products and LEGO connected digital products &mdash; ensuring product and customer requirements were met end-to-end.</li>
<li>Set quality objectives, targets, and performance criteria for firmware releases, app updates, and cloud service deployments &mdash; tracking against KPIs and driving continuous improvement.</li>
<li>Ensured full SW/FW ecosystem qualification &mdash; coordinating verification across embedded firmware, mobile applications, cloud backend, and hardware integration layers as a complete system.</li>
<li>Drove objective release goals for firmware and apps &mdash; establishing release criteria, quality gates, and go/no-go decision frameworks used by project leadership.</li>
<li>Performed defect triaging across the full product stack &mdash; prioritizing by business impact, coordinating cross-team resolution, and tracking systemic patterns for root-cause action.</li>
<li>Recommended Quality go/no-go decisions to project sponsors and leadership &mdash; presenting risk assessments, test evidence, and quality metrics to support informed release decisions.</li>
<li>Connected and supervised project leadership teams on trade-off decisions, priorities, and resource allocation &mdash; balancing quality, schedule, and scope across interdisciplinary development activities.</li>
<li>Reported, escalated, and mitigated risks related to quality and technical feasibility &mdash; maintaining risk registers, driving mitigation plans, and ensuring critical matters reached appropriate leadership.</li>
<li>Engaged and built QA capabilities across Product Development, Engineering, and Quality teams &mdash; mentoring engineers (8&ndash;12 team members), establishing testing standards, and fostering quality ownership.</li>
<li>Managed interdisciplinary development activities &mdash; coordinating quality across hardware, embedded firmware, mobile, cloud, and security teams in a global, cross-functional environment.</li>
<li>Applied cybersecurity quality practices &mdash; ensuring security testing was integrated into the verification strategy for connected products and firmware updates.</li>
<li>Drove end-to-end supply chain quality understanding &mdash; from component-level firmware through manufacturing validation to field operation and post-release monitoring.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; SW Quality &amp; Test Management</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;2008 &ndash; 2013</span></p>
<ul>
<li>Managed SW quality strategy and verification for enterprise financial systems &mdash; setting quality objectives, performing defect triaging, and recommending go/no-go decisions for production releases.</li>
<li>Coordinated cross-functional quality activities across distributed teams &mdash; managing priorities, resource allocation, and capability building within test engineering organizations.</li>
<li>Reported and escalated quality risks to project leadership &mdash; providing structured risk assessments and driving mitigation actions in regulated environments.</li>
<li>Built test automation and tooling (Python, shell scripting) &mdash; establishing continuous integration practices and driving systematic quality improvement.</li>
<li>Engaged with both technical specialists and business stakeholders &mdash; translating quality findings into business-relevant decisions and trade-off recommendations.</li>
</ul>

<h2>TECHNICAL SKILLS &amp; TOOLS</h2>
<p class="tech-line"><b>Quality &amp; V&amp;V:</b> Quality strategy, V&amp;V planning, release criteria, go/no-go frameworks, KPIs, ecosystem qualification</p>
<p class="tech-line"><b>Test Management:</b> Jira, Confluence, Azure DevOps, Zephyr, TestRail, HP ALM &mdash; test planning, tracking, reporting</p>
<p class="tech-line"><b>Automation:</b> Python (pytest, requests, custom frameworks), Selenium, Playwright, CI/CD integration (GitHub Actions, Jenkins)</p>
<p class="tech-line"><b>Embedded/FW:</b> Firmware verification, hardware-software integration testing, IoT device validation, connectivity protocols (BLE, NFC, MQTT)</p>
<p class="tech-line"><b>Cybersecurity:</b> Security testing practices, vulnerability assessment awareness, CEH certified, secure development lifecycle</p>
<p class="tech-line"><b>Project Mgmt:</b> Agile (Scrum/Kanban), SAFe, cross-functional coordination, risk management, resource planning, trade-off facilitation</p>
<p class="tech-line"><b>Platforms:</b> GCP, Azure, AWS, Linux, Docker, Kubernetes, distributed systems</p>

<h2>EDUCATION</h2>
<p><b>Post Graduate Diploma in Operation and Management</b> &mdash; IGNOU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; UP Technical University, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
<tr><td>&bull; ITIL v4 Foundation</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent &mdash; written and verbal)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
