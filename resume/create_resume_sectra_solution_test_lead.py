from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(r"C:/Users/MOKAS10/vcs/csrs-vcs-core")
DOCX_PATH = BASE_DIR / "Mohammad_Kashif_Sectra_Solution_Test_Lead_Resume.docx"
DOC_PATH = BASE_DIR / "Mohammad_Kashif_Sectra_Solution_Test_Lead_Resume.doc"

ACCENT = "1F4E79"
TEXT_DARK = RGBColor(31, 41, 55)
TEXT_MUTED = RGBColor(75, 85, 99)
SECTION_COLOR = RGBColor(31, 78, 121)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color=ACCENT, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_table_borders(table, color="B8CCE4", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def add_text(paragraph, text, *, bold=False, size=10, color=TEXT_DARK, highlight=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = highlight
    return run


def add_section_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, title.upper(), bold=True, size=10.5, color=SECTION_COLOR)
    set_paragraph_bottom_border(paragraph)


HIGHLIGHT_TOKENS = [
    "15+ years",
    "Solution Test Lead",
    "test strategy",
    "system-level",
    "integration testing",
    "end-to-end",
    "E2E",
    "integration point",
    "test plan",
    "test level",
    "test frequency",
    "regression",
    "acceptance testing",
    "continuous integration",
    "automation",
    "test automation",
    "Playwright",
    "CI/CD",
    "security",
    "Threat Modeling",
    "Cyber Jedi",
    "CEH",
    "robustness",
    "safety",
    "non-functional",
    "risk",
    "risk-based",
    "tool chain",
    "test infrastructure",
    "cross-functional",
    "architects",
    "developers",
    "systems engineer",
    "project manager",
    "requirements analysis",
    "validation",
    "ecosystem",
    "complex system",
    "multiple products",
    "Agile",
    "Scrum",
    "DevOps",
    "GitHub Actions",
    "Terraform",
    "mentor",
    "coaching",
    "quality",
    "culture",
]


def add_highlighted_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for token in HIGHLIGHT_TOKENS:
        if token.lower() in text.lower():
            idx = text.lower().index(token.lower())
            before = text[:idx]
            matched = text[idx : idx + len(token)]
            after = text[idx + len(token) :]
            add_text(paragraph, before, size=10)
            add_text(paragraph, matched, bold=True, size=10, highlight=WD_COLOR_INDEX.YELLOW)
            add_text(paragraph, after, size=10)
            return
    add_text(paragraph, text, size=10)


def build_docx():
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    # ── Header ──
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    add_text(header, "MOHAMMAD KASHIF", bold=True, size=18, color=SECTION_COLOR)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(0)
    role.paragraph_format.space_after = Pt(1)
    add_text(
        role,
        "Solution Test Lead | System-Level Test Strategy & Automation | 15+ Years | Security-Focused",
        bold=True,
        size=10.5,
    )

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(2)
    add_text(contact, "Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com", size=9.5, color=TEXT_MUTED)
    add_text(
        contact,
        "  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io",
        size=9.5,
        color=TEXT_MUTED,
    )
    set_paragraph_bottom_border(contact, color="0058A3", size="10")

    # ── Profile ──
    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(2)
    profile.paragraph_format.line_spacing = 1.0
    add_text(
        profile,
        "Solution Test Lead with 15+ years of experience taking overall responsibility for how "
        "complex systems of interconnected products are tested. Proven expertise in designing "
        "system-level test strategies — defining test levels, integration points, test flows, "
        "and test frequency across the full lifecycle. Deep security background: Certified "
        "Ethical Hacker (CEH), Cyber Jedi with Threat Modeling experience, and hands-on work "
        "ensuring solutions meet high standards of security and robustness. Track record of "
        "leading test automation at scale — Playwright E2E, CI/CD pipelines, continuous "
        "integration, regression testing, and test infrastructure improvements. Close "
        "collaborator with developers, architects, systems engineers, and project managers. "
        "Driven by building a culture where testing is a natural part of the entire development "
        "process.",
        size=10,
    )

    # ── How I Match the Role ──
    add_section_heading(document, "How I Match the Role")
    match_lines = [
        (
            "Test Strategy for the System: ",
            "Designed how ecosystems of products should be tested — from individual components "
            "to complete solutions. Defined test levels, test flows, integration points, and "
            "test frequency throughout the lifecycle. Ensured strategy covers functional, "
            "non-functional, and security-related requirements.",
        ),
        (
            "Planning & Coordination: ",
            "Developed test plans covering integration testing, end-to-end testing, and "
            "validation in realistic environments. Prioritized testing efforts based on risk, "
            "complexity, and dependencies between products.",
        ),
        (
            "Automation & Efficiency: ",
            "Led increased automation at system level — Playwright E2E frameworks, CI/CD "
            "pipelines (GitHub Actions), continuous integration, regression testing, and quick "
            "feedback loops. Drove improvements in test infrastructure and tool chains.",
        ),
        (
            "Execution & Analysis: ",
            "Conducted and followed up tests in close collaboration with development teams. "
            "Actively built a culture where tests are a natural part of the entire development "
            "process — from early verification through acceptance testing.",
        ),
        (
            "Security & Robustness: ",
            "Certified Ethical Hacker (CEH). Cyber Jedi — implemented Threat Modeling and "
            "security validation within SDLC. Ensured solutions meet high standards of "
            "security and robustness, aligning with safety-related requirements.",
        ),
    ]
    for label, value in match_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_text(paragraph, label, bold=True, size=10)
        add_text(paragraph, value, size=10)

    # ── Professional Experience ──
    add_section_heading(document, "Professional Experience")

    # -- IKEA VCS --
    cp1 = document.add_paragraph()
    cp1.paragraph_format.space_before = Pt(3)
    cp1.paragraph_format.space_after = Pt(0)
    add_text(cp1, "IKEA IT AB, Malmö — Solution Test Lead / Senior Software Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp1 = document.add_paragraph()
    tp1.paragraph_format.space_after = Pt(0)
    add_text(tp1, "Mar 2022 – Present  |  VCS — Video Collaboration Solution (Multi-Product Ecosystem)", bold=True, size=10)
    ikea_bullets = [
        "Owned the overall test strategy for VCS — a complex system of interconnected products "
        "(API, Data Layer, Token Layer, monitoring) serving 30+ markets. Defined test levels, "
        "integration points, test flows, and test frequency across the full lifecycle.",
        "Developed test plans covering integration testing between VCS components, end-to-end "
        "testing, and validation in realistic environments. Prioritized testing based on risk, "
        "complexity, and inter-product dependencies.",
        "Led automation at system level — built Playwright E2E framework with accessibility, "
        "CI/CD pipelines (GitHub Actions, Terraform), continuous integration, regression "
        "testing, and quick feedback loops. Drove test infrastructure and tool chain improvements.",
        "Conducted and followed up tests in close collaboration with developers, architects, "
        "and project managers across Spain, Sweden, India, and Paris. Built a culture where "
        "testing is integral to the development process.",
        "Security-focused: Cyber Jedi — implemented Threat Modeling, security validation, "
        "and non-functional testing (performance, robustness). Ensured solutions meet high "
        "standards of security alongside functional requirements.",
        "Mapped how products interact — identified critical integration points between VCS API, "
        "Data Layer (ETL/BigQuery), Token Layer, and external IKEA systems. Defined when and "
        "how often test activities are performed in the development cycle.",
    ]
    for b in ikea_bullets:
        add_highlighted_bullet(document, b)

    # -- Truecaller --
    cp2 = document.add_paragraph()
    cp2.paragraph_format.space_before = Pt(3)
    cp2.paragraph_format.space_after = Pt(0)
    add_text(cp2, "Truecaller, Stockholm — Senior QA Lead", bold=True, size=10, color=SECTION_COLOR)
    tp2 = document.add_paragraph()
    tp2.paragraph_format.space_after = Pt(0)
    add_text(tp2, "Sep 2021 – Feb 2022  |  Multi-Platform Product Ecosystem — 300M+ Users", bold=True, size=10)
    tc_bullets = [
        "Defined system-level test strategy for Truecaller's multi-platform ecosystem (iOS, "
        "Android, web) — integration testing between services, end-to-end validation, "
        "regression testing, and continuous integration across platforms.",
        "Prioritized testing efforts based on risk and complexity across a fast-moving product "
        "with 300M+ users. Led test automation improvements and CI/CD pipeline optimization "
        "for quick feedback and high release quality.",
        "Close collaboration with developers, architects, and project managers in a "
        "cross-functional, fast-paced environment. Drove a culture of testing as a natural "
        "part of development.",
    ]
    for b in tc_bullets:
        add_highlighted_bullet(document, b)

    # -- LEGO/IKEA via HCLTech --
    cp3 = document.add_paragraph()
    cp3.paragraph_format.space_before = Pt(3)
    cp3.paragraph_format.space_after = Pt(0)
    add_text(cp3, "LEGO Group & IKEA (via HCLTech) — Test Lead / Senior Test Engineer", bold=True, size=10, color=SECTION_COLOR)
    tp3 = document.add_paragraph()
    tp3.paragraph_format.space_after = Pt(0)
    add_text(tp3, "2016 – 2021  |  E-Commerce & Enterprise Product Ecosystems", bold=True, size=10)
    lego_bullets = [
        "Test Lead at LEGO: owned the test strategy for a complex e-commerce ecosystem — "
        "managed 8–10 engineers, defined test levels and integration points across "
        "interconnected services, API testing, and E2E validation for high-traffic web applications.",
        "IKEA App, One Solution (Genesys), Verint (2018–2021): designed integration test "
        "strategies across multiple IKEA products, planned test activities based on risk and "
        "dependencies, set up CI/CD and regression testing for continuous validation.",
        "Drove test automation improvements, mentored test engineers, and established test "
        "infrastructure and tool chains supporting continuous integration and quick feedback.",
    ]
    for b in lego_bullets:
        add_highlighted_bullet(document, b)

    # -- Banking / Earlier --
    cp4 = document.add_paragraph()
    cp4.paragraph_format.space_before = Pt(3)
    cp4.paragraph_format.space_after = Pt(0)
    add_text(cp4, "HCLTech, Samin TekMindz & Banking — Senior Engineer / Team Lead", bold=True, size=10, color=SECTION_COLOR)
    tp4 = document.add_paragraph()
    tp4.paragraph_format.space_after = Pt(0)
    add_text(tp4, "2008 – 2016  |  Banking, Fintech, Government & Enterprise — India", bold=True, size=10)
    fin_bullets = [
        "Designed test strategies for complex banking systems (Finacle) — integration testing "
        "between core banking modules, payment gateways, and transaction processing systems. "
        "Planned validation activities covering functional, non-functional, and security "
        "requirements.",
        "Led teams of 10+ engineers — mentored test engineers, drove test automation and CI/CD "
        "adoption, and built test infrastructure for continuous integration. Certified Ethical "
        "Hacker (CEH, 2017) — security testing and validation expertise.",
    ]
    for b in fin_bullets:
        add_highlighted_bullet(document, b)

    # ── Key Achievements ──
    add_section_heading(document, "Key Achievements")
    achievements = [
        "System-level test strategy for VCS ecosystem — defined test levels, integration points, "
        "flows, and frequency across 4 interconnected products serving 30+ markets.",
        "Playwright E2E automation at scale — CI/CD (GitHub Actions), continuous integration, "
        "regression, 3x reliability improvement. Built test infrastructure and tool chains.",
        "Cyber Jedi + CEH — Threat Modeling, security validation, non-functional testing. "
        "Ensured solutions meet high standards of security and robustness.",
        "Risk-based test prioritization across complex product ecosystems — Truecaller (300M+ "
        "users), LEGO e-commerce, IKEA multi-product platforms, and banking systems.",
        "Built testing culture across every organization — testing as a natural part of the "
        "entire development process, from early verification through acceptance testing.",
    ]
    for a in achievements:
        add_highlighted_bullet(document, a)

    # ── Education & Certifications ──
    add_section_heading(document, "Education & Certifications")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table, color="9DBAD5")

    for cell, title in zip(table.rows[0].cells, ["Education", "Certifications"]):
        set_cell_shading(cell, "EEF4FA")
        p = cell.paragraphs[0]
        add_text(p, title, bold=True, size=10, color=SECTION_COLOR)

    edu = table.cell(1, 0).paragraphs[0]
    add_text(edu, "M.Tech / B.Tech in Information Technology\nPG Diploma in Operations Management", size=9.5)
    certs = table.cell(1, 1).paragraphs[0]
    add_text(
        certs,
        "Certified Ethical Hacker (CEH) — 2017\n"
        "ISTQB Certified Tester Foundation\n"
        "Google Cloud – Associate Cloud Engineer\n"
        "AWS Cloud Practitioner\nITIL Foundation\n"
        "Six Sigma Green Belt\nUiPath RPA Certified",
        size=9.5,
    )

    # ── Languages ──
    add_section_heading(document, "Languages")
    lp = document.add_paragraph()
    lp.paragraph_format.space_after = Pt(0)
    add_text(lp, "English — Fluent  |  Swedish — Basic", size=10)

    document.save(DOCX_PATH)


def build_doc():
    html = """\
<html>
<head>
  <meta charset="utf-8">
  <title>Mohammad Kashif – Sectra Solution Test Lead Resume</title>
  <style>
    body { font-family: Calibri, Arial, sans-serif; margin: 24px; color: #1f2937; font-size: 10pt; }
    h1 { text-align: center; color: #1f4e79; margin-bottom: 2px; font-size: 18pt; }
    h2 { text-align: center; font-size: 10.5pt; margin-top: 0; margin-bottom: 4px; }
    .contact { text-align: center; color: #4b5563; font-size: 9.5pt; border-bottom: 2px solid #0058a3; padding-bottom: 6px; margin-bottom: 10px; }
    .section { color: #1f4e79; font-weight: 700; font-size: 10.5pt; border-bottom: 1.5px solid #1f4e79; padding-bottom: 2px; margin-top: 8px; margin-bottom: 4px; text-transform: uppercase; }
    .job-title { font-weight: 700; color: #1f4e79; margin-top: 4px; margin-bottom: 0; }
    .job-sub { font-weight: 700; margin-bottom: 2px; }
    ul { margin-top: 2px; margin-bottom: 4px; padding-left: 20px; }
    li { margin-bottom: 2px; }
    .hl { background: #fff59d; font-weight: 700; }
    table { width: 100%; border-collapse: collapse; margin-top: 4px; }
    td { border: 1px solid #b8cce4; padding: 5px; vertical-align: top; font-size: 9.5pt; }
    .tag { background: #d9eaf7; font-weight: 700; }
    p { margin: 2px 0; }
  </style>
</head>
<body>
  <h1>MOHAMMAD KASHIF</h1>
  <h2>Solution Test Lead | System-Level Test Strategy &amp; Automation | 15+ Years | Security-Focused</h2>
  <div class="contact">Malmö, Sweden | +46 702624230 | mo.kashif@gmail.com  |  LinkedIn: linkedin.com/in/md-kashif  |  Blog: ascertain.github.io</div>

  <div class="section">Profile</div>
  <p><span class="hl">Solution Test Lead</span> with <span class="hl">15+ years</span> taking overall responsibility for how <span class="hl">complex system</span>s of interconnected products are tested. Proven expertise in designing <span class="hl">system-level</span> <span class="hl">test strategy</span> — defining <span class="hl">test level</span>s, <span class="hl">integration point</span>s, test flows, and <span class="hl">test frequency</span> across the full lifecycle. Deep <span class="hl">security</span> background: <span class="hl">CEH</span>, <span class="hl">Cyber Jedi</span> with <span class="hl">Threat Modeling</span>. Track record leading <span class="hl">test automation</span> at scale — <span class="hl">Playwright</span> <span class="hl">E2E</span>, <span class="hl">CI/CD</span>, <span class="hl">continuous integration</span>, <span class="hl">regression</span>, and <span class="hl">test infrastructure</span> improvements. Close collaborator with <span class="hl">developers</span>, <span class="hl">architects</span>, <span class="hl">systems engineer</span>s, and <span class="hl">project manager</span>s. Driven by building a <span class="hl">culture</span> where testing is natural to development.</p>

  <div class="section">How I Match the Role</div>
  <p><b>Test Strategy for the System:</b> Designed how ecosystems of products should be tested — test levels, integration points, test flows, frequency. Functional, non-functional, and security requirements.<br>
  <b>Planning &amp; Coordination:</b> Test plans for integration testing, E2E testing, validation in realistic environments. Prioritized by risk, complexity, and dependencies.<br>
  <b>Automation &amp; Efficiency:</b> Led automation at system level — Playwright E2E, CI/CD (GitHub Actions), regression, quick feedback. Test infrastructure and tool chain improvements.<br>
  <b>Execution &amp; Analysis:</b> Tests in collaboration with development teams. Culture where testing is integral to development — early verification through acceptance testing.<br>
  <b>Security &amp; Robustness:</b> CEH certified. Cyber Jedi — Threat Modeling, security validation within SDLC. High standards of security and robustness.</p>

  <div class="section">Professional Experience</div>

  <div class="job-title">IKEA IT AB, Malmö — Solution Test Lead / Senior Software Engineer</div>
  <div class="job-sub">Mar 2022 – Present | VCS — Video Collaboration Solution (Multi-Product Ecosystem)</div>
  <ul>
    <li>Owned overall <span class="hl">test strategy</span> for VCS — <span class="hl">complex system</span> of interconnected products (API, Data Layer, Token Layer, monitoring) serving 30+ markets. Defined <span class="hl">test level</span>s, <span class="hl">integration point</span>s, test flows, <span class="hl">test frequency</span>.</li>
    <li>Developed <span class="hl">test plan</span>s — <span class="hl">integration testing</span> between components, <span class="hl">end-to-end</span> testing, <span class="hl">validation</span> in realistic environments. Prioritized by <span class="hl">risk</span>, complexity, inter-product dependencies.</li>
    <li>Led <span class="hl">automation</span> at <span class="hl">system-level</span> — <span class="hl">Playwright</span> <span class="hl">E2E</span>, <span class="hl">CI/CD</span> (<span class="hl">GitHub Actions</span>, <span class="hl">Terraform</span>), <span class="hl">continuous integration</span>, <span class="hl">regression</span>, quick feedback. <span class="hl">Test infrastructure</span> and <span class="hl">tool chain</span> improvements.</li>
    <li>Close collaboration with <span class="hl">developers</span>, <span class="hl">architects</span>, <span class="hl">project manager</span>s across Spain, Sweden, India, Paris. Built <span class="hl">culture</span> of testing as integral to development.</li>
    <li><span class="hl">Cyber Jedi</span> — <span class="hl">Threat Modeling</span>, <span class="hl">security</span> <span class="hl">validation</span>, <span class="hl">non-functional</span> testing. High standards of <span class="hl">security</span> and <span class="hl">robustness</span>.</li>
    <li>Mapped product interactions — critical <span class="hl">integration point</span>s between VCS API, Data Layer, Token Layer, external systems. Defined test activities across development cycle.</li>
  </ul>

  <div class="job-title">Truecaller, Stockholm — Senior QA Lead</div>
  <div class="job-sub">Sep 2021 – Feb 2022 | Multi-Platform Product Ecosystem — 300M+ Users</div>
  <ul>
    <li><span class="hl">System-level</span> <span class="hl">test strategy</span> for multi-platform <span class="hl">ecosystem</span> (iOS, Android, web) — <span class="hl">integration testing</span>, <span class="hl">E2E</span> <span class="hl">validation</span>, <span class="hl">regression</span>, <span class="hl">continuous integration</span>.</li>
    <li>Prioritized by <span class="hl">risk</span> and complexity across 300M+ users. <span class="hl">Test automation</span> improvements, <span class="hl">CI/CD</span> optimization, quick feedback.</li>
    <li><span class="hl">Cross-functional</span> collaboration with <span class="hl">developers</span>, <span class="hl">architects</span>, <span class="hl">project manager</span>s. Drove testing <span class="hl">culture</span>.</li>
  </ul>

  <div class="job-title">LEGO Group &amp; IKEA (via HCLTech) — Test Lead / Senior Test Engineer</div>
  <div class="job-sub">2016 – 2021 | E-Commerce &amp; Enterprise Product Ecosystems</div>
  <ul>
    <li>LEGO: <span class="hl">test strategy</span> for complex e-commerce <span class="hl">ecosystem</span> — 8–10 engineers, <span class="hl">test level</span>s, <span class="hl">integration point</span>s, API testing, <span class="hl">E2E</span> <span class="hl">validation</span>.</li>
    <li>IKEA App/Genesys/Verint (2018–21): <span class="hl">integration testing</span> across <span class="hl">multiple products</span>, <span class="hl">risk</span>-based planning, <span class="hl">CI/CD</span>, <span class="hl">regression</span>.</li>
    <li><span class="hl">Test automation</span> improvements, <span class="hl">mentor</span>ed engineers, built <span class="hl">test infrastructure</span> and <span class="hl">tool chain</span>s for <span class="hl">continuous integration</span>.</li>
  </ul>

  <div class="job-title">HCLTech, Samin TekMindz &amp; Banking — Senior Engineer / Team Lead</div>
  <div class="job-sub">2008 – 2016 | Banking, Fintech, Government &amp; Enterprise — India</div>
  <ul>
    <li><span class="hl">Test strategy</span> for complex <span class="hl">banking</span> systems (Finacle) — <span class="hl">integration testing</span>, <span class="hl">validation</span> of functional, <span class="hl">non-functional</span>, and <span class="hl">security</span> requirements.</li>
    <li>Led 10+ engineers, <span class="hl">mentor</span>ed, drove <span class="hl">test automation</span>/<span class="hl">CI/CD</span>, built <span class="hl">test infrastructure</span>. <span class="hl">CEH</span> certified (2017).</li>
  </ul>

  <div class="section">Key Achievements</div>
  <ul>
    <li><span class="hl">System-level</span> <span class="hl">test strategy</span> for VCS <span class="hl">ecosystem</span> — <span class="hl">test level</span>s, <span class="hl">integration point</span>s, flows, frequency across 4 products, 30+ markets.</li>
    <li><span class="hl">Playwright</span> <span class="hl">E2E</span> <span class="hl">automation</span> at scale — <span class="hl">CI/CD</span>, <span class="hl">continuous integration</span>, <span class="hl">regression</span>, 3x reliability. <span class="hl">Test infrastructure</span> &amp; <span class="hl">tool chain</span>s.</li>
    <li><span class="hl">Cyber Jedi</span> + <span class="hl">CEH</span> — <span class="hl">Threat Modeling</span>, <span class="hl">security</span> <span class="hl">validation</span>, <span class="hl">non-functional</span> testing. High <span class="hl">security</span> &amp; <span class="hl">robustness</span> standards.</li>
    <li><span class="hl">Risk</span>-based <span class="hl">test</span> prioritization — Truecaller (300M+), LEGO, IKEA, banking <span class="hl">ecosystem</span>s.</li>
    <li>Built testing <span class="hl">culture</span> — testing as natural part of development, early verification through <span class="hl">acceptance testing</span>.</li>
  </ul>

  <div class="section">Education &amp; Certifications</div>
  <table>
    <tr><td class="tag">Education</td><td class="tag">Certifications</td></tr>
    <tr>
      <td>M.Tech / B.Tech in Information Technology<br>PG Diploma in Operations Management</td>
      <td>Certified Ethical Hacker (CEH) — 2017<br>ISTQB Certified Tester Foundation<br>GCP Associate Cloud Engineer<br>AWS Cloud Practitioner<br>ITIL Foundation<br>Six Sigma Green Belt<br>UiPath RPA</td>
    </tr>
  </table>

  <div class="section">Languages</div>
  <p>English — Fluent  |  Swedish — Basic</p>
</body>
</html>
"""
    DOC_PATH.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_doc()
    print(DOCX_PATH)
    print(DOC_PATH)
