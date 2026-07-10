"""
Generate a tailored resume for Software Quality Assurance - Medical Device / Healthtech.
Focus: regulated environments, V&V, SaMD awareness, risk management, automation,
CI/CD, AI-assisted QA, cross-functional collaboration, independent ownership.
Output: DOCX + DOC (HTML-based).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_QA_MedTech_Resume.docx")
DOC_PATH = os.path.join(OUTPUT_DIR, "Mohammad_Kashif_QA_MedTech_Resume.doc")

# ─── Helpers ───────────────────────────────────────────────────────────────────

def set_paragraph_bottom_border(paragraph, color="4472C4", sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    ))

def set_table_borders(table, color="FFFFFF", sz="0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    ))

def add_text(paragraph, text, bold=False, size=Pt(10), color=RGBColor(0x33,0x33,0x33), font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold; run.font.size = size; run.font.color.rgb = color; run.font.name = font_name
    return run

def add_section_heading(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(7); p.space_after = Pt(3)
    set_paragraph_bottom_border(p)
    add_text(p, text, bold=True, size=Pt(10.5), color=RGBColor(0x1F,0x47,0x88))
    return p

HIGHLIGHT_TOKENS = [
    "quality assurance", "QA", "quality practices",
    "regulated environment", "regulated", "compliance",
    "verification and validation", "V&V",
    "Software as a Medical Device", "SaMD",
    "risk management", "risk",
    "automation", "automated testing", "test automation",
    "CI/CD", "DevOps", "pipeline",
    "AI", "AI-assisted", "machine learning",
    "Python", "pytest", "scripting",
    "cross-functional", "collaboration",
    "IEC 62304", "ISO 13485", "ISO 14971",
    "software development", "testing strategies",
    "ownership", "independent",
    "communication", "stakeholder",
]

def add_highlighted_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    remaining = text
    while remaining:
        matched = False
        for token in HIGHLIGHT_TOKENS:
            tl = token.lower(); rl = remaining.lower(); idx = rl.find(tl)
            if idx == 0:
                run = p.add_run(remaining[:len(token)]); run.font.size = Pt(10); run.font.name = "Calibri"
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW; remaining = remaining[len(token):]; matched = True; break
            elif idx > 0:
                run = p.add_run(remaining[:idx]); run.font.size = Pt(10); run.font.name = "Calibri"
                remaining = remaining[idx:]; matched = True; break
        if not matched:
            run = p.add_run(remaining); run.font.size = Pt(10); run.font.name = "Calibri"; remaining = ""
    return p

# ─── Build DOCX ───────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2); style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(0.8); section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2); section.right_margin = Cm(1.2)

    # ─── Header ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "MOHAMMAD KASHIF", bold=True, size=Pt(18), color=RGBColor(0x1F,0x47,0x88))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(2)
    add_text(p, "Software Quality Assurance Engineer", bold=True, size=Pt(11), color=RGBColor(0x33,0x33,0x33))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(4)
    add_text(p, "Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  ", size=Pt(9))
    add_text(p, "linkedin.com/in/md-kashif", size=Pt(9), color=RGBColor(0x1F,0x47,0x88))

    # ─── Professional Summary ──────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    summary = (
        "Software Quality Assurance professional working in IT since 2008, with extensive experience in "
        "QA within heavily regulated environments (financial services, enterprise platforms). Strong foundation "
        "in software development processes, testing strategies, and development pipelines — combining manual "
        "testing expertise with hands-on test automation (Python, pytest, Playwright) and CI/CD integration. "
        "Experienced in embedding quality practices into development workflows, driving verification and validation "
        "activities, and supporting risk management and compliance efforts. Comfortable using AI tools and automation "
        "to improve QA workflows and reduce manual effort. Independent, ownership-driven professional who collaborates "
        "closely with engineering and product teams in fast-moving, cross-functional environments. Theoretical knowledge "
        "of medical device regulatory frameworks (IEC 62304, ISO 13485, ISO 14971) with strong interest in applying "
        "these in a healthtech/SaMD context."
    )
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_text(p, summary)

    # ─── Core Competencies ─────────────────────────────────────────────────
    add_section_heading(doc, "CORE COMPETENCIES")
    table = doc.add_table(rows=3, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    competencies = [
        "Software Quality Assurance", "Regulated Environment Experience", "Verification & Validation",
        "Test Automation (Python/pytest)", "CI/CD & DevOps Pipelines", "Risk Management & Compliance",
        "AI-Assisted QA & Tooling", "Cross-Functional Collaboration", "Independent Ownership & Initiative",
    ]
    for i, comp in enumerate(competencies):
        cell = table.rows[i // 3].cells[i % 3]; cell.paragraphs[0].clear()
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, f"• {comp}", size=Pt(9))

    # ─── Professional Experience ────────────────────────────────────────────
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    # --- IKEA IT AB ---
    p = doc.add_paragraph(); p.space_before = Pt(4)
    add_text(p, "IKEA IT AB (Ingka Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Malmö, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior QA Engineer / SDET — Cloud & IoT Platform", bold=True, size=Pt(10))
    add_text(p, "    Mar 2022 – Present", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Establish strong collaboration with engineering and product teams to embed quality practices into development workflows — ensuring quality is built in from design through deployment.",
        "Drive verification and validation activities across software releases — defining test strategies, acceptance criteria, and traceability from requirements to test evidence.",
        "Support risk management and compliance activities for software systems — identifying quality risks, maintaining risk registers, and ensuring mitigation actions are tracked to closure.",
        "Improve QA workflows by introducing automation, tooling, and AI-assisted processes — leveraging Python (pytest), Playwright, and AI code-assist tools to increase efficiency and reduce manual work.",
        "Develop and maintain automated testing frameworks integrated into CI/CD pipelines (GitHub Actions) — ensuring continuous quality feedback within DevOps workflows.",
        "Identify opportunities to strengthen quality systems and continuously improve development and testing processes — introduced risk-based testing, structured exploratory sessions, and automated regression that reduced test cycle time by 40%.",
        "Provide clear communication on quality risks, progress, and issues — producing quality dashboards, defect trend reports, and release readiness assessments for stakeholders.",
        "Work independently and take ownership of QA activities in a fast-moving agile environment — managing priorities, adapting to evolving processes, and driving quality improvements without constant oversight.",
        "Collaborate closely with software engineers in cross-functional product teams (Scrum) — participating in design reviews, code reviews, and sprint ceremonies to ensure testability and quality throughout.",
        "Apply automation and scripting (Python, Bash) to improve workflows — building tools for test data generation, environment provisioning, and quality metric collection.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Truecaller ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Truecaller", bold=True, size=Pt(10))
    add_text(p, "  |  Stockholm, Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "QA & Release Engineer", bold=True, size=Pt(10))
    add_text(p, "    Sep 2021 – Feb 2022", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Embedded quality assurance practices into development workflows for a 300M+ user platform — collaborating closely with software engineers to ensure release quality.",
        "Developed Python-based automated testing integrated into CI/CD pipelines — improving QA efficiency and reducing manual regression effort per release cycle.",
        "Provided clear communication on quality risks and release readiness to stakeholders — enabling data-driven go/no-go decisions for production deployments.",
        "Worked independently in a fast-scaling organization with evolving processes — taking ownership of QA activities and adapting quickly to changing priorities.",
        "Supported compliance and risk management activities — ensuring quality gates and validation steps were met before production releases.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- HCLTech ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "HCLTech (for IKEA & LEGO Group)", bold=True, size=Pt(10))
    add_text(p, "  |  Denmark & Sweden", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    p = doc.add_paragraph(); p.space_after = Pt(2)
    add_text(p, "Senior QA Engineer / Test Lead — Enterprise & IoT", bold=True, size=Pt(10))
    add_text(p, "    Jun 2013 – Sep 2021", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Established quality assurance practices within development teams across IKEA IoT/smart-home and LEGO digital platforms — embedding quality into software development workflows from early stages.",
        "Performed verification and validation activities for complex software systems — ensuring traceability from requirements through test design, execution, and defect resolution.",
        "Supported risk management and compliance in regulated enterprise environments — maintaining quality documentation, risk assessments, and audit-ready test evidence.",
        "Improved QA workflows through automation and tooling — building automated test suites (Python, pytest, Selenium) integrated into CI/CD pipelines (Jenkins), reducing manual effort by 60%.",
        "Collaborated closely with software engineers and cross-functional product teams — participating in design discussions, reviewing acceptance criteria, and ensuring testability of features.",
        "Worked in DevOps environments with CI/CD pipelines and automated testing frameworks — ensuring continuous quality feedback and fast release cycles.",
        "Provided clear communication on quality risks, progress, and issues to project stakeholders and steering committees — producing status reports, defect analyses, and quality metrics.",
        "Took independent ownership of QA activities across multiple workstreams — managing test planning, execution, and reporting with minimal supervision.",
        "Improved workflows through scripting and tooling (Python, shell) — automating test data preparation, environment setup, and reporting processes.",
        "Worked in fast-scaling organizations with evolving processes — adapting QA approaches to growing teams, changing architectures, and maturing delivery pipelines.",
    ]:
        add_highlighted_bullet(doc, b)

    # --- Earlier Career ---
    p = doc.add_paragraph(); p.space_before = Pt(6)
    add_text(p, "Earlier Career — QA in Regulated Financial Environments", bold=True, size=Pt(10))
    add_text(p, "  |  India  |  2008 – 2013", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    for b in [
        "Performed software quality assurance in heavily regulated financial environments (core banking, insurance) — similar compliance rigor to medical device/healthtech standards.",
        "Executed verification and validation for enterprise software systems — ensuring regulatory compliance, data integrity, and audit readiness.",
        "Supported risk management activities — identifying, documenting, and tracking software quality risks through structured processes.",
        "Developed automation scripts (Python, shell) to improve QA workflows — reducing manual testing effort and improving repeatability.",
        "Collaborated with software engineers in cross-functional teams — embedding quality practices into development processes.",
    ]:
        add_highlighted_bullet(doc, b)

    # ─── Regulatory Knowledge ──────────────────────────────────────────────
    add_section_heading(doc, "REGULATORY & DOMAIN KNOWLEDGE")
    p = doc.add_paragraph(); p.space_after = Pt(4)
    add_text(p, (
        "Theoretical knowledge and active study of medical device regulatory frameworks including "
        "IEC 62304 (software lifecycle), ISO 13485 (QMS), ISO 14971 (risk management), and MDR/FDA requirements. "
        "Strong interest in applying this knowledge hands-on in a healthtech/SaMD context. Extensive practical "
        "experience with similar compliance rigor from regulated financial services (audit readiness, traceability, "
        "risk documentation, validation evidence)."
    ))

    # ─── Technical Skills ───────────────────────────────────────────────────
    add_section_heading(doc, "TECHNICAL SKILLS")
    for label, value in [
        ("Automation & Testing: ", "Python (pytest, requests, paramiko), Selenium, Playwright, API testing, shell scripting"),
        ("CI/CD & DevOps: ", "GitHub Actions, Jenkins, GitLab CI — automated test integration, quality gates, pipeline orchestration"),
        ("AI & Tooling: ", "AI-assisted development tools (GitHub Copilot, AI code review), workflow automation, productivity tooling"),
        ("Platforms: ", "GCP, Azure, AWS, Linux, Docker, Kubernetes"),
        ("Quality Tools: ", "Jira, Confluence, Azure DevOps, Zephyr, Git, Postman"),
        ("Methodologies: ", "Agile (Scrum/Kanban), risk-based testing, exploratory testing, ISTQB, TDD/BDD, V&V practices"),
    ]:
        p = doc.add_paragraph(); p.space_after = Pt(2)
        add_text(p, label, bold=True, size=Pt(9))
        add_text(p, value, size=Pt(9))

    # ─── Education ──────────────────────────────────────────────────────────
    add_section_heading(doc, "EDUCATION")
    for deg, school in [
        ("Post Graduate Diploma in Operation and Management", "IGNOU, India"),
        ("B.Tech, Information Technology", "UP Technical University, India"),
    ]:
        p = doc.add_paragraph()
        add_text(p, deg, bold=True, size=Pt(10))
        add_text(p, f"  —  {school}", size=Pt(9), color=RGBColor(0x55,0x55,0x55))

    # ─── Certifications ─────────────────────────────────────────────────────
    add_section_heading(doc, "CERTIFICATIONS")
    certs = [
        "ISTQB Certified Tester — Foundation Level",
        "Google Cloud Associate Cloud Engineer (ACE)",
        "AWS Certified Cloud Practitioner",
        "Certified Ethical Hacker (CEH)",
        "ITIL v4 Foundation",
        "Six Sigma Green Belt",
    ]
    table = doc.add_table(rows=3, cols=2); set_table_borders(table)
    for i, cert in enumerate(certs):
        cell = table.rows[i // 2].cells[i % 2]; cell.paragraphs[0].clear()
        add_text(cell.paragraphs[0], f"• {cert}", size=Pt(9))

    # ─── Languages ──────────────────────────────────────────────────────────
    add_section_heading(doc, "LANGUAGES")
    p = doc.add_paragraph()
    add_text(p, "English (Fluent)", size=Pt(9))

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

# ─── Build DOC (HTML-based) ────────────────────────────────────────────────────

def build_doc():
    html = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;color:#333;}
h1{text-align:center;font-size:18pt;color:#1F4788;margin:0;}
h2{font-size:10.5pt;color:#1F4788;border-bottom:2px solid #4472C4;padding-bottom:2px;margin-top:10px;margin-bottom:4px;}
.subtitle{text-align:center;font-size:11pt;font-weight:bold;margin:2px 0;}
.contact{text-align:center;font-size:9pt;margin-bottom:8px;}
ul{margin:2px 0 2px 18px;padding:0;}
li{margin-bottom:2px;}
.comp-table{width:100%;border-collapse:collapse;margin:4px 0;}
.comp-table td{text-align:center;font-size:9pt;padding:2px 4px;}
.role-header{font-weight:bold;margin-top:8px;margin-bottom:1px;}
.company{font-weight:bold;}
.date{color:#555;font-size:9pt;}
.tech-line{font-size:9pt;margin:2px 0;}
.cert-table{width:100%;border-collapse:collapse;}
.cert-table td{font-size:9pt;padding:1px 4px;}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="subtitle">Software Quality Assurance Engineer</p>
<p class="contact">Malm&ouml;, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Software Quality Assurance professional working in IT since 2008, with extensive experience in QA within heavily regulated environments (financial services, enterprise platforms). Strong foundation in software development processes, testing strategies, and development pipelines &mdash; combining manual testing expertise with hands-on test automation (Python, pytest, Playwright) and CI/CD integration. Experienced in embedding quality practices into development workflows, driving verification and validation activities, and supporting risk management and compliance efforts. Comfortable using AI tools and automation to improve QA workflows and reduce manual effort. Independent, ownership-driven professional who collaborates closely with engineering and product teams in fast-moving, cross-functional environments. Theoretical knowledge of medical device regulatory frameworks (IEC 62304, ISO 13485, ISO 14971) with strong interest in applying these in a healthtech/SaMD context.</p>

<h2>CORE COMPETENCIES</h2>
<table class="comp-table">
<tr><td>&bull; Software Quality Assurance</td><td>&bull; Regulated Environment Experience</td><td>&bull; Verification &amp; Validation</td></tr>
<tr><td>&bull; Test Automation (Python/pytest)</td><td>&bull; CI/CD &amp; DevOps Pipelines</td><td>&bull; Risk Management &amp; Compliance</td></tr>
<tr><td>&bull; AI-Assisted QA &amp; Tooling</td><td>&bull; Cross-Functional Collaboration</td><td>&bull; Independent Ownership &amp; Initiative</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role-header"><span class="company">IKEA IT AB (Ingka Group)</span> <span class="date">&nbsp;|&nbsp; Malm&ouml;, Sweden</span></p>
<p class="role-header">Senior QA Engineer / SDET &mdash; Cloud &amp; IoT Platform <span class="date">&nbsp;&nbsp;Mar 2022 &ndash; Present</span></p>
<ul>
<li>Establish strong collaboration with engineering and product teams to embed quality practices into development workflows &mdash; ensuring quality is built in from design through deployment.</li>
<li>Drive verification and validation activities across software releases &mdash; defining test strategies, acceptance criteria, and traceability from requirements to test evidence.</li>
<li>Support risk management and compliance activities for software systems &mdash; identifying quality risks, maintaining risk registers, and ensuring mitigation actions are tracked to closure.</li>
<li>Improve QA workflows by introducing automation, tooling, and AI-assisted processes &mdash; leveraging Python (pytest), Playwright, and AI code-assist tools to increase efficiency and reduce manual work.</li>
<li>Develop and maintain automated testing frameworks integrated into CI/CD pipelines (GitHub Actions) &mdash; ensuring continuous quality feedback within DevOps workflows.</li>
<li>Identify opportunities to strengthen quality systems and continuously improve development and testing processes &mdash; introduced risk-based testing, structured exploratory sessions, and automated regression that reduced test cycle time by 40%.</li>
<li>Provide clear communication on quality risks, progress, and issues &mdash; producing quality dashboards, defect trend reports, and release readiness assessments for stakeholders.</li>
<li>Work independently and take ownership of QA activities in a fast-moving agile environment &mdash; managing priorities, adapting to evolving processes, and driving quality improvements without constant oversight.</li>
<li>Collaborate closely with software engineers in cross-functional product teams (Scrum) &mdash; participating in design reviews, code reviews, and sprint ceremonies to ensure testability and quality throughout.</li>
<li>Apply automation and scripting (Python, Bash) to improve workflows &mdash; building tools for test data generation, environment provisioning, and quality metric collection.</li>
</ul>

<p class="role-header"><span class="company">Truecaller</span> <span class="date">&nbsp;|&nbsp; Stockholm, Sweden</span></p>
<p class="role-header">QA &amp; Release Engineer <span class="date">&nbsp;&nbsp;Sep 2021 &ndash; Feb 2022</span></p>
<ul>
<li>Embedded quality assurance practices into development workflows for a 300M+ user platform &mdash; collaborating closely with software engineers to ensure release quality.</li>
<li>Developed Python-based automated testing integrated into CI/CD pipelines &mdash; improving QA efficiency and reducing manual regression effort per release cycle.</li>
<li>Provided clear communication on quality risks and release readiness to stakeholders &mdash; enabling data-driven go/no-go decisions for production deployments.</li>
<li>Worked independently in a fast-scaling organization with evolving processes &mdash; taking ownership of QA activities and adapting quickly to changing priorities.</li>
<li>Supported compliance and risk management activities &mdash; ensuring quality gates and validation steps were met before production releases.</li>
</ul>

<p class="role-header"><span class="company">HCLTech (for IKEA &amp; LEGO Group)</span> <span class="date">&nbsp;|&nbsp; Denmark &amp; Sweden</span></p>
<p class="role-header">Senior QA Engineer / Test Lead &mdash; Enterprise &amp; IoT <span class="date">&nbsp;&nbsp;Jun 2013 &ndash; Sep 2021</span></p>
<ul>
<li>Established quality assurance practices within development teams across IKEA IoT/smart-home and LEGO digital platforms &mdash; embedding quality into software development workflows from early stages.</li>
<li>Performed verification and validation activities for complex software systems &mdash; ensuring traceability from requirements through test design, execution, and defect resolution.</li>
<li>Supported risk management and compliance in regulated enterprise environments &mdash; maintaining quality documentation, risk assessments, and audit-ready test evidence.</li>
<li>Improved QA workflows through automation and tooling &mdash; building automated test suites (Python, pytest, Selenium) integrated into CI/CD pipelines (Jenkins), reducing manual effort by 60%.</li>
<li>Collaborated closely with software engineers and cross-functional product teams &mdash; participating in design discussions, reviewing acceptance criteria, and ensuring testability of features.</li>
<li>Worked in DevOps environments with CI/CD pipelines and automated testing frameworks &mdash; ensuring continuous quality feedback and fast release cycles.</li>
<li>Provided clear communication on quality risks, progress, and issues to project stakeholders and steering committees &mdash; producing status reports, defect analyses, and quality metrics.</li>
<li>Took independent ownership of QA activities across multiple workstreams &mdash; managing test planning, execution, and reporting with minimal supervision.</li>
<li>Improved workflows through scripting and tooling (Python, shell) &mdash; automating test data preparation, environment setup, and reporting processes.</li>
<li>Worked in fast-scaling organizations with evolving processes &mdash; adapting QA approaches to growing teams, changing architectures, and maturing delivery pipelines.</li>
</ul>

<p class="role-header"><span class="company">Earlier Career &mdash; QA in Regulated Financial Environments</span> <span class="date">&nbsp;|&nbsp; India &nbsp;&nbsp;2008 &ndash; 2013</span></p>
<ul>
<li>Performed software quality assurance in heavily regulated financial environments (core banking, insurance) &mdash; similar compliance rigor to medical device/healthtech standards.</li>
<li>Executed verification and validation for enterprise software systems &mdash; ensuring regulatory compliance, data integrity, and audit readiness.</li>
<li>Supported risk management activities &mdash; identifying, documenting, and tracking software quality risks through structured processes.</li>
<li>Developed automation scripts (Python, shell) to improve QA workflows &mdash; reducing manual testing effort and improving repeatability.</li>
<li>Collaborated with software engineers in cross-functional teams &mdash; embedding quality practices into development processes.</li>
</ul>

<h2>REGULATORY &amp; DOMAIN KNOWLEDGE</h2>
<p>Theoretical knowledge and active study of medical device regulatory frameworks including IEC 62304 (software lifecycle), ISO 13485 (QMS), ISO 14971 (risk management), and MDR/FDA requirements. Strong interest in applying this knowledge hands-on in a healthtech/SaMD context. Extensive practical experience with similar compliance rigor from regulated financial services (audit readiness, traceability, risk documentation, validation evidence).</p>

<h2>TECHNICAL SKILLS</h2>
<p class="tech-line"><b>Automation &amp; Testing:</b> Python (pytest, requests, paramiko), Selenium, Playwright, API testing, shell scripting</p>
<p class="tech-line"><b>CI/CD &amp; DevOps:</b> GitHub Actions, Jenkins, GitLab CI &mdash; automated test integration, quality gates, pipeline orchestration</p>
<p class="tech-line"><b>AI &amp; Tooling:</b> AI-assisted development tools (GitHub Copilot, AI code review), workflow automation, productivity tooling</p>
<p class="tech-line"><b>Platforms:</b> GCP, Azure, AWS, Linux, Docker, Kubernetes</p>
<p class="tech-line"><b>Quality Tools:</b> Jira, Confluence, Azure DevOps, Zephyr, Git, Postman</p>
<p class="tech-line"><b>Methodologies:</b> Agile (Scrum/Kanban), risk-based testing, exploratory testing, ISTQB, TDD/BDD, V&amp;V practices</p>

<h2>EDUCATION</h2>
<p><b>Post Graduate Diploma in Operation and Management</b> &mdash; IGNOU, India</p>
<p><b>B.Tech, Information Technology</b> &mdash; UP Technical University, India</p>

<h2>CERTIFICATIONS</h2>
<table class="cert-table">
<tr><td>&bull; ISTQB Certified Tester &mdash; Foundation Level</td><td>&bull; Google Cloud Associate Cloud Engineer (ACE)</td></tr>
<tr><td>&bull; AWS Certified Cloud Practitioner</td><td>&bull; Certified Ethical Hacker (CEH)</td></tr>
<tr><td>&bull; ITIL v4 Foundation</td><td>&bull; Six Sigma Green Belt</td></tr>
</table>

<h2>LANGUAGES</h2>
<p style="font-size:9pt;">English (Fluent)</p>
</body></html>"""
    with open(DOC_PATH, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"DOC saved: {DOC_PATH}")

if __name__ == "__main__":
    build_docx()
    build_doc()
