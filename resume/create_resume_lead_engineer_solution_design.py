"""Generate Lead Engineer Solution Design resume – IKEA Internal (FUO Domain)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_Lead_Engineer_Solution_Design_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0058A3")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x00, 0x58, 0xA3)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "Lead Engineer with 12+ years of experience in software engineering, solution design, and distributed systems "
        "— with a strong focus on setting technical direction, raising engineering quality, and enabling teams across domains. "
        "Proven track record of defining common standards and design patterns, reviewing solution designs, and coaching engineers "
        "on pragmatic, scalable approaches in complex enterprise environments. "
        "Deep experience leading transformation and modernization journeys — migrating legacy systems, re-architecting services, "
        "and evolving platforms in production within the IKEA ecosystem. "
        "Skilled at influencing technical direction without direct authority, working across teams and partnering with architects, "
        "service management, and leadership to ensure solutions are aligned, fit for purpose, and meet non-functional requirements "
        "(scalability, resilience, maintainability, cost) from the start. "
        "Comfortable working in ambiguous environments where the path forward requires creativity, pragmatism, and clear decision-making. "
        "Strong advocate for DevOps culture, cloud-native technologies, continuous improvement, and knowledge sharing across "
        "the broader Ingka engineering community. Aligned with IKEA values — togetherness, simplicity, and leading by example."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("Solution Design", "Distributed systems architecture · Design patterns & standards · Solution reviews & feedback · Non-functional requirements (scalability, resilience, cost, maintainability) · Modularization · Integration patterns"),
        ("Technical Leadership", "Setting technical direction across teams · Influencing without authority · Clear technical decision-making · Constructive pushback · Driving technology agenda · Proof-of-concepts"),
        ("Coaching & Enablement", "Mentoring engineers on solution design · AI approaches · Technical problem-solving · Raising quality across teams · Knowledge sharing · Engineering community building"),
        ("Cloud & Backend", "GCP (Cloud Run, BigQuery, Pub/Sub, Cloud Functions, GCS, Dataflow) · AWS · Cloud-native patterns · Microservices · Event-driven architecture · Serverless · APIs (REST, gRPC)"),
        ("Modernization", "Legacy system migration · Service re-architecture · Platform evolution in production · Incremental modernization strategies · Strangler fig pattern"),
        ("DevOps & Delivery", "CI/CD (GitHub Actions) · Terraform · Docker · Kubernetes · Infrastructure as Code · Monitoring & observability · GitOps · Agile/Scrum · Self-organizing teams"),
        ("Programming", "Python · TypeScript · SQL · Node.js · Java (earlier career) · Shell scripting"),
        ("Certifications", "Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "E8F4FD")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - IKEA VCS
    add_role(doc, "Technical Lead / Team Lead Acting", "IKEA Customer Connect, Ingka Digital", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Define and maintain common standards and design patterns for solution design across the product area — ensuring consistency, quality, and alignment with Ingka's engineering direction.",
        "Review solution designs from engineers across teams — providing clear, constructive feedback that raises quality and drives architectural consistency across the domain.",
        "Maintain a holistic overview of systems and their dependencies — ensuring teams have the context they need when designing new solutions or evolving existing ones.",
        "Ensure non-functional requirements (scalability, resilience, maintainability, cost) are considered from the start — embedded into design reviews and standards rather than treated as afterthoughts.",
        "Run proof-of-concepts to evaluate new technologies and patterns (AI/ML approaches, event-driven architectures, new GCP services) before wider adoption — providing pragmatic recommendations.",
        "Coach and support engineers on solution design, AI approaches, and technical problem-solving — raising the overall technical quality and confidence across the team.",
        "Work closely with PEDX community including Architects, Service Management, and Ingka enabling teams — ensuring solutions are aligned, fit for purpose, and meet enterprise guidelines.",
        "Drive the technology agenda across the product area with a pragmatic, tangible approach — identifying technology opportunities and driving them into our ways of working.",
        "Make clear technical decisions and push back on designs that do not meet the bar — with honest reasoning and a constructive way forward for the teams.",
        "Led a major transformation journey — re-architecting the data layer from batch-based ingestion to event-driven, real-time pipelines serving 32 markets (Python, GCP, BigQuery, Pub/Sub, Terraform).",
        "Share knowledge and good practices across the broader Ingka engineering community — contributing to forums, documentation, and cross-domain collaboration.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "Platform Engineer", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Designed scalable, cloud-native solutions for a globally distributed platform (300M+ users) — working in a complex, ambiguous environment requiring pragmatic design trade-offs.",
        "Contributed to solution design reviews and infrastructure pattern harmonization — improving consistency across engineering teams.",
        "Worked across teams and influenced technical direction — driving CI/CD automation and infrastructure-as-code adoption without direct authority.",
        "Evaluated and adopted new technologies for deployment and observability — running proof-of-concepts before broader rollout.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "Technical Lead — Solution Design & Engineering", "HCLTech (for IKEA & LEGO Group)", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led solution design across multiple enterprise products — defining patterns, reviewing designs, and ensuring scalability, resilience, and maintainability across complex, integrated systems.",
        "Drove transformation and modernization journeys — migrating legacy systems to cloud-native architectures (GCP, AWS), re-architecting monoliths into microservices, and evolving platforms in production.",
        "Coached and mentored a team of 8–12 engineers — raising technical quality through design reviews, pair programming, and structured knowledge sharing.",
        "Worked across teams and influenced technical direction without direct authority — aligning multiple product teams around shared standards, patterns, and technology choices.",
        "Identified technology opportunities and drove them into established ways of working — from proof-of-concept through adoption (containerization, event-driven patterns, Infrastructure as Code).",
        "Collaborated closely with architects, product teams, and business stakeholders — ensuring solution designs were aligned with strategic direction and fit for purpose.",
        "Maintained holistic system oversight and dependency mapping — giving teams the context needed to make informed design decisions in complex, multi-system environments.",
        "Delivered solutions combining Platforms, SaaS (Genesys, Verint), and home-grown systems — demonstrating comfort in complex technology landscapes with diverse constraints.",
        "Pushed back on designs that didn't meet quality standards — providing honest, constructive reasoning and guiding engineers toward better approaches.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Software Engineer", "Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz)", "India", "2008 – 2013")
    bullets_4 = [
        "Built enterprise software solutions across financial and retail domains — gaining foundational experience in distributed systems design, integration patterns, and backend engineering.",
        "Contributed to modernization projects — migrating legacy systems to scalable architectures and establishing reusable design patterns.",
        "Developed platform components, APIs, and integration layers using Java, Python, and SQL — supporting complex business workflows.",
        "Worked in agile, self-organizing, cross-functional teams — building collaborative delivery practices from early career.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── IKEA Ecosystem Experience ─────────────────────────────────────────
    add_heading_block(doc, "IKEA Ecosystem Experience")
    ikea_items = [
        "IKEA Customer Connect (VCS): End-to-end solution design and technical leadership — cloud-native architecture (GCP), data pipelines, APIs, event-driven systems serving 32 markets.",
        "Solution Design Standards: Defining and maintaining common design patterns and standards across product teams — aligned with Ingka's engineering direction.",
        "Transformation & Modernization: Led migration from legacy batch systems to real-time, event-driven architectures — evolving platforms in production with zero-downtime approach.",
        "PEDX Community Collaboration: Working with Architects, Service Management, and enabling teams — ensuring solutions are aligned, scalable, and fit for purpose.",
        "Fulfilment & Core Services Adjacency: Systems integration with downstream fulfilment services, logistics tracking, and operational systems within the IKEA value chain.",
        "IKEA Ways of Working: Product team model, Agile/Scrum, DevOps set-up, self-organizing teams, Ingka DevOps tooling, cross-functional collaboration.",
        "Knowledge Sharing: Active contributor to Ingka engineering community — forums, documentation, design pattern libraries, and cross-domain collaboration.",
    ]
    for item in ikea_items:
        bullet(doc, item, bold_prefix=item.split(":")[0] + ":")

    # ─── IKEA Values ───────────────────────────────────────────────────────
    add_heading_block(doc, "IKEA Culture & Values Alignment")
    values = [
        "Togetherness: Work hands-on and across teams — enabling engineers with support, clarity, and constructive collaboration to deliver better solutions together.",
        "Leading by Example: Coach, mentor, and raise technical quality — setting the standard through design reviews, proof-of-concepts, and knowledge sharing.",
        "Simplicity: Define pragmatic, maintainable standards and patterns that teams can adopt — avoiding over-engineering while meeting real non-functional needs.",
        "Cost-Consciousness: Embed cost as a first-class non-functional requirement — ensuring solutions are economically sustainable from the design phase.",
        "Daring to be Different: Explore new technologies and AI approaches with proof-of-concepts — driving innovation with a pragmatic, tangible approach.",
    ]
    for v in values:
        bullet(doc, v, bold_prefix=v.split(":")[0] + ":")

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#0058A3;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#0058A3;font-size:10.5pt;border-bottom:1px solid #0058A3;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#E8F4FD;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>Lead Engineer with 12+ years of experience in software engineering, solution design, and distributed systems — with a strong focus on setting technical direction, raising engineering quality, and enabling teams across domains. Proven track record of defining common standards and design patterns, reviewing solution designs, and coaching engineers on pragmatic, scalable approaches in complex enterprise environments. Deep experience leading transformation and modernization journeys — migrating legacy systems, re-architecting services, and evolving platforms in production within the IKEA ecosystem. Skilled at influencing technical direction without direct authority, working across teams and partnering with architects, service management, and leadership to ensure solutions are aligned, fit for purpose, and meet non-functional requirements (scalability, resilience, maintainability, cost) from the start. Comfortable working in ambiguous environments where the path forward requires creativity, pragmatism, and clear decision-making. Strong advocate for DevOps culture, cloud-native technologies, continuous improvement, and knowledge sharing across the broader Ingka engineering community. Aligned with IKEA values — togetherness, simplicity, and leading by example.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">Solution Design</td><td>Distributed systems architecture · Design patterns &amp; standards · Solution reviews &amp; feedback · Non-functional requirements (scalability, resilience, cost, maintainability) · Modularization · Integration patterns</td></tr>
<tr><td class="cat">Technical Leadership</td><td>Setting technical direction across teams · Influencing without authority · Clear technical decision-making · Constructive pushback · Driving technology agenda · Proof-of-concepts</td></tr>
<tr><td class="cat">Coaching &amp; Enablement</td><td>Mentoring engineers on solution design · AI approaches · Technical problem-solving · Raising quality across teams · Knowledge sharing · Engineering community building</td></tr>
<tr><td class="cat">Cloud &amp; Backend</td><td>GCP (Cloud Run, BigQuery, Pub/Sub, Cloud Functions, GCS, Dataflow) · AWS · Cloud-native patterns · Microservices · Event-driven architecture · Serverless · APIs (REST, gRPC)</td></tr>
<tr><td class="cat">Modernization</td><td>Legacy system migration · Service re-architecture · Platform evolution in production · Incremental modernization strategies · Strangler fig pattern</td></tr>
<tr><td class="cat">DevOps &amp; Delivery</td><td>CI/CD (GitHub Actions) · Terraform · Docker · Kubernetes · Infrastructure as Code · Monitoring &amp; observability · GitOps · Agile/Scrum · Self-organizing teams</td></tr>
<tr><td class="cat">Programming</td><td>Python · TypeScript · SQL · Node.js · Java (earlier career) · Shell scripting</td></tr>
<tr><td class="cat">Certifications</td><td>Google Cloud Associate Cloud Engineer · AWS Cloud Practitioner · ISTQB · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">Technical Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; IKEA Customer Connect, Ingka Digital &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Define and maintain common standards and design patterns for solution design across the product area — ensuring consistency, quality, and alignment with Ingka's engineering direction.</li>
<li>Review solution designs from engineers across teams — providing clear, constructive feedback that raises quality and drives architectural consistency across the domain.</li>
<li>Maintain a holistic overview of systems and their dependencies — ensuring teams have the context they need when designing new solutions or evolving existing ones.</li>
<li>Ensure non-functional requirements (scalability, resilience, maintainability, cost) are considered from the start — embedded into design reviews and standards rather than treated as afterthoughts.</li>
<li>Run proof-of-concepts to evaluate new technologies and patterns (AI/ML approaches, event-driven architectures, new GCP services) before wider adoption — providing pragmatic recommendations.</li>
<li>Coach and support engineers on solution design, AI approaches, and technical problem-solving — raising the overall technical quality and confidence across the team.</li>
<li>Work closely with PEDX community including Architects, Service Management, and Ingka enabling teams — ensuring solutions are aligned, fit for purpose, and meet enterprise guidelines.</li>
<li>Drive the technology agenda across the product area with a pragmatic, tangible approach — identifying technology opportunities and driving them into our ways of working.</li>
<li>Make clear technical decisions and push back on designs that do not meet the bar — with honest reasoning and a constructive way forward for the teams.</li>
<li>Led a major transformation journey — re-architecting the data layer from batch-based ingestion to event-driven, real-time pipelines serving 32 markets (Python, GCP, BigQuery, Pub/Sub, Terraform).</li>
<li>Share knowledge and good practices across the broader Ingka engineering community — contributing to forums, documentation, and cross-domain collaboration.</li>
</ul>

<p class="role">Platform Engineer <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Designed scalable, cloud-native solutions for a globally distributed platform (300M+ users) — working in a complex, ambiguous environment requiring pragmatic design trade-offs.</li>
<li>Contributed to solution design reviews and infrastructure pattern harmonization — improving consistency across engineering teams.</li>
<li>Worked across teams and influenced technical direction — driving CI/CD automation and infrastructure-as-code adoption without direct authority.</li>
<li>Evaluated and adopted new technologies for deployment and observability — running proof-of-concepts before broader rollout.</li>
</ul>

<p class="role">Technical Lead — Solution Design &amp; Engineering <span class="meta">&nbsp;|&nbsp; HCLTech (for IKEA &amp; LEGO Group) &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led solution design across multiple enterprise products — defining patterns, reviewing designs, and ensuring scalability, resilience, and maintainability across complex, integrated systems.</li>
<li>Drove transformation and modernization journeys — migrating legacy systems to cloud-native architectures (GCP, AWS), re-architecting monoliths into microservices, and evolving platforms in production.</li>
<li>Coached and mentored a team of 8–12 engineers — raising technical quality through design reviews, pair programming, and structured knowledge sharing.</li>
<li>Worked across teams and influenced technical direction without direct authority — aligning multiple product teams around shared standards, patterns, and technology choices.</li>
<li>Identified technology opportunities and drove them into established ways of working — from proof-of-concept through adoption (containerization, event-driven patterns, Infrastructure as Code).</li>
<li>Collaborated closely with architects, product teams, and business stakeholders — ensuring solution designs were aligned with strategic direction and fit for purpose.</li>
<li>Maintained holistic system oversight and dependency mapping — giving teams the context needed to make informed design decisions in complex, multi-system environments.</li>
<li>Delivered solutions combining Platforms, SaaS (Genesys, Verint), and home-grown systems — demonstrating comfort in complex technology landscapes with diverse constraints.</li>
<li>Pushed back on designs that didn't meet quality standards — providing honest, constructive reasoning and guiding engineers toward better approaches.</li>
</ul>

<p class="role">Software Engineer <span class="meta">&nbsp;|&nbsp; Earlier Career (HCL, Ultimate Digital, Marlabs, TekMindz) &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Built enterprise software solutions across financial and retail domains — gaining foundational experience in distributed systems design, integration patterns, and backend engineering.</li>
<li>Contributed to modernization projects — migrating legacy systems to scalable architectures and establishing reusable design patterns.</li>
<li>Developed platform components, APIs, and integration layers using Java, Python, and SQL — supporting complex business workflows.</li>
<li>Worked in agile, self-organizing, cross-functional teams — building collaborative delivery practices from early career.</li>
</ul>

<h2>IKEA ECOSYSTEM EXPERIENCE</h2>
<ul>
<li><strong>IKEA Customer Connect (VCS):</strong> End-to-end solution design and technical leadership — cloud-native architecture (GCP), data pipelines, APIs, event-driven systems serving 32 markets.</li>
<li><strong>Solution Design Standards:</strong> Defining and maintaining common design patterns and standards across product teams — aligned with Ingka's engineering direction.</li>
<li><strong>Transformation &amp; Modernization:</strong> Led migration from legacy batch systems to real-time, event-driven architectures — evolving platforms in production with zero-downtime approach.</li>
<li><strong>PEDX Community Collaboration:</strong> Working with Architects, Service Management, and enabling teams — ensuring solutions are aligned, scalable, and fit for purpose.</li>
<li><strong>Fulfilment &amp; Core Services Adjacency:</strong> Systems integration with downstream fulfilment services, logistics tracking, and operational systems within the IKEA value chain.</li>
<li><strong>IKEA Ways of Working:</strong> Product team model, Agile/Scrum, DevOps set-up, self-organizing teams, Ingka DevOps tooling, cross-functional collaboration.</li>
<li><strong>Knowledge Sharing:</strong> Active contributor to Ingka engineering community — forums, documentation, design pattern libraries, and cross-domain collaboration.</li>
</ul>

<h2>IKEA CULTURE &amp; VALUES ALIGNMENT</h2>
<ul>
<li><strong>Togetherness:</strong> Work hands-on and across teams — enabling engineers with support, clarity, and constructive collaboration to deliver better solutions together.</li>
<li><strong>Leading by Example:</strong> Coach, mentor, and raise technical quality — setting the standard through design reviews, proof-of-concepts, and knowledge sharing.</li>
<li><strong>Simplicity:</strong> Define pragmatic, maintainable standards and patterns that teams can adopt — avoiding over-engineering while meeting real non-functional needs.</li>
<li><strong>Cost-Consciousness:</strong> Embed cost as a first-class non-functional requirement — ensuring solutions are economically sustainable from the design phase.</li>
<li><strong>Daring to be Different:</strong> Explore new technologies and AI approaches with proof-of-concepts — driving innovation with a pragmatic, tangible approach.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
