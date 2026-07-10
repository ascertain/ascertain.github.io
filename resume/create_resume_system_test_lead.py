"""Generate System Test Lead resume – External Application."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT_DIR = pathlib.Path(r"C:\Users\MOKAS10\vcs\csrs-vcs-core")
BASE = "Mohammad_Kashif_System_Test_Lead_Resume"

# ─── helpers ───────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_heading_block(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        text_after = text[len(bold_prefix):]
        r2 = p.add_run(text_after)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)

def add_role(doc, title, company, location, period):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"  |  {company}  |  {location}  |  {period}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── DOCX builder ─────────────────────────────────────────────────────────────
def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.0

    # ─── Name ──────────────────────────────────────────────────────────────
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("MOHAMMAD KASHIF")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ─── Contact ───────────────────────────────────────────────────────────
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.space_after = Pt(2)
    r = contact.add_run("Malmö, Sweden  |  +46 702624230  |  mo.kashif@gmail.com  |  linkedin.com/in/md-kashif")
    r.font.size = Pt(9.5)

    # ─── Profile Summary ───────────────────────────────────────────────────
    add_heading_block(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    r = p.add_run(
        "System Test Lead with 14+ years of experience in system testing, integration testing, "
        "and end-to-end validation of complex, integrated software solutions across multiple industries "
        "(retail/IoT, consumer apps, enterprise systems). Proven ability to plan and lead "
        "system testing activities, define test scope and approach, coordinate acceptance testing, "
        "and drive issue resolution with development teams. Strong hands-on testing background combined "
        "with excellent coordination and communication skills. Experienced in supporting customer-facing "
        "testing activities, pilot phases, and reporting system readiness to stakeholders. "
        "Structured, quality-focused mindset with deep understanding of agile, DevOps, and continuous delivery models."
    )
    r.font.size = Pt(10)

    # ─── Key Skills ────────────────────────────────────────────────────────
    add_heading_block(doc, "Key Skills & Competencies")
    skills_data = [
        ("System & Integration Testing", "System-level test planning · End-to-end testing · Cross-functional workflow validation · Integration testing · Acceptance testing coordination"),
        ("Test Strategy & Planning", "Test scope definition · Test cycles · Entry/exit criteria · Risk-based testing · Defect tracking · Test evidence & reporting"),
        ("Hands-on Testing", "Playwright · Selenium · Python (pytest) · API testing (Postman, REST) · Performance testing · Security testing · Accessibility testing"),
        ("Coordination & Reporting", "Stakeholder communication · Test progress reporting · Risk escalation · Customer acceptance support · Pilot phase coordination"),
        ("DevOps & CI/CD", "GitHub Actions · Jenkins · Terraform · Docker · Kubernetes · Automated quality gates · Continuous delivery pipelines"),
        ("Tools & Platforms", "Jira · Confluence · TestRail · Zephyr · Azure DevOps · GCP · AWS · Git"),
        ("Certifications", "ISTQB Certified Tester · AWS Cloud Practitioner · Google Cloud ACE · CEH · Six Sigma Green Belt"),
    ]
    tbl = doc.add_table(rows=len(skills_data), cols=2)
    tbl.autofit = True
    for i, (cat, detail) in enumerate(skills_data):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.width = Cm(3.8)
        r0 = c0.paragraphs[0].add_run(cat)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r1 = c1.paragraphs[0].add_run(detail)
        r1.font.size = Pt(9.5)
        set_cell_shading(c0, "F2F7FC")

    # ─── Experience ────────────────────────────────────────────────────────
    add_heading_block(doc, "Professional Experience")

    # Role 1 - Ingka Digital
    add_role(doc, "System Test Lead / Team Lead Acting", "Ingka Digital (IKEA Group)", "Malmö, Sweden", "2022 – Present")
    bullets_1 = [
        "Plan and lead system testing activities for complex, integrated customer support solutions serving 32 markets — defining test scope, approach, and test cycles in cooperation with development and release stakeholders.",
        "Execute and oversee hands-on system-level testing, validating end-to-end and cross-functional workflows across APIs, frontend applications, third-party vendor integrations, and data pipelines.",
        "Own and coordinate internal acceptance testing to confirm system readiness before customer/market exposure — ensuring entry criteria are met and clearly communicated.",
        "Support customer acceptance testing and pilot phases by providing test evidence, clarification, and testing-related support to market rollout teams.",
        "Align and integrate team-level testing results into a coherent system-level quality view — consolidating unit, integration, and component test outcomes into release readiness assessments.",
        "Track testing progress, risks, and defects — driving issue resolution together with development teams through daily triage and priority alignment.",
        "Report system testing status, readiness, and risks to product leadership and release stakeholders — providing transparent quality visibility for go/no-go decisions.",
        "Built test automation (Playwright) achieving 90%+ coverage on critical end-to-end flows, integrated into CI/CD pipelines for continuous validation.",
        "Coordinate across multiple teams (platform, product, vendor) ensuring cross-functional dependencies are tested and validated at system level.",
    ]
    for b in bullets_1:
        bullet(doc, b)

    # Role 2 - Truecaller
    add_role(doc, "System Test & Release Coordinator", "Truecaller", "Stockholm, Sweden", "Sep 2021 – Feb 2022")
    bullets_2 = [
        "Planned and executed system-level testing for a globally distributed application (300M+ users) — validating end-to-end workflows across backend services, mobile apps, and third-party integrations.",
        "Defined system test scope and release criteria — coordinating acceptance testing across engineering, QA, and product teams before production releases.",
        "Tracked defects, testing progress, and risks — reporting system readiness to stakeholders and driving issue resolution with development teams.",
        "Supported customer-facing rollout phases with test evidence and quality documentation — ensuring professional and transparent communication.",
        "Built CI/CD quality gates reducing release cycle time by 50% while maintaining system-level quality confidence.",
    ]
    for b in bullets_2:
        bullet(doc, b)

    # Role 3 - HCLTech
    add_role(doc, "System Test Lead / Integration Test Manager", "HCLTech", "Denmark & Sweden", "Jun 2013 – Sep 2021")
    bullets_3 = [
        "Led system testing and integration testing for complex IoT and enterprise solutions (smart-home devices, connected products, enterprise platforms) — coordinating across firmware, cloud, mobile, and hardware teams.",
        "Defined system test strategies, test cycles, and entry/exit criteria for multi-component integrated solutions — ensuring end-to-end validation before customer exposure.",
        "Executed hands-on system-level testing of cross-functional workflows — APIs, device-cloud communication, mobile apps, and backend services as an integrated system.",
        "Owned internal acceptance testing — confirming system readiness, documenting test evidence, and communicating entry criteria for customer acceptance phases.",
        "Supported international customer acceptance testing and pilot deployments — providing test reports, clarification, and defect resolution coordination.",
        "Tracked testing progress, risks, and defects across multiple development teams — driving resolution through daily stand-ups and triage meetings.",
        "Reported system testing status and readiness to project leadership and international stakeholders — enabling informed release decisions.",
        "Built and scaled test automation frameworks (Python, Selenium) for regression and integration testing — reducing manual effort by 70%.",
        "Mentored team members (8–12) on system testing practices, structured test planning, and defect management.",
    ]
    for b in bullets_3:
        bullet(doc, b)

    # Role 4 - Earlier Career
    add_role(doc, "Test Engineer / QA Engineer", "HCL, Ultimate Digital, Marlabs, TekMindz", "India", "2008 – 2013")
    bullets_4 = [
        "Performed system testing and integration testing for enterprise financial and retail applications — validating end-to-end business workflows across multiple integrated modules.",
        "Defined test cases and test cycles for complex multi-tier applications — ensuring cross-functional coverage across frontend, backend, database, and third-party integrations.",
        "Executed hands-on testing of APIs, batch processing, and data flows — identifying and reporting system-level defects with clear reproduction steps.",
        "Supported customer acceptance testing phases — providing test evidence, executing customer scenarios, and coordinating defect resolution with development teams.",
        "Tracked defects and testing progress using structured tools (HP ALM, JIRA) — reporting status to project managers and stakeholders.",
        "Built early test automation (Python, shell scripting, Selenium) — transitioning manual regression into repeatable automated suites.",
        "Collaborated with developers, business analysts, and customers in agile and waterfall environments — adapting testing approach to project delivery models.",
    ]
    for b in bullets_4:
        bullet(doc, b)

    # ─── Education ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Education")
    p = doc.add_paragraph()
    r = p.add_run("Post Graduate Diploma in Operations & Management")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run("  —  IGNOU, India")
    r2.font.size = Pt(10)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("B.Tech, Information Technology")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p2.add_run("  —  UP Technical University, India")
    r4.font.size = Pt(10)

    # ─── Languages ─────────────────────────────────────────────────────────
    add_heading_block(doc, "Languages")
    p = doc.add_paragraph()
    r = p.add_run("English (Fluent)")
    r.font.size = Pt(10)

    # ─── Save ──────────────────────────────────────────────────────────────
    out = OUT_DIR / f"{BASE}.docx"
    doc.save(str(out))
    print(f"DOCX saved: {out}")
    return out

# ─── DOC (HTML) builder ────────────────────────────────────────────────────────
def build_doc():
    content = """<html><head><meta charset="utf-8">
<style>
body{font-family:Calibri,sans-serif;font-size:10pt;margin:0.8cm 1.2cm;line-height:1.3}
h1{text-align:center;color:#1F4E79;font-size:18pt;margin-bottom:2px}
.contact{text-align:center;font-size:9.5pt;margin-bottom:10px}
h2{color:#1F4E79;font-size:10.5pt;border-bottom:1px solid #1F4E79;padding-bottom:2px;margin-top:12px}
.role{font-weight:bold;margin-top:8px} .meta{color:#444;font-size:9.5pt}
ul{margin:2px 0 4px 18px;padding:0} li{margin-bottom:2px}
table{width:100%;border-collapse:collapse;font-size:9.5pt} td{padding:2px 6px;vertical-align:top}
.cat{background:#F2F7FC;font-weight:bold;width:22%}
</style></head><body>
<h1>MOHAMMAD KASHIF</h1>
<p class="contact">Malmö, Sweden &nbsp;|&nbsp; +46 702624230 &nbsp;|&nbsp; mo.kashif@gmail.com &nbsp;|&nbsp; linkedin.com/in/md-kashif</p>

<h2>PROFESSIONAL SUMMARY</h2>
<p>System Test Lead with 14+ years of experience in system testing, integration testing, and end-to-end validation of complex, integrated software solutions across multiple industries (retail/IoT, consumer apps, enterprise systems). Proven ability to plan and lead system testing activities, define test scope and approach, coordinate acceptance testing, and drive issue resolution with development teams. Strong hands-on testing background combined with excellent coordination and communication skills. Experienced in supporting customer-facing testing activities, pilot phases, and reporting system readiness to stakeholders. Structured, quality-focused mindset with deep understanding of agile, DevOps, and continuous delivery models.</p>

<h2>KEY SKILLS &amp; COMPETENCIES</h2>
<table>
<tr><td class="cat">System &amp; Integration Testing</td><td>System-level test planning · End-to-end testing · Cross-functional workflow validation · Integration testing · Acceptance testing coordination</td></tr>
<tr><td class="cat">Test Strategy &amp; Planning</td><td>Test scope definition · Test cycles · Entry/exit criteria · Risk-based testing · Defect tracking · Test evidence &amp; reporting</td></tr>
<tr><td class="cat">Hands-on Testing</td><td>Playwright · Selenium · Python (pytest) · API testing (Postman, REST) · Performance testing · Security testing · Accessibility testing</td></tr>
<tr><td class="cat">Coordination &amp; Reporting</td><td>Stakeholder communication · Test progress reporting · Risk escalation · Customer acceptance support · Pilot phase coordination</td></tr>
<tr><td class="cat">DevOps &amp; CI/CD</td><td>GitHub Actions · Jenkins · Terraform · Docker · Kubernetes · Automated quality gates · Continuous delivery pipelines</td></tr>
<tr><td class="cat">Tools &amp; Platforms</td><td>Jira · Confluence · TestRail · Zephyr · Azure DevOps · GCP · AWS · Git</td></tr>
<tr><td class="cat">Certifications</td><td>ISTQB Certified Tester · AWS Cloud Practitioner · Google Cloud ACE · CEH · Six Sigma Green Belt</td></tr>
</table>

<h2>PROFESSIONAL EXPERIENCE</h2>

<p class="role">System Test Lead / Team Lead Acting <span class="meta">&nbsp;|&nbsp; Ingka Digital (IKEA Group) &nbsp;|&nbsp; Malmö, Sweden &nbsp;|&nbsp; 2022 – Present</span></p>
<ul>
<li>Plan and lead system testing activities for complex, integrated customer support solutions serving 32 markets — defining test scope, approach, and test cycles in cooperation with development and release stakeholders.</li>
<li>Execute and oversee hands-on system-level testing, validating end-to-end and cross-functional workflows across APIs, frontend applications, third-party vendor integrations, and data pipelines.</li>
<li>Own and coordinate internal acceptance testing to confirm system readiness before customer/market exposure — ensuring entry criteria are met and clearly communicated.</li>
<li>Support customer acceptance testing and pilot phases by providing test evidence, clarification, and testing-related support to market rollout teams.</li>
<li>Align and integrate team-level testing results into a coherent system-level quality view — consolidating unit, integration, and component test outcomes into release readiness assessments.</li>
<li>Track testing progress, risks, and defects — driving issue resolution together with development teams through daily triage and priority alignment.</li>
<li>Report system testing status, readiness, and risks to product leadership and release stakeholders — providing transparent quality visibility for go/no-go decisions.</li>
<li>Built test automation (Playwright) achieving 90%+ coverage on critical end-to-end flows, integrated into CI/CD pipelines for continuous validation.</li>
<li>Coordinate across multiple teams (platform, product, vendor) ensuring cross-functional dependencies are tested and validated at system level.</li>
</ul>

<p class="role">System Test &amp; Release Coordinator <span class="meta">&nbsp;|&nbsp; Truecaller &nbsp;|&nbsp; Stockholm, Sweden &nbsp;|&nbsp; Sep 2021 – Feb 2022</span></p>
<ul>
<li>Planned and executed system-level testing for a globally distributed application (300M+ users) — validating end-to-end workflows across backend services, mobile apps, and third-party integrations.</li>
<li>Defined system test scope and release criteria — coordinating acceptance testing across engineering, QA, and product teams before production releases.</li>
<li>Tracked defects, testing progress, and risks — reporting system readiness to stakeholders and driving issue resolution with development teams.</li>
<li>Supported customer-facing rollout phases with test evidence and quality documentation — ensuring professional and transparent communication.</li>
<li>Built CI/CD quality gates reducing release cycle time by 50% while maintaining system-level quality confidence.</li>
</ul>

<p class="role">System Test Lead / Integration Test Manager <span class="meta">&nbsp;|&nbsp; HCLTech &nbsp;|&nbsp; Denmark &amp; Sweden &nbsp;|&nbsp; Jun 2013 – Sep 2021</span></p>
<ul>
<li>Led system testing and integration testing for complex IoT and enterprise solutions (smart-home devices, connected products, enterprise platforms) — coordinating across firmware, cloud, mobile, and hardware teams.</li>
<li>Defined system test strategies, test cycles, and entry/exit criteria for multi-component integrated solutions — ensuring end-to-end validation before customer exposure.</li>
<li>Executed hands-on system-level testing of cross-functional workflows — APIs, device-cloud communication, mobile apps, and backend services as an integrated system.</li>
<li>Owned internal acceptance testing — confirming system readiness, documenting test evidence, and communicating entry criteria for customer acceptance phases.</li>
<li>Supported international customer acceptance testing and pilot deployments — providing test reports, clarification, and defect resolution coordination.</li>
<li>Tracked testing progress, risks, and defects across multiple development teams — driving resolution through daily stand-ups and triage meetings.</li>
<li>Reported system testing status and readiness to project leadership and international stakeholders — enabling informed release decisions.</li>
<li>Built and scaled test automation frameworks (Python, Selenium) for regression and integration testing — reducing manual effort by 70%.</li>
<li>Mentored team members (8–12) on system testing practices, structured test planning, and defect management.</li>
</ul>

<p class="role">Test Engineer / QA Engineer <span class="meta">&nbsp;|&nbsp; HCL, Ultimate Digital, Marlabs, TekMindz &nbsp;|&nbsp; India &nbsp;|&nbsp; 2008 – 2013</span></p>
<ul>
<li>Performed system testing and integration testing for enterprise financial and retail applications — validating end-to-end business workflows across multiple integrated modules.</li>
<li>Defined test cases and test cycles for complex multi-tier applications — ensuring cross-functional coverage across frontend, backend, database, and third-party integrations.</li>
<li>Executed hands-on testing of APIs, batch processing, and data flows — identifying and reporting system-level defects with clear reproduction steps.</li>
<li>Supported customer acceptance testing phases — providing test evidence, executing customer scenarios, and coordinating defect resolution with development teams.</li>
<li>Tracked defects and testing progress using structured tools (HP ALM, JIRA) — reporting status to project managers and stakeholders.</li>
<li>Built early test automation (Python, shell scripting, Selenium) — transitioning manual regression into repeatable automated suites.</li>
<li>Collaborated with developers, business analysts, and customers in agile and waterfall environments — adapting testing approach to project delivery models.</li>
</ul>

<h2>EDUCATION</h2>
<p><strong>Post Graduate Diploma in Operations &amp; Management</strong> — IGNOU, India</p>
<p><strong>B.Tech, Information Technology</strong> — UP Technical University, India</p>

<h2>LANGUAGES</h2>
<p>English (Fluent)</p>

</body></html>"""
    out = OUT_DIR / f"{BASE}.doc"
    out.write_text(content, encoding="utf-8")
    print(f"DOC saved: {out}")

# ─── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_docx()
    build_doc()
